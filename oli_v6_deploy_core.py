"""ILO project quality assessment app (core tabs).

Tabs:
    1. Valoración Preliminar de Calidad de Proyectos
    2. Diagnóstico de Atributos Específicos
    3. Diagnóstico de Sostenibilidad del Proyecto
    4. Pregúntale a tus Documentos
"""

import concurrent.futures
import pickle
import streamlit as st
import pandas as pd
import json
import tempfile
import docx
import numpy as np
import faiss
import os
import re
from matplotlib import pyplot as plt
import seaborn as sns
from io import BytesIO
import torch
import time
import plotly.express as px
import plotly.graph_objects as go
from docx2python import docx2python
from io import BytesIO
import streamlit as st
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import zipfile


def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
        df.to_excel(writer, index=False, sheet_name='Datos Filtrados')
    processed_data = output.getvalue()
    return processed_data

# Use environment variables for API keys
# For local development - use .env file or set environment variables
# For Streamlit Cloud - set these in the app settings
openai_api_key = os.getenv("OPENAI_API_KEY")
# Initialize OpenAI with new SDK (v1.0.0+)
from openai import OpenAI
client = OpenAI(api_key=openai_api_key)

os.environ["KMP_DUPLICATE_LIB_OK"]="TRUE"

# Import tiktoken for accurate token counting
try:
    import tiktoken
    # Use cl100k_base encoding (standard for GPT-4 and GPT-5 models)
    encoding = tiktoken.get_encoding("cl100k_base")
except ImportError:
    encoding = None
    # Warning will be shown when function is first called if needed

# Helper function to truncate text to a token limit
def truncate_to_token_limit(text, max_tokens=110000, encoding_obj=None):
    """
    Truncate text to a maximum number of tokens.
    
    Args:
        text: The text to truncate
        max_tokens: Maximum number of tokens (default 110K for GPT-5-mini's 128K context window)
        encoding_obj: tiktoken encoding object (if None, falls back to character estimation)
    
    Returns:
        Truncated text that fits within the token limit
    """
    if not text:
        return text
    
    if encoding_obj is None:
        # Fallback: use character-based estimation (4 chars per token for Spanish)
        # 110K tokens * 4 = 440K characters
        max_chars = max_tokens * 4
        return text[:max_chars]
    
    # Count tokens
    tokens = encoding_obj.encode(text)
    
    # If within limit, return full text
    if len(tokens) <= max_tokens:
        return text
    
    # Truncate to max_tokens
    truncated_tokens = tokens[:max_tokens]
    truncated_text = encoding_obj.decode(truncated_tokens)

    return truncated_text

# JSON Schemas for structured rubric responses (Tabs 1-4).
# Using strict structured outputs eliminates JSON-parse fallbacks AND gives a
# verifiable 'parte_enfocada' signal that the model committed to Part 2.
RUBRIC_SCHEMA_TWO_PART = {
    "name": "rubric_response_two_part",
    "strict": True,
    "schema": {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "Respuesta": {"type": "string", "enum": ["Yes", "No", "Partial", "Not Found"]},
            "Razonamiento": {"type": "string"},
            "Evidencia": {"type": "string"},
            "parte_enfocada": {"type": "string", "enum": ["Parte 2", "Parte 1", "Ambas"]},
        },
        "required": ["Respuesta", "Razonamiento", "Evidencia", "parte_enfocada"],
    },
}

RUBRIC_SCHEMA_SINGLE = {
    "name": "rubric_response_single",
    "strict": True,
    "schema": {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "Respuesta": {"type": "string", "enum": ["Yes", "No", "Partial", "Not Found"]},
            "Razonamiento": {"type": "string"},
            "Evidencia": {"type": "string"},
        },
        "required": ["Respuesta", "Razonamiento", "Evidencia"],
    },
}

# ============= MAIN APP CODE =============

# Set page config
st.set_page_config(layout="wide")

# ============= ILO GLOBAL STYLES =============
# ILO Color Palette: Red (#C8102E), Blue (#002F6C), White (#FFFFFF)
# Secondary: Light Blue (#0072CE), Light Gray (#F5F5F5)
st.markdown("""
<style>
/* ===== ILO Color Variables ===== */
:root {
    --ilo-red: #C8102E;
    --ilo-blue: #002F6C;
    --ilo-light-blue: #0072CE;
    --ilo-white: #FFFFFF;
    --ilo-light-gray: #F5F5F5;
    --ilo-dark-gray: #333333;
}

/* ===== Global Typography ===== */
h1, h2, h3 {
    color: var(--ilo-blue) !important;
    font-weight: 600 !important;
}

h4, h5, h6 {
    color: var(--ilo-dark-gray) !important;
}

/* ===== Primary Buttons ===== */
.stButton > button {
    background-color: var(--ilo-blue) !important;
    color: var(--ilo-white) !important;
    border: none !important;
    border-radius: 6px !important;
    padding: 0.5rem 1.25rem !important;
    font-weight: 500 !important;
    transition: all 0.2s ease !important;
}

.stButton > button:hover {
    background-color: var(--ilo-light-blue) !important;
    box-shadow: 0 4px 12px rgba(0, 47, 108, 0.25) !important;
}

/* ===== Download Buttons ===== */
.stDownloadButton > button {
    background-color: var(--ilo-red) !important;
    color: var(--ilo-white) !important;
    border: none !important;
    border-radius: 6px !important;
}

.stDownloadButton > button:hover {
    background-color: #A00D24 !important;
    box-shadow: 0 4px 12px rgba(200, 16, 46, 0.3) !important;
}

/* ===== Tabs Styling ===== */
.stTabs [data-baseweb="tab-list"] {
    gap: 8px;
    background-color: var(--ilo-light-gray);
    padding: 0.5rem;
    border-radius: 8px;
}

.stTabs [data-baseweb="tab"] {
    background-color: transparent !important;
    color: var(--ilo-blue) !important;
    border-radius: 6px !important;
    padding: 0.5rem 1rem !important;
    font-weight: 500 !important;
}

.stTabs [aria-selected="true"] {
    background-color: var(--ilo-blue) !important;
    color: var(--ilo-white) !important;
}

/* ===== Expanders ===== */
.streamlit-expanderHeader {
    background-color: var(--ilo-light-gray) !important;
    border-radius: 8px !important;
    color: var(--ilo-blue) !important;
    font-weight: 500 !important;
}

/* ===== Info/Warning/Success Boxes ===== */
.stAlert > div[data-baseweb="notification"] {
    border-radius: 8px !important;
}

div[data-testid="stNotificationContentInfo"] {
    background-color: rgba(0, 114, 206, 0.1) !important;
    border-left: 4px solid var(--ilo-light-blue) !important;
}

div[data-testid="stNotificationContentSuccess"] {
    background-color: rgba(40, 167, 69, 0.1) !important;
    border-left: 4px solid #28a745 !important;
}

div[data-testid="stNotificationContentWarning"] {
    background-color: rgba(200, 16, 46, 0.08) !important;
    border-left: 4px solid var(--ilo-red) !important;
}

/* ===== Checkboxes ===== */
.stCheckbox > label > div[data-testid="stMarkdownContainer"] {
    color: var(--ilo-dark-gray) !important;
}

/* ===== Sliders ===== */
.stSlider > div > div > div > div {
    background-color: var(--ilo-blue) !important;
}

/* ===== Select boxes ===== */
.stSelectbox > div > div {
    border-color: var(--ilo-blue) !important;
}

/* ===== DataFrames ===== */
.stDataFrame {
    border: 1px solid #e0e0e0 !important;
    border-radius: 8px !important;
}

/* ===== Progress bars ===== */
.stProgress > div > div > div > div {
    background-color: var(--ilo-blue) !important;
}

/* ===== Metric cards ===== */
[data-testid="stMetricValue"] {
    color: var(--ilo-blue) !important;
}

/* ===== Reference boxes (custom) ===== */
.reference-box {
    border-left: 4px solid var(--ilo-blue) !important;
    background: linear-gradient(135deg, #f0f4f8, #e8eef5) !important;
    padding: 0.85rem 1rem;
    border-radius: 8px;
    font-size: 0.95rem;
    color: var(--ilo-dark-gray);
    margin-bottom: 0.75rem;
    box-shadow: 0 4px 14px rgba(0, 47, 108, 0.08);
}

/* ===== Horizontal rule ===== */
hr {
    border-top: 2px solid var(--ilo-blue) !important;
}

/* ===== Sidebar ===== */
[data-testid="stSidebar"] {
    background-color: var(--ilo-light-gray) !important;
}

[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {
    color: var(--ilo-blue) !important;
}

/* ===== File uploader ===== */
.stFileUploader > div > div {
    border: 2px dashed var(--ilo-blue) !important;
    border-radius: 8px !important;
}

/* ===== Radio buttons ===== */
.stRadio > div {
    gap: 0.5rem;
}

/* ===== Multiselect ===== */
.stMultiSelect > div > div {
    border-color: var(--ilo-blue) !important;
}

/* ===== Spinner ===== */
.stSpinner > div {
    border-top-color: var(--ilo-blue) !important;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
    <h2 style='text-align:center; color:#002F6C; margin-top:0;'>
        Caja de Herramientas para el Mejor Desempeño de los Proyectos
        <br>
        <span style='font-size:0.8em; font-weight:500;'>Toolkit for Better Project Performance</span>
    </h2>
    <h3 style='text-align:center; color:#002F6C; margin-top:0;'>
        Usando Evidencia de las Evaluaciones
        <br>
        <span style='font-size:0.85em; font-weight:500;'>Using Evidence from Evaluations</span>
    </h3>
    <hr style='border-top: 2px solid #002F6C;'>
""", unsafe_allow_html=True)

tab1, tab2, tab3, tab4 = st.tabs([
    "Valoración Preliminar de Calidad de Proyectos",
    "Diagnóstico de Atributos Específicos",
    "Diagnóstico de Sostenibilidad del Proyecto",
    "Pregúntale a tus Documentos",
])

with tab2:
    st.header("Diagnóstico de Atributos Específicos")

    # Descriptive text box
    st.info("""
    **📋 Descripción de la herramienta:**

    **¿Qué hace esta herramienta?:**
    Profundiza la Valoración Preliminar aplicando rúbricas OIT con niveles de desempeño (1–5) y evidencia trazable (citas y metadatos) sobre un documento .docx que subas. Extrae secciones clave, evalúa contra una matriz de criterios y genera un análisis narrativo por criterio.
    Criterios disponibles y alcance
    -	Metodologías con enfoque participativo → aplicar a informes de evaluación u otros documentos metodológicos.

    -	Integración del enfoque de género → aplicar a documentos de diseño o ejecución de proyecto u otros estudios (p. ej., PRODOC, TPR, etc.).

    -   Integración del enfoque de Transición Justa (enfoque moderno) → aplicar a documentos de diseño o ejecución de proyecto u otros estudios (p. ej., PRODOC, TPR, etc.). 
   
    Puedes exportar a Excel estos resultados (Criterio, Dimensión, Score, Análisis, Evidencia, Error, Rúbrica). Una vez que los resultados son descargados, éstos se dejarán de mostrar en pantalla.
    Si hay vacíos o inconsistencias, se señalan en "Error" para su ajuste.
    
    **¿Para qué usar este diagnóstico?**
    Este diagnóstico en formato EXCEL sirve para **revisar propuestas**, **verificar aspectos puntuales** de informes de evaluación o de ejecución, **comprobar coherencia** con P&B, DWCP y marcos UNSDCF, **elaborar notas técnicas con sustento** y respaldar la rendición de cuentas ante mandantes y donantes.
    """)
    
    # Important requirements reminder
    st.warning("""
    **⚠️ Recordatorio importante:**
    - Los documentos deben estar formateados con **estilos de encabezado de Word** (Heading 1, Heading 2, etc.) para que las secciones se identifiquen correctamente
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Solo se aceptan archivos en formato **.docx**
    """)

    # Read rubrics from Excel files as in megaparse_example.py
    import pandas as pd

    # Function to extract document structure
    def extract_docx_structure(docx_path):
        from docx import Document
        doc = Document(docx_path)
        filename = os.path.basename(docx_path)
        rows = []
        current_headers = {i: '' for i in range(1, 7)}
        para_counter = 0

        def get_header_level(style_name):
            for i in range(1, 7):
                if style_name.lower().startswith(f'heading {i}'.lower()):
                    return i
            return None

        def header_dict():
            return {f'header_{i}': current_headers[i] for i in range(1, 7)}

        for para in doc.paragraphs:
            para_counter += 1
            level = get_header_level(para.style.name)
            if level and 1 <= level <= 6:
                current_headers[level] = para.text.strip()
                for l in range(level+1, 7):
                    current_headers[l] = ''
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'content': '',
                    'source_type': 'heading',
                    'paragraph_number': para_counter,
                    'page_number': None
                })
            elif para.text.strip():
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'content': para.text.strip(),
                    'source_type': 'paragraph',
                    'paragraph_number': para_counter,
                    'page_number': None
                })
        return pd.DataFrame(rows)
    
    # Enhanced extraction function with multi-method header detection, tables, and validation
    def extract_docx_structure_enhanced(docx_path):
        """
        Enhanced document extraction with:
        - Multi-method header detection (style, formatting, pattern)
        - Table extraction
        - List extraction
        - Page estimation
        - Quality metrics
        """
        from docx import Document
        from docx.shared import Pt
        import re
        
        doc = Document(docx_path)
        filename = os.path.basename(docx_path)
        rows = []
        tables_data = []
        current_headers = {i: '' for i in range(1, 7)}
        para_counter = 0
        page_estimate = 1
        words_per_page = 500  # Estimate for page calculation
        
        # Track statistics
        stats = {
            'total_paragraphs': 0,
            'total_tables': 0,
            'headers_detected': 0,
            'headers_by_style': 0,
            'headers_by_formatting': 0,
            'headers_by_pattern': 0,
            'orphaned_paragraphs': 0,
            'empty_sections': 0
        }
        
        def get_header_level_by_style(style_name):
            """Method 1: Style-based detection (original method)"""
            for i in range(1, 7):
                if style_name.lower().startswith(f'heading {i}'.lower()):
                    return i
            return None
        
        def get_header_level_by_formatting(para):
            """Method 2: Formatting-based detection"""
            # Check if paragraph is bold and larger than normal text
            if para.runs:
                first_run = para.runs[0]
                is_bold = first_run.bold
                font_size = first_run.font.size
                
                # If bold and larger font, likely a header
                if is_bold and font_size:
                    if font_size >= Pt(14):  # Larger than normal (11-12pt)
                        # Estimate level based on size
                        if font_size >= Pt(18):
                            return 1
                        elif font_size >= Pt(16):
                            return 2
                        elif font_size >= Pt(14):
                            return 3
                elif is_bold and len(para.text.strip()) < 100:  # Short bold text
                    return 2  # Likely a subheader
            return None
        
        def get_header_level_by_pattern(para_text):
            """Method 3: Pattern-based detection (numbered headings)"""
            text = para_text.strip()
            # Patterns like "1.", "1.1", "1.1.1", "I.", "A.", etc.
            patterns = [
                (r'^\d+\.\s+\w', 1),  # "1. Title"
                (r'^\d+\.\d+\s+\w', 2),  # "1.1 Title"
                (r'^\d+\.\d+\.\d+\s+\w', 3),  # "1.1.1 Title"
                (r'^[IVX]+\.\s+\w', 1),  # "I. Title"
                (r'^[A-Z]\.\s+\w', 2),  # "A. Title"
            ]
            for pattern, level in patterns:
                if re.match(pattern, text):
                    return level
            return None
        
        def detect_header_level(para):
            """Try all methods to detect header level"""
            # Method 1: Style-based
            level = get_header_level_by_style(para.style.name)
            if level:
                stats['headers_by_style'] += 1
                return level
            
            # Method 2: Formatting-based
            level = get_header_level_by_formatting(para)
            if level:
                stats['headers_by_formatting'] += 1
                return level
            
            # Method 3: Pattern-based
            level = get_header_level_by_pattern(para.text)
            if level:
                stats['headers_by_pattern'] += 1
                return level
            
            return None
        
        def header_dict():
            return {f'header_{i}': current_headers[i] for i in range(1, 7)}
        
        def get_full_header_path():
            """Get full hierarchical path of headers"""
            path_parts = []
            for i in range(1, 7):
                if current_headers[i]:
                    path_parts.append(current_headers[i])
            return ' > '.join(path_parts) if path_parts else 'Sin sección'
        
        # Extract paragraphs and headers
        for para in doc.paragraphs:
            para_counter += 1
            stats['total_paragraphs'] += 1
            
            level = detect_header_level(para)
            if level and 1 <= level <= 6:
                current_headers[level] = para.text.strip()
                # Clear lower level headers
                for l in range(level+1, 7):
                    current_headers[l] = ''
                
                stats['headers_detected'] += 1
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'header_path': get_full_header_path(),
                    'content': '',
                    'source_type': 'heading',
                    'paragraph_number': para_counter,
                    'page_estimate': page_estimate,
                    'detection_method': 'style' if get_header_level_by_style(para.style.name) else ('formatting' if get_header_level_by_formatting(para) else 'pattern')
                })
            elif para.text.strip():
                # Estimate page number based on word count
                word_count = len(para.text.split())
                if word_count > 0:
                    page_estimate = max(1, int(stats['total_paragraphs'] * 0.02))  # Rough estimate
                
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'header_path': get_full_header_path(),
                    'content': para.text.strip(),
                    'source_type': 'paragraph',
                    'paragraph_number': para_counter,
                    'page_estimate': page_estimate,
                    'detection_method': None
                })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            stats['total_tables'] += 1
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            # Find which section this table belongs to
            current_section = get_full_header_path()
            
            tables_data.append({
                'filename': filename,
                'table_number': table_idx + 1,
                'section': current_section,
                'header_path': current_section,
                'data': table_data,
                'row_count': len(table_data),
                'col_count': len(table_data[0]) if table_data else 0
            })
            
            # Add table as content row
            table_text = '\n'.join([' | '.join(row) for row in table_data])
            rows.append({
                'filename': filename,
                **header_dict(),
                'header_path': current_section,
                'content': f"[TABLA {table_idx + 1}]\n{table_text}",
                'source_type': 'table',
                'paragraph_number': para_counter + table_idx,
                'page_estimate': page_estimate,
                'detection_method': None
            })
        
        # Calculate quality metrics
        df = pd.DataFrame(rows)
        
        # Count orphaned paragraphs (no header_1)
        stats['orphaned_paragraphs'] = len(df[(df['source_type'] == 'paragraph') & (df['header_1'].isna() | (df['header_1'] == ''))])
        
        # Count empty sections
        if 'header_1' in df.columns:
            sections = df[df['header_1'].notna() & (df['header_1'] != '')]['header_1'].unique()
            for section in sections:
                section_df = df[df['header_1'] == section]
                if len(section_df[section_df['source_type'] == 'paragraph']) == 0:
                    stats['empty_sections'] += 1
        
        return df, tables_data, stats
    
    # Validation function
    def validate_extraction(df, tables_data, stats):
        """
        Validate extraction quality and return metrics, warnings, and recommendations
        """
        warnings = []
        recommendations = []
        quality_score = 100
        
        # Check header detection rate
        if stats['total_paragraphs'] > 0:
            header_rate = (stats['headers_detected'] / stats['total_paragraphs']) * 100
            if header_rate < 2:  # Less than 2% headers might indicate poor detection
                warnings.append(f"⚠️ Tasa de detección de encabezados baja: {header_rate:.1f}%")
                quality_score -= 10
                recommendations.append("Verificar si los encabezados usan estilos estándar de Word")
        
        # Check for orphaned paragraphs
        if stats['orphaned_paragraphs'] > 0:
            orphan_rate = (stats['orphaned_paragraphs'] / stats['total_paragraphs']) * 100
            if orphan_rate > 10:  # More than 10% orphaned
                warnings.append(f"⚠️ {stats['orphaned_paragraphs']} párrafos sin sección asignada ({orphan_rate:.1f}%)")
                quality_score -= 15
                recommendations.append("Revisar párrafos huérfanos y asignarlos manualmente a secciones")
        
        # Check for empty sections
        if stats['empty_sections'] > 0:
            warnings.append(f"⚠️ {stats['empty_sections']} secciones sin contenido")
            quality_score -= 5 * min(stats['empty_sections'], 5)  # Max -25 points
            recommendations.append("Revisar secciones vacías - pueden ser encabezados mal detectados")
        
        # Check table extraction
        if stats['total_tables'] > 0:
            if len(tables_data) < stats['total_tables']:
                warnings.append(f"⚠️ No se extrajeron todas las tablas ({len(tables_data)}/{stats['total_tables']})")
                quality_score -= 10
        
        # Check hierarchy integrity
        if 'header_1' in df.columns and 'header_2' in df.columns:
            # Check for level jumps (H1 -> H3, skipping H2)
            h1_sections = df[df['header_1'].notna() & (df['header_1'] != '')]
            for idx, row in h1_sections.iterrows():
                # Find next header after this one
                next_rows = df.iloc[idx+1:idx+10]  # Check next 10 rows
                next_headers = next_rows[next_rows['source_type'] == 'heading']
                if not next_headers.empty:
                    next_header = next_rows.iloc[0]
                    # Check if there's a level jump
                    if next_header.get('header_3') and not next_header.get('header_2'):
                        warnings.append(f"⚠️ Salto de nivel detectado: {row.get('header_1', '')} → {next_header.get('header_3', '')}")
                        quality_score -= 5
        
        quality_score = max(0, quality_score)  # Don't go below 0
        
        return {
            'quality_score': quality_score,
            'warnings': warnings,
            'recommendations': recommendations,
            'stats': stats
        }
    
    # Function to split text into chunks respecting the token limit
    def split_text_into_chunks(text, max_completion_tokens=7000):
        import re
        # Split by paragraphs first
        paragraphs = text.split('\n')
        chunks = []
        current_chunk = []
        current_length = 0

        # Rough estimate: 1 token ≈ 4 characters in Spanish
        tokens_per_char = 0.25

        for para in paragraphs:
            # Estimate tokens in this paragraph
            para_tokens = len(para) * tokens_per_char

            # If adding this paragraph would exceed the max, start a new chunk
            if current_length + para_tokens > max_completion_tokens and current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = [para]
                current_length = para_tokens
            else:
                current_chunk.append(para)
                current_length += para_tokens

        # Add the last chunk if there's content
        if current_chunk:
            chunks.append('\n'.join(current_chunk))

        return chunks
    engagement_rubric = {}
    performance_rubric = {}
    parteval_rubric = {}
    gender_rubric = {}
    tj_traditional_rubric = {}
    tj_just_transition_rubric = {}

    try:
        df_rubric_engagement = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_engagement')
        df_rubric_engagement.drop(columns=['Unnamed: 0', 'Criterio'], inplace=True, errors='ignore')
        for idx, row in df_rubric_engagement.iterrows():
            indicador = row['Indicador']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
            engagement_rubric[indicador] = {'valores': valores, 'dimension': dimension}

        df_rubric_performance = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_performance')
        df_rubric_performance.drop(columns=['dimension'], inplace=True, errors='ignore')
        for idx, row in df_rubric_performance.iterrows():
            criterio = row['subdim']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['subdim', 'Dimensión'], errors='ignore').values.tolist()
            performance_rubric[criterio] = {'valores': valores, 'dimension': dimension}

        df_rubric_parteval = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_parteval')
        df_rubric_parteval.drop(columns=['Criterio'], inplace=True, errors='ignore')
        for idx, row in df_rubric_parteval.iterrows():
            indicador = row['Indicador']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
            parteval_rubric[indicador] = {'valores': valores, 'dimension': dimension}

        df_rubric_gender = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_gender_')
        df_rubric_gender.drop(columns=['Criterio'], inplace=True, errors='ignore')
        for idx, row in df_rubric_gender.iterrows():
            indicador = row['Indicador']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
            gender_rubric[indicador] = {'valores': valores, 'dimension': dimension}

        try:
            df_rubric_tj_traditional = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_TJ_Traditional')
            df_rubric_tj_traditional.drop(columns=['Criterio'], inplace=True, errors='ignore')
            for idx, row in df_rubric_tj_traditional.iterrows():
                indicador = row['Indicador']
                if pd.notna(indicador) and str(indicador).strip():
                    dimension = row.get('Dimensión', 'No especificada')
                    valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
                    tj_traditional_rubric[indicador] = {'valores': valores, 'dimension': dimension}
        except Exception as e:
            st.error(f"Error cargando TJ Tradicional: {e}")

        try:
            df_rubric_tj_just_transition = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_TJ_TJ')
            df_rubric_tj_just_transition.drop(columns=['Criterio'], inplace=True, errors='ignore')
            for idx, row in df_rubric_tj_just_transition.iterrows():
                indicador = row['Indicador']
                if pd.notna(indicador) and str(indicador).strip():
                    dimension = row.get('Dimensión', 'No especificada')
                    valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
                    tj_just_transition_rubric[indicador] = {'valores': valores, 'dimension': dimension}
        except Exception as e:
            st.error(f"Error cargando TJ Transición Justa: {e}")
    except Exception as e:
        st.error(f"Error leyendo las rúbricas: {e}")

    # Show rubric status
    st.success(f"""
    **ESTADO DE RÚBRICAS:**
    
    - Metodologías con enfoque participativo ({len(parteval_rubric)} criterios disponibles)", 
    - Integración del Enfoque de Género: {len(gender_rubric)} criterios
    - Integración del Enfoque de Transición Justa: Enfoque Moderno: {len(tj_just_transition_rubric)} criterios 
    """)

    # Download button for the rubric file (directly on page, no expander)
    try:
        with open('./Rubricas_6ago2025.xlsx', 'rb') as f:
            st.download_button(
                label="📥 Descargar archivo rúbrica de Atributos específicos",
                data=f,
                file_name="Rubricas_6ago2025.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_rubricas_tab2"
            )
    except FileNotFoundError:
        st.warning("Archivo de rúbricas no disponible para descarga.")

    # Document upload
    st.markdown("### 📄 Carga de Documento")
    
    # Warning box about document requirements
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos:**
    
    **📝 Formato del documento:**
    - Solo se aceptan archivos en formato **.docx** (Word 2007 o posterior)
    - El documento debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del documento deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente y las secciones no se identificarán.
    - Esto es especialmente importante para evaluaciones y PRODOCs que deben tener una estructura clara con secciones bien definidas
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos que excedan este límite serán truncados automáticamente
    - Se recomienda dividir documentos muy extensos (más de ~180 páginas) en secciones más pequeñas si es necesario
    
    **✅ Mejores prácticas:**
    - Usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones principales
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que el documento esté guardado correctamente antes de subirlo
    - Verifica que todas las secciones importantes tengan encabezados antes de procesar
    """)
    
    uploaded_file = st.file_uploader("Suba un archivo DOCX para evaluación:", type=["docx"], key="tab2_file_uploader")

    # Initialize session state for selections and results persistence
    if 'selected_rubrics_tab2' not in st.session_state:
        st.session_state['selected_rubrics_tab2'] = []
    if 'selected_criteria_tab2' not in st.session_state:
        st.session_state['selected_criteria_tab2'] = {}
    if 'tab2_results' not in st.session_state:
        st.session_state['tab2_results'] = None
    if 'document_extracted_tab2' not in st.session_state:
        st.session_state['document_extracted_tab2'] = False

    # Document Extraction Section
    st.markdown("---")
    st.markdown("### 📥 Extracción de Documento")
    
    if uploaded_file is not None:
        file_hash = hash(uploaded_file.getvalue())
        file_changed = st.session_state.get('last_file_hash_tab2') != file_hash
        
        if file_changed:
            st.session_state['document_extracted_tab2'] = False
            st.session_state['last_file_hash_tab2'] = None
        
        if st.button("🔍 Extraer Documento", key="extract_document_tab2", type="primary"):
            if uploaded_file is None:
                st.error("Por favor suba un archivo DOCX primero.")
                st.stop()
            
            with st.spinner("Extrayendo documento..."):
                try:
                    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp_file.write(uploaded_file.read())
                    tmp_file.close()
                    
                    progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                    doc_result = docx2python(tmp_file.name)
                    
                    # Use enhanced extraction
                    df, tables_data, extraction_stats = extract_docx_structure_enhanced(tmp_file.name)
                    progress_bar.progress(0.2, text="Documento cargado. Validando extracción...")
                    
                    # Validate extraction
                    validation_results = validate_extraction(df, tables_data, extraction_stats)
                    progress_bar.progress(0.3, text="Extracción validada. Procesando secciones...")
                    
                    # Extract sections
                    header_1_values = df['header_1'].dropna().unique()
                    llm_summary_rows = []
                    
                    for idx, header in enumerate(header_1_values):
                        section_df = df[df['header_1'] == header].copy()
                        # Extract text directly - already clean from extract_docx_structure_enhanced
                        full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                        # Calculate section stats
                        section_words = len(full_text.split())
                        section_paras = len(section_df[section_df['source_type'] == 'paragraph'])
                        section_tables = len(section_df[section_df['source_type'] == 'table'])
                        
                        llm_summary_rows.append({
                            'header_1': header,
                            'llm_paragraph': full_text if full_text else "",
                            'n_words': section_words,
                            'n_paragraphs': section_paras,
                            'n_tables': section_tables
                        })

                    progress_bar.progress(0.5, text="Secciones extraídas.")
                    
                    # Create exploded dataframe
                    llm_summary_df = pd.DataFrame(llm_summary_rows)
                    exploded_df = llm_summary_df.assign(
                        llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                    ).explode('llm_paragraph')
                    exploded_df = exploded_df.reset_index(drop=True)
                    exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                    
                    # Get full text
                    full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                    
                    # Store in session state
                    file_size = os.path.getsize(tmp_file.name)
                    n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                    n_paragraphs = len(exploded_df)
                    
                    st.session_state['full_document_text_tab2'] = full_document_text
                    st.session_state['document_stats_tab2'] = {
                        'file_size': file_size,
                        'n_words': n_words,
                        'n_paragraphs': n_paragraphs
                    }
                    st.session_state['exploded_df_tab2'] = exploded_df
                    st.session_state['extraction_df_tab2'] = df
                    st.session_state['tables_data_tab2'] = tables_data
                    st.session_state['extraction_stats_tab2'] = extraction_stats
                    st.session_state['validation_results_tab2'] = validation_results
                    st.session_state['sections_df_tab2'] = llm_summary_df
                    st.session_state['selected_sections_tab2'] = list(header_1_values)  # Select all by default
                    st.session_state['last_file_hash_tab2'] = file_hash
                    st.session_state['document_extracted_tab2'] = True
                    
                    try:
                        os.unlink(tmp_file.name)
                    except:
                        pass
                    
                    progress_bar.progress(1.0, text="Extracción completa.")
                    st.rerun()  # Rerun to show extraction results
                    
                except Exception as e:
                    st.error(f"Error procesando el documento: {e}")
                    import traceback
                    st.error(traceback.format_exc())
                    st.stop()
        
        # Show extraction results if document is extracted
        if st.session_state.get('document_extracted_tab2', False) and not file_changed:
            st.success("✅ Documento extraído con éxito")
            
            # Download button for extracted document structure
            extraction_df = st.session_state.get('extraction_df_tab2', pd.DataFrame())
            if not extraction_df.empty:
                excel_data = to_excel(extraction_df)
                # Get filename from extraction_df or use default
                filename_base = extraction_df['filename'].iloc[0] if 'filename' in extraction_df.columns and not extraction_df['filename'].empty else "documento"
                filename_base = filename_base.replace('.docx', '').replace('.doc', '')
                st.download_button(
                    label="📥 Descargar estructura extraída del documento (Excel)",
                    data=excel_data,
                    file_name=f"estructura_documento_tab2_{filename_base}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_extraction_tab2"
                )
                st.caption("El archivo incluye todas las columnas de encabezados (header_1 a header_6), contenido, tipo de fuente, y metadatos de extracción.")
            
            # Display header_1 sections and their content
            extraction_df = st.session_state.get('extraction_df_tab2', pd.DataFrame())
            if not extraction_df.empty:
                with st.expander("📋 Ver estructura extraída del documento (encabezados nivel 1 y contenido)", expanded=False):
                    st.markdown("**Estructura del documento extraído (solo encabezados nivel 1):**")
                    
                    # Get unique header_1 values
                    header_1_sections = extraction_df[extraction_df['header_1'].notna() & (extraction_df['header_1'] != '')]['header_1'].unique()
                    
                    for h1 in header_1_sections:
                        st.markdown(f"### {h1}")
                        
                        # Get all content for this header_1 section
                        section_df = extraction_df[extraction_df['header_1'] == h1]
                        section_content = section_df[section_df['source_type'] == 'paragraph']['content'].tolist()
                        
                        # Display content
                        if section_content:
                            full_text = '\n\n'.join([str(c) for c in section_content if pd.notna(c) and str(c).strip()])
                            if full_text.strip():
                                st.text(full_text)
                                st.caption(f"Total: {len(full_text):,} caracteres")
                        else:
                            st.info("Esta sección no tiene contenido de párrafos extraído.")
                        
                        st.markdown("---")
            
            sections_df = st.session_state.get('sections_df_tab2', pd.DataFrame())
            
            if not sections_df.empty:
                header_1_values = sections_df['header_1'].tolist()
                
                # Section selector - simplified, no warnings or diagnostics
                st.markdown("### 🔍 Selección de Secciones para Evaluación")
                st.info("Selecciona las secciones que deseas incluir en la evaluación. Por defecto, todas las secciones están seleccionadas.")
                
                # Guidance about section extraction
                if len(header_1_values) == 0:
                    st.error("⚠️ **No se detectaron secciones en el documento.** Esto puede deberse a que el documento no usa estilos de encabezado de Word (Heading 1, Heading 2, etc.). Por favor, verifica que tu documento tenga encabezados formateados correctamente.")
                elif len(header_1_values) < 3:
                    st.warning("⚠️ **Se detectaron pocas secciones** en el documento. Si esperabas más secciones, verifica que el documento use estilos de encabezado de Word (Heading 1, Heading 2, etc.) para identificar las secciones principales.")
                
                # Initialize selected sections if not exists
                if 'selected_sections_tab2' not in st.session_state:
                    st.session_state['selected_sections_tab2'] = list(header_1_values)
                
                # Section selection interface
                selected_sections = st.session_state.get('selected_sections_tab2', list(header_1_values)).copy()
                col1, col2 = st.columns([3, 1])
                
                with col2:
                    if st.button("✅ Seleccionar Todas", key="select_all_sections_tab2"):
                        st.session_state['selected_sections_tab2'] = list(header_1_values)
                        st.rerun()
                    
                    if st.button("❌ Deseleccionar Todas", key="deselect_all_sections_tab2"):
                        st.session_state['selected_sections_tab2'] = []
                        st.rerun()
                
                with col1:
                    st.markdown("**Secciones disponibles:**")
                    for section in header_1_values:
                        section_info = sections_df[sections_df['header_1'] == section].iloc[0]
                        is_selected = section in selected_sections
                        
                        checkbox_label = f"**{section}** ({section_info['n_words']:,} palabras, {section_info['n_paragraphs']} párrafos)"
                        
                        checkbox_key = f"section_checkbox_{section}_tab2"
                        new_selection = st.checkbox(checkbox_label, value=is_selected, key=checkbox_key)
                        
                        if new_selection and section not in selected_sections:
                            selected_sections.append(section)
                        elif not new_selection and section in selected_sections:
                            selected_sections.remove(section)
                        
                        # Add expandable preview of extracted content
                        with st.expander(f"👁️ Ver contenido: {section}", expanded=False):
                            section_content = section_info['llm_paragraph']
                            if section_content and section_content.strip():
                                # Show first 500 characters as preview, full content in expandable
                                preview_text = section_content[:500] + "..." if len(section_content) > 500 else section_content
                                st.text_area(
                                    "Contenido extraído:",
                                    value=section_content,
                                    height=200,
                                    key=f"content_preview_{section}_tab2",
                                    label_visibility="collapsed"
                                )
                                st.caption(f"Total: {len(section_content):,} caracteres")
                            else:
                                st.info("Esta sección no tiene contenido extraído.")
                            
                            # Show table-extracted text
                            tables_data = st.session_state.get('tables_data_tab2', [])
                            section_tables = [t for t in tables_data if t.get('section') == section]
                            
                            if section_tables:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                for table_info in section_tables:
                                    table_num = table_info.get('table_number', 'N/A')
                                    table_data = table_info.get('data', [])
                                    if table_data:
                                        # Format table as text
                                        table_text = '\n'.join([' | '.join(str(cell) for cell in row) for row in table_data])
                                        st.text_area(
                                            f"Tabla {table_num}:",
                                            value=table_text,
                                            height=150,
                                            key=f"table_preview_{section}_table{table_num}_tab2",
                                            label_visibility="collapsed"
                                        )
                                        st.caption(f"Tabla {table_num}: {len(table_data)} filas, {len(table_data[0]) if table_data else 0} columnas")
                            else:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                st.info("No se encontraron tablas en esta sección.")
                
                # Update session state
                st.session_state['selected_sections_tab2'] = selected_sections
                
                # Show selection summary - simplified
                if selected_sections:
                    selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
                    total_selected_words = selected_df['n_words'].sum()
                    total_selected_paras = selected_df['n_paragraphs'].sum()
                    
                    # Estimate tokens
                    if encoding:
                        selected_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
                        estimated_tokens = len(encoding.encode(selected_text))
                    else:
                        estimated_tokens = total_selected_words * 1.2  # Rough estimate
                    
                    # Warn if approaching limit
                    if estimated_tokens > 100000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.warning(f"⚠️ **Advertencia de límite de contexto:** Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). El sistema procesa hasta 110,000 tokens (aproximadamente 150-200 páginas). Si el documento excede este límite, será truncado automáticamente.")
                    elif estimated_tokens > 80000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.info(f"ℹ️ Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). Estás dentro del límite de 110,000 tokens (aproximadamente 150-200 páginas).")
                    
                    st.success(f"✅ {len(selected_sections)} secciones seleccionadas | "
                              f"{total_selected_words:,} palabras | "
                              f"~{estimated_tokens:,} tokens estimados")
    
    # Rubric and Criteria Selection Section (moved after document extraction)
    st.markdown("---")
    st.markdown("### 📋 Selección de Rúbricas y Criterios")
    
    # Only show rubric selection if document is extracted
    if st.session_state.get('document_extracted_tab2', False):
        # All available rubrics
        all_rubrics = {
            # "Participación de Actores (durante el proyecto)": engagement_rubric,  # Commented out per user request
            # "Desempeño del proyecto (según informe de evaluación)": performance_rubric,  # Commented out per user request
            "Participación durante la evaluación (metodología)": parteval_rubric,
            "Integración del Enfoque de Género": gender_rubric,
            # "Transición Justa: Enfoque Tradicional": tj_traditional_rubric,  # Commented out per user request
            "Integración del Enfoque de Transición Justa: Enfoque Moderno": tj_just_transition_rubric
        }

        # Step 1: Select Rubrics
        st.markdown("#### 1. Seleccione las rúbricas a aplicar:")
        selected_rubric_names = st.multiselect(
            "Rúbricas:",
            options=list(all_rubrics.keys()),
            default=st.session_state['selected_rubrics_tab2'],
            key='rubric_selector_tab2'
        )
        st.session_state['selected_rubrics_tab2'] = selected_rubric_names

        # Step 2: Select Criteria within each rubric
        if selected_rubric_names:
            st.markdown("#### 2. Seleccione los criterios específicos:")
            
            for rubric_name in selected_rubric_names:
                rubric_dict = all_rubrics[rubric_name]
                
                with st.expander(f"📋 {rubric_name} ({len(rubric_dict)} criterios disponibles)", expanded=True):
                    # Select all checkbox for this rubric
                    select_all_key = f"select_all_{rubric_name}_tab2"
                    select_all = st.checkbox(f"Seleccionar todos los criterios", key=select_all_key)
                    
                    # Initialize criteria selection for this rubric
                    if rubric_name not in st.session_state['selected_criteria_tab2']:
                        st.session_state['selected_criteria_tab2'][rubric_name] = []
                    
                    # Show criteria checkboxes
                    selected_criteria = []
                    for criterion in rubric_dict.keys():
                        # Default checked if select_all or previously selected
                        default_value = select_all or criterion in st.session_state['selected_criteria_tab2'][rubric_name]
                        
                        is_selected = st.checkbox(
                            f"{criterion}",
                            value=default_value,
                            key=f"criterion_{rubric_name}_{criterion}_tab2"
                        )
                        
                        if is_selected:
                            selected_criteria.append(criterion)
                    
                    # Update session state
                    st.session_state['selected_criteria_tab2'][rubric_name] = selected_criteria
                    
                    st.info(f"Criterios seleccionados: {len(selected_criteria)}/{len(rubric_dict)}")

        # Show summary of selections
        if selected_rubric_names:
            total_criteria = sum(len(st.session_state['selected_criteria_tab2'].get(r, [])) for r in selected_rubric_names)
            st.success(f"Total: {len(selected_rubric_names)} rúbricas, {total_criteria} criterios seleccionados")
        else:
            st.info("ℹ️ Selecciona al menos una rúbrica para continuar con la evaluación.")
    else:
        st.info("ℹ️ Por favor extrae el documento primero para poder seleccionar las rúbricas y criterios.")
    
    # Process and Evaluate button
    st.markdown("---")
    st.markdown("### ⚙️ Procesamiento y Evaluación")
    
    # Warning about AI results verification
    st.warning("""
    **⚠️ Importante - Verificación de Resultados:**
    
    Los resultados generados por esta herramienta utilizan inteligencia artificial y deben ser **verificados y corroborados** antes de su uso.
    
    - La IA puede cometer errores, interpretaciones incorrectas o pasar por alto información relevante
    - Los análisis y puntuaciones son **sugerencias** basadas en el contenido del documento, no son definitivos
    - Se recomienda revisar manualmente las evidencias citadas y validar las conclusiones
    - Los resultados deben ser contrastados con conocimiento experto y documentación adicional cuando sea necesario
    
    Esta herramienta es un **asistente de análisis** que facilita la revisión, pero la responsabilidad final de la evaluación recae en el usuario.
    """)
    
    def evaluate_criterion_with_llm(document_text, criterion, descriptions, max_retries=3):
        """Analyze document against criterion with retry logic"""
        import time

        for attempt in range(max_retries):
            try:
                # Truncate to ~110K tokens to maximize context while leaving room for:
                # - System prompt (~50 tokens)
                # - User prompt template + criterion + scoring levels (~500-2000 tokens)
                # - Response tokens (6500 tokens)
                # - Safety buffer (~5000 tokens)
                # Total: 128K - 50 - 2000 - 6500 - 5000 = ~110K tokens for document
                combined_text = truncate_to_token_limit(document_text, max_tokens=110000, encoding_obj=encoding)

                # Now do the expensive analysis on focused content
                prompt = f"""Evaluate this document against: {criterion}

    Scoring levels: {json.dumps(descriptions)}

    Relevant document sections:
    {combined_text}

    IMPORTANTE: Proporciona tu respuesta SIEMPRE en español, incluso si el documento está en inglés.
    
    Provide JSON with:
    {{"analysis": "detailed 2-3 paragraphs IN SPANISH", "score": 1-5, "evidence": ["quote 1", "quote 2", "quote 3", "etc - 5-8 key quotes from the text as an array"]}}"""

                response = client.chat.completions.create(
                    model="gpt-5-mini",
                    messages=[
                        {"role": "system", "content": "Eres un evaluador experto de documentos. Siempre debes responder en español, incluso si el documento está en inglés."},
                        {"role": "user", "content": prompt}
                    ],
                    max_completion_tokens=6500,
                    reasoning_effort="minimal",
                    timeout=120  # 2 minute timeout per request
                )

                content = response.choices[0].message.content.strip()
                # Remove markdown code fences if present
                if content.startswith('```'):
                    # Remove opening fence (```json or ```)
                    content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                    # Remove closing fence
                    if content.endswith('```'):
                        content = content.rsplit('```', 1)[0]
                    content = content.strip()

                result = json.loads(content)
                # Normalize evidence field: convert array to string if needed
                if isinstance(result.get('evidence'), list):
                    result['evidence'] = '\n'.join(result['evidence'])
                return result

            except json.JSONDecodeError as e:
                # If JSON parsing fails, return a default structure
                return {
                    "analysis": f"Failed to parse JSON: {str(e)}. Raw response: {response.choices[0].message.content[:200]}",
                    "score": 3,
                    "evidence": "Unable to parse structured response",
                    "error": f"JSON parsing error: {str(e)}"
                }
            except Exception as e:
                # Check if it's a rate limit error
                error_msg = str(e)
                if "rate_limit" in error_msg.lower() or "429" in error_msg:
                    if attempt < max_retries - 1:
                        wait_time = (2 ** attempt) * 2  # Exponential backoff: 2s, 4s, 8s
                        time.sleep(wait_time)
                        continue

                # If last attempt or non-rate-limit error, return error
                return {
                    "analysis": f"Error during evaluation: {error_msg}",
                    "score": 0,
                    "evidence": "",
                    "error": f"API error (attempt {attempt + 1}/{max_retries}): {error_msg}"
                }

        # If we exhausted all retries
        return {
            "analysis": "Failed after multiple retry attempts",
            "score": 0,
            "evidence": "",
            "error": f"Failed after {max_retries} attempts"
        }

    # Function to evaluate a single text chunk
    def evaluate_single_chunk(text_chunk, criterion, descriptions):
        """Evaluate a single text chunk against a criterion with expanded analysis and evidence"""
        import json

        # Build prompt
        prompt = f"""
        Estás evaluando un documento contra un criterio específico.
        
        Criterio: {criterion}
        
        Descripciones de los niveles de puntuación:
        {json.dumps(descriptions, indent=2)}
        
        Contenido del documento a evaluar:
        {text_chunk}
        
        Analiza qué tan bien el documento cumple con este criterio. Proporciona:
        
        1. Un análisis DETALLADO (2-3 párrafos) que explique a fondo el razonamiento detrás de tu evaluación. Proporciona un razonamiento profundo que abarque los aspectos del criterio.
        
        2. Una puntuación de 1-5 (donde 1 es la más baja y 5 es la más alta).
        
        3. EVIDENCIA del documento que respalde tu puntuación. Incluye entre 5-8 citas textuales del documento, indicando cómo cada fragmento contribuye a tu evaluación.
        
        Formatea tu respuesta como un objeto JSON con las siguientes claves:
        {{"analysis": "tu análisis detallado aquí", "score": puntuación_numérica_entre_1_y_5, "evidence": "citas textuales del documento (5-8 párrafos)"}}
        
        Devuelve solo el objeto JSON, nada más.
        """

        # Call LLM using new Responses API
        try:
            response = client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "system", "content": "Eres un experto evaluador de documentos que proporciona análisis detallados basados en criterios específicos. Tu evidencia cita fragmentos del texto original."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                max_completion_tokens=6500,
                reasoning_effort="minimal"
            )
            raw = response.choices[0].message.content
            if not raw or not raw.strip():
                return {'score': 0, 'analysis': 'Empty response from API', 'evidence': ''}

            # Remove markdown code fences if present
            content = raw.strip()
            if content.startswith('```'):
                content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                if content.endswith('```'):
                    content = content.rsplit('```', 1)[0]
                content = content.strip()

            parsed = json.loads(content)
            # Normalize evidence field: convert array to string if needed
            if isinstance(parsed.get('evidence'), list):
                parsed['evidence'] = '\n'.join(parsed['evidence'])
            return parsed
        except Exception as e:
            return {'score': 0, 'analysis': f'Error: {str(e)}', 'evidence': ''}

    # Function to synthesize evaluations
    def synthesize_evaluations(chunk_results, criterion, descriptions):
        """Synthesize evaluations from multiple document chunks with expanded analysis and evidence"""
        import json

        # Extract and format the individual evaluations for the synthesis
        individual_evals = []
        all_evidence = []

        for i, result in enumerate(chunk_results):
            individual_evals.append(f"Evaluación del fragmento {i+1}:\n" +
                                    f"Puntuación: {result.get('score', 0)}\n" +
                                    f"Análisis: {result.get('analysis', '')}")

            # Collect all evidence
            evidence = result.get('evidence', '')
            if evidence:
                all_evidence.append(f"Evidencia del fragmento {i+1}:\n{evidence}")

        # Define separator outside the f-string to avoid backslash issues
        separator = "\n\n"

        # Create a synthesis prompt
        synthesis_prompt = f"""
        Has evaluado un documento dividido en múltiples fragmentos contra el criterio: {criterion}
        
        Aquí están las evaluaciones individuales de cada fragmento:
        
        {separator.join(individual_evals)}
        
        Basándote en estas evaluaciones individuales, proporciona:
        
        1. Un análisis DETALLADO (2-3 párrafos) que integre los hallazgos clave de todos los fragmentos. Este análisis debe ser comprensivo y abarcar los aspectos relevantes encontrados en el documento.
        
        2. Una puntuación general de 1-5 (puedes promediar las puntuaciones o ajustar según sea necesario)
        
        3. Las evidencias más importantes del documento. Selecciona las 8-10 citas textuales más relevantes de los fragmentos individuales.
        
        Formatea tu respuesta como un objeto JSON con las siguientes claves:
        {{"analysis": "tu análisis global detallado aquí", "score": puntuación_general_entre_1_y_5, "evidence": "las citas textuales más relevantes del documento (8-10 párrafos)"}}
        
        Devuelve solo el objeto JSON, nada más.
        """

        # Call LLM for synthesis using new Responses API
        try:
            response = client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "system", "content": "Eres un experto evaluador de documentos que sintetiza análisis de múltiples fragmentos de texto para producir evaluaciones detalladas con evidencia textual."},
                    {"role": "user", "content": synthesis_prompt}
                ],
                response_format={"type": "json_object"},
                max_completion_tokens=6500,
                reasoning_effort="minimal"
            )
            raw = response.choices[0].message.content
            if not raw or not raw.strip():
                raise ValueError("Empty response from API")

            # Remove markdown code fences if present
            content = raw.strip()
            if content.startswith('```'):
                content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                if content.endswith('```'):
                    content = content.rsplit('```', 1)[0]
                content = content.strip()

            parsed = json.loads(content)
            # Normalize evidence field: convert array to string if needed
            if isinstance(parsed.get('evidence'), list):
                parsed['evidence'] = '\n'.join(parsed['evidence'])
            return parsed
        except Exception as e:
            # If synthesis fails, combine results manually in a more limited way
            avg_score = sum(r.get('score', 0) for r in chunk_results) / len(chunk_results)
            # Take only the first paragraph of each analysis to avoid token limits
            analysis_parts = []
            for r in chunk_results:
                analysis = r.get('analysis', '')
                first_para = analysis.split('\n\n')[0] if '\n\n' in analysis else analysis
                analysis_parts.append(first_para)

            # Take only the first few evidence items
            evidence_parts = []
            evidence_count = 0
            for evidence in all_evidence:
                parts = evidence.split('\n\n')
                # Add up to 2 evidence parts per chunk
                for part in parts[:2]:
                    if evidence_count < 8:  # Limit to 8 total evidence parts
                        evidence_parts.append(part)
                        evidence_count += 1

            return {
                'score': avg_score,
                'analysis': separator.join(analysis_parts),
                'evidence': separator.join(evidence_parts)
            }
    
    if st.button('🚀 Procesar y Evaluar', key='process_evaluate_tab2', type="primary"):
        # Check prerequisites
        if not st.session_state.get('document_extracted_tab2', False):
            st.error("❌ Por favor extrae el documento primero usando el botón 'Extraer Documento'.")
            st.stop()
        
        # Get selected rubrics from session state
        selected_rubric_names = st.session_state.get('selected_rubrics_tab2', [])
        if not selected_rubric_names:
            st.error("Por favor seleccione al menos una rúbrica.")
            st.stop()
        
        # Calculate total criteria
        total_criteria = sum(len(st.session_state.get('selected_criteria_tab2', {}).get(r, [])) for r in selected_rubric_names)
        if total_criteria == 0:
            st.error("Por favor seleccione al menos un criterio.")
            st.stop()
        
        # Check if document needs re-extraction (shouldn't happen, but safety check)
        if uploaded_file is not None:
            file_hash = hash(uploaded_file.getvalue())
            if st.session_state.get('last_file_hash_tab2') != file_hash:
                st.warning("⚠️ El documento ha cambiado. Por favor extrae el documento nuevamente.")
                st.stop()
        
        # Evaluate with selected rubrics and criteria
        # Get selected sections or use full document
        selected_sections = st.session_state.get('selected_sections_tab2', [])
        sections_df = st.session_state.get('sections_df_tab2', pd.DataFrame())
        
        if selected_sections and not sections_df.empty:
            # Filter to selected sections only
            selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
            document_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
            st.info(f"📌 Evaluando {len(selected_sections)} secciones seleccionadas")
        else:
            # Fallback to full document
            document_text = st.session_state.get('full_document_text_tab2', '')
            if not selected_sections:
                st.warning("⚠️ No hay secciones seleccionadas. Usando documento completo.")
        
        if not document_text:
            st.error("No se pudo recuperar el texto del documento.")
            st.stop()

        # Build filtered rubrics based on selection
        # Get selected rubrics from session state
        selected_rubric_names = st.session_state.get('selected_rubrics_tab2', [])
        
        # Define all_rubrics for evaluation
        all_rubrics = {
            "Participación durante la evaluación (metodología)": parteval_rubric,
            "Integración del Enfoque de Género": gender_rubric,
            "Integración del Enfoque de Transición Justa: Enfoque Moderno": tj_just_transition_rubric
        }
        
        rubrics_to_evaluate = []
        for rubric_name in selected_rubric_names:
            selected_criteria_list = st.session_state['selected_criteria_tab2'].get(rubric_name, [])
            if selected_criteria_list:
                # Filter the rubric to only include selected criteria
                full_rubric = all_rubrics[rubric_name]
                filtered_rubric = {k: v for k, v in full_rubric.items() if k in selected_criteria_list}
                rubrics_to_evaluate.append((rubric_name, filtered_rubric))

        if not rubrics_to_evaluate:
            st.error("No hay criterios seleccionados para evaluar.")
            st.stop()

        # Evaluate
        st.info("Iniciando evaluación de criterios...")
        rubric_results = []
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import time
        MAX_WORKERS = 3  # Reduced to avoid rate limiting
        
        def eval_one_criterion(args):
            crit, descriptions, dimension, rubric_name = args
            try:
                result = evaluate_criterion_with_llm(document_text, crit, descriptions)
                # Ensure result is a dictionary
                if not isinstance(result, dict):
                    result = {'score': 0, 'analysis': str(result), 'evidence': '', 'error': 'Invalid result format'}
                return {
                    'Criterio': crit,
                    'Dimensión': dimension,
                    'Score': result.get('score', 0),
                    'Análisis': str(result.get('analysis', '')),
                    'Evidencia': str(result.get('evidence', '')),
                    'Error': str(result.get('error', '')) if 'error' in result else '',
                    'Rúbrica': rubric_name
                }
            except Exception as e:
                return {
                    'Criterio': crit,
                    'Dimensión': dimension,
                    'Score': 0,
                    'Análisis': '',
                    'Evidencia': '',
                    'Error': str(e),
                    'Rúbrica': rubric_name
                }
        
        for rubric_name, rubric_dict in rubrics_to_evaluate:
            if not rubric_dict:
                continue

            rubric_analysis_data = []
            n_criteria = len(rubric_dict)
            progress = st.progress(0, text=f"Preparando evaluación de {n_criteria} criterios para: {rubric_name}...")

            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                progress.progress(0.01, text=f"Enviando {n_criteria} criterios para evaluación paralela...")
                futures = {}
                for idx, (crit, rubric_data) in enumerate(rubric_dict.items()):
                    future = executor.submit(eval_one_criterion, (
                        crit,
                        rubric_data['valores'] if isinstance(rubric_data, dict) else rubric_data,
                        rubric_data.get('dimension', 'No especificada') if isinstance(rubric_data, dict) else 'No especificada',
                        rubric_name
                    ))
                    futures[future] = (crit, idx)
                    # Small delay to avoid overwhelming API at startup
                    if idx > 0 and idx % MAX_WORKERS == 0:
                        time.sleep(0.5)

                progress.progress(0.05, text=f"Esperando resultados de evaluación...")
                completed = 0
                for future in as_completed(futures):
                    result = future.result()
                    rubric_analysis_data.append(result)
                    completed += 1
                    crit, idx = futures[future]
                    progress.progress(0.05 + (completed / n_criteria * 0.95), text=f"Completado {completed}/{n_criteria}: {crit}")
            
            rubric_results.append((rubric_name, pd.DataFrame(rubric_analysis_data)))
        
        # Store results in session state for persistence
        st.session_state['tab2_results'] = rubric_results
        
        # Display results
        if rubric_results:
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''
                    
                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Criterio', 'Dimensión', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Download ZIP
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    # Export as XLSX instead of CSV
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_rubrica_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_rubricas.zip",
                mime="application/zip"
            )
        else:
            st.warning("No se generaron resultados para ninguna rúbrica.")
    else:
        # Check if there are persisted results in session state
        if st.session_state.get('tab2_results') is not None:
            rubric_results = st.session_state['tab2_results']
            
            st.markdown("### 📊 Resultados guardados")
            
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''
                    
                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Criterio', 'Dimensión', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Download ZIP
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_rubrica_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_rubricas.zip",
                mime="application/zip",
                key="tab2_download_persisted"
            )
            
            # Clear results button
            if st.button("🗑️ Limpiar resultados", key="clear_tab2_results"):
                st.session_state['tab2_results'] = None
                st.rerun()
        else:
            st.info("Suba un documento, seleccione rúbricas y criterios, luego presione 'Procesar y Evaluar'.")


with tab4:
    st.header("Pregúntale a tus Documentos")

    # Presentation box
    st.info("""

    Arrastra uno o más DOCX/TXT (max 200MB) y conversa directamente con su contenido. Este chat responde solo con la información de los archivos cargados—no recurre a fuentes externas—para ayudarte a revisar propuestas antes de enviarlas, aclarar pasajes de informes, preparar notas técnicas y comprobar coherencias entre objetivos, actividades y resultados.

    Al formular tus preguntas, indica el nivel de detalle que necesitas y pide que las respuestas incluyan citas breves entre comillas y metadatos (título, sección/página, año) para mantener la trazabilidad. Si un dato no existe en los archivos, se marcará "ND" sin inventar información.

    **Úsalo para:**
    - Aprender de experiencias de evaluación
    - Profundizar en el conocimiento de secciones de los informes
    - Preparar notas técnicas a partir de esta evidencia
    - Verificar la coherencia entre objetivos, actividades y resultados
    - Generar tablas copiables a Excel
    - Comparar varios documentos (hasta de 200MB) en una misma conversación

    *Nota: El chat mantiene memoria de la conversación durante la sesión activa.*
    """)

    # Session state for chat and document
    if 'doc_chat_history' not in st.session_state:
        st.session_state['doc_chat_history'] = []
    if 'doc_chat_docs' not in st.session_state:
        st.session_state['doc_chat_docs'] = []

    # Warning box about document requirements for chat
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos:**
    
    **📝 Formatos aceptados:**
    - Archivos **.docx** (Word 2007 o posterior)
    - Archivos **.txt** (texto plano)
    - Puedes subir **múltiples archivos** (hasta 200MB en total)
    
    **📝 Para documentos Word (.docx):**
    - El documento debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del documento deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente.
    - Para documentos grandes, el sistema usa RAG (Retrieval Augmented Generation) automáticamente
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos muy grandes (>100K caracteres, aproximadamente >40 páginas) se procesan con RAG para mejor eficiencia
    - Documentos pequeños se procesan con contexto completo (más eficiente)
    - Documentos que excedan el límite máximo serán truncados automáticamente
    
    **✅ Mejores prácticas:**
    - Para documentos Word: usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que los documentos estén guardados correctamente antes de subirlos
    """)
    
    uploaded_files = st.file_uploader("Sube uno o más archivos DOCX o TXT para chatear:", type=["docx", "txt"], accept_multiple_files=True)
    if 'doc_chat_docs' not in st.session_state:
        st.session_state['doc_chat_docs'] = []
    if uploaded_files:
        # Only reset docs and chat history if files changed
        uploaded_filenames = sorted([f.name for f in uploaded_files])
        existing_filenames = sorted([doc['filename'] for doc in st.session_state['doc_chat_docs']]) if st.session_state['doc_chat_docs'] else []
        if uploaded_filenames != existing_filenames:
            st.session_state['doc_chat_docs'] = []

            # Load documents
            for uploaded_file in uploaded_files:
                try:
                    if uploaded_file.name.endswith(".docx"):
                        from docx import Document
                        doc = Document(uploaded_file)
                        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    else:
                        full_text = uploaded_file.read().decode("utf-8", errors="ignore")
                    st.session_state['doc_chat_docs'].append({
                        "filename": uploaded_file.name,
                        "text": full_text
                    })
                except Exception as e:
                    st.error(f"Error al procesar el documento '{uploaded_file.name}': {str(e)}")

            # Calculate total document size for smart RAG decision
            total_text = "\n\n".join([doc['text'] for doc in st.session_state['doc_chat_docs']])
            total_chars = len(total_text)
            st.session_state['doc_chat_total_text'] = total_text
            st.session_state['doc_chat_total_chars'] = total_chars

            # Pre-compute embeddings only if documents are large (> 100K chars)
            # For small docs, we'll use full context directly (more efficient)
            if total_chars > 100000:
                with st.spinner("Procesando embeddings de documentos (solo se hace una vez)..."):
                    # Chunk documents
                    def chunk_text(text, chunk_size=2000, overlap=300):
                        chunks = []
                        start = 0
                        while start < len(text):
                            end = min(start + chunk_size, len(text))
                            chunks.append(text[start:end])
                            start += chunk_size - overlap
                        return chunks

                    all_chunks = []
                    for doc in st.session_state['doc_chat_docs']:
                        all_chunks.extend(chunk_text(doc['text']))

                    # Embed all chunks once and cache
                    try:
                        emb_model = "text-embedding-3-large"
                        chunk_embs_resp = client.embeddings.create(input=all_chunks, model=emb_model)
                        chunk_embs = [item.embedding for item in chunk_embs_resp.data]

                        # Cache in session state
                        st.session_state['doc_chat_chunks'] = all_chunks
                        st.session_state['doc_chat_embeddings'] = chunk_embs
                        st.session_state['doc_chat_use_rag'] = True
                        st.info(f"📚 Documentos grandes ({total_chars:,} chars) - usando RAG con {len(all_chunks)} fragmentos en caché")
                    except Exception as e:
                        st.error(f"Error al generar embeddings: {str(e)}")
                        st.session_state['doc_chat_use_rag'] = False
            else:
                st.session_state['doc_chat_use_rag'] = False
                st.info(f"📄 Documentos pequeños ({total_chars:,} chars) - usando contexto completo (más eficiente)")

            st.session_state['doc_chat_history'] = []  # Only reset chat when new files uploaded
            st.success(f"{len(st.session_state['doc_chat_docs'])} documento(s) cargado(s) y listo(s) para chatear.")

    # Show filenames and previews
    if st.session_state['doc_chat_docs']:
        st.info(f"Documentos activos: {', '.join([doc['filename'] for doc in st.session_state['doc_chat_docs']])}")
        for doc in st.session_state['doc_chat_docs']:
            with st.expander(f"Vista previa: {doc['filename']} (primeros 500 caracteres)"):
                st.write(doc['text'][:500] + ("..." if len(doc['text']) > 500 else ""))

    # Chat interface
    if st.session_state['doc_chat_docs']:
        with st.form("doc_chat_form", clear_on_submit=True):
            user_input = st.text_area("Escribe tu preguntas (para las respuestas se considera el texto completo de los documentos cargados):", key="doc_chat_input")
            submitted = st.form_submit_button("Enviar pregunta")
            if submitted and user_input.strip():
                # Add user message to history
                if 'doc_chat_history' not in st.session_state:
                    st.session_state['doc_chat_history'] = []
                st.session_state['doc_chat_history'].append({"role": "user", "content": user_input.strip()})

                question = user_input.strip()
                use_rag = st.session_state.get('doc_chat_use_rag', False)

                # Smart context selection: use full text for small docs, RAG for large docs
                if not use_rag:
                    # Small documents: use full context (up to 110K tokens)
                    # This is more efficient - no chunking, no embeddings, just direct LLM call
                    context = truncate_to_token_limit(
                        st.session_state.get('doc_chat_total_text', ''), 
                        max_tokens=110000, 
                        encoding_obj=encoding
                    )

                    messages = [
                        {"role": "system", "content": "Eres un asistente experto en análisis documental. Responde usando solo la información del documento proporcionado."},
                        {"role": "system", "content": f"Texto del documento:\n{context}"}
                    ]
                    for msg in st.session_state['doc_chat_history'][-5:]:
                        messages.append(msg)

                    try:
                        response = client.chat.completions.create(
                            model="gpt-5-mini",
                            messages=messages,
                            max_completion_tokens=4096,
                            reasoning_effort="minimal"
                        )
                        answer = response.choices[0].message.content.strip()
                        st.session_state['doc_chat_history'].append({"role": "assistant", "content": answer})
                    except Exception as e:
                        st.session_state['doc_chat_history'].append({"role": "assistant", "content": f"[Error al obtener respuesta: {str(e)}]"})

                else:
                    # Large documents: use RAG with cached embeddings
                    import numpy as np
                    import faiss
                    import re

                    # Get cached chunks and embeddings
                    all_chunks = st.session_state.get('doc_chat_chunks', [])
                    chunk_embs = st.session_state.get('doc_chat_embeddings', [])

                    if not all_chunks or not chunk_embs:
                        st.session_state['doc_chat_history'].append({"role": "assistant", "content": "[Error: No se encontraron fragmentos en caché.]"})
                    else:
                        # Embed user question (only embedding we need per query!)
                        try:
                            emb_model = "text-embedding-3-large"
                            question_emb = client.embeddings.create(input=question, model=emb_model).data[0].embedding
                        except Exception as e:
                            st.session_state['doc_chat_history'].append({"role": "assistant", "content": f"[Error al obtener embedding de la pregunta: {str(e)}]"})
                            question_emb = None

                        # Retrieve relevant chunks using FAISS
                        if question_emb:
                            dim = len(chunk_embs[0])
                            xb = np.array(chunk_embs).astype('float32')
                            index = faiss.IndexFlatIP(dim)

                            # Normalize for cosine similarity
                            faiss.normalize_L2(xb)
                            index.add(xb)  # FIX: Actually add vectors to the index!

                            xq = np.array([question_emb]).astype('float32')
                            faiss.normalize_L2(xq)

                            # Reduced from 50 to 15 for efficiency (15 chunks = ~30K chars)
                            top_n = min(15, len(all_chunks))
                            D, I = index.search(xq, top_n)
                            selected_chunks = [all_chunks[i] for i in I[0] if i < len(all_chunks)]

                            # Also retrieve chunks with exact keyword matches
                            stopwords = set([
                                'el','la','los','las','de','del','y','en','a','un','una','que','por','con','para','es','al','se','su','sus','o','u','como','más','menos','le','lo','su','the','and','of','in','to','for','is','on','at','by','an','or','as','be','are','was','were','from','it','this','that','with','but','not','can','may','do','does'
                            ])
                            qwords = [w for w in re.findall(r'\w+', question.lower()) if w not in stopwords and len(w) > 2]
                            keyword_chunks = []
                            for chunk in all_chunks:
                                chunk_lc = chunk.lower()
                                if any(qw in chunk_lc for qw in qwords):
                                    keyword_chunks.append(chunk)
                                    if len(keyword_chunks) >= 5:  # Limit keyword matches
                                        break

                            # Merge and deduplicate
                            seen = set()
                            merged_chunks = []
                            for chunk in selected_chunks + keyword_chunks:
                                if chunk not in seen:
                                    merged_chunks.append(chunk)
                                    seen.add(chunk)

                            context = '\n---\n'.join(merged_chunks)

                            # Send to LLM
                            messages = [
                                {"role": "system", "content": "Eres un asistente experto en análisis documental. Responde usando solo la información del documento proporcionado. Si la información no es explícita, infiere la respuesta usando pistas contextuales y tu capacidad de síntesis."},
                                {"role": "system", "content": f"Fragmentos relevantes del documento:\n{context}"}
                            ]
                            for msg in st.session_state['doc_chat_history'][-5:]:
                                messages.append(msg)

                            try:
                                response = client.chat.completions.create(
                                    model="gpt-5-mini",
                                    messages=messages,
                                    max_completion_tokens=4096,
                                    reasoning_effort="minimal"
                                )
                                answer = response.choices[0].message.content.strip()
                                st.session_state['doc_chat_history'].append({"role": "assistant", "content": answer})
                            except Exception as e:
                                st.session_state['doc_chat_history'].append({"role": "assistant", "content": f"[Error al obtener respuesta: {str(e)}]"})

        # Display chat history in a persistent, scrollable container
        st.markdown(
            '''
            <div style="height:600px; overflow-y:auto; border:1px solid #333; border-radius:8px; padding:1em; background:#18191a; margin-bottom:1em;">
            '''
            +
            "".join(
                f"<div style='margin-bottom:1em; color:#2980b9;'><b>Tú:</b> {msg['content']}</div>" if msg['role']=='user'
                else f"<div style='margin-bottom:1em; color:#27ae60;'><b>Asistente:</b> {msg['content']}</div>"
                for msg in st.session_state['doc_chat_history']
            )
            +
            '</div>',
            unsafe_allow_html=True
        )
    else:
        st.info("Sube uno o más documentos válidos para comenzar el chat.")

with tab3:
    st.header("Diagnóstico de Sostenibilidad del Proyecto")
    
    # Read rubric from Excel file
    import pandas as pd
    prodoc_rubric = {}
    
    try:
        # Load rubric from PRODOC_rubric.xlsx
        # df_rubric_prodoc = pd.read_excel('./PRODOC_rubric.xlsx', sheet_name='rubric')
        df_rubric_prodoc = pd.read_excel('./Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx', sheet_name='rubric')

        # Verify required columns exist
        required_cols = ['Dimensión', 'Criterio', 'Indicador']
        missing_cols = [col for col in required_cols if col not in df_rubric_prodoc.columns]
        if missing_cols:
            st.error(f"Columnas faltantes en el archivo Excel: {missing_cols}")
            st.info(f"Columnas disponibles: {list(df_rubric_prodoc.columns)}")
        
        # Helper function to extract numeric prefix for sorting (e.g., "1.1" -> (1, 1), "2.3" -> (2, 3))
        def extract_sort_key(text):
            import re
            if pd.isna(text):
                return (999, 999, 999)
            text = str(text).strip()
            # Match patterns like "1.", "1.1", "1.1.", "(3.1)", etc.
            match = re.match(r'[\(]?(\d+)(?:\.(\d+))?(?:\.(\d+))?', text)
            if match:
                parts = [int(p) if p else 0 for p in match.groups()]
                return tuple(parts)
            return (999, 999, 999)
        
        # Process each row to extract criteria and values
        for idx, row in df_rubric_prodoc.iterrows():
            indicador = row.get('Indicador', '')
            criterio = row.get('Criterio', '')
            dimension = row.get('Dimensión', 'No especificada')
            
            if pd.isna(indicador) or str(indicador).strip() == '':
                continue
            
            indicador = str(indicador).strip()
            criterio = str(criterio).strip() if not pd.isna(criterio) else ''
            dimension = str(dimension).strip() if not pd.isna(dimension) else 'No especificada'
            
            # Get level columns (Nivel 1, Nivel 2, etc.)
            level_cols = sorted([col for col in df_rubric_prodoc.columns if col.startswith('Nivel')])
            
            valores = []
            for col in level_cols:
                val = row[col]
                if not pd.isna(val) and str(val).strip() != '':
                    valores.append(str(val).strip())
            
            # Store with unique key (dimension + criterio + indicador) to avoid overwrites
            # since same criterio and indicador text can appear in different dimensions
            unique_key = f"{dimension}|{criterio}|{indicador}"
            prodoc_rubric[unique_key] = {
                'valores': valores, 
                'dimension': dimension,
                'criterio': criterio,
                'indicador': indicador,  # Keep original indicador text for display
                'sort_key': extract_sort_key(indicador)
            }
        
        # Show dimension breakdown
        dim_counts = {}
        for ind, data in prodoc_rubric.items():
            dim = data.get('dimension', 'N/A')
            dim_counts[dim] = dim_counts.get(dim, 0) + 1
        st.success(f"Rúbrica cargada: {len(prodoc_rubric)} indicadores en {len(dim_counts)} dimensiones.")
        
        # Download button for the rubric file (directly on page, no expander)
        try:
            with open('./Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx', 'rb') as f:
                st.download_button(
                    label="📥 Descargar archivo rúbrica de Sostenibilidad del proyecto",
                    data=f,
                    file_name="Evaluacion_sostenibilidad_rubric.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_sostenibilidad_rubric"
                )
        except FileNotFoundError:
            st.warning("Archivo de rúbrica no disponible para descarga.")
    except FileNotFoundError:
        st.error("No se encontró el archivo PRODOC_rubric.xlsx. Por favor, asegúrese de que existe en el directorio de la aplicación.")
    except Exception as e:
        st.error(f"Error al cargar la rúbrica desde PRODOC_rubric.xlsx: {str(e)}")
    
    # Instrucciones generales
    st.info("""
    **Instrucciones:**
    1.	Seleccione la rúbrica y los criterios específicos que desea analizar.
    2.	Seleccione las secciones adecuadas de diagnóstico según el documento cargado (algunos criterios son relevantes a Documentos de Proyecto y otras a informes de progreso).
    3.	Suba el archivo en formato DOCX correspondiente.
    4.	Presione el botón de Procesar y Analizar para analizar el documento.
    5.	Revise los resultados de cada rúbrica en las tablas interactivas.
    6.	Visualice las puntuaciones promedio por dimensión en los gráficos de barras.
    7.	Descargue todos los resultados y evidencias en archivos ZIP.

    """)

    # Initialize session state for selections and results persistence
    if 'selected_criteria_tab3' not in st.session_state:
        st.session_state['selected_criteria_tab3'] = []
    if 'selected_dimensions_tab3' not in st.session_state:
        st.session_state['selected_dimensions_tab3'] = []
    if 'tab3_results' not in st.session_state:
        st.session_state['tab3_results'] = None
    if 'document_extracted_tab3' not in st.session_state:
        st.session_state['document_extracted_tab3'] = False

    # Document upload
    st.markdown("### 📄 Carga de Documento")
    
    # Warning box about document requirements
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos PRODOC:**
    
    **📝 Formato del documento:**
    - Solo se aceptan archivos en formato **.docx** (Word 2007 o posterior)
    - El documento PRODOC debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del PRODOC deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente y las secciones no se identificarán.
    - Los PRODOCs deben tener una estructura clara con secciones bien definidas (Marco Lógico, Presupuesto, Cronograma, etc.)
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos que excedan este límite serán truncados automáticamente
    - Se recomienda dividir documentos muy extensos (más de ~180 páginas) en secciones más pequeñas si es necesario
    
    **✅ Mejores prácticas:**
    - Usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones principales del PRODOC
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que todas las secciones importantes (Marco Lógico, Presupuesto, etc.) tengan encabezados claros
    - Verifica que el documento esté guardado correctamente antes de subirlo
    """)
    
    uploaded_file_prodoc = st.file_uploader("Suba un archivo DOCX para evaluación:", type=["docx"], key="prodoc_file_uploader_tab3")
    
    # Document Extraction Section
    st.markdown("---")
    st.markdown("### 📥 Extracción de Documento")
    
    if uploaded_file_prodoc is not None:
        file_hash = hash(uploaded_file_prodoc.getvalue())
        file_changed = st.session_state.get('last_file_hash_tab3') != file_hash
        
        if file_changed:
            st.session_state['document_extracted_tab3'] = False
            st.session_state['last_file_hash_tab3'] = None
        
        if st.button("🔍 Extraer Documento", key="extract_document_tab3", type="primary"):
            if uploaded_file_prodoc is None:
                st.error("Por favor suba un archivo DOCX primero.")
                st.stop()
            
            with st.spinner("Extrayendo documento..."):
                try:
                    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp_file.write(uploaded_file_prodoc.read())
                    tmp_file.close()
                    
                    progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                    doc_result = docx2python(tmp_file.name)
                    
                    # Use enhanced extraction
                    df, tables_data, extraction_stats = extract_docx_structure_enhanced(tmp_file.name)
                    progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")
                    
                    # Extract sections
                    header_1_values = df['header_1'].dropna().unique()
                    llm_summary_rows = []
                    
                    for idx, header in enumerate(header_1_values):
                        section_df = df[df['header_1'] == header].copy()
                        full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                        section_words = len(full_text.split())
                        section_paras = len(section_df[section_df['source_type'] == 'paragraph'])
                        section_tables = len(section_df[section_df['source_type'] == 'table'])
                        
                        llm_summary_rows.append({
                            'header_1': header,
                            'llm_paragraph': full_text if full_text else "",
                            'n_words': section_words,
                            'n_paragraphs': section_paras,
                            'n_tables': section_tables
                        })

                    progress_bar.progress(0.5, text="Secciones extraídas.")
                    
                    # Create exploded dataframe
                    llm_summary_df = pd.DataFrame(llm_summary_rows)
                    exploded_df = llm_summary_df.assign(
                        llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                    ).explode('llm_paragraph')
                    exploded_df = exploded_df.reset_index(drop=True)
                    exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                    
                    # Get full text
                    full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                    
                    # Store in session state
                    file_size = os.path.getsize(tmp_file.name)
                    n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                    n_paragraphs = len(exploded_df)
                    
                    st.session_state['full_document_text_tab3'] = full_document_text
                    st.session_state['prodoc_document_stats_tab3'] = {
                        'file_size': file_size,
                        'n_words': n_words,
                        'n_paragraphs': n_paragraphs
                    }
                    st.session_state['exploded_df_tab3'] = exploded_df
                    st.session_state['extraction_df_tab3'] = df
                    st.session_state['tables_data_tab3'] = tables_data
                    st.session_state['extraction_stats_tab3'] = extraction_stats
                    st.session_state['sections_df_tab3'] = llm_summary_df
                    st.session_state['selected_sections_tab3'] = list(header_1_values)  # Select all by default
                    st.session_state['last_file_hash_tab3'] = file_hash
                    st.session_state['document_extracted_tab3'] = True
                    
                    try:
                        os.unlink(tmp_file.name)
                    except:
                        pass
                    
                    progress_bar.progress(1.0, text="Extracción completa.")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error procesando el documento: {e}")
                    import traceback
                    st.error(traceback.format_exc())
                    st.stop()
        
        # Show extraction results if document is extracted
        if st.session_state.get('document_extracted_tab3', False) and not file_changed:
            st.success("✅ Documento extraído con éxito")
            
            # Download button for extracted document structure
            extraction_df = st.session_state.get('extraction_df_tab3', pd.DataFrame())
            if not extraction_df.empty:
                excel_data = to_excel(extraction_df)
                # Get filename from extraction_df or use default
                filename_base = extraction_df['filename'].iloc[0] if 'filename' in extraction_df.columns and not extraction_df['filename'].empty else "documento"
                filename_base = filename_base.replace('.docx', '').replace('.doc', '')
                st.download_button(
                    label="📥 Descargar estructura extraída del documento (Excel)",
                    data=excel_data,
                    file_name=f"estructura_documento_tab3_{filename_base}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_extraction_tab3"
                )
                st.caption("El archivo incluye todas las columnas de encabezados (header_1 a header_6), contenido, tipo de fuente, y metadatos de extracción.")
            
            # Display header_1 sections and their content
            extraction_df = st.session_state.get('extraction_df_tab3', pd.DataFrame())
            if not extraction_df.empty:
                with st.expander("📋 Ver estructura extraída del documento (encabezados nivel 1 y contenido)", expanded=False):
                    st.markdown("**Estructura del documento extraído (solo encabezados nivel 1):**")
                    
                    # Get unique header_1 values
                    header_1_sections = extraction_df[extraction_df['header_1'].notna() & (extraction_df['header_1'] != '')]['header_1'].unique()
                    
                    for h1 in header_1_sections:
                        st.markdown(f"### {h1}")
                        
                        # Get all content for this header_1 section
                        section_df = extraction_df[extraction_df['header_1'] == h1]
                        section_content = section_df[section_df['source_type'] == 'paragraph']['content'].tolist()
                        
                        # Display content
                        if section_content:
                            full_text = '\n\n'.join([str(c) for c in section_content if pd.notna(c) and str(c).strip()])
                            if full_text.strip():
                                st.text(full_text)
                                st.caption(f"Total: {len(full_text):,} caracteres")
                        else:
                            st.info("Esta sección no tiene contenido de párrafos extraído.")
                        
                        st.markdown("---")
            
            sections_df = st.session_state.get('sections_df_tab3', pd.DataFrame())
            
            if not sections_df.empty:
                header_1_values = sections_df['header_1'].tolist()
                
                # Section selector
                st.markdown("### 🔍 Selección de Secciones para Evaluación")
                st.info("Selecciona las secciones que deseas incluir en la evaluación. Por defecto, todas las secciones están seleccionadas.")
                
                # Guidance about section extraction
                if len(header_1_values) == 0:
                    st.error("⚠️ **No se detectaron secciones en el documento.** Esto puede deberse a que el documento no usa estilos de encabezado de Word (Heading 1, Heading 2, etc.). Por favor, verifica que tu documento PRODOC tenga encabezados formateados correctamente.")
                elif len(header_1_values) < 3:
                    st.warning("⚠️ **Se detectaron pocas secciones** en el documento PRODOC. Si esperabas más secciones, verifica que el documento use estilos de encabezado de Word (Heading 1, Heading 2, etc.) para identificar las secciones principales como Marco Lógico, Presupuesto, Cronograma, etc.")
                
                # Initialize selected sections if not exists
                if 'selected_sections_tab3' not in st.session_state:
                    st.session_state['selected_sections_tab3'] = list(header_1_values)
                
                # Section selection interface
                selected_sections = st.session_state.get('selected_sections_tab3', list(header_1_values)).copy()
                col1, col2 = st.columns([3, 1])
                
                with col2:
                    if st.button("✅ Seleccionar Todas", key="select_all_sections_tab3"):
                        st.session_state['selected_sections_tab3'] = list(header_1_values)
                        st.rerun()
                    
                    if st.button("❌ Deseleccionar Todas", key="deselect_all_sections_tab3"):
                        st.session_state['selected_sections_tab3'] = []
                        st.rerun()
                
                with col1:
                    st.markdown("**Secciones disponibles:**")
                    for section in header_1_values:
                        section_info = sections_df[sections_df['header_1'] == section].iloc[0]
                        is_selected = section in selected_sections
                        
                        checkbox_label = f"**{section}** ({section_info['n_words']:,} palabras, {section_info['n_paragraphs']} párrafos)"
                        
                        checkbox_key = f"section_checkbox_{section}_tab3"
                        new_selection = st.checkbox(checkbox_label, value=is_selected, key=checkbox_key)
                        
                        if new_selection and section not in selected_sections:
                            selected_sections.append(section)
                        elif not new_selection and section in selected_sections:
                            selected_sections.remove(section)
                        
                        # Add expandable preview of extracted content
                        with st.expander(f"👁️ Ver contenido: {section}", expanded=False):
                            section_content = section_info['llm_paragraph']
                            if section_content and section_content.strip():
                                st.text_area(
                                    "Contenido extraído:",
                                    value=section_content,
                                    height=200,
                                    key=f"content_preview_{section}_tab3",
                                    label_visibility="collapsed"
                                )
                                st.caption(f"Total: {len(section_content):,} caracteres")
                            else:
                                st.info("Esta sección no tiene contenido extraído.")
                            
                            # Show table-extracted text
                            tables_data = st.session_state.get('tables_data_tab3', [])
                            section_tables = [t for t in tables_data if t.get('section') == section]
                            
                            if section_tables:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                for table_info in section_tables:
                                    table_num = table_info.get('table_number', 'N/A')
                                    table_data = table_info.get('data', [])
                                    if table_data:
                                        # Format table as text
                                        table_text = '\n'.join([' | '.join(str(cell) for cell in row) for row in table_data])
                                        st.text_area(
                                            f"Tabla {table_num}:",
                                            value=table_text,
                                            height=150,
                                            key=f"table_preview_{section}_table{table_num}_tab3",
                                            label_visibility="collapsed"
                                        )
                                        st.caption(f"Tabla {table_num}: {len(table_data)} filas, {len(table_data[0]) if table_data else 0} columnas")
                            else:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                st.info("No se encontraron tablas en esta sección.")
                
                # Update session state
                st.session_state['selected_sections_tab3'] = selected_sections
                
                # Show selection summary
                if selected_sections:
                    selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
                    total_selected_words = selected_df['n_words'].sum()
                    total_selected_paras = selected_df['n_paragraphs'].sum()
                    
                    # Estimate tokens
                    if encoding:
                        selected_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
                        estimated_tokens = len(encoding.encode(selected_text))
                    else:
                        estimated_tokens = total_selected_words * 1.2
                    
                    # Warn if approaching limit
                    if estimated_tokens > 100000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.warning(f"⚠️ **Advertencia de límite de contexto:** Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). El sistema procesa hasta 110,000 tokens (aproximadamente 150-200 páginas). Si el documento excede este límite, será truncado automáticamente.")
                    elif estimated_tokens > 80000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.info(f"ℹ️ Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). Estás dentro del límite de 110,000 tokens (aproximadamente 150-200 páginas).")
                    
                    st.success(f"✅ {len(selected_sections)} secciones seleccionadas | "
                              f"{total_selected_words:,} palabras | "
                              f"~{estimated_tokens:,} tokens estimados")

    # Rubric and Criteria Selection Section (moved after document extraction)
    st.markdown("---")
    st.markdown("### 📋 Selección de Criterios")
    
    # Group indicadores by dimension and criterio, maintaining order
    # Structure: {dimension: {criterio: [(unique_key, indicador_text, sort_key)]}}
    criteria_by_dimension = {}
    for unique_key, data in prodoc_rubric.items():
        dimension = data.get('dimension', 'No especificada')
        criterio = data.get('criterio', 'Sin criterio')
        indicador_text = data.get('indicador', unique_key)  # Use indicador text for display
        sort_key = data.get('sort_key', (999, 999, 999))
        
        if dimension not in criteria_by_dimension:
            criteria_by_dimension[dimension] = {}
        if criterio not in criteria_by_dimension[dimension]:
            criteria_by_dimension[dimension][criterio] = []
        criteria_by_dimension[dimension][criterio].append((unique_key, indicador_text, sort_key))
    
    # Sort indicadores within each criterio by their numeric prefix (sort_key is index 2)
    for dimension in criteria_by_dimension:
        for criterio in criteria_by_dimension[dimension]:
            criteria_by_dimension[dimension][criterio].sort(key=lambda x: x[2])
    
    # Define dimension order
    dimension_order = ['Diseño', 'Implementación', 'Pre-Cierre']
    
    # Helper to extract criterio number for sorting (e.g., "1. Participación..." -> 1)
    def get_criterio_order(criterio_name):
        import re
        match = re.match(r'(\d+)\.', criterio_name)
        return int(match.group(1)) if match else 999
    
    # Only show rubric selection if document is extracted
    if st.session_state.get('document_extracted_tab3', False):
        # Display the loaded rubric with selection grouped by dimension
        with st.expander("Ver y seleccionar criterios de evaluación", expanded=True):
            st.subheader("Criterios de Evaluación PRODOC")
        st.markdown(
            """
            <div class='reference-box'>
            Analiza si el Documento de Proyecto (PRODOC) incorpora, desde el inicio, los factores que favorecen la continuidad de resultados: participación de mandantes/socios y gestión de riesgos, bases de sostenibilidad institucional y política, consideraciones de género y, cuando aplique, transición justa. Usa como evidencia el Documento de Proyecto y anexos; los puntajes sirven para ajustar estrategias y definir tempranamente el plan de sostenibilidad.
            </div>
            """,
            unsafe_allow_html=True
        )
        
        # Select all checkbox
        select_all_tab3 = st.checkbox("Seleccionar todos los criterios", key='select_all_tab3')
        
        selected_criteria = []
        
        # Display criteria grouped by dimension
        dimension_descriptions = {
            "Diseño": (
                "Analiza si el Documento de Proyecto (PRODOC) incorpora, desde el inicio, los factores que "
                "favorecen la continuidad de resultados: participación de mandantes/socios y gestión de riesgos, "
                "bases de sostenibilidad institucional y política, consideraciones de género y, cuando aplique, "
                "transición justa. Usa como evidencia el Documento de Proyecto y anexos; los puntajes sirven para "
                "ajustar estrategias y definir tempranamente el plan de sostenibilidad."
            ),
            "Implementación": (
                "Contrasta avances reportados con los criterios de la Matriz de Criterios (capacidades desarrolladas, "
                "alianzas, recursos movilizados, integración en políticas/planes, gestión del conocimiento, etc.) "
                "para identificar riesgos, cuellos de botella y acciones de mitigación. Usa informes de progreso, actas, "
                "convenios y otros documentos de ejecución; los resultados orientan ajustes y fortalecen la trazabilidad "
                "de decisiones."
            ),
            "Pre-Cierre": (
                "Aplica una revisión ex post (idealmente en el último trimestre de un proyecto) para verificar qué "
                "elementos efectivamente aseguran la continuidad de resultados y documentar lecciones. Si se aplica "
                "antes del cierre, algunos puntajes serán referenciales (“lo esperado”); en todos los casos, el diagnóstico "
                "alimenta el plan/estrategia de sostenibilidad y su seguimiento. Usa un informe de evaluación para "
                "realizar este análisis."
            )
        }

        # Note: .reference-box styles are defined in the global ILO styles section

        # Sort dimensions according to defined order
        sorted_dimensions = sorted(
            criteria_by_dimension.keys(), 
            key=lambda d: dimension_order.index(d) if d in dimension_order else 999
        )
        
        for dimension in sorted_dimensions:
            st.markdown(f"#### 📊 {dimension}")
            if dimension in dimension_descriptions:
                st.markdown(
                    f"<div class='reference-box'>{dimension_descriptions[dimension]}</div>",
                    unsafe_allow_html=True
                )
            
            # Checkbox to select entire dimension
            dimension_key = f"dimension_tab3_{dimension}"
            select_dimension = st.checkbox(
                f"✅ Seleccionar toda la dimensión '{dimension}'",
                value=select_all_tab3 or dimension in st.session_state.get('selected_dimensions_tab3', []),
                key=dimension_key
            )
            
            # Get criterios for this dimension, sorted by their number
            criterios_in_dimension = criteria_by_dimension[dimension]
            sorted_criterios = sorted(criterios_in_dimension.keys(), key=get_criterio_order)
            
            # Show criterios and their indicadores within this dimension
            with st.container():
                for criterio in sorted_criterios:
                    # Display criterio header
                    st.markdown(f"**{criterio}**")
                    
                    # Get sorted indicadores for this criterio
                    # Each item is (unique_key, indicador_text, sort_key)
                    indicadores_data = criterios_in_dimension[criterio]
                    
                    for unique_key, indicador_text, sort_key in indicadores_data:
                        # If dimension is selected, auto-select all its indicadores
                        default_value = select_all_tab3 or select_dimension or unique_key in st.session_state['selected_criteria_tab3']
                        
                        is_selected = st.checkbox(
                            f"  ↳ {indicador_text}",
                            value=default_value,
                            key=f"criterion_tab3_{unique_key}",
                            disabled=select_dimension  # Disable individual selection if dimension is selected
                        )
                        
                        if is_selected or select_dimension:
                            selected_criteria.append(unique_key)
            
            st.markdown("---")  # Separator between dimensions
        
        # Update session state
        st.session_state['selected_criteria_tab3'] = selected_criteria
        selected_dimensions = [dim for dim in criteria_by_dimension.keys() 
                              if st.session_state.get(f"dimension_tab3_{dim}", False)]
        st.session_state['selected_dimensions_tab3'] = selected_dimensions
        
        # Count total indicadores
        total_indicadores = sum(
            len(indicadores) 
            for criterios in criteria_by_dimension.values() 
            for indicadores in criterios.values()
        )
        
        st.info(f"📌 Indicadores seleccionados: {len(selected_criteria)}/{total_indicadores} | Dimensiones seleccionadas: {len(selected_dimensions)}/{len(criteria_by_dimension)}")
    else:
        st.info("ℹ️ Por favor extrae el documento primero para poder seleccionar los criterios.")

    # Process and Evaluate button
    st.markdown("---")
    st.markdown("### ⚙️ Procesamiento y Evaluación")
    
    # Warning about AI results verification
    st.warning("""
    **⚠️ Importante - Verificación de Resultados:**
    
    Los resultados generados por esta herramienta utilizan inteligencia artificial y deben ser **verificados y corroborados** antes de su uso.
    
    - La IA puede cometer errores, interpretaciones incorrectas o pasar por alto información relevante
    - Los análisis y puntuaciones son **sugerencias** basadas en el contenido del documento, no son definitivos
    - Se recomienda revisar manualmente las evidencias citadas y validar las conclusiones
    - Los resultados deben ser contrastados con conocimiento experto y documentación adicional cuando sea necesario
    
    Esta herramienta es un **asistente de análisis** que facilita la revisión, pero la responsabilidad final de la evaluación recae en el usuario.
    """)
    
    if st.button('🚀 Procesar y Evaluar', key="prodoc_process_button_tab3", type="primary"):
        # Check prerequisites
        if not st.session_state.get('document_extracted_tab3', False):
            st.error("❌ Por favor extrae el documento primero usando el botón 'Extraer Documento'.")
            st.stop()
        
        if uploaded_file_prodoc is None:
            st.error("Por favor suba un archivo DOCX primero.")
            st.stop()
        
        # Get selected criteria from session state
        selected_criteria = st.session_state.get('selected_criteria_tab3', [])
        if not selected_criteria:
            st.error("Por favor seleccione al menos un criterio.")
            st.stop()
        
        # Get selected sections or use full document
        selected_sections = st.session_state.get('selected_sections_tab3', [])
        sections_df = st.session_state.get('sections_df_tab3', pd.DataFrame())
        
        if selected_sections and not sections_df.empty:
            # Filter to selected sections only
            selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
            document_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
            st.info(f"📌 Evaluando {len(selected_sections)} secciones seleccionadas")
        else:
            # Fallback to full document
            document_text = st.session_state.get('full_document_text_tab3', '')
            if not selected_sections:
                st.warning("⚠️ No hay secciones seleccionadas. Usando documento completo.")
        
        if not document_text:
            st.error("No se pudo recuperar el texto del documento.")
            st.stop()
        
        # Skip the old extraction logic - document is already extracted
        # Evaluate with selected criteria
        if False:  # Disable old extraction logic
            uploaded_file = uploaded_file_prodoc
            st.markdown("#### Procesando documento...")

            file_hash = hash(uploaded_file.getvalue())
            if st.session_state.get('prodoc_last_file_hash_tab3') != file_hash:
                with st.spinner("Procesando documento..."):
                    try:
                        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        tmp_file.write(uploaded_file.read())
                        tmp_file.close()
                        
                        progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                        doc_result = docx2python(tmp_file.name)
                        df = extract_docx_structure(tmp_file.name)
                        
                        progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")

                        # Extract sections directly (no LLM needed - text is already well-formatted)
                        header_1_values = df['header_1'].dropna().unique()
                        llm_summary_rows = []
                        progress_bar.progress(0.3, text="Extrayendo secciones del documento...")

                        for idx, header in enumerate(header_1_values):
                            section_df = df[df['header_1'] == header].copy()
                            # Extract text directly - already clean from extract_docx_structure
                            full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                            llm_summary_rows.append({'header_1': header, 'llm_paragraph': full_text if full_text else ""})

                        progress_bar.progress(0.5, text="Secciones extraídas.")
                        llm_summary_df = pd.DataFrame(llm_summary_rows)
                        llm_summary_df['n_words'] = llm_summary_df['llm_paragraph'].str.split().str.len()
                        exploded_df = llm_summary_df.assign(
                            llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                        ).explode('llm_paragraph')
                        exploded_df = exploded_df.reset_index(drop=True)
                        exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                        
                        full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                        file_size = os.path.getsize(tmp_file.name)
                        n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                        n_paragraphs = len(exploded_df)
                        
                        st.session_state['prodoc_full_document_text_tab3'] = full_document_text
                        st.session_state['prodoc_document_stats_tab3'] = {
                            'file_size': file_size,
                            'n_words': n_words,
                            'n_paragraphs': n_paragraphs
                        }
                        st.session_state['prodoc_exploded_df_tab3'] = exploded_df
                        st.session_state['prodoc_last_file_hash_tab3'] = file_hash
                        
                        try:
                            os.unlink(tmp_file.name)
                        except:
                            pass
                        
                        progress_bar.progress(0.8, text="Documento procesado. Listo para evaluación.")
                        st.info(f"**Resumen del documento:**\n\n" + 
                                f"- Tamaño del archivo: {file_size/1024:.2f} KB\n" + 
                                f"- Número de palabras: {n_words}\n" + 
                                f"- Número de párrafos: {n_paragraphs}")
                        st.markdown("#### Estructura extraída del documento:")
                        st.dataframe(exploded_df, use_container_width=True)
                        progress_bar.progress(1.0, text="Procesamiento completo.")
                        
                    except Exception as e:
                        st.error(f"Error procesando el documento: {e}")
                        import traceback
                        st.error(traceback.format_exc())
                        st.stop()
        
        # Define the evaluation function for tab3 with COUNTING support
        def evaluate_criterion_with_llm(document_text, criterion, descriptions, max_retries=3):
            """
            Analyze document against criterion with retry logic.
            TAB3-SPECIFIC: Enhanced to handle COUNT-BASED rubrics (stakeholders, participation forms, etc.)
            """
            import time

            # Detect if this is a stakeholder participation counting rubric
            is_stakeholder_counting = any(keyword in criterion.lower() for keyword in
                ['mandante', 'participan', 'stakeholder', 'actores', 'gobierno', 'empleador', 'trabajador'])

            for attempt in range(max_retries):
                try:
                    # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
                    combined_text = truncate_to_token_limit(document_text, max_tokens=110000, encoding_obj=encoding)

                    # Enhanced prompt for count-based rubrics
                    if is_stakeholder_counting:
                        system_content = """Eres un evaluador experto de documentos especialmente capacitado para CONTAR y aplicar LÓGICA DE UMBRALES.

**INSTRUCCIONES CRÍTICAS PARA EVALUACIÓN BASADA EN CONTEOS CON UMBRALES:**

Los niveles de la rúbrica especifican requisitos como "AL MENOS X mandantes participando en AL MENOS Y formas".
Debes aplicar lógica de umbrales estricta:

1. **IDENTIFICA y CUENTA** cada tipo de mandante/actor mencionado:
   - **Gobierno** (autoridades, funcionarios públicos, ministerios, etc.)
   - **Empleadores** (empresas, organizaciones patronales, cámaras de comercio, etc.)
   - **Trabajadores** (sindicatos, organizaciones de trabajadores, trabajadores individuales, etc.)

2. **CUENTA las formas de participación** para CADA mandante individualmente:
   - Diseño, reuniones, discusiones, comentarios, provisión de información, co-implementación, compromisos, etc.
   - Cada forma debe ser DISTINTA y VERIFICABLE en el documento

3. **FILTRO DE EVIDENCIA Y TIEMPO VERBAL (CRÍTICO - DEBE APLICARSE ESTRICTAMENTE):**
   
   **REGLA ABSOLUTA: SOLO PARTICIPACIÓN FACTUAL, PASADA O PRESENTE**
   
   - **SOLO cuenta participación que YA OCURRIÓ o ESTÁ OCURRIENDO** (hechos verificables, acciones concretas realizadas).
   - **NUNCA CUENTES** participaciones futuras, prometidas, planificadas, hipotéticas o esperadas.
   
   **VERBOS Y EXPRESIONES QUE DEBES RECHAZAR (NO CUENTAN):**
   - Futuro simple: "participará", "asistirá", "consultará", "se invitará", "se convocará", "se reunirá"
   - Futuro compuesto: "habrá participado", "habrá asistido"
   - Condicional: "participaría", "asistiría", "se consultaría"
   - Expresiones de planificación: "se espera que", "se prevé que", "está previsto", "se planifica", "se programará"
   - Expresiones de intención: "se pretende", "se busca", "se tiene como objetivo", "se propone"
   - Expresiones de compromiso futuro: "se compromete a", "acordó participar" (si no hay evidencia de participación real)
   - Marcadores temporales futuros: "en el futuro", "durante la implementación", "en las próximas fases", "posteriormente"
   - Participaciones hipotéticas: "podría participar", "sería consultado", "tendría la oportunidad"
   
   **VERBOS Y EXPRESIONES QUE SÍ CUENTAN (participación factual):**
   - Pasado simple: "participó", "asistió", "validó", "revisó", "comentó", "aprobó", "contribuyó", "colaboró"
   - Pasado compuesto: "ha participado", "ha asistido", "ha validado"
   - Presente: "participa", "asiste", "es miembro de", "forma parte de", "colabora en"
   - Gerundio de acciones completadas: "habiendo participado", "habiendo asistido"
   - Expresiones de hecho realizado: "fue consultado", "fue invitado y asistió", "se reunió con"
   
   **EJEMPLOS ESPECÍFICOS:**
   - ❌ "Se consultará a los mandantes durante la implementación" -> NO CUENTA (futuro)
   - ❌ "Se espera que los empleadores participen en el diseño" -> NO CUENTA (expectativa futura)
   - ❌ "Los trabajadores serán invitados a las reuniones" -> NO CUENTA (futuro)
   - ❌ "Está previsto que el gobierno valide la propuesta" -> NO CUENTA (planificación futura)
   - ❌ "Se tiene como objetivo involucrar a los actores" -> NO CUENTA (intención futura)
   - ✅ "El gobierno participó en la reunión de diseño del 15 de marzo" -> SÍ CUENTA (pasado factual)
   - ✅ "Los empleadores asistieron a las consultas y proporcionaron comentarios" -> SÍ CUENTA (pasado factual)
   - ✅ "Los trabajadores son miembros del comité de diseño" -> SÍ CUENTA (presente factual)
   - ✅ "Los mandantes fueron consultados y validaron el marco lógico" -> SÍ CUENTA (pasado factual)
   
   **VERIFICACIÓN OBLIGATORIA:**
   Antes de contar cualquier participación, pregunta: "¿Esta acción YA OCURRIÓ o está ocurriendo AHORA?" 
   Si la respuesta es NO o es incierta → NO CUENTES.
   Si la respuesta es SÍ y hay evidencia clara → SÍ CUENTA.

4. **APLICA LÓGICA DE UMBRALES** según los niveles de la rúbrica:
   - Si el nivel requiere "AL MENOS 2 mandantes en AL MENOS 2 formas":
     * CUENTA cuántos mandantes tienen 2 o más formas de participación
     * Si al menos 2 mandantes alcanzan ese umbral → cumple el nivel
     * Si solo 1 mandante alcanza ese umbral → NO cumple el nivel

   - Ejemplo:
     * Gobierno: 3 formas ✓ (cumple umbral de 2+)
     * Empleadores: 2 formas ✓ (cumple umbral de 2+)
     * Trabajadores: 1 forma ✗ (NO cumple umbral de 2+)
     * Resultado: 2 mandantes cumplen el umbral → SÍ califica para "al menos 2 mandantes en al menos 2 formas"

4. **ESTRUCTURA TU ANÁLISIS** así:
   ```
   CONTEO POR MANDANTE:
   - Gobierno: [N] formas identificadas → [LISTAR formas]
   - Empleadores: [N] formas identificadas → [LISTAR formas]
   - Trabajadores: [N] formas identificadas → [LISTAR formas]

   EVALUACIÓN DE UMBRALES (según nivel de la rúbrica):
   - Nivel X requiere: "AL MENOS [A] mandantes en AL MENOS [B] formas"
   - Mandantes que cumplen umbral de [B]+ formas: [LISTA DE MANDANTES]
   - Total de mandantes que cumplen umbral: [NÚMERO]
   - ¿Cumple requisito de [A]+ mandantes?: [SÍ/NO]

   JUSTIFICACIÓN DEL PUNTAJE:
   [Explicar EXPLÍCITAMENTE cómo los conteos y umbrales determinan el nivel]
   ```

5. **ASIGNA EL PUNTAJE** basándote ESTRICTAMENTE en:
   - Cuántos mandantes cumplen el umbral mínimo de formas
   - Si ese número cumple el requisito de "al menos X mandantes"

Siempre responde en español, incluso si el documento está en inglés."""

                        user_content = f"""Evalúa este documento contra el siguiente indicador (REQUIERE CONTEO CON LÓGICA DE UMBRALES):

**Indicador:** {criterion}

**Niveles de puntuación:** {json.dumps(descriptions, ensure_ascii=False, indent=2)}

**Documento a evaluar:**
{combined_text}

**INSTRUCCIONES CRÍTICAS:**
1. **PRIMERO: APLICA EL FILTRO DE TIEMPO VERBAL** - Revisa CADA mención de participación y verifica que sea factual (pasado o presente). RECHAZA cualquier participación futura, prometida o planificada.
2. CUENTA SOLO las formas de participación FACTUALES (pasadas o presentes) para CADA mandante
3. IDENTIFICA el umbral mínimo de formas requerido en cada nivel (ej: "al menos 2 formas")
4. CUENTA cuántos mandantes CUMPLEN ese umbral con participación REAL (no prometida)
5. VERIFICA si el número de mandantes que cumplen el umbral alcanza el requisito del nivel
6. Justifica el puntaje EXPLÍCITAMENTE basándote en la lógica de umbrales y menciona explícitamente que solo se contó participación factual

**EJEMPLO DE RAZONAMIENTO CORRECTO:**
"APLICACIÓN DEL FILTRO DE TIEMPO VERBAL: Se revisaron todas las menciones de participación. Se excluyeron las siguientes por ser futuras/planificadas: 'se consultará a los trabajadores durante la implementación', 'se espera que los empleadores participen'. Solo se contaron participaciones factuales (pasadas o presentes).

CONTEO POR MANDANTE (solo participación real):
- Gobierno: 3 formas identificadas → participó en reunión de diseño (15/03), validó marco lógico, proporcionó comentarios escritos (✓ todas factuales, cumple umbral de 2+)
- Empleadores: 2 formas identificadas → asistieron a consulta (20/03), revisaron propuesta técnica (✓ ambas factuales, cumple umbral de 2+)
- Trabajadores: 1 forma identificada → mencionados en documento pero sin evidencia de participación real (✗ NO cumple umbral de 2+)

EVALUACIÓN DE UMBRALES:
- Nivel 4 requiere: AL MENOS 2 mandantes en AL MENOS 2 formas cada uno
- Mandantes que cumplen umbral de 2+ formas: Gobierno (3 formas), Empleadores (2 formas)
- Total de mandantes que cumplen umbral: 2
- ¿Cumple requisito de 2+ mandantes?: SÍ

JUSTIFICACIÓN DEL PUNTAJE: Dado que 2 mandantes (Gobierno y Empleadores) cumplen el umbral de 2+ formas cada uno con participación REAL verificable, SÍ se alcanza el Nivel 4. Se excluyeron participaciones futuras/planificadas del conteo."

**RECORDATORIO FINAL CRÍTICO:**
- Si encuentras participaciones futuras o planificadas en el documento, MENCIONA explícitamente en tu análisis que fueron EXCLUIDAS del conteo
- Solo incluye en "evidence" citas que demuestren participación REAL (pasada o presente)
- Si el documento solo menciona participación futura/planificada y no hay evidencia de participación real, el puntaje debe reflejar esto (probablemente nivel 1 o 0)

Proporciona tu respuesta como JSON:
{{"analysis": "COMIENZA indicando si aplicaste el filtro de tiempo verbal y qué participaciones fueron excluidas (si las hubo). Luego presenta el conteo por mandante SOLO con participación factual. EVALÚA EXPLÍCITAMENTE los umbrales según cada nivel. Finalmente JUSTIFICA el puntaje con la lógica de umbrales, asegurándote de mencionar que solo se contó participación real. 2-3 párrafos en ESPAÑOL", "score": 1-5, "evidence": ["cita 1 que evidencia mandante y forma de participación específica REAL (pasada o presente)", "cita 2", "etc - 5-8 citas clave como array, SOLO participaciones factuales"]}}"""

                    else:
                        # Original prompt for non-counting rubrics
                        system_content = "Eres un evaluador experto de documentos. Siempre debes responder en español, incluso si el documento está en inglés."

                        user_content = f"""Evaluate this document against: {criterion}

    Scoring levels: {json.dumps(descriptions)}

    Relevant document sections:
    {combined_text}

    IMPORTANTE: Proporciona tu respuesta SIEMPRE en español, incluso si el documento está en inglés.

    Provide JSON with:
    {{"analysis": "detailed 2-3 paragraphs IN SPANISH", "score": 1-5, "evidence": ["quote 1", "quote 2", "quote 3", "etc - 5-8 key quotes from the text as an array"]}}"""

                    response = client.chat.completions.create(
                        model="gpt-5-mini",
                        messages=[
                            {"role": "system", "content": system_content},
                            {"role": "user", "content": user_content}
                        ],
                        max_completion_tokens=6500,
                        reasoning_effort="minimal",
                        timeout=120  # 2 minute timeout per request
                    )

                    content = response.choices[0].message.content.strip()
                    # Remove markdown code fences if present
                    if content.startswith('```'):
                        # Remove opening fence (```json or ```)
                        content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                        # Remove closing fence
                        if content.endswith('```'):
                            content = content.rsplit('```', 1)[0]
                        content = content.strip()

                    result = json.loads(content)
                    # Normalize evidence field: convert array to string if needed
                    if isinstance(result.get('evidence'), list):
                        result['evidence'] = '\n'.join(result['evidence'])
                    return result

                except json.JSONDecodeError as e:
                    # If JSON parsing fails, return a default structure
                    return {
                        "analysis": f"Failed to parse JSON: {str(e)}. Raw response: {response.choices[0].message.content[:200]}",
                        "score": 3,
                        "evidence": "Unable to parse structured response",
                        "error": f"JSON parsing error: {str(e)}"
                    }
                except Exception as e:
                    # Check if it's a rate limit error
                    error_msg = str(e)
                    if "rate_limit" in error_msg.lower() or "429" in error_msg:
                        if attempt < max_retries - 1:
                            wait_time = (2 ** attempt) * 2  # Exponential backoff: 2s, 4s, 8s
                            time.sleep(wait_time)
                            continue

                    # If last attempt or non-rate-limit error, return error
                    return {
                        "analysis": f"Error during evaluation: {error_msg}",
                        "score": 0,
                        "evidence": "",
                        "error": f"API error (attempt {attempt + 1}/{max_retries}): {error_msg}"
                    }

            # If we exhausted all retries
            return {
                "analysis": "Failed after multiple retry attempts",
                "score": 0,
                "evidence": "",
                "error": f"Failed after {max_retries} attempts"
            }
        
        # Evaluate with selected criteria
        document_text = st.session_state.get('full_document_text_tab3', '')
        if not document_text:
            st.error("No se pudo recuperar el texto del documento. Por favor, vuelva a cargar el archivo.")
            st.stop()
        
        # Build filtered rubric based on selection
        filtered_rubric = {k: v for k, v in prodoc_rubric.items() if k in selected_criteria}
        
        if not filtered_rubric:
            st.error("No hay criterios seleccionados para evaluar.")
            st.stop()
        
        rubrics = [("Evaluación PRODOC", filtered_rubric)]
        
        rubric_results = []
        from concurrent.futures import ThreadPoolExecutor, as_completed
        MAX_WORKERS = 8
        
        # Define dimension order for sorting
        dimension_order_map = {'Diseño': 0, 'Implementación': 1, 'Evaluación': 2}
        
        def eval_one_criterion_tab3(args):
            unique_key, descriptions, dimension, criterio_text, indicador_text, sort_key, rubric_name = args
            try:
                result = evaluate_criterion_with_llm(document_text, indicador_text, descriptions)
                # Ensure result is a dictionary
                if not isinstance(result, dict):
                    result = {'score': 0, 'analysis': str(result), 'evidence': '', 'error': 'Invalid result format'}
                return {
                    'Dimensión': dimension,
                    'Criterio': criterio_text,
                    'Indicador': indicador_text,
                    'Score': result.get('score', 0),
                    'Análisis': str(result.get('analysis', '')),
                    'Evidencia': str(result.get('evidence', '')),
                    'Error': str(result.get('error', '')) if 'error' in result else '',
                    'Rúbrica': rubric_name,
                    '_dim_order': dimension_order_map.get(dimension, 99),
                    '_sort_key': sort_key
                }
            except Exception as e:
                return {
                    'Dimensión': dimension,
                    'Criterio': criterio_text,
                    'Indicador': indicador_text,
                    'Score': 0,
                    'Análisis': '',
                    'Evidencia': '',
                    'Error': str(e),
                    'Rúbrica': rubric_name,
                    '_dim_order': dimension_order_map.get(dimension, 99),
                    '_sort_key': sort_key
                }
        
        for rubric_name, rubric_dict in rubrics:
            rubric_analysis_data = []
            n_criteria = len(rubric_dict)
            progress = st.progress(0, text=f"Iniciando evaluación por rúbrica: {rubric_name}...")
            
            with st.spinner(f'Evaluando documento por rúbrica: {rubric_name}...'):
                with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    futures = {
                        executor.submit(eval_one_criterion_tab3, (
                            unique_key,
                            rubric_data['valores'] if isinstance(rubric_data, dict) else rubric_data,
                            rubric_data.get('dimension', 'No especificada') if isinstance(rubric_data, dict) else 'No especificada',
                            rubric_data.get('criterio', '') if isinstance(rubric_data, dict) else '',
                            rubric_data.get('indicador', unique_key) if isinstance(rubric_data, dict) else unique_key,
                            rubric_data.get('sort_key', (999, 999, 999)) if isinstance(rubric_data, dict) else (999, 999, 999),
                            rubric_name
                        )): (unique_key, idx)
                        for idx, (unique_key, rubric_data) in enumerate(rubric_dict.items())
                    }
                    
                    completed = 0
                    for future in as_completed(futures):
                        result = future.result()
                        rubric_analysis_data.append(result)
                        completed += 1
                        unique_key, idx = futures[future]
                        progress.progress(completed / n_criteria, text=f"Evaluando indicador...")
            
            # Create DataFrame and sort by dimension order and sort_key
            df_result = pd.DataFrame(rubric_analysis_data)
            if not df_result.empty:
                df_result = df_result.sort_values(by=['_dim_order', '_sort_key'])
                df_result = df_result.drop(columns=['_dim_order', '_sort_key'], errors='ignore')
            rubric_results.append((rubric_name, df_result))
        
        # Store results in session state for persistence
        st.session_state['tab3_results'] = rubric_results
        
        # Show and allow download of results
        if rubric_results:
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''

                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Dimensión', 'Criterio', 'Indicador', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Create visualizations for the results
            st.markdown("### Visualización de Resultados")
            
            all_scores = []
            for rubric_name, df in rubric_results:
                for _, row in df.iterrows():
                    all_scores.append({
                        'Indicador': row.get('Indicador', row.get('Criterio', '')),
                        'Criterio': row.get('Criterio', ''),
                        'Dimensión': row['Dimensión'],
                        'Puntuación': row['Score']
                    })
            
            scores_df = pd.DataFrame(all_scores)
            overall_avg = scores_df['Puntuación'].mean()
            
            # Visualization 1: Average Score by Dimension
            st.markdown("#### 📊 Puntuación Promedio por Dimensión")
            dimension_avg = scores_df.groupby('Dimensión')['Puntuación'].mean().reset_index()
            dimension_avg = dimension_avg.sort_values(by='Puntuación', ascending=False)
            
            fig_dim = go.Figure()
            fig_dim.add_trace(go.Bar(
                x=dimension_avg['Dimensión'],
                y=dimension_avg['Puntuación'],
                text=dimension_avg['Puntuación'].round(2),
                textposition='auto',
                marker_color='#002F6C',
                name='Promedio por Dimensión'
            ))
            
            fig_dim.add_trace(go.Scatter(
                x=dimension_avg['Dimensión'],
                y=[overall_avg] * len(dimension_avg),
                mode='lines',
                line=dict(color='#C8102E', width=2, dash='dash'),
                name=f'Promedio General: {overall_avg:.2f}'
            ))
            
            fig_dim.update_layout(
                title='Puntuación Promedio por Dimensión',
                xaxis_title='Dimensión',
                yaxis_title='Puntuación Promedio',
                yaxis=dict(range=[0, 5.5]),
                height=500,
                legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                margin=dict(l=20, r=20, t=80, b=50),
                hoverlabel=dict(bgcolor="white", font_size=12, font_family="Arial")
            )
            
            fig_dim.update_xaxes(tickangle=-45)
            fig_dim.update_yaxes(showgrid=True, gridwidth=1, gridcolor='LightGrey')
            
            st.plotly_chart(fig_dim, use_container_width=True)
            
            # Visualization 2: Individual Criteria Scores
            st.markdown("#### 📈 Puntuación por Criterio Individual")
            scores_df_sorted = scores_df.sort_values(by='Puntuación', ascending=False)
            
            fig = go.Figure()
            scores_df_sorted['Criterio_ID'] = [f"Criterio {i+1}" for i in range(len(scores_df_sorted))]
            scores_df_sorted['Hover_Text'] = scores_df_sorted.apply(
                lambda row: f"<b>{row['Criterio_ID']}</b><br>Dimensión: {row['Dimensión']}<br>{row['Criterio']}<br>Puntuación: {row['Puntuación']:.2f}", 
                axis=1
            )
            
            fig.add_trace(go.Bar(
                y=scores_df_sorted['Criterio_ID'],
                x=scores_df_sorted['Puntuación'],
                text=scores_df_sorted['Puntuación'].round(2),
                textposition='auto',
                marker_color='#0072CE',
                orientation='h',
                name='Puntuación',
                hovertext=scores_df_sorted['Hover_Text'],
                hoverinfo='text'
            ))
            
            fig.add_trace(go.Scatter(
                y=scores_df_sorted['Criterio_ID'],
                x=[overall_avg] * len(scores_df_sorted),
                mode='lines',
                line=dict(color='#C8102E', width=2, dash='dash'),
                name=f'Promedio General: {overall_avg:.2f}'
            ))
            
            fig.update_layout(
                title='Puntuación por Criterio (Ordenado de Mayor a Menor)',
                xaxis_title='Puntuación',
                yaxis_title='',
                xaxis=dict(range=[0, 5.5]),
                height=max(400, len(scores_df_sorted) * 35),
                width=800,
                legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                margin=dict(l=20, r=20, t=80, b=50),
                hoverlabel=dict(bgcolor="white", font_size=12, font_family="Arial")
            )
            
            fig.update_yaxes(automargin=True)
            fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor='LightGrey')
            
            st.plotly_chart(fig, use_container_width=True)
            
            # Download ZIP with XLSX files
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    # Export as XLSX instead of CSV
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_prodoc_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_evaluacion_prodoc.zip",
                mime="application/zip",
                key="prodoc_download_button_tab3"
            )
        else:
            st.warning("No se generaron resultados para ninguna rúbrica.")
    else:
        # Check if there are persisted results in session state
        if st.session_state.get('tab3_results') is not None:
            rubric_results = st.session_state['tab3_results']
            
            st.markdown("### 📊 Resultados guardados")
            
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''

                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Dimensión', 'Criterio', 'Indicador', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Create visualizations for the results
            st.markdown("### Visualización de Resultados")
            
            all_scores = []
            for rubric_name, df in rubric_results:
                for _, row in df.iterrows():
                    all_scores.append({
                        'Indicador': row.get('Indicador', row.get('Criterio', '')),
                        'Criterio': row.get('Criterio', ''),
                        'Dimensión': row['Dimensión'],
                        'Puntuación': row['Score']
                    })
            
            if all_scores:
                scores_df = pd.DataFrame(all_scores)
                overall_avg = scores_df['Puntuación'].mean()
                
                # Visualization: Average Score by Dimension
                st.markdown("#### 📊 Puntuación Promedio por Dimensión")
                dimension_avg = scores_df.groupby('Dimensión')['Puntuación'].mean().reset_index()
                dimension_avg = dimension_avg.sort_values(by='Puntuación', ascending=False)
                
                fig_dim = go.Figure()
                fig_dim.add_trace(go.Bar(
                    x=dimension_avg['Dimensión'],
                    y=dimension_avg['Puntuación'],
                    text=dimension_avg['Puntuación'].round(2),
                    textposition='auto',
                    marker_color='#002F6C',
                    name='Promedio por Dimensión'
                ))
                
                fig_dim.add_trace(go.Scatter(
                    x=dimension_avg['Dimensión'],
                    y=[overall_avg] * len(dimension_avg),
                    mode='lines',
                    line=dict(color='#C8102E', width=2, dash='dash'),
                    name=f'Promedio General: {overall_avg:.2f}'
                ))
                
                fig_dim.update_layout(
                    title='Puntuación Promedio por Dimensión',
                    xaxis_title='Dimensión',
                    yaxis_title='Puntuación Promedio',
                    yaxis=dict(range=[0, 5.5]),
                    height=500,
                    legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                    margin=dict(l=20, r=20, t=80, b=50),
                    hoverlabel=dict(bgcolor="white", font_size=12, font_family="Arial")
                )
                
                fig_dim.update_xaxes(tickangle=-45)
                fig_dim.update_yaxes(showgrid=True, gridwidth=1, gridcolor='LightGrey')
                
                st.plotly_chart(fig_dim, use_container_width=True)
            
            # Download ZIP
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_prodoc_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_evaluacion_prodoc.zip",
                mime="application/zip",
                key="prodoc_download_button_tab3_persisted"
            )
            
            # Clear results button
            if st.button("🗑️ Limpiar resultados", key="clear_tab3_results"):
                st.session_state['tab3_results'] = None
                st.rerun()
        else:
            st.info("Seleccione criterios, suba un documento y presione 'Procesar y Evaluar'.")


# ================== TAB 7: APPRAISAL CHECKLIST (IMPROVED) =====================
# Configuration
from concurrent.futures import ThreadPoolExecutor, as_completed
import io
MAX_WORKERS = 48  # Reduced from 48 to avoid rate limits
OPENAI_MODEL = "gpt-5-mini"  # GPT-5 mini model with reasoning

@st.cache_data
def load_appraisal_questions():
    """Load and cache appraisal questions from Excel file"""
    try:
        df = pd.read_excel('./Appraisal Checklist_2025 es-419.xlsx', sheet_name='rubric')
        if 'Pregunta_Realizada' not in df.columns:
            return None, "La columna 'Pregunta_Realizada' no se encontró en el archivo de Excel."
        
        # Replace first 3 characters of Pregunta_Realizada with Tema numbering
        if 'Tema' in df.columns:
            df['Pregunta_Realizada'] = df.apply(
                lambda row: str(row['Tema']) + ' ' + str(row['Pregunta_Realizada'])[3:].strip() 
                if pd.notna(row['Tema']) and pd.notna(row['Pregunta_Realizada']) and len(str(row['Pregunta_Realizada'])) > 3
                else row['Pregunta_Realizada'], 
                axis=1
            )
        
        return df, None
    except FileNotFoundError:
        return None, "No se encontró el archivo Appraisal Checklist_2025 es-419.xlsx. Asegúrate de que exista en el directorio de la aplicación."
    except Exception as e:
        return None, f"Error al cargar el archivo de preguntas: {str(e)}"

def extract_document_content(uploaded_file):
    """Extract and process content from uploaded DOCX file"""
    try:
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_file_path = tmp_file.name
        
        # Extract content using docx2python
        doc_result = docx2python(tmp_file_path)
        
        # Get file stats
        file_size = os.path.getsize(tmp_file_path)
        
        # Extract structured content (you'll need to implement extract_docx_structure)
        # For now, using simple text extraction
        full_text = doc_result.text
        word_count = len(full_text.split())
        
        # Clean up temp file
        os.unlink(tmp_file_path)
        
        return {
            'text': full_text,
            'file_size': file_size,
            'word_count': word_count,
            'success': True
        }
    
    except Exception as e:
        return {
            'text': '',
            'file_size': 0,
            'word_count': 0,
            'success': False,
            'error': str(e)
        }


def parse_critic_verdict(critic_text):
    """Extract the VEREDICTO tag from the first line of a critic response.

    Returns (verdict, body) where:
      - verdict ∈ {"Yes", "No", "Partial", "Not Found", "Keep", None}
      - body is the critic text with the verdict line stripped (original text if no verdict parsed)
    """
    if not isinstance(critic_text, str) or not critic_text.strip():
        return None, critic_text or ""
    m = re.match(
        r'\s*VEREDICTO\s*:\s*(Not\s*Found|Partial|Yes|No|Keep)\s*[\.\n:]?',
        critic_text,
        flags=re.IGNORECASE,
    )
    if not m:
        return None, critic_text.strip()
    raw = m.group(1).strip()
    canonical = {
        'yes': 'Yes',
        'no': 'No',
        'partial': 'Partial',
        'not found': 'Not Found',
        'notfound': 'Not Found',
        'keep': 'Keep',
    }
    verdict = canonical.get(re.sub(r'\s+', ' ', raw.lower()), None)
    body = critic_text[m.end():].lstrip()
    return verdict, body


def parse_two_part_question(question):
    """
    Detect and parse two-part rubric questions where:
    - Part 1 (broader): Sets the general context
    - Part 2 (specific): Asks the critical detail that needs primary focus

    Returns dict with 'is_two_part', 'part1', 'part2', 'full_question'.
    """
    import re

    if not isinstance(question, str) or question.count('?') < 2:
        return {'is_two_part': False, 'part1': None, 'part2': None, 'full_question': question}

    # Pattern 1: Explicit keyword separator ("específicamente", "además", ...).
    keywords = ['específicamente', 'en particular', 'en concreto', 'puntualmente', 'además']
    for keyword in keywords:
        pattern = rf'(.+?\?)\s*{keyword}[,:]?\s*¿\s*(.+?\?)'
        match = re.search(pattern, question, re.IGNORECASE | re.UNICODE)
        if match:
            return {
                'is_two_part': True,
                'part1': match.group(1).strip(),
                'part2': '¿' + match.group(2).strip(),
                'full_question': question,
            }

    # Pattern 2a: Two clauses joined with explicit ¿ separator (Spanish convention).
    # Unicode-safe: matches lowercase, accented capitals (Á/Í/Ó/Ú/Ñ), digits, etc.
    match = re.search(r'(.+?\?)\s*¿\s*(.+\?)', question, re.UNICODE)
    if match:
        return {
            'is_two_part': True,
            'part1': match.group(1).strip(),
            'part2': '¿' + match.group(2).strip(),
            'full_question': question,
        }

    # Pattern 2b: Run-together — first clause ends in '?', second starts immediately
    # with a non-whitespace character (rubric formatting quirk, e.g. "...gráfico?Se identifican...").
    match = re.search(r'(.+?\?)(\S.+\?)', question, re.UNICODE)
    if match:
        part2 = match.group(2).strip()
        if not part2.startswith('¿'):
            part2 = '¿' + part2
        return {
            'is_two_part': True,
            'part1': match.group(1).strip(),
            'part2': part2,
            'full_question': question,
        }

    return {'is_two_part': False, 'part1': None, 'part2': None, 'full_question': question}

def analyze_question_with_llm_tab1(question, document_text):
    """
    Analyze a single question against the document using LLM with full context (up to 110K tokens).
    TAB1-SPECIFIC VERSION: Explicitly states when 2+ parts are detected in the question.
    """
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        combined_text = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        if not combined_text.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Not Found',
                'Razonamiento': 'No se encontró contenido en el documento para analizar.',
                'Evidencia': '',
                'Status': 'Success'
            }

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize system prompt based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert document analyst specializing in ILO project quality assessment.

**CRITICAL INSTRUCTION FOR TWO-PART QUESTIONS:**

This question has TWO PARTS with different priorities:
- **Part 2 (Specific Focus - PRIMARY)**: The critical question that requires detailed analysis. Drives the final answer.
- **Part 1 (Broader Context)**: Provides the general framework.

**YOUR ANALYSIS APPROACH:**
1. **Answer Part 2 IN DEPTH FIRST.** This drives the final Respuesta.
2. **Then relate Part 2's findings back to Part 1** (broader context).
3. The final Respuesta is determined PRIMARILY by Part 2; Part 1 provides framing only.

**Response fields (returned via structured output):**
- "Respuesta": Yes / No / Partial / Not Found — based PRIMARILY on Part 2.
- "Razonamiento": MUST begin with "Se identificaron 2 partes en esta pregunta." Then answer Part 2 concisely, then briefly connect to Part 1. Max 120 words. Be terse — no filler, no restating the question.
- "Evidencia": Quotes supporting Part 2 FIRST, then supporting evidence for Part 1. Max 180 words. Trim quotes to the essential span; avoid long blocks.
- "parte_enfocada": Set to "Parte 2" when your analysis was driven by Part 2 (the expected default). Use "Ambas" only if both parts truly required equal weight. Use "Parte 1" only if Part 2 was unanswerable from the document.

**SCOPE LOCK (CRITICAL):**
La pregunta nombra un sujeto o alcance específico (p.ej. "personas con discapacidad", "mujeres rurales", "trabajo infantil en minería", una región, un sector o un período concreto). Tu análisis debe limitarse ESTRICTAMENTE a ese sujeto exacto.
1. Identifica el sujeto exacto de la pregunta antes de responder.
2. Evidencia sobre sujetos relacionados pero distintos (otras poblaciones vulnerables, otros grupos, otros sectores, otras regiones, otros períodos) NO cuenta como respuesta. Ejemplo: si la pregunta es sobre discapacidad, evidencia sobre mujeres, pueblos indígenas o juventud NO la satisface.
3. Si el sujeto nombrado está ausente del documento pero sujetos relacionados están ampliamente cubiertos, Respuesta debe ser "Not Found" o "No" — NO "Partial" ni "Yes". Sub-inclusión es preferible a sobre-inclusión.
4. En tu Razonamiento, declara explícitamente cuál sujeto exacto se analizó y confirma que la evidencia trata sobre ese sujeto (no sobre uno relacionado).

**EVIDENCE-ROLE FILTER & DECISION GATE FOR PARTE 2 (APPLY FIRST — DRIVES THE Respuesta):**

Before choosing a Respuesta, classify every candidate evidence passage about the Parte-2 subject (e.g., personas con discapacidad, género, pueblos indígenas, mujeres rurales) as either FRAMING or DEDICATED.

FRAMING mention (does NOT count toward Parte 2, even if the subject is named):
  - Overall objective / impact statement naming multiple groups
  - Stakeholder, consultation, or research participant lists
  - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
  - Boilerplate inclusion language
  - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

DEDICATED element (counts toward Parte 2):
  A. A sub-objective, outcome, or output whose title/purpose names the subject
  B. An indicator disaggregated by or specifically targeting the subject
  C. An activity whose primary purpose addresses the subject
  D. A budget line or resource allocation for the subject
  E. A quantifiable target for the subject

Count of DEDICATED elements (A–E) — NOT raw mention count — drives Respuesta:
  - 0 dedicated elements (only FRAMING mentions)      → "No" or "Not Found"
  - 1–2 dedicated elements                            → "Partial"
  - 3–4 dedicated elements                            → "Partial" (or "Yes" only if substantive and at-par with any claim/label in Parte 1)
  - 5 dedicated elements with substantive evidence    → "Yes"

If every citable quote is a FRAMING mention, Respuesta MUST be "No" or "Not Found" — regardless of how many times the subject is named. Passing mentions in stakeholder lists, consultation rosters, or "among others" phrases DO NOT qualify for "Partial".

In "Evidencia", prefix each cited quote with [DEDICATED] or [FRAMING]. You may NOT justify "Partial" or "Yes" using only [FRAMING] quotes.

Siempre responder en español. Enfoque: 95% en Parte 2, 5% en su relación con Parte 1. La Parte 1 solo existe para enmarcar brevemente; NO analices la Parte 1 de forma independiente."""

            # Part 2 listed FIRST to exploit primacy bias toward the priority clause.
            # Full original question deliberately omitted — including it re-anchors the model on Part 1.
            user_content = f"""PREGUNTA CON DOS PARTES — RESPONDE PRIMERO LA PARTE 2.

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**PARTE 1 (Contexto General — solo para enmarcar):**
{parsed['part1']}

**Texto del Documento:**
{combined_text}

RECUERDA:
1. COMENZAR tu Razonamiento con "Se identificaron 2 partes en esta pregunta."
2. Enfoca tu análisis PRINCIPALMENTE en la Parte 2 (pregunta específica).
3. Luego explica cómo se relaciona con la Parte 1 (contexto general).
4. Marca "parte_enfocada" como "Parte 2" salvo que el documento no permita responderla."""

            response_format = {"type": "json_schema", "json_schema": RUBRIC_SCHEMA_TWO_PART}

        else:
            # Original single-question prompt
            system_content = """You are an expert document analyst. Analyze the document against the given question and provide a structured JSON response with exactly this format and always respond in Spanish:
            {
                "Respuesta": "Yes/No/Partial/Not Found",
                "Razonamiento": "Concise explanation (max 120 words, terse)",
                "Evidencia": "Trimmed text excerpts supporting the answer (max 180 words)"
            }

**SCOPE LOCK (CRITICAL):**
The question names a specific subject/scope (e.g., a specific population like people with disabilities; a specific sector, region, or period). Your analysis must be STRICTLY limited to that exact named subject.
1. Identify the exact subject of the question before answering.
2. Evidence about related-but-different subjects (other vulnerable populations, other groups, other sectors/regions/periods) does NOT count. Example: if the question is about disability, evidence about women, indigenous peoples, or youth does NOT satisfy it.
3. If the named subject is absent but related subjects are extensively covered, Respuesta must be "Not Found" or "No" — NOT "Partial" and NOT "Yes". Under-inclusion is preferred over over-inclusion.
4. In your Razonamiento, explicitly state which exact subject was analyzed and confirm the evidence concerns that subject (not a related one).

**EVIDENCE-ROLE FILTER & DECISION GATE (APPLY FIRST — DRIVES Respuesta):**

Classify every candidate evidence passage about the named specific subject as either FRAMING or DEDICATED.

FRAMING mention (does NOT count, even if the subject is named):
  - Overall objective / impact statement naming multiple groups
  - Stakeholder, consultation, or research participant lists
  - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
  - Boilerplate inclusion language
  - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

DEDICATED element (counts):
  A. A sub-objective, outcome, or output whose title/purpose names the subject
  B. An indicator disaggregated by or targeting the subject
  C. An activity whose primary purpose addresses the subject
  D. A budget line or resource allocation for the subject
  E. A quantifiable target for the subject

Count of DEDICATED elements (A–E) — NOT raw mention count — drives Respuesta:
  - 0 (only FRAMING mentions)                        → "No" or "Not Found"
  - 1–2                                              → "Partial"
  - 3–4                                              → "Partial" (or "Yes" only if substantive)
  - 5 with substantive evidence                      → "Yes"

If every citable quote is a FRAMING mention, Respuesta MUST be "No" or "Not Found". In "Evidencia", prefix each quote with [DEDICATED] or [FRAMING]. You may NOT justify "Partial" or "Yes" using only [FRAMING] quotes."""

            user_content = f"Question: {question}\n\nDocument Text: {combined_text}"
            response_format = {"type": "json_schema", "json_schema": RUBRIC_SCHEMA_SINGLE}

        # Bump reasoning effort for two-part: the 70/30 weighting needs more deliberation.
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {"role": "system", "content": system_content},
                {"role": "user", "content": user_content},
            ],
            max_completion_tokens=3000,
            reasoning_effort="medium" if parsed['is_two_part'] else "minimal",
            response_format=response_format,
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'No se recibió respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

        # Structured outputs guarantees valid JSON conforming to the schema.
        try:
            result = json.loads(content.strip())
        except json.JSONDecodeError:
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'Error al procesar la respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

        # Flag rows where the model failed to commit to Part 2 focus despite a two-part question.
        # 'Partial' is the analyst's signal to spot-check the row.
        status = 'Success'
        if parsed['is_two_part'] and result.get('parte_enfocada') != 'Parte 2':
            status = 'Partial'

        return {
            'Pregunta': question,
            'Respuesta': result.get('Respuesta', 'Not Found'),
            'Razonamiento': result.get('Razonamiento', ''),
            'Evidencia': result.get('Evidencia', ''),
            'Status': status,
        }

    except Exception as e:
        return {
            'Pregunta': question,
            'Respuesta': 'Error',
            'Razonamiento': f'Analysis failed: {str(e)}',
            'Evidencia': '',
            'Status': 'Error'
        }

def analyze_question_with_llm(question, document_text):
    """Analyze a single question against the document using LLM with full context (up to 110K tokens)."""
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        combined_text = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        if not combined_text.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Not Found',
                'Razonamiento': 'No se encontró contenido en el documento para analizar.',
                'Evidencia': '',
                'Status': 'Success'
            }

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize system prompt based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert document analyst specializing in ILO project quality assessment.

**CRITICAL INSTRUCTION FOR TWO-PART QUESTIONS:**

This question has TWO PARTS with different priorities:
- **Part 2 (Specific Focus - PRIMARY)**: The critical question that requires detailed analysis. Drives the final answer.
- **Part 1 (Broader Context)**: Provides the general framework.

**YOUR ANALYSIS APPROACH:**
1. **Answer Part 2 IN DEPTH FIRST.** This drives the final Respuesta.
2. **Then relate Part 2's findings back to Part 1** (broader context).
3. The final Respuesta is determined PRIMARILY by Part 2; Part 1 provides framing only.

**Response fields (returned via structured output):**
- "Respuesta": Yes / No / Partial / Not Found — based PRIMARILY on Part 2.
- "Razonamiento": Begin by answering Part 2 concisely, then briefly connect to Part 1. Max 120 words. Be terse — no filler, no restating the question.
- "Evidencia": Quotes supporting Part 2 FIRST, then supporting evidence for Part 1. Max 180 words. Trim quotes to the essential span; avoid long blocks.
- "parte_enfocada": Set to "Parte 2" when your analysis was driven by Part 2 (the expected default). Use "Ambas" only if both parts truly required equal weight. Use "Parte 1" only if Part 2 was unanswerable from the document.

**SCOPE LOCK (CRITICAL):**
The question names a specific subject or scope (e.g., "people with disabilities", "rural women", "child labor in mining", a specific region/sector/period). Your analysis must be STRICTLY limited to that exact named subject.
1. Identify the exact subject of the question before answering.
2. Evidence about related-but-different subjects (other vulnerable populations, other groups, other sectors, other regions, other periods) does NOT count as answering the question. Example: if the question is about disability, evidence about women, indigenous peoples, or youth does NOT satisfy it.
3. If the named subject is absent from the document but related subjects are extensively covered, Respuesta must be "Not Found" or "No" — NOT "Partial" and NOT "Yes". Under-inclusion is preferred over over-inclusion.
4. In your Razonamiento, explicitly state which exact subject was analyzed and confirm the evidence concerns that subject (not a related one).

**EVIDENCE-ROLE FILTER & DECISION GATE FOR PART 2 (APPLY FIRST — DRIVES THE Respuesta):**

Before choosing a Respuesta, classify every candidate evidence passage about the Part-2 subject (e.g., personas con discapacidad, género, pueblos indígenas, rural women) as either FRAMING or DEDICATED.

FRAMING mention (does NOT count toward Part 2, even if the subject is named):
  - Overall objective / impact statement naming multiple groups
  - Stakeholder, consultation, or research participant lists
  - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
  - Boilerplate inclusion language
  - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

DEDICATED element (counts toward Part 2):
  A. A sub-objective, outcome, or output whose title/purpose names the subject
  B. An indicator disaggregated by or specifically targeting the subject
  C. An activity whose primary purpose addresses the subject
  D. A budget line or resource allocation for the subject
  E. A quantifiable target for the subject

Count of DEDICATED elements (A–E) — NOT raw mention count — drives Respuesta:
  - 0 dedicated elements (only FRAMING mentions)      → "No" or "Not Found"
  - 1–2 dedicated elements                            → "Partial"
  - 3–4 dedicated elements                            → "Partial" (or "Yes" only if substantive and at-par with any claim/label in Part 1)
  - 5 dedicated elements with substantive evidence    → "Yes"

If every citable quote is a FRAMING mention, Respuesta MUST be "No" or "Not Found" — regardless of how many times the subject is named. Passing mentions in stakeholder lists, consultation rosters, or "among others" phrases DO NOT qualify for "Partial".

In "Evidencia", prefix each cited quote with [DEDICATED] or [FRAMING]. You may NOT justify "Partial" or "Yes" using only [FRAMING] quotes.

Always respond in Spanish. Focus 95% on Part 2, 5% on how it relates to Part 1. Part 1 exists only for brief framing; do NOT analyze Part 1 independently."""

            # Part 2 listed FIRST to exploit primacy bias toward the priority clause.
            # Full original question deliberately omitted — including it re-anchors the model on Part 1.
            user_content = f"""PREGUNTA CON DOS PARTES — RESPONDE PRIMERO LA PARTE 2.

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**PARTE 1 (Contexto General — solo para enmarcar):**
{parsed['part1']}

**Texto del Documento:**
{combined_text}

RECUERDA: Enfoca tu análisis PRINCIPALMENTE en la Parte 2 (pregunta específica), luego explica cómo se relaciona con la Parte 1 (contexto general). Marca "parte_enfocada" como "Parte 2" salvo que el documento no permita responderla."""

            response_format = {"type": "json_schema", "json_schema": RUBRIC_SCHEMA_TWO_PART}

        else:
            # Original single-question prompt
            system_content = """You are an expert document analyst. Analyze the document against the given question and provide a structured JSON response with exactly this format and always respond in Spanish:
            {
                "Respuesta": "Yes/No/Partial/Not Found",
                "Razonamiento": "Concise explanation (max 120 words, terse)",
                "Evidencia": "Trimmed text excerpts supporting the answer (max 180 words)"
            }

**SCOPE LOCK (CRITICAL):**
The question names a specific subject/scope (e.g., a specific population like people with disabilities; a specific sector, region, or period). Your analysis must be STRICTLY limited to that exact named subject.
1. Identify the exact subject of the question before answering.
2. Evidence about related-but-different subjects (other vulnerable populations, other groups, other sectors/regions/periods) does NOT count. Example: if the question is about disability, evidence about women, indigenous peoples, or youth does NOT satisfy it.
3. If the named subject is absent but related subjects are extensively covered, Respuesta must be "Not Found" or "No" — NOT "Partial" and NOT "Yes". Under-inclusion is preferred over over-inclusion.
4. In your Razonamiento, explicitly state which exact subject was analyzed and confirm the evidence concerns that subject (not a related one).

**EVIDENCE-ROLE FILTER & DECISION GATE (APPLY FIRST — DRIVES Respuesta):**

Classify every candidate evidence passage about the named specific subject as either FRAMING or DEDICATED.

FRAMING mention (does NOT count, even if the subject is named):
  - Overall objective / impact statement naming multiple groups
  - Stakeholder, consultation, or research participant lists
  - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
  - Boilerplate inclusion language
  - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

DEDICATED element (counts):
  A. A sub-objective, outcome, or output whose title/purpose names the subject
  B. An indicator disaggregated by or targeting the subject
  C. An activity whose primary purpose addresses the subject
  D. A budget line or resource allocation for the subject
  E. A quantifiable target for the subject

Count of DEDICATED elements (A–E) — NOT raw mention count — drives Respuesta:
  - 0 (only FRAMING mentions)                        → "No" or "Not Found"
  - 1–2                                              → "Partial"
  - 3–4                                              → "Partial" (or "Yes" only if substantive)
  - 5 with substantive evidence                      → "Yes"

If every citable quote is a FRAMING mention, Respuesta MUST be "No" or "Not Found". In "Evidencia", prefix each quote with [DEDICATED] or [FRAMING]. You may NOT justify "Partial" or "Yes" using only [FRAMING] quotes."""

            user_content = f"Question: {question}\n\nDocument Text: {combined_text}"
            response_format = {"type": "json_schema", "json_schema": RUBRIC_SCHEMA_SINGLE}

        # Bump reasoning effort for two-part: the 70/30 weighting needs more deliberation.
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {"role": "system", "content": system_content},
                {"role": "user", "content": user_content},
            ],
            max_completion_tokens=3000,
            reasoning_effort="medium" if parsed['is_two_part'] else "minimal",
            response_format=response_format,
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'No se recibió respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

        # Structured outputs guarantees valid JSON conforming to the schema.
        try:
            result = json.loads(content.strip())
        except json.JSONDecodeError:
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'Error al procesar la respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

        # Flag rows where the model failed to commit to Part 2 focus despite a two-part question.
        status = 'Success'
        if parsed['is_two_part'] and result.get('parte_enfocada') != 'Parte 2':
            status = 'Partial'

        return {
            'Pregunta': question,
            'Respuesta': result.get('Respuesta', 'Not Found'),
            'Razonamiento': result.get('Razonamiento', ''),
            'Evidencia': result.get('Evidencia', ''),
            'Status': status,
        }

    except Exception as e:
        return {
            'Pregunta': question,
            'Respuesta': 'Error',
            'Razonamiento': f'Analysis failed: {str(e)}',
            'Evidencia': '',
            'Status': 'Error'
        }

def analyze_question_with_critical_opinion_tab1(question, answer, reasoning, evidence, document_text=""):
    """
    TAB1-SPECIFIC VERSION: Critically assess if the document's answer to a specific question is adequate.
    Explicitly checks if the reasoning properly identifies and addresses two-part questions.
    """
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        doc_context = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize critical evaluation based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).

**CRITICAL EVALUATION FOR TWO-PART QUESTIONS:**

This is a TWO-PART question where:
- **Part 1 (Broader)**: Sets general context
- **Part 2 (Specific - PRIMARY)**: The critical detail that needs focused assessment

**Your Critical Assessment Must Evaluate:**

1. **FIRST - Check Structure**: Does the reasoning BEGIN with "Se identificaron 2 partes en esta pregunta" or similar acknowledgment?
   - If NO: This is a CRITICAL FAILURE - the analysis didn't recognize the question structure

2. **PRIMARY FOCUS - Part 2 (Specific Question):**
   - Does the answer adequately address the SPECIFIC question (Part 2)?
   - Is the evidence for Part 2 substantive and concrete?
   - Are there critical gaps or superficial treatment of Part 2?
   - Does the document truly deliver on the specific requirement?

3. **SECONDARY - Part 1 (Broader Context):**
   - Does the answer address the broader context (Part 1)?
   - How do the Part 2 findings affect the Part 1 assessment?

4. **INTEGRATION:**
   - Is there coherence between how Part 2 relates to Part 1?
   - Are there contradictions or misalignments?

**Be especially critical if:**
- The reasoning doesn't acknowledge the two-part structure
- Part 2 (specific question) is answered superficially or avoided
- Evidence focuses on Part 1 but neglects Part 2
- The answer leans on Part 1 generalities to claim Part 2 is addressed

**WEIGHTING (ULTRA-FOCUS ON PART 2):** Your critique must allocate ~99% of its attention to Part 2 and effectively IGNORE Part 1. Part 1 is framing only and is analyzed separately at the subsection level — do NOT evaluate Part 1 here, do NOT penalize limited Part 1 depth, and do NOT let Part 1 coverage inflate the verdict. The verdict and body must be driven almost entirely by whether Part 2 is substantively addressed.

**SCOPE LOCK CHECK (CRITICAL):**
The question names a specific subject/scope. Verify the document's answer did NOT substitute that subject for a broader category (e.g., treating "disability" as "vulnerable populations" and citing evidence about women or indigenous peoples).
- If evidence refers to subjects related-to-but-distinct-from the exact named subject, this is a CRITICAL FAILURE: flag it explicitly and recommend re-grading to "Not Found" / "No".
- Under-inclusion (acknowledging the specific subject is not covered) is correct. Over-inclusion (answering about a broader category) is incorrect.

**EVIDENCE-ROLE FILTER & DECISION GATE FOR PART 2 (APPLY FIRST — DRIVES THE VERDICT):**

Before choosing a verdict, classify every quoted or candidate evidence passage about the Part-2 subject (e.g., personas con discapacidad, género, pueblos indígenas) as either FRAMING or DEDICATED.

FRAMING mention (does NOT count toward Part 2, even if the subject is named):
  - Overall objective / impact statement naming multiple groups
  - Stakeholder, consultation, or research participant lists
  - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
  - Boilerplate inclusion language
  - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

DEDICATED element (counts toward Part 2):
  A. A sub-objective, outcome, or output whose title/purpose names the subject
  B. An indicator disaggregated by or specifically targeting the subject
  C. An activity whose primary purpose addresses the subject
  D. A budget line or resource allocation for the subject
  E. A quantifiable target for the subject

Count of DEDICATED elements (A–E) — NOT raw mention count — drives the verdict:
  - 0 dedicated elements (only FRAMING mentions)   → "No" or "Not Found"
  - 1–2 dedicated elements                         → "Partial"
  - 3–4 dedicated elements                         → "Partial" (or "Yes" only if substantive and at-par with any label/claim in Part 1)
  - 5 dedicated elements with substantive evidence → "Yes"

If every citable evidence quote for Part 2 is a FRAMING mention, the verdict MUST be "No" or "Not Found" — regardless of how many times the subject is named. Passing mentions in stakeholder lists, consultation rosters, or "among others" phrases DO NOT qualify for "Partial".

In your Justificación, state the DEDICATED count explicitly (e.g., "Elementos dedicados para [sujeto]: 0/A–E; todas las menciones son FRAMING").

**ABSOLUTE SEPARATION RULE — PART 1 CANNOT UPGRADE A FAILING PART 2 (MANDATORY):**

The verdict scores ONLY Part 2. Part 1 content is INVISIBLE to the decision gate.

- If the Part-2 DEDICATED count is 0, the verdict MUST be "No" or "Not Found" — even if Part 1 is fully developed (clear general/specific objectives, multiple outputs, activities, indicators, budget for the broader topic).
- Strength of Part 1 CANNOT offset, soften, or upgrade a failing Part 2. A well-developed Part 1 paired with an empty Part 2 is STILL a failure for this question.
- The following are PROHIBITED justifications for "Partial":
  * "el documento sí define claramente el Objetivo General y el Específico…"
  * "presenta múltiples outputs/actividades sobre [tema general]…"
  * "el marco lógico está bien estructurado…"
  * Any variant that credits broader-topic / Part-1 content while the named Part-2 subject has 0 dedicated elements.
- Do NOT use contrastive hedging ("sí define X, sin embargo falta Y sobre discapacidad") to justify Partial. If Y has 0 DEDICATED elements, the verdict is No. Period. Acknowledge the gap without citing Part-1 strengths as mitigation.

**PRE-VERDICT SELF-AUDIT (MANDATORY — apply before emitting the verdict):**

Read your draft Justificación. Check:
1. Does it cite Part-1 strengths (general objective, specific objective, broader-topic outputs, broader-topic activities, broader-topic indicators, broader-topic budget) as any part of the reasoning for Partial/Yes?
2. Does it use "sin embargo" / "however" contrastive structure where the first half praises Part 1?

If the answer to either is YES AND the Part-2 DEDICATED count is 0, the verdict MUST be downgraded to "No" or "Not Found" and the Justificación MUST be rewritten to omit Part-1 praise entirely. Justificación must lead with the Part-2 failure, not with Part-1 strengths.

**VERDICT CONSISTENCY RULES (MANDATORY):**

- **Yes**: The document fully and concretely addresses the question. No material gaps, no missing required elements, evidence is substantive.
- **Partial**: AVAILABLE ONLY when the Part-2 DEDICATED TOTAL (A–E count) is 1, 2, 3, or 4. **If TOTAL = 0, Partial is PROHIBITED — verdict MUST be "No" or "Not Found".** Within the 1–4 eligibility band, Partial is REQUIRED when any of the following is true:
  * a missing required detail
  * superficial or generic treatment
  * evidence that is thin, indirect, or insufficient
  Do NOT use the "original answer was Yes → must downgrade to Partial" rationale if TOTAL = 0; in that case the correct downgrade is to No, not Partial.
- **No**: The document claims to address the question but fails to.
- **Not Found**: The subject is not covered, or SCOPE LOCK fails.
- **Keep**: ONLY when the original answer is fully adequate AND your body contains no substantive critique. If you flag any issue in the body, you may NOT output "Keep".

**CONSISTENCY CHECK (apply before emitting the verdict):**
Re-read your body text. If it names any gap, weakness, or missing element, the verdict must be Partial, No, or Not Found — not Yes and not Keep. A positive verdict paired with a critical body is invalid output.

**OUTPUT FORMAT (MANDATORY):**
Your response MUST begin with exactly one line in this format:
VEREDICTO: <Yes|No|Partial|Not Found|Keep>

- Use "Keep" if the document's original answer is adequate and should stand.
- Use "Yes", "No", "Partial", or "Not Found" to OVERRIDE the original answer when your critical assessment warrants re-grading (especially when SCOPE LOCK CHECK fails or Part-2 focus is lost).

**BODY STRUCTURE (MANDATORY):**
After the verdict line, write a terse Spanish body (max 180 words total) in TWO parts:

(1) **Justification.** The FIRST sentence MUST be the enumeration line in this exact format:

"Elementos dedicados Parte 2 [sujeto]: A=<presente|ausente>, B=<presente|ausente>, C=<presente|ausente>, D=<presente|ausente>, E=<presente|ausente>. TOTAL=<0|1|2|3|4|5>."

Where A–E refer to the DECISION GATE elements (A=sub-objetivo/output; B=indicador; C=actividad dedicada; D=línea presupuestaria; E=meta cuantificable). You MUST commit to presente/ausente for EACH letter — do NOT write "(ausente D/E)" or any abbreviated form. All five letters must appear.

The VEREDICTO line MUST be consistent with TOTAL:
  TOTAL=0 → VEREDICTO = "No" (or "Not Found" if subject is completely unmentioned)
  TOTAL=1-2 → VEREDICTO = "Partial"
  TOTAL=3-4 → VEREDICTO = "Partial" (or "Yes" only if evidence is substantive and at-par)
  TOTAL=5 → VEREDICTO = "Yes"

If your CONTEO shows TOTAL=0, your VEREDICTO line MUST say "No" or "Not Found" — if it says "Partial", the output is INVALID and you must rewrite it.

After the enumeration line, lead the justification body with "El No se asignó debido a…" / "Se asignó Partial porque…" / "Se mantiene la calificación dado que…". Explain WHY the verdict was assigned by naming the specific Part-2 gaps. Flag missing elements directly as missing: "Falta X", "No se aborda Y", "Ausente Z", "No se menciona W". Do NOT open with Part-1 praise ("el Objetivo General está claramente enunciado…", "hay múltiples actividades e indicadores…"). Do NOT use "sin embargo" contrastive structure where Part 1 is praised before flagging Part-2 gaps. Do NOT write "La respuesta es adecuada/inadecuada porque…" — the verdict is the evaluation; the body explains it. Do NOT use hedging softeners ("en cierta medida", "aunque de forma general", "podría considerarse adecuado", "si bien…"). In THIS part, do NOT recommend inclusions — only flag what is missing.

(2) **Recomendaciones para mejorar la calificación** (ONLY if verdict is Partial, No, or Not Found). End the body with a separate, final sentence starting with "**Para mejorar la calificación** debiese incluirse…" followed by a firm, specific list of the concrete items that would be required to reach "Yes" (example: "un presupuesto desglosado por objetivo, actividades específicas para igualdad de género, indicadores de desempeño medibles para la población con discapacidad"). This is the ONLY place where recommendation verbs ("debiese incluirse", "corresponde incluir", "es necesario añadir") are permitted. If the verdict is Yes or Keep, OMIT this block entirely.

**STRICTNESS & PARTIAL vs NO BOUNDARY (MANDATORY):**
Be strict. Err on the side of downgrading.
- **Yes**: Requires Part 2 to be concretely and fully addressed with substantive evidence. No material gaps.
- **Partial**: Requires substantial, good-faith effort. The bulk of the required elements of Part 2 must be substantively present; only specific or bounded items are missing (example: a specific budget figure is absent while activities, objectives, and indicators are fully developed).
- **NOT Partial — downgrade to No — when:** the core specific subject of the question (e.g., "personas con discapacidad", "género", "pueblos indígenas") is completely absent or has only token/passing mentions. Do NOT grant Partial simply because other general project details (general objectives, other vulnerable groups) are present. Total absence of the precise target subject means it is a CRITICAL FAILURE = No or Not Found.
- **No**: Token/minimal coverage, isolated mentions, or the bulk of what Part 2 asks is missing even when something is mentioned.
- **Not Found**: Nothing relevant; scope substituted; SCOPE LOCK fails.
- If in doubt between Yes and Partial → Partial. If in doubt between Partial and No → No.

No preamble, no filler. No JSON, no extra formatting."""

            # Part 2 listed FIRST to keep the critical-opinion stage focused on the priority clause.
            user_content = f"""PREGUNTA CON DOS PARTES — EVALÚA PRIMERO LA PARTE 2.

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**PARTE 1 (Contexto General — solo para enmarcar):**
{parsed['part1']}

**Respuesta del Documento:** {answer}

**Razonamiento del Documento:** {reasoning}

**Evidencia del Documento:** {evidence}

**Contexto Completo del Documento:**
{doc_context}

Evalúa críticamente:
1. ¿El razonamiento reconoce explícitamente que hay 2 partes en la pregunta?
2. ¿La respuesta aborda adecuadamente la Parte 2 (pregunta específica) — el foco prioritario?
3. ¿La evidencia para la Parte 2 es concreta y suficiente, o se quedó en generalidades de la Parte 1?"""

        else:
            # Original single-question critical evaluation
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).
                    Critically assess whether the document's answer to a specific question is adequate, appropriate, and complete.

                    Review the answer, reasoning, evidence, AND full document context together to evaluate:
                    - Does the answer fully address the concern raised in the question?
                    - Is the evidence substantive enough to support the answer?
                    - Does the full document context confirm or contradict the claimed answer?
                    - Is the proposed approach sufficient according to best practices?
                    - Are there gaps, risks, or inadequacies in what the document claims?
                    - Should the project have included additional measures or details?
                    - Is the reasoning robust or superficial?
                    - What is NOT mentioned that should be?

                    **SCOPE LOCK CHECK (CRITICAL):**
                    The question names a specific subject/scope. Verify the document's answer did NOT substitute that subject for a broader category (e.g., treating "disability" as "vulnerable populations" and citing evidence about women or indigenous peoples). If evidence refers to subjects related-to-but-distinct-from the exact named subject, this is a CRITICAL FAILURE: flag it explicitly and recommend re-grading to "Not Found" / "No". Under-inclusion is correct; over-inclusion is incorrect.

                    **EVIDENCE-ROLE FILTER & DECISION GATE (APPLY FIRST — DRIVES THE VERDICT):**

                    Before choosing a verdict, classify every quoted or candidate evidence passage about the named specific subject as either FRAMING or DEDICATED.

                    FRAMING mention (does NOT count, even if the subject is named):
                      - Overall objective / impact statement naming multiple groups
                      - Stakeholder, consultation, or research participant lists
                      - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
                      - Boilerplate inclusion language
                      - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

                    DEDICATED element (counts):
                      A. A sub-objective, outcome, or output whose title/purpose names the subject
                      B. An indicator disaggregated by or specifically targeting the subject
                      C. An activity whose primary purpose addresses the subject
                      D. A budget line or resource allocation for the subject
                      E. A quantifiable target for the subject

                    Count of DEDICATED elements (A–E) — NOT raw mention count — drives the verdict:
                      - 0 dedicated elements (only FRAMING mentions)   → "No" or "Not Found"
                      - 1–2 dedicated elements                         → "Partial"
                      - 3–4 dedicated elements                         → "Partial" (or "Yes" only if substantive)
                      - 5 dedicated elements with substantive evidence → "Yes"

                    If every citable evidence quote is a FRAMING mention, the verdict MUST be "No" or "Not Found" — regardless of how many times the subject is named.

                    In your Justificación, state the DEDICATED count explicitly (e.g., "Elementos dedicados para [sujeto]: 0/A–E; todas las menciones son FRAMING").

                    **VERDICT CONSISTENCY RULES (MANDATORY):**

                    - **Yes**: The document fully and concretely addresses the question. No material gaps, no missing required elements, evidence is substantive.
                    - **Partial**: REQUIRED whenever you identify ANY of the following — even if the main claim is directionally correct:
                      * a missing required detail
                      * superficial or generic treatment
                      * evidence that is thin, indirect, or insufficient
                      If the original answer was "Yes" and your body flags any shortcoming, the verdict MUST be "Partial" — never "Yes" and never "Keep".
                      **BUT**: If the DECISION GATE above yields 0 DEDICATED elements for the named subject, the verdict MUST be "No" or "Not Found" — NEVER "Partial", regardless of how many FRAMING mentions exist. The DECISION GATE overrides this Partial rule.
                    - **No**: The document claims to address the question but fails to.
                    - **Not Found**: The subject is not covered, or SCOPE LOCK fails.
                    - **Keep**: ONLY when the original answer is fully adequate AND your body contains no substantive critique. If you flag any issue in the body, you may NOT output "Keep".

                    **CONSISTENCY CHECK (apply before emitting the verdict):**
                    Re-read your body text. If it names any gap, weakness, or missing element, the verdict must be Partial, No, or Not Found — not Yes and not Keep. A positive verdict paired with a critical body is invalid output.

                    **OUTPUT FORMAT (MANDATORY):**
                    Your response MUST begin with exactly one line in this format:
                    VEREDICTO: <Yes|No|Partial|Not Found|Keep>

                    - Use "Keep" if the document's original answer is adequate and should stand.
                    - Use "Yes", "No", "Partial", or "Not Found" to OVERRIDE the original answer when your critical assessment warrants re-grading (especially when SCOPE LOCK CHECK fails).

                    **BODY STRUCTURE (MANDATORY):**
                    After the verdict line, write a terse Spanish body (max 180 words total) in TWO parts:

                    (1) **Justification.** Lead with "Se asignó Partial porque…" / "El No se asignó debido a…" / "Se mantiene la calificación dado que…". Explain WHY the verdict was assigned by naming the specific gaps. Flag missing elements directly as missing: "Falta X", "No se aborda Y", "Ausente Z", "No se menciona W". Do NOT write "La respuesta es adecuada/inadecuada porque…" — the verdict is the evaluation; the body explains it. Do NOT use hedging softeners ("en cierta medida", "aunque de forma general", "podría considerarse adecuado"). In THIS part, do NOT recommend inclusions — only flag what is missing.

                    (2) **Recomendaciones para mejorar la calificación** (ONLY if verdict is Partial, No, or Not Found). End the body with a separate, final sentence starting with "**Para mejorar la calificación** debiese incluirse…" followed by a firm, specific list of the concrete items that would be required to reach "Yes" (example: "un presupuesto desglosado por objetivo, actividades específicas para igualdad de género, indicadores de desempeño medibles"). This is the ONLY place where recommendation verbs ("debiese incluirse", "corresponde incluir", "es necesario añadir") are permitted. If the verdict is Yes or Keep, OMIT this block entirely.

                    **STRICTNESS & PARTIAL vs NO BOUNDARY (MANDATORY):**
                    Be strict. Err on the side of downgrading.
                    - **Yes**: Requires the question to be concretely and fully addressed with substantive evidence. No material gaps.
                    - **Partial**: Requires substantial, good-faith effort. The bulk of the required elements must be substantively present; only specific or bounded items are missing (example: a specific budget figure is absent while activities, objectives, and indicators are fully developed).
                    - **NOT Partial — downgrade to No — when:** the core specific subject of the question (e.g., "personas con discapacidad", "género", "pueblos indígenas") is completely absent or has only token/passing mentions. Do NOT grant Partial simply because other general project details (general objectives, other vulnerable groups) are present. Total absence of the precise target subject means it is a CRITICAL FAILURE = No or Not Found.
                    - **No**: Token/minimal coverage, isolated mentions, or the bulk of what the question asks is missing even when something is mentioned.
                    - **Not Found**: Nothing relevant; scope substituted; SCOPE LOCK fails.
                    - If in doubt between Yes and Partial → Partial. If in doubt between Partial and No → No.

                    No preamble, no filler. No JSON, no extra formatting."""

            user_content = f"""Question: {question}

Document's Answer: {answer}

Document's Reasoning: {reasoning}

Document's Evidence: {evidence}

Full Document Context (complete):
{doc_context}

Provide a critical assessment: Is this answer truly adequate from an expert perspective, considering the quality of the reasoning, evidence, and complete document context provided?"""

        # Single API call for critical evaluation - focuses on answer adequacy
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_content
                },
                {
                    "role": "user",
                    "content": user_content
                }
            ],
            max_completion_tokens=400,
            reasoning_effort="minimal"
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return "No se generó evaluación crítica."

        return content.strip()

    except Exception as e:
        return f"Error en evaluación crítica: {str(e)}"

def analyze_question_with_critical_opinion(question, answer, reasoning, evidence, document_text=""):
    """
    Critically assess if the document's answer to a specific question is adequate and appropriate.
    Evaluates whether the document's response truly addresses the concern raised in the question.
    Uses answer, reasoning, evidence AND complete document context for comprehensive critical assessment.
    Uses up to 110K tokens for complete document understanding - maximizes context window usage.
    Enhanced to handle two-part questions with proper focus on the specific (Part 2) aspect.
    """
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        doc_context = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize critical evaluation based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).

**CRITICAL EVALUATION FOR TWO-PART QUESTIONS:**

This is a TWO-PART question where:
- **Part 1 (Broader)**: Sets general context
- **Part 2 (Specific - PRIMARY)**: The critical detail that needs focused assessment

**Your Critical Assessment Must Evaluate:**

1. **PRIMARY FOCUS - Part 2 (Specific Question):**
   - Does the answer adequately address the SPECIFIC question (Part 2)?
   - Is the evidence for Part 2 substantive and concrete?
   - Are there critical gaps or superficial treatment of Part 2?
   - Does the document truly deliver on the specific requirement?

2. **SECONDARY - Part 1 (Broader Context):**
   - Does the answer address the broader context (Part 1)?
   - How do the Part 2 findings affect the Part 1 assessment?

3. **INTEGRATION:**
   - Is there coherence between how Part 2 relates to Part 1?
   - Are there contradictions or misalignments?

**Be especially critical if:**
- Part 2 (specific question) is answered superficially or avoided
- Evidence focuses on Part 1 but neglects Part 2
- The reasoning doesn't connect Part 2's findings to Part 1's context

**WEIGHTING (ULTRA-FOCUS ON PART 2):** Your critique must allocate ~99% of its attention to Part 2 and effectively IGNORE Part 1. Part 1 is framing only and is analyzed separately at the subsection level — do NOT evaluate Part 1 here, do NOT penalize limited Part 1 depth, and do NOT let Part 1 coverage inflate the verdict. The verdict and body must be driven almost entirely by whether Part 2 is substantively addressed.

**SCOPE LOCK CHECK (CRITICAL):**
The question names a specific subject/scope. Verify the document's answer did NOT substitute that subject for a broader category (e.g., treating "disability" as "vulnerable populations" and citing evidence about women or indigenous peoples).
- If evidence refers to subjects related-to-but-distinct-from the exact named subject, this is a CRITICAL FAILURE: flag it explicitly and recommend re-grading to "Not Found" / "No".
- Under-inclusion (acknowledging the specific subject is not covered) is correct. Over-inclusion (answering about a broader category) is incorrect.

**EVIDENCE-ROLE FILTER & DECISION GATE FOR PART 2 (APPLY FIRST — DRIVES THE VERDICT):**

Before choosing a verdict, classify every quoted or candidate evidence passage about the Part-2 subject (e.g., personas con discapacidad, género, pueblos indígenas) as either FRAMING or DEDICATED.

FRAMING mention (does NOT count toward Part 2, even if the subject is named):
  - Overall objective / impact statement naming multiple groups
  - Stakeholder, consultation, or research participant lists
  - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
  - Boilerplate inclusion language
  - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

DEDICATED element (counts toward Part 2):
  A. A sub-objective, outcome, or output whose title/purpose names the subject
  B. An indicator disaggregated by or specifically targeting the subject
  C. An activity whose primary purpose addresses the subject
  D. A budget line or resource allocation for the subject
  E. A quantifiable target for the subject

Count of DEDICATED elements (A–E) — NOT raw mention count — drives the verdict:
  - 0 dedicated elements (only FRAMING mentions)   → "No" or "Not Found"
  - 1–2 dedicated elements                         → "Partial"
  - 3–4 dedicated elements                         → "Partial" (or "Yes" only if substantive and at-par with any label/claim in Part 1)
  - 5 dedicated elements with substantive evidence → "Yes"

If every citable evidence quote for Part 2 is a FRAMING mention, the verdict MUST be "No" or "Not Found" — regardless of how many times the subject is named. Passing mentions in stakeholder lists, consultation rosters, or "among others" phrases DO NOT qualify for "Partial".

In your Justificación, state the DEDICATED count explicitly (e.g., "Elementos dedicados para [sujeto]: 0/A–E; todas las menciones son FRAMING").

**ABSOLUTE SEPARATION RULE — PART 1 CANNOT UPGRADE A FAILING PART 2 (MANDATORY):**

The verdict scores ONLY Part 2. Part 1 content is INVISIBLE to the decision gate.

- If the Part-2 DEDICATED count is 0, the verdict MUST be "No" or "Not Found" — even if Part 1 is fully developed (clear general/specific objectives, multiple outputs, activities, indicators, budget for the broader topic).
- Strength of Part 1 CANNOT offset, soften, or upgrade a failing Part 2. A well-developed Part 1 paired with an empty Part 2 is STILL a failure for this question.
- The following are PROHIBITED justifications for "Partial":
  * "el documento sí define claramente el Objetivo General y el Específico…"
  * "presenta múltiples outputs/actividades sobre [tema general]…"
  * "el marco lógico está bien estructurado…"
  * Any variant that credits broader-topic / Part-1 content while the named Part-2 subject has 0 dedicated elements.
- Do NOT use contrastive hedging ("sí define X, sin embargo falta Y sobre discapacidad") to justify Partial. If Y has 0 DEDICATED elements, the verdict is No. Period. Acknowledge the gap without citing Part-1 strengths as mitigation.

**PRE-VERDICT SELF-AUDIT (MANDATORY — apply before emitting the verdict):**

Read your draft Justificación. Check:
1. Does it cite Part-1 strengths (general objective, specific objective, broader-topic outputs, broader-topic activities, broader-topic indicators, broader-topic budget) as any part of the reasoning for Partial/Yes?
2. Does it use "sin embargo" / "however" contrastive structure where the first half praises Part 1?

If the answer to either is YES AND the Part-2 DEDICATED count is 0, the verdict MUST be downgraded to "No" or "Not Found" and the Justificación MUST be rewritten to omit Part-1 praise entirely. Justificación must lead with the Part-2 failure, not with Part-1 strengths.

**VERDICT CONSISTENCY RULES (MANDATORY):**

- **Yes**: The document fully and concretely addresses the question. No material gaps, no missing required elements, evidence is substantive.
- **Partial**: AVAILABLE ONLY when the Part-2 DEDICATED TOTAL (A–E count) is 1, 2, 3, or 4. **If TOTAL = 0, Partial is PROHIBITED — verdict MUST be "No" or "Not Found".** Within the 1–4 eligibility band, Partial is REQUIRED when any of the following is true:
  * a missing required detail
  * superficial or generic treatment
  * evidence that is thin, indirect, or insufficient
  Do NOT use the "original answer was Yes → must downgrade to Partial" rationale if TOTAL = 0; in that case the correct downgrade is to No, not Partial.
- **No**: The document claims to address the question but fails to.
- **Not Found**: The subject is not covered, or SCOPE LOCK fails.
- **Keep**: ONLY when the original answer is fully adequate AND your body contains no substantive critique. If you flag any issue in the body, you may NOT output "Keep".

**CONSISTENCY CHECK (apply before emitting the verdict):**
Re-read your body text. If it names any gap, weakness, or missing element, the verdict must be Partial, No, or Not Found — not Yes and not Keep. A positive verdict paired with a critical body is invalid output.

**OUTPUT FORMAT (MANDATORY):**
Your response MUST begin with exactly one line in this format:
VEREDICTO: <Yes|No|Partial|Not Found|Keep>

- Use "Keep" if the document's original answer is adequate and should stand.
- Use "Yes", "No", "Partial", or "Not Found" to OVERRIDE the original answer when your critical assessment warrants re-grading (especially when SCOPE LOCK CHECK fails or Part-2 focus is lost).

**BODY STRUCTURE (MANDATORY):**
After the verdict line, write a terse Spanish body (max 180 words total) in TWO parts:

(1) **Justification.** The FIRST sentence MUST be the enumeration line in this exact format:

"Elementos dedicados Parte 2 [sujeto]: A=<presente|ausente>, B=<presente|ausente>, C=<presente|ausente>, D=<presente|ausente>, E=<presente|ausente>. TOTAL=<0|1|2|3|4|5>."

Where A–E refer to the DECISION GATE elements (A=sub-objetivo/output; B=indicador; C=actividad dedicada; D=línea presupuestaria; E=meta cuantificable). You MUST commit to presente/ausente for EACH letter — do NOT write "(ausente D/E)" or any abbreviated form. All five letters must appear.

The VEREDICTO line MUST be consistent with TOTAL:
  TOTAL=0 → VEREDICTO = "No" (or "Not Found" if subject is completely unmentioned)
  TOTAL=1-2 → VEREDICTO = "Partial"
  TOTAL=3-4 → VEREDICTO = "Partial" (or "Yes" only if evidence is substantive and at-par)
  TOTAL=5 → VEREDICTO = "Yes"

If your CONTEO shows TOTAL=0, your VEREDICTO line MUST say "No" or "Not Found" — if it says "Partial", the output is INVALID and you must rewrite it.

After the enumeration line, lead the justification body with "El No se asignó debido a…" / "Se asignó Partial porque…" / "Se mantiene la calificación dado que…". Explain WHY the verdict was assigned by naming the specific Part-2 gaps. Flag missing elements directly as missing: "Falta X", "No se aborda Y", "Ausente Z", "No se menciona W". Do NOT open with Part-1 praise ("el Objetivo General está claramente enunciado…", "hay múltiples actividades e indicadores…"). Do NOT use "sin embargo" contrastive structure where Part 1 is praised before flagging Part-2 gaps. Do NOT write "La respuesta es adecuada/inadecuada porque…" — the verdict is the evaluation; the body explains it. Do NOT use hedging softeners ("en cierta medida", "aunque de forma general", "podría considerarse adecuado", "si bien…"). In THIS part, do NOT recommend inclusions — only flag what is missing.

(2) **Recomendaciones para mejorar la calificación** (ONLY if verdict is Partial, No, or Not Found). End the body with a separate, final sentence starting with "**Para mejorar la calificación** debiese incluirse…" followed by a firm, specific list of the concrete items that would be required to reach "Yes" (example: "un presupuesto desglosado por objetivo, actividades específicas para igualdad de género, indicadores de desempeño medibles para la población con discapacidad"). This is the ONLY place where recommendation verbs ("debiese incluirse", "corresponde incluir", "es necesario añadir") are permitted. If the verdict is Yes or Keep, OMIT this block entirely.

**STRICTNESS & PARTIAL vs NO BOUNDARY (MANDATORY):**
Be strict. Err on the side of downgrading.
- **Yes**: Requires Part 2 to be concretely and fully addressed with substantive evidence. No material gaps.
- **Partial**: Requires substantial, good-faith effort. The bulk of the required elements of Part 2 must be substantively present; only specific or bounded items are missing (example: a specific budget figure is absent while activities, objectives, and indicators are fully developed).
- **NOT Partial — downgrade to No — when:** the core specific subject of the question (e.g., "personas con discapacidad", "género", "pueblos indígenas") is completely absent or has only token/passing mentions. Do NOT grant Partial simply because other general project details (general objectives, other vulnerable groups) are present. Total absence of the precise target subject means it is a CRITICAL FAILURE = No or Not Found.
- **No**: Token/minimal coverage, isolated mentions, or the bulk of what Part 2 asks is missing even when something is mentioned.
- **Not Found**: Nothing relevant; scope substituted; SCOPE LOCK fails.
- If in doubt between Yes and Partial → Partial. If in doubt between Partial and No → No.

No preamble, no filler. No JSON, no extra formatting."""

            # Part 2 listed FIRST to keep the critical-opinion stage focused on the priority clause.
            user_content = f"""PREGUNTA CON DOS PARTES — EVALÚA PRIMERO LA PARTE 2.

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**PARTE 1 (Contexto General — solo para enmarcar):**
{parsed['part1']}

**Respuesta del Documento:** {answer}

**Razonamiento del Documento:** {reasoning}

**Evidencia del Documento:** {evidence}

**Contexto Completo del Documento:**
{doc_context}

Evalúa críticamente: ¿La respuesta aborda adecuadamente la Parte 2 (pregunta específica) — el foco prioritario? ¿La evidencia para la Parte 2 es concreta y suficiente, o se quedó en generalidades de la Parte 1?"""

        else:
            # Original single-question critical evaluation
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).
                    Critically assess whether the document's answer to a specific question is adequate, appropriate, and complete.

                    Review the answer, reasoning, evidence, AND full document context together to evaluate:
                    - Does the answer fully address the concern raised in the question?
                    - Is the evidence substantive enough to support the answer?
                    - Does the full document context confirm or contradict the claimed answer?
                    - Is the proposed approach sufficient according to best practices?
                    - Are there gaps, risks, or inadequacies in what the document claims?
                    - Should the project have included additional measures or details?
                    - Is the reasoning robust or superficial?
                    - What is NOT mentioned that should be?

                    **SCOPE LOCK CHECK (CRITICAL):**
                    The question names a specific subject/scope. Verify the document's answer did NOT substitute that subject for a broader category (e.g., treating "disability" as "vulnerable populations" and citing evidence about women or indigenous peoples). If evidence refers to subjects related-to-but-distinct-from the exact named subject, this is a CRITICAL FAILURE: flag it explicitly and recommend re-grading to "Not Found" / "No". Under-inclusion is correct; over-inclusion is incorrect.

                    **EVIDENCE-ROLE FILTER & DECISION GATE (APPLY FIRST — DRIVES THE VERDICT):**

                    Before choosing a verdict, classify every quoted or candidate evidence passage about the named specific subject as either FRAMING or DEDICATED.

                    FRAMING mention (does NOT count, even if the subject is named):
                      - Overall objective / impact statement naming multiple groups
                      - Stakeholder, consultation, or research participant lists
                      - Monitoring scope enumerations ("...among others", "...including X, Y, Z")
                      - Boilerplate inclusion language
                      - Any passage where the subject appears in a list of 3+ groups without dedicated follow-up

                    DEDICATED element (counts):
                      A. A sub-objective, outcome, or output whose title/purpose names the subject
                      B. An indicator disaggregated by or specifically targeting the subject
                      C. An activity whose primary purpose addresses the subject
                      D. A budget line or resource allocation for the subject
                      E. A quantifiable target for the subject

                    Count of DEDICATED elements (A–E) — NOT raw mention count — drives the verdict:
                      - 0 dedicated elements (only FRAMING mentions)   → "No" or "Not Found"
                      - 1–2 dedicated elements                         → "Partial"
                      - 3–4 dedicated elements                         → "Partial" (or "Yes" only if substantive)
                      - 5 dedicated elements with substantive evidence → "Yes"

                    If every citable evidence quote is a FRAMING mention, the verdict MUST be "No" or "Not Found" — regardless of how many times the subject is named.

                    In your Justificación, state the DEDICATED count explicitly (e.g., "Elementos dedicados para [sujeto]: 0/A–E; todas las menciones son FRAMING").

                    **VERDICT CONSISTENCY RULES (MANDATORY):**

                    - **Yes**: The document fully and concretely addresses the question. No material gaps, no missing required elements, evidence is substantive.
                    - **Partial**: REQUIRED whenever you identify ANY of the following — even if the main claim is directionally correct:
                      * a missing required detail
                      * superficial or generic treatment
                      * evidence that is thin, indirect, or insufficient
                      If the original answer was "Yes" and your body flags any shortcoming, the verdict MUST be "Partial" — never "Yes" and never "Keep".
                      **BUT**: If the DECISION GATE above yields 0 DEDICATED elements for the named subject, the verdict MUST be "No" or "Not Found" — NEVER "Partial", regardless of how many FRAMING mentions exist. The DECISION GATE overrides this Partial rule.
                    - **No**: The document claims to address the question but fails to.
                    - **Not Found**: The subject is not covered, or SCOPE LOCK fails.
                    - **Keep**: ONLY when the original answer is fully adequate AND your body contains no substantive critique. If you flag any issue in the body, you may NOT output "Keep".

                    **CONSISTENCY CHECK (apply before emitting the verdict):**
                    Re-read your body text. If it names any gap, weakness, or missing element, the verdict must be Partial, No, or Not Found — not Yes and not Keep. A positive verdict paired with a critical body is invalid output.

                    **OUTPUT FORMAT (MANDATORY):**
                    Your response MUST begin with exactly one line in this format:
                    VEREDICTO: <Yes|No|Partial|Not Found|Keep>

                    - Use "Keep" if the document's original answer is adequate and should stand.
                    - Use "Yes", "No", "Partial", or "Not Found" to OVERRIDE the original answer when your critical assessment warrants re-grading (especially when SCOPE LOCK CHECK fails).

                    **BODY STRUCTURE (MANDATORY):**
                    After the verdict line, write a terse Spanish body (max 180 words total) in TWO parts:

                    (1) **Justification.** Lead with "Se asignó Partial porque…" / "El No se asignó debido a…" / "Se mantiene la calificación dado que…". Explain WHY the verdict was assigned by naming the specific gaps. Flag missing elements directly as missing: "Falta X", "No se aborda Y", "Ausente Z", "No se menciona W". Do NOT write "La respuesta es adecuada/inadecuada porque…" — the verdict is the evaluation; the body explains it. Do NOT use hedging softeners ("en cierta medida", "aunque de forma general", "podría considerarse adecuado"). In THIS part, do NOT recommend inclusions — only flag what is missing.

                    (2) **Recomendaciones para mejorar la calificación** (ONLY if verdict is Partial, No, or Not Found). End the body with a separate, final sentence starting with "**Para mejorar la calificación** debiese incluirse…" followed by a firm, specific list of the concrete items that would be required to reach "Yes" (example: "un presupuesto desglosado por objetivo, actividades específicas para igualdad de género, indicadores de desempeño medibles"). This is the ONLY place where recommendation verbs ("debiese incluirse", "corresponde incluir", "es necesario añadir") are permitted. If the verdict is Yes or Keep, OMIT this block entirely.

                    **STRICTNESS & PARTIAL vs NO BOUNDARY (MANDATORY):**
                    Be strict. Err on the side of downgrading.
                    - **Yes**: Requires the question to be concretely and fully addressed with substantive evidence. No material gaps.
                    - **Partial**: Requires substantial, good-faith effort. The bulk of the required elements must be substantively present; only specific or bounded items are missing (example: a specific budget figure is absent while activities, objectives, and indicators are fully developed).
                    - **NOT Partial — downgrade to No — when:** the core specific subject of the question (e.g., "personas con discapacidad", "género", "pueblos indígenas") is completely absent or has only token/passing mentions. Do NOT grant Partial simply because other general project details (general objectives, other vulnerable groups) are present. Total absence of the precise target subject means it is a CRITICAL FAILURE = No or Not Found.
                    - **No**: Token/minimal coverage, isolated mentions, or the bulk of what the question asks is missing even when something is mentioned.
                    - **Not Found**: Nothing relevant; scope substituted; SCOPE LOCK fails.
                    - If in doubt between Yes and Partial → Partial. If in doubt between Partial and No → No.

                    No preamble, no filler. No JSON, no extra formatting."""

            user_content = f"""Question: {question}

Document's Answer: {answer}

Document's Reasoning: {reasoning}

Document's Evidence: {evidence}

Full Document Context (complete):
{doc_context}

Provide a critical assessment: Is this answer truly adequate from an expert perspective, considering the quality of the reasoning, evidence, and complete document context provided?"""

        # Single API call for critical evaluation - focuses on answer adequacy
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_content
                },
                {
                    "role": "user",
                    "content": user_content
                }
            ],
            max_completion_tokens=400,
            reasoning_effort="minimal"
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return "No se generó evaluación crítica."

        return content.strip()

    except Exception as e:
        return f"Error en evaluación crítica: {str(e)}"

def extract_section_number(question_text):
    """Extract section number from question text (e.g., '1.1 ¿Pregunta?' -> 1)"""
    import re
    match = re.match(r'(\d+)\.', str(question_text).strip())
    return int(match.group(1)) if match else None

def extract_subsection_number(question_text):
    """Extract full section.subsection from question text (e.g., '1.1 ¿Pregunta?' -> '1.1', '2.3 ¿Pregunta?' -> '2.3')"""
    import re
    match = re.match(r'(\d+\.\d+)', str(question_text).strip())
    return match.group(1) if match else None

def parse_subsection_for_sorting(subsection_str):
    """Convert subsection string to tuple for proper sorting (e.g., '1.1' -> (1, 1), '2.10' -> (2, 10))"""
    try:
        parts = str(subsection_str).split('.')
        return (int(parts[0]), int(parts[1]))
    except:
        return (999, 999)

def parse_question_sort_key(question_text):
    """Extract the full numeric prefix from question text as a tuple for fine-grained sorting.

    Examples: '1.1 ¿...?' -> (1, 1); '1.1.2 ¿...?' -> (1, 1, 2); '2.10 ¿...?' -> (2, 10).
    Unlike parse_subsection_for_sorting which caps at two levels, this preserves every
    numeric segment so nested numbering (1.1, 1.1.1, 1.1.2, 1.2) sorts naturally and
    does not collapse to a single subsection bucket.
    """
    import re
    match = re.match(r'^\s*(\d+(?:\.\d+)*)', str(question_text).strip())
    if not match:
        return (999, 999, 999)
    try:
        return tuple(int(p) for p in match.group(1).split('.'))
    except ValueError:
        return (999, 999, 999)

def synthesize_subsection_analysis(subsection_id, subsection_questions_df):
    """Synthesize subsection-level analysis from individual question answers within that subsection"""
    try:
        # Build context from individual question Q&A pairs
        qa_context = "\n\n".join([
            f"Pregunta {row['Pregunta']}:\nRespuesta: {row['Respuesta']}\nRazonamiento: {row['Razonamiento']}"
            for _, row in subsection_questions_df.iterrows()
        ])
        
        # Single efficient LLM call per subsection
        prompt = f"""Based on the following individual question answers from subsection {subsection_id} of a document evaluation, 
synthesize a comprehensive subsection-level analysis.

Subsection {subsection_id} - Individual Q&A:
{qa_context}

Provide a concise subsection-level analysis (1-2 paragraphs) that:
1. Integrates findings across all questions in this subsection
2. Identifies key strengths and gaps
3. Provides a clear assessment

Format as JSON with exactly this structure:
{{"subsection_analysis": "your analysis here (1-2 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert document analyst. Synthesize individual question findings into clear subsection-level insights. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=1500,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('subsection_analysis', 'Error en síntesis')
    
    except Exception as e:
        return f"Error generando análisis de subsección: {str(e)}"

def synthesize_section_analysis(section_num, subsection_analyses_dict):
    """Synthesize section-level analysis from subsection analyses"""
    try:
        # Build context from subsection analyses
        subsection_context = "\n\n".join([
            f"Subsección {subsec_id}:\n{analysis}"
            for subsec_id, analysis in sorted(subsection_analyses_dict.items(), key=lambda x: parse_subsection_for_sorting(x[0]))
        ])
        
        # Single efficient LLM call per section
        prompt = f"""Based on the following subsection analyses from section {section_num} of a document evaluation, 
synthesize a comprehensive section-level analysis.

Section {section_num} - Subsection Analyses:
{subsection_context}

Provide a detailed section-level analysis (2-3 paragraphs) that:
1. Integrates key findings across all subsections
2. Identifies overarching patterns, strengths, and gaps
3. Provides strategic, actionable insights for improvement

Format as JSON with exactly this structure:
{{"section_analysis": "your detailed analysis here (2-3 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert document analyst. Synthesize subsection findings into clear, actionable section-level insights. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=2000,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('section_analysis', 'Error en síntesis')
    
    except Exception as e:
        return f"Error generando análisis de sección: {str(e)}"

def synthesize_critical_evaluation_subsection(subsection_id, critical_opinions_df):
    """Synthesize critical evaluations at subsection level from individual question critical opinions"""
    try:
        # Build context from individual critical opinions
        critical_context = "\n\n".join([
            f"Pregunta {row['Pregunta']}:\nEvaluación Crítica: {row['Evaluación Crítica']}"
            for _, row in critical_opinions_df.iterrows()
        ])
        
        # Single efficient LLM call per subsection
        prompt = f"""Based on the following individual critical evaluations from subsection {subsection_id}, 
synthesize a comprehensive subsection-level critical assessment.

Subsection {subsection_id} - Individual Critical Evaluations:
{critical_context}

Provide a concise subsection-level critical assessment (1-2 paragraphs) that:
1. Integrates critical findings across all questions in this subsection
2. Identifies major gaps, risks, and inadequacies
3. Highlights patterns in insufficient responses
4. Provides clear recommendations for improvement

Format as JSON with exactly this structure:
{{"critical_evaluation": "your critical assessment here (1-2 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert critic of project quality. Synthesize individual critical opinions into clear subsection-level critical assessment. Be direct about inadequacies. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=1500,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('critical_evaluation', 'Error en síntesis crítica')
    
    except Exception as e:
        return f"Error generando evaluación crítica de subsección: {str(e)}"

def synthesize_critical_evaluation_section(section_num, critical_subsection_dict):
    """Synthesize critical evaluations at section level from subsection critical assessments"""
    try:
        # Build context from subsection critical evaluations
        critical_context = "\n\n".join([
            f"Subsección {subsec_id}:\n{evaluation}"
            for subsec_id, evaluation in sorted(critical_subsection_dict.items(), key=lambda x: parse_subsection_for_sorting(x[0]))
        ])
        
        # Single efficient LLM call per section
        prompt = f"""Based on the following subsection critical evaluations from section {section_num}, 
synthesize a comprehensive section-level critical assessment.

Section {section_num} - Subsection Critical Evaluations:
{critical_context}

Provide a detailed section-level critical assessment (2-3 paragraphs) that:
1. Integrates critical findings across all subsections
2. Identifies overarching gaps, risks, and systemic inadequacies
3. Highlights critical patterns of insufficient responses
4. Provides strategic recommendations for improving the overall section

Format as JSON with exactly this structure:
{{"critical_evaluation": "your critical assessment here (2-3 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert critic of project quality. Synthesize subsection critical findings into clear, actionable section-level critical assessment. Be direct and strategic. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=2000,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('critical_evaluation', 'Error en síntesis crítica')
    
    except Exception as e:
        return f"Error generando evaluación crítica de sección: {str(e)}"

def create_results_download_with_sections(results_df, subsection_analyses, subsection_critical_evaluations, section_analyses, section_critical_evaluations, filename_base="appraisal_checklist"):
    """Create ZIP file with results including three separate sheets for questions, subsections, and sections"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        excel_buffer = io.BytesIO()
        
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
            workbook = writer.book
            
            # Format definitions
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#002F6C',
                'font_color': 'white',
                'border': 1,
                'valign': 'vcenter',
                'text_wrap': True
            })
            section_header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#0072CE',
                'font_color': 'white',
                'border': 1,
                'valign': 'vcenter',
                'text_wrap': True,
                'font_size': 12
            })
            subsection_header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#4A90E2',
                'font_color': 'white',
                'border': 1,
                'valign': 'vcenter',
                'text_wrap': True,
                'font_size': 11
            })
            merged_format = workbook.add_format({
                'border': 1,
                'valign': 'top',
                'text_wrap': True,
                'bg_color': '#F5F5F5'
            })
            subsection_merged_format = workbook.add_format({
                'border': 1,
                'valign': 'top',
                'text_wrap': True,
                'bg_color': '#E8F4F8',
                'italic': True
            })
            normal_format = workbook.add_format({
                'border': 1,
                'valign': 'top',
                'text_wrap': True
            })
            
            # Extract and ensure sorting columns exist
            if '_section' not in results_df.columns:
                results_df['_section'] = results_df['Pregunta'].apply(extract_section_number)
            if '_subsection' not in results_df.columns:
                results_df['_subsection'] = results_df['Pregunta'].apply(extract_subsection_number)
            if '_sort_key' not in results_df.columns:
                results_df['_sort_key'] = results_df['_subsection'].apply(parse_subsection_for_sorting)
            
            # Always re-derive _sort_key from the question text so persisted session-state
            # DataFrames (which may have been built with the old two-level sort key) are
            # re-sorted by the full numeric prefix.
            results_df['_sort_key'] = results_df['Pregunta'].apply(parse_question_sort_key)

            # Primary sort by full numeric prefix; _orig_idx breaks ties if available.
            sort_cols = ['_sort_key']
            if '_orig_idx' in results_df.columns:
                sort_cols.append('_orig_idx')
            results_df_sorted = results_df.sort_values(by=sort_cols).reset_index(drop=True)
            
            # ===== SHEET 1: Questions (Preguntas) =====
            sheet_questions = workbook.add_worksheet('1. Preguntas')
            questions_data = results_df_sorted[['_subsection', 'Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia', 'Status']].copy()
            questions_data.columns = ['Subsección', 'Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia', 'Status']
            questions_data.to_excel(writer, index=False, sheet_name='1. Preguntas', startrow=0)

            # Format headers
            for col_num, value in enumerate(questions_data.columns.values):
                writer.sheets['1. Preguntas'].write(0, col_num, value, header_format)

            # Set column widths
            writer.sheets['1. Preguntas'].set_column('A:A', 12)  # Subsección
            writer.sheets['1. Preguntas'].set_column('B:B', 35)  # Pregunta
            writer.sheets['1. Preguntas'].set_column('C:C', 12)  # Respuesta
            writer.sheets['1. Preguntas'].set_column('D:D', 55)  # Razonamiento (includes critical evaluation)
            writer.sheets['1. Preguntas'].set_column('E:E', 40)  # Evidencia
            writer.sheets['1. Preguntas'].set_column('F:F', 10)  # Status
            
            # ===== SHEET 2: Subsection Analysis (Análisis por Subsección) =====
            sheet_subsections = workbook.add_worksheet('2. Análisis Subsecciones')
            
            # Get unique subsections in sorted order
            subsections_sorted = sorted(results_df_sorted['_subsection'].dropna().unique(), key=parse_subsection_for_sorting)
            
            # Write headers
            subsec_headers = ['Subsección', 'Sección', 'Análisis de Subsección']
            for col_num, header in enumerate(subsec_headers):
                sheet_subsections.write(0, col_num, header, header_format)
            
            current_row = 1
            for subsection_id in subsections_sorted:
                section_num = extract_section_number(subsection_id)
                analysis_text = subsection_analyses.get(subsection_id, "No se generó análisis")
                critical_text = subsection_critical_evaluations.get(subsection_id, "No se generó evaluación crítica")
                
                sheet_subsections.write(current_row, 0, subsection_id, normal_format)
                sheet_subsections.write(current_row, 1, f"Sección {section_num}", normal_format)
                sheet_subsections.write(current_row, 2, analysis_text, subsection_merged_format)
                sheet_subsections.write(current_row, 3, critical_text, subsection_merged_format)
                current_row += 1
            
            # Set column widths
            sheet_subsections.set_column('A:A', 12)  # Subsección
            sheet_subsections.set_column('B:B', 15)  # Sección
            sheet_subsections.set_column('C:C', 80)  # Análisis
            sheet_subsections.set_column('D:D', 80)  # Evaluación Crítica
            
            # ===== SHEET 3: Section Analysis (Análisis por Sección) =====
            sheet_sections = workbook.add_worksheet('3. Análisis Secciones')
            
            # Get unique sections in sorted order
            sections_sorted = sorted(results_df_sorted['_section'].dropna().unique())
            
            # Write headers
            section_headers = ['Sección', 'Análisis de Sección', 'Evaluación Crítica']
            for col_num, header in enumerate(section_headers):
                sheet_sections.write(0, col_num, header, header_format)
            
            current_row = 1
            for section_num in sections_sorted:
                section_num = int(section_num)
                analysis_text = section_analyses.get(section_num, "No se generó análisis")
                critical_text = section_critical_evaluations.get(section_num, "No se generó evaluación crítica")
                
                sheet_sections.write(current_row, 0, f"Sección {section_num}", section_header_format)
                sheet_sections.write(current_row, 1, analysis_text, merged_format)
                sheet_sections.write(current_row, 2, critical_text, merged_format)
                current_row += 1
            
            # Set column widths
            sheet_sections.set_column('A:A', 15)  # Sección
            sheet_sections.set_column('B:B', 90)  # Análisis
            sheet_sections.set_column('C:C', 90)  # Evaluación Crítica
        
        excel_buffer.seek(0)
        zipf.writestr(f"{filename_base}_results.xlsx", excel_buffer.getvalue())

        # Add original template if available
        try:
            with open('./Appraisal Checklist_2025 es-419.xlsx', 'rb') as f:
                template_data = f.read()
                zipf.writestr(f"{filename_base}_rubric_template.xlsx", template_data)
        except FileNotFoundError:
            pass
        
        # Add summary report
        summary = f"""
Resumen del Análisis de la Lista de la Valoración Preliminar de la Calidad
===========================================================

Archivo Excel generado con 3 hojas:
  1. Preguntas: Análisis individual de cada pregunta (ordenado por subsección)
  2. Análisis Subsecciones: Síntesis de preguntas agrupadas por subsección
  3. Análisis Secciones: Síntesis ejecutiva por sección

Total de preguntas analizadas: {len(results_df)}
Análisis exitosos: {len(results_df[results_df['Status'] == 'Success'])}
Análisis fallidos: {len(results_df[results_df['Status'] == 'Error'])}

Subsecciones analizadas: {len(subsection_analyses)}
Secciones analizadas: {len(section_analyses)}

Distribución de respuestas:
{results_df['Respuesta'].value_counts().to_string()}
        """
        zipf.writestr(f"{filename_base}_summary.txt", summary)
    
    zip_buffer.seek(0)
    return zip_buffer

with tab1:
    st.header("📋 Valoración Preliminar de Calidad de Proyectos (Preliminary Project Quality Appraisal)")
    
    # Load questions
    df_appraisal, error_msg = load_appraisal_questions()
    
    if error_msg:
        st.error(error_msg)
        st.stop()
    
    # Download button for the rubric file (directly on page, no expander)
    # Download button for the rubric file (directly on page, no expander)
    try:
        with open('./Appraisal Checklist_2025 es-419.xlsx', 'rb') as f:
            st.download_button(
                label="📥 Descargar archivo rúbrica de Valoración preliminar",
                data=f,
                file_name="Appraisal Checklist_2025 es-419.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_appraisal_rubric"
            )
    except FileNotFoundError:
        st.warning("Archivo de rúbrica no disponible para descarga.")
    
    # Instrucciones
    st.markdown("""
    ### Instrucciones
    
    1. **Subir documento**: Selecciona un Documento de Diseño de Proyecto en archivo de formato DOCX para el análisis de valoración preliminar de calidad.
    2. **Procesar**: Haz clic en 'Analizar documento' para iniciar la evaluación
    3. **Revisar resultados**: Examina los resultados del análisis en la tabla interactiva
    4. **Descargar**: Obtén todos los resultados y las pruebas en un archivo ZIP
    
    ---
    """)
    
  
    # Initialize session state for Tab 1 results persistence
    if 'tab1_results_df' not in st.session_state:
        st.session_state['tab1_results_df'] = None
    if 'tab1_doc_stats' not in st.session_state:
        st.session_state['tab1_doc_stats'] = None

    # Initialize session state for Tab 1
    if 'document_extracted_tab1' not in st.session_state:
        st.session_state['document_extracted_tab1'] = False
    if 'selected_sections_tab1' not in st.session_state:
        st.session_state['selected_sections_tab1'] = []
    
    # Document upload
    st.subheader("📄 Carga de documento")
    
    # Warning box about document requirements
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos:**
    
    **📝 Formato del documento:**
    - Solo se aceptan archivos en formato **.docx** (Word 2007 o posterior)
    - El documento debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del documento deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente y las secciones no se identificarán.
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos que excedan este límite serán truncados automáticamente
    - Se recomienda dividir documentos muy extensos (más de ~180 páginas) en secciones más pequeñas si es necesario
    
    **✅ Mejores prácticas:**
    - Usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones principales
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que el documento esté guardado correctamente antes de subirlo
    """)
    
    uploaded_file = st.file_uploader(
        "Sube un DOCX para la evaluación:",
        type=["docx"],
        key="appraisal_file_uploader",
        help="Selecciona un documento de Word (.docx) para el análisis de valoración preliminar de calidad"
    )
    
    # Document Extraction Section
    st.markdown("---")
    st.markdown("### 📥 Extracción de Documento")
    
    if uploaded_file is not None:
        file_hash = hash(uploaded_file.getvalue())
        file_changed = st.session_state.get('last_file_hash_tab1') != file_hash
        
        if file_changed:
            st.session_state['document_extracted_tab1'] = False
            st.session_state['last_file_hash_tab1'] = None
        
        if st.button("🔍 Extraer Documento", key="extract_document_tab1", type="primary"):
            if uploaded_file is None:
                st.error("Por favor suba un archivo DOCX primero.")
                st.stop()
            
            with st.spinner("Extrayendo documento..."):
                try:
                    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp_file.write(uploaded_file.read())
                    tmp_file.close()
                    
                    progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                    doc_result = docx2python(tmp_file.name)
                    
                    # Use enhanced extraction
                    df, tables_data, extraction_stats = extract_docx_structure_enhanced(tmp_file.name)
                    progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")
                    
                    # Extract sections
                    header_1_values = df['header_1'].dropna().unique()
                    llm_summary_rows = []
                    
                    for idx, header in enumerate(header_1_values):
                        section_df = df[df['header_1'] == header].copy()
                        full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                        section_words = len(full_text.split())
                        section_paras = len(section_df[section_df['source_type'] == 'paragraph'])
                        section_tables = len(section_df[section_df['source_type'] == 'table'])
                        
                        llm_summary_rows.append({
                            'header_1': header,
                            'llm_paragraph': full_text if full_text else "",
                            'n_words': section_words,
                            'n_paragraphs': section_paras,
                            'n_tables': section_tables
                        })

                    progress_bar.progress(0.5, text="Secciones extraídas.")
                    
                    # Create exploded dataframe
                    llm_summary_df = pd.DataFrame(llm_summary_rows)
                    exploded_df = llm_summary_df.assign(
                        llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                    ).explode('llm_paragraph')
                    exploded_df = exploded_df.reset_index(drop=True)
                    exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                    
                    # Get full text
                    full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                    
                    # Store in session state
                    file_size = os.path.getsize(tmp_file.name)
                    n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                    n_paragraphs = len(exploded_df)
                    
                    st.session_state['full_document_text_tab1'] = full_document_text
                    st.session_state['appraisal_document_stats'] = {
                        'file_size': file_size,
                        'word_count': n_words,
                        'n_words': n_words
                    }
                    st.session_state['exploded_df_tab1'] = exploded_df
                    st.session_state['extraction_df_tab1'] = df
                    st.session_state['tables_data_tab1'] = tables_data
                    st.session_state['extraction_stats_tab1'] = extraction_stats
                    st.session_state['sections_df_tab1'] = llm_summary_df
                    st.session_state['selected_sections_tab1'] = list(header_1_values)  # Select all by default
                    st.session_state['last_file_hash_tab1'] = file_hash
                    st.session_state['document_extracted_tab1'] = True
                    
                    try:
                        os.unlink(tmp_file.name)
                    except:
                        pass
                    
                    progress_bar.progress(1.0, text="Extracción completa.")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error procesando el documento: {e}")
                    import traceback
                    st.error(traceback.format_exc())
                    st.stop()
        
        # Show extraction results if document is extracted
        if st.session_state.get('document_extracted_tab1', False) and not file_changed:
            st.success("✅ Documento extraído con éxito")
            
            # Download button for extracted document structure
            extraction_df = st.session_state.get('extraction_df_tab1', pd.DataFrame())
            if not extraction_df.empty:
                excel_data = to_excel(extraction_df)
                # Get filename from extraction_df or use default
                filename_base = extraction_df['filename'].iloc[0] if 'filename' in extraction_df.columns and not extraction_df['filename'].empty else "documento"
                filename_base = filename_base.replace('.docx', '').replace('.doc', '')
                st.download_button(
                    label="📥 Descargar estructura extraída del documento (Excel)",
                    data=excel_data,
                    file_name=f"estructura_documento_tab1_{filename_base}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_extraction_tab1"
                )
                st.caption("El archivo incluye todas las columnas de encabezados (header_1 a header_6), contenido, tipo de fuente, y metadatos de extracción.")
            
            # Display header_1 sections and their content
            extraction_df = st.session_state.get('extraction_df_tab1', pd.DataFrame())
            if not extraction_df.empty:
                with st.expander("📋 Ver estructura extraída del documento (encabezados nivel 1 y contenido)", expanded=False):
                    st.markdown("**Estructura del documento extraído (solo encabezados nivel 1):**")
                    
                    # Get unique header_1 values
                    header_1_sections = extraction_df[extraction_df['header_1'].notna() & (extraction_df['header_1'] != '')]['header_1'].unique()
                    
                    for h1 in header_1_sections:
                        st.markdown(f"### {h1}")
                        
                        # Get all content for this header_1 section
                        section_df = extraction_df[extraction_df['header_1'] == h1]
                        section_content = section_df[section_df['source_type'] == 'paragraph']['content'].tolist()
                        
                        # Display content
                        if section_content:
                            full_text = '\n\n'.join([str(c) for c in section_content if pd.notna(c) and str(c).strip()])
                            if full_text.strip():
                                st.text(full_text)
                                st.caption(f"Total: {len(full_text):,} caracteres")
                        else:
                            st.info("Esta sección no tiene contenido de párrafos extraído.")
                        
                        st.markdown("---")
            
            sections_df = st.session_state.get('sections_df_tab1', pd.DataFrame())
            
            if not sections_df.empty:
                header_1_values = sections_df['header_1'].tolist()
                
                # Section selector
                st.markdown("### 🔍 Selección de Secciones para Evaluación")
                st.info("Selecciona las secciones que deseas incluir en la evaluación. Por defecto, todas las secciones están seleccionadas.")
                
                # Guidance about section extraction
                if len(header_1_values) == 0:
                    st.error("⚠️ **No se detectaron secciones en el documento.** Esto puede deberse a que el documento no usa estilos de encabezado de Word (Heading 1, Heading 2, etc.). Por favor, verifica que tu documento tenga encabezados formateados correctamente.")
                elif len(header_1_values) < 3:
                    st.warning("⚠️ **Se detectaron pocas secciones** en el documento. Si esperabas más secciones, verifica que el documento use estilos de encabezado de Word (Heading 1, Heading 2, etc.) para identificar las secciones principales.")
                
                # Initialize selected sections if not exists
                if 'selected_sections_tab1' not in st.session_state:
                    st.session_state['selected_sections_tab1'] = list(header_1_values)
                
                # Section selection interface
                selected_sections = st.session_state.get('selected_sections_tab1', list(header_1_values)).copy()
                col1, col2 = st.columns([3, 1])
                
                with col2:
                    if st.button("✅ Seleccionar Todas", key="select_all_sections_tab1"):
                        st.session_state['selected_sections_tab1'] = list(header_1_values)
                        st.rerun()
                    
                    if st.button("❌ Deseleccionar Todas", key="deselect_all_sections_tab1"):
                        st.session_state['selected_sections_tab1'] = []
                        st.rerun()
                
                with col1:
                    st.markdown("**Secciones disponibles:**")
                    for section in header_1_values:
                        section_info = sections_df[sections_df['header_1'] == section].iloc[0]
                        is_selected = section in selected_sections
                        
                        checkbox_label = f"**{section}** ({section_info['n_words']:,} palabras, {section_info['n_paragraphs']} párrafos)"
                        
                        checkbox_key = f"section_checkbox_{section}_tab1"
                        new_selection = st.checkbox(checkbox_label, value=is_selected, key=checkbox_key)
                        
                        if new_selection and section not in selected_sections:
                            selected_sections.append(section)
                        elif not new_selection and section in selected_sections:
                            selected_sections.remove(section)
                        
                        # Add expandable preview of extracted content
                        with st.expander(f"👁️ Ver contenido: {section}", expanded=False):
                            section_content = section_info['llm_paragraph']
                            if section_content and section_content.strip():
                                st.text_area(
                                    "Contenido extraído:",
                                    value=section_content,
                                    height=200,
                                    key=f"content_preview_{section}_tab1",
                                    label_visibility="collapsed"
                                )
                                st.caption(f"Total: {len(section_content):,} caracteres")
                            else:
                                st.info("Esta sección no tiene contenido extraído.")
                            
                            # Show table-extracted text
                            tables_data = st.session_state.get('tables_data_tab1', [])
                            section_tables = [t for t in tables_data if t.get('section') == section]
                            
                            if section_tables:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                for table_info in section_tables:
                                    table_num = table_info.get('table_number', 'N/A')
                                    table_data = table_info.get('data', [])
                                    if table_data:
                                        # Format table as text
                                        table_text = '\n'.join([' | '.join(str(cell) for cell in row) for row in table_data])
                                        st.text_area(
                                            f"Tabla {table_num}:",
                                            value=table_text,
                                            height=150,
                                            key=f"table_preview_{section}_table{table_num}_tab1",
                                            label_visibility="collapsed"
                                        )
                                        st.caption(f"Tabla {table_num}: {len(table_data)} filas, {len(table_data[0]) if table_data else 0} columnas")
                            else:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                st.info("No se encontraron tablas en esta sección.")
                
                # Update session state
                st.session_state['selected_sections_tab1'] = selected_sections
                
                # Show selection summary
                if selected_sections:
                    selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
                    total_selected_words = selected_df['n_words'].sum()
                    total_selected_paras = selected_df['n_paragraphs'].sum()
                    
                    # Estimate tokens
                    if encoding:
                        selected_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
                        estimated_tokens = len(encoding.encode(selected_text))
                    else:
                        estimated_tokens = total_selected_words * 1.2
                    
                    st.success(f"✅ {len(selected_sections)} secciones seleccionadas | "
                              f"{total_selected_words:,} palabras | "
                              f"~{estimated_tokens:,} tokens estimados")
    
    # Filter section for questions (optional, for filtering which questions to analyze)
    st.markdown("---")
    st.subheader("🔍 Filtros de Preguntas (opcional)")
    st.info("Selecciona secciones y subsecciones de preguntas para un análisis enfocado. Deja en blanco para analizar todas las preguntas.")
    
    # Get unique sections and subsections from questions
    df_with_sections = df_appraisal.copy()
    df_with_sections['_section'] = df_with_sections['Pregunta_Realizada'].apply(extract_section_number)
    df_with_sections['_subsection'] = df_with_sections['Pregunta_Realizada'].apply(extract_subsection_number)
    
    # Section Names Mapping
    SECTION_NAMES = {
        1: "Pertinencia",
        2: "Apropiación y sostenibilidad",
        3: "Gestión orientada a los resultados",
        4: "Transparencia y rendición de cuentas",
        5: "Presentación de la propuesta"
    }

    all_sections = sorted(df_with_sections['_section'].dropna().unique())
    
    # Create two columns for filters
    col_filter1, col_filter2 = st.columns(2)
    
    with col_filter1:
        # Format section options with names
        section_options = [f"{int(s)}. {SECTION_NAMES.get(int(s), 'Sección ' + str(int(s)))}" for s in all_sections]
        
        selected_sections_filter = st.multiselect(
            "Selecciona Secciones:",
            options=section_options,
            key="filter_sections_tab1",
            help="Selecciona una o más secciones para analizar"
        )
    
    # Extract section numbers from filter selections
    filtered_section_nums = []
    if selected_sections_filter:
        for s in selected_sections_filter:
            try:
                # Extract number from "1. Name" format
                filtered_section_nums.append(int(s.split('.')[0]))
            except:
                pass

    # Filter available subsections based on selected sections
    if filtered_section_nums:
        # Show only subsections belonging to selected sections
        available_subsections_df = df_with_sections[df_with_sections['_section'].isin(filtered_section_nums)]
    else:
        # Show all subsections if no section is selected
        available_subsections_df = df_with_sections
        
    all_subsections = sorted(available_subsections_df['_subsection'].dropna().unique(), 
                            key=lambda x: parse_subsection_for_sorting(x) if pd.notna(x) else (999, 999))

    with col_filter2:
        selected_subsections_filter = st.multiselect(
            "Selecciona Subsecciones:",
            options=[f"Subsección {s}" for s in all_subsections],
            key="filter_subsections_tab1",
            help="Selecciona una o más subsecciones para analizar"
        )
    
    # Extract subsection IDs
    filtered_subsection_ids = [s.replace("Subsección ", "") for s in selected_subsections_filter] if selected_subsections_filter else None
    
    # Processing button
    st.markdown("---")
    st.markdown("### ⚙️ Procesamiento y Evaluación")
    
    # Warning about AI results verification
    st.warning("""
    **⚠️ Importante - Verificación de Resultados:**
    
    Los resultados generados por esta herramienta utilizan inteligencia artificial y deben ser **verificados y corroborados** antes de su uso.
    
    - La IA puede cometer errores, interpretaciones incorrectas o pasar por alto información relevante
    - Los análisis y puntuaciones son **sugerencias** basadas en el contenido del documento, no son definitivos
    - Se recomienda revisar manualmente las evidencias citadas y validar las conclusiones
    - Los resultados deben ser contrastados con conocimiento experto y documentación adicional cuando sea necesario
    
    Esta herramienta es un **asistente de análisis** que facilita la revisión, pero la responsabilidad final de la evaluación recae en el usuario.
    """)
    
    if st.button('🔍 Analizar documento', key="appraisal_process_button", type="primary"):
        # Check prerequisites
        if not st.session_state.get('document_extracted_tab1', False):
            st.error("❌ Por favor extrae el documento primero usando el botón 'Extraer Documento'.")
            st.stop()
        
        if uploaded_file is None:
            st.warning("⚠️ Por favor suba un archivo DOCX primero.")
            st.stop()
        
        # Get selected sections or use full document
        selected_sections = st.session_state.get('selected_sections_tab1', [])
        sections_df = st.session_state.get('sections_df_tab1', pd.DataFrame())
        
        if selected_sections and not sections_df.empty:
            # Filter to selected sections only
            selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
            document_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
            st.info(f"📌 Evaluando {len(selected_sections)} secciones seleccionadas")
        else:
            # Fallback to full document
            document_text = st.session_state.get('full_document_text_tab1', '')
            if not selected_sections:
                st.warning("⚠️ No hay secciones seleccionadas. Usando documento completo.")
        
        if not document_text:
            st.error("No se pudo recuperar el texto del documento.")
            st.stop()
        
        # Get questions for analysis
        questions = df_appraisal['Pregunta_Realizada'].dropna().unique().tolist()
        
        # Apply filters if selected
        if filtered_section_nums or filtered_subsection_ids:
            questions_df_temp = df_appraisal[df_appraisal['Pregunta_Realizada'].isin(questions)].copy()
            questions_df_temp['_section'] = questions_df_temp['Pregunta_Realizada'].apply(extract_section_number)
            questions_df_temp['_subsection'] = questions_df_temp['Pregunta_Realizada'].apply(extract_subsection_number)
            
            # Filter by selected sections and/or subsections
            if filtered_section_nums and filtered_subsection_ids:
                # Both filters selected - apply AND logic
                filtered_questions = questions_df_temp[
                    (questions_df_temp['_section'].isin(filtered_section_nums)) |
                    (questions_df_temp['_subsection'].isin(filtered_subsection_ids))
                ]
            elif filtered_section_nums:
                # Only sections selected
                filtered_questions = questions_df_temp[questions_df_temp['_section'].isin(filtered_section_nums)]
            else:
                # Only subsections selected
                filtered_questions = questions_df_temp[questions_df_temp['_subsection'].isin(filtered_subsection_ids)]
            
            questions = filtered_questions['Pregunta_Realizada'].dropna().unique().tolist()
            
            # Show filter info
            st.info(f"📌 Análisis filtrado: {len(questions)} preguntas seleccionadas de {len(df_appraisal)}")
        
        if not questions:
            st.error("❌ No se encontraron preguntas para el análisis con los filtros aplicados.")
            st.stop()
        
        # Analyze questions
        st.markdown("### 🔍 Progreso del análisis")
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        results = []
        critical_opinions = {}
        
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            # Submit all questions for document-based analysis first
            # TAB1: Using tab1-specific function that explicitly states when 2 parts are detected
            # Tag each future with its original index so within-subsection ordering survives
            # the parallel as_completed dispatch.
            future_to_question_doc = {
                executor.submit(analyze_question_with_llm_tab1, q, document_text): (idx, q)
                for idx, q in enumerate(questions)
            }

            # Process document-based analyses
            completed = 0
            for future in as_completed(future_to_question_doc):
                idx, question = future_to_question_doc[future]
                result = future.result()
                result['_orig_idx'] = idx
                results.append(result)
                completed += 1

                # Update progress
                progress = completed / len(questions)
                progress_bar.progress(progress * 0.5)  # First half of progress bar
                status_text.text(f"Análisis documentales: {completed}/{len(questions)} preguntas")

            # Create results DataFrame to get answers for critical evaluation
            results_df_temp = pd.DataFrame(results)

            # Now submit critical evaluations with answer, reasoning, evidence AND document context
            # TAB1: Using tab1-specific function that checks for proper two-part acknowledgment
            future_to_question_critical = {
                executor.submit(
                    analyze_question_with_critical_opinion_tab1,
                    row['Pregunta'],
                    row['Respuesta'],
                    row['Razonamiento'],
                    row['Evidencia'],
                    document_text  # Pass selected sections text for context
                ): row['Pregunta']
                for _, row in results_df_temp.iterrows()
            }
            
            # Process critical evaluations
            completed = 0
            for future in as_completed(future_to_question_critical):
                question = future_to_question_critical[future]
                critical_opinion = future.result()
                critical_opinions[question] = critical_opinion
                completed += 1
                
                # Update progress
                progress = completed / len(questions)
                progress_bar.progress(0.5 + progress * 0.5)  # Second half of progress bar
                status_text.text(f"Evaluaciones críticas: {completed}/{len(questions)} preguntas")

        # Apply critic verdict. The critic is authoritative and its body becomes the single
        # coherent Razonamiento — whether it overrides the answer or confirms it. This gives the
        # reader one reasoning column instead of an original reasoning plus a separate critique.
        # Respuesta Original preserves the pre-critic grade for audit.
        override_count = 0
        unparsed_count = 0
        for result in results:
            q = result.get('Pregunta')
            critic_raw = critical_opinions.get(q, "") or ""
            verdict, critic_body = parse_critic_verdict(critic_raw)
            original_answer = result.get('Respuesta', '')
            result['Respuesta Original'] = original_answer

            if verdict is None:
                unparsed_count += 1
            elif verdict != 'Keep' and verdict != original_answer:
                result['Respuesta'] = verdict
                override_count += 1

            if critic_body:
                result['Razonamiento'] = critic_body

            critical_opinions[q] = critic_body if critic_body else critic_raw

        if override_count:
            st.info(f"🔄 {override_count} respuesta(s) ajustada(s) por la evaluación crítica.")
        if unparsed_count:
            st.warning(
                f"⚠️ {unparsed_count} evaluación(es) crítica(s) sin veredicto parseable; "
                "se mantuvo la respuesta original."
            )

        # Create results DataFrame - sort by subsection for proper ordering
        results_df = pd.DataFrame(results)

        # Add critical opinions to dataframe
        results_df['Evaluación Crítica'] = results_df['Pregunta'].map(critical_opinions).fillna("No disponible")
        results_df['_section_num'] = results_df['Pregunta'].apply(extract_section_number)
        results_df['_subsection'] = results_df['Pregunta'].apply(extract_subsection_number)
        # Sort key uses the FULL numeric prefix (e.g., "1.1.2" -> (1,1,2)) so nested
        # numbering sorts naturally instead of collapsing to the two-level subsection.
        results_df['_sort_key'] = results_df['Pregunta'].apply(parse_question_sort_key)

        # Primary sort by full numeric prefix; _orig_idx breaks ties for questions that
        # share the same numeric prefix, keeping rubric order within that group.
        results_df = results_df.sort_values(by=['_sort_key', '_orig_idx']).reset_index(drop=True)
        
        # THREE-LEVEL SYNTHESIS: Question -> Subsection -> Section
        st.markdown("### 📈 Síntesis Multinivel")
        
        # Level 1: Subsection synthesis (grouping individual questions)
        st.info("Generando análisis a nivel de subsección...")
        subsections = sorted(results_df['_subsection'].dropna().unique(), key=parse_subsection_for_sorting)
        subsection_analyses = {}
        
        subsection_progress = st.progress(0)
        for idx, subsection_id in enumerate(subsections):
            subsection_df = results_df[results_df['_subsection'] == subsection_id].copy()
            
            # Generate subsection analysis from individual questions
            subsection_analysis = synthesize_subsection_analysis(
                subsection_id,
                subsection_df[['Pregunta', 'Respuesta', 'Razonamiento']]
            )
            subsection_analyses[subsection_id] = subsection_analysis
            
            subsection_progress.progress((idx + 1) / len(subsections))
        
        # Level 2: Critical evaluation synthesis at subsection level
        st.info("Generando evaluaciones críticas a nivel de subsección...")
        subsection_critical_evaluations = {}
        
        critical_subsection_progress = st.progress(0)
        for idx, subsection_id in enumerate(subsections):
            subsection_df = results_df[results_df['_subsection'] == subsection_id].copy()
            
            # Generate critical evaluation synthesis from individual critical opinions
            critical_evaluation = synthesize_critical_evaluation_subsection(
                subsection_id,
                subsection_df[['Pregunta', 'Evaluación Crítica']]
            )
            subsection_critical_evaluations[subsection_id] = critical_evaluation
            
            critical_subsection_progress.progress((idx + 1) / len(subsections))
        
        # Level 3: Section synthesis (grouping subsections)
        st.info("Generando análisis a nivel de sección...")
        sections = sorted(results_df['_section_num'].dropna().unique())
        section_analyses = {}
        
        section_progress = st.progress(0)
        for idx, section_num in enumerate(sections):
            section_num = int(section_num)
            # Get all subsections for this section
            section_subsections = {
                subsec_id: subsection_analyses[subsec_id]
                for subsec_id in subsections
                if subsec_id.startswith(f"{section_num}.")
            }
            
            # Generate section analysis from subsection analyses
            section_analysis = synthesize_section_analysis(
                section_num,
                section_subsections
            )
            section_analyses[section_num] = section_analysis
            
            section_progress.progress((idx + 1) / len(sections))
        
        # Level 4: Critical evaluation synthesis at section level
        st.info("Generando evaluaciones críticas a nivel de sección...")
        section_critical_evaluations = {}
        
        critical_section_progress = st.progress(0)
        for idx, section_num in enumerate(sections):
            section_num = int(section_num)
            # Get all critical subsection evaluations for this section
            section_critical_subsections = {
                subsec_id: subsection_critical_evaluations[subsec_id]
                for subsec_id in subsections
                if subsec_id.startswith(f"{section_num}.")
            }
            
            # Generate critical evaluation from critical subsection evaluations
            critical_evaluation = synthesize_critical_evaluation_section(
                section_num,
                section_critical_subsections
            )
            section_critical_evaluations[section_num] = critical_evaluation
            
            critical_section_progress.progress((idx + 1) / len(sections))
        
        # Store all results in session state
        st.session_state['tab1_results_df'] = results_df
        st.session_state['tab1_subsection_analyses'] = subsection_analyses
        st.session_state['tab1_subsection_critical_evaluations'] = subsection_critical_evaluations
        st.session_state['tab1_section_analyses'] = section_analyses
        st.session_state['tab1_section_critical_evaluations'] = section_critical_evaluations

        # Get document stats from session state
        doc_stats = st.session_state.get('appraisal_document_stats', {})
        st.session_state['tab1_doc_stats'] = {
            'file_size': doc_stats.get('file_size', 0),
            'word_count': doc_stats.get('word_count', 0)
        }
        
        # Display results
        st.markdown("### 📊 Resultados del análisis")
        
        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total de preguntas", len(results_df))
        with col2:
            success_count = len(results_df[results_df['Status'] == 'Success'])
            st.metric("Exitosas", success_count)
        with col3:
            error_count = len(results_df[results_df['Status'] == 'Error'])
            st.metric("Errores", error_count)
        with col4:
            if success_count > 0:
                yes_count = len(results_df[results_df['Respuesta'] == 'Yes'])
                st.metric("Respuestas 'Sí'", yes_count)
        
        # Results table
        st.markdown("#### 📋 Resultados detallados")
        
        # Filter options
        col1, col2 = st.columns([1, 1])
        with col1:
            response_filter = st.selectbox(
                "Filtrar por respuesta:",
                ['Todas'] + results_df['Respuesta'].unique().tolist(),
                key="response_filter"
            )
        
        # Apply filter
        filtered_df = results_df.copy()
        if response_filter != 'Todas':
            filtered_df = filtered_df[filtered_df['Respuesta'] == response_filter]
        
        # Display filtered results
        st.dataframe(
            filtered_df[['Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia']], 
            use_container_width=True,
            height=400
        )
        
        # Download section
        st.markdown("### 📥 Descargar resultados")
        
        if len(results_df) > 0:
            # Get analyses from session state
            subsection_analyses = st.session_state.get('tab1_subsection_analyses', {})
            subsection_critical_evaluations = st.session_state.get('tab1_subsection_critical_evaluations', {})
            section_analyses = st.session_state.get('tab1_section_analyses', {})
            section_critical_evaluations = st.session_state.get('tab1_section_critical_evaluations', {})
            zip_buffer = create_results_download_with_sections(results_df, subsection_analyses, subsection_critical_evaluations, section_analyses, section_critical_evaluations)
            
            st.download_button(
                label="📦 Descargar resultados en ZIP",
                data=zip_buffer,
                file_name="appraisal_checklist_results.zip",
                mime="application/zip",
                key="appraisal_download_button"
            )
            
            st.success("✅ ¡Análisis completo! Usa el botón de descarga para obtener todos los resultados.")
        else:
            st.warning("⚠️ No hay resultados para descargar.")
    
    else:
        # Check if there are persisted results in session state
        if st.session_state.get('tab1_results_df') is not None:
            results_df = st.session_state['tab1_results_df']
            doc_stats = st.session_state.get('tab1_doc_stats', {})
            
            # Display persisted results
            st.markdown("### 📊 Resultados del análisis (guardados)")
            
            if doc_stats:
                st.info(f"""
                **Resumen del documento analizado:**
                - Tamaño del archivo: {doc_stats.get('file_size', 0)/1024:.2f} KB
                - Número de palabras: {doc_stats.get('word_count', 0):,}
                """)
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Total de preguntas", len(results_df))
            with col2:
                success_count = len(results_df[results_df['Status'] == 'Success'])
                st.metric("Exitosas", success_count)
            with col3:
                error_count = len(results_df[results_df['Status'] == 'Error'])
                st.metric("Errores", error_count)
            with col4:
                if success_count > 0:
                    yes_count = len(results_df[results_df['Respuesta'] == 'Yes'])
                    st.metric("Respuestas 'Sí'", yes_count)
            
            # Results table
            st.markdown("#### 📋 Resultados detallados")
            
            # Filter options
            col1, col2 = st.columns([1, 1])
            with col1:
                response_filter = st.selectbox(
                    "Filtrar por respuesta:",
                    ['Todas'] + results_df['Respuesta'].unique().tolist(),
                    key="response_filter_persisted"
                )
            
            # Apply filter
            filtered_df = results_df.copy()
            if response_filter != 'Todas':
                filtered_df = filtered_df[filtered_df['Respuesta'] == response_filter]
            
            # Display filtered results
            st.dataframe(
                filtered_df[['Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia']], 
                use_container_width=True,
                height=400
            )
            
            # Download section
            st.markdown("### 📥 Descargar resultados")
            
            if len(results_df) > 0:
                # Get analyses from session state
                subsection_analyses = st.session_state.get('tab1_subsection_analyses', {})
                subsection_critical_evaluations = st.session_state.get('tab1_subsection_critical_evaluations', {})
                section_analyses = st.session_state.get('tab1_section_analyses', {})
                section_critical_evaluations = st.session_state.get('tab1_section_critical_evaluations', {})
                zip_buffer = create_results_download_with_sections(results_df, subsection_analyses, subsection_critical_evaluations, section_analyses, section_critical_evaluations)
                
                st.download_button(
                    label="📦 Descargar resultados en ZIP",
                    data=zip_buffer,
                    file_name="appraisal_checklist_results.zip",
                    mime="application/zip",
                    key="appraisal_download_button_persisted"
                )
            
            # Clear results button
            if st.button("🗑️ Limpiar resultados", key="clear_tab1_results"):
                st.session_state['tab1_results_df'] = None
                st.session_state['tab1_doc_stats'] = None
                st.rerun()
        else:
            if uploaded_file:
                st.info("👆 Haz clic en 'Analizar documento' para iniciar la evaluación de la valoración preliminar de calidad.")
            else:
                st.info("📁 Por favor sube un archivo DOCX para comenzar.")
