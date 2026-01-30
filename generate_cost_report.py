from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cost_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Cost Estimation Report: App Usage Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('This report provides a cost estimation for a single 200-page document processed through the application tabs (Preliminary Appraisal, Specific Attributes, and Sustainability).')

    # Usage Summary
    doc.add_heading('Usage Summary (200-Page Document)', level=1)
    
    usage_data = [
        ('Document Size', '~200 pages / 100,000 words'),
        ('Total Input Tokens', '~20.9 Million'),
        ('Total Output Tokens', '~95,000'),
        ('Total API Calls', '~190 calls')
    ]
    
    for item, value in usage_data:
        p = doc.add_paragraph()
        p.add_run(f'{item}: ').bold = True
        p.add_run(value)

    # Cost Comparison Table
    doc.add_heading('Cost Comparison by Model', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    headers = ['Model', 'Class', 'Input Price\n(per 1M)', 'Output Price\n(per 1M)', 'Est. Total Cost', 'Multiplier']
    for i, text in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(text)
        run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data
    # Costs calculated based on ~20.9M input and ~0.095M output
    # Formula: (20.9 * Input_Price) + (0.095 * Output_Price)
    
    # GPT-5-mini: (20.9 * 0.25) + (0.095 * 2.00) = 5.225 + 0.19 = $5.42
    # GPT-4.1-mini: (20.9 * 0.40) + (0.095 * 1.60) = 8.36 + 0.152 = $8.51
    
    models = [
        ('GPT-5-mini', 'Efficiency (New Pricing)', '$0.25', '$2.00', '$5.42', '1x'),
        ('GPT-4.1-mini', 'Efficiency (Proposed)', '$0.40', '$1.60', '$8.51', '~1.6x'),
        ('GPT-4o-mini', 'Efficiency (Standard)', '$0.15', '$0.60', '$3.20', '0.6x'),
        ('GPT-4o', 'Flagship', '$2.50', '$10.00', '$53.20', '~10x'),
        ('o1-mini', 'Reasoning (Fast)', '$3.00', '$12.00', '$63.85', '~12x'),
        ('o1', 'Reasoning (High)', '$15.00', '$60.00', '$319.20', '~60x')
    ]

    for model_name, model_class, in_price, out_price, total, mult in models:
        row_cells = table.add_row().cells
        row_cells[0].text = model_name
        row_cells[1].text = model_class
        row_cells[2].text = in_price
        row_cells[3].text = out_price
        row_cells[4].text = total
        row_cells[5].text = mult
        
        # Center numerical columns
        for i in range(2, 6):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Detailed Analysis
    doc.add_heading('Detailed Usage Analysis', level=1)
    
    doc.add_paragraph('1. Architecture Bottleneck', style='List Number')
    doc.add_paragraph('The application currently uses a full-context truncation strategy (up to 110k tokens) for every single evaluation question. It does not utilize Retrieval Augmented Generation (RAG) for the specific criteria evaluations in Tabs 1, 2, and 3. This means the input costs scale linearly with the number of questions, irrespective of whether the question only needs a small part of the document.')
    
    doc.add_paragraph('2. Token Breakdown', style='List Number')
    p = doc.add_paragraph()
    p.add_run('Tab 1 (Prelim Appraisal): ').bold = True
    p.add_run('~30 questions × 2 calls × 110k tokens = 6.6M tokens')
    
    p = doc.add_paragraph()
    p.add_run('Tab 2 (Specific Attributes): ').bold = True
    p.add_run('~85 criteria × 110k tokens = 9.35M tokens')
    
    p = doc.add_paragraph()
    p.add_run('Tab 3 (Sustainability): ').bold = True
    p.add_run('~45 criteria × 110k tokens = 4.95M tokens')

    # Recommendations
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph('1. For Cost Optmization:', style='List Bullet')
    doc.add_paragraph('Refactor the evaluation logic to use RAG (Vector Search) for Tabs 1 and 3. This would reduce input tokens from ~110,000 per call to ~4,000 per call, potentially reducing costs by 95% regardless of the model used.')
    
    doc.add_paragraph('2. Model Selection:', style='List Bullet')
    doc.add_paragraph('Continue using "mini" class models (GPT-5-mini / GPT-4o-mini). Using full flagship or reasoning models with the current architecture is economically impractical ($50-$300 per document).')

    doc.save('Cost_Estimation_Report.docx')

if __name__ == "__main__":
    create_cost_report()
