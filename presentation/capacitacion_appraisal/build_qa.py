"""Q&A briefing DOCX for the 30-min ILO training block (Ahmed)."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(0x00, 0x3E, 0x7E)
CYAN  = RGBColor(0x00, 0x72, 0xBC)
RED   = RGBColor(0xD6, 0x00, 0x1C)
GRAY  = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT = "QA_Sesion_Appraisal_GPT.docx"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def body(doc, text, bold=False, italic=False, size=10.5, after=6, color=None, before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    return p


def section(doc, letra, titulo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{letra}  ·  {titulo}")
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = BLUE
    # thin rule
    pr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "12")
    bt.set(qn("w:space"), "3"); bt.set(qn("w:color"), "D6001C")
    bd.append(bt); pr.append(bd)
    return p


def pregunta(doc, n, texto, quien=None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    set_cell_bg(c, "003E7E")
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(f"{n}.  {texto}")
    r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = WHITE
    if quien:
        r2 = p.add_run(f"     [{quien}]")
        r2.font.size = Pt(9.5); r2.font.bold = True; r2.font.color.rgb = RGBColor(0xFF, 0xC8, 0xCE)
    c.paragraphs[0].paragraph_format.space_after = Pt(2)
    c.paragraphs[0].paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def bloque(doc, etiqueta, texto, color=BLUE, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(etiqueta + "  ")
    r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = color
    r2 = p.add_run(texto)
    r2.font.size = Pt(10.5); r2.font.italic = italic
    return p


def decir(doc, texto):
    """The line to actually say out loud."""
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, "E0ECF6")
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run("QUÉ DECIR:  ")
    r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = BLUE
    r2 = p.add_run(texto)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ── Portada ───────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("PREGUNTAS Y RESPUESTAS")
r.font.bold = True; r.font.size = Pt(19); r.font.color.rgb = BLUE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)

p2 = doc.add_paragraph()
r2 = p2.add_run("Capacitación · Agente GPT Appraisal Checklist de PRODOCs")
r2.font.size = Pt(12); r2.font.bold = True; r2.font.color.rgb = CYAN
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(10)

meta = doc.add_table(rows=3, cols=2)
meta.style = "Table Grid"
for i, (k, v) in enumerate([
    ("Bloque", "Punto 4 — «Cómo funciona la tecnología y cuáles son sus límites»"),
    ("Expositor", "Ahmed Eid · Día 1, jueves 27 de agosto de 2026"),
    ("Uso", "Documento de apoyo para responder en vivo. 13 min de preguntas integrados al bloque."),
]):
    row = meta.rows[i]
    row.cells[0].text = k; row.cells[1].text = v
    set_cell_bg(row.cells[0], "003E7E"); set_cell_bg(row.cells[1], "F0F4F8")
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9.5)
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(9.5)
    row.cells[0].width = Cm(2.8); row.cells[1].width = Cm(13.5)

doc.add_paragraph()
body(doc, "Cómo usar este documento: la línea azul es lo que conviene decir en voz alta, en 20–30 segundos. "
          "El «respaldo» son los datos para sostener la respuesta si insisten. «Cuidado» marca la trampa de "
          "cada pregunta: lo que no conviene afirmar.", italic=True, size=9.5, color=GRAY)


def ejemplo(doc, titulo, lineas):
    """Worked example box: title + list of (etiqueta, texto) lines."""
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, "FFF4F5")
    c.text = ""
    p0 = c.paragraphs[0]
    r = p0.add_run("EJEMPLO · " + titulo)
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RED
    p0.paragraph_format.space_before = Pt(4); p0.paragraph_format.space_after = Pt(3)
    for etiqueta, texto in lineas:
        pl = c.add_paragraph()
        pl.paragraph_format.space_after = Pt(2)
        pl.paragraph_format.left_indent = Cm(0.25)
        if etiqueta:
            re_ = pl.add_run(etiqueta + "  ")
            re_.font.size = Pt(10); re_.font.bold = True; re_.font.color.rgb = BLUE
        rt = pl.add_run(texto)
        rt.font.size = Pt(10)
    c.paragraphs[-1].paragraph_format.space_after = Pt(5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

# ══ A · VARIABILIDAD Y CONFIANZA ═══════════════════════════════════════
section(doc, "A", "Variabilidad y confianza en el resultado")

pregunta(doc, 1, "¿Por qué puede cambiar una respuesta?")
decir(doc, "Porque un modelo de lenguaje no es determinista. No ejecuta una regla fija: en cada consulta "
           "construye la respuesta eligiendo entre continuaciones posibles del texto, y esa elección tiene "
           "un componente probabilístico. La misma pregunta sobre el mismo documento puede producir "
           "respuestas distintas. No es un error de programación ni una falla de configuración: es cómo "
           "funciona esta tecnología, y cualquier herramienta construida sobre un modelo de lenguaje —"
           "incluido el GPT Empresarial— tiene esta propiedad. La diferencia es que nosotros la medimos "
           "en lugar de ocultarla. Por eso no preguntamos una sola vez: cada criterio se evalúa diez veces "
           "de forma independiente, se toma el resultado más frecuente, y se les informa cuántas de las "
           "diez coincidieron.")
ejemplo(doc, "un criterio sobre supuestos del marco lógico", [
    ("Criterio:", "«El PRODOC identifica los supuestos que conectan productos con resultados»."),
    ("Situación:", "el documento menciona supuestos en un párrafo narrativo de la sección de riesgos, pero "
                   "la columna de supuestos del marco lógico está vacía."),
    ("Qué pasó:", "de las 10 corridas, 7 respondieron «Partial» —hay supuestos, pero no en el instrumento— "
                  "y 3 respondieron «No» —el marco lógico no los tiene—."),
    ("Resultado entregado:", "Partial, con estabilidad 70% y «Resultado Alternativo: No»."),
    ("Lectura correcta:", "el criterio es genuinamente limítrofe. El desacuerdo del modelo está reflejando "
                          "una ambigüedad real del documento, no un capricho. Ahí decide el appraiser."),
])
bloque(doc, "Respaldo:", "además de la aleatoriedad propia del modelo, hay tres fuentes que amplifican la "
      "variación: criterios que admiten más de una lectura, evidencia dispersa en distintas partes del "
      "documento, y documentos muy extensos. Las diez corridas y el resultado modal absorben esa variación; "
      "el porcentaje de estabilidad la hace visible.")
bloque(doc, "Cuidado:", "no prometer que el resultado es idéntico siempre. No lo es, y prometerlo destruye "
      "la credibilidad de toda la sesión la primera vez que alguien vea una diferencia entre dos corridas.",
      color=RED)

pregunta(doc, 2, "¿Por qué aparece una estabilidad de 50–70%? ¿Qué significa?")
decir(doc, "La estabilidad es el porcentaje de las diez corridas que coincidió con el resultado final. Un "
           "60% quiere decir que seis corridas dijeron una cosa y cuatro dijeron otra. El umbral "
           "institucional es 80%: por debajo de eso el criterio se marca automáticamente en la columna "
           "«Revisión humana recomendada». Lo importante es entender qué les está diciendo ese número. No "
           "les dice que el sistema falló: les dice que ese criterio, sobre ese documento concreto, no "
           "admite una lectura única. Casi siempre eso significa una de dos cosas: o el criterio es "
           "genuinamente ambiguo, o la evidencia del documento es insuficiente para resolverlo. En ambos "
           "casos la conclusión práctica es la misma: ese es un criterio donde su juicio profesional pesa "
           "más que el diagnóstico automático.")
ejemplo(doc, "cómo se ve en el Excel", [
    ("Respuesta:", "Partial"),
    ("Estabilidad (%):", "60"),
    ("Estable (≥80%):", "No"),
    ("Resultado Alternativo:", "No  —es decir, las cuatro corridas discrepantes dijeron «No»"),
    ("Revisión humana recomendada:", "Sí (estabilidad bajo umbral)"),
    ("Qué hacer:", "abrir la evidencia citada y decidir entre Partial y No. El rango de discrepancia ya está "
                   "acotado: el modelo nunca consideró «Yes», así que la discusión no es si el criterio se "
                   "cumple, sino si se cumple parcialmente o no se cumple."),
])
bloque(doc, "Respaldo:", "el Excel reporta también la distribución completa de las diez respuestas, no solo "
      "la moda y la alternativa. Eso permite distinguir un 60/40 limpio de un caso donde las diez corridas "
      "se repartieron entre tres respuestas distintas, que es una señal mucho más fuerte de ambigüedad.")
bloque(doc, "Cuidado:", "no presentar la baja estabilidad como una debilidad del sistema. Es información que "
      "otras herramientas simplemente no entregan: el GPT Empresarial da una respuesta única y no dice nada "
      "sobre su propia consistencia. Aquí, la incertidumbre es visible y accionable.", color=RED)

pregunta(doc, 3, "Si un criterio sale con 100% de estabilidad, ¿puedo confiar sin revisar la evidencia?")
decir(doc, "No, y esta es probablemente la confusión más importante que quiero evitar hoy. La estabilidad "
           "mide cuánta confianza interna tuvo el modelo, no si acertó. Son dos cosas distintas. Si el "
           "modelo interpretó mal un criterio, lo va a interpretar mal las diez veces, con total "
           "consistencia, y les va a entregar un 100% de estabilidad sobre una conclusión equivocada. La "
           "estabilidad alta les dice «el modelo no dudó»; no les dice «el modelo tiene razón». Lo único "
           "que convierte un resultado automático en un diagnóstico defendible es la evidencia citada, y "
           "por eso está en el Excel, en cada fila.")
ejemplo(doc, "100% estable y discutible", [
    ("Criterio:", "1.5.6 — «la propuesta promueve un enfoque transformador en materia de género»."),
    ("Texto del PRODOC:", "«el proyecto asegurará la participación equitativa de mujeres en los comités "
                          "sectoriales, con una meta de 50% de representación femenina»."),
    ("Qué hizo el modelo:", "las diez corridas marcaron el test de acciones dedicadas como verdadero —hay "
                            "una acción concreta, con meta cuantificable— y devolvieron Partial con 100% "
                            "de estabilidad."),
    ("Por qué es discutible:", "asegurar participación numérica es inclusión, no necesariamente "
                               "transformación de relaciones de poder. Un especialista en género podría "
                               "sostener que ese texto no acredita enfoque transformador."),
    ("Conclusión:", "el modelo fue perfectamente consistente y aun así el resultado merece discusión. La "
                    "estabilidad no lo protege de eso."),
])
bloque(doc, "Respaldo:", "estabilidad alta significa consistencia; corrección significa verificación contra "
      "el documento. El Excel las separa deliberadamente en dos columnas distintas para que no se confundan.")
bloque(doc, "Cuidado:", "si el público sale de esta sesión creyendo que 100% equivale a «correcto», habremos "
      "generado exceso de confianza, que es peor que no haber capacitado. Decir esta respuesta despacio.",
      color=RED)

pregunta(doc, 4, "¿Qué significa que un criterio sea subjetivo y qué implica?")
decir(doc, "Significa que dos expertos competentes podrían discrepar legítimamente sobre él leyendo el mismo "
           "documento. La rúbrica etiqueta cada criterio como subjetividad alta, media o baja, y esa "
           "etiqueta hace dos cosas. Primero, el modelo dedica más razonamiento a los criterios de alta "
           "subjetividad antes de responder. Segundo, esos criterios se marcan siempre para revisión "
           "humana, independientemente de la estabilidad que hayan obtenido. Hay dos disparadores de esa "
           "marca: subjetividad alta, o estabilidad por debajo de 80%. Un criterio puede activar los dos.")
ejemplo(doc, "los dos extremos de la escala", [
    ("Subjetividad baja:", "«¿el PRODOC incluye un marco lógico?» — se verifica mirando el documento; dos "
                           "revisores coincidirán siempre."),
    ("Subjetividad alta:", "«¿el enfoque de género es transformador?» — depende de qué se entienda por "
                           "transformador, del contexto del país y del sector. Dos especialistas pueden "
                           "discrepar de buena fe."),
    ("Implicación práctica:", "en el primer caso el diagnóstico automático es casi definitivo. En el "
                              "segundo es un punto de partida para su lectura, nunca una conclusión."),
])
bloque(doc, "Respaldo:", "la columna «Revisión humana recomendada» funciona como cola de trabajo priorizada. "
      "Si tienen tiempo limitado, empezar por ahí concentra el esfuerzo humano donde el diagnóstico "
      "automático es menos concluyente.")
bloque(doc, "Cuidado:", "«subjetivo» no equivale a «poco confiable». Significa que el juicio depende del "
      "contexto institucional, que es precisamente donde el appraiser aporta lo que la herramienta no "
      "puede aportar.", color=RED)

# ══ B · CÓMO LEER LOS RESULTADOS ═══════════════════════════════════════
section(doc, "B", "Cómo leer los resultados")

pregunta(doc, 5, "¿Por qué no encontró algo que yo sé que existe en el PRODOC?")
decir(doc, "Es la pregunta más frecuente y casi siempre tiene la misma explicación. El Agente aplica un "
           "filtro deliberado: distingue entre menciones de marco y menciones dedicadas. Una mención de "
           "marco es el tema apareciendo sin desarrollo propio: una lista de grupos poblacionales, lenguaje "
           "genérico de inclusión, una enumeración de tipo «entre otros». Una mención dedicada es el tema "
           "con espacio propio en el diseño: un producto que lo nombra, un indicador desagregado, una "
           "actividad cuyo propósito principal es ese, una partida presupuestaria, una meta cuantificable. "
           "Si toda la evidencia disponible es de marco, el resultado debe ser «No» o «Not Found» aunque el "
           "tema aparezca quince veces en el documento. Muchas veces, lo que el especialista recuerda haber "
           "escrito es exactamente marco. Y cuando no sea así —cuando la evidencia dedicada esté y el "
           "Agente no la haya visto— verifíquenlo, descarten ese diagnóstico y avísennos, porque ese caso "
           "es información directa para ajustar la rúbrica.")
ejemplo(doc, "por qué «sí cubrimos discapacidad» puede salir como No", [
    ("Criterio:", "1.5.1 — «el proyecto ha tenido en cuenta las necesidades de las personas con discapacidad»."),
    ("Lo que dice el PRODOC:", "en el objetivo general, «contribuir al trabajo decente de mujeres, jóvenes, "
                               "personas con discapacidad y pueblos indígenas en el sector rural»."),
    ("Qué vio el Agente:", "una lista de cuatro grupos poblacionales, sin ningún seguimiento posterior. "
                           "Ningún producto, indicador, actividad ni partida menciona discapacidad. "
                           "Clasificación: MARCO."),
    ("Resultado:", "No — con la evidencia citada, que es exactamente esa frase del objetivo."),
    ("Por qué está bien:", "el diagnóstico es correcto y útil. El PRODOC nombra el grupo pero no lo aterriza "
                           "en el diseño. Eso es precisamente la brecha que el Appraisal debería detectar."),
    ("Qué haría cambiar el resultado:", "un indicador desagregado por discapacidad, una actividad de "
                                        "accesibilidad, o una línea presupuestaria específica."),
])
bloque(doc, "Respaldo:", "el umbral es concreto: un pasaje donde el sujeto aparece en una lista de tres o "
      "más grupos, sin seguimiento dedicado, se clasifica como marco automáticamente. Esta regla existe "
      "porque el problema clásico del appraisal es justamente el proyecto que nombra un tema transversal "
      "sin hacer nada al respecto.")
bloque(doc, "Cuidado:", "no ponerse a la defensiva cuando alguien plantee este caso. Que el appraiser "
      "detecte y descarte un diagnóstico incorrecto es el uso correcto de la herramienta, no una queja. "
      "Conviene decirlo en esos términos.", color=RED)

pregunta(doc, 6, "¿«Not Found» y «No» son lo mismo?")
decir(doc, "No, y confundirlos lleva a la acción equivocada. «No» afirma algo sobre el proyecto: el criterio "
           "no se cumple. Eso apunta normalmente a una brecha de diseño, y se resuelve rediseñando. «Not "
           "Found» afirma algo sobre el documento: no hay información suficiente para determinarlo. Eso "
           "apunta a una brecha documental, y puede resolverse escribiendo lo que el equipo ya sabe. La "
           "distinción importa porque son conversaciones distintas con el equipo formulador: en un caso hay "
           "que revisar el diseño, en el otro basta con documentar una decisión que ya está tomada.")
ejemplo(doc, "el mismo criterio, dos resultados distintos", [
    ("Criterio:", "«el PRODOC define una estrategia de sostenibilidad o salida al término del proyecto»."),
    ("Caso «No»:", "el documento dice «la continuidad de las actividades dependerá de la disponibilidad de "
                   "fondos futuros». Hay respuesta, y es inadecuada: no hay estrategia, hay una expectativa. "
                   "→ brecha de diseño, hay que construir la estrategia."),
    ("Caso «Not Found»:", "el documento no contiene ninguna sección, párrafo ni mención sobre qué ocurre "
                          "después del mes 36. → brecha documental; posiblemente el equipo sí tiene un plan "
                          "y no lo escribió."),
    ("Por qué importa:", "en el primer caso el appraisal debe cuestionar el diseño. En el segundo, basta "
                         "preguntar al equipo y pedir que lo incorpore."),
])
bloque(doc, "Respaldo:", "la escala completa es Yes / Partial / No / Not Found / N/A. «N/A» aparece solo "
      "cuando la aplicabilidad condicional del criterio no se satisface —por ejemplo, un criterio que solo "
      "aplica si el proyecto contrata trabajadores directamente—.")

pregunta(doc, 7, "¿Por qué no me da un informe con comentarios de mejora al PRODOC?")
decir(doc, "Por una decisión de diseño, no por una carencia. Una recomendación útil necesita conocimiento "
           "que el Agente no tiene: el contexto del país, la relación con el mandante, el historial del "
           "programa, lo que ya se intentó y falló. El Agente solo tiene el documento. Si le pidiéramos "
           "recomendaciones, produciría texto plausible y genérico, y ese texto sería indistinguible de una "
           "recomendación fundada: eso es peor que no darla. Preferimos un diagnóstico acotado y auditable. "
           "Ahora bien, esto no cierra la puerta: sobre el resultado ya generado pueden seguir conversando "
           "con el Agente y pedirle exactamente eso, sabiendo de dónde sale cada afirmación. La "
           "recomendación real surge de la triangulación entre el hallazgo automatizado y el conocimiento "
           "del appraiser.")
ejemplo(doc, "del hallazgo a la pregunta útil", [
    ("Hallazgo automatizado:", "«no se identifican claramente los supuestos que conectan productos con "
                               "resultados». Resultado: PARTIAL."),
    ("Lo que el Agente NO hace:", "escribir «se recomienda incorporar una matriz de supuestos en la sección "
                                  "4». Suena bien y no está fundado en nada que el Agente sepa del proyecto."),
    ("Lo que sí permite hacer:", "convertir el hallazgo en preguntas al equipo formulador — ¿qué condiciones "
                                 "deben darse para que los productos previstos generen el cambio esperado? "
                                 "¿están esas condiciones bajo control del proyecto? ¿dónde están "
                                 "documentadas en el PRODOC?"),
    ("Resultado:", "la responsabilidad vuelve al revisor, y las respuestas permiten distinguir una "
                   "deficiencia real de diseño de una simple deficiencia documental."),
])
bloque(doc, "Cuidado:", "no prometer que esta funcionalidad llegará en una versión futura. Es una decisión "
      "deliberada, y anunciarla como pendiente genera una expectativa que no vamos a satisfacer.", color=RED)

# ══ C · COSTO, TIEMPO Y ALCANCE ════════════════════════════════════════
section(doc, "C", "Costo, tiempo y alcance")

pregunta(doc, 8, "¿Cuánto cuesta evaluar un PRODOC y de qué depende el costo?")
decir(doc, "El costo depende de dos variables: el tamaño del documento y cuántos criterios se evalúen. La "
           "aritmética es sencilla y conviene tenerla clara. Una valoración completa son 76 criterios, y "
           "cada criterio se ejecuta diez veces: del orden de 760 consultas al modelo. Cada consulta "
           "incluye el texto del documento, así que un PRODOC extenso cuesta más que uno breve en cada una "
           "de esas 760 consultas. Evaluar solo la sección 3, que tiene 27 criterios, son 270 consultas: "
           "aproximadamente un tercio. Por eso insisto tanto en filtrar: no es solo para que llegue antes, "
           "es la palanca directa sobre el costo. Todavía no tenemos una cifra en dólares por documento "
           "medida en uso real, y medirla es justamente uno de los objetivos del piloto.")
ejemplo(doc, "la aritmética del alcance", [
    ("Rúbrica completa:", "76 criterios × 10 corridas ≈ 760 consultas"),
    ("Solo sección 3 (Marco de resultados):", "27 criterios × 10 ≈ 270 consultas — un tercio del costo"),
    ("Solo sección 1 (Pertinencia):", "20 criterios × 10 ≈ 200 consultas"),
    ("Solo subsección 1.5:", "6 criterios × 10 = 60 consultas — útil para verificar un punto concreto"),
    ("Recomendación:", "para una primera lectura de un PRODOC nuevo, una sección. La rúbrica completa "
                       "cuando el documento ya esté maduro y se quiera el diagnóstico integral."),
])
bloque(doc, "Respaldo:", "el piloto tiene asignado un presupuesto de USD 500 en consumo de API, destinado "
      "íntegramente a ese consumo. Una de sus finalidades explícitas es medir el costo unitario real por "
      "documento, para poder estimar con base empírica un despliegue mayor.")
bloque(doc, "Cuidado:", "no dar una cifra en dólares por PRODOC. No la hemos medido en uso real, y un número "
      "improvisado hoy se va a citar después como si fuera oficial. Si insisten, la respuesta es: «esa "
      "medición es precisamente uno de los objetivos del piloto».", color=RED)

pregunta(doc, 9, "¿Cuánto tarda una evaluación?")
decir(doc, "Varios minutos para la rúbrica completa, y bastante menos si filtran por sección. No son 760 "
           "consultas en fila una detrás de otra: el motor lanza hasta 48 en paralelo, por eso el tiempo "
           "total es de minutos y no de horas. Hay un detalle operativo que conviene conocer: la primera "
           "evaluación del día puede tardar más en arrancar, porque el servicio se suspende cuando lleva "
           "un rato sin uso y necesita unos segundos para despertar. Si ven que la primera consulta demora "
           "en responder, es eso, no una falla.")
bloque(doc, "Respaldo:", "conviene no cerrar la conversación mientras corre. Si el servicio se reinicia "
      "durante una evaluación, esa evaluación se pierde y hay que relanzarla: los trabajos en curso no se "
      "recuperan.")

pregunta(doc, 10, "¿Se pueden evaluar otros documentos distintos de PRODOCs?")
decir(doc, "Con este Agente no, porque está construido sobre el Appraisal Checklist, que es una rúbrica de "
           "diseño de proyecto. Pero se desarrollaron otros dos agentes con la misma arquitectura. Uno de "
           "atributos específicos, que aplica rúbricas temáticas: metodologías participativas, integración "
           "del enfoque de género, y transición justa. Y otro de sostenibilidad, que se aplica según la "
           "etapa del ciclo: diseño para PRODOCs, implementación para informes de avance, y pre-cierre "
           "para documentos de cierre o evaluación final.")
ejemplo(doc, "cuándo usar cuál", [
    ("Appraisal Checklist (este):", "un PRODOC en fase de diseño, diagnóstico integral de calidad."),
    ("Atributos específicos:", "profundizar un tema que el Checklist toca solo de forma transversal. Si el "
                               "1.5.6 salió Partial y quieren un análisis detallado de género, ese es el "
                               "agente."),
    ("Sostenibilidad:", "un informe de avance a medio término, o un documento de cierre."),
    ("Advertencia:", "las escalas son distintas —Yes/Partial/No aquí, 1 a 5 en atributos específicos, 0 a 3 "
                     "en sostenibilidad—. No son intercambiables ni promediables entre sí."),
])

# ══ D · TECNOLOGÍA, ACCESO Y CONFIDENCIALIDAD ══════════════════════════
section(doc, "D", "Tecnología, acceso y confidencialidad")

pregunta(doc, 11, "Si uso el GPT Empresarial de la OIT y le subo la rúbrica, ¿obtengo los mismos resultados?")
decir(doc, "No, y la diferencia no es de marca sino de método. Si le suben el Appraisal Checklist al GPT "
           "Empresarial, hará una lectura general del archivo en una sola pasada y les devolverá un texto "
           "razonable. Lo que no hará es lo que hace este Agente: descomponer cada criterio en tests "
           "cerrados con una regla de decisión explícita, ejecutar cada uno diez veces, consolidar el "
           "resultado y decirles cuánto coincidieron las corridas. Tampoco les dará un Excel con la "
           "evidencia citada criterio por criterio. Y hay un punto práctico adicional: hoy no existe acceso "
           "a este Agente desde el entorno empresarial. El valor agregado no es «usar inteligencia "
           "artificial» —eso lo hacen los dos—, es la rúbrica institucional aplicada de forma mecánica, "
           "repetida y medida, con la misma versión para todos los proyectos.")
ejemplo(doc, "la misma pregunta, dos respuestas", [
    ("GPT Empresarial:", "«El proyecto aborda parcialmente el enfoque de género. Se recomienda fortalecer "
                         "la perspectiva transformadora.» — Plausible, pero: ¿qué criterio? ¿con qué "
                         "evidencia? ¿lo diría igual si se lo preguntan otra vez?"),
    ("Agente Appraisal:", "«Criterio 1.5.6 — Partial. Estabilidad 90%. T1 verdadero, T2 falso, T3 verdadero. "
                          "Evidencia: “el proyecto asegurará la participación equitativa…”. Revisión humana "
                          "recomendada: sí, subjetividad alta.»"),
    ("La diferencia:", "el segundo es verificable, comparable entre proyectos y archivable. El primero es "
                       "una opinión bien redactada."),
])
bloque(doc, "Cuidado:", "no descalificar el GPT Empresarial en general: es útil para muchas otras tareas. "
      "La diferencia que estamos señalando es específica a esta tarea.", color=RED)

pregunta(doc, 12, "¿Qué pasa con la confidencialidad del PRODOC que subo?")
decir(doc, "Es una pregunta importante y merece una respuesta precisa, no tranquilizadora. El recorrido del "
           "documento es este: ustedes lo suben a la conversación de ChatGPT, que lo aloja en sus "
           "servidores; nuestro servicio lo descarga desde ahí para procesarlo; extrae el texto y lo envía "
           "al modelo para la evaluación. Nuestro servicio no guarda el documento: lo procesa y lo "
           "descarta, y no hay base de datos ni registro histórico. Pero la conversación sí queda en la "
           "cuenta de ChatGPT de quien la hizo, con el archivo adjunto, y su retención depende de la "
           "configuración de esa cuenta. Además, el Excel que descargan contiene citas textuales del "
           "PRODOC, así que debe manejarse con la misma confidencialidad que el documento original. La "
           "recomendación operativa es simple: no subir documentos que no estén autorizados para "
           "procesamiento por servicios externos.")
ejemplo(doc, "el recorrido del documento, en orden", [
    ("1.", "Usted adjunta el .docx en la conversación de ChatGPT → queda alojado en servidores de OpenAI."),
    ("2.", "El Agente entrega a nuestro servicio un enlace temporal de descarga."),
    ("3.", "Nuestro servicio descarga el archivo, extrae el texto y lo descarta al terminar."),
    ("4.", "El texto se envía al modelo en cada una de las consultas de evaluación."),
    ("5.", "El resultado vive en memoria del servicio hasta que se reinicia; no se archiva."),
    ("6.", "El Excel queda en su equipo. La conversación queda en su cuenta de ChatGPT."),
])
bloque(doc, "Cuidado:", "no afirmar que «los datos no salen de la OIT» —es falso— ni que «no se usan para "
      "entrenar» —depende de las condiciones contractuales de la cuenta, que no corresponde afirmar desde "
      "esta sesión—. Si presionan en ese punto, derivar a la instancia institucional competente.", color=RED)

pregunta(doc, 13, "¿Se puede vincular a otros recursos o aplicaciones de la OIT?")
decir(doc, "Sí, y ahí está buena parte del potencial que todavía no hemos explotado. Lo inmediato: sobre el "
           "resultado ya generado pueden seguir conversando con el Agente, pedirle que resuma las brechas "
           "de una sección, que las ordene por prioridad, que las reformule como preguntas para el equipo. "
           "Lo más interesante es el vínculo con el chatbot i-EVal: ante una brecha detectada, pueden "
           "consultar qué han encontrado evaluaciones anteriores en proyectos similares de la región. Eso "
           "convierte un hallazgo de diseño en una estrategia implementable, apoyada en evidencia "
           "acumulada.")
ejemplo(doc, "encadenar el diagnóstico con evidencia de evaluación", [
    ("Hallazgo del Agente:", "criterio de sostenibilidad institucional — Partial, estabilidad 80%."),
    ("Pregunta a i-EVal:", "«¿qué problemas recurrentes han identificado las evaluaciones en proyectos que "
                           "dependen de capacidades de ministerios de trabajo?»"),
    ("Uso del resultado:", "incorporar al diseño una medida de mitigación fundada en lo que ya se aprendió "
                           "en la región, en lugar de una declaración genérica de sostenibilidad."),
])

# ══ E · INSTITUCIONALES ════════════════════════════════════════════════
section(doc, "E", "Institucionales — dar el marco y derivar")
body(doc, "Estas cuatro no son preguntas técnicas. La estrategia es la misma en las cuatro: responder con "
          "precisión la parte que sí es del expositor, y devolver explícitamente la parte que es una "
          "decisión institucional. Responder de más aquí compromete definiciones que no son suyas.",
     italic=True, size=9.5, color=GRAY, after=8)

pregunta(doc, 14, "¿Cuándo lo puedo usar?", quien="derivar a Cybele · bloque 3")
decir(doc, "La secuencia de implementación para 2026 y el perfil de usuario los cubrió Cybele en el bloque "
           "anterior: hoy el uso está acotado a EVAL-ROLAC como primera etapa. Lo que sí corresponde a mi "
           "parte es el requisito técnico de uso, y quiero ser explícito: esta herramienta supone que quien "
           "la usa conoce el proceso de appraisal y sabe interpretar el Checklist. No es una herramienta "
           "para cualquiera que tenga un PRODOC en la mano. Si quien la usa no puede juzgar si un "
           "diagnóstico es razonable, la herramienta deja de ser un apoyo y se convierte en un riesgo: "
           "produce resultados con apariencia de autoridad que nadie está en condiciones de contrastar.")

pregunta(doc, 15, "¿Está en línea con las IGDS?", quien="derivar")
decir(doc, "Esa es una definición institucional que excede lo técnico y no me corresponde pronunciarla. Lo "
           "que sí puedo afirmar con precisión es lo siguiente: la rúbrica que aplica el Agente es el "
           "Appraisal Checklist institucional, sin modificaciones metodológicas de nuestra parte. No "
           "inventamos criterios, no cambiamos umbrales y no alteramos la escala. Lo que hicimos fue "
           "mecanizar la aplicación de esa rúbrica. La alineación formal con las IGDS corresponde definirla "
           "a la instancia competente, y con gusto damos toda la información técnica que necesiten para "
           "esa evaluación.")
bloque(doc, "Cuidado:", "no afirmar que sí está alineado ni que no lo está. Es un pronunciamiento "
      "institucional; una respuesta improvisada en una capacitación puede citarse después como posición "
      "del proyecto.", color=RED)

pregunta(doc, 16, "¿Cómo es y cuánto cuesta el mantenimiento?", quien="derivar · marco general")
decir(doc, "El mantenimiento tiene tres componentes distintos y conviene no mezclarlos. Primero, el consumo "
           "de API, que es variable y depende directamente de cuánto se use. Segundo, el servicio donde "
           "corre el backend, que hoy está en un plan sin costo con las limitaciones que eso implica. Y "
           "tercero, la actualización de la rúbrica: si el Appraisal Checklist cambia, hay que actualizarlo "
           "en el servidor y volver a desplegar. Ese tercer punto es el más importante desde el punto de "
           "vista institucional, porque significa que la herramienta necesita una persona o un equipo "
           "responsable de mantenerla sincronizada con la metodología. El esquema y el costo para después "
           "del piloto son una decisión institucional pendiente.")

pregunta(doc, 17, "¿Cómo puedo ayudar a que se institucionalice la herramienta?", quien="derivar a Cybele / día 2")
decir(doc, "Lo más útil que pueden hacer es usarla sobre PRODOCs reales y decirnos dónde falla. En "
           "particular dos tipos de caso: cuando el Agente no reconoció algo que sí estaba en el documento, "
           "y cuando marcó como brecha algo que en realidad no lo era. Cada uno de esos casos es "
           "información directa para ajustar la rúbrica, y de hecho ya hemos hecho ajustes de ese tipo. El "
           "piloto está diseñado para recoger exactamente eso de forma sistemática: brecha detectada, "
           "recomendación, respuesta del equipo, acción tomada y situación posterior.")
bloque(doc, "Respaldo:", "si preguntan por un ejemplo concreto de ajuste: un criterio de género calificaba "
      "«No» cuando el PRODOC tenía acciones dedicadas a transformar relaciones pero no usaba la "
      "terminología de la taxonomía. Se corrigió para que la sustancia prevalezca sobre la nomenclatura. "
      "Ese cambio salió de una revisión como las que ustedes pueden hacer.")

# ══ Cierre ═════════════════════════════════════════════════════════════
section(doc, "F", "Tres frases para no perder el hilo")
for txt in [
    "El Agente lee el documento, no el proyecto. Un «No» puede ser una brecha de diseño o simplemente algo no documentado.",
    "La estabilidad mide consistencia, no acierto. La evidencia citada es lo único que hace defendible el diagnóstico.",
    "El valor de la herramienta depende de la calidad de la revisión humana posterior.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(txt)
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(6)

body(doc, "Si una pregunta se sale del guion y no tiene la respuesta: «No lo tengo medido / definido, lo "
          "verifico y les respondo». Es siempre preferible a improvisar una cifra, un plazo o un "
          "compromiso institucional.", italic=True, size=9.5, color=GRAY, before=10)

doc.save(OUTPUT)
print("Guardado: " + OUTPUT)
