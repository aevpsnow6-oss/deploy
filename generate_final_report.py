"""Generate the final consultancy deliverable report for ILO PO 40551829."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

BLUE = RGBColor(0x00, 0x3E, 0x7E)   # ILO dark blue
CYAN = RGBColor(0x00, 0x9A, 0xD2)   # ILO light blue
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT = "Informe_Final_Consultoría_OIT_PO40551829.docx"


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=BLUE):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table_header_row(table, headers, bg="003E7E"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9.5)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_data_row(table, values, row_idx):
    row = table.add_row()
    bg = "F0F4F8" if row_idx % 2 == 0 else "FFFFFF"
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)


doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ── Default body font ────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ════════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(36)
run = p.add_run("INFORME FINAL DE CONSULTORÍA")
run.font.bold = True
run.font.size = Pt(20)
run.font.color.rgb = BLUE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
run2 = p2.add_run("Optimización de la Aplicación de IA para la Gestión y\nProcesamiento de Documentos de Diseño de Proyectos")
run2.font.size = Pt(14)
run2.font.color.rgb = CYAN
run2.font.bold = True
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(10)

doc.add_paragraph()

meta_table = doc.add_table(rows=7, cols=2)
meta_table.style = "Table Grid"
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Organización", "Organización Internacional del Trabajo (OIT) – Oficina Regional para América Latina y el Caribe"),
    ("Orden de Compra", "40551829 / 0"),
    ("Consultor", "Ahmed Guillermo Eid Valdiviezo"),
    ("Período de ejecución", "25 de mayo de 2026 – 8 de junio de 2026"),
    ("Aprobadores técnicos", "Cybele Burga (Oficial Regional de Evaluación) · Carlos Castañeda (Oficial de Evaluación)"),
    ("Fecha de entrega", "8 de junio de 2026"),
    ("Versión", "1.0 – Final"),
]
for i, (k, v) in enumerate(meta_data):
    row = meta_table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], "003E7E")
    set_cell_bg(row.cells[1], "F0F4F8")
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(9.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Resumen Ejecutivo", level=1)

add_body(doc,
    "La presente consultoría tuvo como objetivo mejorar y ampliar las funcionalidades de la aplicación de "
    "inteligencia artificial desarrollada en Streamlit para la OIT, destinada a la gestión y procesamiento "
    "de documentos de diseño de proyectos (PRODOCs) y al análisis de recomendaciones de evaluaciones. "
    "El trabajo se ejecutó entre el 25 de mayo y el 8 de junio de 2026, en estrecha colaboración con el "
    "equipo técnico de la OIT.")

add_body(doc,
    "Se cumplieron íntegramente las tres tareas contractuales —ajustes de outputs, modularización del código "
    "y actualización de documentación— y se entregaron además tres GPTs personalizados en ChatGPT que "
    "permiten a los equipos de la OIT realizar evaluaciones directamente desde su interfaz conversacional, "
    "sin necesidad de acceder al entorno Streamlit.")

add_body(doc, "Los entregables principales son:", space_after=2)
for item in [
    ("Rúbrica v3 mecanizada (Tab 1):", " Evaluación automatizada de 76 criterios de calidad de PRODOCs con motor de razonamiento adaptativo y salida XLSX estructurada."),
    ("Aplicación Streamlit modularizada:", " Separación del código fuente en dos módulos independientes: valoración de calidad de proyectos (oli_v6_deploy_core.py) y clasificación de recomendaciones (oli_v6_deploy_recommendations.py)."),
    ("Tres GPTs publicados:", " GPT de valoración preliminar de calidad (Tab 1 v3), GPT de diagnóstico de sostenibilidad (Tab 3) y GPT de diagnóstico de atributos específicos (Tab 2), todos con backend FastAPI desplegado en Render."),
    ("Documentación técnica:", " Guías de configuración, transferencia y uso para cada módulo y GPT."),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 2. CONTEXTO Y ARQUITECTURA DE LA APLICACIÓN
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Contexto y Arquitectura de la Aplicación", level=1)

add_body(doc,
    "La aplicación recibida al inicio de la consultoría era un único fichero monolítico (oli_v6_deploy.py, "
    "≈ 550 KB) que integraba todas las funcionalidades en un mismo módulo Streamlit. La arquitectura "
    "resultante de la consultoría organiza el sistema en cuatro capas diferenciadas:")

arch_table = doc.add_table(rows=1, cols=3)
arch_table.style = "Table Grid"
add_table_header_row(arch_table, ["Capa", "Componente", "Responsabilidad"])
arch_rows = [
    ("Interfaz Streamlit – Calidad", "oli_v6_deploy_core.py", "Tabs 1–4: valoración preliminar, atributos específicos, sostenibilidad, RAG documental"),
    ("Interfaz Streamlit – Recomendaciones", "oli_v6_deploy_recommendations.py", "Tabs 5–6: clasificación de recomendaciones en español e inglés"),
    ("Motor de evaluación (sin Streamlit)", "tab1_v3_core.py · sustainability_core.py", "Lógica pura reutilizable por Streamlit y por los GPT Actions"),
    ("Backend GPT Actions", "gpt_action_api.py (FastAPI, Render)", "API HTTPS con endpoints asíncronos para los tres GPTs"),
]
for i, row_data in enumerate(arch_rows):
    add_data_row(arch_table, row_data, i)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 3. TAREA 1 – AJUSTES DE LA LÓGICA DE GENERACIÓN DE OUTPUTS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Tarea 1 – Ajustes de la Lógica de Generación de Outputs", level=1)

add_body(doc,
    "Se realizaron ajustes iterativos en prompts, rúbricas y estructura de resultados para mejorar la "
    "coherencia, reducir redundancias y aumentar la utilidad operativa de las salidas descargables. "
    "El trabajo se organizó en tres áreas: la rúbrica de Tab 1, los criterios de innovación y las "
    "salidas de recomendaciones.")

# 3.1 Rubric iterations
add_heading(doc, "3.1 Evolución de la Rúbrica de Tab 1 (Valoración Preliminar de Calidad)", level=2)

add_body(doc,
    "La rúbrica de valoración de PRODOCs evolucionó en tres versiones durante la consultoría, cada una "
    "incorporando observaciones del equipo técnico de la OIT:")

rub_table = doc.add_table(rows=1, cols=3)
rub_table.style = "Table Grid"
add_table_header_row(rub_table, ["Versión", "Archivo", "Cambios principales"])
rub_rows = [
    ("v1 (base)", "Rubrica_Tab1_Detallada_Full.xlsx", "Rúbrica original con 76 criterios, estructura de secciones 1–5, escala SI/PARCIAL/NO/NA"),
    ("v2 (revisión cliente)", "Rubrica_Tab1_Detallada_Full_v2.xlsx", "Aplicación de 30+ comentarios del equipo OIT: correcciones de redacción, ajuste de ejemplos orientadores, nuevas notas operativas, columna de trazabilidad de cambios"),
    ("v3 (mecanizada)", "Rubrica_Tab1_Detallada_Full_v3.xlsx", "Versión final mecanizada: columna 'Pregunta orientadora' (no evaluada), definición DEDICADO/MARCO alineada con Tab 1, orden jerárquico explícito, tipos de criterio con descripciones enriquecidas"),
]
for i, row_data in enumerate(rub_rows):
    add_data_row(rub_table, row_data, i)

doc.add_paragraph()

add_heading(doc, "3.2 Motor de Evaluación v3 – Diseño del Prompt y Razonamiento Adaptativo", level=2)

add_body(doc,
    "El motor de evaluación v3 (tab1_v3_core.py) implementa un sistema de prompt estructurado y "
    "razonamiento diferenciado por nivel de subjetividad del criterio:")

for item in [
    ("Prompt de sistema (V3_SYSTEM_PROMPT):", " Define el rol del modelo como analista experto de PRODOCs OIT. Instruye al modelo a emitir exactamente uno de los veredictos aceptados (SI, PARCIAL, NO, NA), a fundamentar la evidencia en citas textuales del documento, y a evitar interpretaciones subjetivas no respaldadas en texto."),
    ("Razonamiento adaptativo:", " Los criterios con subjetividad 'Alta' se procesan con reasoning_effort='medium', activando la capacidad de razonamiento extendido del modelo. Los criterios de subjetividad normal o baja usan reasoning_effort='minimal', optimizando costo y velocidad sin sacrificar calidad."),
    ("Columna 'Pregunta orientadora':", " Se añadió como guía de lectura para el evaluador humano, explícitamente excluida de la evaluación automática para no sesgar el criterio principal."),
    ("Esfuerzo paralelo controlado:", " Hasta 48 llamadas concurrentes mediante ThreadPoolExecutor, con callback de progreso en tiempo real en la interfaz Streamlit."),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

doc.add_paragraph()

add_heading(doc, "3.3 Estructura de la Salida XLSX", level=2)

add_body(doc,
    "La salida descargable fue rediseñada para servir dos audiencias diferenciadas dentro del mismo fichero:")

out_table = doc.add_table(rows=1, cols=3)
out_table.style = "Table Grid"
add_table_header_row(out_table, ["Hoja", "Audiencia", "Columnas principales"])
out_rows = [
    ("Lectura amigable", "Evaluadores y gestores de proyectos", "ID, Criterio, Resultado de valoración, Lectura rápida, Principal oportunidad de mejora, Evidencia clave, Revisión humana recomendada"),
    ("Auditoría técnica", "Revisores técnicos y desarrolladores", "Respuesta interna, Razonamiento del modelo, Evidencia completa, Status de ejecución"),
    ("Resumen", "Dirección / aprobadores", "Conteo por veredicto, criterios de alta subjetividad, tasa de error, nombre y extensión del documento"),
]
for i, row_data in enumerate(out_rows):
    add_data_row(out_table, row_data, i)

doc.add_paragraph()

add_heading(doc, "3.4 Ajustes a los Criterios de Innovación", level=2)

add_body(doc,
    "En respuesta a las observaciones del equipo técnico sobre la inconsistencia en los criterios de "
    "innovación, se realizaron los siguientes ajustes:")

for item in [
    "Redefinición de la escala SI/PARCIAL/NO para los criterios de innovación, con ejemplos orientadores diferenciados según tipo de intervención (proyectos DEDICADOS vs. de MARCO).",
    "Reducción de repeticiones en los textos de diagnóstico mediante instrucciones explícitas en el prompt del sistema para evitar parafrasear el criterio evaluado.",
    "Alineación de la definición DEDICADO/MARCO entre la rúbrica v3 y la lógica de Tab 1 existente, corrigiendo una divergencia identificada en el proceso de revisión.",
]:
    add_bullet(doc, item)

doc.add_paragraph()

add_heading(doc, "3.5 Mejoras en los Outputs de Clasificación de Recomendaciones", level=2)

add_body(doc,
    "La funcionalidad de clasificación de recomendaciones (Tabs 5–6) recibió mejoras orientadas a aumentar "
    "la claridad y utilidad de los archivos descargables:")

for item in [
    "Enriquecimiento de los campos exportados: se añadieron columnas de contexto, relevancia y prioridad sugerida en la descarga XLSX.",
    "Reducción de redundancias en las narrativas generadas, mediante ajustes de temperatura y longitud máxima de respuesta.",
    "Módulo separado oli_v6_deploy_recommendations.py que incluye ambas interfaces (español e inglés) como tabs independientes dentro de la misma aplicación.",
]:
    add_bullet(doc, item)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 4. TAREA 2 – MODULARIZACIÓN DEL CÓDIGO FUENTE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Tarea 2 – Modularización del Código Fuente", level=1)

add_body(doc,
    "La separación del código monolítico en módulos independientes fue una de las tareas centrales de la "
    "consultoría. La estrategia seguida priorizó la separación funcional limpia sobre la refactorización "
    "exhaustiva, garantizando que cada módulo sea transferible y mantenible de forma independiente.")

add_heading(doc, "4.1 Módulo de Valoración de Calidad de Proyectos (Core)", level=2)

add_body(doc, "Archivo: oli_v6_deploy_core.py", bold=True)
add_body(doc,
    "Contiene las cuatro pestañas orientadas al diseño y calidad de proyectos. Es autocontenido: "
    "se ejecuta con streamlit run oli_v6_deploy_core.py sin dependencias del módulo de recomendaciones.")

mod_table = doc.add_table(rows=1, cols=3)
mod_table.style = "Table Grid"
add_table_header_row(mod_table, ["Tab", "Nombre", "Funcionalidad"])
mod_rows = [
    ("Tab 1", "Valoración Preliminar de Calidad", "Evaluación de PRODOCs con rúbrica estructurada. Descarga XLSX con diagnóstico por criterio."),
    ("Tab 2", "Diagnóstico de Atributos Específicos", "Análisis profundo de atributos como transversalidad, género, sostenibilidad con lógica RAG jerárquico."),
    ("Tab 3", "Diagnóstico de Sostenibilidad", "Evaluación contra rúbrica de sostenibilidad (escala 0–3) con dimensiones Diseño, Implementación y Evaluación."),
    ("Tab 4", "Pregúntale a tus Documentos", "Motor RAG conversacional para consultas abiertas sobre el corpus documental del proyecto."),
]
for i, row_data in enumerate(mod_rows):
    add_data_row(mod_table, row_data, i)

doc.add_paragraph()

add_heading(doc, "4.2 Módulo de Clasificación de Recomendaciones", level=2)

add_body(doc, "Archivo: oli_v6_deploy_recommendations.py", bold=True)
add_body(doc,
    "Contiene las pestañas de análisis y clasificación de recomendaciones de evaluaciones. Desplegable "
    "de forma completamente independiente al módulo Core, con su propia gestión de estado y cache.")

rec_table = doc.add_table(rows=1, cols=3)
rec_table.style = "Table Grid"
add_table_header_row(rec_table, ["Tab", "Nombre", "Funcionalidad"])
rec_rows = [
    ("Tab 5 (ES)", "Clasificación de Recomendaciones", "Clasifica, prioriza y analiza recomendaciones de evaluaciones en español. Descarga XLSX enriquecida."),
    ("Tab 6 (EN)", "Recommendation Classification", "Versión en inglés del mismo flujo. Permite trabajar con corpus de evaluaciones en inglés."),
]
for i, row_data in enumerate(rec_rows):
    add_data_row(rec_table, row_data, i)

doc.add_paragraph()

add_heading(doc, "4.3 Motores Core sin Streamlit", level=2)

add_body(doc,
    "Como parte de la arquitectura de GPT Actions, se extrajeron los motores de evaluación en módulos "
    "completamente independientes de Streamlit, reutilizables por cualquier backend o entorno de ejecución:")

for item in [
    ("tab1_v3_core.py:", " Motor de evaluación v3. Contiene la lógica de carga de rúbrica, extracción de texto DOCX, construcción de prompts, evaluación paralela con progreso y serialización de resultados a XLSX. Importado tanto por tab1_v3.py (Streamlit) como por gpt_action_api.py (FastAPI)."),
    ("sustainability_core.py:", " Motor de diagnóstico de sostenibilidad. Mismo patrón de diseño: Streamlit-free, importado por la pestaña Streamlit y por el endpoint /sustainability/ del backend GPT."),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

doc.add_paragraph()

add_heading(doc, "4.4 Nuevo Tab 7 – Evaluación v3 Experimental en Streamlit", level=2)

add_body(doc,
    "Se implementó un séptimo tab experimental en la aplicación Streamlit (tab1_v3.py) que expone la "
    "rúbrica mecanizada v3 directamente en la interfaz existente, sin modificar el Tab 1 original. "
    "La integración requiere únicamente tres cambios en oli_v6_deploy_core.py: import del módulo, "
    "adición del tab al tuple de st.tabs() y el bloque with tab7: tab1_v3.render(client). "
    "El tab puede retirarse revertiendo esos tres cambios sin afectar el resto de la aplicación.")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 5. GPTES PERSONALIZADOS EN CHATGPT
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. GPTs Personalizados en ChatGPT", level=1)

add_body(doc,
    "Se diseñó, desarrolló y desplegó una arquitectura de GPT Actions que permite utilizar los motores "
    "de evaluación directamente desde ChatGPT, sin necesidad de acceder al entorno Streamlit. Los tres "
    "GPTs comparten un único backend FastAPI desplegado en Render (render.com), lo que minimiza costos "
    "de infraestructura y simplifica el mantenimiento.")

add_body(doc, "Arquitectura común de los GPTs:", bold=True)
for item in [
    "El usuario sube un documento DOCX directamente en el chat de ChatGPT.",
    "El GPT invoca el endpoint de inicio de trabajo (POST /v3/jobs o POST /sustainability/jobs) pasando la referencia del archivo.",
    "El backend descarga el archivo desde los servidores de OpenAI, lo procesa con el motor de evaluación y ejecuta las llamadas al modelo en paralelo.",
    "El GPT sondea el estado del trabajo hasta que el proceso finaliza (GET /jobs/{id}).",
    "Una vez completado, el GPT recupera el resultado (GET /jobs/{id}/result) y presenta al usuario un resumen narrativo junto con el archivo XLSX descargable.",
]:
    add_bullet(doc, item)

doc.add_paragraph()

add_heading(doc, "5.1 GPT de Valoración Preliminar de Calidad (Tab 1 v3)", level=2)

for item in [
    ("Nombre:", " ILO PRODOC Quality Appraisal"),
    ("Enlace:", " https://chatgpt.com/g/g-6a2643b11e708191adc1c03e64260a25-ilo-prodoc-quality-appraisal"),
    ("Endpoint:", " POST /v3/jobs → GET /v3/jobs/{id} → GET /v3/jobs/{id}/result"),
    ("Motor:", " tab1_v3_core.py + Rubrica_Tab1_Detallada_Full_v3.xlsx"),
    ("Criterios:", " 76 criterios distribuidos en 5 secciones; evaluación completa o filtrada por sección/subsección"),
    ("Escala:", " Yes / Partial / No / Not Found / N/A por criterio"),
    ("Estabilidad:", " 10 corridas por criterio (STABILITY_REPEATS) consolidadas por moda, umbral de 80%; del orden de 760 llamadas al modelo en una valoración completa"),
    ("Modelo:", " gpt-5-mini con razonamiento adaptativo (effort medium/minimal según subjetividad del criterio)"),
    ("Estado:", " Publicado con visibilidad \"Anyone with a link\"; backend operativo en https://ilo-prodoc-appraisal-v3.onrender.com"),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

doc.add_paragraph()

add_heading(doc, "5.2 GPT de Diagnóstico de Sostenibilidad (Tab 3)", level=2)

for item in [
    ("Nombre:", " OIT - Diagnóstico de Sostenibilidad del Proyecto"),
    ("Enlace:", " https://chatgpt.com/g/g-6a43d24307b88191bb362632f133c0f3-oit-diagnostico-de-sostenibilidad-del-proyecto"),
    ("Endpoint:", " POST /sustainability/jobs → GET /sustainability/jobs/{id} → GET /sustainability/jobs/{id}/result"),
    ("Motor:", " sustainability_core.py + Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx"),
    ("Escala:", " 0–3 por indicador, con dimensiones Diseño, Implementación y Pre-Cierre"),
    ("Estabilidad:", " Esquema de 5 corridas por indicador consolidadas por moda, con porcentaje de estabilidad reportado en el XLSX"),
    ("Salida XLSX:", " Dos hojas: 'Lectura amigable' para gestores y 'Auditoría técnica' para revisores"),
    ("Estado:", " Publicado con visibilidad \"Anyone with a link\"; mismo backend compartido"),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

doc.add_paragraph()

add_heading(doc, "5.3 GPT de Diagnóstico de Atributos Específicos (Tab 2)", level=2)

for item in [
    ("Nombre:", " OIT - Diagnóstico de Atributos Específicos"),
    ("Enlace:", " https://chatgpt.com/g/g-6a43cc82d4d08191bf6e60357453e336-oit-diagnostico-de-atributos-especificos"),
    ("Endpoint:", " POST /attributes/jobs → GET /attributes/jobs/{id} → GET /attributes/jobs/{id}/result"),
    ("Motor:", " tab2_core.py + Rubricas_6ago2025.xlsx"),
    ("Rúbricas seleccionables:", " métodos participativos, género o transición justa (una o varias por evaluación)"),
    ("Escala:", " Puntuación 1–5 por criterio, con esquema de estabilidad de 5 corridas para reducir variabilidad del modelo"),
    ("Modelo:", " gpt-5-mini; hasta 8 llamadas paralelas por corrida"),
    ("Estado:", " Publicado con visibilidad \"Anyone with a link\"; integrado al mismo backend compartido"),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

doc.add_paragraph()

add_heading(doc, "5.4 Infraestructura y Seguridad del Backend", level=2)

infra_table = doc.add_table(rows=1, cols=2)
infra_table.style = "Table Grid"
add_table_header_row(infra_table, ["Componente", "Detalle"])
infra_data = [
    ("Plataforma", "Render.com, free tier"),
    ("Runtime", "Docker, Python 3.10-slim"),
    ("Framework", "FastAPI + uvicorn[standard]"),
    ("Autenticacion", "X-API-Key en cabecera; /health y /privacy sin autenticacion"),
    ("Politica de privacidad", "/privacy HTML - requerido por ChatGPT para GPTs publicos"),
    ("Persistencia", "In-memory (pilot); migrar a Redis para uso institucional"),
    ("Paralelismo", "Hasta 48 llamadas OpenAI concurrentes por evaluacion"),
    ("Despliegue CI/CD", "render.yaml Blueprint; redespliegue automatico al hacer push a main"),
]
for i, row_data in enumerate(infra_data):
    add_data_row(infra_table, row_data, i)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 6. TAREA 3 – DOCUMENTACIÓN TÉCNICA
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Tarea 3 – Documentación Técnica para Transferencia", level=1)

add_body(doc,
    "Se actualizó y creó documentación técnica orientada a facilitar la transferencia, mantenimiento "
    "y futura ampliación de la aplicación por parte del equipo de la OIT.")

doc_table = doc.add_table(rows=1, cols=3)
doc_table.style = "Table Grid"
add_table_header_row(doc_table, ["Documento", "Formato", "Contenido"])
doc_table = doc.add_table(rows=1, cols=3)
doc_table.style = "Table Grid"
add_table_header_row(doc_table, ["Documento", "Formato", "Contenido"])
doc_data = [
    ("1_Documentacion_Tecnica.docx", "DOCX", "Arquitectura general, stack tecnologico, flujo de datos"),
    ("2_Documentacion_Codigo_Fuente.docx", "DOCX", "Modulos, funciones principales y dependencias"),
    ("3_Documentacion_Base_Datos.docx", "DOCX", "Ficheros de embeddings y bases de conocimiento"),
    ("4_Documentacion_Configuracion_Seguridad.docx", "DOCX", "Variables de entorno, claves API, seguridad"),
    ("5_Documentacion_Funcional_Operativa.docx", "DOCX", "Guia de uso para evaluadores y flujos de trabajo"),
    ("6_Otros_Contacto_Recomendaciones.docx", "DOCX", "Recomendaciones de evolucion futura y contacto tecnico"),
    ("docs/gpt_action_v3_setup.md", "Markdown", "Configuracion completa del GPT de valoracion v3"),
    ("docs/gpt_action_sustainability_setup.md", "Markdown", "Configuracion completa del GPT de sostenibilidad"),
]
for i, row_data in enumerate(doc_data):
    add_data_row(doc_table, row_data, i)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 7. INVENTARIO DE ENTREGABLES
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Inventario de Entregables", level=1)

add_body(doc, "A. Aplicación Streamlit", bold=True)
for f in [
    "oli_v6_deploy_core.py – Módulo Streamlit: valoración de calidad de proyectos (Tabs 1–4)",
    "oli_v6_deploy_recommendations.py – Módulo Streamlit: clasificación de recomendaciones (Tabs 5–6)",
    "tab1_v3.py – Tab 7 experimental v3 (add-on para oli_v6_deploy_core.py)",
]:
    add_bullet(doc, f)

doc.add_paragraph()
add_body(doc, "B. Motores de Evaluación (sin Streamlit)", bold=True)
for f in [
    "tab1_v3_core.py – Motor de evaluación v3: rúbrica Tab 1, 76 criterios",
    "sustainability_core.py – Motor de diagnóstico de sostenibilidad",
]:
    add_bullet(doc, f)

doc.add_paragraph()
add_body(doc, "C. Backend GPT Actions", bold=True)
for f in [
    "gpt_action_api.py – FastAPI: endpoints /v3/ y /sustainability/ compartidos",
    "openapi_gpt_action_v3.yaml – Esquema OpenAPI para el GPT de valoración",
    "openapi_gpt_action_sustainability.yaml – Esquema OpenAPI para el GPT de sostenibilidad",
    "Dockerfile.gpt-action – Imagen Docker de producción",
    "docker-compose.gpt-action.yml – Composición para pruebas locales",
    "render.yaml – Blueprint de despliegue en Render",
    "requirements.gpt-action.txt – Dependencias del backend GPT",
]:
    add_bullet(doc, f)

doc.add_paragraph()
add_body(doc, "D. Rúbricas", bold=True)
for f in [
    "Rubrica_Tab1_Detallada_Full_v3.xlsx – Rúbrica mecanizada final (76 criterios)",
    "Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx – Rúbrica de sostenibilidad",
]:
    add_bullet(doc, f)

doc.add_paragraph()
add_body(doc, "E. Documentación", bold=True)
for f in [
    "documentation/ – Seis documentos técnicos de transferencia (DOCX)",
    "docs/gpt_action_v3_setup.md – Guía de configuración GPT valoración v3",
    "docs/gpt_action_sustainability_setup.md – Guía de configuración GPT sostenibilidad",
    "Informe_Final_Consultoría_OIT_PO40551829.docx – Este documento",
]:
    add_bullet(doc, f)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 8. CUMPLIMIENTO DE LAS TAREAS CONTRACTUALES
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Cumplimiento de las Tareas Contractuales", level=1)

comp_table = doc.add_table(rows=1, cols=4)
comp_table.style = "Table Grid"
add_table_header_row(comp_table, ["Tarea contractual", "Entregable(s)", "Estado", "Observaciones"])
comp_data = [
    ("Ajustes en lógica de outputs (coherencia, repeticiones, innovación, recomendaciones)",
     "tab1_v3_core.py, Rubrica_Tab1_Detallada_Full_v3.xlsx, oli_v6_deploy_recommendations.py",
     "✅ Completado",
     "Tres versiones de rúbrica; razonamiento adaptativo; outputs con dos niveles de detalle"),
    ("Modularización del código fuente (módulo calidad + módulo recomendaciones)",
     "oli_v6_deploy_core.py, oli_v6_deploy_recommendations.py",
     "✅ Completado",
     "Separación completa; cada módulo ejecutable de forma independiente"),
    ("Documentación técnica para transferencia",
     "documentation/ (6 docs), docs/ (2 guías Markdown)",
     "✅ Completado",
     "Incluye guías de despliegue, configuración y uso de GPTs"),
    ("GPT Actions (valoración v3, sostenibilidad, atributos específicos) [valor añadido]",
     "gpt_action_api.py, openapi_*.yaml, render.yaml",
     "✅ Publicado",
     "Tres GPTs en ChatGPT; backend FastAPI operativo en Render"),
]
for i, row_data in enumerate(comp_data):
    add_data_row(comp_table, row_data, i)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 9. ACCESOS A LA APLICACIÓN
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Accesos a la Aplicación", level=1)

add_body(doc,
    "Los accesos y credenciales de las aplicaciones se proporcionan directamente a los aprobadores "
    "técnicos (Cybele Burga y Carlos Castañeda) por canal seguro. A continuación se indica la "
    "naturaleza de cada acceso:")

acc_table = doc.add_table(rows=1, cols=3)
acc_table.style = "Table Grid"
add_table_header_row(acc_table, ["Aplicación", "Tipo de acceso", "Nota"])
acc_data = [
    ("Streamlit (módulo Core)", "URL de Streamlit Cloud + contraseña de entorno", "Requiere OPENAI_API_KEY configurada"),
    ("Streamlit (módulo Recomendaciones)", "URL de Streamlit Cloud + contraseña de entorno", "Requiere OPENAI_API_KEY configurada"),
    ("GPT Valoración v3", "https://chatgpt.com/g/g-6a2643b11e708191adc1c03e64260a25-ilo-prodoc-quality-appraisal", "Anyone with a link; requiere cuenta ChatGPT"),
    ("GPT Sostenibilidad", "https://chatgpt.com/g/g-6a43d24307b88191bb362632f133c0f3-oit-diagnostico-de-sostenibilidad-del-proyecto", "Anyone with a link; requiere cuenta ChatGPT"),
    ("GPT Atributos Específicos", "https://chatgpt.com/g/g-6a43cc82d4d08191bf6e60357453e336-oit-diagnostico-de-atributos-especificos", "Anyone with a link; requiere cuenta ChatGPT"),
    ("Backend API (Render)", "URL + ILO_GPT_ACTION_API_KEY", "Proporcionados por canal seguro"),
]
for i, row_data in enumerate(acc_data):
    add_data_row(acc_table, row_data, i)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 10. RECOMENDACIONES PARA CONTINUIDAD
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. Recomendaciones para Continuidad", level=1)

recs = [
    ("Migrar el almacenamiento de jobs a Redis o base de datos",
     "El backend piloto guarda los trabajos en memoria. Para uso institucional con múltiples usuarios simultáneos, se recomienda reemplazar el dict JOBS por una capa de persistencia (Redis, PostgreSQL). Esfuerzo estimado: 1–2 días."),
    ("Establecer un límite de gasto en la API de OpenAI",
     "Cada evaluación completa consume ~76 llamadas al modelo. Se recomienda configurar un límite mensual en la cuenta OpenAI para evitar gasto no controlado si el enlace del GPT se distribuye ampliamente."),
    ("Elevar el plan de Render a Starter si hay uso intensivo",
     "El free tier duerme tras 15 minutos de inactividad (inicio en frío ~30–60s). Para equipos que usen el GPT en horario laboral continuo, el plan Starter ($7/mes) elimina las esperas y garantiza disponibilidad 24/7."),
    ("Migrar el repositorio GitHub a privado",
     "El repositorio de despliegue actualmente es público y contiene ficheros de proyecto de la OIT. Se recomienda convertirlo a privado; Render soporta repositorios privados sin costo adicional."),
    ("Evaluar la incorporación de OAuth por usuario",
     "El esquema actual de clave compartida no permite distinguir qué usuario realizó cada evaluación. Para auditoría institucional, implementar OAuth (OpenID Connect con la cuenta ILO) permitiría registro por usuario."),
]
for title, desc in recs:
    add_bullet(doc, f" {desc}", bold_prefix=f"{title}:")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    f"Informe preparado por Ahmed Guillermo Eid Valdiviezo · Contrato OIT PO 40551829 · "
    f"Entregado el 8 de junio de 2026"
)
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUTPUT)
print(f"Report saved: {OUTPUT}")
