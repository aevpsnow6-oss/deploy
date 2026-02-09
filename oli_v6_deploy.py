import concurrent.futures
import pickle
import streamlit as st
import pandas as pd
import json
import tempfile
import docx
import numpy as np
import faiss
import os
import re
from matplotlib import pyplot as plt
import seaborn as sns
from io import BytesIO
import torch
import time
import plotly.express as px
import plotly.graph_objects as go
from docx2python import docx2python
from io import BytesIO
import streamlit as st
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import zipfile

# --- Utility function for Excel export ---
def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
        df.to_excel(writer, index=False, sheet_name='Datos Filtrados')
    processed_data = output.getvalue()
    return processed_data

# Use environment variables for API keys
# For local development - use .env file or set environment variables
# For Streamlit Cloud - set these in the app settings
openai_api_key = os.getenv("OPENAI_API_KEY")
# Initialize OpenAI with new SDK (v1.0.0+)
from openai import OpenAI
client = OpenAI(api_key=openai_api_key)

os.environ["KMP_DUPLICATE_LIB_OK"]="TRUE"

# Import tiktoken for accurate token counting
try:
    import tiktoken
    # Use cl100k_base encoding (standard for GPT-4 and GPT-5 models)
    encoding = tiktoken.get_encoding("cl100k_base")
except ImportError:
    encoding = None
    # Warning will be shown when function is first called if needed

# Helper function to truncate text to a token limit
def truncate_to_token_limit(text, max_tokens=110000, encoding_obj=None):
    """
    Truncate text to a maximum number of tokens.
    
    Args:
        text: The text to truncate
        max_tokens: Maximum number of tokens (default 110K for GPT-5-mini's 128K context window)
        encoding_obj: tiktoken encoding object (if None, falls back to character estimation)
    
    Returns:
        Truncated text that fits within the token limit
    """
    if not text:
        return text
    
    if encoding_obj is None:
        # Fallback: use character-based estimation (4 chars per token for Spanish)
        # 110K tokens * 4 = 440K characters
        max_chars = max_tokens * 4
        return text[:max_chars]
    
    # Count tokens
    tokens = encoding_obj.encode(text)
    
    # If within limit, return full text
    if len(tokens) <= max_tokens:
        return text
    
    # Truncate to max_tokens
    truncated_tokens = tokens[:max_tokens]
    truncated_text = encoding_obj.decode(truncated_tokens)
    
    return truncated_text

# Function to get embeddings from OpenAI
def get_embedding_with_retry(text, model='text-embedding-3-large', max_retries=3, delay=1):
    if not openai_api_key:
        st.error("No se encontró la clave de API de OpenAI. Por favor, configura la variable de entorno OPENAI_API_KEY.")
        return None
    for attempt in range(max_retries):
        try:
            response = client.embeddings.create(input=text, model=model)
            return np.array(response.data[0].embedding)
        except Exception as e:
            st.warning(f"Intento {attempt + 1} fallido: {str(e)}")
            time.sleep(delay)
    return None

# Function to find similar recommendations using embeddings

def find_similar_recommendations(query_embedding, index, doc_embeddings, structured_embeddings, score_threshold=0.5, top_n=20):
    # Normalize query embedding for cosine similarity
    query_embedding = np.array(query_embedding).reshape(1, -1)
    # Search the index
    try:
        distances, indices = index.search(query_embedding, index.ntotal)

        # Filter results based on the score threshold
        filtered_recommendations = []
        for idx, dist in zip(indices[0], distances[0]):
            if idx < len(structured_embeddings) and dist >= score_threshold:
                metadata = structured_embeddings[idx]
                recommendation = {
                    "recommendation": metadata["text"],
                    "similarity": float(dist),  # Convert to float for JSON serialization
                    "country": metadata["country"],
                    "year": metadata["year"],
                    "eval_id": metadata["eval_id"]
                }
                filtered_recommendations.append(recommendation)
            if len(filtered_recommendations) >= top_n:
                break
        return filtered_recommendations
    except Exception as e:
        st.error(f"Error en la búsqueda por similitud: {str(e)}")
        return []

# Function to find recommendations by term matching

def find_recommendations_by_term_matching(query, doc_texts, structured_embeddings, top_n=10):
    try:
        matched_recommendations = []
        query_lower = query.lower()
        for idx, text in enumerate(doc_texts):
            if isinstance(text, str) and query_lower in text.lower():
                if idx < len(structured_embeddings):
                    metadata = structured_embeddings[idx]
                    matched_recommendations.append({
                        "recommendation": text,
                        "country": metadata["country"],
                        "year": metadata["year"],
                        "eval_id": metadata["eval_id"]
                    })
        matched_recommendations = sorted(matched_recommendations, key=lambda x: len(str(x["recommendation"])))
        return matched_recommendations[:top_n]
    except Exception as e:
        st.error(f"Error en la coincidencia de términos: {str(e)}")
        return []

# Function to generate executive summary using LLM
def generate_executive_summary(recommendations_text, max_output_tokens=3000):
    """
    Generates an executive summary from a list of recommendations using OpenAI.
    """
    if not recommendations_text:
        return "No hay recomendaciones para resumir."

    # Truncate to avoid context limit (approx 120k chars is safe for gpt-4o-mini's 128k tokens, keeping room for output)
    # Using existing helper
    truncated_input = truncate_to_token_limit(recommendations_text, max_tokens=100000, encoding_obj=encoding)

    system_prompt = """Eres un asistente experto en análisis de evaluaciones de proyectos de desarrollo. 
    Tu tarea es generar un Resumen Ejecutivo exhaustivo, informativo y detallado basado en las recomendaciones, respuestas de gestión y planes de acción proporcionados.
    
    El resumen debe:
    1. Identificar los temas principales y patrones recurrentes.
    2. Resaltar las acciones críticas sugeridas.
    3. Incorporar la perspectiva de la respuesta de gestión si está disponible.
    4. Estar estructurado con títulos claros y puntos clave (bullet points).
    5. Ser profesional y directo.
    6. Estar en Español.
    """
    
    user_prompt = f"Aquí están los datos de las recomendaciones (incluyendo descripción, respuesta de gestión, comentarios, etc.):\n\n{truncated_input}\n\nPor favor, genera el Resumen Ejecutivo ahora."

    if not openai_api_key:
        return "Error: No se encontró la clave API de OpenAI."

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini", # Cost-effective model
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=max_output_tokens
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error generando el resumen: {str(e)}"

def verify_match_with_llm(rec_text, candidates, api_key):
    """
    Uses LLM to verify which of the candidates is the best match for the recommendation.
    candidates: list of dicts {'dimension':..., 'subdim':..., 'definition':...}
    Returns: index of the best match in candidates list, or -1 if none fit well.
    """
    if not api_key: 
        return 0 # Fallback to top 1 if no key
        
    candidates_str = ""
    for i, c in enumerate(candidates):
        # We assume 'texto_merged' acts as the definition/description
        candidates_str += f"Option {i+1}:\nCategory: {c['dimension']} - {c['subdim']}\nDefinition/Context: {c['texto_merged']}\n\n"
        
    system_prompt = """You are an expert in classifying development project recommendations.
    You will be given a 'Recommendation' and a list of 'Options' (categories with definitions).
    Your task is to select the Option that BEST fits the Recommendation based on the definition.
    
    CRITICAL INSTRUCTIONS:
    1. Read the Recommendation and the Definitions carefully.
    2. **STRICT SEMANTIC MATCH**: The recommendation must describe the *same action* and *same object* as the definition.
    3. **IGNORE SUPERFICIAL KEYWORDS**: Do not select an option just because it shares a word (e.g., "Pensions") if the *context* is different (e.g., "Political backing" vs "Financial funding").
    4. **REQUIREMENT CHECK**: If a definition requires a specific mechanism (e.g., "Financing", "Legislation", "Training"), the recommendation MUST explicitly mention it.
    5. Return ONLY the number of the best option (1, 2, 3...).
    6. If none of them fit reasonably well, return 0.
    """
    
    user_prompt = f"Recommendation: \"{rec_text}\"\n\n{candidates_str}\n\nWhich option is the best fit? (Return just the number, e.g. 1)"

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.0, # Deterministic
            max_tokens=6
        )
        content = response.choices[0].message.content.strip()
        # Parse number
        import re
        match = re.search(r'\d+', content)
        if match:
            val = int(match.group())
            if 1 <= val <= len(candidates):
                return val - 1 # 0-indexed
            else:
                return -1 # None or 0
        return 0 # Fallback to Top 1 if parse fails
    except:
        return 0 # Fallback to Top 1 on error

# --- Analysis Cache (for Deep Analysis) ---
class AnalysisCache:
    def __init__(self, cache_file="analysis_cache.pkl"):
        self.cache_file = cache_file
        self.cache = {}
        self.load()

    def load(self):
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'rb') as f:
                    self.cache = pickle.load(f)
            except:
                self.cache = {}

    def save(self):
        with open(self.cache_file, 'wb') as f:
            pickle.dump(self.cache, f)

    def get(self, key):
        return self.cache.get(key)
    
    def set(self, key, value):
        self.cache[key] = value

    def generate_key(self, rec, plan):
        # Create a hash based on rec + plan content
        content = f"{str(rec).strip()}|{str(plan).strip()}"
        return content 

# --- Deep Analysis Logic ---
def analyze_recommendation_plan_pair(recommendation, action_plan, comments, client, model="gpt-4o-mini"):
    """
    Send recommendation, action plan, and comments to OpenAI API for analysis
    """
    if pd.isna(recommendation) or pd.isna(action_plan):
        return None # Skip empty
    
    # Handle NaN comments
    if pd.isna(comments):
        comments = "No additional comments provided."
    
    # Define ILO development project relevant tags
    ilo_tags_list = [
        "capacity_building", "training", "employment_creation", "gender_equality", 
        "decent_work", "labor_rights", "social_dialogue", "social_protection",
        "occupational_safety", "child_labor", "forced_labor", "labor_migration",
        "working_conditions", "skills_development", "social_inclusion", "monitoring_evaluation",
        "institutional_strengthening", "policy_development", "knowledge_management",
        "sustainability", "stakeholder_engagement", "data_collection", "technical_assistance",
        "project_design", "implementation_methodology", "resource_allocation", 
        "coordination", "partnership_building", "innovation", "digital_transformation"
    ]
    
    rec_tags_list = ['governance', 'participation', 'gender_issues', 'just_transition', 'institutional_strenghtening', 
                'public_policy_incidence', 'financial_sustainability']
    
    prompt = f"""
    Analyze the following recommendation, its corresponding action plan, and additional comments:
    
    RECOMMENDATION:
    {recommendation}
    
    ACTION PLAN:
    {action_plan}
    
    ADDITIONAL COMMENTS:
    {comments}
    
    Please extract and return ONLY a JSON object with the following fields:
    1. extracted_actions_from_rec: List all specific actions requested in the recommendation
    2. actions_proposed_in_plan: List all specific actions mentioned in the action plan
    3. difficulties_mentioned: Any difficulties or challenges mentioned in implementing the recommendation (look carefully in the action plan AND comments for these)
    4. reasons_for_rejection: If the recommendation wasn't fully accepted, reasons given (pay special attention to justifications in the comments)
    5. rejection_difficulty_classification: Classify the reasons (Financial, Technical, Political, Low priority, Unjustified, Third party dependency, Time constraints, Cultural/behavioral, Local operational constraints, Mandate limitations, Other). At most 3.
    6. coherence_score: Score from 0-10 how well the action plan addresses the recommendation
    7. coherence_rationale: Detailed explanation (6-8 sentences).
    8. plan_quality_score: Score from 0-10 (specificity, feasibility).
    9. plan_quality_rationale: Detailed explanation (6-8 sentences).
    10. attention_level_score: Score from 0-10 (priority given).
    11. attention_level_rationale: Detailed explanation (6-8 sentences).
    12. overall_score: Score from 0-10.
    13. overall_score_rationale: Extensive analysis (6-8 sentences).
    14. tags: Select 2-5 most relevant tags from: {", ".join(ilo_tags_list)}
    
    Now, analyze the recommendation on its own:
    15. rec_innovation_score: (Very low, Low, Medium, High, Very High).
    16. rec_innovation_rationale: Explanation (3-4 sentences).
    17. rec_precision_and_clarity: (Very low, Low, Medium, High, Very High).
    18. rec_precision_and_clarity_rationale: Explanation (3-4 sentences).
    19. rec_additional_tags: Select 2-5 relevant tags from: {", ".join(rec_tags_list)}
    20. rec_operational_feasibility: (Very low, Low, Medium, High, Very High).
    21. rec_operational_feasibility_rationale: Explanation (3-4 sentences).
    22. rec_timeline: (short, medium, long).
    23. rec_timeline_rationale: Explanation (3-4 sentences).
    24. rec_expected_impact: (Very low, Low, Medium, High, Very High).
    25. rec_expected_impact_rationale: Explanation (3-4 sentences).
    26. rec_intervention_approach: (processes, results, policy).
    27. rec_intervention_approach_rationale: Explanation (3-4 sentences).
    
    Respond ONLY with the JSON object.
    """
    
    try:
        response = client.chat.completions.create(
            model=model,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You analyze recommendations and action plans, returning structured JSON results."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        
        return json.loads(response.choices[0].message.content)
    
    except Exception as e:
        return {
            "error": str(e),
            "coherence_score": 0,
            "overall_score": 0
        }

def run_row_analysis(args):
    """
    Worker for parallel analysis
    """
    idx, rec, plan, comments, api_key, cache_obj = args
    
    # Check cache first
    key = cache_obj.generate_key(rec, plan)
    cached = cache_obj.get(key)
    if cached:
        return idx, cached, True # True = from cache
        
    # If not in cache, call API
    client = OpenAI(api_key=api_key)
    result = analyze_recommendation_plan_pair(rec, plan, comments, client)
    
    return idx, result, False # False = new call

# --- Begin: SimpleHierarchicalStore and RAG logic from megaparse_example.py ---
import pickle
from typing import List, Dict, Any
import numpy as np
import json
import os
import openai

class SimpleHierarchicalStore:
    def __init__(self, use_cache=True, cache_dir=None):
        self.documents = {}
        self.sections = {}
        self.paragraphs = []
        self.use_cache = use_cache
        self.cache_dir = cache_dir or os.getcwd()
        self.embedding_cache = {}
        self.query_cache = {}
        self.storage_dir = cache_dir or os.path.join(os.path.expanduser("~"), "document_store")
        os.makedirs(self.storage_dir, exist_ok=True)
        self.cache_file = os.path.join(self.storage_dir, "embedding_cache.pkl")
        if use_cache:
            self._load_cache()
    def _load_cache(self):
        if not self.use_cache:
            return
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'rb') as f:
                    self.embedding_cache = pickle.load(f)
            except Exception:
                self.embedding_cache = {}
        else:
            self.embedding_cache = {}
    def _hash_text(self, text):
        import hashlib
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    def _save_cache(self):
        if not self.use_cache:
            return
        try:
            with open(self.cache_file, 'wb') as f:
                pickle.dump(self.embedding_cache, f)
        except Exception:
            pass
    def get_embedding(self, text: str):
        if not text or text.isspace():
            return [0.0] * 1536
        if self.use_cache:
            text_hash = hash(text)
            if text_hash in self.embedding_cache:
                return self.embedding_cache[text_hash]
        try:
            response = openai.embeddings.create(
                input=text,
                model="text-embedding-3-small"
            )
            embedding = response.data[0].embedding
            if self.use_cache:
                self.embedding_cache[hash(text)] = embedding
                if len(self.embedding_cache) % 100 == 0:
                    self._save_cache()
            return embedding
        except Exception:
            return [0.0] * 1536
    def add_documents(self, df, content_column='content', section_column='header_1', batch_size=20):
        doc_id = df['document_id'].iloc[0] if 'document_id' in df.columns else 'doc1'
        self.documents[doc_id] = {'embedding': self.get_embedding(' '.join(df[content_column].astype(str).tolist()))}
        for _, row in df.iterrows():
            section_id = row.get(section_column, '')
            if pd.isna(section_id):
                section_id = '_default_section'
            section_text = str(row.get(content_column, ''))
            if not section_text.strip():
                continue
            section_embedding = self.get_embedding(section_text)
            self.sections[(doc_id, section_id)] = {
                'text': section_text,
                'embedding': section_embedding
            }
            self.paragraphs.append({
                'text': section_text,
                'embedding': section_embedding,
                'document_id': doc_id,
                'section_id': section_id,
                'position': row.get('paragraph_number', 0)
            })
    def cosine_similarity(self, embedding1, embedding2):
        if not embedding1 or not embedding2:
            return 0
        dot_product = sum(a * b for a, b in zip(embedding1, embedding2))
        norm1 = sum(a * a for a in embedding1) ** 0.5
        norm2 = sum(b * b for b in embedding2) ** 0.5
        return dot_product / (norm1 * norm2) if norm1 > 0 and norm2 > 0 else 0
    def score_rubric_directly(self, rubric_elements: Dict, top_n_paragraphs: int = 10) -> Dict:
        print(f"[score_rubric_directly] Evaluating criterion: {rubric_elements}")
        results = {}
        for criterion, descriptions in rubric_elements.items():
            print(f"[score_rubric_directly] Evaluating criterion: {criterion}")
            criterion_embedding = self.get_embedding(criterion)
            paragraph_scores = []
            for p in self.paragraphs:
                similarity = self.cosine_similarity(criterion_embedding, p['embedding'])
                paragraph_scores.append((p, similarity))
            paragraph_scores.sort(key=lambda x: x[1], reverse=True)
            top_paragraphs = paragraph_scores[:top_n_paragraphs]
            context_text = '\n\n---\n\n'.join([p[0]['text'] for p in top_paragraphs])
            try:
                analysis = self.analyze_criterion(criterion, context_text, descriptions)
                print(f"[score_rubric_directly] Analysis result for '{criterion}': {analysis}")
                results[criterion] = {
                    'analysis': analysis,
                    'context': context_text,
                    'score': analysis.get('score', 0),
                    'confidence': analysis.get('confidence', 0),
                    'top_paragraphs': [{'text': p[0]['text'], 'similarity': p[1]} for p in top_paragraphs[:3]]
                }
            except Exception as e:
                print(f"[score_rubric_directly] Exception for '{criterion}': {e}")
                results[criterion] = {
                    'analysis': {'error': str(e)},
                    'context': context_text,
                    'score': 0,
                    'confidence': 0
                }
        print(f"[score_rubric_directly] Final results: {results}")
        return results
    def analyze_criterion(self, criterion: str, context: str, descriptions: list) -> dict:
        print(f"[analyze_criterion] Called with criterion: {criterion}")
        print(f"[analyze_criterion] Context: {context[:200]} ...")
        print(f"[analyze_criterion] Descriptions: {descriptions}")
        prompt = f"""
        You are evaluating a document against a specific criterion. 
        Criterion: {criterion}
        Descriptions of scoring levels:
        {json.dumps(descriptions, indent=2)}
        Document content to evaluate:
        {context}
        Please analyze how well the document meets this criterion. Provide:
        1. A detailed analysis (2-3 paragraphs)
        2. A score from 1-5 (where 1 is lowest and 5 is highest)
        3. Key evidence from the document that supports your score
        4. Any recommendations for improvement
        5. A confidence level (0-1) indicating how confident you are in this assessment
        Format your response as a JSON object with the following keys:
        {"analysis": "your detailed analysis here", "score": numeric_score_between_1_and_5, "evidence": "key evidence from the document", "recommendations": "your recommendations for improvement", "confidence": confidence_level_between_0_and_1}
        Return only the JSON object, nothing else.
        """
        try:
            response = openai.chat.completions.create(
                model="gpt.5.1-mini",
                messages=[
                    {"role": "system", "content": "You are an expert document evaluator that provides detailed analysis and scoring based on specific criteria."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"}
            )
            raw = response.choices[0].message.content.strip()
            print(f"[analyze_criterion] Raw response: {raw}")
            parsed = json.loads(raw)
            print(f"[analyze_criterion] Parsed response: {parsed}")
            return parsed
        except Exception as e:
            print(f"[analyze_criterion] Exception: {e}")
            return {'score': 0, 'analysis': f'Error: {str(e)}'}
# --- End: SimpleHierarchicalStore and RAG logic ---

# --- HIERARCHICAL RAG RUBRIC EVALUATION ---
def add_rubric_evaluation_section(exploded_df, toc, toc_hierarchy):
    """
    Add a new section for rubric-based evaluation of the document using hierarchical RAG pipeline.
    Allows users to select rubric type, choose criteria, evaluate selected sections, view results, and download as CSV/Excel.
    """
    import streamlit as st
    import pandas as pd
    from io import BytesIO
    from collections import defaultdict

    st.markdown("### Evaluación por Rúbrica (Hierarchical RAG)")

    # Initialize session state keys with unique prefixes to avoid conflicts
    session_keys = {
        'selected_sections': 'rubric_eval_selected_sections',
        'sections_confirmed': 'rubric_eval_sections_confirmed',
        'selected_criteria': 'rubric_eval_selected_criteria'
    }

    # --- Section Selection ---
    if session_keys['selected_sections'] not in st.session_state:
        st.session_state[session_keys['selected_sections']] = []
    if session_keys['sections_confirmed'] not in st.session_state:
        st.session_state[session_keys['sections_confirmed']] = False

    st.markdown("#### 1. Selección de Secciones para Evaluación")
    main_sections = []
    for level, headings in sorted(toc_hierarchy.items()):
        if headings and not main_sections:
            main_sections = headings
    valid_selected_sections = [s for s in st.session_state[session_keys['selected_sections']] if s in main_sections]
    if not valid_selected_sections and main_sections:
        valid_selected_sections = [main_sections[0]]
    selected_sections = st.multiselect(
        "Seleccione las secciones del documento que desea evaluar:",
        options=main_sections,
        default=valid_selected_sections,
        key="rubric_section_multiselect_eval"
    )
    if st.button("Confirmar Secciones para Evaluación"):
        if not selected_sections:
            st.warning("Por favor seleccione al menos una sección para evaluar.")
        else:
            st.session_state[session_keys['selected_sections']] = selected_sections
            st.session_state[session_keys['sections_confirmed']] = True
            st.success(f"Secciones confirmadas para evaluación: {', '.join(selected_sections)}")
    if st.session_state[session_keys['selected_sections']]:
        st.info(f"Secciones actualmente seleccionadas para evaluación: {', '.join(st.session_state[session_keys['selected_sections']])}")

    # --- Rubric Selection & Evaluation ---
    if st.session_state[session_keys['sections_confirmed']]:
        st.markdown("#### 2. Selección de Rúbrica y Criterios")
        rubric_type = st.selectbox(
            "Seleccione tipo de rúbrica para evaluación:",
            ["Participación (Engagement)", "Desempeño (Performance)"],
            index=0,
            key='rubric_type_tab2'
        )
        if rubric_type == "Participación (Engagement)":
            rubric_df = load_engagement_rubric()
            criteria_col = "Criterion"
            short_col = "crit_short"
            group_col = "Criterio"
        else:
            rubric_df = load_performance_rubric()
            criteria_col = "subdim"
            short_col = None
            group_col = "dimension"
        st.markdown("##### Estructura de la Rúbrica")
        rubric_groups = rubric_df.groupby(group_col)
        cols = st.columns([1, 3])
        with cols[0]:
            st.markdown("**Categorías**")
            categories = list(rubric_groups.groups.keys())
            selected_category = st.radio(
                "Seleccione una categoría:",
                categories,
                label_visibility="collapsed"
            )
        with cols[1]:
            st.markdown("**Criterios**")
            if selected_category:
                category_criteria = rubric_df[rubric_df[group_col] == selected_category]
                if session_keys['selected_criteria'] not in st.session_state:
                    st.session_state[session_keys['selected_criteria']] = {}
                all_criteria_in_category = st.checkbox(f"Seleccionar todos los criterios en '{selected_category}'")
                selected_criteria_ids = []
                for _, criterion_row in category_criteria.iterrows():
                    criterion_id = criterion_row[criteria_col]
                    criterion_name = criterion_row[short_col] if short_col and short_col in criterion_row else criterion_id
                    if criterion_id not in st.session_state[session_keys['selected_criteria']]:
                        st.session_state[session_keys['selected_criteria']][criterion_id] = False
                    if all_criteria_in_category:
                        st.session_state[session_keys['selected_criteria']][criterion_id] = True
                    is_selected = st.checkbox(
                        criterion_name,
                        value=st.session_state[session_keys['selected_criteria']][criterion_id],
                        key=f"criterion_{criterion_id}"
                    )
                    st.session_state[session_keys['selected_criteria']][criterion_id] = is_selected
                    if is_selected:
                        selected_criteria_ids.append(criterion_id)
        with st.expander("Opciones avanzadas de selección de criterios"):
            select_all_criteria = st.checkbox("Seleccionar TODOS los criterios de todas las categorías")
            if select_all_criteria:
                for _, row in rubric_df.iterrows():
                    criterion_id = row[criteria_col]
                    st.session_state[session_keys['selected_criteria']][criterion_id] = True
                st.success("Todos los criterios han sido seleccionados.")
            if st.button("Limpiar todas las selecciones"):
                for criterion_id in st.session_state[session_keys['selected_criteria']]:
                    st.session_state[session_keys['selected_criteria']][criterion_id] = False
                st.success("Todas las selecciones han sido limpiadas.")
        if st.button("Ver Detalles de Criterios Seleccionados"):
            selected_any = any(st.session_state[session_keys['selected_criteria']].values())
            if not selected_any:
                st.warning("Por favor seleccione al menos un criterio para ver sus detalles.")
            else:
                st.markdown("##### Detalles de Criterios Seleccionados")
                selected_criteria_df = rubric_df[rubric_df[criteria_col].isin(
                    [cid for cid, selected in st.session_state[session_keys['selected_criteria']].items() if selected]
                )]
                for _, criterion_row in selected_criteria_df.iterrows():
                    criterion_id = criterion_row[criteria_col]
                    criterion_name = criterion_row[short_col] if short_col and short_col in criterion_row else criterion_id
                    with st.expander(f"{criterion_name}", expanded=True):
                        levels_df = rubric_to_levels_df(criterion_row, criteria_col)
                        st.table(levels_df)
        if st.button("Iniciar Evaluación de Criterios Seleccionados (RAG)"):
            selected_any = any(st.session_state[session_keys['selected_criteria']].values())
            if not selected_any:
                st.warning("Por favor seleccione al menos un criterio para evaluar.")
            else:
                st.markdown("#### 3. Evaluación de Criterios (RAG)")
                selected_criteria_ids = [cid for cid, selected in st.session_state[session_keys['selected_criteria']].items() if selected]
                filtered_df = exploded_df[exploded_df['header_1'].isin(st.session_state[session_keys['selected_sections']])].copy()
                if filtered_df.empty:
                    st.warning("No se encontraron párrafos en las secciones seleccionadas.")
                    return
                rubric_dict = {}
                for cid in selected_criteria_ids:
                    crit_row = rubric_df[rubric_df[criteria_col] == cid].iloc[0]
                    levels = rubric_to_levels_df(crit_row, criteria_col)
                    rubric_dict[cid] = levels['Description'].tolist()
                st.info(f"Evaluando {len(selected_criteria_ids)} criterios sobre {len(filtered_df)} párrafos.")

                # Progress bar for rubric evaluation (per section and per criterion)
                section_list = list(filtered_df['header_1'].unique())
                total_steps = len(section_list) * len(selected_criteria_ids)
                progress_bar = st.progress(0, text="Progreso de evaluación por rúbrica")
                progress_count = 0
                results = {}
                store = SimpleHierarchicalStore(use_cache=True)
                filtered_df['document_id'] = 'doc1'  # Single doc context
                store.add_documents(filtered_df)

                # Evaluate per section and per criterion
                for section in section_list:
                    section_df = filtered_df[filtered_df['header_1'] == section]
                    for cid in selected_criteria_ids:
                        rubric_dict_single = {cid: rubric_dict[cid]}
                        # Use top_n_paragraphs=10 or as appropriate
                        result = store.score_rubric_directly(rubric_dict_single, top_n_paragraphs=10)
                        # result is a dict keyed by cid
                        if cid not in results:
                            results[cid] = {'score': 0, 'context': '', 'analysis': {}}
                        # Merge/accumulate results per criterion
                        if cid in result:
                            # If multiple sections, you may want to aggregate or just take the last/first
                            # Here, we take the last section's result for simplicity
                            results[cid] = result[cid]
                        progress_count += 1
                        progress_bar.progress(progress_count / total_steps, text=f"Evaluando sección '{section}' y criterio '{cid}'")
                progress_bar.empty()

                eval_rows = []
                for cid, result in results.items():
                    criterion_name = cid
                    if short_col:
                        crit_row = rubric_df[rubric_df[criteria_col] == cid].iloc[0]
                        criterion_name = crit_row[short_col] if short_col in crit_row else cid
                    score = result.get('score', 0)

                    if create_filtered:
                        filtered_sections = {}
                        progress_bar = st.progress(0, text="Progreso de filtrado de secciones")
                        total_sections = len(selected_sections)
                        for idx, section in enumerate(selected_sections):
                            if section in sections_content:
                                filtered_sections[section] = sections_content[section]
                            progress_bar.progress((idx + 1) / total_sections, text=f"Filtrando sección '{section}' ({idx + 1}/{total_sections})")
                        progress_bar.empty()
                        st.session_state.filtered_sections = filtered_sections
                        # Convert to a dataframe for Excel export
                        filtered_data = []
                        for section, paragraphs in filtered_sections.items():
                            # Get the section level from TOC
                            section_level = 0
                            for heading, level in toc:
                                if heading == section:
                                    section_level = level
                                    break
                            # Process paragraphs based on content type
                            in_table = False
                            table_content = []
                            table_rows = []
                            header_row = None
                            for text in paragraphs:
                                if text == '[TABLE_START]':
                                    in_table = True
                                    table_content = []
                                    table_rows = []
                                    header_row = None
                                elif text == '[TABLE_END]':
                                    in_table = False
                                    # Process collected table content - store the processed table
                                    if table_rows:
                                        # Create a JSON representation of the table
                                        table_data = {
                                            'header': header_row if header_row else [],
                                            'rows': table_rows
                                        }
                                        filtered_data.append({
                                            'section': section,
                                            'level': section_level,
                                            'content_type': 'table',
                                            'text': json.dumps(table_data)
                                        })
                                elif text.startswith('[TABLE_HEADER]'):
                                    # Process header row
                                    cells = text[14:].split('|')
                                    header_row = cells
                                    table_content.append(text)
                                elif text.startswith('[TABLE_ROW]'):
                                    # Process data row
                                    cells = text[11:].split('|')
                                    table_rows.append(cells)
                                    table_content.append(text)
                                else:
                                    # Regular paragraph
                                    filtered_data.append({
                                        'section': section,
                                        'level': section_level,
                                        'content_type': 'paragraph',
                                        'text': text
                                    })
                                filtered_df = pd.DataFrame(filtered_data)
                                
                                if not filtered_df.empty:
                                    st.success(f"Salida filtrada creada con {len(filtered_df)} elementos de {len(filtered_sections)} secciones.")
                                    
                                    # Show a preview
                                    with st.expander("Vista Previa del Contenido Filtrado"):
                                        st.dataframe(filtered_df[['section', 'level', 'content_type', 'text']])
                                    
                                    # Download button for the filtered document
                                    excel_data = BytesIO()
                                    with pd.ExcelWriter(excel_data, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                                        filtered_df.to_excel(writer, index=False)
                                    excel_data.seek(0)
                                    
                                    st.download_button(
                                        label="Descargar Documento Filtrado",
                                        data=excel_data,
                                        file_name="filtered_document.xlsx",
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                    )
                                else:
                                    st.warning("La salida filtrada está vacía. Por favor seleccione al menos una sección con contenido.")
                            else:
                                st.warning("Por favor seleccione al menos una sección para crear una salida filtrada.")
                        else:
                            st.warning("No se encontraron secciones con encabezados en el documento.")
                    else:
                        st.warning("No se encontró contenido en el documento.")
                
                # Tab 3: Rubric Evaluation
                with doc_tabs[2]:
                    try:
                        # Call the rubric evaluation function with the document sections
                        if 'add_rubric_evaluation_section' in globals():
                            add_rubric_evaluation_section(sections_content, toc, toc_hierarchy)
                        else:
                            st.info("La función de evaluación por rúbrica no está disponible. Por favor actualice el código con la implementación de esta función.")
                    except Exception as e:
                        st.error(f"Error processing document: {str(e)}")
                        import traceback
                        st.error(traceback.format_exc())

# ============= VISUALIZATION FUNCTIONS =============

# Function to prepare additional data for new visualizations
def prepare_additional_data(df):
    """
    Prepare additional columns needed for the new visualizations.
    This should be called after the original data loading.
    """
    # Make sure we have datetime and year
    if 'Recommendation_date' in df.columns:
        df['year'] = pd.to_datetime(df['Recommendation_date']).dt.year
        df['year'] = df['year'].fillna(2023).astype(int)
    
    # Standardize categorical fields (if they exist in the dataframe)
    categorical_mappings = {
        'rec_innovation_score': {
            'very high': 'Very High', 'high': 'High', 'High': 'High',
            'medium': 'Medium', 'Medium': 'Medium',
            'low': 'Low', 'very low': 'Very Low'
        },
        'rec_precision_and_clarity': {
            'high': 'High', 'High': 'High',
            'medium': 'Medium', 'Medium': 'Medium',
            'low': 'Low', 'Low': 'Low'
        },
        'rec_expected_impact': {
            'high': 'High', 'High': 'High',
            'medium': 'Medium', 'Medium': 'Medium',
            'low': 'Low', 'Low': 'Low'
        },
        'rec_intervention_approach': {
            'policy': 'Policy', 'Policy': 'Policy',
            'process': 'Process', 'Process': 'Process'
        },
        'rec_operational_feasibility': {
            'high': 'High', 'High': 'High',
            'medium': 'Medium', 'Medium': 'Medium',
            'low': 'Low', 'Low': 'Low'
        },
        'rec_timeline': {
            'short': 'Short', 'Short': 'Short',
            'medium': 'Medium', 'Medium': 'Medium',
            'long': 'Long', 'Long': 'Long'
        }
    }
    
    # Apply standardization to columns that exist in the dataframe
    for col, mapping in categorical_mappings.items():
        if col in df.columns:
            df[col] = df[col].replace(mapping)
    
    # Process tags if they exist
    if 'tags' in df.columns:
        # Convert tags to list format if needed
        df['tags'] = df['tags'].apply(
            lambda x: [x] if isinstance(x, str) and not pd.isna(x) else 
                     x if isinstance(x, list) else []
        )
    
    # Process rejection difficulty classification if it exists
    if 'rejection_difficulty_classification' in df.columns:
        # Clean brackets and prepare for analysis
        df['clean_classification'] = df['rejection_difficulty_classification'].apply(
            lambda x: x.replace('[', '').replace(']', '').strip() if isinstance(x, str) else ""
        )
        
        # Split tags
        df['clean_tags'] = df['clean_classification'].apply(
            lambda x: [tag.strip().strip("'").strip('"') for tag in x.split(',') if tag.strip()] 
            if isinstance(x, str) and x else []
        )
    
    return df

# Function to plot score evolution over time - simpler fix
def plot_score_evolution(filtered_df):
    """
    Creates a line plot showing the evolution of scores over time.
    Fixed to handle potential string concatenation issues.
    """
    # Only include columns that end with '_score' and exclude any problematic columns
    score_columns = [col for col in filtered_df.columns if col.endswith('_score') and col != 'clean_tags']
    
    if not score_columns or filtered_df.empty:
        st.warning("No hay datos de puntuación disponibles para los filtros seleccionados.")
        return
    
    # Create a copy of the filtered dataframe to avoid modifying the original
    df_scores = filtered_df.copy()
    
    # Clean up each score column to handle possible string concatenation issues
    for col in score_columns:
        # Check if the column contains long concatenated strings (like 'MediumMediumMedium...')
        mask = df_scores[col].astype(str).str.len() > 15
        if mask.any():
            # Set problematic values to NaN
            df_scores.loc[mask, col] = np.nan
        
        # Try to convert to numeric, coercing errors to NaN
        df_scores[col] = pd.to_numeric(df_scores[col], errors='coerce')
    
    # Calculate yearly averages for each score type
    yearly_scores = df_scores.groupby('year')[score_columns].mean().reset_index()
    
    # Only proceed if we have data
    if yearly_scores.empty or yearly_scores[score_columns].isna().all().all():
        st.warning("No hay datos de puntuación utilizables para los años seleccionados.")
        return
    
    # Create a Plotly line chart
    fig = go.Figure()
    
    # Add a line for each score
    for column in score_columns:
        # Skip columns with all NaN values
        if yearly_scores[column].isna().all():
            continue
            
        # Create a more readable label
        label = column.replace('_score', '').replace('_', ' ').title()
        
        fig.add_trace(go.Scatter(
            x=yearly_scores['year'], 
            y=yearly_scores[column],
            mode='lines+markers+text',
            name=label,
            text=yearly_scores[column].round(2),
            textposition="top center",
            line=dict(width=3)
        ))
    
    # Update layout
    fig.update_layout(
        title='Evolución de Puntuaciones Promedio por Año',
        xaxis_title='Año',
        yaxis_title='Puntuación Promedio (0-10)',
        yaxis=dict(range=[0, 10]),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.3,
            xanchor="center",
            x=0.5
        ),
        hovermode="x unified",
        height=500
    )
    
    # Add grid
    fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor='LightGray')
    fig.update_yaxes(showgrid=True, gridwidth=1, gridcolor='LightGray')
    
    return fig

# Function to create a composition plot (stacked bar chart)
def create_composition_plot(filtered_df, var_name, title):
    """
    Creates a stacked bar chart showing the composition of a variable over time.
    Harmonizes 'process' and 'processes' if var_name is 'dimension'.
    """
    if var_name not in filtered_df.columns or filtered_df.empty:
        st.warning(f"No hay datos disponibles para {var_name} con los filtros seleccionados.")
        return

    # Group by year and variable, then count
    var_by_year = filtered_df.groupby(['year', var_name]).size().unstack(fill_value=0)
    # Calculate percentages
    var_by_year_pct = var_by_year.div(var_by_year.sum(axis=1), axis=0) * 100
    # Only proceed if we have data
    if var_by_year_pct.empty:
        st.warning(f"No hay datos disponibles para {var_name} con los años seleccionados.")
        return
    
    # Create a Plotly stacked bar chart
    fig = go.Figure()
    
    # Calculate cumulative percentages for text positioning
    cumulative = pd.DataFrame(0, index=var_by_year_pct.index, columns=['cum'])
    
    # Add a bar for each category
    for category in var_by_year_pct.columns:
        fig.add_trace(go.Bar(
            x=var_by_year_pct.index,
            y=var_by_year_pct[category],
            name=category,
            text=var_by_year_pct[category].round(1).astype(str) + '%',
            textposition='inside',
            insidetextanchor='middle',
            textfont=dict(size=20),  # Increased label font size
            hoverinfo='name+y'
        ))
        
        # Update cumulative for next bar
        cumulative['cum'] += var_by_year_pct[category]
    
    # Update layout
    # Remove undefined or empty title
    # Fix undefined or empty title: if title is None, '', or 'undefined' (any case/whitespace), do not show a title
    if title is None or str(title).strip() == '' or str(title).strip().lower() == 'undefined':
        layout_title = ''
    else:
        layout_title = title
    fig.update_layout(
        title=layout_title,
        xaxis_title='Año',
        yaxis_title='Porcentaje (%)',
        barmode='stack',
        uniformtext=dict(mode='hide', minsize=10),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.3,
            xanchor="center",
            x=0.5
        ),
        height=500
    )
    
    # Add sample count annotations on top of each bar
    total_counts = filtered_df.groupby('year').size()
    for year in var_by_year_pct.index:
        if year in total_counts.index:
            fig.add_annotation(
                x=year,
                y=105,
                text=f"n={total_counts[year]}",
                showarrow=False,
                font=dict(size=14, color="black")
            )
    
    return fig

# Function to create tag composition plot
def create_tag_composition_plot(filtered_df, top_n=8):
    """
    Creates a stacked bar chart showing the composition of tags over time.
    """
    if 'tags' not in filtered_df.columns or filtered_df.empty:
        st.warning("No hay datos de etiquetas disponibles para los filtros seleccionados.")
        return
    
    # Explode the dataframe by tags
    exploded_df = filtered_df.explode('tags')
    
    # Remove rows with empty tags
    exploded_df = exploded_df.dropna(subset=['tags']).reset_index(drop=True)
    
    if exploded_df.empty:
        st.warning("No hay datos de etiquetas disponibles después del filtrado.")
        return
    
    # Count occurrences of each tag
    tag_counts = exploded_df['tags'].value_counts()
    
    # Get the top N tags
    top_tags = tag_counts.head(top_n).index.tolist()
    
    # Replace non-top tags with 'Other tags'
    exploded_df['tag_category'] = exploded_df['tags'].apply(
        lambda x: x if x in top_tags else 'Otras etiquetas'
    )
    
    # Count yearly occurrences for each tag category
    yearly_tag_counts = exploded_df.groupby(['year', 'tag_category']).size().unstack(fill_value=0)
    
    # Calculate the percentage of each tag category per year
    yearly_tag_percentages = yearly_tag_counts.div(yearly_tag_counts.sum(axis=1), axis=0) * 100
    
    # Only proceed if we have data
    if yearly_tag_percentages.empty:
        st.warning("No hay datos de etiquetas disponibles para los años seleccionados.")
        return
    
    # Create a Plotly stacked bar chart
    fig = go.Figure()
    
    # Add a bar for each tag category
    for category in yearly_tag_percentages.columns:
        fig.add_trace(go.Bar(
            x=yearly_tag_percentages.index,
            y=yearly_tag_percentages[category],
            name=category,
            text=yearly_tag_percentages[category].round(1).astype(str) + '%',
            textposition='inside',
            insidetextanchor='middle',
            textfont=dict(size=14),
            hoverinfo='name+y'
        ))
    
    # Update layout
    fig.update_layout(
        title='Evolución de la Composición de Etiquetas por Año',
        xaxis_title='Año',
        yaxis_title='Porcentaje (%)',
        barmode='stack',
        uniformtext=dict(mode='hide', minsize=10),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.4,
            xanchor="center",
            x=0.5
        ),
        height=500
    )
    
    # Add sample count annotations on top of each bar
    total_counts = exploded_df.groupby('year').size()
    for year in yearly_tag_percentages.index:
        if year in total_counts.index:
            fig.add_annotation(
                x=year,
                y=105,
                text=f"n={total_counts[year]}",
                showarrow=False,
                font=dict(size=14, color="black")
            )
    
    return fig

# Function to create rejection difficulty classification plot
def create_difficulty_classification_plot(filtered_df, top_n=8):
    """
    Creates a stacked bar chart showing the composition of rejection difficulty classifications over time.
    """
    if 'clean_tags' not in filtered_df.columns or filtered_df.empty:
        st.warning("No hay datos de clasificación disponibles para los filtros seleccionados.")
        return
    
    # Explode the dataframe by tags
    exploded_df = filtered_df.explode('clean_tags')
    
    # Remove rows with empty tags
    exploded_df = exploded_df[
        exploded_df['clean_tags'].apply(lambda x: isinstance(x, str) and len(x) > 0)
    ].reset_index(drop=True)
    
    if exploded_df.empty:
        st.warning("No hay datos de clasificación disponibles después del filtrado.")
        return
    
    # Count occurrences of each tag
    tag_counts = exploded_df['clean_tags'].value_counts()
    
    # Get the top N tags
    top_tags = tag_counts.head(top_n).index.tolist()
    
    # Replace non-top tags with 'Other tags'
    exploded_df['tag_category'] = exploded_df['clean_tags'].apply(
        lambda x: x if x in top_tags else 'Otras etiquetas'
    )
    
    # Count yearly occurrences for each tag category
    yearly_tag_counts = exploded_df.groupby(['year', 'tag_category']).size().unstack(fill_value=0)
    
    # Calculate the percentage of each tag category per year
    yearly_tag_percentages = yearly_tag_counts.div(yearly_tag_counts.sum(axis=1), axis=0) * 100
    
    # Only proceed if we have data
    if yearly_tag_percentages.empty:
        st.warning("No hay datos de clasificación disponibles para los años seleccionados.")
        return
    
    # Create a Plotly stacked bar chart
    fig = go.Figure()
    
    # Add a bar for each classification category
    for category in yearly_tag_percentages.columns:
        fig.add_trace(go.Bar(
            x=yearly_tag_percentages.index,
            y=yearly_tag_percentages[category],
            name=category,
            text=yearly_tag_percentages[category].round(1).astype(str) + '%',
            textposition='inside',
            insidetextanchor='middle',
            textfont=dict(size=14),
            hoverinfo='name+y'
        ))
    
    # Update layout
    fig.update_layout(
        title='Evolución de la Composición de Clasificaciones de Dificultad de Rechazo por Año',
        xaxis_title='Año',
        yaxis_title='Porcentaje (%)',
        barmode='stack',
        uniformtext=dict(mode='hide', minsize=10),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.4,
            xanchor="center",
            x=0.5
        ),
        height=500
    )
    
    # Add sample count annotations on top of each bar
    total_counts = exploded_df.groupby('year').size()
    for year in yearly_tag_percentages.index:
        if year in total_counts.index:
            fig.add_annotation(
                x=year,
                y=105,
                text=f"n={total_counts[year]}",
                showarrow=False,
                font=dict(size=14, color="black")
            )
    
    return fig

# # Function to add advanced visualization section to the Streamlit app
# def add_advanced_visualization_section(filtered_df):
#     """
#     Adds an advanced visualization section to the Streamlit app.
#     """
#     st.markdown("#### Visualizaciones Avanzadas")

#     # --- Score Evolution ---
#     st.markdown("<h4 style='margin-top: 2em;'>Evolución de Puntuaciones Promedio por Año</h4>", unsafe_allow_html=True)
#     score_fig = plot_score_evolution(filtered_df)
#     if score_fig:
#         # Remove inline title from plot
#         st.plotly_chart(score_fig, use_container_width=True)

#     # --- Variable Composition ---
#     st.markdown("<h4 style='margin-top: 2em;'>Composición por Variable</h4>", unsafe_allow_html=True)
#     available_vars = [col for col in filtered_df.columns if col.startswith('rec_')]
#     if available_vars:
#         var_mapping = {
#             'rec_innovation_score': 'Nivel de Innovación',
#             'rec_precision_and_clarity': 'Precisión y Claridad',
#             'rec_expected_impact': 'Impacto Esperado',
#             'rec_intervention_approach': 'Enfoque de Intervención',
#             'rec_operational_feasibility': 'Factibilidad Operativa',
#             'rec_timeline': 'Plazo de Implementación'
#         }
#         var_options = {var_mapping.get(var, var): var for var in available_vars if var in var_mapping}
#         if var_options:
#             selected_var_label = st.selectbox(
#                 "Seleccione una variable para visualizar:", 
#                 options=list(var_options.keys())
#             )
#             selected_var = var_options[selected_var_label]
#             var_titles = {
#                 'rec_innovation_score': 'Composición de Niveles de Innovación por Año',
#                 'rec_precision_and_clarity': 'Composición de Niveles de Precisión y Claridad por Año',
#                 'rec_expected_impact': 'Composición de Niveles de Impacto Esperado por Año',
#                 'rec_intervention_approach': 'Composición de Enfoques de Intervención por Año',
#                 'rec_operational_feasibility': 'Composición de Niveles de Factibilidad Operativa por Año',
#                 'rec_timeline': 'Composición de Plazos de Implementación por Año'
#             }
#             composition_fig = create_composition_plot(
#                 filtered_df, 
#                 selected_var, 
#                 var_titles.get(selected_var, f'Composición de {selected_var_label} por Año')
#             )
#             if composition_fig:
#                 st.plotly_chart(composition_fig, use_container_width=True)
#         else:
#             st.warning("No se encontraron variables de composición en los datos filtrados.")
#     else:
#         st.warning("No se encontraron variables de composición en los datos filtrados.")

#     # --- Difficulty Classification ---
#     st.markdown("<h4 style='margin-top: 2em;'>Clasificación de Dificultad de Rechazo</h4>", unsafe_allow_html=True)
#     if 'clean_tags' in filtered_df.columns:
#         top_n = st.slider("Número de clasificaciones principales a mostrar:", min_value=3, max_value=15, value=8, key='diff_class_slider')
#         diff_fig = create_difficulty_classification_plot(filtered_df, top_n)
#         if diff_fig:
#             st.plotly_chart(diff_fig, use_container_width=True)
#     else:
#         st.warning("No se encontraron datos de clasificación de dificultad en los datos filtrados.")

# Fixed version of the function with tab-specific keys
def add_advanced_visualization_section(filtered_df, tab_id="tab1"):
    """
    Adds an advanced visualization section to the Streamlit app.
    
    Parameters:
    -----------
    filtered_df : pandas DataFrame
        The filtered dataframe to visualize
    tab_id : str
        The tab identifier to make widget keys unique (default: "tab1")
    """
    st.markdown("#### Calidad de  Respuesta Institucional a Recomendaciones")

    # --- Score Evolution ---
    st.markdown("<h4 style='margin-top: 2em;'>Evolución de Puntuaciones Promedio por Año</h4>", unsafe_allow_html=True)
    col1, col2 = st.columns([2, 1])
    with col1:
        score_fig = plot_score_evolution(filtered_df)
        if score_fig:
            # Remove inline title from plot
            st.plotly_chart(score_fig, use_container_width=True)
    with col2:
        # Make the legend box visually match the plot height
        st.markdown("""
            <style>
            textarea[data-baseweb="textarea"] {
                min-height: 500px !important;
                height: 500px !important;
                max-height: 700px;
            }
            </style>
        """, unsafe_allow_html=True)
        legend_definitions = st.text_area(
            "Leyenda:",
            value="""
            • Coherencia: Nivel de alineación entre la recomendación y el plan de acción.
            • Calidad del plan: Nivel de especificidad y claridad del plan de acción.
            • Nivel de atención: Nivel de consideración hacia las prioridades de la recomendación, en el plan de acción.
            • Puntuación agregada: Promedio de las puntuaciones de coherencia, calidad del plan y nivel de atención.
            """,
            key=f"legend_definitions_{tab_id}",
            height=500
        )


    # --- Variable Composition ---
    st.markdown("<h4 style='margin-top: 2em;'>Composición de Recomendaciones por Atributo</h4>", unsafe_allow_html=True)
    available_vars = [col for col in filtered_df.columns if col.startswith('rec_')]
    if available_vars:
        var_mapping = {
            'rec_innovation_score': 'Nivel de Innovación',
            'rec_precision_and_clarity': 'Precisión y Claridad',
            'rec_expected_impact': 'Impacto Esperado',
            'rec_intervention_approach': 'Enfoque de Intervención',
            'rec_operational_feasibility': 'Factibilidad Operativa',
            'rec_timeline': 'Plazo de Implementación'
        }
        var_options = {var_mapping.get(var, var): var for var in available_vars if var in var_mapping}
        if var_options:
            # Use tab_id parameter to create unique key for this selectbox
            selected_var_label = st.selectbox(
                "Seleccione una variable para visualizar:", 
                options=list(var_options.keys()),
                key=f"variable_{tab_id}"  # This ensures unique keys across tabs
            )
            selected_var = var_options[selected_var_label]
            var_titles = {
                'rec_innovation_score': 'Composición de Niveles de Innovación por Año',
                'rec_precision_and_clarity': 'Composición de Niveles de Precisión y Claridad por Año',
                'rec_expected_impact': 'Composición de Niveles de Impacto Esperado por Año',
                'rec_intervention_approach': 'Composición de Enfoques de Intervención por Año',
                'rec_operational_feasibility': 'Composición de Niveles de Factibilidad Operativa por Año',
                'rec_timeline': 'Composición de Plazos de Implementación por Año'
            }
            composition_fig = create_composition_plot(
                filtered_df, 
                selected_var, 
                var_titles.get(selected_var, f'Composición de {selected_var_label} por Año')
            )
            if composition_fig:
                st.plotly_chart(composition_fig, use_container_width=True)
        else:
            st.warning("No se encontraron variables de composición en los datos filtrados.")
    else:
        st.warning("No se encontraron variables de composición en los datos filtrados.")

    # --- Difficulty Classification ---
    st.markdown("<h4 style='margin-top: 2em;'>Principales Barreras de Implementación</h4>", unsafe_allow_html=True)
    if 'clean_tags' in filtered_df.columns:
        # Use tab_id parameter to create unique key for this slider
        top_n = st.slider(
            "Número de clasificaciones principales a mostrar:", 
            min_value=3, max_value=15, value=8, 
            key=f"diff_class_slider_{tab_id}"  # This ensures unique keys across tabs
        )
        diff_fig = create_difficulty_classification_plot(filtered_df, top_n)
        if diff_fig:
            st.plotly_chart(diff_fig, use_container_width=True)
    else:
        st.warning("No se encontraron datos de barreras en los datos filtrados.")
# ============= DATA LOADING FUNCTIONS =============

# Load data - use relative paths for deployment
# Define paths as relative to the current directory or using st.secrets for Streamlit Cloud
@st.cache_data
def load_data():
    # Replace with a check for environment, use st.secrets for paths in production
    if os.getenv("STREAMLIT_ENV") == "production":
        # Use st.secrets for file paths in production
        df_path = st.secrets["df_path"]
        df_raw_path = st.secrets["df_raw_path"]
        embeddings_path = st.secrets["embeddings_path"]
    else:
        # Use relative paths for local development
        df_path = "./df_complete_all_full.xlsx"
        df_raw_path = "./df_split_actions.xlsx"
        embeddings_path = "./emb_Recomm_rec_cl_4.pt"
    
    df = pd.read_excel(df_path)
    df['index_df'] = df['ID_Recomendacion']
    # Replace spaces and dots with underscore in column names
    df.columns = df.columns.str.replace(' ', '_').str.replace('.', '_')
    df.rename(columns={'Dimension': 'dimension', 
                       'Subdimension': 'subdim'}, inplace=True)

    # Raw data
    df_raw = pd.read_excel(df_raw_path)
    df_raw['year'] = pd.to_datetime(df_raw['Recommendation date'], format='%Y-%m-%d').dt.year
    df_raw.columns = df_raw.columns.str.replace(' ', '_').str.replace('.', '_')
    missing_index_df = df_raw[~df_raw['index_df'].isin(df['index_df'])]
    
    # Reset index of both DataFrames before concatenation
    df = df.reset_index(drop=True)
    
    # Concatenate with ignore_index=True to avoid index conflicts
    df = pd.concat([df, missing_index_df], axis=0)
    
    # Processes
    df['year'] = pd.to_datetime(df['Recommendation_date'], format='%Y-%m-%d').dt.year
    df['year'] = df['year'].fillna(2023).astype(int)
    df['year'] = df['year'].astype(int)
    df['dimension'] = df['dimension'].fillna('Sin Clasificar')
    df['subdim'] = df['subdim'].fillna('Sin Clasificar')
    df['Management_response'] = df['Management_response'].fillna('Sin respuesta')
    df['Management_response'] = df['Management_response'].replace('Partially completed', 'Partially Completed')
    
    return df, df_raw

@st.cache_data
def load_extended_data():
    """
    Load and prepare additional data for enhanced visualizations.
    This extends the existing load_data function with new data processing.
    """
    # First load the original data
    df, df_raw = load_data()
    
    # Load additional analysis data if available
    try:
        if os.getenv("STREAMLIT_ENV") == "production":
            analyzed_path = st.secrets.get("analyzed_recommendations_path", None)
        else:
            # Use a relative or absolute path based on your setup
            analyzed_path = "./analyzed_recommendations_plans_v5.csv"
        
        if analyzed_path and os.path.exists(analyzed_path):
            # Load the analyzed recommendations with pipe separator
            analyzed_df = pd.read_csv(analyzed_path, sep='|')
            
            # Convert dates and ensure year column
            analyzed_df['Recommendation date'] = pd.to_datetime(analyzed_df['Recommendation date'])
            analyzed_df['year'] = analyzed_df['Recommendation date'].dt.year
            
            # Change years prior to 2018 to 2018 to match the original analysis
            analyzed_df.loc[analyzed_df['year'] < 2018, 'year'] = 2018
            
            # Prepare additional columns needed for the visualizations
            analyzed_df = prepare_additional_data(analyzed_df)
            
            # Merge with the original dataframe if needed (based on a common key)
            if 'index_df' in df.columns and 'index_df' in analyzed_df.columns:
                # Select only new columns from analyzed_df to avoid duplicates
                new_cols = [col for col in analyzed_df.columns if col not in df.columns]
                if new_cols:  # Only proceed if there are new columns to add
                    # Include the key column
                    merge_cols = ['index_df'] + new_cols
                    
                    # Merge the new data
                    df = pd.merge(df, analyzed_df[merge_cols], on='index_df', how='left')
            
            # Store the analyzed dataframe in session state for potential use elsewhere
            st.session_state['analyzed_df'] = analyzed_df
        else:
            st.warning("No se encontró el archivo de datos de análisis adicional. Es posible que algunas visualizaciones no estén disponibles.")
            st.session_state['analyzed_df'] = None
    except Exception as e:
        st.warning(f"Nota: No se pudieron cargar los datos de análisis adicionales. Algunas visualizaciones pueden no estar disponibles. Error: {str(e)}")
        st.session_state['analyzed_df'] = None
    # Process the main dataframe with the additional preparation
    df = prepare_additional_data(df)
    
    return df, df_raw

# Load embeddings
@st.cache_data
def load_embeddings():
    if os.getenv("STREAMLIT_ENV") == "production":
        embeddings_path = st.secrets["embeddings_path"]
        structured_embeddings_path = st.secrets["structured_embeddings_path"]
    else:
        embeddings_path = "./emb_Recomm_rec_cl_4.pt"
        structured_embeddings_path = "./Recommendation_RAG_Metadata.pt"
    
    doc_embeddings = torch.load(embeddings_path)
    doc_embeddings = np.array(doc_embeddings)
    
    structured_embeddings = torch.load(structured_embeddings_path)
    
    # Create a FAISS index
    dimension = doc_embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)
    index.add(doc_embeddings)
    
    return doc_embeddings, structured_embeddings, index

import concurrent.futures

def process_text_analysis(combined_text, map_template, combine_template_prefix, user_template_part):
    """
    Process text analysis in chunks and combine results (parallelized).
    
    Parameters:
    -----------
    combined_text : str
        The combined text to analyze
    map_template : str
        Template for the initial summarization of chunks
    combine_template_prefix : str
        Prefix for the template used to combine summaries
    user_template_part : str
        User-defined part of the template for final analysis
        
    Returns:
    --------
    str
        Analyzed and summarized text
    """
    if not combined_text:
        return None
    
    text_chunks = split_text(combined_text)
    chunk_summaries = []
    MAX_WORKERS = min(8, len(text_chunks)) if text_chunks else 1

    def summarize_chunk(chunk):
        return summarize_text(chunk, map_template)

    # Parallelize chunk summarization
    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = [executor.submit(summarize_chunk, chunk) for chunk in text_chunks]
        for future in concurrent.futures.as_completed(futures):
            summary = future.result()
            if summary:
                chunk_summaries.append(summary)

    # Preserve original chunk order (since as_completed doesn't)
    chunk_summaries_ordered = []
    if chunk_summaries and len(chunk_summaries) == len(text_chunks):
        # If all succeeded, sort by chunk order
        chunk_map = {futures[i]: i for i in range(len(futures))}
        # But as_completed gives no order, so instead:
        # Re-run in order
        for i in range(len(text_chunks)):
            result = futures[i].result()
            if result:
                chunk_summaries_ordered.append(result)
    else:
        chunk_summaries_ordered = chunk_summaries

    if chunk_summaries_ordered:
        combined_summaries = " ".join(chunk_summaries_ordered)
        final_template = combine_template_prefix + user_template_part
        return summarize_text(combined_summaries, final_template)
    
    return None


def split_text(text, max_length=1500):
    """
    Split text into chunks of specified maximum length.
    
    Parameters:
    -----------
    text : str
        The text to split
    max_length : int, optional
        Maximum length of each chunk (default: 1500)
        
    Returns:
    --------
    list
        List of text chunks
    """
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        current_length += len(word) + 1
        if current_length > max_length:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word) + 1
        else:
            current_chunk.append(word)
    
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

def build_combined_text(df, selections):
    """
    Build combined text from selected text sources.
    
    Parameters:
    -----------
    df : pandas.DataFrame
        DataFrame containing the text columns
    selections : dict
        Dictionary indicating which text sources to include
        
    Returns:
    --------
    str
        Combined text from all selected sources
    """
    texts = []
    if selections['recommendations']:
        texts.append(" ".join(df['Recommendation_description'].astype(str).dropna().unique()))
    if selections['lessons']:
        texts.append(" ".join(df['Lessons_learned_description'].astype(str).dropna().unique()))
    if selections['practices']:
        texts.append(" ".join(df['Good_practices_description'].astype(str).dropna().unique()))
    if selections['plans']:
        texts.append(" ".join(df['Action_plan'].astype(str).dropna().unique()))
    return " ".join(texts)

def summarize_text(text, prompt_template):
    """
    Summarize text using OpenAI API.
    
    Parameters:
    -----------
    text : str
        The text to summarize
    prompt_template : str
        Template for the prompt
        
    Returns:
    --------
    str
        Summarized text
    """
    if not openai_api_key:
        st.error("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
        return None
        
    prompt = prompt_template.format(text=text)
    try:
        # Use new Responses API with GPT-5
        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that summarizes and analyzes texts."},
                {"role": "user", "content": prompt}
            ],
            reasoning_effort="minimal"
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error al llamar a la API de OpenAI: {e}")
        return None

# ============= MAIN APP CODE =============

# Set page config
st.set_page_config(layout="wide")

# ============= ILO GLOBAL STYLES =============
# ILO Color Palette: Red (#C8102E), Blue (#002F6C), White (#FFFFFF)
# Secondary: Light Blue (#0072CE), Light Gray (#F5F5F5)
st.markdown("""
<style>
/* ===== ILO Color Variables ===== */
:root {
    --ilo-red: #C8102E;
    --ilo-blue: #002F6C;
    --ilo-light-blue: #0072CE;
    --ilo-white: #FFFFFF;
    --ilo-light-gray: #F5F5F5;
    --ilo-dark-gray: #333333;
}

/* ===== Global Typography ===== */
h1, h2, h3 {
    color: var(--ilo-blue) !important;
    font-weight: 600 !important;
}

h4, h5, h6 {
    color: var(--ilo-dark-gray) !important;
}

/* ===== Primary Buttons ===== */
.stButton > button {
    background-color: var(--ilo-blue) !important;
    color: var(--ilo-white) !important;
    border: none !important;
    border-radius: 6px !important;
    padding: 0.5rem 1.25rem !important;
    font-weight: 500 !important;
    transition: all 0.2s ease !important;
}

.stButton > button:hover {
    background-color: var(--ilo-light-blue) !important;
    box-shadow: 0 4px 12px rgba(0, 47, 108, 0.25) !important;
}

/* ===== Download Buttons ===== */
.stDownloadButton > button {
    background-color: var(--ilo-red) !important;
    color: var(--ilo-white) !important;
    border: none !important;
    border-radius: 6px !important;
}

.stDownloadButton > button:hover {
    background-color: #A00D24 !important;
    box-shadow: 0 4px 12px rgba(200, 16, 46, 0.3) !important;
}

/* ===== Tabs Styling ===== */
.stTabs [data-baseweb="tab-list"] {
    gap: 8px;
    background-color: var(--ilo-light-gray);
    padding: 0.5rem;
    border-radius: 8px;
}

.stTabs [data-baseweb="tab"] {
    background-color: transparent !important;
    color: var(--ilo-blue) !important;
    border-radius: 6px !important;
    padding: 0.5rem 1rem !important;
    font-weight: 500 !important;
}

.stTabs [aria-selected="true"] {
    background-color: var(--ilo-blue) !important;
    color: var(--ilo-white) !important;
}

/* ===== Expanders ===== */
.streamlit-expanderHeader {
    background-color: var(--ilo-light-gray) !important;
    border-radius: 8px !important;
    color: var(--ilo-blue) !important;
    font-weight: 500 !important;
}

/* ===== Info/Warning/Success Boxes ===== */
.stAlert > div[data-baseweb="notification"] {
    border-radius: 8px !important;
}

div[data-testid="stNotificationContentInfo"] {
    background-color: rgba(0, 114, 206, 0.1) !important;
    border-left: 4px solid var(--ilo-light-blue) !important;
}

div[data-testid="stNotificationContentSuccess"] {
    background-color: rgba(40, 167, 69, 0.1) !important;
    border-left: 4px solid #28a745 !important;
}

div[data-testid="stNotificationContentWarning"] {
    background-color: rgba(200, 16, 46, 0.08) !important;
    border-left: 4px solid var(--ilo-red) !important;
}

/* ===== Checkboxes ===== */
.stCheckbox > label > div[data-testid="stMarkdownContainer"] {
    color: var(--ilo-dark-gray) !important;
}

/* ===== Sliders ===== */
.stSlider > div > div > div > div {
    background-color: var(--ilo-blue) !important;
}

/* ===== Select boxes ===== */
.stSelectbox > div > div {
    border-color: var(--ilo-blue) !important;
}

/* ===== DataFrames ===== */
.stDataFrame {
    border: 1px solid #e0e0e0 !important;
    border-radius: 8px !important;
}

/* ===== Progress bars ===== */
.stProgress > div > div > div > div {
    background-color: var(--ilo-blue) !important;
}

/* ===== Metric cards ===== */
[data-testid="stMetricValue"] {
    color: var(--ilo-blue) !important;
}

/* ===== Reference boxes (custom) ===== */
.reference-box {
    border-left: 4px solid var(--ilo-blue) !important;
    background: linear-gradient(135deg, #f0f4f8, #e8eef5) !important;
    padding: 0.85rem 1rem;
    border-radius: 8px;
    font-size: 0.95rem;
    color: var(--ilo-dark-gray);
    margin-bottom: 0.75rem;
    box-shadow: 0 4px 14px rgba(0, 47, 108, 0.08);
}

/* ===== Horizontal rule ===== */
hr {
    border-top: 2px solid var(--ilo-blue) !important;
}

/* ===== Sidebar ===== */
[data-testid="stSidebar"] {
    background-color: var(--ilo-light-gray) !important;
}

[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {
    color: var(--ilo-blue) !important;
}

/* ===== File uploader ===== */
.stFileUploader > div > div {
    border: 2px dashed var(--ilo-blue) !important;
    border-radius: 8px !important;
}

/* ===== Radio buttons ===== */
.stRadio > div {
    gap: 0.5rem;
}

/* ===== Multiselect ===== */
.stMultiSelect > div > div {
    border-color: var(--ilo-blue) !important;
}

/* ===== Spinner ===== */
.stSpinner > div {
    border-top-color: var(--ilo-blue) !important;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
    <h2 style='text-align:center; color:#002F6C; margin-top:0;'>Caja de Herramientas para el Mejor Desempeño de los Proyectos</h2>
    <h3 style='text-align:center; color:#002F6C; margin-top:0;'>Usando Evidencia de las Evaluaciones</h3>
    <hr style='border-top: 2px solid #002F6C;'>
""", unsafe_allow_html=True)

# Initialize data and embeddings - wrap in try/except for better error handling
try:
    # Use the extended data loading function that includes the new visualizations data
    df, df_raw = load_extended_data()
    doc_embeddings, structured_embeddings, index = load_embeddings()
    doc_texts = df_raw['Recommendation_description'].tolist()
except Exception as e:
    st.error(f"Error cargando datos: {str(e)}")
    st.stop()


# Check for API key before running the app
if not openai_api_key:
    st.warning("No se encontró la clave de API de OpenAI. Configura la variable de entorno OPENAI_API_KEY en Streamlit Cloud.")
    st.info("Para desarrollo local, puedes usar un archivo .env o configurar la variable de entorno.")
    # Continue with limited functionality or show instructions on setup

# Initialize data and embeddings - wrap in try/except for better error handling
try:
    # Use the extended data loading function that includes the new visualizations data
    df, df_raw = load_extended_data()
    doc_embeddings, structured_embeddings, index = load_embeddings()
    doc_texts = df_raw['Recommendation_description'].tolist()
except Exception as e:
    st.error(f"Error cargando datos: {str(e)}")
    st.stop()

# Function to get embeddings for lessons learned
def get_lessons_embedding_with_retry(text, model='text-embedding-3-large', max_retries=3, delay=1):
    if not openai_api_key:
        st.error("No se encontró la clave de API de OpenAI. Por favor, configura la variable de entorno OPENAI_API_KEY.")
        return None
    for attempt in range(max_retries):
        try:
            response = client.embeddings.create(input=text, model=model)
            return np.array(response.data[0].embedding)
        except Exception as e:
            st.warning(f"Intento {attempt + 1} fallido: {str(e)}")
            time.sleep(delay)
    return None

# Function to find similar lessons learned using embeddings
def find_similar_lessons(query_embedding, lessons_index, lessons_embeddings, structured_lessons, score_threshold=0.5, top_n=20):
    # Normalize query embedding for cosine similarity
    query_embedding = np.array(query_embedding).reshape(1, -1)
    # Search the index
    try:
        distances, indices = lessons_index.search(query_embedding, lessons_index.ntotal)

        # Filter results based on the score threshold
        filtered_lessons = []
        for idx, dist in zip(indices[0], distances[0]):
            if idx < len(structured_lessons) and dist >= score_threshold:
                metadata = structured_lessons[idx]
                lesson = {
                    "lesson": metadata["text"],
                    "similarity": float(dist),  # Convert to float for JSON serialization
                    "country": metadata["country"],
                    "year": metadata["year"],
                    "eval_id": metadata["eval_id"]
                }
                filtered_lessons.append(lesson)
            if len(filtered_lessons) >= top_n:
                break
        return filtered_lessons
    except Exception as e:
        st.error(f"Error in similarity search: {str(e)}")
        return []

# Function to find lessons by term matching
def find_lessons_by_term_matching(query, lessons_texts, structured_lessons, top_n=10):
    try:
        matched_lessons = []
        query_lower = query.lower()
        for idx, text in enumerate(lessons_texts):
            if isinstance(text, str) and query_lower in text.lower():
                if idx < len(structured_lessons):
                    metadata = structured_lessons[idx]
                    matched_lessons.append({
                        "lesson": text,
                        "country": metadata["country"],
                        "year": metadata["year"],
                        "eval_id": metadata["eval_id"]
                    })
        matched_lessons = sorted(matched_lessons, key=lambda x: len(str(x["lesson"])))
        return matched_lessons[:top_n]
    except Exception as e:
        st.error(f"Error in term matching: {str(e)}")
        return []

# Function to load and prepare lessons embeddings
@st.cache_data
def load_lessons_embeddings():
    try:
        if os.getenv("STREAMLIT_ENV") == "production":
            lessons_embeddings_path = st.secrets["lessons_embeddings_path"]
            structured_lessons_path = st.secrets["structured_lessons_path"]
        else:
            lessons_embeddings_path = "./emb_LL_ll_cl_4.pt"
            structured_lessons_path = "./lessons_metadata.pt"
        
        # Load embeddings and metadata
        lessons_embeddings = torch.load(lessons_embeddings_path)
        lessons_embeddings = np.array(lessons_embeddings)
        
        structured_lessons = torch.load(structured_lessons_path)
        
        # Create a FAISS index
        dimension = lessons_embeddings.shape[1]
        lessons_index = faiss.IndexFlatIP(dimension)
        lessons_index.add(lessons_embeddings)
        
        return lessons_embeddings, structured_lessons, lessons_index
    except Exception as e:
        st.error(f"Error cargando embeddings de lecciones: {str(e)}")
        # Return placeholder data to avoid errors
        return np.array([]), [], None
    
    
# Tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs([ "Valoración Preliminar de Calidad de Proyectos",
                                         "Diagnóstico de Atributos Específicos",
                                         "Diagnóstico de Sostenibilidad del Proyecto",
                                         "Pregúntale a tus Documentos",
                                        #  "Estadísticas sobre Recomendaciones de Evaluaciones y sus Planes de Acción",
                                         "Clasificación de Recomendaciones (World)"])

#-----------------------#-----------------------#
#-----------------------#-----------------------#
# Tab 2: Revisión por criterios con trazabilidad
# with tab2:
#     st.header("Acercamiento a Valoración de Calidad de Proyectos")

#     # Descriptive text box
#     st.info("""
#     **📋 Descripción de la herramienta:**

#     Sube un Word (.docx) para evaluarlo con criterios y niveles de desempeño (rúbricas) alineados a la OIT. La herramienta extrae secciones clave, aplica la matriz de criterios y asigna puntajes 1–5 con análisis narrativo y evidencia trazable (citas + metadatos). Los criterios de Participación de Actores, Género y Transición Justa se aplican a un documento de proyecto. En cuanto a los criterios de "Desempeño del Proyecto" se deberán aplicar exclusivamente a informes de evaluación ya que se basan en la metodología de meta-análisis de la Oficina de Evaluación de la OIT. Finalmente, los criterios de Metodologías con Enfoque Participativo se aplican a informes de evaluación u otros tipos de documentos. 
    
#     Puedes exportar a Excel estos resultados (Criterio, Dimensión, Score, Análisis, Evidencia, Error, Rúbrica). Una vez que los resultados son descargados, éstos se dejarán de mostrar en pantalla.

#     Si hay vacíos o inconsistencias, se señalan en "Error" para su ajuste. Este diagnóstico en formato EXCEL sirve para revisar propuestas antes de enviarlas a donantes, verificar aspectos puntuales de informes de evaluación o de ejecución, comprobar coherencia con P&B, DWCP y marcos UNSDCF, elaborar notas técnicas con sustento y respaldar la rendición de cuentas ante mandantes y donantes.
#     """)

#     # Read rubrics from Excel files as in megaparse_example.py
#     import pandas as pd
#     engagement_rubric = {}
#     performance_rubric = {}
#     parteval_rubric = {}
#     gender_rubric = {}
#     tj_traditional_rubric = {}
#     tj_just_transition_rubric = {}

#     try:
#         df_rubric_engagement = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_engagement')
#         df_rubric_engagement.drop(columns=['Unnamed: 0', 'Criterio'], inplace=True, errors='ignore')
#         for idx, row in df_rubric_engagement.iterrows():
#             indicador = row['Indicador']
#             dimension = row.get('Dimensión', 'No especificada')  # Get dimension, default if not present
#             valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
#             engagement_rubric[indicador] = {'valores': valores, 'dimension': dimension}

#         df_rubric_performance = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_performance')
#         df_rubric_performance.drop(columns=['dimension'], inplace=True, errors='ignore')
#         for idx, row in df_rubric_performance.iterrows():
#             criterio = row['subdim']
#             dimension = row.get('Dimensión', 'No especificada')  # Get dimension, default if not present
#             valores = row.drop(['subdim', 'Dimensión'], errors='ignore').values.tolist()
#             performance_rubric[criterio] = {'valores': valores, 'dimension': dimension}

#         df_rubric_parteval = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_parteval')
#         df_rubric_parteval.drop(columns=['Criterio'], inplace=True, errors='ignore')
#         for idx, row in df_rubric_parteval.iterrows():
#             indicador = row['Indicador']
#             dimension = row.get('Dimensión', 'No especificada')  # Get dimension, default if not present
#             valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
#             parteval_rubric[indicador] = {'valores': valores, 'dimension': dimension}

#         df_rubric_gender = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_gender_')
#         df_rubric_gender.drop(columns=['Criterio'], inplace=True, errors='ignore')
#         for idx, row in df_rubric_gender.iterrows():
#             indicador = row['Indicador']
#             dimension = row.get('Dimensión', 'No especificada')  # Get dimension, default if not present
#             valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
#             gender_rubric[indicador] = {'valores': valores, 'dimension': dimension}

#         # Load TJ Traditional rubric from Rubricas_6ago2025.xlsx
#         try:
#             df_rubric_tj_traditional = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_TJ_Traditional')
#             df_rubric_tj_traditional.drop(columns=['Criterio'], inplace=True, errors='ignore')
#             for idx, row in df_rubric_tj_traditional.iterrows():
#                 indicador = row['Indicador']
#                 if pd.notna(indicador) and str(indicador).strip():
#                     dimension = row.get('Dimensión', 'No especificada')  # Get dimension, default if not present
#                     valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
#                     tj_traditional_rubric[indicador] = {'valores': valores, 'dimension': dimension}
#         except Exception as e:
#             st.error(f"❌ Error cargando TJ Tradicional: {e}")
#             st.write("Sheets disponibles:", list(pd.ExcelFile('./Rubricas_6ago2025.xlsx').sheet_names))

#         # Load TJ Just Transition rubric from Rubricas_6ago2025.xlsx
#         try:
#             df_rubric_tj_just_transition = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_TJ_TJ')
#             df_rubric_tj_just_transition.drop(columns=['Criterio'], inplace=True, errors='ignore')
#             for idx, row in df_rubric_tj_just_transition.iterrows():
#                 indicador = row['Indicador']
#                 if pd.notna(indicador) and str(indicador).strip():
#                     dimension = row.get('Dimensión', 'No especificada')  # Get dimension, default if not present
#                     valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
#                     tj_just_transition_rubric[indicador] = {'valores': valores, 'dimension': dimension}
#         except Exception as e:
#             st.error(f"❌ Error cargando TJ Transición Justa: {e}")
#             st.write("Sheets disponibles:", list(pd.ExcelFile('./Rubricas_6ago2025.xlsx').sheet_names))
#     except Exception as e:
#         st.error(f"Error leyendo las rúbricas: {e}")

#     # Debug: Show Excel sheet names and rubric status
#     try:
#         available_sheets = pd.ExcelFile('./Rubricas_6ago2025.xlsx').sheet_names
#         st.warning(f"**Sheets disponibles en Rubricas_6ago2025.xlsx:** {', '.join(available_sheets)}")
#     except Exception as e:
#         st.error(f"Error leyendo Excel: {e}")

#     st.error(f"""
#     **ESTADO DE RÚBRICAS:**
#     - Participación de Actores (durante el proyecto): {len(engagement_rubric)} criterios
#     - Desempeño del Proyecto (según informes de evaluación): {len(performance_rubric)} criterios
#     - Metodologías con Enfoque Participativo: {len(parteval_rubric)} criterios
#     - Enfoque de Género: {len(gender_rubric)} criterios
#     - Transición Justa: Enfoque Tradicional: {len(tj_traditional_rubric)} criterios
#     - Transición Justa: Enfoque Moderno: {len(tj_just_transition_rubric)} criterios 
#     """)

    # # Function to extract document structure
    #     def extract_docx_structure(docx_path):
    #         from docx import Document
    #         doc = Document(docx_path)
    #         filename = os.path.basename(docx_path)
    #         rows = []
    #         current_headers = {i: '' for i in range(1, 7)}
    #         para_counter = 0
    
    #         def get_header_level(style_name):
    #             for i in range(1, 7):
    #                 if style_name.lower().startswith(f'heading {i}'.lower()):
    #                     return i
    #             return None
    
    #         def header_dict():
    #             return {f'header_{i}': current_headers[i] for i in range(1, 7)}
    
    #         for para in doc.paragraphs:
    #             para_counter += 1
    #             level = get_header_level(para.style.name)
    #             if level and 1 <= level <= 6:
    #                 current_headers[level] = para.text.strip()
    #                 for l in range(level+1, 7):
    #                     current_headers[l] = ''
    #                 rows.append({
    #                     'filename': filename,
    #                     **header_dict(),
    #                     'content': '',
    #                     'source_type': 'heading',
    #                     'paragraph_number': para_counter,
    #                     'page_number': None
    #                 })
    #             elif para.text.strip():
    #                 rows.append({
    #                     'filename': filename,
    #                     **header_dict(),
    #                     'content': para.text.strip(),
    #                     'source_type': 'paragraph',
    #                     'paragraph_number': para_counter,
    #                     'page_number': None
    #                 })
    #         return pd.DataFrame(rows)
    
    #     # Function to split text into chunks respecting the token limit
    #     def split_text_into_chunks(text, max_completion_tokens=7000):
    #         import re
    #         # Split by paragraphs first
    #         paragraphs = text.split('\n')
    #         chunks = []
    #         current_chunk = []
    #         current_length = 0
    
    #         # Rough estimate: 1 token ≈ 4 characters in Spanish
    #         tokens_per_char = 0.25
    
    #         for para in paragraphs:
    #             # Estimate tokens in this paragraph
    #             para_tokens = len(para) * tokens_per_char
    
    #             # If adding this paragraph would exceed the max, start a new chunk
    #             if current_length + para_tokens > max_tokens and current_chunk:
    #                 chunks.append('\n'.join(current_chunk))
    #                 current_chunk = [para]
    #                 current_length = para_tokens
    #             else:
    #                 current_chunk.append(para)
    #                 current_length += para_tokens
    
    #         # Add the last chunk if there's content
    #         if current_chunk:
    #             chunks.append('\n'.join(current_chunk))
    
    #         return chunks

#     # Function to directly evaluate content against a criterion using LLM
#     # def evaluate_criterion_with_llm(document_text, criterion, descriptions):
#     #     """Evaluate document against a criterion directly with LLM"""

#     #     # Split document into manageable chunks if needed
#     #     chunks = split_text_into_chunks(document_text)

#     #     # If text fits in one chunk, evaluate directly
#     #     if len(chunks) == 1:
#     #         return evaluate_single_chunk(chunks[0], criterion, descriptions)

#     #     # For multiple chunks, evaluate each and then synthesize
#     #     chunk_results = []
#     #     for i, chunk in enumerate(chunks):
#     #         st.info(f"Evaluando criterio '{criterion}' - Fragmento {i+1}/{len(chunks)}")
#     #         result = evaluate_single_chunk(chunk, criterion, descriptions)
#     #         chunk_results.append(result)

#     #     # Synthesize results from all chunks
#     #     return synthesize_evaluations(chunk_results, criterion, descriptions)

#     def evaluate_criterion_with_llm(document_text, criterion, descriptions):
#         """Analyze complete document efficiently using a two-stage approach"""
        
#         # Stage 1: Extract relevant sections (cheap, fast model)
#         chunks = split_text_into_chunks(document_text, max_completion_tokens=7000)
        
#         relevant_chunks = []
#         for chunk in chunks:
#             # Quick relevance check with cheap model
#             check_prompt = f"Does this text mention or relate to '{criterion}'? Answer only YES or NO.\n\n{chunk[:1000]}"
            
#             response = openai.ChatCompletion.create(
#                 model="gpt-5-mini",  # Cheap model for filtering
#                 messages=[{"role": "user", "content": check_prompt}],
#                 max_completion_tokens=5,
#                 temperature=0
#             )
            
#             if "YES" in response["choices"][0]["message"]["content"].upper():
#                 relevant_chunks.append(chunk)
        
#         # Stage 2: Deep analysis only on relevant chunks
#         if not relevant_chunks:
#             # If nothing relevant found, use first and last chunks as context
#             relevant_chunks = [chunks[0], chunks[-1]] if len(chunks) > 1 else chunks
        
#         # Combine relevant chunks (limit to ~10k chars)
#         combined_text = "\n\n---\n\n".join(relevant_chunks)[:10000]
        
#         # Now do the expensive analysis on focused content
#         prompt = f"""Evaluate this document against: {criterion}
    
#     Scoring levels: {json.dumps(descriptions)}
    
#     Relevant document sections:
#     {combined_text}
    
#     Provide JSON with:
#     {{"analysis": "detailed 2-3 paragraphs", "score": 1-5, "evidence": "5-8 key quotes from the text"}}"""
    
#         response = openai.ChatCompletion.create(
#             model="gpt.5.1-mini",
#             messages=[
#                 {"role": "system", "content": "You are an expert document evaluator."},
#                 {"role": "user", "content": prompt}
#             ],
#             max_completion_tokens=1500,
#             temperature=0.1
#         )
        
#         return json.loads(response["choices"][0]["message"]["content"])

#     # Function to evaluate a single text chunk
#     def evaluate_single_chunk(text_chunk, criterion, descriptions):
#         """Evaluate a single text chunk against a criterion with expanded analysis and evidence"""
#         import json

#         # Build prompt
#         prompt = f"""
#         Estás evaluando un documento contra un criterio específico.
        
#         Criterio: {criterion}
        
#         Descripciones de los niveles de puntuación:
#         {json.dumps(descriptions, indent=2)}
        
#         Contenido del documento a evaluar:
#         {text_chunk}
        
#         Analiza qué tan bien el documento cumple con este criterio. Proporciona:
        
#         1. Un análisis DETALLADO (2-3 párrafos) que explique a fondo el razonamiento detrás de tu evaluación. Proporciona un razonamiento profundo que abarque los aspectos del criterio.
        
#         2. Una puntuación de 1-5 (donde 1 es la más baja y 5 es la más alta).
        
#         3. EVIDENCIA del documento que respalde tu puntuación. Incluye entre 5-8 citas textuales del documento, indicando cómo cada fragmento contribuye a tu evaluación.
        
#         Formatea tu respuesta como un objeto JSON con las siguientes claves:
#         {{"analysis": "tu análisis detallado aquí", "score": puntuación_numérica_entre_1_y_5, "evidence": "citas textuales del documento (5-8 párrafos)"}}
        
#         Devuelve solo el objeto JSON, nada más.
#         """

#         # Call LLM using OpenAI v0.28 syntax
#         try:
#             response = openai.ChatCompletion.create(
#                 model="gpt.5.1-mini",
#                 messages=[
#                     {"role": "system", "content": "Eres un experto evaluador de documentos que proporciona análisis detallados basados en criterios específicos. Tu evidencia cita fragmentos del texto original."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 response_format={"type": "json_object"},
#                 max_completion_tokens=7000
#             )
#             raw = response["choices"][0]["message"]["content"].strip()
#             parsed = json.loads(raw)
#             return parsed
#         except Exception as e:
#             return {'score': 0, 'analysis': f'Error: {str(e)}', 'evidence': ''}

#     # Function to synthesize evaluations
#     def synthesize_evaluations(chunk_results, criterion, descriptions):
#         """Synthesize evaluations from multiple document chunks with expanded analysis and evidence"""
#         import json

#         # Extract and format the individual evaluations for the synthesis
#         individual_evals = []
#         all_evidence = []

#         for i, result in enumerate(chunk_results):
#             individual_evals.append(f"Evaluación del fragmento {i+1}:\n" +
#                                     f"Puntuación: {result.get('score', 0)}\n" +
#                                     f"Análisis: {result.get('analysis', '')}")

#             # Collect all evidence
#             evidence = result.get('evidence', '')
#             if evidence:
#                 all_evidence.append(f"Evidencia del fragmento {i+1}:\n{evidence}")

#         # Define separator outside the f-string to avoid backslash issues
#         separator = "\n\n"

#         # Create a synthesis prompt
#         synthesis_prompt = f"""
#         Has evaluado un documento dividido en múltiples fragmentos contra el criterio: {criterion}
        
#         Aquí están las evaluaciones individuales de cada fragmento:
        
#         {separator.join(individual_evals)}
        
#         Basándote en estas evaluaciones individuales, proporciona:
        
#         1. Un análisis DETALLADO (2-3 párrafos) que integre los hallazgos clave de todos los fragmentos. Este análisis debe ser comprensivo y abarcar los aspectos relevantes encontrados en el documento.
        
#         2. Una puntuación general de 1-5 (puedes promediar las puntuaciones o ajustar según sea necesario)
        
#         3. Las evidencias más importantes del documento. Selecciona las 8-10 citas textuales más relevantes de los fragmentos individuales.
        
#         Formatea tu respuesta como un objeto JSON con las siguientes claves:
#         {{"analysis": "tu análisis global detallado aquí", "score": puntuación_general_entre_1_y_5, "evidence": "las citas textuales más relevantes del documento (8-10 párrafos)"}}
        
#         Devuelve solo el objeto JSON, nada más.
#         """

#         # Call LLM for synthesis using OpenAI v0.28 syntax
#         try:
#             response = openai.ChatCompletion.create(
#                 model="gpt.5.1-mini",
#                 messages=[
#                     {"role": "system", "content": "Eres un experto evaluador de documentos que sintetiza análisis de múltiples fragmentos de texto para producir evaluaciones detalladas con evidencia textual."},
#                     {"role": "user", "content": synthesis_prompt}
#                 ],
#                 response_format={"type": "json_object"},
#                 max_completion_tokens=7000
#             )
#             raw = response["choices"][0]["message"]["content"].strip()
#             parsed = json.loads(raw)
#             return parsed
#         except Exception as e:
#             # If synthesis fails, combine results manually in a more limited way
#             avg_score = sum(r.get('score', 0) for r in chunk_results) / len(chunk_results)
#             # Take only the first paragraph of each analysis to avoid token limits
#             analysis_parts = []
#             for r in chunk_results:
#                 analysis = r.get('analysis', '')
#                 first_para = analysis.split('\n\n')[0] if '\n\n' in analysis else analysis
#                 analysis_parts.append(first_para)

#             # Take only the first few evidence items
#             evidence_parts = []
#             evidence_count = 0
#             for evidence in all_evidence:
#                 parts = evidence.split('\n\n')
#                 # Add up to 2 evidence parts per chunk
#                 for part in parts[:2]:
#                     if evidence_count < 8:  # Limit to 8 total evidence parts
#                         evidence_parts.append(part)
#                         evidence_count += 1

#             return {
#                 'score': avg_score,
#                 'analysis': separator.join(analysis_parts),
#                 'evidence': separator.join(evidence_parts)
#             }

#     # Document upload interface
#     uploaded_file = st.file_uploader("Suba un archivo DOCX para evaluación:", type=["docx"])

#     # Move instructions/info to the top of the tab
#     st.info("""
#     **Instrucciones:**
#     1. Cargue un informe de evaluación en formato .DOCX y presione el botón 'Procesar y Evaluar'.
#     2. Revise los resultados de cada rúbrica en la tabla interactiva.
#     3. Visualice las puntuaciones promedio por dimensión y subdimensión en los gráficos de barras.
#     4. Descargue todos los resultados y evidencias en un archivo ZIP.
#     """)
#     # Unified process, evaluate, and download button
#     st.markdown("#### Procesamiento y Evaluación de Documento")
#     st.markdown('---')
#     if st.button('Procesar y Evaluar'):
#         # Only process if file is uploaded and not already processed for this file
#         if uploaded_file is not None:
#             file_hash = hash(uploaded_file.getvalue())
#             if st.session_state.get('last_file_hash') != file_hash:
#                 with st.spinner("Procesando documento..."):
#                     try:
#                         tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
#                         tmp_file.write(uploaded_file.read())
#                         tmp_file.close()
#                         progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
#                         doc_result = docx2python(tmp_file.name)
#                         df = extract_docx_structure(tmp_file.name)
#                         progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")
#                         header_1_values = df['header_1'].dropna().unique()
#                         llm_summary_rows = []
#                         llm_progress = st.progress(0, text="Procesando secciones con LLM...")
#                         total_sections = len(header_1_values)
#                         for idx, header in enumerate(header_1_values):
#                             section_df = df[df['header_1'] == header].copy()
#                             full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
#                             if not full_text:
#                                 llm_output = ""
#                             else:
#                                 llm_progress.progress((idx+1)/total_sections, text=f"Procesando sección: {header}")
#                                 try:
#                                     response = openai.ChatCompletion.create(
#                                         model="gpt.5.1-mini",
#                                         messages=[
#                                             {"role": "system", "content": "You are a helpful assistant that rewrites extracted document content into well-structured, formal paragraphs. Do not rewrite the original content, just reconstruct it in proper, coherent paragraphs, without rephrasing or paraphrasing or rewording."},
#                                             {"role": "user", "content": full_text}
#                                         ],
#                                         max_completion_tokens=4096,
#                                         temperature=0.01,
#                                     )
#                                     llm_output = response["choices"][0]["message"]["content"].strip()
#                                 except Exception as e:
#                                     llm_output = f"[LLM ERROR: {e}]"
#                             llm_summary_rows.append({'header_1': header, 'llm_paragraph': llm_output})
#                         llm_progress.progress(1.0, text="LLM parsing completado.")
#                         llm_summary_df = pd.DataFrame(llm_summary_rows)
#                         llm_summary_df['n_words'] = llm_summary_df['llm_paragraph'].str.split().str.len()
#                         exploded_df = llm_summary_df.assign(
#                             llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
#                         ).explode('llm_paragraph')
#                         exploded_df = exploded_df.reset_index(drop=True)
#                         exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
#                         full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
#                         file_size = os.path.getsize(tmp_file.name)
#                         n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
#                         n_paragraphs = len(exploded_df)
#                         st.session_state['full_document_text'] = full_document_text
#                         st.session_state['document_stats'] = {
#                             'file_size': file_size,
#                             'n_words': n_words,
#                             'n_paragraphs': n_paragraphs
#                         }
#                         st.session_state['exploded_df'] = exploded_df
#                         st.session_state['last_file_hash'] = file_hash
#                         try:
#                             os.unlink(tmp_file.name)
#                         except:
#                             pass
#                         progress_bar.progress(0.8, text="Documento procesado. Listo para evaluación.")
#                         st.info(f"**Resumen del documento:**\n\n" + 
#                                 f"- Tamaño del archivo: {file_size/1024:.2f} KB\n" + 
#                                 f"- Número de palabras: {n_words}\n" + 
#                                 f"- Número de párrafos: {n_paragraphs}")
#                         st.markdown("#### Estructura extraída del documento:")
#                         st.dataframe(exploded_df, use_container_width=True)
#                         progress_bar.progress(1.0, text="Procesamiento completo.")
#                     except Exception as e:
#                         st.error(f"Error procesando el documento: {e}")
#                         import traceback
#                         st.error(traceback.format_exc())
#                         st.stop()
#         # Now, always run rubric evaluation if document is processed
#         document_text = st.session_state.get('full_document_text', '')
#         if not document_text:
#             st.error("No se pudo recuperar el texto del documento. Por favor, vuelva a cargar el archivo.")
#             st.stop()
#         rubrics = [
#             ("Participación de Actores (durante el proyecto)", engagement_rubric),
#             ("Desempeño del proyecto (según informe de evaluación)", performance_rubric),
#             ("Participación durante la evaluación (metodología)", parteval_rubric),
#             ("Enfoque de Género", gender_rubric),
#             ("Transición Justa: Enfoque Tradicional", tj_traditional_rubric),
#             ("Transición Justa: Enfoque Moderno", tj_just_transition_rubric)
#         ]
#         rubric_results = []
#         from concurrent.futures import ThreadPoolExecutor, as_completed
#         MAX_WORKERS = 8
#         def eval_one_criterion(args):
#             crit, descriptions, dimension, rubric_name = args
#             try:
#                 result = evaluate_criterion_with_llm(document_text, crit, descriptions)
#                 return {
#                     'Criterio': crit,
#                     'Dimensión': dimension,
#                     'Score': result.get('score', 0),
#                     'Análisis': result.get('analysis', ''),
#                     'Evidencia': result.get('evidence', ''),
#                     'Error': result.get('error', '') if 'error' in result else '',
#                     'Rúbrica': rubric_name
#                 }
#             except Exception as e:
#                 return {
#                     'Criterio': crit,
#                     'Dimensión': dimension,
#                     'Score': 0,
#                     'Análisis': '',
#                     'Evidencia': '',
#                     'Error': str(e),
#                     'Rúbrica': rubric_name
#                 }
#         for rubric_name, rubric_dict in rubrics:
#             # Skip empty rubrics
#             if not rubric_dict:
#                 st.warning(f"Saltando rúbrica {rubric_name}: sin criterios cargados")
#                 continue

#             rubric_analysis_data = []
#             n_criteria = len(rubric_dict)
#             progress = st.progress(0, text=f"Iniciando evaluación por rúbrica: {rubric_name}...")
#             with st.spinner(f'Evaluando documento por rúbrica: {rubric_name}...'):
#                 with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
#                     futures = {
#                         executor.submit(eval_one_criterion, (
#                             crit,
#                             rubric_data['valores'] if isinstance(rubric_data, dict) else rubric_data,
#                             rubric_data.get('dimension', 'No especificada') if isinstance(rubric_data, dict) else 'No especificada',
#                             rubric_name
#                         )): (crit, idx)
#                         for idx, (crit, rubric_data) in enumerate(rubric_dict.items())
#                     }
#                     completed = 0
#                     for future in as_completed(futures):
#                         result = future.result()
#                         rubric_analysis_data.append(result)
#                         completed += 1
#                         crit, idx = futures[future]
#                         progress.progress(completed / n_criteria, text=f"Evaluando criterio: {crit}")
#             rubric_results.append((rubric_name, pd.DataFrame(rubric_analysis_data)))
#         # Show and allow download of both results only after evaluation
#         if rubric_results:
#             for rubric_name, rubric_analysis_df in rubric_results:
#                 st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
#                 if not rubric_analysis_df.empty:
#                     # Ensure 'Evidencia' column is present
#                     if 'Evidencia' not in rubric_analysis_df.columns:
#                         rubric_analysis_df['Evidencia'] = ''

#                     # Reorder columns to show in logical order: Criterio, Dimensión, Score, Análisis, Evidencia, Error, Rúbrica
#                     cols = rubric_analysis_df.columns.tolist()
#                     desired_order = ['Criterio', 'Dimensión', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
#                     new_order = [col for col in desired_order if col in cols]
#                     remaining_cols = [col for col in cols if col not in desired_order]
#                     final_order = new_order + remaining_cols
#                     rubric_analysis_df = rubric_analysis_df[final_order]
#                     # Normalize 'Evidencia' column to always be a string
#                     if 'Evidencia' in rubric_analysis_df.columns:
#                         rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
#                             lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
#                         )
#                     st.dataframe(rubric_analysis_df, use_container_width=True)
#                 else:
#                     st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
#             # Provide a zip download for both results
#             import io, zipfile
#             zip_buffer = io.BytesIO()
#             with zipfile.ZipFile(zip_buffer, "w") as zipf:
#                 for rubric_name, rubric_analysis_df in rubric_results:
#                     # Normalize 'Evidencia' column to always be a string before exporting
#                     if 'Evidencia' in rubric_analysis_df.columns:
#                         rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
#                             lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
#                         )
#                     csv = rubric_analysis_df.to_csv(index=False)
#                     arcname = f"evaluacion_rubrica_{rubric_name.replace(' ', '_').lower()}.csv"
#                     zipf.writestr(arcname, csv)
#             zip_buffer.seek(0)
#             st.download_button(
#                 label="Descargar ambos resultados como ZIP",
#                 data=zip_buffer,
#                 file_name="resultados_rubricas.zip",
#                 mime="application/zip"
#             )
#         else:
#             st.warning("No se generaron resultados para ninguna rúbrica.")
#     else:
#         st.info("Por favor suba un archivo DOCX para comenzar y pulse el botón para procesar y evaluar.")

# Tab 2: Revisión por criterios con trazabilidad
with tab2:
    st.header("Diagnóstico de Atributos Específicos")

    # Descriptive text box
    st.info("""
    **📋 Descripción de la herramienta:**

    **¿Qué hace esta herramienta?:**
    Profundiza la Valoración Preliminar aplicando rúbricas OIT con niveles de desempeño (1–5) y evidencia trazable (citas y metadatos) sobre un documento .docx que subas. Extrae secciones clave, evalúa contra una matriz de criterios y genera un análisis narrativo por criterio.
    Criterios disponibles y alcance
    -	Metodologías con enfoque participativo → aplicar a informes de evaluación u otros documentos metodológicos.

    -	Integración del enfoque de género → aplicar a documentos de diseño o ejecución de proyecto u otros estudios (p. ej., PRODOC, TPR, etc.).

    -   Integración del enfoque de Transición Justa (enfoque moderno) → aplicar a documentos de diseño o ejecución de proyecto u otros estudios (p. ej., PRODOC, TPR, etc.). 
   
    Puedes exportar a Excel estos resultados (Criterio, Dimensión, Score, Análisis, Evidencia, Error, Rúbrica). Una vez que los resultados son descargados, éstos se dejarán de mostrar en pantalla.
    Si hay vacíos o inconsistencias, se señalan en "Error" para su ajuste.
    
    **¿Para qué usar este diagnóstico?**
    Este diagnóstico en formato EXCEL sirve para **revisar propuestas**, **verificar aspectos puntuales** de informes de evaluación o de ejecución, **comprobar coherencia** con P&B, DWCP y marcos UNSDCF, **elaborar notas técnicas con sustento** y respaldar la rendición de cuentas ante mandantes y donantes.
    """)
    
    # Important requirements reminder
    st.warning("""
    **⚠️ Recordatorio importante:**
    - Los documentos deben estar formateados con **estilos de encabezado de Word** (Heading 1, Heading 2, etc.) para que las secciones se identifiquen correctamente
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Solo se aceptan archivos en formato **.docx**
    """)

    # Read rubrics from Excel files as in megaparse_example.py
    import pandas as pd

    # Function to extract document structure
    def extract_docx_structure(docx_path):
        from docx import Document
        doc = Document(docx_path)
        filename = os.path.basename(docx_path)
        rows = []
        current_headers = {i: '' for i in range(1, 7)}
        para_counter = 0

        def get_header_level(style_name):
            for i in range(1, 7):
                if style_name.lower().startswith(f'heading {i}'.lower()):
                    return i
            return None

        def header_dict():
            return {f'header_{i}': current_headers[i] for i in range(1, 7)}

        for para in doc.paragraphs:
            para_counter += 1
            level = get_header_level(para.style.name)
            if level and 1 <= level <= 6:
                current_headers[level] = para.text.strip()
                for l in range(level+1, 7):
                    current_headers[l] = ''
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'content': '',
                    'source_type': 'heading',
                    'paragraph_number': para_counter,
                    'page_number': None
                })
            elif para.text.strip():
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'content': para.text.strip(),
                    'source_type': 'paragraph',
                    'paragraph_number': para_counter,
                    'page_number': None
                })
        return pd.DataFrame(rows)
    
    # Enhanced extraction function with multi-method header detection, tables, and validation
    def extract_docx_structure_enhanced(docx_path):
        """
        Enhanced document extraction with:
        - Multi-method header detection (style, formatting, pattern)
        - Table extraction
        - List extraction
        - Page estimation
        - Quality metrics
        """
        from docx import Document
        from docx.shared import Pt
        import re
        
        doc = Document(docx_path)
        filename = os.path.basename(docx_path)
        rows = []
        tables_data = []
        current_headers = {i: '' for i in range(1, 7)}
        para_counter = 0
        page_estimate = 1
        words_per_page = 500  # Estimate for page calculation
        
        # Track statistics
        stats = {
            'total_paragraphs': 0,
            'total_tables': 0,
            'headers_detected': 0,
            'headers_by_style': 0,
            'headers_by_formatting': 0,
            'headers_by_pattern': 0,
            'orphaned_paragraphs': 0,
            'empty_sections': 0
        }
        
        def get_header_level_by_style(style_name):
            """Method 1: Style-based detection (original method)"""
            for i in range(1, 7):
                if style_name.lower().startswith(f'heading {i}'.lower()):
                    return i
            return None
        
        def get_header_level_by_formatting(para):
            """Method 2: Formatting-based detection"""
            # Check if paragraph is bold and larger than normal text
            if para.runs:
                first_run = para.runs[0]
                is_bold = first_run.bold
                font_size = first_run.font.size
                
                # If bold and larger font, likely a header
                if is_bold and font_size:
                    if font_size >= Pt(14):  # Larger than normal (11-12pt)
                        # Estimate level based on size
                        if font_size >= Pt(18):
                            return 1
                        elif font_size >= Pt(16):
                            return 2
                        elif font_size >= Pt(14):
                            return 3
                elif is_bold and len(para.text.strip()) < 100:  # Short bold text
                    return 2  # Likely a subheader
            return None
        
        def get_header_level_by_pattern(para_text):
            """Method 3: Pattern-based detection (numbered headings)"""
            text = para_text.strip()
            # Patterns like "1.", "1.1", "1.1.1", "I.", "A.", etc.
            patterns = [
                (r'^\d+\.\s+\w', 1),  # "1. Title"
                (r'^\d+\.\d+\s+\w', 2),  # "1.1 Title"
                (r'^\d+\.\d+\.\d+\s+\w', 3),  # "1.1.1 Title"
                (r'^[IVX]+\.\s+\w', 1),  # "I. Title"
                (r'^[A-Z]\.\s+\w', 2),  # "A. Title"
            ]
            for pattern, level in patterns:
                if re.match(pattern, text):
                    return level
            return None
        
        def detect_header_level(para):
            """Try all methods to detect header level"""
            # Method 1: Style-based
            level = get_header_level_by_style(para.style.name)
            if level:
                stats['headers_by_style'] += 1
                return level
            
            # Method 2: Formatting-based
            level = get_header_level_by_formatting(para)
            if level:
                stats['headers_by_formatting'] += 1
                return level
            
            # Method 3: Pattern-based
            level = get_header_level_by_pattern(para.text)
            if level:
                stats['headers_by_pattern'] += 1
                return level
            
            return None
        
        def header_dict():
            return {f'header_{i}': current_headers[i] for i in range(1, 7)}
        
        def get_full_header_path():
            """Get full hierarchical path of headers"""
            path_parts = []
            for i in range(1, 7):
                if current_headers[i]:
                    path_parts.append(current_headers[i])
            return ' > '.join(path_parts) if path_parts else 'Sin sección'
        
        # Extract paragraphs and headers
        for para in doc.paragraphs:
            para_counter += 1
            stats['total_paragraphs'] += 1
            
            level = detect_header_level(para)
            if level and 1 <= level <= 6:
                current_headers[level] = para.text.strip()
                # Clear lower level headers
                for l in range(level+1, 7):
                    current_headers[l] = ''
                
                stats['headers_detected'] += 1
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'header_path': get_full_header_path(),
                    'content': '',
                    'source_type': 'heading',
                    'paragraph_number': para_counter,
                    'page_estimate': page_estimate,
                    'detection_method': 'style' if get_header_level_by_style(para.style.name) else ('formatting' if get_header_level_by_formatting(para) else 'pattern')
                })
            elif para.text.strip():
                # Estimate page number based on word count
                word_count = len(para.text.split())
                if word_count > 0:
                    page_estimate = max(1, int(stats['total_paragraphs'] * 0.02))  # Rough estimate
                
                rows.append({
                    'filename': filename,
                    **header_dict(),
                    'header_path': get_full_header_path(),
                    'content': para.text.strip(),
                    'source_type': 'paragraph',
                    'paragraph_number': para_counter,
                    'page_estimate': page_estimate,
                    'detection_method': None
                })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            stats['total_tables'] += 1
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            # Find which section this table belongs to
            current_section = get_full_header_path()
            
            tables_data.append({
                'filename': filename,
                'table_number': table_idx + 1,
                'section': current_section,
                'header_path': current_section,
                'data': table_data,
                'row_count': len(table_data),
                'col_count': len(table_data[0]) if table_data else 0
            })
            
            # Add table as content row
            table_text = '\n'.join([' | '.join(row) for row in table_data])
            rows.append({
                'filename': filename,
                **header_dict(),
                'header_path': current_section,
                'content': f"[TABLA {table_idx + 1}]\n{table_text}",
                'source_type': 'table',
                'paragraph_number': para_counter + table_idx,
                'page_estimate': page_estimate,
                'detection_method': None
            })
        
        # Calculate quality metrics
        df = pd.DataFrame(rows)
        
        # Count orphaned paragraphs (no header_1)
        stats['orphaned_paragraphs'] = len(df[(df['source_type'] == 'paragraph') & (df['header_1'].isna() | (df['header_1'] == ''))])
        
        # Count empty sections
        if 'header_1' in df.columns:
            sections = df[df['header_1'].notna() & (df['header_1'] != '')]['header_1'].unique()
            for section in sections:
                section_df = df[df['header_1'] == section]
                if len(section_df[section_df['source_type'] == 'paragraph']) == 0:
                    stats['empty_sections'] += 1
        
        return df, tables_data, stats
    
    # Validation function
    def validate_extraction(df, tables_data, stats):
        """
        Validate extraction quality and return metrics, warnings, and recommendations
        """
        warnings = []
        recommendations = []
        quality_score = 100
        
        # Check header detection rate
        if stats['total_paragraphs'] > 0:
            header_rate = (stats['headers_detected'] / stats['total_paragraphs']) * 100
            if header_rate < 2:  # Less than 2% headers might indicate poor detection
                warnings.append(f"⚠️ Tasa de detección de encabezados baja: {header_rate:.1f}%")
                quality_score -= 10
                recommendations.append("Verificar si los encabezados usan estilos estándar de Word")
        
        # Check for orphaned paragraphs
        if stats['orphaned_paragraphs'] > 0:
            orphan_rate = (stats['orphaned_paragraphs'] / stats['total_paragraphs']) * 100
            if orphan_rate > 10:  # More than 10% orphaned
                warnings.append(f"⚠️ {stats['orphaned_paragraphs']} párrafos sin sección asignada ({orphan_rate:.1f}%)")
                quality_score -= 15
                recommendations.append("Revisar párrafos huérfanos y asignarlos manualmente a secciones")
        
        # Check for empty sections
        if stats['empty_sections'] > 0:
            warnings.append(f"⚠️ {stats['empty_sections']} secciones sin contenido")
            quality_score -= 5 * min(stats['empty_sections'], 5)  # Max -25 points
            recommendations.append("Revisar secciones vacías - pueden ser encabezados mal detectados")
        
        # Check table extraction
        if stats['total_tables'] > 0:
            if len(tables_data) < stats['total_tables']:
                warnings.append(f"⚠️ No se extrajeron todas las tablas ({len(tables_data)}/{stats['total_tables']})")
                quality_score -= 10
        
        # Check hierarchy integrity
        if 'header_1' in df.columns and 'header_2' in df.columns:
            # Check for level jumps (H1 -> H3, skipping H2)
            h1_sections = df[df['header_1'].notna() & (df['header_1'] != '')]
            for idx, row in h1_sections.iterrows():
                # Find next header after this one
                next_rows = df.iloc[idx+1:idx+10]  # Check next 10 rows
                next_headers = next_rows[next_rows['source_type'] == 'heading']
                if not next_headers.empty:
                    next_header = next_rows.iloc[0]
                    # Check if there's a level jump
                    if next_header.get('header_3') and not next_header.get('header_2'):
                        warnings.append(f"⚠️ Salto de nivel detectado: {row.get('header_1', '')} → {next_header.get('header_3', '')}")
                        quality_score -= 5
        
        quality_score = max(0, quality_score)  # Don't go below 0
        
        return {
            'quality_score': quality_score,
            'warnings': warnings,
            'recommendations': recommendations,
            'stats': stats
        }
    
    # Function to split text into chunks respecting the token limit
    def split_text_into_chunks(text, max_completion_tokens=7000):
        import re
        # Split by paragraphs first
        paragraphs = text.split('\n')
        chunks = []
        current_chunk = []
        current_length = 0

        # Rough estimate: 1 token ≈ 4 characters in Spanish
        tokens_per_char = 0.25

        for para in paragraphs:
            # Estimate tokens in this paragraph
            para_tokens = len(para) * tokens_per_char

            # If adding this paragraph would exceed the max, start a new chunk
            if current_length + para_tokens > max_completion_tokens and current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = [para]
                current_length = para_tokens
            else:
                current_chunk.append(para)
                current_length += para_tokens

        # Add the last chunk if there's content
        if current_chunk:
            chunks.append('\n'.join(current_chunk))

        return chunks
    engagement_rubric = {}
    performance_rubric = {}
    parteval_rubric = {}
    gender_rubric = {}
    tj_traditional_rubric = {}
    tj_just_transition_rubric = {}

    try:
        df_rubric_engagement = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_engagement')
        df_rubric_engagement.drop(columns=['Unnamed: 0', 'Criterio'], inplace=True, errors='ignore')
        for idx, row in df_rubric_engagement.iterrows():
            indicador = row['Indicador']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
            engagement_rubric[indicador] = {'valores': valores, 'dimension': dimension}

        df_rubric_performance = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_performance')
        df_rubric_performance.drop(columns=['dimension'], inplace=True, errors='ignore')
        for idx, row in df_rubric_performance.iterrows():
            criterio = row['subdim']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['subdim', 'Dimensión'], errors='ignore').values.tolist()
            performance_rubric[criterio] = {'valores': valores, 'dimension': dimension}

        df_rubric_parteval = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_parteval')
        df_rubric_parteval.drop(columns=['Criterio'], inplace=True, errors='ignore')
        for idx, row in df_rubric_parteval.iterrows():
            indicador = row['Indicador']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
            parteval_rubric[indicador] = {'valores': valores, 'dimension': dimension}

        df_rubric_gender = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_gender_')
        df_rubric_gender.drop(columns=['Criterio'], inplace=True, errors='ignore')
        for idx, row in df_rubric_gender.iterrows():
            indicador = row['Indicador']
            dimension = row.get('Dimensión', 'No especificada')
            valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
            gender_rubric[indicador] = {'valores': valores, 'dimension': dimension}

        try:
            df_rubric_tj_traditional = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_TJ_Traditional')
            df_rubric_tj_traditional.drop(columns=['Criterio'], inplace=True, errors='ignore')
            for idx, row in df_rubric_tj_traditional.iterrows():
                indicador = row['Indicador']
                if pd.notna(indicador) and str(indicador).strip():
                    dimension = row.get('Dimensión', 'No especificada')
                    valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
                    tj_traditional_rubric[indicador] = {'valores': valores, 'dimension': dimension}
        except Exception as e:
            st.error(f"Error cargando TJ Tradicional: {e}")

        try:
            df_rubric_tj_just_transition = pd.read_excel('./Rubricas_6ago2025.xlsx', sheet_name='rubric_TJ_TJ')
            df_rubric_tj_just_transition.drop(columns=['Criterio'], inplace=True, errors='ignore')
            for idx, row in df_rubric_tj_just_transition.iterrows():
                indicador = row['Indicador']
                if pd.notna(indicador) and str(indicador).strip():
                    dimension = row.get('Dimensión', 'No especificada')
                    valores = row.drop(['Indicador', 'Dimensión'], errors='ignore').values.tolist()
                    tj_just_transition_rubric[indicador] = {'valores': valores, 'dimension': dimension}
        except Exception as e:
            st.error(f"Error cargando TJ Transición Justa: {e}")
    except Exception as e:
        st.error(f"Error leyendo las rúbricas: {e}")

    # Show rubric status
    st.success(f"""
    **ESTADO DE RÚBRICAS:**
    
    - Metodologías con enfoque participativo ({len(parteval_rubric)} criterios disponibles)", 
    - Integración del Enfoque de Género: {len(gender_rubric)} criterios
    - Integración del Enfoque de Transición Justa: Enfoque Moderno: {len(tj_just_transition_rubric)} criterios 
    """)

    # Download button for the rubric file (directly on page, no expander)
    try:
        with open('./Rubricas_6ago2025.xlsx', 'rb') as f:
            st.download_button(
                label="📥 Descargar archivo rúbrica de Atributos específicos",
                data=f,
                file_name="Rubricas_6ago2025.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_rubricas_tab2"
            )
    except FileNotFoundError:
        st.warning("Archivo de rúbricas no disponible para descarga.")

    # Document upload
    st.markdown("### 📄 Carga de Documento")
    
    # Warning box about document requirements
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos:**
    
    **📝 Formato del documento:**
    - Solo se aceptan archivos en formato **.docx** (Word 2007 o posterior)
    - El documento debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del documento deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente y las secciones no se identificarán.
    - Esto es especialmente importante para evaluaciones y PRODOCs que deben tener una estructura clara con secciones bien definidas
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos que excedan este límite serán truncados automáticamente
    - Se recomienda dividir documentos muy extensos (más de ~180 páginas) en secciones más pequeñas si es necesario
    
    **✅ Mejores prácticas:**
    - Usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones principales
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que el documento esté guardado correctamente antes de subirlo
    - Verifica que todas las secciones importantes tengan encabezados antes de procesar
    """)
    
    uploaded_file = st.file_uploader("Suba un archivo DOCX para evaluación:", type=["docx"], key="tab2_file_uploader")

    # Initialize session state for selections and results persistence
    if 'selected_rubrics_tab2' not in st.session_state:
        st.session_state['selected_rubrics_tab2'] = []
    if 'selected_criteria_tab2' not in st.session_state:
        st.session_state['selected_criteria_tab2'] = {}
    if 'tab2_results' not in st.session_state:
        st.session_state['tab2_results'] = None
    if 'document_extracted_tab2' not in st.session_state:
        st.session_state['document_extracted_tab2'] = False

    # Document Extraction Section
    st.markdown("---")
    st.markdown("### 📥 Extracción de Documento")
    
    if uploaded_file is not None:
        file_hash = hash(uploaded_file.getvalue())
        file_changed = st.session_state.get('last_file_hash_tab2') != file_hash
        
        if file_changed:
            st.session_state['document_extracted_tab2'] = False
            st.session_state['last_file_hash_tab2'] = None
        
        if st.button("🔍 Extraer Documento", key="extract_document_tab2", type="primary"):
            if uploaded_file is None:
                st.error("Por favor suba un archivo DOCX primero.")
                st.stop()
            
            with st.spinner("Extrayendo documento..."):
                try:
                    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp_file.write(uploaded_file.read())
                    tmp_file.close()
                    
                    progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                    doc_result = docx2python(tmp_file.name)
                    
                    # Use enhanced extraction
                    df, tables_data, extraction_stats = extract_docx_structure_enhanced(tmp_file.name)
                    progress_bar.progress(0.2, text="Documento cargado. Validando extracción...")
                    
                    # Validate extraction
                    validation_results = validate_extraction(df, tables_data, extraction_stats)
                    progress_bar.progress(0.3, text="Extracción validada. Procesando secciones...")
                    
                    # Extract sections
                    header_1_values = df['header_1'].dropna().unique()
                    llm_summary_rows = []
                    
                    for idx, header in enumerate(header_1_values):
                        section_df = df[df['header_1'] == header].copy()
                        # Extract text directly - already clean from extract_docx_structure_enhanced
                        full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                        # Calculate section stats
                        section_words = len(full_text.split())
                        section_paras = len(section_df[section_df['source_type'] == 'paragraph'])
                        section_tables = len(section_df[section_df['source_type'] == 'table'])
                        
                        llm_summary_rows.append({
                            'header_1': header,
                            'llm_paragraph': full_text if full_text else "",
                            'n_words': section_words,
                            'n_paragraphs': section_paras,
                            'n_tables': section_tables
                        })

                    progress_bar.progress(0.5, text="Secciones extraídas.")
                    
                    # Create exploded dataframe
                    llm_summary_df = pd.DataFrame(llm_summary_rows)
                    exploded_df = llm_summary_df.assign(
                        llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                    ).explode('llm_paragraph')
                    exploded_df = exploded_df.reset_index(drop=True)
                    exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                    
                    # Get full text
                    full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                    
                    # Store in session state
                    file_size = os.path.getsize(tmp_file.name)
                    n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                    n_paragraphs = len(exploded_df)
                    
                    st.session_state['full_document_text_tab2'] = full_document_text
                    st.session_state['document_stats_tab2'] = {
                        'file_size': file_size,
                        'n_words': n_words,
                        'n_paragraphs': n_paragraphs
                    }
                    st.session_state['exploded_df_tab2'] = exploded_df
                    st.session_state['extraction_df_tab2'] = df
                    st.session_state['tables_data_tab2'] = tables_data
                    st.session_state['extraction_stats_tab2'] = extraction_stats
                    st.session_state['validation_results_tab2'] = validation_results
                    st.session_state['sections_df_tab2'] = llm_summary_df
                    st.session_state['selected_sections_tab2'] = list(header_1_values)  # Select all by default
                    st.session_state['last_file_hash_tab2'] = file_hash
                    st.session_state['document_extracted_tab2'] = True
                    
                    try:
                        os.unlink(tmp_file.name)
                    except:
                        pass
                    
                    progress_bar.progress(1.0, text="Extracción completa.")
                    st.rerun()  # Rerun to show extraction results
                    
                except Exception as e:
                    st.error(f"Error procesando el documento: {e}")
                    import traceback
                    st.error(traceback.format_exc())
                    st.stop()
        
        # Show extraction results if document is extracted
        if st.session_state.get('document_extracted_tab2', False) and not file_changed:
            st.success("✅ Documento extraído con éxito")
            
            sections_df = st.session_state.get('sections_df_tab2', pd.DataFrame())
            
            if not sections_df.empty:
                header_1_values = sections_df['header_1'].tolist()
                
                # Section selector - simplified, no warnings or diagnostics
                st.markdown("### 🔍 Selección de Secciones para Evaluación")
                st.info("Selecciona las secciones que deseas incluir en la evaluación. Por defecto, todas las secciones están seleccionadas.")
                
                # Guidance about section extraction
                if len(header_1_values) == 0:
                    st.error("⚠️ **No se detectaron secciones en el documento.** Esto puede deberse a que el documento no usa estilos de encabezado de Word (Heading 1, Heading 2, etc.). Por favor, verifica que tu documento tenga encabezados formateados correctamente.")
                elif len(header_1_values) < 3:
                    st.warning("⚠️ **Se detectaron pocas secciones** en el documento. Si esperabas más secciones, verifica que el documento use estilos de encabezado de Word (Heading 1, Heading 2, etc.) para identificar las secciones principales.")
                
                # Initialize selected sections if not exists
                if 'selected_sections_tab2' not in st.session_state:
                    st.session_state['selected_sections_tab2'] = list(header_1_values)
                
                # Section selection interface
                selected_sections = st.session_state.get('selected_sections_tab2', list(header_1_values)).copy()
                col1, col2 = st.columns([3, 1])
                
                with col2:
                    if st.button("✅ Seleccionar Todas", key="select_all_sections_tab2"):
                        st.session_state['selected_sections_tab2'] = list(header_1_values)
                        st.rerun()
                    
                    if st.button("❌ Deseleccionar Todas", key="deselect_all_sections_tab2"):
                        st.session_state['selected_sections_tab2'] = []
                        st.rerun()
                
                with col1:
                    st.markdown("**Secciones disponibles:**")
                    for section in header_1_values:
                        section_info = sections_df[sections_df['header_1'] == section].iloc[0]
                        is_selected = section in selected_sections
                        
                        checkbox_label = f"**{section}** ({section_info['n_words']:,} palabras, {section_info['n_paragraphs']} párrafos)"
                        
                        checkbox_key = f"section_checkbox_{section}_tab2"
                        new_selection = st.checkbox(checkbox_label, value=is_selected, key=checkbox_key)
                        
                        if new_selection and section not in selected_sections:
                            selected_sections.append(section)
                        elif not new_selection and section in selected_sections:
                            selected_sections.remove(section)
                        
                        # Add expandable preview of extracted content
                        with st.expander(f"👁️ Ver contenido: {section}", expanded=False):
                            section_content = section_info['llm_paragraph']
                            if section_content and section_content.strip():
                                # Show first 500 characters as preview, full content in expandable
                                preview_text = section_content[:500] + "..." if len(section_content) > 500 else section_content
                                st.text_area(
                                    "Contenido extraído:",
                                    value=section_content,
                                    height=200,
                                    key=f"content_preview_{section}_tab2",
                                    label_visibility="collapsed"
                                )
                                st.caption(f"Total: {len(section_content):,} caracteres")
                            else:
                                st.info("Esta sección no tiene contenido extraído.")
                            
                            # Show table-extracted text
                            tables_data = st.session_state.get('tables_data_tab2', [])
                            section_tables = [t for t in tables_data if t.get('section') == section]
                            
                            if section_tables:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                for table_info in section_tables:
                                    table_num = table_info.get('table_number', 'N/A')
                                    table_data = table_info.get('data', [])
                                    if table_data:
                                        # Format table as text
                                        table_text = '\n'.join([' | '.join(str(cell) for cell in row) for row in table_data])
                                        st.text_area(
                                            f"Tabla {table_num}:",
                                            value=table_text,
                                            height=150,
                                            key=f"table_preview_{section}_table{table_num}_tab2",
                                            label_visibility="collapsed"
                                        )
                                        st.caption(f"Tabla {table_num}: {len(table_data)} filas, {len(table_data[0]) if table_data else 0} columnas")
                            else:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                st.info("No se encontraron tablas en esta sección.")
                
                # Update session state
                st.session_state['selected_sections_tab2'] = selected_sections
                
                # Show selection summary - simplified
                if selected_sections:
                    selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
                    total_selected_words = selected_df['n_words'].sum()
                    total_selected_paras = selected_df['n_paragraphs'].sum()
                    
                    # Estimate tokens
                    if encoding:
                        selected_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
                        estimated_tokens = len(encoding.encode(selected_text))
                    else:
                        estimated_tokens = total_selected_words * 1.2  # Rough estimate
                    
                    # Warn if approaching limit
                    if estimated_tokens > 100000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.warning(f"⚠️ **Advertencia de límite de contexto:** Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). El sistema procesa hasta 110,000 tokens (aproximadamente 150-200 páginas). Si el documento excede este límite, será truncado automáticamente.")
                    elif estimated_tokens > 80000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.info(f"ℹ️ Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). Estás dentro del límite de 110,000 tokens (aproximadamente 150-200 páginas).")
                    
                    st.success(f"✅ {len(selected_sections)} secciones seleccionadas | "
                              f"{total_selected_words:,} palabras | "
                              f"~{estimated_tokens:,} tokens estimados")
    
    # Rubric and Criteria Selection Section (moved after document extraction)
    st.markdown("---")
    st.markdown("### 📋 Selección de Rúbricas y Criterios")
    
    # Only show rubric selection if document is extracted
    if st.session_state.get('document_extracted_tab2', False):
        # All available rubrics
        all_rubrics = {
            # "Participación de Actores (durante el proyecto)": engagement_rubric,  # Commented out per user request
            # "Desempeño del proyecto (según informe de evaluación)": performance_rubric,  # Commented out per user request
            "Participación durante la evaluación (metodología)": parteval_rubric,
            "Integración del Enfoque de Género": gender_rubric,
            # "Transición Justa: Enfoque Tradicional": tj_traditional_rubric,  # Commented out per user request
            "Integración del Enfoque de Transición Justa: Enfoque Moderno": tj_just_transition_rubric
        }

        # Step 1: Select Rubrics
        st.markdown("#### 1. Seleccione las rúbricas a aplicar:")
        selected_rubric_names = st.multiselect(
            "Rúbricas:",
            options=list(all_rubrics.keys()),
            default=st.session_state['selected_rubrics_tab2'],
            key='rubric_selector_tab2'
        )
        st.session_state['selected_rubrics_tab2'] = selected_rubric_names

        # Step 2: Select Criteria within each rubric
        if selected_rubric_names:
            st.markdown("#### 2. Seleccione los criterios específicos:")
            
            for rubric_name in selected_rubric_names:
                rubric_dict = all_rubrics[rubric_name]
                
                with st.expander(f"📋 {rubric_name} ({len(rubric_dict)} criterios disponibles)", expanded=True):
                    # Select all checkbox for this rubric
                    select_all_key = f"select_all_{rubric_name}_tab2"
                    select_all = st.checkbox(f"Seleccionar todos los criterios", key=select_all_key)
                    
                    # Initialize criteria selection for this rubric
                    if rubric_name not in st.session_state['selected_criteria_tab2']:
                        st.session_state['selected_criteria_tab2'][rubric_name] = []
                    
                    # Show criteria checkboxes
                    selected_criteria = []
                    for criterion in rubric_dict.keys():
                        # Default checked if select_all or previously selected
                        default_value = select_all or criterion in st.session_state['selected_criteria_tab2'][rubric_name]
                        
                        is_selected = st.checkbox(
                            f"{criterion}",
                            value=default_value,
                            key=f"criterion_{rubric_name}_{criterion}_tab2"
                        )
                        
                        if is_selected:
                            selected_criteria.append(criterion)
                    
                    # Update session state
                    st.session_state['selected_criteria_tab2'][rubric_name] = selected_criteria
                    
                    st.info(f"Criterios seleccionados: {len(selected_criteria)}/{len(rubric_dict)}")

        # Show summary of selections
        if selected_rubric_names:
            total_criteria = sum(len(st.session_state['selected_criteria_tab2'].get(r, [])) for r in selected_rubric_names)
            st.success(f"Total: {len(selected_rubric_names)} rúbricas, {total_criteria} criterios seleccionados")
        else:
            st.info("ℹ️ Selecciona al menos una rúbrica para continuar con la evaluación.")
    else:
        st.info("ℹ️ Por favor extrae el documento primero para poder seleccionar las rúbricas y criterios.")
    
    # Process and Evaluate button
    st.markdown("---")
    st.markdown("### ⚙️ Procesamiento y Evaluación")

    def evaluate_criterion_with_llm(document_text, criterion, descriptions, max_retries=3):
        """Analyze document against criterion with retry logic"""
        import time

        for attempt in range(max_retries):
            try:
                # Truncate to ~110K tokens to maximize context while leaving room for:
                # - System prompt (~50 tokens)
                # - User prompt template + criterion + scoring levels (~500-2000 tokens)
                # - Response tokens (6500 tokens)
                # - Safety buffer (~5000 tokens)
                # Total: 128K - 50 - 2000 - 6500 - 5000 = ~110K tokens for document
                combined_text = truncate_to_token_limit(document_text, max_tokens=110000, encoding_obj=encoding)

                # Now do the expensive analysis on focused content
                prompt = f"""Evaluate this document against: {criterion}

    Scoring levels: {json.dumps(descriptions)}

    Relevant document sections:
    {combined_text}

    IMPORTANTE: Proporciona tu respuesta SIEMPRE en español, incluso si el documento está en inglés.
    
    Provide JSON with:
    {{"analysis": "detailed 2-3 paragraphs IN SPANISH", "score": 1-5, "evidence": ["quote 1", "quote 2", "quote 3", "etc - 5-8 key quotes from the text as an array"]}}"""

                response = client.chat.completions.create(
                    model="gpt-5-mini",
                    messages=[
                        {"role": "system", "content": "Eres un evaluador experto de documentos. Siempre debes responder en español, incluso si el documento está en inglés."},
                        {"role": "user", "content": prompt}
                    ],
                    max_completion_tokens=6500,
                    reasoning_effort="minimal",
                    timeout=120  # 2 minute timeout per request
                )

                content = response.choices[0].message.content.strip()
                # Remove markdown code fences if present
                if content.startswith('```'):
                    # Remove opening fence (```json or ```)
                    content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                    # Remove closing fence
                    if content.endswith('```'):
                        content = content.rsplit('```', 1)[0]
                    content = content.strip()

                result = json.loads(content)
                # Normalize evidence field: convert array to string if needed
                if isinstance(result.get('evidence'), list):
                    result['evidence'] = '\n'.join(result['evidence'])
                return result

            except json.JSONDecodeError as e:
                # If JSON parsing fails, return a default structure
                return {
                    "analysis": f"Failed to parse JSON: {str(e)}. Raw response: {response.choices[0].message.content[:200]}",
                    "score": 3,
                    "evidence": "Unable to parse structured response",
                    "error": f"JSON parsing error: {str(e)}"
                }
            except Exception as e:
                # Check if it's a rate limit error
                error_msg = str(e)
                if "rate_limit" in error_msg.lower() or "429" in error_msg:
                    if attempt < max_retries - 1:
                        wait_time = (2 ** attempt) * 2  # Exponential backoff: 2s, 4s, 8s
                        time.sleep(wait_time)
                        continue

                # If last attempt or non-rate-limit error, return error
                return {
                    "analysis": f"Error during evaluation: {error_msg}",
                    "score": 0,
                    "evidence": "",
                    "error": f"API error (attempt {attempt + 1}/{max_retries}): {error_msg}"
                }

        # If we exhausted all retries
        return {
            "analysis": "Failed after multiple retry attempts",
            "score": 0,
            "evidence": "",
            "error": f"Failed after {max_retries} attempts"
        }

    # Function to evaluate a single text chunk
    def evaluate_single_chunk(text_chunk, criterion, descriptions):
        """Evaluate a single text chunk against a criterion with expanded analysis and evidence"""
        import json

        # Build prompt
        prompt = f"""
        Estás evaluando un documento contra un criterio específico.
        
        Criterio: {criterion}
        
        Descripciones de los niveles de puntuación:
        {json.dumps(descriptions, indent=2)}
        
        Contenido del documento a evaluar:
        {text_chunk}
        
        Analiza qué tan bien el documento cumple con este criterio. Proporciona:
        
        1. Un análisis DETALLADO (2-3 párrafos) que explique a fondo el razonamiento detrás de tu evaluación. Proporciona un razonamiento profundo que abarque los aspectos del criterio.
        
        2. Una puntuación de 1-5 (donde 1 es la más baja y 5 es la más alta).
        
        3. EVIDENCIA del documento que respalde tu puntuación. Incluye entre 5-8 citas textuales del documento, indicando cómo cada fragmento contribuye a tu evaluación.
        
        Formatea tu respuesta como un objeto JSON con las siguientes claves:
        {{"analysis": "tu análisis detallado aquí", "score": puntuación_numérica_entre_1_y_5, "evidence": "citas textuales del documento (5-8 párrafos)"}}
        
        Devuelve solo el objeto JSON, nada más.
        """

        # Call LLM using new Responses API
        try:
            response = client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "system", "content": "Eres un experto evaluador de documentos que proporciona análisis detallados basados en criterios específicos. Tu evidencia cita fragmentos del texto original."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                max_completion_tokens=6500,
                reasoning_effort="minimal"
            )
            raw = response.choices[0].message.content
            if not raw or not raw.strip():
                return {'score': 0, 'analysis': 'Empty response from API', 'evidence': ''}

            # Remove markdown code fences if present
            content = raw.strip()
            if content.startswith('```'):
                content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                if content.endswith('```'):
                    content = content.rsplit('```', 1)[0]
                content = content.strip()

            parsed = json.loads(content)
            # Normalize evidence field: convert array to string if needed
            if isinstance(parsed.get('evidence'), list):
                parsed['evidence'] = '\n'.join(parsed['evidence'])
            return parsed
        except Exception as e:
            return {'score': 0, 'analysis': f'Error: {str(e)}', 'evidence': ''}

    # Function to synthesize evaluations
    def synthesize_evaluations(chunk_results, criterion, descriptions):
        """Synthesize evaluations from multiple document chunks with expanded analysis and evidence"""
        import json

        # Extract and format the individual evaluations for the synthesis
        individual_evals = []
        all_evidence = []

        for i, result in enumerate(chunk_results):
            individual_evals.append(f"Evaluación del fragmento {i+1}:\n" +
                                    f"Puntuación: {result.get('score', 0)}\n" +
                                    f"Análisis: {result.get('analysis', '')}")

            # Collect all evidence
            evidence = result.get('evidence', '')
            if evidence:
                all_evidence.append(f"Evidencia del fragmento {i+1}:\n{evidence}")

        # Define separator outside the f-string to avoid backslash issues
        separator = "\n\n"

        # Create a synthesis prompt
        synthesis_prompt = f"""
        Has evaluado un documento dividido en múltiples fragmentos contra el criterio: {criterion}
        
        Aquí están las evaluaciones individuales de cada fragmento:
        
        {separator.join(individual_evals)}
        
        Basándote en estas evaluaciones individuales, proporciona:
        
        1. Un análisis DETALLADO (2-3 párrafos) que integre los hallazgos clave de todos los fragmentos. Este análisis debe ser comprensivo y abarcar los aspectos relevantes encontrados en el documento.
        
        2. Una puntuación general de 1-5 (puedes promediar las puntuaciones o ajustar según sea necesario)
        
        3. Las evidencias más importantes del documento. Selecciona las 8-10 citas textuales más relevantes de los fragmentos individuales.
        
        Formatea tu respuesta como un objeto JSON con las siguientes claves:
        {{"analysis": "tu análisis global detallado aquí", "score": puntuación_general_entre_1_y_5, "evidence": "las citas textuales más relevantes del documento (8-10 párrafos)"}}
        
        Devuelve solo el objeto JSON, nada más.
        """

        # Call LLM for synthesis using new Responses API
        try:
            response = client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "system", "content": "Eres un experto evaluador de documentos que sintetiza análisis de múltiples fragmentos de texto para producir evaluaciones detalladas con evidencia textual."},
                    {"role": "user", "content": synthesis_prompt}
                ],
                response_format={"type": "json_object"},
                max_completion_tokens=6500,
                reasoning_effort="minimal"
            )
            raw = response.choices[0].message.content
            if not raw or not raw.strip():
                raise ValueError("Empty response from API")

            # Remove markdown code fences if present
            content = raw.strip()
            if content.startswith('```'):
                content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                if content.endswith('```'):
                    content = content.rsplit('```', 1)[0]
                content = content.strip()

            parsed = json.loads(content)
            # Normalize evidence field: convert array to string if needed
            if isinstance(parsed.get('evidence'), list):
                parsed['evidence'] = '\n'.join(parsed['evidence'])
            return parsed
        except Exception as e:
            # If synthesis fails, combine results manually in a more limited way
            avg_score = sum(r.get('score', 0) for r in chunk_results) / len(chunk_results)
            # Take only the first paragraph of each analysis to avoid token limits
            analysis_parts = []
            for r in chunk_results:
                analysis = r.get('analysis', '')
                first_para = analysis.split('\n\n')[0] if '\n\n' in analysis else analysis
                analysis_parts.append(first_para)

            # Take only the first few evidence items
            evidence_parts = []
            evidence_count = 0
            for evidence in all_evidence:
                parts = evidence.split('\n\n')
                # Add up to 2 evidence parts per chunk
                for part in parts[:2]:
                    if evidence_count < 8:  # Limit to 8 total evidence parts
                        evidence_parts.append(part)
                        evidence_count += 1

            return {
                'score': avg_score,
                'analysis': separator.join(analysis_parts),
                'evidence': separator.join(evidence_parts)
            }
    
    if st.button('🚀 Procesar y Evaluar', key='process_evaluate_tab2', type="primary"):
        # Check prerequisites
        if not st.session_state.get('document_extracted_tab2', False):
            st.error("❌ Por favor extrae el documento primero usando el botón 'Extraer Documento'.")
            st.stop()
        
        # Get selected rubrics from session state
        selected_rubric_names = st.session_state.get('selected_rubrics_tab2', [])
        if not selected_rubric_names:
            st.error("Por favor seleccione al menos una rúbrica.")
            st.stop()
        
        # Calculate total criteria
        total_criteria = sum(len(st.session_state.get('selected_criteria_tab2', {}).get(r, [])) for r in selected_rubric_names)
        if total_criteria == 0:
            st.error("Por favor seleccione al menos un criterio.")
            st.stop()
        
        # Check if document needs re-extraction (shouldn't happen, but safety check)
        if uploaded_file is not None:
            file_hash = hash(uploaded_file.getvalue())
            if st.session_state.get('last_file_hash_tab2') != file_hash:
                st.warning("⚠️ El documento ha cambiado. Por favor extrae el documento nuevamente.")
                st.stop()
        
        # Evaluate with selected rubrics and criteria
        # Get selected sections or use full document
        selected_sections = st.session_state.get('selected_sections_tab2', [])
        sections_df = st.session_state.get('sections_df_tab2', pd.DataFrame())
        
        if selected_sections and not sections_df.empty:
            # Filter to selected sections only
            selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
            document_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
            st.info(f"📌 Evaluando {len(selected_sections)} secciones seleccionadas")
        else:
            # Fallback to full document
            document_text = st.session_state.get('full_document_text_tab2', '')
            if not selected_sections:
                st.warning("⚠️ No hay secciones seleccionadas. Usando documento completo.")
        
        if not document_text:
            st.error("No se pudo recuperar el texto del documento.")
            st.stop()

        # Build filtered rubrics based on selection
        # Get selected rubrics from session state
        selected_rubric_names = st.session_state.get('selected_rubrics_tab2', [])
        
        # Define all_rubrics for evaluation
        all_rubrics = {
            "Participación durante la evaluación (metodología)": parteval_rubric,
            "Integración del Enfoque de Género": gender_rubric,
            "Integración del Enfoque de Transición Justa: Enfoque Moderno": tj_just_transition_rubric
        }
        
        rubrics_to_evaluate = []
        for rubric_name in selected_rubric_names:
            selected_criteria_list = st.session_state['selected_criteria_tab2'].get(rubric_name, [])
            if selected_criteria_list:
                # Filter the rubric to only include selected criteria
                full_rubric = all_rubrics[rubric_name]
                filtered_rubric = {k: v for k, v in full_rubric.items() if k in selected_criteria_list}
                rubrics_to_evaluate.append((rubric_name, filtered_rubric))

        if not rubrics_to_evaluate:
            st.error("No hay criterios seleccionados para evaluar.")
            st.stop()

        # Evaluate
        st.info("Iniciando evaluación de criterios...")
        rubric_results = []
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import time
        MAX_WORKERS = 3  # Reduced to avoid rate limiting
        
        def eval_one_criterion(args):
            crit, descriptions, dimension, rubric_name = args
            try:
                result = evaluate_criterion_with_llm(document_text, crit, descriptions)
                # Ensure result is a dictionary
                if not isinstance(result, dict):
                    result = {'score': 0, 'analysis': str(result), 'evidence': '', 'error': 'Invalid result format'}
                return {
                    'Criterio': crit,
                    'Dimensión': dimension,
                    'Score': result.get('score', 0),
                    'Análisis': str(result.get('analysis', '')),
                    'Evidencia': str(result.get('evidence', '')),
                    'Error': str(result.get('error', '')) if 'error' in result else '',
                    'Rúbrica': rubric_name
                }
            except Exception as e:
                return {
                    'Criterio': crit,
                    'Dimensión': dimension,
                    'Score': 0,
                    'Análisis': '',
                    'Evidencia': '',
                    'Error': str(e),
                    'Rúbrica': rubric_name
                }
        
        for rubric_name, rubric_dict in rubrics_to_evaluate:
            if not rubric_dict:
                continue

            rubric_analysis_data = []
            n_criteria = len(rubric_dict)
            progress = st.progress(0, text=f"Preparando evaluación de {n_criteria} criterios para: {rubric_name}...")

            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                progress.progress(0.01, text=f"Enviando {n_criteria} criterios para evaluación paralela...")
                futures = {}
                for idx, (crit, rubric_data) in enumerate(rubric_dict.items()):
                    future = executor.submit(eval_one_criterion, (
                        crit,
                        rubric_data['valores'] if isinstance(rubric_data, dict) else rubric_data,
                        rubric_data.get('dimension', 'No especificada') if isinstance(rubric_data, dict) else 'No especificada',
                        rubric_name
                    ))
                    futures[future] = (crit, idx)
                    # Small delay to avoid overwhelming API at startup
                    if idx > 0 and idx % MAX_WORKERS == 0:
                        time.sleep(0.5)

                progress.progress(0.05, text=f"Esperando resultados de evaluación...")
                completed = 0
                for future in as_completed(futures):
                    result = future.result()
                    rubric_analysis_data.append(result)
                    completed += 1
                    crit, idx = futures[future]
                    progress.progress(0.05 + (completed / n_criteria * 0.95), text=f"Completado {completed}/{n_criteria}: {crit}")
            
            rubric_results.append((rubric_name, pd.DataFrame(rubric_analysis_data)))
        
        # Store results in session state for persistence
        st.session_state['tab2_results'] = rubric_results
        
        # Display results
        if rubric_results:
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''
                    
                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Criterio', 'Dimensión', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Download ZIP
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    # Export as XLSX instead of CSV
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_rubrica_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_rubricas.zip",
                mime="application/zip"
            )
        else:
            st.warning("No se generaron resultados para ninguna rúbrica.")
    else:
        # Check if there are persisted results in session state
        if st.session_state.get('tab2_results') is not None:
            rubric_results = st.session_state['tab2_results']
            
            st.markdown("### 📊 Resultados guardados")
            
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''
                    
                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Criterio', 'Dimensión', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Download ZIP
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_rubrica_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_rubricas.zip",
                mime="application/zip",
                key="tab2_download_persisted"
            )
            
            # Clear results button
            if st.button("🗑️ Limpiar resultados", key="clear_tab2_results"):
                st.session_state['tab2_results'] = None
                st.rerun()
        else:
            st.info("Suba un documento, seleccione rúbricas y criterios, luego presione 'Procesar y Evaluar'.")

#===================######################=====================
# ================== TAB 5: DOCUMENT CHAT =====================
with tab4:
    st.header("Pregúntale a tus Documentos")

    # Presentation box
    st.info("""

    Arrastra uno o más DOCX/TXT (max 200MB) y conversa directamente con su contenido. Este chat responde solo con la información de los archivos cargados—no recurre a fuentes externas—para ayudarte a revisar propuestas antes de enviarlas, aclarar pasajes de informes, preparar notas técnicas y comprobar coherencias entre objetivos, actividades y resultados.

    Al formular tus preguntas, indica el nivel de detalle que necesitas y pide que las respuestas incluyan citas breves entre comillas y metadatos (título, sección/página, año) para mantener la trazabilidad. Si un dato no existe en los archivos, se marcará "ND" sin inventar información.

    **Úsalo para:**
    - Aprender de experiencias de evaluación
    - Profundizar en el conocimiento de secciones de los informes
    - Preparar notas técnicas a partir de esta evidencia
    - Verificar la coherencia entre objetivos, actividades y resultados
    - Generar tablas copiables a Excel
    - Comparar varios documentos (hasta de 200MB) en una misma conversación

    *Nota: El chat mantiene memoria de la conversación durante la sesión activa.*
    """)

    # Session state for chat and document
    if 'doc_chat_history' not in st.session_state:
        st.session_state['doc_chat_history'] = []
    if 'doc_chat_docs' not in st.session_state:
        st.session_state['doc_chat_docs'] = []

    # Warning box about document requirements for chat
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos:**
    
    **📝 Formatos aceptados:**
    - Archivos **.docx** (Word 2007 o posterior)
    - Archivos **.txt** (texto plano)
    - Puedes subir **múltiples archivos** (hasta 200MB en total)
    
    **📝 Para documentos Word (.docx):**
    - El documento debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del documento deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente.
    - Para documentos grandes, el sistema usa RAG (Retrieval Augmented Generation) automáticamente
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos muy grandes (>100K caracteres, aproximadamente >40 páginas) se procesan con RAG para mejor eficiencia
    - Documentos pequeños se procesan con contexto completo (más eficiente)
    - Documentos que excedan el límite máximo serán truncados automáticamente
    
    **✅ Mejores prácticas:**
    - Para documentos Word: usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que los documentos estén guardados correctamente antes de subirlos
    """)
    
    uploaded_files = st.file_uploader("Sube uno o más archivos DOCX o TXT para chatear:", type=["docx", "txt"], accept_multiple_files=True)
    if 'doc_chat_docs' not in st.session_state:
        st.session_state['doc_chat_docs'] = []
    if uploaded_files:
        # Only reset docs and chat history if files changed
        uploaded_filenames = sorted([f.name for f in uploaded_files])
        existing_filenames = sorted([doc['filename'] for doc in st.session_state['doc_chat_docs']]) if st.session_state['doc_chat_docs'] else []
        if uploaded_filenames != existing_filenames:
            st.session_state['doc_chat_docs'] = []

            # Load documents
            for uploaded_file in uploaded_files:
                try:
                    if uploaded_file.name.endswith(".docx"):
                        from docx import Document
                        doc = Document(uploaded_file)
                        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    else:
                        full_text = uploaded_file.read().decode("utf-8", errors="ignore")
                    st.session_state['doc_chat_docs'].append({
                        "filename": uploaded_file.name,
                        "text": full_text
                    })
                except Exception as e:
                    st.error(f"Error al procesar el documento '{uploaded_file.name}': {str(e)}")

            # Calculate total document size for smart RAG decision
            total_text = "\n\n".join([doc['text'] for doc in st.session_state['doc_chat_docs']])
            total_chars = len(total_text)
            st.session_state['doc_chat_total_text'] = total_text
            st.session_state['doc_chat_total_chars'] = total_chars

            # Pre-compute embeddings only if documents are large (> 100K chars)
            # For small docs, we'll use full context directly (more efficient)
            if total_chars > 100000:
                with st.spinner("Procesando embeddings de documentos (solo se hace una vez)..."):
                    # Chunk documents
                    def chunk_text(text, chunk_size=2000, overlap=300):
                        chunks = []
                        start = 0
                        while start < len(text):
                            end = min(start + chunk_size, len(text))
                            chunks.append(text[start:end])
                            start += chunk_size - overlap
                        return chunks

                    all_chunks = []
                    for doc in st.session_state['doc_chat_docs']:
                        all_chunks.extend(chunk_text(doc['text']))

                    # Embed all chunks once and cache
                    try:
                        emb_model = "text-embedding-3-large"
                        chunk_embs_resp = client.embeddings.create(input=all_chunks, model=emb_model)
                        chunk_embs = [item.embedding for item in chunk_embs_resp.data]

                        # Cache in session state
                        st.session_state['doc_chat_chunks'] = all_chunks
                        st.session_state['doc_chat_embeddings'] = chunk_embs
                        st.session_state['doc_chat_use_rag'] = True
                        st.info(f"📚 Documentos grandes ({total_chars:,} chars) - usando RAG con {len(all_chunks)} fragmentos en caché")
                    except Exception as e:
                        st.error(f"Error al generar embeddings: {str(e)}")
                        st.session_state['doc_chat_use_rag'] = False
            else:
                st.session_state['doc_chat_use_rag'] = False
                st.info(f"📄 Documentos pequeños ({total_chars:,} chars) - usando contexto completo (más eficiente)")

            st.session_state['doc_chat_history'] = []  # Only reset chat when new files uploaded
            st.success(f"{len(st.session_state['doc_chat_docs'])} documento(s) cargado(s) y listo(s) para chatear.")

    # Show filenames and previews
    if st.session_state['doc_chat_docs']:
        st.info(f"Documentos activos: {', '.join([doc['filename'] for doc in st.session_state['doc_chat_docs']])}")
        for doc in st.session_state['doc_chat_docs']:
            with st.expander(f"Vista previa: {doc['filename']} (primeros 500 caracteres)"):
                st.write(doc['text'][:500] + ("..." if len(doc['text']) > 500 else ""))

    # Chat interface
    if st.session_state['doc_chat_docs']:
        with st.form("doc_chat_form", clear_on_submit=True):
            user_input = st.text_area("Escribe tu preguntas (para las respuestas se considera el texto completo de los documentos cargados):", key="doc_chat_input")
            submitted = st.form_submit_button("Enviar pregunta")
            if submitted and user_input.strip():
                # Add user message to history
                if 'doc_chat_history' not in st.session_state:
                    st.session_state['doc_chat_history'] = []
                st.session_state['doc_chat_history'].append({"role": "user", "content": user_input.strip()})

                question = user_input.strip()
                use_rag = st.session_state.get('doc_chat_use_rag', False)

                # Smart context selection: use full text for small docs, RAG for large docs
                if not use_rag:
                    # Small documents: use full context (up to 110K tokens)
                    # This is more efficient - no chunking, no embeddings, just direct LLM call
                    context = truncate_to_token_limit(
                        st.session_state.get('doc_chat_total_text', ''), 
                        max_tokens=110000, 
                        encoding_obj=encoding
                    )

                    messages = [
                        {"role": "system", "content": "Eres un asistente experto en análisis documental. Responde usando solo la información del documento proporcionado."},
                        {"role": "system", "content": f"Texto del documento:\n{context}"}
                    ]
                    for msg in st.session_state['doc_chat_history'][-5:]:
                        messages.append(msg)

                    try:
                        response = client.chat.completions.create(
                            model="gpt-5-mini",
                            messages=messages,
                            max_completion_tokens=4096,
                            reasoning_effort="minimal"
                        )
                        answer = response.choices[0].message.content.strip()
                        st.session_state['doc_chat_history'].append({"role": "assistant", "content": answer})
                    except Exception as e:
                        st.session_state['doc_chat_history'].append({"role": "assistant", "content": f"[Error al obtener respuesta: {str(e)}]"})

                else:
                    # Large documents: use RAG with cached embeddings
                    import numpy as np
                    import faiss
                    import re

                    # Get cached chunks and embeddings
                    all_chunks = st.session_state.get('doc_chat_chunks', [])
                    chunk_embs = st.session_state.get('doc_chat_embeddings', [])

                    if not all_chunks or not chunk_embs:
                        st.session_state['doc_chat_history'].append({"role": "assistant", "content": "[Error: No se encontraron fragmentos en caché.]"})
                    else:
                        # Embed user question (only embedding we need per query!)
                        try:
                            emb_model = "text-embedding-3-large"
                            question_emb = client.embeddings.create(input=question, model=emb_model).data[0].embedding
                        except Exception as e:
                            st.session_state['doc_chat_history'].append({"role": "assistant", "content": f"[Error al obtener embedding de la pregunta: {str(e)}]"})
                            question_emb = None

                        # Retrieve relevant chunks using FAISS
                        if question_emb:
                            dim = len(chunk_embs[0])
                            xb = np.array(chunk_embs).astype('float32')
                            index = faiss.IndexFlatIP(dim)

                            # Normalize for cosine similarity
                            faiss.normalize_L2(xb)
                            index.add(xb)  # FIX: Actually add vectors to the index!

                            xq = np.array([question_emb]).astype('float32')
                            faiss.normalize_L2(xq)

                            # Reduced from 50 to 15 for efficiency (15 chunks = ~30K chars)
                            top_n = min(15, len(all_chunks))
                            D, I = index.search(xq, top_n)
                            selected_chunks = [all_chunks[i] for i in I[0] if i < len(all_chunks)]

                            # Also retrieve chunks with exact keyword matches
                            stopwords = set([
                                'el','la','los','las','de','del','y','en','a','un','una','que','por','con','para','es','al','se','su','sus','o','u','como','más','menos','le','lo','su','the','and','of','in','to','for','is','on','at','by','an','or','as','be','are','was','were','from','it','this','that','with','but','not','can','may','do','does'
                            ])
                            qwords = [w for w in re.findall(r'\w+', question.lower()) if w not in stopwords and len(w) > 2]
                            keyword_chunks = []
                            for chunk in all_chunks:
                                chunk_lc = chunk.lower()
                                if any(qw in chunk_lc for qw in qwords):
                                    keyword_chunks.append(chunk)
                                    if len(keyword_chunks) >= 5:  # Limit keyword matches
                                        break

                            # Merge and deduplicate
                            seen = set()
                            merged_chunks = []
                            for chunk in selected_chunks + keyword_chunks:
                                if chunk not in seen:
                                    merged_chunks.append(chunk)
                                    seen.add(chunk)

                            context = '\n---\n'.join(merged_chunks)

                            # Send to LLM
                            messages = [
                                {"role": "system", "content": "Eres un asistente experto en análisis documental. Responde usando solo la información del documento proporcionado. Si la información no es explícita, infiere la respuesta usando pistas contextuales y tu capacidad de síntesis."},
                                {"role": "system", "content": f"Fragmentos relevantes del documento:\n{context}"}
                            ]
                            for msg in st.session_state['doc_chat_history'][-5:]:
                                messages.append(msg)

                            try:
                                response = client.chat.completions.create(
                                    model="gpt-5-mini",
                                    messages=messages,
                                    max_completion_tokens=4096,
                                    reasoning_effort="minimal"
                                )
                                answer = response.choices[0].message.content.strip()
                                st.session_state['doc_chat_history'].append({"role": "assistant", "content": answer})
                            except Exception as e:
                                st.session_state['doc_chat_history'].append({"role": "assistant", "content": f"[Error al obtener respuesta: {str(e)}]"})

        # Display chat history in a persistent, scrollable container
        st.markdown(
            '''
            <div style="height:600px; overflow-y:auto; border:1px solid #333; border-radius:8px; padding:1em; background:#18191a; margin-bottom:1em;">
            '''
            +
            "".join(
                f"<div style='margin-bottom:1em; color:#2980b9;'><b>Tú:</b> {msg['content']}</div>" if msg['role']=='user'
                else f"<div style='margin-bottom:1em; color:#27ae60;'><b>Asistente:</b> {msg['content']}</div>"
                for msg in st.session_state['doc_chat_history']
            )
            +
            '</div>',
            unsafe_allow_html=True
        )
    else:
        st.info("Sube uno o más documentos válidos para comenzar el chat.")

# # ================== TAB 6: EVALUACIÓN DE PRODOCS =====================
# with tab3:
#     st.header("Diagnóstico de Sostenibilidad del Proyecto")
    
#     # Read rubric from Excel file
#     import pandas as pd
#     prodoc_rubric = {}
    
#     try:
#         # Load rubric from PRODOC_rubric.xlsx
#         df_rubric_prodoc = pd.read_excel('./PRODOC_rubric.xlsx', sheet_name='rubric')
        
#         # Check if 'Indicador' column exists
#         if 'Indicador' not in df_rubric_prodoc.columns:
#             st.error("La columna 'Indicador' no existe en el archivo Excel.")
#             # Try to use the first column as 'Indicador' if it exists
#             if len(df_rubric_prodoc.columns) > 0:
#                 indicador_col = df_rubric_prodoc.columns[0]
#                 st.warning(f"Usando la columna '{indicador_col}' como columna de indicadores.")
#                 df_rubric_prodoc.rename(columns={indicador_col: 'Indicador'}, inplace=True)
#             else:
#                 prodoc_rubric = {}
#                 st.error("No se pudo encontrar una columna para los criterios.")
        
#         # Process each row to extract criteria and values
#         for idx, row in df_rubric_prodoc.iterrows():
#             # Get the indicator value
#             indicador = row['Indicador']
            
#             # Skip empty indicators
#             if pd.isna(indicador) or str(indicador).strip() == '':
#                 continue
            
#             # Convert to string if it's not already
#             indicador = str(indicador).strip()
            
#             # Get level columns (Nivel 0, Nivel 1, etc.)
#             level_cols = [col for col in df_rubric_prodoc.columns if col.startswith('Nivel')]
            
#             # Extract values from level columns
#             valores = []
#             for col in level_cols:
#                 val = row[col]
#                 if not pd.isna(val) and str(val).strip() != '':
#                     valores.append(str(val).strip())
            
#             # Store in our rubric dictionary
#             prodoc_rubric[indicador] = valores
        
#         # Success message with count of loaded criteria
#         st.success(f"Rúbrica cargada correctamente desde PRODOC_rubric.xlsx: {len(prodoc_rubric)} criterios cargados.")
#     except FileNotFoundError:
#         st.error("No se encontró el archivo PRODOC_rubric.xlsx. Por favor, asegúrese de que existe en el directorio de la aplicación.")
#     except Exception as e:
#         st.error(f"Error al cargar la rúbrica desde PRODOC_rubric.xlsx: {str(e)}")
    
#     # Display the loaded rubric
#     with st.expander("Ver rúbrica cargada"):
#         st.subheader("Criterios de Evaluación PRODOC")
#         for criterion, values in prodoc_rubric.items():
#             st.markdown(f"**{criterion}**: {values}")
    
#     # Instrucciones generales
#     st.info("""
#     **Instrucciones:**
#     1. Suba los archivos DOCX en las secciones correspondientes.
#     2. Presione el botón de evaluación en cada sección para analizar el documento.
#     3. Revise los resultados de cada rúbrica en las tablas interactivas.
#     4. Visualice las puntuaciones promedio por dimensión en los gráficos de barras.
#     5. Descargue todos los resultados y evidencias en archivos ZIP.
#     """)

#     # Sección 1: Documento PRODOC
#     st.markdown("### 📄 Sección 1: Documento PRODOC")
#     st.info("Suba el documento PRODOC para evaluación de sostenibilidad del proyecto.")
#     uploaded_file_prodoc = st.file_uploader("Suba un archivo DOCX del PRODOC:", type=["docx"], key="prodoc_file_uploader")

#     if st.button('🔍 Procesar y Evaluar PRODOC', key="prodoc_process_button"):
#         # Only process if file is uploaded
#         if uploaded_file_prodoc is not None:
#             uploaded_file = uploaded_file_prodoc
#             document_type = "PRODOC"
#             st.markdown(f"#### Procesando documento {document_type}...")

#             file_hash = hash(uploaded_file.getvalue())
#             if st.session_state.get('prodoc_last_file_hash') != file_hash:
#                 with st.spinner("Procesando documento..."):
#                     try:
#                         tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
#                         tmp_file.write(uploaded_file.read())
#                         tmp_file.close()
#                         progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
#                         doc_result = docx2python(tmp_file.name)
#                         df = extract_docx_structure(tmp_file.name)
#                         progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")
#                         header_1_values = df['header_1'].dropna().unique()
#                         llm_summary_rows = []
#                         llm_progress = st.progress(0, text="Procesando secciones con LLM...")
#                         total_sections = len(header_1_values)
#                         for idx, header in enumerate(header_1_values):
#                             section_df = df[df['header_1'] == header].copy()
#                             full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
#                             if not full_text:
#                                 llm_output = ""
#                             else:
#                                 llm_progress.progress((idx+1)/total_sections, text=f"Procesando sección: {header}")
#                                 try:
#                                     response = openai.ChatCompletion.create(
#                                         model="gpt.5.1-mini",
#                                         messages=[
#                                             {"role": "system", "content": "You are a helpful assistant that rewrites extracted document content into well-structured, formal paragraphs. Do not rewrite the original content, just reconstruct it in proper, coherent paragraphs, without rephrasing or paraphrasing or rewording."},
#                                             {"role": "user", "content": full_text}
#                                         ],
#                                         max_completion_tokens=4096,
#                                         temperature=0.01,
#                                     )
#                                     llm_output = response["choices"][0]["message"]["content"].strip()
#                                 except Exception as e:
#                                     llm_output = f"[LLM ERROR: {e}]"
#                             llm_summary_rows.append({'header_1': header, 'llm_paragraph': llm_output})
#                         llm_progress.progress(1.0, text="LLM parsing completado.")
#                         llm_summary_df = pd.DataFrame(llm_summary_rows)
#                         llm_summary_df['n_words'] = llm_summary_df['llm_paragraph'].str.split().str.len()
#                         exploded_df = llm_summary_df.assign(
#                             llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
#                         ).explode('llm_paragraph')
#                         exploded_df = exploded_df.reset_index(drop=True)
#                         exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
#                         full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
#                         file_size = os.path.getsize(tmp_file.name)
#                         n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
#                         n_paragraphs = len(exploded_df)
#                         st.session_state['prodoc_full_document_text'] = full_document_text
#                         st.session_state['prodoc_document_stats'] = {
#                             'file_size': file_size,
#                             'n_words': n_words,
#                             'n_paragraphs': n_paragraphs
#                         }
#                         st.session_state['prodoc_exploded_df'] = exploded_df
#                         st.session_state['prodoc_last_file_hash'] = file_hash
#                         try:
#                             os.unlink(tmp_file.name)
#                         except:
#                             pass
#                         progress_bar.progress(0.8, text="Documento procesado. Listo para evaluación.")
#                         st.info(f"**Resumen del documento:**\n\n" + 
#                                 f"- Tamaño del archivo: {file_size/1024:.2f} KB\n" + 
#                                 f"- Número de palabras: {n_words}\n" + 
#                                 f"- Número de párrafos: {n_paragraphs}")
#                         st.markdown("#### Estructura extraída del documento:")
#                         st.dataframe(exploded_df, use_container_width=True)
#                         progress_bar.progress(1.0, text="Procesamiento completo.")
#                     except Exception as e:
#                         st.error(f"Error procesando el documento: {e}")
#                         import traceback
#                         st.error(traceback.format_exc())
#                         st.stop()
            
#             # Now, always run rubric evaluation if document is processed
#             document_text = st.session_state.get('prodoc_full_document_text', '')
#             if not document_text:
#                 st.error("No se pudo recuperar el texto del documento. Por favor, vuelva a cargar el archivo.")
#                 st.stop()
                
#             # Define the rubric to evaluate, using the same structure as tab4
#             rubrics = [
#                 ("Evaluación PRODOC", prodoc_rubric)
#             ]
            
#             rubric_results = []
#             from concurrent.futures import ThreadPoolExecutor, as_completed
#             MAX_WORKERS = 48
            
#             def eval_one_criterion(args):
#                 crit, descriptions, dimension, rubric_name = args
#                 try:
#                     result = evaluate_criterion_with_llm(document_text, crit, descriptions)
#                     return {
#                         'Criterio': crit,
#                         'Dimensión': dimension,
#                         'Score': result.get('score', 0),
#                         'Análisis': result.get('analysis', ''),
#                         'Evidencia': result.get('evidence', ''),
#                         'Error': result.get('error', '') if 'error' in result else '',
#                         'Rúbrica': rubric_name
#                     }
#                 except Exception as e:
#                     return {
#                         'Criterio': crit,
#                         'Dimensión': dimension,
#                         'Score': 0,
#                         'Análisis': '',
#                         'Evidencia': '',
#                         'Error': str(e),
#                         'Rúbrica': rubric_name
#                     }
            
#             for rubric_name, rubric_dict in rubrics:
#                 rubric_analysis_data = []
#                 n_criteria = len(rubric_dict)
#                 progress = st.progress(0, text=f"Iniciando evaluación por rúbrica: {rubric_name}...")
#                 with st.spinner(f'Evaluando documento por rúbrica: {rubric_name}...'):
#                     with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
#                         futures = {
#                             executor.submit(eval_one_criterion, (
#                                 crit,
#                                 rubric_data['valores'] if isinstance(rubric_data, dict) else rubric_data,
#                                 rubric_data.get('dimension', 'No especificada') if isinstance(rubric_data, dict) else 'No especificada',
#                                 rubric_name
#                             )): (crit, idx)
#                             for idx, (crit, rubric_data) in enumerate(rubric_dict.items())
#                         }
#                         completed = 0
#                         for future in as_completed(futures):
#                             result = future.result()
#                             rubric_analysis_data.append(result)
#                             completed += 1
#                             crit, idx = futures[future]
#                             progress.progress(completed / n_criteria, text=f"Evaluando criterio: {crit}")
#                 rubric_results.append((rubric_name, pd.DataFrame(rubric_analysis_data)))
            
#             # Show and allow download of results only after evaluation
#             if rubric_results:
#                 for rubric_name, rubric_analysis_df in rubric_results:
#                     st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
#                     if not rubric_analysis_df.empty:
#                         # Ensure 'Evidencia' column is present
#                         if 'Evidencia' not in rubric_analysis_df.columns:
#                             rubric_analysis_df['Evidencia'] = ''

#                         # Reorder columns to show in logical order: Criterio, Dimensión, Score, Análisis, Evidencia, Error, Rúbrica
#                         cols = rubric_analysis_df.columns.tolist()
#                         desired_order = ['Criterio', 'Dimensión', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
#                         new_order = [col for col in desired_order if col in cols]
#                         remaining_cols = [col for col in cols if col not in desired_order]
#                         final_order = new_order + remaining_cols
#                         rubric_analysis_df = rubric_analysis_df[final_order]
#                         # Normalize 'Evidencia' column to always be a string
#                         if 'Evidencia' in rubric_analysis_df.columns:
#                             rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
#                                 lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
#                             )
#                         st.dataframe(rubric_analysis_df, use_container_width=True)
#                     else:
#                         st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
                
#                 # Create visualizations for the results
#                 st.markdown("### Visualización de Resultados")
                
#                 # Prepare data for visualization
#                 all_scores = []
#                 for rubric_name, df in rubric_results:
#                     for _, row in df.iterrows():
#                         all_scores.append({
#                             'Criterio': row['Criterio'],
#                             'Puntuación': row['Score']
#                         })
                
#                 scores_df = pd.DataFrame(all_scores)
                
#                 # Calculate the overall average score
#                 overall_avg = scores_df['Puntuación'].mean()
                
#                 # Sort the dataframe by score in descending order
#                 scores_df = scores_df.sort_values(by='Puntuación', ascending=False)
                
#                 # Create a horizontal bar chart instead of vertical for better label readability
#                 fig = go.Figure()
                
#                 # Create short identifiers for criteria (e.g., "Criterio 1", "Criterio 2", etc.)
#                 scores_df['Criterio_ID'] = [f"Criterio {i+1}" for i in range(len(scores_df))]
                
#                 # Create custom hover text with full criteria description
#                 scores_df['Hover_Text'] = scores_df.apply(
#                     lambda row: f"<b>{row['Criterio_ID']}</b><br>{row['Criterio']}<br>Puntuación: {row['Puntuación']:.2f}", 
#                     axis=1
#                 )
                
#                 # Add the bars - horizontal orientation with hover text
#                 fig.add_trace(go.Bar(
#                     y=scores_df['Criterio_ID'],  # Using short identifiers
#                     x=scores_df['Puntuación'],
#                     text=scores_df['Puntuación'].round(2),
#                     textposition='auto',
#                     marker_color='#3498db',
#                     orientation='h',  # Horizontal bars
#                     name='Puntuación',
#                     hovertext=scores_df['Hover_Text'],
#                     hoverinfo='text'
#                 ))
                
#                 # Add the average line - vertical for horizontal chart
#                 fig.add_trace(go.Scatter(
#                     y=scores_df['Criterio_ID'],
#                     x=[overall_avg] * len(scores_df),
#                     mode='lines',
#                     line=dict(color='red', width=2, dash='dash'),
#                     name=f'Promedio General: {overall_avg:.2f}'
#                 ))
                
#                 # Update layout for better readability
#                 fig.update_layout(
#                     title='Puntuación por Criterio (Ordenado de Mayor a Menor)',
#                     xaxis_title='Puntuación',
#                     yaxis_title='',  # No need for y-axis title
#                     xaxis=dict(range=[0, 5.5]),  # Now x-axis has the scores
#                     height=max(400, len(scores_df) * 35),  # Reduced height since we're using shorter labels
#                     width=800,
#                     legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
#                     margin=dict(l=20, r=20, t=80, b=50),
#                     hoverlabel=dict(
#                         bgcolor="white",
#                         font_size=12,
#                         font_family="Arial"
#                     )
#                 )
                
#                 # Configure y-axis for cleaner look
#                 fig.update_yaxes(
#                     automargin=True  # Automatically adjust margins to fit labels
#                 )
                
#                 # Add grid lines for better score reference
#                 fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor='LightGrey')
                
#                 # Display the chart
#                 st.plotly_chart(fig, use_container_width=True)
                
#                 # Provide a zip download for all results
#                 import io, zipfile
#                 zip_buffer = io.BytesIO()
#                 with zipfile.ZipFile(zip_buffer, "w") as zipf:
#                     for rubric_name, rubric_analysis_df in rubric_results:
#                         # Normalize 'Evidencia' column to always be a string before exporting
#                         if 'Evidencia' in rubric_analysis_df.columns:
#                             rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
#                                 lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
#                             )
#                         csv = rubric_analysis_df.to_csv(index=False)
#                         arcname = f"evaluacion_prodoc_{rubric_name.replace(' ', '_').lower()}.csv"
#                         zipf.writestr(arcname, csv)
#                 zip_buffer.seek(0)
#                 st.download_button(
#                     label="Descargar resultados como ZIP",
#                     data=zip_buffer,
#                     file_name="resultados_evaluacion_prodoc.zip",
#                     mime="application/zip",
#                     key="prodoc_download_button"
#                 )
#             else:
#                 st.warning("No se generaron resultados para ninguna rúbrica.")
#         else:
#             st.info("Por favor suba un archivo DOCX para comenzar y pulse el botón para procesar y evaluar.")
#     else:
#         st.info("Por favor suba un archivo DOCX del PRODOC para comenzar.")

#     # Sección 2: Documento de Evaluación
#     st.markdown("### 📋 Sección 2: Documento de Evaluación")
#     st.info("Suba el documento de evaluación para análisis comparativo con la misma rúbrica de sostenibilidad.")
#     uploaded_file_evaluation = st.file_uploader("Suba un archivo DOCX de evaluación:", type=["docx"], key="evaluation_file_uploader")

#     if st.button('🔍 Procesar y Evaluar Documento de Evaluación', key="evaluation_process_button"):
#         # Only process if file is uploaded
#         if uploaded_file_evaluation is not None:
#             uploaded_file = uploaded_file_evaluation
#             document_type = "Evaluación"
#             st.markdown(f"#### Procesando documento {document_type}...")

#             file_hash = hash(uploaded_file.getvalue())
#             if st.session_state.get('evaluation_last_file_hash') != file_hash:
#                 with st.spinner("Procesando documento de evaluación..."):
#                     try:
#                         tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
#                         tmp_file.write(uploaded_file.read())
#                         tmp_file.close()
#                         progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
#                         doc_result = docx2python(tmp_file.name)
#                         df = extract_docx_structure(tmp_file.name)

#                         # Use the same rubric processing logic as PRODOC
#                         all_sections = extract_all_sections(doc_result)

#                         # Process with the same rubric (prodoc_rubric)
#                         results_evaluation = process_with_rubrics(all_sections, prodoc_rubric, document_type, progress_bar)

#                         # Store in session state with different key
#                         st.session_state['evaluation_results'] = results_evaluation
#                         st.session_state['evaluation_sections'] = all_sections
#                         st.session_state['evaluation_last_file_hash'] = file_hash
#                         st.session_state['evaluation_document_processed'] = True

#                         os.unlink(tmp_file.name)
#                         progress_bar.progress(1.0, text="¡Análisis completado!")

#                         st.success("✅ Documento de evaluación procesado exitosamente")

#                     except Exception as e:
#                         st.error(f"Error al procesar el documento de evaluación: {str(e)}")
#                         st.session_state['evaluation_document_processed'] = False
#                         import traceback
#                         st.error(f"Detalles del error: {traceback.format_exc()}")
#             else:
#                 st.info("✅ Este archivo ya ha sido procesado. Mostrando resultados existentes.")
#                 st.session_state['evaluation_document_processed'] = True
#         else:
#             st.info("Por favor suba un archivo DOCX de evaluación para comenzar.")
#     else:
#         st.info("Por favor suba un archivo DOCX de evaluación para comenzar.")

#     # Display evaluation results if processed
#     if st.session_state.get('evaluation_document_processed', False) and 'evaluation_results' in st.session_state:
#         st.markdown("#### 📊 Resultados de Evaluación - Documento de Evaluación")
#         results_evaluation = st.session_state['evaluation_results']

#         # Same display logic as PRODOC
#         if results_evaluation:
#             display_results(results_evaluation, document_type="Evaluación", tab_key="evaluation")
#         else:
#             st.warning("No se generaron resultados para ninguna rúbrica del documento de evaluación.")

# ================== TAB 6: EVALUACIÓN DE PRODOCS =====================
with tab3:
    st.header("Diagnóstico de Sostenibilidad del Proyecto")
    
    # Read rubric from Excel file
    import pandas as pd
    prodoc_rubric = {}
    
    try:
        # Load rubric from PRODOC_rubric.xlsx
        # df_rubric_prodoc = pd.read_excel('./PRODOC_rubric.xlsx', sheet_name='rubric')
        df_rubric_prodoc = pd.read_excel('./Evaluación de sostenibilidad del proyecto_rubric_7nov.xlsx', sheet_name='rubric')

        # Verify required columns exist
        required_cols = ['Dimensión', 'Criterio', 'Indicador']
        missing_cols = [col for col in required_cols if col not in df_rubric_prodoc.columns]
        if missing_cols:
            st.error(f"Columnas faltantes en el archivo Excel: {missing_cols}")
            st.info(f"Columnas disponibles: {list(df_rubric_prodoc.columns)}")
        
        # Helper function to extract numeric prefix for sorting (e.g., "1.1" -> (1, 1), "2.3" -> (2, 3))
        def extract_sort_key(text):
            import re
            if pd.isna(text):
                return (999, 999, 999)
            text = str(text).strip()
            # Match patterns like "1.", "1.1", "1.1.", "(3.1)", etc.
            match = re.match(r'[\(]?(\d+)(?:\.(\d+))?(?:\.(\d+))?', text)
            if match:
                parts = [int(p) if p else 0 for p in match.groups()]
                return tuple(parts)
            return (999, 999, 999)
        
        # Process each row to extract criteria and values
        for idx, row in df_rubric_prodoc.iterrows():
            indicador = row.get('Indicador', '')
            criterio = row.get('Criterio', '')
            dimension = row.get('Dimensión', 'No especificada')
            
            if pd.isna(indicador) or str(indicador).strip() == '':
                continue
            
            indicador = str(indicador).strip()
            criterio = str(criterio).strip() if not pd.isna(criterio) else ''
            dimension = str(dimension).strip() if not pd.isna(dimension) else 'No especificada'
            
            # Get level columns (Nivel 1, Nivel 2, etc.)
            level_cols = sorted([col for col in df_rubric_prodoc.columns if col.startswith('Nivel')])
            
            valores = []
            for col in level_cols:
                val = row[col]
                if not pd.isna(val) and str(val).strip() != '':
                    valores.append(str(val).strip())
            
            # Store with unique key (dimension + criterio + indicador) to avoid overwrites
            # since same criterio and indicador text can appear in different dimensions
            unique_key = f"{dimension}|{criterio}|{indicador}"
            prodoc_rubric[unique_key] = {
                'valores': valores, 
                'dimension': dimension,
                'criterio': criterio,
                'indicador': indicador,  # Keep original indicador text for display
                'sort_key': extract_sort_key(indicador)
            }
        
        # Show dimension breakdown
        dim_counts = {}
        for ind, data in prodoc_rubric.items():
            dim = data.get('dimension', 'N/A')
            dim_counts[dim] = dim_counts.get(dim, 0) + 1
        st.success(f"Rúbrica cargada: {len(prodoc_rubric)} indicadores en {len(dim_counts)} dimensiones.")
        
        # Download button for the rubric file (directly on page, no expander)
        try:
            with open('./Evaluación de sostenibilidad del proyecto_rubric_7nov.xlsx', 'rb') as f:
                st.download_button(
                    label="📥 Descargar archivo rúbrica de Sostenibilidad del proyecto",
                    data=f,
                    file_name="Evaluacion_sostenibilidad_rubric.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_sostenibilidad_rubric"
                )
        except FileNotFoundError:
            st.warning("Archivo de rúbrica no disponible para descarga.")
    except FileNotFoundError:
        st.error("No se encontró el archivo PRODOC_rubric.xlsx. Por favor, asegúrese de que existe en el directorio de la aplicación.")
    except Exception as e:
        st.error(f"Error al cargar la rúbrica desde PRODOC_rubric.xlsx: {str(e)}")
    
    # Instrucciones generales
    st.info("""
    **Instrucciones:**
    1.	Seleccione la rúbrica y los criterios específicos que desea analizar.
    2.	Seleccione las secciones adecuadas de diagnóstico según el documento cargado (algunos criterios son relevantes a Documentos de Proyecto y otras a informes de progreso).
    3.	Suba el archivo en formato DOCX correspondiente.
    4.	Presione el botón de Procesar y Analizar para analizar el documento.
    5.	Revise los resultados de cada rúbrica en las tablas interactivas.
    6.	Visualice las puntuaciones promedio por dimensión en los gráficos de barras.
    7.	Descargue todos los resultados y evidencias en archivos ZIP.

    """)

    # Initialize session state for selections and results persistence
    if 'selected_criteria_tab3' not in st.session_state:
        st.session_state['selected_criteria_tab3'] = []
    if 'selected_dimensions_tab3' not in st.session_state:
        st.session_state['selected_dimensions_tab3'] = []
    if 'tab3_results' not in st.session_state:
        st.session_state['tab3_results'] = None
    if 'document_extracted_tab3' not in st.session_state:
        st.session_state['document_extracted_tab3'] = False

    # Document upload
    st.markdown("### 📄 Carga de Documento")
    
    # Warning box about document requirements
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos PRODOC:**
    
    **📝 Formato del documento:**
    - Solo se aceptan archivos en formato **.docx** (Word 2007 o posterior)
    - El documento PRODOC debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del PRODOC deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente y las secciones no se identificarán.
    - Los PRODOCs deben tener una estructura clara con secciones bien definidas (Marco Lógico, Presupuesto, Cronograma, etc.)
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos que excedan este límite serán truncados automáticamente
    - Se recomienda dividir documentos muy extensos (más de ~180 páginas) en secciones más pequeñas si es necesario
    
    **✅ Mejores prácticas:**
    - Usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones principales del PRODOC
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que todas las secciones importantes (Marco Lógico, Presupuesto, etc.) tengan encabezados claros
    - Verifica que el documento esté guardado correctamente antes de subirlo
    """)
    
    uploaded_file_prodoc = st.file_uploader("Suba un archivo DOCX para evaluación:", type=["docx"], key="prodoc_file_uploader_tab3")
    
    # Document Extraction Section
    st.markdown("---")
    st.markdown("### 📥 Extracción de Documento")
    
    if uploaded_file_prodoc is not None:
        file_hash = hash(uploaded_file_prodoc.getvalue())
        file_changed = st.session_state.get('last_file_hash_tab3') != file_hash
        
        if file_changed:
            st.session_state['document_extracted_tab3'] = False
            st.session_state['last_file_hash_tab3'] = None
        
        if st.button("🔍 Extraer Documento", key="extract_document_tab3", type="primary"):
            if uploaded_file_prodoc is None:
                st.error("Por favor suba un archivo DOCX primero.")
                st.stop()
            
            with st.spinner("Extrayendo documento..."):
                try:
                    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp_file.write(uploaded_file_prodoc.read())
                    tmp_file.close()
                    
                    progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                    doc_result = docx2python(tmp_file.name)
                    
                    # Use enhanced extraction
                    df, tables_data, extraction_stats = extract_docx_structure_enhanced(tmp_file.name)
                    progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")
                    
                    # Extract sections
                    header_1_values = df['header_1'].dropna().unique()
                    llm_summary_rows = []
                    
                    for idx, header in enumerate(header_1_values):
                        section_df = df[df['header_1'] == header].copy()
                        full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                        section_words = len(full_text.split())
                        section_paras = len(section_df[section_df['source_type'] == 'paragraph'])
                        section_tables = len(section_df[section_df['source_type'] == 'table'])
                        
                        llm_summary_rows.append({
                            'header_1': header,
                            'llm_paragraph': full_text if full_text else "",
                            'n_words': section_words,
                            'n_paragraphs': section_paras,
                            'n_tables': section_tables
                        })

                    progress_bar.progress(0.5, text="Secciones extraídas.")
                    
                    # Create exploded dataframe
                    llm_summary_df = pd.DataFrame(llm_summary_rows)
                    exploded_df = llm_summary_df.assign(
                        llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                    ).explode('llm_paragraph')
                    exploded_df = exploded_df.reset_index(drop=True)
                    exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                    
                    # Get full text
                    full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                    
                    # Store in session state
                    file_size = os.path.getsize(tmp_file.name)
                    n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                    n_paragraphs = len(exploded_df)
                    
                    st.session_state['full_document_text_tab3'] = full_document_text
                    st.session_state['prodoc_document_stats_tab3'] = {
                        'file_size': file_size,
                        'n_words': n_words,
                        'n_paragraphs': n_paragraphs
                    }
                    st.session_state['exploded_df_tab3'] = exploded_df
                    st.session_state['extraction_df_tab3'] = df
                    st.session_state['tables_data_tab3'] = tables_data
                    st.session_state['extraction_stats_tab3'] = extraction_stats
                    st.session_state['sections_df_tab3'] = llm_summary_df
                    st.session_state['selected_sections_tab3'] = list(header_1_values)  # Select all by default
                    st.session_state['last_file_hash_tab3'] = file_hash
                    st.session_state['document_extracted_tab3'] = True
                    
                    try:
                        os.unlink(tmp_file.name)
                    except:
                        pass
                    
                    progress_bar.progress(1.0, text="Extracción completa.")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error procesando el documento: {e}")
                    import traceback
                    st.error(traceback.format_exc())
                    st.stop()
        
        # Show extraction results if document is extracted
        if st.session_state.get('document_extracted_tab3', False) and not file_changed:
            st.success("✅ Documento extraído con éxito")
            
            sections_df = st.session_state.get('sections_df_tab3', pd.DataFrame())
            
            if not sections_df.empty:
                header_1_values = sections_df['header_1'].tolist()
                
                # Section selector
                st.markdown("### 🔍 Selección de Secciones para Evaluación")
                st.info("Selecciona las secciones que deseas incluir en la evaluación. Por defecto, todas las secciones están seleccionadas.")
                
                # Guidance about section extraction
                if len(header_1_values) == 0:
                    st.error("⚠️ **No se detectaron secciones en el documento.** Esto puede deberse a que el documento no usa estilos de encabezado de Word (Heading 1, Heading 2, etc.). Por favor, verifica que tu documento PRODOC tenga encabezados formateados correctamente.")
                elif len(header_1_values) < 3:
                    st.warning("⚠️ **Se detectaron pocas secciones** en el documento PRODOC. Si esperabas más secciones, verifica que el documento use estilos de encabezado de Word (Heading 1, Heading 2, etc.) para identificar las secciones principales como Marco Lógico, Presupuesto, Cronograma, etc.")
                
                # Initialize selected sections if not exists
                if 'selected_sections_tab3' not in st.session_state:
                    st.session_state['selected_sections_tab3'] = list(header_1_values)
                
                # Section selection interface
                selected_sections = st.session_state.get('selected_sections_tab3', list(header_1_values)).copy()
                col1, col2 = st.columns([3, 1])
                
                with col2:
                    if st.button("✅ Seleccionar Todas", key="select_all_sections_tab3"):
                        st.session_state['selected_sections_tab3'] = list(header_1_values)
                        st.rerun()
                    
                    if st.button("❌ Deseleccionar Todas", key="deselect_all_sections_tab3"):
                        st.session_state['selected_sections_tab3'] = []
                        st.rerun()
                
                with col1:
                    st.markdown("**Secciones disponibles:**")
                    for section in header_1_values:
                        section_info = sections_df[sections_df['header_1'] == section].iloc[0]
                        is_selected = section in selected_sections
                        
                        checkbox_label = f"**{section}** ({section_info['n_words']:,} palabras, {section_info['n_paragraphs']} párrafos)"
                        
                        checkbox_key = f"section_checkbox_{section}_tab3"
                        new_selection = st.checkbox(checkbox_label, value=is_selected, key=checkbox_key)
                        
                        if new_selection and section not in selected_sections:
                            selected_sections.append(section)
                        elif not new_selection and section in selected_sections:
                            selected_sections.remove(section)
                        
                        # Add expandable preview of extracted content
                        with st.expander(f"👁️ Ver contenido: {section}", expanded=False):
                            section_content = section_info['llm_paragraph']
                            if section_content and section_content.strip():
                                st.text_area(
                                    "Contenido extraído:",
                                    value=section_content,
                                    height=200,
                                    key=f"content_preview_{section}_tab3",
                                    label_visibility="collapsed"
                                )
                                st.caption(f"Total: {len(section_content):,} caracteres")
                            else:
                                st.info("Esta sección no tiene contenido extraído.")
                            
                            # Show table-extracted text
                            tables_data = st.session_state.get('tables_data_tab3', [])
                            section_tables = [t for t in tables_data if t.get('section') == section]
                            
                            if section_tables:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                for table_info in section_tables:
                                    table_num = table_info.get('table_number', 'N/A')
                                    table_data = table_info.get('data', [])
                                    if table_data:
                                        # Format table as text
                                        table_text = '\n'.join([' | '.join(str(cell) for cell in row) for row in table_data])
                                        st.text_area(
                                            f"Tabla {table_num}:",
                                            value=table_text,
                                            height=150,
                                            key=f"table_preview_{section}_table{table_num}_tab3",
                                            label_visibility="collapsed"
                                        )
                                        st.caption(f"Tabla {table_num}: {len(table_data)} filas, {len(table_data[0]) if table_data else 0} columnas")
                            else:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                st.info("No se encontraron tablas en esta sección.")
                
                # Update session state
                st.session_state['selected_sections_tab3'] = selected_sections
                
                # Show selection summary
                if selected_sections:
                    selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
                    total_selected_words = selected_df['n_words'].sum()
                    total_selected_paras = selected_df['n_paragraphs'].sum()
                    
                    # Estimate tokens
                    if encoding:
                        selected_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
                        estimated_tokens = len(encoding.encode(selected_text))
                    else:
                        estimated_tokens = total_selected_words * 1.2
                    
                    # Warn if approaching limit
                    if estimated_tokens > 100000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.warning(f"⚠️ **Advertencia de límite de contexto:** Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). El sistema procesa hasta 110,000 tokens (aproximadamente 150-200 páginas). Si el documento excede este límite, será truncado automáticamente.")
                    elif estimated_tokens > 80000:
                        estimated_pages = (estimated_tokens / 110000) * 180  # Approximate pages based on 180 pages = 110K tokens
                        st.info(f"ℹ️ Las secciones seleccionadas contienen aproximadamente {estimated_tokens:,.0f} tokens estimados (~{estimated_pages:.0f} páginas aproximadas). Estás dentro del límite de 110,000 tokens (aproximadamente 150-200 páginas).")
                    
                    st.success(f"✅ {len(selected_sections)} secciones seleccionadas | "
                              f"{total_selected_words:,} palabras | "
                              f"~{estimated_tokens:,} tokens estimados")

    # Rubric and Criteria Selection Section (moved after document extraction)
    st.markdown("---")
    st.markdown("### 📋 Selección de Criterios")
    
    # Group indicadores by dimension and criterio, maintaining order
    # Structure: {dimension: {criterio: [(unique_key, indicador_text, sort_key)]}}
    criteria_by_dimension = {}
    for unique_key, data in prodoc_rubric.items():
        dimension = data.get('dimension', 'No especificada')
        criterio = data.get('criterio', 'Sin criterio')
        indicador_text = data.get('indicador', unique_key)  # Use indicador text for display
        sort_key = data.get('sort_key', (999, 999, 999))
        
        if dimension not in criteria_by_dimension:
            criteria_by_dimension[dimension] = {}
        if criterio not in criteria_by_dimension[dimension]:
            criteria_by_dimension[dimension][criterio] = []
        criteria_by_dimension[dimension][criterio].append((unique_key, indicador_text, sort_key))
    
    # Sort indicadores within each criterio by their numeric prefix (sort_key is index 2)
    for dimension in criteria_by_dimension:
        for criterio in criteria_by_dimension[dimension]:
            criteria_by_dimension[dimension][criterio].sort(key=lambda x: x[2])
    
    # Define dimension order
    dimension_order = ['Diseño', 'Implementación', 'Pre-Cierre']
    
    # Helper to extract criterio number for sorting (e.g., "1. Participación..." -> 1)
    def get_criterio_order(criterio_name):
        import re
        match = re.match(r'(\d+)\.', criterio_name)
        return int(match.group(1)) if match else 999
    
    # Only show rubric selection if document is extracted
    if st.session_state.get('document_extracted_tab3', False):
        # Display the loaded rubric with selection grouped by dimension
        with st.expander("Ver y seleccionar criterios de evaluación", expanded=True):
            st.subheader("Criterios de Evaluación PRODOC")
        st.markdown(
            """
            <div class='reference-box'>
            Analiza si el Documento de Proyecto (PRODOC) incorpora, desde el inicio, los factores que favorecen la continuidad de resultados: participación de mandantes/socios y gestión de riesgos, bases de sostenibilidad institucional y política, consideraciones de género y, cuando aplique, transición justa. Usa como evidencia el Documento de Proyecto y anexos; los puntajes sirven para ajustar estrategias y definir tempranamente el plan de sostenibilidad.
            </div>
            """,
            unsafe_allow_html=True
        )
        
        # Select all checkbox
        select_all_tab3 = st.checkbox("Seleccionar todos los criterios", key='select_all_tab3')
        
        selected_criteria = []
        
        # Display criteria grouped by dimension
        dimension_descriptions = {
            "Diseño": (
                "Analiza si el Documento de Proyecto (PRODOC) incorpora, desde el inicio, los factores que "
                "favorecen la continuidad de resultados: participación de mandantes/socios y gestión de riesgos, "
                "bases de sostenibilidad institucional y política, consideraciones de género y, cuando aplique, "
                "transición justa. Usa como evidencia el Documento de Proyecto y anexos; los puntajes sirven para "
                "ajustar estrategias y definir tempranamente el plan de sostenibilidad."
            ),
            "Implementación": (
                "Contrasta avances reportados con los criterios de la Matriz de Criterios (capacidades desarrolladas, "
                "alianzas, recursos movilizados, integración en políticas/planes, gestión del conocimiento, etc.) "
                "para identificar riesgos, cuellos de botella y acciones de mitigación. Usa informes de progreso, actas, "
                "convenios y otros documentos de ejecución; los resultados orientan ajustes y fortalecen la trazabilidad "
                "de decisiones."
            ),
            "Pre-Cierre": (
                "Aplica una revisión ex post (idealmente en el último trimestre de un proyecto) para verificar qué "
                "elementos efectivamente aseguran la continuidad de resultados y documentar lecciones. Si se aplica "
                "antes del cierre, algunos puntajes serán referenciales (“lo esperado”); en todos los casos, el diagnóstico "
                "alimenta el plan/estrategia de sostenibilidad y su seguimiento. Usa un informe de evaluación para "
                "realizar este análisis."
            )
        }

        # Note: .reference-box styles are defined in the global ILO styles section

        # Sort dimensions according to defined order
        sorted_dimensions = sorted(
            criteria_by_dimension.keys(), 
            key=lambda d: dimension_order.index(d) if d in dimension_order else 999
        )
        
        for dimension in sorted_dimensions:
            st.markdown(f"#### 📊 {dimension}")
            if dimension in dimension_descriptions:
                st.markdown(
                    f"<div class='reference-box'>{dimension_descriptions[dimension]}</div>",
                    unsafe_allow_html=True
                )
            
            # Checkbox to select entire dimension
            dimension_key = f"dimension_tab3_{dimension}"
            select_dimension = st.checkbox(
                f"✅ Seleccionar toda la dimensión '{dimension}'",
                value=select_all_tab3 or dimension in st.session_state.get('selected_dimensions_tab3', []),
                key=dimension_key
            )
            
            # Get criterios for this dimension, sorted by their number
            criterios_in_dimension = criteria_by_dimension[dimension]
            sorted_criterios = sorted(criterios_in_dimension.keys(), key=get_criterio_order)
            
            # Show criterios and their indicadores within this dimension
            with st.container():
                for criterio in sorted_criterios:
                    # Display criterio header
                    st.markdown(f"**{criterio}**")
                    
                    # Get sorted indicadores for this criterio
                    # Each item is (unique_key, indicador_text, sort_key)
                    indicadores_data = criterios_in_dimension[criterio]
                    
                    for unique_key, indicador_text, sort_key in indicadores_data:
                        # If dimension is selected, auto-select all its indicadores
                        default_value = select_all_tab3 or select_dimension or unique_key in st.session_state['selected_criteria_tab3']
                        
                        is_selected = st.checkbox(
                            f"  ↳ {indicador_text}",
                            value=default_value,
                            key=f"criterion_tab3_{unique_key}",
                            disabled=select_dimension  # Disable individual selection if dimension is selected
                        )
                        
                        if is_selected or select_dimension:
                            selected_criteria.append(unique_key)
            
            st.markdown("---")  # Separator between dimensions
        
        # Update session state
        st.session_state['selected_criteria_tab3'] = selected_criteria
        selected_dimensions = [dim for dim in criteria_by_dimension.keys() 
                              if st.session_state.get(f"dimension_tab3_{dim}", False)]
        st.session_state['selected_dimensions_tab3'] = selected_dimensions
        
        # Count total indicadores
        total_indicadores = sum(
            len(indicadores) 
            for criterios in criteria_by_dimension.values() 
            for indicadores in criterios.values()
        )
        
        st.info(f"📌 Indicadores seleccionados: {len(selected_criteria)}/{total_indicadores} | Dimensiones seleccionadas: {len(selected_dimensions)}/{len(criteria_by_dimension)}")
    else:
        st.info("ℹ️ Por favor extrae el documento primero para poder seleccionar los criterios.")

    # Process and Evaluate button
    st.markdown("---")
    st.markdown("### ⚙️ Procesamiento y Evaluación")
    if st.button('🚀 Procesar y Evaluar', key="prodoc_process_button_tab3", type="primary"):
        # Check prerequisites
        if not st.session_state.get('document_extracted_tab3', False):
            st.error("❌ Por favor extrae el documento primero usando el botón 'Extraer Documento'.")
            st.stop()
        
        if uploaded_file_prodoc is None:
            st.error("Por favor suba un archivo DOCX primero.")
            st.stop()
        
        # Get selected criteria from session state
        selected_criteria = st.session_state.get('selected_criteria_tab3', [])
        if not selected_criteria:
            st.error("Por favor seleccione al menos un criterio.")
            st.stop()
        
        # Get selected sections or use full document
        selected_sections = st.session_state.get('selected_sections_tab3', [])
        sections_df = st.session_state.get('sections_df_tab3', pd.DataFrame())
        
        if selected_sections and not sections_df.empty:
            # Filter to selected sections only
            selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
            document_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
            st.info(f"📌 Evaluando {len(selected_sections)} secciones seleccionadas")
        else:
            # Fallback to full document
            document_text = st.session_state.get('full_document_text_tab3', '')
            if not selected_sections:
                st.warning("⚠️ No hay secciones seleccionadas. Usando documento completo.")
        
        if not document_text:
            st.error("No se pudo recuperar el texto del documento.")
            st.stop()
        
        # Skip the old extraction logic - document is already extracted
        # Evaluate with selected criteria
        if False:  # Disable old extraction logic
            uploaded_file = uploaded_file_prodoc
            st.markdown("#### Procesando documento...")

            file_hash = hash(uploaded_file.getvalue())
            if st.session_state.get('prodoc_last_file_hash_tab3') != file_hash:
                with st.spinner("Procesando documento..."):
                    try:
                        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        tmp_file.write(uploaded_file.read())
                        tmp_file.close()
                        
                        progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                        doc_result = docx2python(tmp_file.name)
                        df = extract_docx_structure(tmp_file.name)
                        
                        progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")

                        # Extract sections directly (no LLM needed - text is already well-formatted)
                        header_1_values = df['header_1'].dropna().unique()
                        llm_summary_rows = []
                        progress_bar.progress(0.3, text="Extrayendo secciones del documento...")

                        for idx, header in enumerate(header_1_values):
                            section_df = df[df['header_1'] == header].copy()
                            # Extract text directly - already clean from extract_docx_structure
                            full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                            llm_summary_rows.append({'header_1': header, 'llm_paragraph': full_text if full_text else ""})

                        progress_bar.progress(0.5, text="Secciones extraídas.")
                        llm_summary_df = pd.DataFrame(llm_summary_rows)
                        llm_summary_df['n_words'] = llm_summary_df['llm_paragraph'].str.split().str.len()
                        exploded_df = llm_summary_df.assign(
                            llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                        ).explode('llm_paragraph')
                        exploded_df = exploded_df.reset_index(drop=True)
                        exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                        
                        full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                        file_size = os.path.getsize(tmp_file.name)
                        n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                        n_paragraphs = len(exploded_df)
                        
                        st.session_state['prodoc_full_document_text_tab3'] = full_document_text
                        st.session_state['prodoc_document_stats_tab3'] = {
                            'file_size': file_size,
                            'n_words': n_words,
                            'n_paragraphs': n_paragraphs
                        }
                        st.session_state['prodoc_exploded_df_tab3'] = exploded_df
                        st.session_state['prodoc_last_file_hash_tab3'] = file_hash
                        
                        try:
                            os.unlink(tmp_file.name)
                        except:
                            pass
                        
                        progress_bar.progress(0.8, text="Documento procesado. Listo para evaluación.")
                        st.info(f"**Resumen del documento:**\n\n" + 
                                f"- Tamaño del archivo: {file_size/1024:.2f} KB\n" + 
                                f"- Número de palabras: {n_words}\n" + 
                                f"- Número de párrafos: {n_paragraphs}")
                        st.markdown("#### Estructura extraída del documento:")
                        st.dataframe(exploded_df, use_container_width=True)
                        progress_bar.progress(1.0, text="Procesamiento completo.")
                        
                    except Exception as e:
                        st.error(f"Error procesando el documento: {e}")
                        import traceback
                        st.error(traceback.format_exc())
                        st.stop()
        
        # Define the evaluation function for tab3 with COUNTING support
        def evaluate_criterion_with_llm(document_text, criterion, descriptions, max_retries=3):
            """
            Analyze document against criterion with retry logic.
            TAB3-SPECIFIC: Enhanced to handle COUNT-BASED rubrics (stakeholders, participation forms, etc.)
            """
            import time

            # Detect if this is a stakeholder participation counting rubric
            is_stakeholder_counting = any(keyword in criterion.lower() for keyword in
                ['mandante', 'participan', 'stakeholder', 'actores', 'gobierno', 'empleador', 'trabajador'])

            for attempt in range(max_retries):
                try:
                    # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
                    combined_text = truncate_to_token_limit(document_text, max_tokens=110000, encoding_obj=encoding)

                    # Enhanced prompt for count-based rubrics
                    if is_stakeholder_counting:
                        system_content = """Eres un evaluador experto de documentos especialmente capacitado para CONTAR y aplicar LÓGICA DE UMBRALES.

**INSTRUCCIONES CRÍTICAS PARA EVALUACIÓN BASADA EN CONTEOS CON UMBRALES:**

Los niveles de la rúbrica especifican requisitos como "AL MENOS X mandantes participando en AL MENOS Y formas".
Debes aplicar lógica de umbrales estricta:

1. **IDENTIFICA y CUENTA** cada tipo de mandante/actor mencionado:
   - **Gobierno** (autoridades, funcionarios públicos, ministerios, etc.)
   - **Empleadores** (empresas, organizaciones patronales, cámaras de comercio, etc.)
   - **Trabajadores** (sindicatos, organizaciones de trabajadores, trabajadores individuales, etc.)

2. **CUENTA las formas de participación** para CADA mandante individualmente:
   - Diseño, reuniones, discusiones, comentarios, provisión de información, co-implementación, compromisos, etc.
   - Cada forma debe ser DISTINTA y VERIFICABLE en el documento

3. **FILTRO DE EVIDENCIA Y TIEMPO VERBAL (CRÍTICO - DEBE APLICARSE ESTRICTAMENTE):**
   
   **REGLA ABSOLUTA: SOLO PARTICIPACIÓN FACTUAL, PASADA O PRESENTE**
   
   - **SOLO cuenta participación que YA OCURRIÓ o ESTÁ OCURRIENDO** (hechos verificables, acciones concretas realizadas).
   - **NUNCA CUENTES** participaciones futuras, prometidas, planificadas, hipotéticas o esperadas.
   
   **VERBOS Y EXPRESIONES QUE DEBES RECHAZAR (NO CUENTAN):**
   - Futuro simple: "participará", "asistirá", "consultará", "se invitará", "se convocará", "se reunirá"
   - Futuro compuesto: "habrá participado", "habrá asistido"
   - Condicional: "participaría", "asistiría", "se consultaría"
   - Expresiones de planificación: "se espera que", "se prevé que", "está previsto", "se planifica", "se programará"
   - Expresiones de intención: "se pretende", "se busca", "se tiene como objetivo", "se propone"
   - Expresiones de compromiso futuro: "se compromete a", "acordó participar" (si no hay evidencia de participación real)
   - Marcadores temporales futuros: "en el futuro", "durante la implementación", "en las próximas fases", "posteriormente"
   - Participaciones hipotéticas: "podría participar", "sería consultado", "tendría la oportunidad"
   
   **VERBOS Y EXPRESIONES QUE SÍ CUENTAN (participación factual):**
   - Pasado simple: "participó", "asistió", "validó", "revisó", "comentó", "aprobó", "contribuyó", "colaboró"
   - Pasado compuesto: "ha participado", "ha asistido", "ha validado"
   - Presente: "participa", "asiste", "es miembro de", "forma parte de", "colabora en"
   - Gerundio de acciones completadas: "habiendo participado", "habiendo asistido"
   - Expresiones de hecho realizado: "fue consultado", "fue invitado y asistió", "se reunió con"
   
   **EJEMPLOS ESPECÍFICOS:**
   - ❌ "Se consultará a los mandantes durante la implementación" -> NO CUENTA (futuro)
   - ❌ "Se espera que los empleadores participen en el diseño" -> NO CUENTA (expectativa futura)
   - ❌ "Los trabajadores serán invitados a las reuniones" -> NO CUENTA (futuro)
   - ❌ "Está previsto que el gobierno valide la propuesta" -> NO CUENTA (planificación futura)
   - ❌ "Se tiene como objetivo involucrar a los actores" -> NO CUENTA (intención futura)
   - ✅ "El gobierno participó en la reunión de diseño del 15 de marzo" -> SÍ CUENTA (pasado factual)
   - ✅ "Los empleadores asistieron a las consultas y proporcionaron comentarios" -> SÍ CUENTA (pasado factual)
   - ✅ "Los trabajadores son miembros del comité de diseño" -> SÍ CUENTA (presente factual)
   - ✅ "Los mandantes fueron consultados y validaron el marco lógico" -> SÍ CUENTA (pasado factual)
   
   **VERIFICACIÓN OBLIGATORIA:**
   Antes de contar cualquier participación, pregunta: "¿Esta acción YA OCURRIÓ o está ocurriendo AHORA?" 
   Si la respuesta es NO o es incierta → NO CUENTES.
   Si la respuesta es SÍ y hay evidencia clara → SÍ CUENTA.

4. **APLICA LÓGICA DE UMBRALES** según los niveles de la rúbrica:
   - Si el nivel requiere "AL MENOS 2 mandantes en AL MENOS 2 formas":
     * CUENTA cuántos mandantes tienen 2 o más formas de participación
     * Si al menos 2 mandantes alcanzan ese umbral → cumple el nivel
     * Si solo 1 mandante alcanza ese umbral → NO cumple el nivel

   - Ejemplo:
     * Gobierno: 3 formas ✓ (cumple umbral de 2+)
     * Empleadores: 2 formas ✓ (cumple umbral de 2+)
     * Trabajadores: 1 forma ✗ (NO cumple umbral de 2+)
     * Resultado: 2 mandantes cumplen el umbral → SÍ califica para "al menos 2 mandantes en al menos 2 formas"

4. **ESTRUCTURA TU ANÁLISIS** así:
   ```
   CONTEO POR MANDANTE:
   - Gobierno: [N] formas identificadas → [LISTAR formas]
   - Empleadores: [N] formas identificadas → [LISTAR formas]
   - Trabajadores: [N] formas identificadas → [LISTAR formas]

   EVALUACIÓN DE UMBRALES (según nivel de la rúbrica):
   - Nivel X requiere: "AL MENOS [A] mandantes en AL MENOS [B] formas"
   - Mandantes que cumplen umbral de [B]+ formas: [LISTA DE MANDANTES]
   - Total de mandantes que cumplen umbral: [NÚMERO]
   - ¿Cumple requisito de [A]+ mandantes?: [SÍ/NO]

   JUSTIFICACIÓN DEL PUNTAJE:
   [Explicar EXPLÍCITAMENTE cómo los conteos y umbrales determinan el nivel]
   ```

5. **ASIGNA EL PUNTAJE** basándote ESTRICTAMENTE en:
   - Cuántos mandantes cumplen el umbral mínimo de formas
   - Si ese número cumple el requisito de "al menos X mandantes"

Siempre responde en español, incluso si el documento está en inglés."""

                        user_content = f"""Evalúa este documento contra el siguiente indicador (REQUIERE CONTEO CON LÓGICA DE UMBRALES):

**Indicador:** {criterion}

**Niveles de puntuación:** {json.dumps(descriptions, ensure_ascii=False, indent=2)}

**Documento a evaluar:**
{combined_text}

**INSTRUCCIONES CRÍTICAS:**
1. **PRIMERO: APLICA EL FILTRO DE TIEMPO VERBAL** - Revisa CADA mención de participación y verifica que sea factual (pasado o presente). RECHAZA cualquier participación futura, prometida o planificada.
2. CUENTA SOLO las formas de participación FACTUALES (pasadas o presentes) para CADA mandante
3. IDENTIFICA el umbral mínimo de formas requerido en cada nivel (ej: "al menos 2 formas")
4. CUENTA cuántos mandantes CUMPLEN ese umbral con participación REAL (no prometida)
5. VERIFICA si el número de mandantes que cumplen el umbral alcanza el requisito del nivel
6. Justifica el puntaje EXPLÍCITAMENTE basándote en la lógica de umbrales y menciona explícitamente que solo se contó participación factual

**EJEMPLO DE RAZONAMIENTO CORRECTO:**
"APLICACIÓN DEL FILTRO DE TIEMPO VERBAL: Se revisaron todas las menciones de participación. Se excluyeron las siguientes por ser futuras/planificadas: 'se consultará a los trabajadores durante la implementación', 'se espera que los empleadores participen'. Solo se contaron participaciones factuales (pasadas o presentes).

CONTEO POR MANDANTE (solo participación real):
- Gobierno: 3 formas identificadas → participó en reunión de diseño (15/03), validó marco lógico, proporcionó comentarios escritos (✓ todas factuales, cumple umbral de 2+)
- Empleadores: 2 formas identificadas → asistieron a consulta (20/03), revisaron propuesta técnica (✓ ambas factuales, cumple umbral de 2+)
- Trabajadores: 1 forma identificada → mencionados en documento pero sin evidencia de participación real (✗ NO cumple umbral de 2+)

EVALUACIÓN DE UMBRALES:
- Nivel 4 requiere: AL MENOS 2 mandantes en AL MENOS 2 formas cada uno
- Mandantes que cumplen umbral de 2+ formas: Gobierno (3 formas), Empleadores (2 formas)
- Total de mandantes que cumplen umbral: 2
- ¿Cumple requisito de 2+ mandantes?: SÍ

JUSTIFICACIÓN DEL PUNTAJE: Dado que 2 mandantes (Gobierno y Empleadores) cumplen el umbral de 2+ formas cada uno con participación REAL verificable, SÍ se alcanza el Nivel 4. Se excluyeron participaciones futuras/planificadas del conteo."

**RECORDATORIO FINAL CRÍTICO:**
- Si encuentras participaciones futuras o planificadas en el documento, MENCIONA explícitamente en tu análisis que fueron EXCLUIDAS del conteo
- Solo incluye en "evidence" citas que demuestren participación REAL (pasada o presente)
- Si el documento solo menciona participación futura/planificada y no hay evidencia de participación real, el puntaje debe reflejar esto (probablemente nivel 1 o 0)

Proporciona tu respuesta como JSON:
{{"analysis": "COMIENZA indicando si aplicaste el filtro de tiempo verbal y qué participaciones fueron excluidas (si las hubo). Luego presenta el conteo por mandante SOLO con participación factual. EVALÚA EXPLÍCITAMENTE los umbrales según cada nivel. Finalmente JUSTIFICA el puntaje con la lógica de umbrales, asegurándote de mencionar que solo se contó participación real. 2-3 párrafos en ESPAÑOL", "score": 1-5, "evidence": ["cita 1 que evidencia mandante y forma de participación específica REAL (pasada o presente)", "cita 2", "etc - 5-8 citas clave como array, SOLO participaciones factuales"]}}"""

                    else:
                        # Original prompt for non-counting rubrics
                        system_content = "Eres un evaluador experto de documentos. Siempre debes responder en español, incluso si el documento está en inglés."

                        user_content = f"""Evaluate this document against: {criterion}

    Scoring levels: {json.dumps(descriptions)}

    Relevant document sections:
    {combined_text}

    IMPORTANTE: Proporciona tu respuesta SIEMPRE en español, incluso si el documento está en inglés.

    Provide JSON with:
    {{"analysis": "detailed 2-3 paragraphs IN SPANISH", "score": 1-5, "evidence": ["quote 1", "quote 2", "quote 3", "etc - 5-8 key quotes from the text as an array"]}}"""

                    response = client.chat.completions.create(
                        model="gpt-5-mini",
                        messages=[
                            {"role": "system", "content": system_content},
                            {"role": "user", "content": user_content}
                        ],
                        max_completion_tokens=6500,
                        reasoning_effort="minimal",
                        timeout=120  # 2 minute timeout per request
                    )

                    content = response.choices[0].message.content.strip()
                    # Remove markdown code fences if present
                    if content.startswith('```'):
                        # Remove opening fence (```json or ```)
                        content = content.split('\n', 1)[1] if '\n' in content else content[3:]
                        # Remove closing fence
                        if content.endswith('```'):
                            content = content.rsplit('```', 1)[0]
                        content = content.strip()

                    result = json.loads(content)
                    # Normalize evidence field: convert array to string if needed
                    if isinstance(result.get('evidence'), list):
                        result['evidence'] = '\n'.join(result['evidence'])
                    return result

                except json.JSONDecodeError as e:
                    # If JSON parsing fails, return a default structure
                    return {
                        "analysis": f"Failed to parse JSON: {str(e)}. Raw response: {response.choices[0].message.content[:200]}",
                        "score": 3,
                        "evidence": "Unable to parse structured response",
                        "error": f"JSON parsing error: {str(e)}"
                    }
                except Exception as e:
                    # Check if it's a rate limit error
                    error_msg = str(e)
                    if "rate_limit" in error_msg.lower() or "429" in error_msg:
                        if attempt < max_retries - 1:
                            wait_time = (2 ** attempt) * 2  # Exponential backoff: 2s, 4s, 8s
                            time.sleep(wait_time)
                            continue

                    # If last attempt or non-rate-limit error, return error
                    return {
                        "analysis": f"Error during evaluation: {error_msg}",
                        "score": 0,
                        "evidence": "",
                        "error": f"API error (attempt {attempt + 1}/{max_retries}): {error_msg}"
                    }

            # If we exhausted all retries
            return {
                "analysis": "Failed after multiple retry attempts",
                "score": 0,
                "evidence": "",
                "error": f"Failed after {max_retries} attempts"
            }
        
        # Evaluate with selected criteria
        document_text = st.session_state.get('full_document_text_tab3', '')
        if not document_text:
            st.error("No se pudo recuperar el texto del documento. Por favor, vuelva a cargar el archivo.")
            st.stop()
        
        # Build filtered rubric based on selection
        filtered_rubric = {k: v for k, v in prodoc_rubric.items() if k in selected_criteria}
        
        if not filtered_rubric:
            st.error("No hay criterios seleccionados para evaluar.")
            st.stop()
        
        rubrics = [("Evaluación PRODOC", filtered_rubric)]
        
        rubric_results = []
        from concurrent.futures import ThreadPoolExecutor, as_completed
        MAX_WORKERS = 8
        
        # Define dimension order for sorting
        dimension_order_map = {'Diseño': 0, 'Implementación': 1, 'Evaluación': 2}
        
        def eval_one_criterion_tab3(args):
            unique_key, descriptions, dimension, criterio_text, indicador_text, sort_key, rubric_name = args
            try:
                result = evaluate_criterion_with_llm(document_text, indicador_text, descriptions)
                # Ensure result is a dictionary
                if not isinstance(result, dict):
                    result = {'score': 0, 'analysis': str(result), 'evidence': '', 'error': 'Invalid result format'}
                return {
                    'Dimensión': dimension,
                    'Criterio': criterio_text,
                    'Indicador': indicador_text,
                    'Score': result.get('score', 0),
                    'Análisis': str(result.get('analysis', '')),
                    'Evidencia': str(result.get('evidence', '')),
                    'Error': str(result.get('error', '')) if 'error' in result else '',
                    'Rúbrica': rubric_name,
                    '_dim_order': dimension_order_map.get(dimension, 99),
                    '_sort_key': sort_key
                }
            except Exception as e:
                return {
                    'Dimensión': dimension,
                    'Criterio': criterio_text,
                    'Indicador': indicador_text,
                    'Score': 0,
                    'Análisis': '',
                    'Evidencia': '',
                    'Error': str(e),
                    'Rúbrica': rubric_name,
                    '_dim_order': dimension_order_map.get(dimension, 99),
                    '_sort_key': sort_key
                }
        
        for rubric_name, rubric_dict in rubrics:
            rubric_analysis_data = []
            n_criteria = len(rubric_dict)
            progress = st.progress(0, text=f"Iniciando evaluación por rúbrica: {rubric_name}...")
            
            with st.spinner(f'Evaluando documento por rúbrica: {rubric_name}...'):
                with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    futures = {
                        executor.submit(eval_one_criterion_tab3, (
                            unique_key,
                            rubric_data['valores'] if isinstance(rubric_data, dict) else rubric_data,
                            rubric_data.get('dimension', 'No especificada') if isinstance(rubric_data, dict) else 'No especificada',
                            rubric_data.get('criterio', '') if isinstance(rubric_data, dict) else '',
                            rubric_data.get('indicador', unique_key) if isinstance(rubric_data, dict) else unique_key,
                            rubric_data.get('sort_key', (999, 999, 999)) if isinstance(rubric_data, dict) else (999, 999, 999),
                            rubric_name
                        )): (unique_key, idx)
                        for idx, (unique_key, rubric_data) in enumerate(rubric_dict.items())
                    }
                    
                    completed = 0
                    for future in as_completed(futures):
                        result = future.result()
                        rubric_analysis_data.append(result)
                        completed += 1
                        unique_key, idx = futures[future]
                        progress.progress(completed / n_criteria, text=f"Evaluando indicador...")
            
            # Create DataFrame and sort by dimension order and sort_key
            df_result = pd.DataFrame(rubric_analysis_data)
            if not df_result.empty:
                df_result = df_result.sort_values(by=['_dim_order', '_sort_key'])
                df_result = df_result.drop(columns=['_dim_order', '_sort_key'], errors='ignore')
            rubric_results.append((rubric_name, df_result))
        
        # Store results in session state for persistence
        st.session_state['tab3_results'] = rubric_results
        
        # Show and allow download of results
        if rubric_results:
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''

                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Dimensión', 'Criterio', 'Indicador', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Create visualizations for the results
            st.markdown("### Visualización de Resultados")
            
            all_scores = []
            for rubric_name, df in rubric_results:
                for _, row in df.iterrows():
                    all_scores.append({
                        'Indicador': row.get('Indicador', row.get('Criterio', '')),
                        'Criterio': row.get('Criterio', ''),
                        'Dimensión': row['Dimensión'],
                        'Puntuación': row['Score']
                    })
            
            scores_df = pd.DataFrame(all_scores)
            overall_avg = scores_df['Puntuación'].mean()
            
            # Visualization 1: Average Score by Dimension
            st.markdown("#### 📊 Puntuación Promedio por Dimensión")
            dimension_avg = scores_df.groupby('Dimensión')['Puntuación'].mean().reset_index()
            dimension_avg = dimension_avg.sort_values(by='Puntuación', ascending=False)
            
            fig_dim = go.Figure()
            fig_dim.add_trace(go.Bar(
                x=dimension_avg['Dimensión'],
                y=dimension_avg['Puntuación'],
                text=dimension_avg['Puntuación'].round(2),
                textposition='auto',
                marker_color='#002F6C',
                name='Promedio por Dimensión'
            ))
            
            fig_dim.add_trace(go.Scatter(
                x=dimension_avg['Dimensión'],
                y=[overall_avg] * len(dimension_avg),
                mode='lines',
                line=dict(color='#C8102E', width=2, dash='dash'),
                name=f'Promedio General: {overall_avg:.2f}'
            ))
            
            fig_dim.update_layout(
                title='Puntuación Promedio por Dimensión',
                xaxis_title='Dimensión',
                yaxis_title='Puntuación Promedio',
                yaxis=dict(range=[0, 5.5]),
                height=500,
                legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                margin=dict(l=20, r=20, t=80, b=50),
                hoverlabel=dict(bgcolor="white", font_size=12, font_family="Arial")
            )
            
            fig_dim.update_xaxes(tickangle=-45)
            fig_dim.update_yaxes(showgrid=True, gridwidth=1, gridcolor='LightGrey')
            
            st.plotly_chart(fig_dim, use_container_width=True)
            
            # Visualization 2: Individual Criteria Scores
            st.markdown("#### 📈 Puntuación por Criterio Individual")
            scores_df_sorted = scores_df.sort_values(by='Puntuación', ascending=False)
            
            fig = go.Figure()
            scores_df_sorted['Criterio_ID'] = [f"Criterio {i+1}" for i in range(len(scores_df_sorted))]
            scores_df_sorted['Hover_Text'] = scores_df_sorted.apply(
                lambda row: f"<b>{row['Criterio_ID']}</b><br>Dimensión: {row['Dimensión']}<br>{row['Criterio']}<br>Puntuación: {row['Puntuación']:.2f}", 
                axis=1
            )
            
            fig.add_trace(go.Bar(
                y=scores_df_sorted['Criterio_ID'],
                x=scores_df_sorted['Puntuación'],
                text=scores_df_sorted['Puntuación'].round(2),
                textposition='auto',
                marker_color='#0072CE',
                orientation='h',
                name='Puntuación',
                hovertext=scores_df_sorted['Hover_Text'],
                hoverinfo='text'
            ))
            
            fig.add_trace(go.Scatter(
                y=scores_df_sorted['Criterio_ID'],
                x=[overall_avg] * len(scores_df_sorted),
                mode='lines',
                line=dict(color='#C8102E', width=2, dash='dash'),
                name=f'Promedio General: {overall_avg:.2f}'
            ))
            
            fig.update_layout(
                title='Puntuación por Criterio (Ordenado de Mayor a Menor)',
                xaxis_title='Puntuación',
                yaxis_title='',
                xaxis=dict(range=[0, 5.5]),
                height=max(400, len(scores_df_sorted) * 35),
                width=800,
                legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                margin=dict(l=20, r=20, t=80, b=50),
                hoverlabel=dict(bgcolor="white", font_size=12, font_family="Arial")
            )
            
            fig.update_yaxes(automargin=True)
            fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor='LightGrey')
            
            st.plotly_chart(fig, use_container_width=True)
            
            # Download ZIP with XLSX files
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    # Export as XLSX instead of CSV
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_prodoc_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_evaluacion_prodoc.zip",
                mime="application/zip",
                key="prodoc_download_button_tab3"
            )
        else:
            st.warning("No se generaron resultados para ninguna rúbrica.")
    else:
        # Check if there are persisted results in session state
        if st.session_state.get('tab3_results') is not None:
            rubric_results = st.session_state['tab3_results']
            
            st.markdown("### 📊 Resultados guardados")
            
            for rubric_name, rubric_analysis_df in rubric_results:
                st.markdown(f'#### Resultados de la evaluación por rúbrica: {rubric_name}')
                if not rubric_analysis_df.empty:
                    if 'Evidencia' not in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = ''

                    cols = rubric_analysis_df.columns.tolist()
                    desired_order = ['Dimensión', 'Criterio', 'Indicador', 'Score', 'Análisis', 'Evidencia', 'Error', 'Rúbrica']
                    new_order = [col for col in desired_order if col in cols]
                    remaining_cols = [col for col in cols if col not in desired_order]
                    final_order = new_order + remaining_cols
                    rubric_analysis_df = rubric_analysis_df[final_order]
                    
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    
                    st.dataframe(rubric_analysis_df, use_container_width=True)
                else:
                    st.warning(f"No se generaron resultados para la rúbrica: {rubric_name}")
            
            # Create visualizations for the results
            st.markdown("### Visualización de Resultados")
            
            all_scores = []
            for rubric_name, df in rubric_results:
                for _, row in df.iterrows():
                    all_scores.append({
                        'Indicador': row.get('Indicador', row.get('Criterio', '')),
                        'Criterio': row.get('Criterio', ''),
                        'Dimensión': row['Dimensión'],
                        'Puntuación': row['Score']
                    })
            
            if all_scores:
                scores_df = pd.DataFrame(all_scores)
                overall_avg = scores_df['Puntuación'].mean()
                
                # Visualization: Average Score by Dimension
                st.markdown("#### 📊 Puntuación Promedio por Dimensión")
                dimension_avg = scores_df.groupby('Dimensión')['Puntuación'].mean().reset_index()
                dimension_avg = dimension_avg.sort_values(by='Puntuación', ascending=False)
                
                fig_dim = go.Figure()
                fig_dim.add_trace(go.Bar(
                    x=dimension_avg['Dimensión'],
                    y=dimension_avg['Puntuación'],
                    text=dimension_avg['Puntuación'].round(2),
                    textposition='auto',
                    marker_color='#002F6C',
                    name='Promedio por Dimensión'
                ))
                
                fig_dim.add_trace(go.Scatter(
                    x=dimension_avg['Dimensión'],
                    y=[overall_avg] * len(dimension_avg),
                    mode='lines',
                    line=dict(color='#C8102E', width=2, dash='dash'),
                    name=f'Promedio General: {overall_avg:.2f}'
                ))
                
                fig_dim.update_layout(
                    title='Puntuación Promedio por Dimensión',
                    xaxis_title='Dimensión',
                    yaxis_title='Puntuación Promedio',
                    yaxis=dict(range=[0, 5.5]),
                    height=500,
                    legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                    margin=dict(l=20, r=20, t=80, b=50),
                    hoverlabel=dict(bgcolor="white", font_size=12, font_family="Arial")
                )
                
                fig_dim.update_xaxes(tickangle=-45)
                fig_dim.update_yaxes(showgrid=True, gridwidth=1, gridcolor='LightGrey')
                
                st.plotly_chart(fig_dim, use_container_width=True)
            
            # Download ZIP
            import io, zipfile
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for rubric_name, rubric_analysis_df in rubric_results:
                    if 'Evidencia' in rubric_analysis_df.columns:
                        rubric_analysis_df['Evidencia'] = rubric_analysis_df['Evidencia'].apply(
                            lambda x: "\n".join(x) if isinstance(x, list) else (str(x) if x is not None else "")
                        )
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
                        rubric_analysis_df.to_excel(writer, index=False, sheet_name='Resultados')
                    excel_buffer.seek(0)
                    arcname = f"evaluacion_prodoc_{rubric_name.replace(' ', '_').lower()}.xlsx"
                    zipf.writestr(arcname, excel_buffer.getvalue())
            zip_buffer.seek(0)
            
            st.download_button(
                label="Descargar resultados como ZIP",
                data=zip_buffer,
                file_name="resultados_evaluacion_prodoc.zip",
                mime="application/zip",
                key="prodoc_download_button_tab3_persisted"
            )
            
            # Clear results button
            if st.button("🗑️ Limpiar resultados", key="clear_tab3_results"):
                st.session_state['tab3_results'] = None
                st.rerun()
        else:
            st.info("Seleccione criterios, suba un documento y presione 'Procesar y Evaluar'.")

# ================== TAB 7: APPRAISAL CHECKLIST (IMPROVED) =====================
# Configuration
from concurrent.futures import ThreadPoolExecutor, as_completed
import io
MAX_WORKERS = 48  # Reduced from 48 to avoid rate limits
OPENAI_MODEL = "gpt-5-mini"  # GPT-5 mini model with reasoning

@st.cache_data
def load_appraisal_questions():
    """Load and cache appraisal questions from Excel file"""
    try:
        df = pd.read_excel('./Appraisal Checklist_2025 es-419.xlsx', sheet_name='rubric')
        if 'Pregunta_Realizada' not in df.columns:
            return None, "La columna 'Pregunta_Realizada' no se encontró en el archivo de Excel."
        
        # Replace first 3 characters of Pregunta_Realizada with Tema numbering
        if 'Tema' in df.columns:
            df['Pregunta_Realizada'] = df.apply(
                lambda row: str(row['Tema']) + ' ' + str(row['Pregunta_Realizada'])[3:].strip() 
                if pd.notna(row['Tema']) and pd.notna(row['Pregunta_Realizada']) and len(str(row['Pregunta_Realizada'])) > 3
                else row['Pregunta_Realizada'], 
                axis=1
            )
        
        return df, None
    except FileNotFoundError:
        return None, "No se encontró el archivo Appraisal Checklist_2025 es-419.xlsx. Asegúrate de que exista en el directorio de la aplicación."
    except Exception as e:
        return None, f"Error al cargar el archivo de preguntas: {str(e)}"

def extract_document_content(uploaded_file):
    """Extract and process content from uploaded DOCX file"""
    try:
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_file_path = tmp_file.name
        
        # Extract content using docx2python
        doc_result = docx2python(tmp_file_path)
        
        # Get file stats
        file_size = os.path.getsize(tmp_file_path)
        
        # Extract structured content (you'll need to implement extract_docx_structure)
        # For now, using simple text extraction
        full_text = doc_result.text
        word_count = len(full_text.split())
        
        # Clean up temp file
        os.unlink(tmp_file_path)
        
        return {
            'text': full_text,
            'file_size': file_size,
            'word_count': word_count,
            'success': True
        }
    
    except Exception as e:
        return {
            'text': '',
            'file_size': 0,
            'word_count': 0,
            'success': False,
            'error': str(e)
        }

# def analyze_question_with_llm(question, document_text):
#     """Analyze a single question against the document using LLM"""
#     try:
#         response = openai.ChatCompletion.create(
#             model=OPENAI_MODEL,
#             messages=[
#                 {
#                     "role": "system", 
#                     "content": """You are an expert document analyst. Analyze the document against the given question and provide a structured JSON response with exactly this format:
#                     {
#                         "Respuesta": "Yes/No/Partial/Not Found",
#                         "Razonamiento": "Brief explanation of your analysis (max 200 words)",
#                         "Evidencia": "Specific text excerpts that support your answer (max 300 words)"
#                     }"""
#                 },
#                 {
#                     "role": "user", 
#                     "content": f"Question: {question}\n\nDocument Text: {document_text[:4000]}..."  # Limit context
#                 }
#             ],
#             max_completion_tokens=800,
#             temperature=0.1,
#         )
        
#         content = response["choices"][0]["message"]["content"].strip()
#         result = json.loads(content)
        
#         return {
#             'Pregunta': question,
#             'Respuesta': result.get('Respuesta', 'Error'),
#             'Razonamiento': result.get('Razonamiento', ''),
#             'Evidencia': result.get('Evidencia', ''),
#             'Status': 'Success'
#         }
    
#     except Exception as e:
#         return {
#             'Pregunta': question,
#             'Respuesta': 'Error',
#             'Razonamiento': f'Analysis failed: {str(e)}',
#             'Evidencia': '',
#             'Status': 'Error'
#         }

def parse_two_part_question(question):
    """
    Detect and parse two-part rubric questions where:
    - Part 1 (broader): Sets the general context
    - Part 2 (specific): Asks the critical detail that needs primary focus

    Returns dict with 'is_two_part', 'part1', 'part2', 'full_question'
    """
    import re

    # Pattern 1: Questions with explicit keywords separating them (check this first)
    # Keywords like "Específicamente" act as clear separators
    keywords = ['específicamente', 'en particular', 'en concreto', 'puntualmente', 'además']
    for keyword in keywords:
        # The keyword is a separator, not part of either question
        pattern = rf'(.*?\?)\s*{keyword}[,:]?\s*¿\s*(.*?\?)'
        match = re.search(pattern, question, re.IGNORECASE)
        if match:
            part2 = match.group(2).strip()
            # Ensure part2 starts with ¿
            if not part2.startswith('¿'):
                part2 = '¿' + part2
            return {
                'is_two_part': True,
                'part1': match.group(1).strip(),
                'part2': part2,
                'full_question': question
            }

    # Pattern 2: Questions run together like "...gráfico?Se identifican..." or "...?¿Se..."
    # This is the most common pattern in the rubrics (no keyword separator)
    match = re.search(r'(.*?\?)\s*¿?\s*([A-ZSÉA].*?\?)', question)
    if match:
        part1 = match.group(1).strip()
        part2 = match.group(2).strip()

        # Ensure part2 starts with ¿ if it doesn't already
        if not part2.startswith('¿'):
            part2 = '¿' + part2

        return {
            'is_two_part': True,
            'part1': part1,
            'part2': part2,
            'full_question': question
        }

    # Not a two-part question
    return {
        'is_two_part': False,
        'part1': None,
        'part2': None,
        'full_question': question
    }

def analyze_question_with_llm_tab1(question, document_text):
    """
    Analyze a single question against the document using LLM with full context (up to 110K tokens).
    TAB1-SPECIFIC VERSION: Explicitly states when 2+ parts are detected in the question.
    """
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        combined_text = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        if not combined_text.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Not Found',
                'Razonamiento': 'No se encontró contenido en el documento para analizar.',
                'Evidencia': '',
                'Status': 'Success'
            }

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize system prompt based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert document analyst specializing in ILO project quality assessment.

**CRITICAL INSTRUCTION FOR TWO-PART QUESTIONS:**

This question has TWO PARTS with different priorities:
1. **Part 1 (Broader Context)**: Provides the general framework
2. **Part 2 (Specific Focus - PRIMARY)**: The critical question that requires detailed analysis

**YOUR ANALYSIS APPROACH:**
1. **START with Part 2**: Analyze the specific question FIRST and IN DEPTH
2. **Then Part 1**: Consider how Part 2's findings relate to the broader context in Part 1
3. **Final Assessment**: The overall answer should be driven by Part 2, but explained in context of Part 1

**Response Structure (JSON):**
{
    "Respuesta": "Yes/No/Partial/Not Found (based primarily on Part 2)",
    "Razonamiento": "**DEBE COMENZAR CON**: 'Se identificaron 2 partes en esta pregunta.' LUEGO responder la pregunta específica (Parte 2) con análisis detallado. DESPUÉS explicar cómo esto se relaciona con la pregunta general (Parte 1). Máximo 200 palabras.",
    "Evidencia": "Proporcionar evidencia PRINCIPALMENTE para la Parte 2 (pregunta específica), luego evidencia de apoyo para la Parte 1. Máximo 300 palabras con citas directas."
}

**Estructura OBLIGATORIA para Razonamiento:**
"Se identificaron 2 partes en esta pregunta. [Respuesta específica a Parte 2 con análisis detallado]. Esto [afecta/se relaciona con] [Parte 1] porque [explicación de la relación]."

Siempre responder en español. Enfoque: 70% en Parte 2, 30% en cómo se relaciona con Parte 1."""

            user_content = f"""PREGUNTA CON DOS PARTES:

**PARTE 1 (Contexto General):**
{parsed['part1']}

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**Pregunta Completa:**
{question}

**Texto del Documento:**
{combined_text}

RECUERDA:
1. COMENZAR tu Razonamiento con "Se identificaron 2 partes en esta pregunta."
2. Enfoca tu análisis PRINCIPALMENTE en la Parte 2 (pregunta específica)
3. Luego explica cómo se relaciona con la Parte 1 (contexto general)"""

        else:
            # Original single-question prompt
            system_content = """You are an expert document analyst. Analyze the document against the given question and provide a structured JSON response with exactly this format and always respond in Spanish:
            {
                "Respuesta": "Yes/No/Partial/Not Found",
                "Razonamiento": "Brief explanation of your analysis (max 200 words)",
                "Evidencia": "Specific text excerpts that support your answer (max 300 words)"
            }"""

            user_content = f"Question: {question}\n\nDocument Text: {combined_text}"

        # Single API call per question - much more efficient than chunking
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_content
                },
                {
                    "role": "user",
                    "content": user_content
                }
            ],
            max_completion_tokens=8000,
            reasoning_effort="minimal"
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'No se recibió respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

        # Parse JSON response
        try:
            result = json.loads(content.strip())
            return {
                'Pregunta': question,
                'Respuesta': result.get('Respuesta', 'Not Found'),
                'Razonamiento': result.get('Razonamiento', ''),
                'Evidencia': result.get('Evidencia', ''),
                'Status': 'Success'
            }
        except json.JSONDecodeError:
            # If JSON parsing fails, return error
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'Error al procesar la respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

    except Exception as e:
        return {
            'Pregunta': question,
            'Respuesta': 'Error',
            'Razonamiento': f'Analysis failed: {str(e)}',
            'Evidencia': '',
            'Status': 'Error'
        }

def analyze_question_with_llm(question, document_text):
    """Analyze a single question against the document using LLM with full context (up to 110K tokens)."""
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        combined_text = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        if not combined_text.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Not Found',
                'Razonamiento': 'No se encontró contenido en el documento para analizar.',
                'Evidencia': '',
                'Status': 'Success'
            }

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize system prompt based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert document analyst specializing in ILO project quality assessment.

**CRITICAL INSTRUCTION FOR TWO-PART QUESTIONS:**

This question has TWO PARTS with different priorities:
1. **Part 1 (Broader Context)**: Provides the general framework
2. **Part 2 (Specific Focus - PRIMARY)**: The critical question that requires detailed analysis

**YOUR ANALYSIS APPROACH:**
1. **START with Part 2**: Analyze the specific question FIRST and IN DEPTH
2. **Then Part 1**: Consider how Part 2's findings relate to the broader context in Part 1
3. **Final Assessment**: The overall answer should be driven by Part 2, but explained in context of Part 1

**Response Structure (JSON):**
{
    "Respuesta": "Yes/No/Partial/Not Found (based primarily on Part 2)",
    "Razonamiento": "BEGIN by answering the specific question (Part 2) with detailed analysis. THEN explain how this relates to the broader question (Part 1). Maximum 200 words.",
    "Evidencia": "Provide evidence PRIMARILY for Part 2 (the specific question), then supporting evidence for Part 1. Maximum 300 words with direct quotes."
}

**Example Structure for Razonamiento:**
"Respecto a [specific Part 2]: [detailed analysis of Part 2]... Esto [connects to/affects] [broader Part 1] porque [explanation of relationship]."

Always respond in Spanish. Focus 70% on Part 2, 30% on how it relates to Part 1."""

            user_content = f"""PREGUNTA CON DOS PARTES:

**PARTE 1 (Contexto General):**
{parsed['part1']}

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**Pregunta Completa:**
{question}

**Texto del Documento:**
{combined_text}

RECUERDA: Enfoca tu análisis PRINCIPALMENTE en la Parte 2 (pregunta específica), luego explica cómo se relaciona con la Parte 1 (contexto general)."""

        else:
            # Original single-question prompt
            system_content = """You are an expert document analyst. Analyze the document against the given question and provide a structured JSON response with exactly this format and always respond in Spanish:
            {
                "Respuesta": "Yes/No/Partial/Not Found",
                "Razonamiento": "Brief explanation of your analysis (max 200 words)",
                "Evidencia": "Specific text excerpts that support your answer (max 300 words)"
            }"""

            user_content = f"Question: {question}\n\nDocument Text: {combined_text}"

        # Single API call per question - much more efficient than chunking
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_content
                },
                {
                    "role": "user",
                    "content": user_content
                }
            ],
            max_completion_tokens=8000,
            reasoning_effort="minimal"
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'No se recibió respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

        # Parse JSON response
        try:
            result = json.loads(content.strip())
            return {
                'Pregunta': question,
                'Respuesta': result.get('Respuesta', 'Not Found'),
                'Razonamiento': result.get('Razonamiento', ''),
                'Evidencia': result.get('Evidencia', ''),
                'Status': 'Success'
            }
        except json.JSONDecodeError:
            # If JSON parsing fails, return error
            return {
                'Pregunta': question,
                'Respuesta': 'Error',
                'Razonamiento': 'Error al procesar la respuesta del modelo.',
                'Evidencia': '',
                'Status': 'Error'
            }

    except Exception as e:
        return {
            'Pregunta': question,
            'Respuesta': 'Error',
            'Razonamiento': f'Analysis failed: {str(e)}',
            'Evidencia': '',
            'Status': 'Error'
        }

def analyze_question_with_critical_opinion_tab1(question, answer, reasoning, evidence, document_text=""):
    """
    TAB1-SPECIFIC VERSION: Critically assess if the document's answer to a specific question is adequate.
    Explicitly checks if the reasoning properly identifies and addresses two-part questions.
    """
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        doc_context = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize critical evaluation based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).

**CRITICAL EVALUATION FOR TWO-PART QUESTIONS:**

This is a TWO-PART question where:
- **Part 1 (Broader)**: Sets general context
- **Part 2 (Specific - PRIMARY)**: The critical detail that needs focused assessment

**Your Critical Assessment Must Evaluate:**

1. **FIRST - Check Structure**: Does the reasoning BEGIN with "Se identificaron 2 partes en esta pregunta" or similar acknowledgment?
   - If NO: This is a CRITICAL FAILURE - the analysis didn't recognize the question structure

2. **PRIMARY FOCUS - Part 2 (Specific Question):**
   - Does the answer adequately address the SPECIFIC question (Part 2)?
   - Is the evidence for Part 2 substantive and concrete?
   - Are there critical gaps or superficial treatment of Part 2?
   - Does the document truly deliver on the specific requirement?

3. **SECONDARY - Part 1 (Broader Context):**
   - Does the answer address the broader context (Part 1)?
   - How do the Part 2 findings affect the Part 1 assessment?

4. **INTEGRATION:**
   - Is there coherence between how Part 2 relates to Part 1?
   - Are there contradictions or misalignments?

**Be especially critical if:**
- The reasoning doesn't acknowledge the two-part structure
- Part 2 (specific question) is answered superficially or avoided
- Evidence focuses on Part 1 but neglects Part 2
- The reasoning doesn't connect Part 2's findings to Part 1's context

Respond in Spanish with a brief (max 150 words) critical assessment. Be direct about shortcomings, especially regarding Part 2.
Respond ONLY with the critical opinion text, no JSON, no formatting."""

            user_content = f"""PREGUNTA CON DOS PARTES:

**PARTE 1 (Contexto General):**
{parsed['part1']}

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**Pregunta Completa:**
{question}

**Respuesta del Documento:** {answer}

**Razonamiento del Documento:** {reasoning}

**Evidencia del Documento:** {evidence}

**Contexto Completo del Documento:**
{doc_context}

Evalúa críticamente:
1. ¿El razonamiento reconoce explícitamente que hay 2 partes en la pregunta?
2. ¿La respuesta aborda adecuadamente AMBAS partes, especialmente la Parte 2 (pregunta específica)?
3. ¿La evidencia y el razonamiento son suficientes para ambas partes?"""

        else:
            # Original single-question critical evaluation
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).
                    Critically assess whether the document's answer to a specific question is adequate, appropriate, and complete.

                    Review the answer, reasoning, evidence, AND full document context together to evaluate:
                    - Does the answer fully address the concern raised in the question?
                    - Is the evidence substantive enough to support the answer?
                    - Does the full document context confirm or contradict the claimed answer?
                    - Is the proposed approach sufficient according to best practices?
                    - Are there gaps, risks, or inadequacies in what the document claims?
                    - Should the project have included additional measures or details?
                    - Is the reasoning robust or superficial?
                    - What is NOT mentioned that should be?

                    Respond in Spanish with a brief (max 150 words) critical assessment. Be direct about any shortcomings.
                    Respond ONLY with the critical opinion text, no JSON, no formatting."""

            user_content = f"""Question: {question}

Document's Answer: {answer}

Document's Reasoning: {reasoning}

Document's Evidence: {evidence}

Full Document Context (complete):
{doc_context}

Provide a critical assessment: Is this answer truly adequate from an expert perspective, considering the quality of the reasoning, evidence, and complete document context provided?"""

        # Single API call for critical evaluation - focuses on answer adequacy
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_content
                },
                {
                    "role": "user",
                    "content": user_content
                }
            ],
            max_completion_tokens=500,
            reasoning_effort="minimal"
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return "No se generó evaluación crítica."

        return content.strip()

    except Exception as e:
        return f"Error en evaluación crítica: {str(e)}"

def analyze_question_with_critical_opinion(question, answer, reasoning, evidence, document_text=""):
    """
    Critically assess if the document's answer to a specific question is adequate and appropriate.
    Evaluates whether the document's response truly addresses the concern raised in the question.
    Uses answer, reasoning, evidence AND complete document context for comprehensive critical assessment.
    Uses up to 110K tokens for complete document understanding - maximizes context window usage.
    Enhanced to handle two-part questions with proper focus on the specific (Part 2) aspect.
    """
    try:
        # Truncate to ~110K tokens to maximize context while leaving room for prompts and response
        doc_context = truncate_to_token_limit(document_text or "", max_tokens=110000, encoding_obj=encoding)

        # Parse question to detect two-part structure
        parsed = parse_two_part_question(question)

        # Customize critical evaluation based on question structure
        if parsed['is_two_part']:
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).

**CRITICAL EVALUATION FOR TWO-PART QUESTIONS:**

This is a TWO-PART question where:
- **Part 1 (Broader)**: Sets general context
- **Part 2 (Specific - PRIMARY)**: The critical detail that needs focused assessment

**Your Critical Assessment Must Evaluate:**

1. **PRIMARY FOCUS - Part 2 (Specific Question):**
   - Does the answer adequately address the SPECIFIC question (Part 2)?
   - Is the evidence for Part 2 substantive and concrete?
   - Are there critical gaps or superficial treatment of Part 2?
   - Does the document truly deliver on the specific requirement?

2. **SECONDARY - Part 1 (Broader Context):**
   - Does the answer address the broader context (Part 1)?
   - How do the Part 2 findings affect the Part 1 assessment?

3. **INTEGRATION:**
   - Is there coherence between how Part 2 relates to Part 1?
   - Are there contradictions or misalignments?

**Be especially critical if:**
- Part 2 (specific question) is answered superficially or avoided
- Evidence focuses on Part 1 but neglects Part 2
- The reasoning doesn't connect Part 2's findings to Part 1's context

Respond in Spanish with a brief (max 150 words) critical assessment. Be direct about shortcomings, especially regarding Part 2.
Respond ONLY with the critical opinion text, no JSON, no formatting."""

            user_content = f"""PREGUNTA CON DOS PARTES:

**PARTE 1 (Contexto General):**
{parsed['part1']}

**PARTE 2 (Enfoque Específico - PRIORITARIO):**
{parsed['part2']}

**Pregunta Completa:**
{question}

**Respuesta del Documento:** {answer}

**Razonamiento del Documento:** {reasoning}

**Evidencia del Documento:** {evidence}

**Contexto Completo del Documento:**
{doc_context}

Evalúa críticamente: ¿La respuesta aborda adecuadamente AMBAS partes, especialmente la Parte 2 (pregunta específica)? ¿La evidencia y el razonamiento son suficientes para ambas partes?"""

        else:
            # Original single-question critical evaluation
            system_content = """You are an expert in project quality appraisal and international development (ILO standards).
                    Critically assess whether the document's answer to a specific question is adequate, appropriate, and complete.

                    Review the answer, reasoning, evidence, AND full document context together to evaluate:
                    - Does the answer fully address the concern raised in the question?
                    - Is the evidence substantive enough to support the answer?
                    - Does the full document context confirm or contradict the claimed answer?
                    - Is the proposed approach sufficient according to best practices?
                    - Are there gaps, risks, or inadequacies in what the document claims?
                    - Should the project have included additional measures or details?
                    - Is the reasoning robust or superficial?
                    - What is NOT mentioned that should be?

                    Respond in Spanish with a brief (max 150 words) critical assessment. Be direct about any shortcomings.
                    Respond ONLY with the critical opinion text, no JSON, no formatting."""

            user_content = f"""Question: {question}

Document's Answer: {answer}

Document's Reasoning: {reasoning}

Document's Evidence: {evidence}

Full Document Context (complete):
{doc_context}

Provide a critical assessment: Is this answer truly adequate from an expert perspective, considering the quality of the reasoning, evidence, and complete document context provided?"""

        # Single API call for critical evaluation - focuses on answer adequacy
        resp = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_content
                },
                {
                    "role": "user",
                    "content": user_content
                }
            ],
            max_completion_tokens=500,
            reasoning_effort="minimal"
        )

        content = resp.choices[0].message.content
        if not content or not content.strip():
            return "No se generó evaluación crítica."

        return content.strip()

    except Exception as e:
        return f"Error en evaluación crítica: {str(e)}"

def extract_section_number(question_text):
    """Extract section number from question text (e.g., '1.1 ¿Pregunta?' -> 1)"""
    import re
    match = re.match(r'(\d+)\.', str(question_text).strip())
    return int(match.group(1)) if match else None

def extract_subsection_number(question_text):
    """Extract full section.subsection from question text (e.g., '1.1 ¿Pregunta?' -> '1.1', '2.3 ¿Pregunta?' -> '2.3')"""
    import re
    match = re.match(r'(\d+\.\d+)', str(question_text).strip())
    return match.group(1) if match else None

def parse_subsection_for_sorting(subsection_str):
    """Convert subsection string to tuple for proper sorting (e.g., '1.1' -> (1, 1), '2.10' -> (2, 10))"""
    try:
        parts = str(subsection_str).split('.')
        return (int(parts[0]), int(parts[1]))
    except:
        return (999, 999)

def synthesize_subsection_analysis(subsection_id, subsection_questions_df):
    """Synthesize subsection-level analysis from individual question answers within that subsection"""
    try:
        # Build context from individual question Q&A pairs
        qa_context = "\n\n".join([
            f"Pregunta {row['Pregunta']}:\nRespuesta: {row['Respuesta']}\nRazonamiento: {row['Razonamiento']}"
            for _, row in subsection_questions_df.iterrows()
        ])
        
        # Single efficient LLM call per subsection
        prompt = f"""Based on the following individual question answers from subsection {subsection_id} of a document evaluation, 
synthesize a comprehensive subsection-level analysis.

Subsection {subsection_id} - Individual Q&A:
{qa_context}

Provide a concise subsection-level analysis (1-2 paragraphs) that:
1. Integrates findings across all questions in this subsection
2. Identifies key strengths and gaps
3. Provides a clear assessment

Format as JSON with exactly this structure:
{{"subsection_analysis": "your analysis here (1-2 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert document analyst. Synthesize individual question findings into clear subsection-level insights. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=1500,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('subsection_analysis', 'Error en síntesis')
    
    except Exception as e:
        return f"Error generando análisis de subsección: {str(e)}"

def synthesize_section_analysis(section_num, subsection_analyses_dict):
    """Synthesize section-level analysis from subsection analyses"""
    try:
        # Build context from subsection analyses
        subsection_context = "\n\n".join([
            f"Subsección {subsec_id}:\n{analysis}"
            for subsec_id, analysis in sorted(subsection_analyses_dict.items(), key=lambda x: parse_subsection_for_sorting(x[0]))
        ])
        
        # Single efficient LLM call per section
        prompt = f"""Based on the following subsection analyses from section {section_num} of a document evaluation, 
synthesize a comprehensive section-level analysis.

Section {section_num} - Subsection Analyses:
{subsection_context}

Provide a detailed section-level analysis (2-3 paragraphs) that:
1. Integrates key findings across all subsections
2. Identifies overarching patterns, strengths, and gaps
3. Provides strategic, actionable insights for improvement

Format as JSON with exactly this structure:
{{"section_analysis": "your detailed analysis here (2-3 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert document analyst. Synthesize subsection findings into clear, actionable section-level insights. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=2000,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('section_analysis', 'Error en síntesis')
    
    except Exception as e:
        return f"Error generando análisis de sección: {str(e)}"

def synthesize_critical_evaluation_subsection(subsection_id, critical_opinions_df):
    """Synthesize critical evaluations at subsection level from individual question critical opinions"""
    try:
        # Build context from individual critical opinions
        critical_context = "\n\n".join([
            f"Pregunta {row['Pregunta']}:\nEvaluación Crítica: {row['Evaluación Crítica']}"
            for _, row in critical_opinions_df.iterrows()
        ])
        
        # Single efficient LLM call per subsection
        prompt = f"""Based on the following individual critical evaluations from subsection {subsection_id}, 
synthesize a comprehensive subsection-level critical assessment.

Subsection {subsection_id} - Individual Critical Evaluations:
{critical_context}

Provide a concise subsection-level critical assessment (1-2 paragraphs) that:
1. Integrates critical findings across all questions in this subsection
2. Identifies major gaps, risks, and inadequacies
3. Highlights patterns in insufficient responses
4. Provides clear recommendations for improvement

Format as JSON with exactly this structure:
{{"critical_evaluation": "your critical assessment here (1-2 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert critic of project quality. Synthesize individual critical opinions into clear subsection-level critical assessment. Be direct about inadequacies. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=1500,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('critical_evaluation', 'Error en síntesis crítica')
    
    except Exception as e:
        return f"Error generando evaluación crítica de subsección: {str(e)}"

def synthesize_critical_evaluation_section(section_num, critical_subsection_dict):
    """Synthesize critical evaluations at section level from subsection critical assessments"""
    try:
        # Build context from subsection critical evaluations
        critical_context = "\n\n".join([
            f"Subsección {subsec_id}:\n{evaluation}"
            for subsec_id, evaluation in sorted(critical_subsection_dict.items(), key=lambda x: parse_subsection_for_sorting(x[0]))
        ])
        
        # Single efficient LLM call per section
        prompt = f"""Based on the following subsection critical evaluations from section {section_num}, 
synthesize a comprehensive section-level critical assessment.

Section {section_num} - Subsection Critical Evaluations:
{critical_context}

Provide a detailed section-level critical assessment (2-3 paragraphs) that:
1. Integrates critical findings across all subsections
2. Identifies overarching gaps, risks, and systemic inadequacies
3. Highlights critical patterns of insufficient responses
4. Provides strategic recommendations for improving the overall section

Format as JSON with exactly this structure:
{{"critical_evaluation": "your critical assessment here (2-3 paragraphs)"}}

Return ONLY the JSON, no other text."""

        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert critic of project quality. Synthesize subsection critical findings into clear, actionable section-level critical assessment. Be direct and strategic. Always respond in Spanish."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_completion_tokens=2000,
            reasoning_effort="minimal"
        )
        
        content = response.choices[0].message.content.strip()
        result = json.loads(content)
        return result.get('critical_evaluation', 'Error en síntesis crítica')
    
    except Exception as e:
        return f"Error generando evaluación crítica de sección: {str(e)}"

def create_results_download_with_sections(results_df, subsection_analyses, subsection_critical_evaluations, section_analyses, section_critical_evaluations, filename_base="appraisal_checklist"):
    """Create ZIP file with results including three separate sheets for questions, subsections, and sections"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        excel_buffer = io.BytesIO()
        
        with pd.ExcelWriter(excel_buffer, engine='xlsxwriter', engine_kwargs={'options': {'strings_to_urls': False}}) as writer:
            workbook = writer.book
            
            # Format definitions
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#002F6C',
                'font_color': 'white',
                'border': 1,
                'valign': 'vcenter',
                'text_wrap': True
            })
            section_header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#0072CE',
                'font_color': 'white',
                'border': 1,
                'valign': 'vcenter',
                'text_wrap': True,
                'font_size': 12
            })
            subsection_header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#4A90E2',
                'font_color': 'white',
                'border': 1,
                'valign': 'vcenter',
                'text_wrap': True,
                'font_size': 11
            })
            merged_format = workbook.add_format({
                'border': 1,
                'valign': 'top',
                'text_wrap': True,
                'bg_color': '#F5F5F5'
            })
            subsection_merged_format = workbook.add_format({
                'border': 1,
                'valign': 'top',
                'text_wrap': True,
                'bg_color': '#E8F4F8',
                'italic': True
            })
            normal_format = workbook.add_format({
                'border': 1,
                'valign': 'top',
                'text_wrap': True
            })
            
            # Extract and ensure sorting columns exist
            if '_section' not in results_df.columns:
                results_df['_section'] = results_df['Pregunta'].apply(extract_section_number)
            if '_subsection' not in results_df.columns:
                results_df['_subsection'] = results_df['Pregunta'].apply(extract_subsection_number)
            if '_sort_key' not in results_df.columns:
                results_df['_sort_key'] = results_df['_subsection'].apply(parse_subsection_for_sorting)
            
            # Sort the entire dataframe by sort_key to ensure consistent ordering
            results_df_sorted = results_df.sort_values(by='_sort_key').reset_index(drop=True)
            
            # ===== SHEET 1: Questions (Preguntas) =====
            sheet_questions = workbook.add_worksheet('1. Preguntas')
            questions_data = results_df_sorted[['_subsection', 'Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia', 'Evaluación Crítica', 'Status']].copy()
            questions_data.columns = ['Subsección', 'Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia', 'Evaluación Crítica', 'Status']
            questions_data.to_excel(writer, index=False, sheet_name='1. Preguntas', startrow=0)
            
            # Format headers
            for col_num, value in enumerate(questions_data.columns.values):
                writer.sheets['1. Preguntas'].write(0, col_num, value, header_format)
            
            # Set column widths
            writer.sheets['1. Preguntas'].set_column('A:A', 12)  # Subsección
            writer.sheets['1. Preguntas'].set_column('B:B', 35)  # Pregunta
            writer.sheets['1. Preguntas'].set_column('C:C', 12)  # Respuesta
            writer.sheets['1. Preguntas'].set_column('D:D', 50)  # Razonamiento
            writer.sheets['1. Preguntas'].set_column('E:E', 40)  # Evidencia
            writer.sheets['1. Preguntas'].set_column('F:F', 45)  # Evaluación Crítica
            writer.sheets['1. Preguntas'].set_column('G:G', 10)  # Status
            
            # ===== SHEET 2: Subsection Analysis (Análisis por Subsección) =====
            sheet_subsections = workbook.add_worksheet('2. Análisis Subsecciones')
            
            # Get unique subsections in sorted order
            subsections_sorted = sorted(results_df_sorted['_subsection'].dropna().unique(), key=parse_subsection_for_sorting)
            
            # Write headers
            subsec_headers = ['Subsección', 'Sección', 'Análisis de Subsección']
            for col_num, header in enumerate(subsec_headers):
                sheet_subsections.write(0, col_num, header, header_format)
            
            current_row = 1
            for subsection_id in subsections_sorted:
                section_num = extract_section_number(subsection_id)
                analysis_text = subsection_analyses.get(subsection_id, "No se generó análisis")
                critical_text = subsection_critical_evaluations.get(subsection_id, "No se generó evaluación crítica")
                
                sheet_subsections.write(current_row, 0, subsection_id, normal_format)
                sheet_subsections.write(current_row, 1, f"Sección {section_num}", normal_format)
                sheet_subsections.write(current_row, 2, analysis_text, subsection_merged_format)
                sheet_subsections.write(current_row, 3, critical_text, subsection_merged_format)
                current_row += 1
            
            # Set column widths
            sheet_subsections.set_column('A:A', 12)  # Subsección
            sheet_subsections.set_column('B:B', 15)  # Sección
            sheet_subsections.set_column('C:C', 80)  # Análisis
            sheet_subsections.set_column('D:D', 80)  # Evaluación Crítica
            
            # ===== SHEET 3: Section Analysis (Análisis por Sección) =====
            sheet_sections = workbook.add_worksheet('3. Análisis Secciones')
            
            # Get unique sections in sorted order
            sections_sorted = sorted(results_df_sorted['_section'].dropna().unique())
            
            # Write headers
            section_headers = ['Sección', 'Análisis de Sección', 'Evaluación Crítica']
            for col_num, header in enumerate(section_headers):
                sheet_sections.write(0, col_num, header, header_format)
            
            current_row = 1
            for section_num in sections_sorted:
                section_num = int(section_num)
                analysis_text = section_analyses.get(section_num, "No se generó análisis")
                critical_text = section_critical_evaluations.get(section_num, "No se generó evaluación crítica")
                
                sheet_sections.write(current_row, 0, f"Sección {section_num}", section_header_format)
                sheet_sections.write(current_row, 1, analysis_text, merged_format)
                sheet_sections.write(current_row, 2, critical_text, merged_format)
                current_row += 1
            
            # Set column widths
            sheet_sections.set_column('A:A', 15)  # Sección
            sheet_sections.set_column('B:B', 90)  # Análisis
            sheet_sections.set_column('C:C', 90)  # Evaluación Crítica
        
        excel_buffer.seek(0)
        zipf.writestr(f"{filename_base}_results.xlsx", excel_buffer.getvalue())

        # Add original template if available
        try:
            with open('./Appraisal Checklist_2025 es-419.xlsx', 'rb') as f:
                template_data = f.read()
                zipf.writestr(f"{filename_base}_rubric_template.xlsx", template_data)
        except FileNotFoundError:
            pass
        
        # Add summary report
        summary = f"""
Resumen del Análisis de la Lista de la Valoración Preliminar de la Calidad
===========================================================

Archivo Excel generado con 3 hojas:
  1. Preguntas: Análisis individual de cada pregunta (ordenado por subsección)
  2. Análisis Subsecciones: Síntesis de preguntas agrupadas por subsección
  3. Análisis Secciones: Síntesis ejecutiva por sección

Total de preguntas analizadas: {len(results_df)}
Análisis exitosos: {len(results_df[results_df['Status'] == 'Success'])}
Análisis fallidos: {len(results_df[results_df['Status'] == 'Error'])}

Subsecciones analizadas: {len(subsection_analyses)}
Secciones analizadas: {len(section_analyses)}

Distribución de respuestas:
{results_df['Respuesta'].value_counts().to_string()}
        """
        zipf.writestr(f"{filename_base}_summary.txt", summary)
    
    zip_buffer.seek(0)
    return zip_buffer

# Main tab interface
with tab1:
    st.header("📋 Valoración Preliminar de Calidad de Proyectos")
    
    # Load questions
    df_appraisal, error_msg = load_appraisal_questions()
    
    if error_msg:
        st.error(error_msg)
        st.stop()
    
    # Download button for the rubric file (directly on page, no expander)
    # Download button for the rubric file (directly on page, no expander)
    try:
        with open('./Appraisal Checklist_2025 es-419.xlsx', 'rb') as f:
            st.download_button(
                label="📥 Descargar archivo rúbrica de Valoración preliminar",
                data=f,
                file_name="Appraisal Checklist_2025 es-419.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_appraisal_rubric"
            )
    except FileNotFoundError:
        st.warning("Archivo de rúbrica no disponible para descarga.")
    
    # Instrucciones
    st.markdown("""
    ### Instrucciones
    
    1. **Subir documento**: Selecciona un Documento de Diseño de Proyecto en archivo de formato DOCX para el análisis de valoración preliminar de calidad.
    2. **Procesar**: Haz clic en 'Analizar documento' para iniciar la evaluación
    3. **Revisar resultados**: Examina los resultados del análisis en la tabla interactiva
    4. **Descargar**: Obtén todos los resultados y las pruebas en un archivo ZIP
    
    ---
    """)
    
  
    # Initialize session state for Tab 1 results persistence
    if 'tab1_results_df' not in st.session_state:
        st.session_state['tab1_results_df'] = None
    if 'tab1_doc_stats' not in st.session_state:
        st.session_state['tab1_doc_stats'] = None

    # Initialize session state for Tab 1
    if 'document_extracted_tab1' not in st.session_state:
        st.session_state['document_extracted_tab1'] = False
    if 'selected_sections_tab1' not in st.session_state:
        st.session_state['selected_sections_tab1'] = []
    
    # Document upload
    st.subheader("📄 Carga de documento")
    
    # Warning box about document requirements
    st.warning("""
    **⚠️ Requisitos importantes para la carga de documentos:**
    
    **📝 Formato del documento:**
    - Solo se aceptan archivos en formato **.docx** (Word 2007 o posterior)
    - El documento debe estar **correctamente formateado** usando los estilos de encabezado de Word (Heading 1, Heading 2, etc.)
    - **CRÍTICO:** Las secciones del documento deben estar identificadas con **encabezados usando estilos estándar de Word**. Sin encabezados apropiados, el texto no se extraerá correctamente y las secciones no se identificarán.
    
    **📊 Límites de contexto:**
    - El sistema procesa hasta **110,000 tokens** (~440,000 caracteres, aproximadamente **150-200 páginas**) por documento
    - Documentos que excedan este límite serán truncados automáticamente
    - Se recomienda dividir documentos muy extensos (más de ~180 páginas) en secciones más pequeñas si es necesario
    
    **✅ Mejores prácticas:**
    - Usa estilos de Word (Título 1, Título 2, etc.) para identificar secciones principales
    - Evita usar texto en negrita o mayúsculas como sustituto de encabezados
    - Asegúrate de que el documento esté guardado correctamente antes de subirlo
    """)
    
    uploaded_file = st.file_uploader(
        "Sube un DOCX para la evaluación:",
        type=["docx"],
        key="appraisal_file_uploader",
        help="Selecciona un documento de Word (.docx) para el análisis de valoración preliminar de calidad"
    )
    
    # Document Extraction Section
    st.markdown("---")
    st.markdown("### 📥 Extracción de Documento")
    
    if uploaded_file is not None:
        file_hash = hash(uploaded_file.getvalue())
        file_changed = st.session_state.get('last_file_hash_tab1') != file_hash
        
        if file_changed:
            st.session_state['document_extracted_tab1'] = False
            st.session_state['last_file_hash_tab1'] = None
        
        if st.button("🔍 Extraer Documento", key="extract_document_tab1", type="primary"):
            if uploaded_file is None:
                st.error("Por favor suba un archivo DOCX primero.")
                st.stop()
            
            with st.spinner("Extrayendo documento..."):
                try:
                    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp_file.write(uploaded_file.read())
                    tmp_file.close()
                    
                    progress_bar = st.progress(0, text="Leyendo y extrayendo contenido del DOCX...")
                    doc_result = docx2python(tmp_file.name)
                    
                    # Use enhanced extraction
                    df, tables_data, extraction_stats = extract_docx_structure_enhanced(tmp_file.name)
                    progress_bar.progress(0.2, text="Documento cargado. Procesando estructura...")
                    
                    # Extract sections
                    header_1_values = df['header_1'].dropna().unique()
                    llm_summary_rows = []
                    
                    for idx, header in enumerate(header_1_values):
                        section_df = df[df['header_1'] == header].copy()
                        full_text = '\n'.join(section_df['content'].astype(str).tolist()).strip()
                        section_words = len(full_text.split())
                        section_paras = len(section_df[section_df['source_type'] == 'paragraph'])
                        section_tables = len(section_df[section_df['source_type'] == 'table'])
                        
                        llm_summary_rows.append({
                            'header_1': header,
                            'llm_paragraph': full_text if full_text else "",
                            'n_words': section_words,
                            'n_paragraphs': section_paras,
                            'n_tables': section_tables
                        })

                    progress_bar.progress(0.5, text="Secciones extraídas.")
                    
                    # Create exploded dataframe
                    llm_summary_df = pd.DataFrame(llm_summary_rows)
                    exploded_df = llm_summary_df.assign(
                        llm_paragraph=llm_summary_df['llm_paragraph'].str.split('\n')
                    ).explode('llm_paragraph')
                    exploded_df = exploded_df.reset_index(drop=True)
                    exploded_df = exploded_df[exploded_df['llm_paragraph'].str.strip() != '']
                    
                    # Get full text
                    full_document_text = "\n\n".join(exploded_df['llm_paragraph'].tolist())
                    
                    # Store in session state
                    file_size = os.path.getsize(tmp_file.name)
                    n_words = exploded_df['llm_paragraph'].str.split().str.len().sum()
                    n_paragraphs = len(exploded_df)
                    
                    st.session_state['full_document_text_tab1'] = full_document_text
                    st.session_state['appraisal_document_stats'] = {
                        'file_size': file_size,
                        'word_count': n_words,
                        'n_words': n_words
                    }
                    st.session_state['exploded_df_tab1'] = exploded_df
                    st.session_state['extraction_df_tab1'] = df
                    st.session_state['tables_data_tab1'] = tables_data
                    st.session_state['extraction_stats_tab1'] = extraction_stats
                    st.session_state['sections_df_tab1'] = llm_summary_df
                    st.session_state['selected_sections_tab1'] = list(header_1_values)  # Select all by default
                    st.session_state['last_file_hash_tab1'] = file_hash
                    st.session_state['document_extracted_tab1'] = True
                    
                    try:
                        os.unlink(tmp_file.name)
                    except:
                        pass
                    
                    progress_bar.progress(1.0, text="Extracción completa.")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error procesando el documento: {e}")
                    import traceback
                    st.error(traceback.format_exc())
                    st.stop()
        
        # Show extraction results if document is extracted
        if st.session_state.get('document_extracted_tab1', False) and not file_changed:
            st.success("✅ Documento extraído con éxito")
            
            sections_df = st.session_state.get('sections_df_tab1', pd.DataFrame())
            
            if not sections_df.empty:
                header_1_values = sections_df['header_1'].tolist()
                
                # Section selector
                st.markdown("### 🔍 Selección de Secciones para Evaluación")
                st.info("Selecciona las secciones que deseas incluir en la evaluación. Por defecto, todas las secciones están seleccionadas.")
                
                # Guidance about section extraction
                if len(header_1_values) == 0:
                    st.error("⚠️ **No se detectaron secciones en el documento.** Esto puede deberse a que el documento no usa estilos de encabezado de Word (Heading 1, Heading 2, etc.). Por favor, verifica que tu documento tenga encabezados formateados correctamente.")
                elif len(header_1_values) < 3:
                    st.warning("⚠️ **Se detectaron pocas secciones** en el documento. Si esperabas más secciones, verifica que el documento use estilos de encabezado de Word (Heading 1, Heading 2, etc.) para identificar las secciones principales.")
                
                # Initialize selected sections if not exists
                if 'selected_sections_tab1' not in st.session_state:
                    st.session_state['selected_sections_tab1'] = list(header_1_values)
                
                # Section selection interface
                selected_sections = st.session_state.get('selected_sections_tab1', list(header_1_values)).copy()
                col1, col2 = st.columns([3, 1])
                
                with col2:
                    if st.button("✅ Seleccionar Todas", key="select_all_sections_tab1"):
                        st.session_state['selected_sections_tab1'] = list(header_1_values)
                        st.rerun()
                    
                    if st.button("❌ Deseleccionar Todas", key="deselect_all_sections_tab1"):
                        st.session_state['selected_sections_tab1'] = []
                        st.rerun()
                
                with col1:
                    st.markdown("**Secciones disponibles:**")
                    for section in header_1_values:
                        section_info = sections_df[sections_df['header_1'] == section].iloc[0]
                        is_selected = section in selected_sections
                        
                        checkbox_label = f"**{section}** ({section_info['n_words']:,} palabras, {section_info['n_paragraphs']} párrafos)"
                        
                        checkbox_key = f"section_checkbox_{section}_tab1"
                        new_selection = st.checkbox(checkbox_label, value=is_selected, key=checkbox_key)
                        
                        if new_selection and section not in selected_sections:
                            selected_sections.append(section)
                        elif not new_selection and section in selected_sections:
                            selected_sections.remove(section)
                        
                        # Add expandable preview of extracted content
                        with st.expander(f"👁️ Ver contenido: {section}", expanded=False):
                            section_content = section_info['llm_paragraph']
                            if section_content and section_content.strip():
                                st.text_area(
                                    "Contenido extraído:",
                                    value=section_content,
                                    height=200,
                                    key=f"content_preview_{section}_tab1",
                                    label_visibility="collapsed"
                                )
                                st.caption(f"Total: {len(section_content):,} caracteres")
                            else:
                                st.info("Esta sección no tiene contenido extraído.")
                            
                            # Show table-extracted text
                            tables_data = st.session_state.get('tables_data_tab1', [])
                            section_tables = [t for t in tables_data if t.get('section') == section]
                            
                            if section_tables:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                for table_info in section_tables:
                                    table_num = table_info.get('table_number', 'N/A')
                                    table_data = table_info.get('data', [])
                                    if table_data:
                                        # Format table as text
                                        table_text = '\n'.join([' | '.join(str(cell) for cell in row) for row in table_data])
                                        st.text_area(
                                            f"Tabla {table_num}:",
                                            value=table_text,
                                            height=150,
                                            key=f"table_preview_{section}_table{table_num}_tab1",
                                            label_visibility="collapsed"
                                        )
                                        st.caption(f"Tabla {table_num}: {len(table_data)} filas, {len(table_data[0]) if table_data else 0} columnas")
                            else:
                                st.markdown("---")
                                st.markdown("#### 📊 Texto extraído desde tablas")
                                st.info("No se encontraron tablas en esta sección.")
                
                # Update session state
                st.session_state['selected_sections_tab1'] = selected_sections
                
                # Show selection summary
                if selected_sections:
                    selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
                    total_selected_words = selected_df['n_words'].sum()
                    total_selected_paras = selected_df['n_paragraphs'].sum()
                    
                    # Estimate tokens
                    if encoding:
                        selected_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
                        estimated_tokens = len(encoding.encode(selected_text))
                    else:
                        estimated_tokens = total_selected_words * 1.2
                    
                    st.success(f"✅ {len(selected_sections)} secciones seleccionadas | "
                              f"{total_selected_words:,} palabras | "
                              f"~{estimated_tokens:,} tokens estimados")
    
    # Filter section for questions (optional, for filtering which questions to analyze)
    st.markdown("---")
    st.subheader("🔍 Filtros de Preguntas (opcional)")
    st.info("Selecciona secciones y subsecciones de preguntas para un análisis enfocado. Deja en blanco para analizar todas las preguntas.")
    
    # Get unique sections and subsections from questions
    df_with_sections = df_appraisal.copy()
    df_with_sections['_section'] = df_with_sections['Pregunta_Realizada'].apply(extract_section_number)
    df_with_sections['_subsection'] = df_with_sections['Pregunta_Realizada'].apply(extract_subsection_number)
    
    # Section Names Mapping
    SECTION_NAMES = {
        1: "Pertinencia",
        2: "Apropiación y sostenibilidad",
        3: "Gestión orientada a los resultados",
        4: "Transparencia y rendición de cuentas",
        5: "Presentación de la propuesta"
    }

    all_sections = sorted(df_with_sections['_section'].dropna().unique())
    
    # Create two columns for filters
    col_filter1, col_filter2 = st.columns(2)
    
    with col_filter1:
        # Format section options with names
        section_options = [f"{int(s)}. {SECTION_NAMES.get(int(s), 'Sección ' + str(int(s)))}" for s in all_sections]
        
        selected_sections_filter = st.multiselect(
            "Selecciona Secciones:",
            options=section_options,
            key="filter_sections_tab1",
            help="Selecciona una o más secciones para analizar"
        )
    
    # Extract section numbers from filter selections
    filtered_section_nums = []
    if selected_sections_filter:
        for s in selected_sections_filter:
            try:
                # Extract number from "1. Name" format
                filtered_section_nums.append(int(s.split('.')[0]))
            except:
                pass

    # Filter available subsections based on selected sections
    if filtered_section_nums:
        # Show only subsections belonging to selected sections
        available_subsections_df = df_with_sections[df_with_sections['_section'].isin(filtered_section_nums)]
    else:
        # Show all subsections if no section is selected
        available_subsections_df = df_with_sections
        
    all_subsections = sorted(available_subsections_df['_subsection'].dropna().unique(), 
                            key=lambda x: parse_subsection_for_sorting(x) if pd.notna(x) else (999, 999))

    with col_filter2:
        selected_subsections_filter = st.multiselect(
            "Selecciona Subsecciones:",
            options=[f"Subsección {s}" for s in all_subsections],
            key="filter_subsections_tab1",
            help="Selecciona una o más subsecciones para analizar"
        )
    
    # Extract subsection IDs
    filtered_subsection_ids = [s.replace("Subsección ", "") for s in selected_subsections_filter] if selected_subsections_filter else None
    
    # Processing button
    st.markdown("---")
    st.markdown("### ⚙️ Procesamiento y Evaluación")
    if st.button('🔍 Analizar documento', key="appraisal_process_button", type="primary"):
        # Check prerequisites
        if not st.session_state.get('document_extracted_tab1', False):
            st.error("❌ Por favor extrae el documento primero usando el botón 'Extraer Documento'.")
            st.stop()
        
        if uploaded_file is None:
            st.warning("⚠️ Por favor suba un archivo DOCX primero.")
            st.stop()
        
        # Get selected sections or use full document
        selected_sections = st.session_state.get('selected_sections_tab1', [])
        sections_df = st.session_state.get('sections_df_tab1', pd.DataFrame())
        
        if selected_sections and not sections_df.empty:
            # Filter to selected sections only
            selected_df = sections_df[sections_df['header_1'].isin(selected_sections)]
            document_text = "\n\n".join(selected_df['llm_paragraph'].tolist())
            st.info(f"📌 Evaluando {len(selected_sections)} secciones seleccionadas")
        else:
            # Fallback to full document
            document_text = st.session_state.get('full_document_text_tab1', '')
            if not selected_sections:
                st.warning("⚠️ No hay secciones seleccionadas. Usando documento completo.")
        
        if not document_text:
            st.error("No se pudo recuperar el texto del documento.")
            st.stop()
        
        # Get questions for analysis
        questions = df_appraisal['Pregunta_Realizada'].dropna().unique().tolist()
        
        # Apply filters if selected
        if filtered_section_nums or filtered_subsection_ids:
            questions_df_temp = df_appraisal[df_appraisal['Pregunta_Realizada'].isin(questions)].copy()
            questions_df_temp['_section'] = questions_df_temp['Pregunta_Realizada'].apply(extract_section_number)
            questions_df_temp['_subsection'] = questions_df_temp['Pregunta_Realizada'].apply(extract_subsection_number)
            
            # Filter by selected sections and/or subsections
            if filtered_section_nums and filtered_subsection_ids:
                # Both filters selected - apply AND logic
                filtered_questions = questions_df_temp[
                    (questions_df_temp['_section'].isin(filtered_section_nums)) |
                    (questions_df_temp['_subsection'].isin(filtered_subsection_ids))
                ]
            elif filtered_section_nums:
                # Only sections selected
                filtered_questions = questions_df_temp[questions_df_temp['_section'].isin(filtered_section_nums)]
            else:
                # Only subsections selected
                filtered_questions = questions_df_temp[questions_df_temp['_subsection'].isin(filtered_subsection_ids)]
            
            questions = filtered_questions['Pregunta_Realizada'].dropna().unique().tolist()
            
            # Show filter info
            st.info(f"📌 Análisis filtrado: {len(questions)} preguntas seleccionadas de {len(df_appraisal)}")
        
        if not questions:
            st.error("❌ No se encontraron preguntas para el análisis con los filtros aplicados.")
            st.stop()
        
        # Analyze questions
        st.markdown("### 🔍 Progreso del análisis")
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        results = []
        critical_opinions = {}
        
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            # Submit all questions for document-based analysis first
            # TAB1: Using tab1-specific function that explicitly states when 2 parts are detected
            future_to_question_doc = {
                executor.submit(analyze_question_with_llm_tab1, q, document_text): q
                for q in questions
            }

            # Process document-based analyses
            completed = 0
            for future in as_completed(future_to_question_doc):
                question = future_to_question_doc[future]
                result = future.result()
                results.append(result)
                completed += 1

                # Update progress
                progress = completed / len(questions)
                progress_bar.progress(progress * 0.5)  # First half of progress bar
                status_text.text(f"Análisis documentales: {completed}/{len(questions)} preguntas")

            # Create results DataFrame to get answers for critical evaluation
            results_df_temp = pd.DataFrame(results)

            # Now submit critical evaluations with answer, reasoning, evidence AND document context
            # TAB1: Using tab1-specific function that checks for proper two-part acknowledgment
            future_to_question_critical = {
                executor.submit(
                    analyze_question_with_critical_opinion_tab1,
                    row['Pregunta'],
                    row['Respuesta'],
                    row['Razonamiento'],
                    row['Evidencia'],
                    document_text  # Pass selected sections text for context
                ): row['Pregunta']
                for _, row in results_df_temp.iterrows()
            }
            
            # Process critical evaluations
            completed = 0
            for future in as_completed(future_to_question_critical):
                question = future_to_question_critical[future]
                critical_opinion = future.result()
                critical_opinions[question] = critical_opinion
                completed += 1
                
                # Update progress
                progress = completed / len(questions)
                progress_bar.progress(0.5 + progress * 0.5)  # Second half of progress bar
                status_text.text(f"Evaluaciones críticas: {completed}/{len(questions)} preguntas")
        
        # Create results DataFrame - sort by subsection for proper ordering
        results_df = pd.DataFrame(results)
        
        # Add critical opinions to dataframe
        results_df['Evaluación Crítica'] = results_df['Pregunta'].map(critical_opinions).fillna("No disponible")
        results_df['_section_num'] = results_df['Pregunta'].apply(extract_section_number)
        results_df['_subsection'] = results_df['Pregunta'].apply(extract_subsection_number)
        results_df['_sort_key'] = results_df['_subsection'].apply(parse_subsection_for_sorting)
        
        # Sort results by section and subsection
        results_df = results_df.sort_values(by=['_sort_key']).reset_index(drop=True)
        
        # THREE-LEVEL SYNTHESIS: Question -> Subsection -> Section
        st.markdown("### 📈 Síntesis Multinivel")
        
        # Level 1: Subsection synthesis (grouping individual questions)
        st.info("Generando análisis a nivel de subsección...")
        subsections = sorted(results_df['_subsection'].dropna().unique(), key=parse_subsection_for_sorting)
        subsection_analyses = {}
        
        subsection_progress = st.progress(0)
        for idx, subsection_id in enumerate(subsections):
            subsection_df = results_df[results_df['_subsection'] == subsection_id].copy()
            
            # Generate subsection analysis from individual questions
            subsection_analysis = synthesize_subsection_analysis(
                subsection_id,
                subsection_df[['Pregunta', 'Respuesta', 'Razonamiento']]
            )
            subsection_analyses[subsection_id] = subsection_analysis
            
            subsection_progress.progress((idx + 1) / len(subsections))
        
        # Level 2: Critical evaluation synthesis at subsection level
        st.info("Generando evaluaciones críticas a nivel de subsección...")
        subsection_critical_evaluations = {}
        
        critical_subsection_progress = st.progress(0)
        for idx, subsection_id in enumerate(subsections):
            subsection_df = results_df[results_df['_subsection'] == subsection_id].copy()
            
            # Generate critical evaluation synthesis from individual critical opinions
            critical_evaluation = synthesize_critical_evaluation_subsection(
                subsection_id,
                subsection_df[['Pregunta', 'Evaluación Crítica']]
            )
            subsection_critical_evaluations[subsection_id] = critical_evaluation
            
            critical_subsection_progress.progress((idx + 1) / len(subsections))
        
        # Level 3: Section synthesis (grouping subsections)
        st.info("Generando análisis a nivel de sección...")
        sections = sorted(results_df['_section_num'].dropna().unique())
        section_analyses = {}
        
        section_progress = st.progress(0)
        for idx, section_num in enumerate(sections):
            section_num = int(section_num)
            # Get all subsections for this section
            section_subsections = {
                subsec_id: subsection_analyses[subsec_id]
                for subsec_id in subsections
                if subsec_id.startswith(f"{section_num}.")
            }
            
            # Generate section analysis from subsection analyses
            section_analysis = synthesize_section_analysis(
                section_num,
                section_subsections
            )
            section_analyses[section_num] = section_analysis
            
            section_progress.progress((idx + 1) / len(sections))
        
        # Level 4: Critical evaluation synthesis at section level
        st.info("Generando evaluaciones críticas a nivel de sección...")
        section_critical_evaluations = {}
        
        critical_section_progress = st.progress(0)
        for idx, section_num in enumerate(sections):
            section_num = int(section_num)
            # Get all critical subsection evaluations for this section
            section_critical_subsections = {
                subsec_id: subsection_critical_evaluations[subsec_id]
                for subsec_id in subsections
                if subsec_id.startswith(f"{section_num}.")
            }
            
            # Generate critical evaluation from critical subsection evaluations
            critical_evaluation = synthesize_critical_evaluation_section(
                section_num,
                section_critical_subsections
            )
            section_critical_evaluations[section_num] = critical_evaluation
            
            critical_section_progress.progress((idx + 1) / len(sections))
        
        # Store all results in session state
        st.session_state['tab1_results_df'] = results_df
        st.session_state['tab1_subsection_analyses'] = subsection_analyses
        st.session_state['tab1_subsection_critical_evaluations'] = subsection_critical_evaluations
        st.session_state['tab1_section_analyses'] = section_analyses
        st.session_state['tab1_section_critical_evaluations'] = section_critical_evaluations

        # Get document stats from session state
        doc_stats = st.session_state.get('appraisal_document_stats', {})
        st.session_state['tab1_doc_stats'] = {
            'file_size': doc_stats.get('file_size', 0),
            'word_count': doc_stats.get('word_count', 0)
        }
        
        # Display results
        st.markdown("### 📊 Resultados del análisis")
        
        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total de preguntas", len(results_df))
        with col2:
            success_count = len(results_df[results_df['Status'] == 'Success'])
            st.metric("Exitosas", success_count)
        with col3:
            error_count = len(results_df[results_df['Status'] == 'Error'])
            st.metric("Errores", error_count)
        with col4:
            if success_count > 0:
                yes_count = len(results_df[results_df['Respuesta'] == 'Yes'])
                st.metric("Respuestas 'Sí'", yes_count)
        
        # Results table
        st.markdown("#### 📋 Resultados detallados")
        
        # Filter options
        col1, col2 = st.columns([1, 1])
        with col1:
            response_filter = st.selectbox(
                "Filtrar por respuesta:",
                ['Todas'] + results_df['Respuesta'].unique().tolist(),
                key="response_filter"
            )
        
        # Apply filter
        filtered_df = results_df.copy()
        if response_filter != 'Todas':
            filtered_df = filtered_df[filtered_df['Respuesta'] == response_filter]
        
        # Display filtered results
        st.dataframe(
            filtered_df[['Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia']], 
            use_container_width=True,
            height=400
        )
        
        # Download section
        st.markdown("### 📥 Descargar resultados")
        
        if len(results_df) > 0:
            # Get analyses from session state
            subsection_analyses = st.session_state.get('tab1_subsection_analyses', {})
            subsection_critical_evaluations = st.session_state.get('tab1_subsection_critical_evaluations', {})
            section_analyses = st.session_state.get('tab1_section_analyses', {})
            section_critical_evaluations = st.session_state.get('tab1_section_critical_evaluations', {})
            zip_buffer = create_results_download_with_sections(results_df, subsection_analyses, subsection_critical_evaluations, section_analyses, section_critical_evaluations)
            
            st.download_button(
                label="📦 Descargar resultados en ZIP",
                data=zip_buffer,
                file_name="appraisal_checklist_results.zip",
                mime="application/zip",
                key="appraisal_download_button"
            )
            
            st.success("✅ ¡Análisis completo! Usa el botón de descarga para obtener todos los resultados.")
        else:
            st.warning("⚠️ No hay resultados para descargar.")
    
    else:
        # Check if there are persisted results in session state
        if st.session_state.get('tab1_results_df') is not None:
            results_df = st.session_state['tab1_results_df']
            doc_stats = st.session_state.get('tab1_doc_stats', {})
            
            # Display persisted results
            st.markdown("### 📊 Resultados del análisis (guardados)")
            
            if doc_stats:
                st.info(f"""
                **Resumen del documento analizado:**
                - Tamaño del archivo: {doc_stats.get('file_size', 0)/1024:.2f} KB
                - Número de palabras: {doc_stats.get('word_count', 0):,}
                """)
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Total de preguntas", len(results_df))
            with col2:
                success_count = len(results_df[results_df['Status'] == 'Success'])
                st.metric("Exitosas", success_count)
            with col3:
                error_count = len(results_df[results_df['Status'] == 'Error'])
                st.metric("Errores", error_count)
            with col4:
                if success_count > 0:
                    yes_count = len(results_df[results_df['Respuesta'] == 'Yes'])
                    st.metric("Respuestas 'Sí'", yes_count)
            
            # Results table
            st.markdown("#### 📋 Resultados detallados")
            
            # Filter options
            col1, col2 = st.columns([1, 1])
            with col1:
                response_filter = st.selectbox(
                    "Filtrar por respuesta:",
                    ['Todas'] + results_df['Respuesta'].unique().tolist(),
                    key="response_filter_persisted"
                )
            
            # Apply filter
            filtered_df = results_df.copy()
            if response_filter != 'Todas':
                filtered_df = filtered_df[filtered_df['Respuesta'] == response_filter]
            
            # Display filtered results
            st.dataframe(
                filtered_df[['Pregunta', 'Respuesta', 'Razonamiento', 'Evidencia']], 
                use_container_width=True,
                height=400
            )
            
            # Download section
            st.markdown("### 📥 Descargar resultados")
            
            if len(results_df) > 0:
                # Get analyses from session state
                subsection_analyses = st.session_state.get('tab1_subsection_analyses', {})
                subsection_critical_evaluations = st.session_state.get('tab1_subsection_critical_evaluations', {})
                section_analyses = st.session_state.get('tab1_section_analyses', {})
                section_critical_evaluations = st.session_state.get('tab1_section_critical_evaluations', {})
                zip_buffer = create_results_download_with_sections(results_df, subsection_analyses, subsection_critical_evaluations, section_analyses, section_critical_evaluations)
                
                st.download_button(
                    label="📦 Descargar resultados en ZIP",
                    data=zip_buffer,
                    file_name="appraisal_checklist_results.zip",
                    mime="application/zip",
                    key="appraisal_download_button_persisted"
                )
            
            # Clear results button
            if st.button("🗑️ Limpiar resultados", key="clear_tab1_results"):
                st.session_state['tab1_results_df'] = None
                st.session_state['tab1_doc_stats'] = None
                st.rerun()
        else:
            if uploaded_file:
                st.info("👆 Haz clic en 'Analizar documento' para iniciar la evaluación de la valoración preliminar de calidad.")
            else:
                st.info("📁 Por favor sube un archivo DOCX para comenzar.")

#--------------------------#-------------------------------#
#--------------------------#-------------------------------#
# # Tab 8: Tablero de seguimiento de recomendaciones y planes de acción
# with tab5:
#     st.header("📊 Estadísticas sobre Recomendaciones de Evaluaciones y sus Planes de Acción")

#     # Comprehensive presentation box
#     st.info("""

#     En esta ventana se presentan gráficos y estadísticas derivados de las respuestas institucionales a las recomendaciones (incluidos sus planes de acción) ya procesadas por la herramienta. El panel permite seguir el estado de respuesta (completadas, parcialmente completadas, acción no planificada, acción no tomada, rechazadas, sin respuesta), analizar la calidad de la respuesta (coherencia con el plan, calidad del plan, nivel de atención), observar la evolución en el tiempo y la composición por país, año y dimensión (gobernanza, participación, género, transición justa, capacidades, sostenibilidad financiera, incidencia).

#     También muestra atributos de las recomendaciones (innovación, precisión/claridad, viabilidad, horizonte temporal, impacto esperado) e identifica barreras de implementación.

#     **🎯 Orientado al nivel directivo y unidades de programación para:**
#     - Conducir conversaciones de seguimiento: qué ocurrió después de decidir implementar (o no) cada recomendación
#     - Identificar por qué quedaron pendientes y qué ajustes corresponden
#     - Facilitar acuerdos operativos (responsables y plazos)
#     - Realizar reprogramaciones cuando haya cuellos de botella
#     - Escalar barreras críticas
#     - Preparar notas técnicas/minutas con evidencia gráfica para rendición de cuentas ante mandantes

#     **💡 En síntesis:** El tablero transforma la evidencia en prioridades de acción verificables y ciclos de aprendizaje para cerrar brechas y sostener avances.

#     **🔍 Análisis por similitud:** Use la opción "Por similitud" para estimar la correspondencia entre el núcleo de la recomendación y la respuesta/plan. Una similitud baja puede señalar un desajuste de alcance, actores o resultados esperados; tómelo como insumo para discutir ajustes y seguimiento con los equipos.

#     *Sugerencia: comience con un umbral de 0.70 y ajústelo según el contexto.*
#     """)

#     # Chat section for querying similar recommendations
#     st.header("Búsqueda de Recomendaciones")
#     st.markdown("### Búsqueda")

#     # Input for user query
#     user_query = st.text_input("Búsqueda en recomendaciones:", value="¿Qué aspectos deben mejorarse sobre coordinación con partes interesadas?", key='user_query_tab8')

#     # Search method selection
#     search_method = st.radio("Método de búsqueda:", ["Por Similitud", "Por Coincidencia de Términos"], key='search_method_tab8')

#     # Slider for similarity score threshold (only relevant for similarity search)
#     if search_method == "Por Similitud":
#         score_threshold = st.slider("Umbral de similitud:", min_value=0.0, max_value=1.0, value=0.5, step=0.01, key='score_threshold_tab8')

#     # Function to display results
#     def display_results(results):
#         if results:
#             st.markdown("#### Recomendaciones similares")
#             for i, result in enumerate(results):
#                 st.markdown(f"**Recomendación {i+1}:**")
#                 st.markdown(f"**Texto:** {result['recommendation']}")
#                 if "similarity" in result:
#                     st.markdown(f"**Puntuación de similitud:** {result['similarity']:.2f}")
#                 st.markdown(f"**País:** {result['country']}")
#                 st.markdown(f"**Año:** {result['year']}")
#                 st.markdown(f"**Número de evaluación:** {result['eval_id']}")
#                 st.markdown("---")
#         else:
#             st.write("No se encontraron recomendaciones para la búsqueda.")

#     # Button to search for recommendations
#     if st.button("Buscar Recomendaciones", key='search_button_tab8'):
#         if user_query:
#             with st.spinner('Buscando recomendaciones...'):
#                 if search_method == "Por Similitud":
#                     query_embedding = get_embedding_with_retry(user_query)
#                     if query_embedding is not None:
#                         results = find_similar_recommendations(query_embedding, index, doc_embeddings, structured_embeddings, score_threshold)
#                         display_results(results)
#                     else:
#                         st.error("No se pudo generar el embedding para la consulta.")
#                 else:
#                     results = find_recommendations_by_term_matching(user_query, doc_texts, structured_embeddings)
#                     display_results(results)

#     # Use the main recommendations dataset
#     filtered_df = df.copy()

#     # Check if this is the recommendations dataset (has required columns)
#     required_cols = ['Theme_cl', 'Recommendation_theme', 'Management_response']
#     has_required_cols = all(col in df.columns for col in required_cols)

#     if not has_required_cols:
#         st.warning("⚠️ Esta sección requiere datos de recomendaciones. Por favor, carga el archivo de recomendaciones en la configuración inicial.")
#     else:
#         # --- FILTER ROW WITHIN TAB ---
#         st.markdown("### Filtros")
#         filter_container = st.container()

#         with filter_container:
#             # Create 3 columns for filters
#             col1, col2, col3 = st.columns(3)

#             with col1:
#                 # # Administrative unit filter
#                 # office_options = ['Todas'] + sorted([str(x) for x in df['Recommendation_administrative_unit'].unique() if not pd.isna(x)])
#                 # selected_offices_viz = st.multiselect('Unidad Administrativa:',
#                 #                                  options=office_options,
#                 #                                  default='Todas',
#                 #                                  key='unidad_administrativa_viz')
    
#                 # Country filter
#                 country_col_candidates = ['Country(ies)', 'Country', 'Countries', 'country', 'Country (ies)', 'Country/ies']
#                 country_col = next((c for c in country_col_candidates if c in df.columns), None)
#                 if country_col:
#                     country_options = ['Todas'] + sorted([str(x) for x in df[country_col].unique() if not pd.isna(x)])
#                 else:
#                     st.warning("No se encontró una columna de país en los datos.")
#                     country_options = ['Todas']
#                 selected_countries_viz = st.multiselect('País:',
#                                                   options=country_options,
#                                                   default='Todas' if 'Todas' in country_options else [],
#                                                   key='pais_viz')
    
#             with col2:
#                 # Year filter with slider (only if year column exists)
#                 if 'year' in df.columns:
#                     min_year = int(df['year'].min())
#                     max_year = max(int(df['year'].max()), 2025)
#                     selected_year_range_viz = st.slider('Rango de Años:',
#                                                   min_value=min_year,
#                                                   max_value=max_year,
#                                                   value=(min_year, max_year),
#                                                   key='rango_anos_viz')
#                 else:
#                     selected_year_range_viz = None
    
#                 # Evaluation theme filter
#                 evaltheme_options = ['Todas'] + sorted([str(x) for x in df['Theme_cl'].unique() if not pd.isna(x)])
#                 selected_evaltheme_viz = st.multiselect('Tema (Evaluación):',
#                                                   options=evaltheme_options,
#                                                   default='Todas',
#                                                   key='tema_eval_viz')
    
#             with col3:
#                 # Recommendation theme filter
#                 rectheme_options = ['Todas'] + sorted([str(x) for x in df['Recommendation_theme'].unique() if not pd.isna(x)])
#                 selected_rectheme_viz = st.multiselect('Tema (Recomendación):',
#                                                  options=rectheme_options,
#                                                  default='Todas',
#                                                  key='tema_recomendacion_viz')
    
#                 # Management response filter
#                 mgtres_options = ['Todas'] + sorted([str(x) for x in df['Management_response'].unique() if not pd.isna(x)])
#                 selected_mgtres_viz = st.multiselect('Respuesta de gerencia:',
#                                                options=mgtres_options,
#                                                default='Todas',
#                                                key='respuesta_gerencia_viz')
    
#         # Apply filters
#         # if 'Todas' not in selected_offices_viz and selected_offices_viz:
#         #     filtered_df = filtered_df[filtered_df['Recommendation_administrative_unit'].isin(selected_offices_viz)]
    
#         if 'Todas' not in selected_countries_viz and selected_countries_viz and 'country_col' in locals() and country_col:
#             filtered_df = filtered_df[filtered_df[country_col].isin(selected_countries_viz)]
    
#         if selected_year_range_viz is not None and 'year' in filtered_df.columns:
#             filtered_df = filtered_df[
#                 (filtered_df['year'] >= selected_year_range_viz[0]) &
#                 (filtered_df['year'] <= selected_year_range_viz[1])
#             ]
    
#         if 'Todas' not in selected_evaltheme_viz and selected_evaltheme_viz:
#             filtered_df = filtered_df[filtered_df['Theme_cl'].isin(selected_evaltheme_viz)]
    
#         if 'Todas' not in selected_rectheme_viz and selected_rectheme_viz:
#             filtered_df = filtered_df[filtered_df['Recommendation_theme'].isin(selected_rectheme_viz)]
    
#         if 'Todas' not in selected_mgtres_viz and selected_mgtres_viz:
#             filtered_df = filtered_df[filtered_df['Management_response'].isin(selected_mgtres_viz)]
    
#         # Remove duplicates for plotting
#         filtered_df_unique = filtered_df.drop_duplicates(subset=['index_df'])
    
#         # --- KPIs FOR VISUALIZATION TAB ---
#         if not filtered_df_unique.empty:
#             # Main KPIs
#             total_kpi_labels = [
#                 "Total Recomendaciones",
#                 "Países",
#                 "Evaluaciones",
#                 "Años Cubiertos"
#             ]
#             total_kpi_values = [
#                 filtered_df_unique.shape[0],
#                 (filtered_df_unique[country_col].nunique() if 'country_col' in locals() and country_col and country_col in filtered_df_unique.columns else 0),
#                 filtered_df_unique['Evaluation_number'].nunique() if 'Evaluation_number' in filtered_df_unique.columns else filtered_df_unique['Evaluation number'].nunique() if 'Evaluation number' in filtered_df_unique.columns else 0,
#                 filtered_df_unique['year'].nunique()
#             ]
    
#             total_cols = st.columns(4)
#             total_kpi_html = [
#                 f"""
#                 <div style='text-align:center;'>
#                     <span style='font-size:1.2em; font-weight:700;'>{label}</span><br>
#                     <span style='font-size:2.3em; font-weight:700; color:#002F6C;'>{value}</span>
#                 </div>
#                 """
#                 for label, value in zip(total_kpi_labels, total_kpi_values)
#             ]
    
#             for col, html in zip(total_cols, total_kpi_html):
#                 col.markdown(html, unsafe_allow_html=True)
    
#             st.markdown("<hr style='border-top: 1px solid #e1e4e8;'>", unsafe_allow_html=True)
    
#             # KPIs for management response statuses
#             mgmt_labels = [
#                 ("Completadas", filtered_df_unique[filtered_df_unique['Management_response'] == 'Completed'].shape[0], '#27ae60'),
#                 ("Parcialmente Completadas", filtered_df_unique[filtered_df_unique['Management_response'] == 'Partially Completed'].shape[0], '#f7b731'),
#                 ("Acción no tomada aún", filtered_df_unique[filtered_df_unique['Management_response'] == 'Action not yet taken'].shape[0], '#fd9644'),
#                 ("Rechazadas", filtered_df_unique[filtered_df_unique['Management_response'] == 'Rejected'].shape[0], '#8854d0'),
#                 ("Acción no planificada", filtered_df_unique[filtered_df_unique['Management_response'] == 'No Action Planned'].shape[0], '#3867d6'),
#                 ("Sin respuesta", filtered_df_unique[filtered_df_unique['Management_response'] == 'Sin respuesta'].shape[0], '#eb3b5a'),
#             ]
    
#             st.markdown("<span style='font-size:1.6em; font-weight:700;'>Respuesta de Gerencia</span>", unsafe_allow_html=True)
#             kpi_cols = st.columns(3)
#             for i, (label, value, color) in enumerate(mgmt_labels):
#                 kpi_cols[i % 3].markdown(
#                     f"""
#                     <div style='text-align:center;'>
#                         <span style='font-size:1.2em; font-weight:700;'>{label}</span><br>
#                         <span style='font-size:2.3em; font-weight:700; color:{color};'>{value}</span>
#                     </div>
#                     """,
#                     unsafe_allow_html=True
#                 )
    
#             # Display plots if data is available
#             if not filtered_df.empty:
#                 country_counts = (
#                     filtered_df_unique[country_col].value_counts()
#                     if 'country_col' in locals() and country_col and country_col in filtered_df_unique.columns
#                     else pd.Series(dtype=int)
#                 )
    
#                 # Add CSS for dashboard styling
#                 st.markdown('<style>.dashboard-subtitle {font-size: 1.3rem; font-weight: 600; margin-bottom: 0.2em; margin-top: 1.2em; color: #002F6C;}</style>', unsafe_allow_html=True)
    
#                 # Create two columns for charts
#                 row1_col1, row1_col2 = st.columns(2)
    
#                 with row1_col1:
#                     st.markdown('<div class="dashboard-subtitle">Número de Recomendaciones por País</div>', unsafe_allow_html=True)
#                     fig1 = go.Figure()
#                     if not country_counts.empty:
#                         fig1.add_trace(go.Bar(
#                             y=country_counts.index.tolist(),
#                             x=country_counts.values.tolist(),
#                             orientation='h',
#                             text=country_counts.values.tolist(),
#                             textposition='auto',
#                             marker_color='#002F6C',
#                             hovertemplate='%{y}: %{x} recomendaciones'
#                         ))
#                     else:
#                         fig1.add_trace(go.Bar(y=[], x=[]))
    
#                     # Fixed height for alignment with year plot
#                     fixed_height = 500
#                     fig1.update_layout(
#                         xaxis_title='Número de Recomendaciones',
#                         yaxis_title='País',
#                         margin=dict(t=10, l=10, r=10, b=40),
#                         font=dict(size=22),
#                         height=fixed_height,
#                         plot_bgcolor='white',
#                         showlegend=False
#                     )
#                     fig1.update_xaxes(showgrid=True, gridcolor='LightGray')
#                     fig1.update_yaxes(showgrid=False)
#                     st.plotly_chart(fig1, use_container_width=True)
    
#                 with row1_col2:
#                     st.markdown('<div class="dashboard-subtitle">Número de Recomendaciones por Año</div>', unsafe_allow_html=True)
#                     year_counts = filtered_df_unique['year'].value_counts().sort_index()
#                     fig2 = go.Figure()
#                     fig2.add_trace(go.Bar(
#                         x=year_counts.index.astype(str).tolist(),
#                         y=year_counts.values.tolist(),
#                         text=year_counts.values.tolist(),
#                         textposition='auto',
#                         marker_color='#002F6C',
#                         hovertemplate='Año %{x}: %{y} recomendaciones',
#                         textfont=dict(size=22)
#                     ))
#                     fig2.update_layout(
#                         xaxis_title='Año',
#                         yaxis_title='Número de Recomendaciones',
#                         margin=dict(t=10, l=10, r=10, b=40),
#                         font=dict(size=22),
#                         height=500,
#                         plot_bgcolor='white',
#                         showlegend=False
#                     )
#                     fig2.update_xaxes(showgrid=True, gridcolor='LightGray', tickangle=45, title_font=dict(size=22), tickfont=dict(size=20))
#                     fig2.update_yaxes(showgrid=True, gridcolor='LightGray', title_font=dict(size=22), tickfont=dict(size=20))
#                     st.plotly_chart(fig2, use_container_width=True)
    
#                 # Dimension treemap
#                 st.markdown('<div class="dashboard-subtitle">Composición de Recomendaciones por Dimensión</div>', unsafe_allow_html=True)
    
#                 # Clean and prepare dimension data - use the same unique rows used for KPIs
#                 import numpy as np
#                 treemap_df = filtered_df_unique.copy()
#                 treemap_df['dimension'] = treemap_df['dimension'].astype(str).str.strip().str.lower().replace({
#                     'processes': 'process', 'process': 'process', 'nan': np.nan, 'none': np.nan, '': np.nan
#                 })
#                 treemap_df['dimension'] = treemap_df['dimension'].replace({'process': 'Process'})
#                 treemap_df = treemap_df[treemap_df['dimension'].notna()]

#                 # treemap_df['rec_intervention_approach'] = treemap_df['rec_intervention_approach'].astype(str).str.strip().str.lower().replace({
#                 #     'processes': 'process', 'process': 'process', 'nan': np.nan, 'none': np.nan, '': np.nan
#                 # })
#                 # treemap_df['rec_intervention_approach'] = treemap_df['rec_intervention_approach'].replace({'process': 'Process'})
#                 # treemap_df = treemap_df[treemap_df['rec_intervention_approach'].notna()]

#                 # Count recommendations by dimension
#                 dimension_counts = treemap_df.groupby('dimension').agg({
#                     'index_df': 'nunique'
#                 }).reset_index()
    
#                 # Calculate percentages and format text
#                 dimension_counts['percentage'] = dimension_counts['index_df'] / dimension_counts['index_df'].sum() * 100
#                 dimension_counts['text'] = dimension_counts.apply(
#                     lambda row: f"{row['dimension']}<br>Recomendaciones: {row['index_df']}<br>Porcentaje: {row['percentage']:.2f}%",
#                     axis=1
#                 )
#                 dimension_counts['font_size'] = dimension_counts['index_df'] / dimension_counts['index_df'].max() * 30 + 10  # Scale font size
    
#                 # Remove 'Sin Clasificar' and capitalize dimension labels
#                 dimension_counts = dimension_counts[dimension_counts['dimension'].str.lower() != 'sin clasificar']
#                 dimension_counts['dimension'] = dimension_counts['dimension'].astype(str).str.title()
    
#                 # Create treemap
#                 fig3 = px.treemap(
#                     dimension_counts,
#                     path=['dimension'],
#                     values='index_df',
#                     title='Composición de Recomendaciones por Dimensión',
#                     hover_data={'text': True, 'index_df': False, 'percentage': False},
#                     custom_data=['text']
#                 )
#                 fig3.update_traces(
#                     textinfo='label+value',
#                     hovertemplate='%{customdata[0]}',
#                     textfont_size=32
#                 )
#                 fig3.update_layout(
#                     margin=dict(t=50, l=25, r=25, b=25),
#                     width=900,
#                     height=500,
#                     title_font_size=32,
#                     font=dict(size=28),
#                     legend_font_size=28
#                 )
#                 st.plotly_chart(fig3, use_container_width=True)
    
#                 # Subdimension treemap - use treemap_df which was already cleaned above
#                 # Harmonize process/processes before plotting subdimensions
#                 treemap_df['dimension'] = treemap_df['dimension'].replace({'processes': 'Process', 'process': 'Process', 'Process': 'Process'})

#                 # Remove 'Sin Clasificar' from both dimension and subdimension
#                 treemap_df = treemap_df[treemap_df['dimension'].str.lower() != 'sin clasificar']
#                 treemap_df = treemap_df[treemap_df['subdim'].str.lower() != 'sin clasificar']

#                 # Capitalize dimension and subdimension labels
#                 treemap_df['dimension'] = treemap_df['dimension'].astype(str).str.title()
#                 treemap_df['subdim'] = treemap_df['subdim'].astype(str).str.title()

#                 # Count by subdimension
#                 subdimension_counts = treemap_df.groupby(['dimension', 'subdim']).agg({
#                     'index_df': 'nunique'
#                 }).reset_index()
    
#                 # Calculate percentages and format text
#                 subdimension_counts['percentage'] = subdimension_counts['index_df'] / subdimension_counts['index_df'].sum() * 100
#                 subdimension_counts['text'] = subdimension_counts.apply(
#                     lambda row: f"{row['subdim']}<br>Recomendaciones: {row['index_df']}<br>Porcentaje: {row['percentage']:.2f}%",
#                     axis=1
#                 )
#                 subdimension_counts['font_size'] = subdimension_counts['index_df'] / subdimension_counts['index_df'].max() * 30 + 10
    
#                 # Create treemap
#                 fig4 = px.treemap(
#                     subdimension_counts,
#                     path=['dimension', 'subdim'],
#                     values='index_df',
#                     title='Composición de Recomendaciones por Subdimensión',
#                     hover_data={'text': True, 'index_df': False, 'percentage': False},
#                     custom_data=['text']
#                 )
#                 fig4.update_traces(
#                     textinfo='label+value',
#                     hovertemplate='%{customdata[0]}',
#                     textfont_size=32
#                 )
#                 fig4.update_layout(
#                     margin=dict(t=50, l=25, r=25, b=25),
#                     width=900,
#                     height=500,
#                     title_font_size=32,
#                     font=dict(size=28),
#                     legend_font_size=28
#                 )

#                 st.plotly_chart(fig4, use_container_width=True)
    
#                 # Add the advanced visualization section
#                 add_advanced_visualization_section(filtered_df, tab_id="tab8")
#             else:
#                 st.warning("No hay datos disponibles para los filtros seleccionados.")
#         else:
#             st.warning("No hay datos disponibles para los filtros seleccionados.")

#--------------------------#-------------------------------#
#--------------------------#-------------------------------#
# Tab 6: Clasificación de Recomendaciones (World)
with tab5:
    st.header("Clasificación de Recomendaciones (World)")
    st.info("""
    **🌍 Clasificación Inteligente de Recomendaciones**

    Esta herramienta permite clasificar automáticamente las recomendaciones contenidas en el archivo `Recommendations_World.xlsx` 
    asignándolas a las subdimensiones definidas en el marco de trabajo `Frame_Recommendations_English.xlsx`.

    **Optimizaciones:**
    1.  **Filtrado Previo:** Filtra por región/país antes de procesar para ahorrar costos y tiempo.
    2.  **Caché Persistente:** Los embeddings se guardan localmente para no regenerarlos en futuras ejecuciones.
    3.  **Deduplicación:** Textos idénticos se procesan una sola vez.

    **Funcionamiento:**
    1.  Carga los datos de recomendaciones y el marco de referencia.
    2.  Calcula la similitud semántica (usando OpenAI Embeddings) entre cada recomendación y las definiciones de subdimensiones.
    3.  Asigna la subdimensión más relevante.
    4.  Visualiza la distribución mediante Treemaps interactivos.
    """)

    # --- Persistent Cache Class ---
    class EmbeddingsCache:
        def __init__(self, cache_file="embeddings_cache.pkl"):
            self.cache_file = cache_file
            self.cache = {}
            self.load_cache()

        def load_cache(self):
            if os.path.exists(self.cache_file):
                try:
                    with open(self.cache_file, 'rb') as f:
                        self.cache = pickle.load(f)
                except Exception:
                    self.cache = {}

        def save_cache(self):
            try:
                with open(self.cache_file, 'wb') as f:
                    pickle.dump(self.cache, f)
            except Exception:
                pass

        def get(self, text):
            return self.cache.get(text)

        def set(self, text, embedding):
            self.cache[text] = embedding

        def save(self):
            self.save_cache()

    # Initialize cache in session state if not exists
    if 'embeddings_cache_obj' not in st.session_state:
        st.session_state['embeddings_cache_obj'] = EmbeddingsCache()

    # --- Load Data ---
    @st.cache_data
    def load_world_data():
        try:
            # Check if files exist
            if not os.path.exists('Recommendations_World.xlsx') or not os.path.exists('Frame_Recommendations_English.xlsx'):
                return None, None
            
            rec_df = pd.read_excel('Recommendations_World.xlsx')
            frame_df = pd.read_excel('Frame_Recommendations_English.xlsx')
            return rec_df, frame_df
        except Exception as e:
            st.error(f"Error cargando los archivos de datos: {e}")
            return None, None

    
    # --- Data Source Selection ---
    st.markdown("### Fuente de Datos")
    data_source = st.radio("Selecciona origen de las recomendaciones:", 
                          ["Archivos Predeterminados (World)", "Cargar Archivo Propio (.xlsx)"], 
                          horizontal=True,
                          key="data_source_radio")
    
    rec_df_world = None
    frame_df_world = None
    
    # helper to just get frame (reusing cached function for now)
    _, frame_default = load_world_data()
    frame_df_world = frame_default

    if data_source == "Archivos Predeterminados (World)":
        rec_default, _ = load_world_data()
        rec_df_world = rec_default
        
        if rec_df_world is None:
             st.warning("⚠️ No se encontró el archivo 'Recommendations_World.xlsx' por defecto.")

    else:
        uploaded_file = st.file_uploader("Sube tu archivo de recomendaciones (Excel):", type=["xlsx"], key="custom_rec_upload")
        if uploaded_file:
            try:
                rec_df_world = pd.read_excel(uploaded_file)
                # Validation
                if 'Recommendation description' not in rec_df_world.columns:
                    st.error("❌ El archivo debe contener una columna llamada 'Recommendation description'.")
                    rec_df_world = None
                else:
                    st.success(f"Archivo cargado correctamente: {len(rec_df_world)} filas.")
            except Exception as e:
                st.error(f"Error al leer el archivo: {e}")
        else:
            st.info("Esperando archivo...")

    if rec_df_world is None or frame_df_world is None:
        if frame_df_world is None:
             st.warning("⚠️ No se encontró el archivo 'Frame_Recommendations_English.xlsx'. es necesario para la clasificación.")
    else:
        
        # --- DATA PREP: Extract Year ---
        if 'Recommendation date' in rec_df_world.columns:
            # Try parsing with existing format or coerse
            rec_df_world['Recommendation date'] = pd.to_datetime(rec_df_world['Recommendation date'], errors='coerce')
            rec_df_world['Year'] = rec_df_world['Recommendation date'].dt.year
            # Fill NaNs with a placeholder if needed, or leave as NaN
            # rec_df_world['Year'] = rec_df_world['Year'].fillna(0).astype(int)
        
        # --- Filters (PRE-PROCESSING) ---
        st.markdown("### 1. Filtros de Selección")
        st.caption("Selecciona los filtros para definir el subconjunto de datos a clasificar. Esto optimiza el tiempo y costo de procesamiento.")
        
        # Initialize filter selection dictionaries
        filters = {}
        
        # --- Group 1: Ubicación y Tiempo ---
        with st.expander("📍 Ubicación y Tiempo", expanded=True):
            col_grp1_1, col_grp1_2 = st.columns(2)
            
            with col_grp1_1:
                # Region Filter
                regions = sorted([str(x) for x in rec_df_world['Region(s)'].unique() if not pd.isna(x)]) if 'Region(s)' in rec_df_world.columns else []
                filters['Region(s)'] = st.multiselect("Región(es):", options=regions, default=[], key="tab6_region_filter")
                
                # Admin Unit
                if 'Recommendation administrative unit' in rec_df_world.columns:
                    admin_units = sorted([str(x) for x in rec_df_world['Recommendation administrative unit'].unique() if not pd.isna(x)])
                    filters['Recommendation administrative unit'] = st.multiselect("Unidad Administrativa:", options=admin_units, default=[], key="tab6_admin_unit")
            
            with col_grp1_2:
                # Country Filter (Cascading)
                available_countries_df = rec_df_world.copy()
                if filters['Region(s)']:
                     if 'Region(s)' in available_countries_df.columns:
                        available_countries_df = available_countries_df[available_countries_df['Region(s)'].isin(filters['Region(s)'])]
                
                countries = sorted([str(x) for x in available_countries_df['Country(ies)'].unique() if not pd.isna(x)]) if 'Country(ies)' in available_countries_df.columns else []
                filters['Country(ies)'] = st.multiselect("País(es):", options=countries, default=[], key="tab6_country_filter")
                
                # Year Filter
                if 'Year' in rec_df_world.columns:
                    years = sorted([int(x) for x in rec_df_world['Year'].unique() if not pd.isna(x)])
                    filters['Year'] = st.multiselect("Año:", options=years, default=[], key="tab6_year_filter")

        # --- Group 2: Temática y Técnica ---
        with st.expander("📚 Temática y Técnica", expanded=False):
            col_grp2_1, col_grp2_2 = st.columns(2)
            
            with col_grp2_1:
                if 'Evaluation theme(s)' in rec_df_world.columns:
                    themes = sorted([str(x) for x in rec_df_world['Evaluation theme(s)'].unique() if not pd.isna(x)])
                    filters['Evaluation theme(s)'] = st.multiselect("Temática de Evaluación:", options=themes, default=[], key="tab6_eval_theme")
                
                if 'Recommendation theme' in rec_df_world.columns:
                    rec_themes = sorted([str(x) for x in rec_df_world['Recommendation theme'].unique() if not pd.isna(x)])
                    filters['Recommendation theme'] = st.multiselect("Temática de Recomendación:", options=rec_themes, default=[], key="tab6_rec_theme")

                if 'Technical unit(s)' in rec_df_world.columns:
                    tech_units = sorted([str(x) for x in rec_df_world['Technical unit(s)'].unique() if not pd.isna(x)])
                    filters['Technical unit(s)'] = st.multiselect("Unidad Técnica:", options=tech_units, default=[], key="tab6_tech_unit")

            with col_grp2_2:
                if 'Funding source(s)' in rec_df_world.columns:
                    fundings = sorted([str(x) for x in rec_df_world['Funding source(s)'].unique() if not pd.isna(x)])
                    filters['Funding source(s)'] = st.multiselect("Fuente de Financiamiento:", options=fundings, default=[], key="tab6_funding")
                
                if 'Evaluation nature' in rec_df_world.columns:
                    natures = sorted([str(x) for x in rec_df_world['Evaluation nature'].unique() if not pd.isna(x)])
                    filters['Evaluation nature'] = st.multiselect("Naturaleza de Evaluación:", options=natures, default=[], key="tab6_eval_nature")
                
                if 'Evaluation type' in rec_df_world.columns:
                    types = sorted([str(x) for x in rec_df_world['Evaluation type'].unique() if not pd.isna(x)])
                    filters['Evaluation type'] = st.multiselect("Tipo de Evaluación:", options=types, default=[], key="tab6_eval_type")

        # --- Group 3: Gestión y Respuesta ---
        with st.expander("⚙️ Gestión y Respuesta", expanded=False):
             col_grp3_1, col_grp3_2 = st.columns(2)
             
             with col_grp3_1:
                 if 'Evaluation timing' in rec_df_world.columns:
                    timings = sorted([str(x) for x in rec_df_world['Evaluation timing'].unique() if not pd.isna(x)])
                    filters['Evaluation timing'] = st.multiselect("Momento de Evaluación:", options=timings, default=[], key="tab6_eval_timing")

                 if 'Progress' in rec_df_world.columns:
                    progresses = sorted([str(x) for x in rec_df_world['Progress'].unique() if not pd.isna(x)])
                    filters['Progress'] = st.multiselect("Progreso:", options=progresses, default=[], key="tab6_progress")

             with col_grp3_2:
                 if 'Management response' in rec_df_world.columns:
                     responses = sorted([str(x) for x in rec_df_world['Management response'].unique() if not pd.isna(x)])
                     filters['Management response'] = st.multiselect("Respuesta de Administración:", options=responses, default=[], key="tab6_mgmt_response")
                 
                 # Optional: Evaluation document type if needed
                 if 'Evaluation document type' in rec_df_world.columns:
                     doc_types = sorted([str(x) for x in rec_df_world['Evaluation document type'].unique() if not pd.isna(x)])
                     filters['Evaluation document type'] = st.multiselect("Tipo de Documento:", options=doc_types, default=[], key="tab6_doc_type")


        # Apply ALL filters
        start_count = len(rec_df_world)
        filtered_rec_df = rec_df_world.copy()
        
        for col_name, selected_values in filters.items():
            if selected_values and col_name in filtered_rec_df.columns:
                # Handle Year specifically if it's float/int comparison issues, but isin handles standard types well
                filtered_rec_df = filtered_rec_df[filtered_rec_df[col_name].isin(selected_values)]

            
        end_count = len(filtered_rec_df)
        
        st.markdown(f"**Registros seleccionados:** {end_count} de {start_count}")

        # --- Classification Logic ---
        
        # Persistent state for classified data
        if 'classified_world_df' not in st.session_state:
            st.session_state['classified_world_df'] = None

        def classify_recommendations(target_df, reference_df, cache_obj, use_llm_verification=False):
            """
            Classifies recommendations in target_df based on similarity to texts in reference_df.
            Uses persistent cache and deduplication.
            """
            
            progress_bar = st.progress(0, text="Iniciando proceso...")

            # 1. Embed Reference Frame
            ref_embeddings = []
            ref_metadata = []
            
            # Check if we have 'texto_merged'
            if 'texto_merged' not in reference_df.columns:
                st.error("El archivo Frame debe tener la columna 'texto_merged'.")
                return None

            # Generate embeddings for Frame (with cache)
            for idx, row in reference_df.iterrows():
                text = str(row['texto_merged'])
                
                # Try cache first
                emb = cache_obj.get(text)
                if emb is None:
                    emb = get_embedding_with_retry(text) 
                    if emb is not None:
                        cache_obj.set(text, emb)
                
                if emb is not None:
                    ref_embeddings.append(emb)
                    ref_metadata.append({
                        'dimension': row.get('dimension', 'Unknown'),
                        'subdim': row.get('subdim', 'Unknown'),
                        'texto_merged': text
                    })
                
                progress_val = 0.1 * (idx + 1) / len(reference_df)
                progress_bar.progress(progress_val, text=f"Procesando marco de referencia {idx+1}/{len(reference_df)}")
            
            cache_obj.save() # Save frame embeddings
            
            if not ref_embeddings:
                st.error("No se pudieron generar embeddings para el marco de referencia.")
                return None

            ref_embeddings = np.array(ref_embeddings)
            
            # Normalize frame embeddings
            ref_norms = np.linalg.norm(ref_embeddings, axis=1, keepdims=True)
            ref_embeddings_norm = ref_embeddings / (ref_norms + 1e-10)
            
            # 2. Embed Unique Recommendations (Deduplication)
            start_time = time.time()
            total_recs = len(target_df)
            
            # Identify unique texts to embed
            unique_texts = target_df['Recommendation description'].dropna().unique()
            unique_embeddings = {}
            
            new_embeddings_count = 0
            
            # Embed unique texts
            BATCH_SIZE = 20
            for i in range(0, len(unique_texts), BATCH_SIZE):
                batch_texts = unique_texts[i:i+BATCH_SIZE]
                for text in batch_texts:
                    text_str = str(text)
                    # Try cache
                    emb = cache_obj.get(text_str)
                    if emb is None:
                        emb = get_embedding_with_retry(text_str)
                        if emb is not None:
                            cache_obj.set(text_str, emb)
                            new_embeddings_count += 1
                    
                    if emb is not None:
                        unique_embeddings[text_str] = emb
                
                progress_pct = 0.1 + (0.4 * (i + len(batch_texts)) / len(unique_texts))
                progress_bar.progress(progress_pct, text=f"Generando embeddings únicos: {min(i+len(batch_texts), len(unique_texts))}/{len(unique_texts)}")
                
                # Save cache periodically
                if new_embeddings_count > 50:
                    cache_obj.save()
                    new_embeddings_count = 0

            cache_obj.save() # Final save

        # 3. Classify Rows (Match) - Top 3
            classified_rows = []
            
            progress_bar.progress(0.5, text="Asignando clasificaciones (Top 3) " + ("y verificando con LLM..." if use_llm_verification else "..."))
            
            for i, (idx, row) in enumerate(target_df.iterrows()):
                rec_text = str(row.get('Recommendation description', ''))
                
                # Default values
                vals = {
                    'assigned_dimension': 'Unclassified',
                    'assigned_subdim': 'Unclassified',
                    'matched_frame_text': '',
                    'similarity_score': 0.0,
                    'Otras dimensiones': '',
                    'Otras subdimensiones': ''
                }

                if rec_text in unique_embeddings:
                    rec_emb = unique_embeddings[rec_text]
                    
                    # Cosine similarity
                    rec_emb_norm = rec_emb / (np.linalg.norm(rec_emb) + 1e-10)
                    similarities = np.dot(ref_embeddings_norm, rec_emb_norm)
                    
                    # Get Top Matches for context
                    # Even if verifying, we start with top candidates from embeddings
                    
                    if use_llm_verification:
                        top_n_candidates = 10
                    else:
                        top_n_candidates = 3
                        
                    top_k_indices = np.argsort(similarities)[-top_n_candidates:][::-1]
                    
                    best_idx = top_k_indices[0] # Default to embedding top 1
                    
                    # LLM Verification (Reranking)
                    if use_llm_verification:
                        # Construct candidates list (Top 10)
                        candidates = []
                        for k in top_k_indices:
                             candidates.append(ref_metadata[k])
                        
                        # Call LLM
                        verified_idx = verify_match_with_llm(rec_text, candidates, openai_api_key)
                        
                        if verified_idx >= 0 and verified_idx < len(top_k_indices):
                            best_idx = top_k_indices[verified_idx]
                        # Else fallback to embedding best_idx (0)
                    
                    # Assign Best Match (Verified or Embedding)
                    best_match = ref_metadata[best_idx]
                    
                    vals['assigned_dimension'] = best_match['dimension']
                    vals['assigned_subdim'] = best_match['subdim']
                    vals['matched_frame_text'] = best_match['texto_merged']
                    vals['similarity_score'] = similarities[best_idx]
                    
                    # Process Top 2 & 3 (Alternatives)
                    # Note: If LLM picked #2, then #1 and #3 become alternatives.
                    # For simplicity, we stick to the Embedding Ranking for "Others" 
                    # OR we could just list the other 2 from the Top 3 set.
                    # Let's simple: List the *other* indices from top_k_indices that are NOT best_idx
                    
                    secondary_dims = []
                    secondary_subdims = []
                    
                    similarity_threshold = 0.60
                    
                    for k_idx in top_k_indices:
                        if k_idx == best_idx: continue # Skip the chosen one
                        
                        if similarities[k_idx] >= similarity_threshold:
                            match_k = ref_metadata[k_idx]
                            secondary_dims.append(match_k['dimension'])
                            secondary_subdims.append(match_k['subdim'])
                        
                    vals['Otras dimensiones'] = "; ".join(secondary_dims)
                    vals['Otras subdimensiones'] = "; ".join(secondary_subdims)
                
                # Merge original row data with new classification data
                # We want to keep ALL original columns
                row_dict = row.to_dict()
                row_dict.update(vals)
                classified_rows.append(row_dict)
                
                if i % 10 == 0: # Update faster if slow LLM
                     progress_bar.progress(0.5 + (0.5 * (i + 1) / total_recs), text=f"Clasificando registros {i+1}/{total_recs}")

            progress_bar.empty()
            return pd.DataFrame(classified_rows)

        # Button to run classification
        st.markdown("### 2. Ejecución")
        
        # LLM Verification Toggle
        use_llm_chk = st.checkbox("✅ Usar Verificación con LLM (Alta Precisión, Más Lento)", value=False, help="Si se activa, el sistema usará IA para leer las definiciones y corregir la clasificación (ej. distinguir 'Pensiones' financiera de 'Pensiones' política).")

        if start_count == 0:
            st.warning("El archivo de recomendaciones está vacío.")
        elif end_count == 0:
             st.warning("No hay registros que coincidan con los filtros seleccionados. Ajusta los filtros.")
        else:
            if st.button("🚀 Iniciar Clasificación", key="start_classification_tab6"):
                with st.spinner("Ejecutando clasificación optimizada..."):
                    classified_df = classify_recommendations(filtered_rec_df, frame_df_world, st.session_state['embeddings_cache_obj'], use_llm_verification=use_llm_chk)
                    if classified_df is not None and not classified_df.empty:
                        st.session_state['classified_world_df'] = classified_df
                        st.success(f"¡Clasificación completada! {len(classified_df)} recomendaciones procesadas.")
                    else:
                        st.error("Hubo un problema durante la clasificación.")

        # --- Visualization ---
        df_viz = st.session_state['classified_world_df']
        
        if df_viz is not None:
            st.markdown("---")
            st.subheader("📊 Visualización de Resultados")
            
            
            # --- Visual Filters (Post-Classification) ---
            st.markdown("### 🔎 Filtros Visuales")
            st.warning("Estos filtros afectan solo a los gráficos y tablas a continuación, sin re-ejecutar la clasificación.")
            
            # Prepare filter options (based on current results to be safe)
            # Create a localized copy for filtering
            df_filtered_viz = df_viz.copy()

            # 1. Year Slider (Range)
            if 'Year' in df_viz.columns:
                 # Handle NaNs or zeros logic 
                 valid_years = sorted([int(x) for x in df_viz['Year'].unique() if pd.notna(x) and x > 1900])
                 if valid_years:
                     min_year, max_year = min(valid_years), max(valid_years)
                     year_range = st.slider(
                         "Intervalo de Años:",
                         min_value=min_year,
                         max_value=max_year,
                         value=(min_year, max_year),
                         key="viz_year_slider"
                     )
                     df_filtered_viz = df_filtered_viz[
                         (df_filtered_viz['Year'] >= year_range[0]) & 
                         (df_filtered_viz['Year'] <= year_range[1])
                     ]

            # 2. Metadata Filters (multiselects) - Grouped
            with st.expander("Más Filtros Visuales (Metadatos)", expanded=False):
                col_vf1, col_vf2, col_vf3 = st.columns(3)
                
                with col_vf1:
                    # Admin Unit
                    if 'Recommendation administrative unit' in df_filtered_viz.columns:
                         opts = sorted([str(x) for x in df_viz['Recommendation administrative unit'].unique() if pd.notna(x)])
                         sel_admin = st.multiselect("Unidad Admin:", options=opts, key="viz_admin_unit")
                         if sel_admin:
                             df_filtered_viz = df_filtered_viz[df_filtered_viz['Recommendation administrative unit'].isin(sel_admin)]

                with col_vf2:
                    # Country
                    if 'Country(ies)' in df_filtered_viz.columns:
                         opts = sorted([str(x) for x in df_viz['Country(ies)'].unique() if pd.notna(x)])
                         sel_country = st.multiselect("País:", options=opts, key="viz_country")
                         if sel_country:
                             df_filtered_viz = df_filtered_viz[df_filtered_viz['Country(ies)'].isin(sel_country)]

                with col_vf3:
                    # Dimension (Pre-filter for everything?? Or just let the treemap handle it? User asked for filters)
                    # Let's add Dimension here too affects everything
                    opts = sorted([str(x) for x in df_viz['assigned_dimension'].unique() if pd.notna(x)])
                    sel_dim_global = st.multiselect("Dimensión:", options=opts, key="viz_dim_global")
                    if sel_dim_global:
                         df_filtered_viz = df_filtered_viz[df_filtered_viz['assigned_dimension'].isin(sel_dim_global)]


            # Count Unique Recommendations
            # Use 'Recommendation ID' if available, otherwise 'Recommendation description'
            id_col = 'Recommendation ID' if 'Recommendation ID' in df_filtered_viz.columns else 'Recommendation description'
            unique_count = df_filtered_viz[id_col].nunique()
            
            st.markdown(f"**Registros visualizados:** {len(df_filtered_viz)} | **Recomendaciones Únicas:** {unique_count}")
            st.warning("⚠️ Nota: El número de registros puede ser mayor que el número de recomendaciones únicas debido a que algunas recomendaciones tienen múltiples atributos (ej. múltiples temas), generando filas duplicadas en el archivo original.")

            st.markdown("---") 
            
            # Helper to get deduplicated DF for strict counts
            # We use this for Treemaps and Single-Value Evolution
            df_viz_dedup = df_filtered_viz.drop_duplicates(subset=[id_col]).copy() 

            
            # --- Chart Logic (Updated to use df_filtered_viz) ---
            
            # --- Treemap 1: Dimension ---
            available_dims = sorted(df_filtered_viz['assigned_dimension'].unique())
            
            # We use a selectbox for explicit control (MOVED TO TOP)
            manual_dim = st.selectbox(
                "Filtrar Subdimensión por Dimensión (Treemaps):", 
                options=["Todos"] + available_dims, 
                index=0,
                key="manual_dim_filter_top"
            )
            
            selected_dim_label = None
            if manual_dim != "Todos":
                selected_dim_label = manual_dim
            
            # --- Treemap 1: Dimension ---
            st.markdown("#### Por Dimensión")
            st.caption("Selecciona una dimensión en el gráfico para ver detalles (opcional).")
            
            # USE DEDUPLICATED DATA FOR TREEMAPS
            dim_counts = df_viz_dedup['assigned_dimension'].value_counts().reset_index()
            dim_counts.columns = ['dimension', 'count']
            
            fig_dim = px.treemap(
                dim_counts,
                path=['dimension'],
                values='count',
                title='Recomendaciones Únicas por Dimensión',
                color='dimension'
            )
            fig_dim.update_traces(textinfo="label+value", textfont_size=20)
            
            selection_dim = st.plotly_chart(fig_dim, on_select="rerun", key="treemap_dim_select", use_container_width=True)
            
            if isinstance(selection_dim, dict) and "selection" in selection_dim and "points" in selection_dim["selection"]:
                 points = selection_dim["selection"]["points"]
                 if points:
                     pt = points[0]
                     clicked_label = pt.get('label') or pt.get('x') or pt.get('id')
                     if clicked_label and clicked_label in available_dims and manual_dim == "Todos":
                         selected_dim_label = clicked_label
                         st.info(f"Filtro aplicado por selección en gráfico: {selected_dim_label}")

            # --- Treemap 2: Subdim ---
            st.markdown("#### Por Subdimensión")
            
            # USE DEDUPLICATED DATA FOR TREEMAPS
            df_sub = df_viz_dedup.copy()
            title_suffix = ""
            
            if selected_dim_label:
                if selected_dim_label in df_sub['assigned_dimension'].unique():
                    df_sub = df_sub[df_sub['assigned_dimension'] == selected_dim_label]
                    title_suffix = f" ({selected_dim_label})"
            
            if not df_sub.empty:
                subdim_counts = df_sub['assigned_subdim'].value_counts().reset_index()
                subdim_counts.columns = ['subdim', 'count']
                
                fig_sub = px.treemap(
                    subdim_counts,
                    path=['subdim'],
                    values='count',
                    title=f'Recomendaciones Únicas por Subdimensión{title_suffix}',
                    color='subdim'
                )
                fig_sub.update_traces(textinfo="label+value", textfont_size=20)
                st.plotly_chart(fig_sub, use_container_width=True)
            else:
                st.info("No hay datos para mostrar.")

            st.markdown("---")
            
            # --- Evolution Plots (New) ---
            st.subheader("📈 Evolución Temporal")
            
            evo_toggle = st.radio("Tipo de Gráfico:", ["Absoluto", "Porcentaje (100%)"], horizontal=True, key="evo_chart_type")
            is_percent = (evo_toggle == "Porcentaje (100%)")
            bar_norm_val = 'percent' if is_percent else None
            
            # Helper to plot evolution
            def plot_evolution(df_in, cat_col, title_prefix):
                if 'Year' not in df_in.columns or df_in.empty:
                    st.info("No hay datos de Año para mostrar evolución.")
                    return
                    
                # Filter out bad years (0, nan)
                df_clean = df_in[df_in['Year'] > 1900].copy()
                if df_clean.empty:
                     st.info("No hay datos válidos de año.")
                     return

                # Check cardinality
                unique_vals = df_clean[cat_col].nunique()
                if unique_vals > 20: 
                     st.warning(f"Hay muchos valores únicos ({unique_vals}) en {cat_col}. Se muestran los Top 20.")
                     top_20 = df_clean[cat_col].value_counts().nlargest(20).index
                     df_clean = df_clean[df_clean[cat_col].isin(top_20)]

                # Group
                df_grouped = df_clean.groupby(['Year', cat_col]).size().reset_index(name='count')
                
                # Determine title suffix
                t_suffix = " (%)" if is_percent else ""
                
                # Create figure
                # We use 'relative' barmode. If is_percent is true, we update layout with barnorm='percent'
                fig = px.bar(
                    df_grouped, 
                    x="Year", 
                    y="count", 
                    color=cat_col, 
                    title=f"Evolución de {title_prefix}{t_suffix}",
                    barmode='relative', 
                )
                
                if is_percent:
                     fig.update_layout(barnorm='percent')
                
                st.plotly_chart(fig, use_container_width=True)
            
            
            # Define the tabs - Expanded List
            tab_labels = [
                "Por País", "Por Unidad Admin", "Por Dimensión", "Por Subdimensión",
                "Resp. Gestión", "Temática Eval.", "Temática Rec.", "Unidad Técnica",
                "Fuente Fondo", "Naturaleza", "Tipo Eval.", "Momento", "Progreso", "Tipo Doc."
            ]
            
            tabs = st.tabs(tab_labels)
            
            # 0. Por País
            with tabs[0]:
                 if 'Country(ies)' in df_viz_dedup.columns:
                     plot_evolution(df_viz_dedup, 'Country(ies)', "País")
                 else:
                     st.info("Columna 'Country(ies)' no disponible.")

            # 1. Por Unidad Admin
            with tabs[1]:
                 if 'Recommendation administrative unit' in df_viz_dedup.columns:
                     plot_evolution(df_viz_dedup, 'Recommendation administrative unit', "Unidad Administrativa")
                 else:
                     st.info("Columna 'Recommendation administrative unit' no disponible.")

            # 2. Por Dimensión
            with tabs[2]:
                 plot_evolution(df_viz_dedup, 'assigned_dimension', "Dimensión")

            # 3. Por Subdimensión
            with tabs[3]:
                 plot_evolution(df_viz_dedup, 'assigned_subdim', "Subdimensión")
                 
            # 4. Resp. Gestión (Management response) - Single value usually
            with tabs[4]:
                if 'Management response' in df_viz_dedup.columns:
                    plot_evolution(df_viz_dedup, 'Management response', "Respuesta de Gestión")
                else:
                    st.info("Columna 'Management response' no disponible.")
            
            # 5. Temática Eval. (Evaluation theme(s)) - MULTI VALUE (Keep duplicates to show frequency)
            with tabs[5]:
                if 'Evaluation theme(s)' in df_filtered_viz.columns:
                    plot_evolution(df_filtered_viz, 'Evaluation theme(s)', "Temática de Evaluación")
                else:
                    st.info("Columna 'Evaluation theme(s)' no disponible.")

            # 6. Temática Rec. (Recommendation theme) - MULTI VALUE (Keep duplicates)
            with tabs[6]:
                if 'Recommendation theme' in df_filtered_viz.columns:
                    plot_evolution(df_filtered_viz, 'Recommendation theme', "Temática de Recomendación")
                else:
                    st.info("Columna 'Recommendation theme' no disponible.")

            # 7. Unidad Técnica (Technical unit(s)) - MULTI VALUE ? (Likely yes if it has (s))
            with tabs[7]:
                if 'Technical unit(s)' in df_filtered_viz.columns:
                    plot_evolution(df_filtered_viz, 'Technical unit(s)', "Unidad Técnica")
                else:
                    st.info("Columna 'Technical unit(s)' no disponible.")

            # 8. Fuente Fondo (Funding source(s)) - MULTI VALUE
            with tabs[8]:
                if 'Funding source(s)' in df_filtered_viz.columns:
                    plot_evolution(df_filtered_viz, 'Funding source(s)', "Fuente de Financiamiento")
                else:
                    st.info("Columna 'Funding source(s)' no disponible.")

            # 9. Naturaleza (Evaluation nature) - Single
            with tabs[9]:
                if 'Evaluation nature' in df_viz_dedup.columns:
                    plot_evolution(df_viz_dedup, 'Evaluation nature', "Naturaleza de Evaluación")
                else:
                    st.info("Columna 'Evaluation nature' no disponible.")
            
            # 10. Tipo Eval. (Evaluation type) - Single
            with tabs[10]:
                if 'Evaluation type' in df_viz_dedup.columns:
                    plot_evolution(df_viz_dedup, 'Evaluation type', "Tipo de Evaluación")
                else:
                    st.info("Columna 'Evaluation type' no disponible.")

            # 11. Momento (Evaluation timing) - Single
            with tabs[11]:
                if 'Evaluation timing' in df_viz_dedup.columns:
                    plot_evolution(df_viz_dedup, 'Evaluation timing', "Momento de Evaluación")
                else:
                    st.info("Columna 'Evaluation timing' no disponible.")

            # 12. Progreso (Progress) - Single
            with tabs[12]:
                if 'Progress' in df_viz_dedup.columns:
                    plot_evolution(df_viz_dedup, 'Progress', "Progreso")
                else:
                    st.info("Columna 'Progress' no disponible.")

            # 13. Tipo Doc. (Evaluation document type) - Single
            with tabs[13]:
                if 'Evaluation document type' in df_viz_dedup.columns:
                    plot_evolution(df_viz_dedup, 'Evaluation document type', "Tipo de Documento")
                else:
                    st.info("Columna 'Evaluation document type' no disponible.")

            # --- SECCION 3: HERRAMIENTAS AVANZADAS DE IA (NUEVA) ---
            st.markdown("---")
            st.subheader("🤖 3. Herramientas Avanzadas de IA")
            st.info("Utiliza herramientas de Inteligencia Artificial para analizar en profundidad o resumir las recomendaciones filtradas.")

            # --- Shared Filters for AI Tools ---
            st.markdown("##### 1. Definir Subconjunto para Análisis")
            
            df_ai_base = df_filtered_viz.copy()
            ai_filters = {}
            
            with st.expander("🔎 Filtros Globales para Análisis AI", expanded=True):
                 c_ai1, c_ai2 = st.columns(2)
                 
                 with c_ai1:
                      if 'Region(s)' in df_viz.columns:
                          opts = sorted([str(x) for x in df_viz['Region(s)'].unique() if pd.notna(x)])
                          ai_filters['Region(s)'] = st.multiselect("Región:", opts, key="ai_region")
                      
                      if 'Country(ies)' in df_viz.columns:
                           opts = sorted([str(x) for x in df_viz['Country(ies)'].unique() if pd.notna(x)])
                           ai_filters['Country(ies)'] = st.multiselect("País:", opts, key="ai_country")
                      
                      if 'Evaluation theme(s)' in df_viz.columns:
                           opts = sorted([str(x) for x in df_viz['Evaluation theme(s)'].unique() if pd.notna(x)])
                           ai_filters['Evaluation theme(s)'] = st.multiselect("Temática Eval:", opts, key="ai_eval_theme")

                 with c_ai2:
                      if 'Year' in df_viz.columns:
                          opts = sorted([int(x) for x in df_viz['Year'].unique() if pd.notna(x)])
                          ai_filters['Year'] = st.multiselect("Año:", opts, key="ai_year")
                      
                      if 'Management response' in df_viz.columns:
                          opts = sorted([str(x) for x in df_viz['Management response'].unique() if pd.notna(x)])
                          ai_filters['Management response'] = st.multiselect("Resp. Gestión:", opts, key="ai_mgmt")
                      
                      if 'assigned_dimension' in df_viz.columns:
                          opts = sorted([str(x) for x in df_viz['assigned_dimension'].unique() if pd.notna(x)])
                          ai_filters['assigned_dimension'] = st.multiselect("Dimensión:", opts, key="ai_dim")

            # Apply Filters
            for col, vals in ai_filters.items():
                if vals and col in df_ai_base.columns:
                    df_ai_base = df_ai_base[df_ai_base[col].isin(vals)]
            
            st.write(f"**Total Registros Filtrados:** {len(df_ai_base)}")
            
            st.markdown("---")
            
            # --- Dual Action Buttons ---
            col_deep, col_summ = st.columns(2)
            
            # --- LEFT: Deep Analysis ---
            with col_deep:
                st.markdown("#### 🧠 Análisis Profundo")
                st.caption("Evalúa coherencia, calidad e innovación de planes de acción.")
                
                # Prepare Deep Data
                df_deep_ready = df_ai_base.copy()
                if 'Action plan' in df_deep_ready.columns:
                    df_deep_ready = df_deep_ready[df_deep_ready['Action plan'].notna() & (df_deep_ready['Action plan'].astype(str).str.strip() != "")]
                valid_deep_count = len(df_deep_ready)
                
                st.metric("Válidos (con Plan)", valid_deep_count)
                
                with st.expander("Configuración"):
                    deep_model = st.selectbox("Modelo:", ["gpt-4o-mini", "gpt-4o"], index=0, key="deep_model_sel")
                    limit_rows_deep = st.number_input("Límite (0=Todos):", min_value=0, value=0, step=10, key="deep_limit")
                
                if st.button("🚀 Iniciar Análisis", key="btn_run_deep"):
                     if df_deep_ready.empty:
                         st.warning("No hay datos válidos.")
                     else:
                         id_col_deep = 'Recommendation ID' if 'Recommendation ID' in df_deep_ready.columns else 'Recommendation description'
                         df_deep_input = df_deep_ready.drop_duplicates(subset=[id_col_deep]).copy()
                         if limit_rows_deep > 0:
                             df_deep_input = df_deep_input.head(limit_rows_deep)
                        
                         st.info(f"Procesando all {len(df_deep_input)} recomendaciones..." if limit_rows_deep == 0 else f"Procesando {len(df_deep_input)} recomendaciones...")
                         
                         analysis_cache = AnalysisCache()
                         args_list = []
                         for idx, row in df_deep_input.iterrows():
                             rec = str(row.get('Recommendation description', ''))
                             plan = str(row.get('Action plan', ''))
                             comments = str(row.get('Comments', '')) if 'Comments' in row else ""
                             args_list.append((idx, rec, plan, comments, openai_api_key, analysis_cache))
                         
                         results_map = {}
                         pbar_deep = st.progress(0, text="Analizando...")
                         completed_count = 0
                         
                         with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
                             futures = [executor.submit(run_row_analysis, arg) for arg in args_list]
                             for future in concurrent.futures.as_completed(futures):
                                 idx, result, _ = future.result()
                                 if result:
                                    results_map[idx] = result
                                 completed_count += 1
                                 # Update progress every time, but with higher concurrency it will feel 'batched'
                                 pbar_deep.progress(completed_count / len(args_list), text=f"Analizando... {completed_count}/{len(args_list)}")
                         
                         pbar_deep.empty()
                         analysis_cache.save()
                         
                         analysis_list = []
                         for idx, res in results_map.items():
                             res['original_index'] = idx
                             analysis_list.append(res)
                         
                         if analysis_list:
                             df_res = pd.DataFrame(analysis_list)
                             df_res.set_index('original_index', inplace=True)
                             df_final_deep = df_deep_input.join(df_res)
                             
                             list_cols = ['extracted_actions_from_rec', 'actions_proposed_in_plan', 'rejection_difficulty_classification', 'tags', 'rec_additional_tags']
                             for col in list_cols:
                                 if col in df_final_deep.columns:
                                     df_final_deep[col] = df_final_deep[col].apply(lambda x: ", ".join(x) if isinstance(x, list) else str(x))
                             
                             st.session_state['deep_analysis_df'] = df_final_deep
                             st.success("¡Análisis completado!")

            # --- RIGHT: Summary Generation ---
            with col_summ:
                st.markdown("#### ✨ Resumen Ejecutivo")
                st.caption("Genera una síntesis narrativa de los hallazgos.")
                
                st.metric("Total a Resumir", len(df_ai_base))
                
                with st.expander("Configuración"):
                    max_tokens_val = st.slider("Longitud (Tokens):", 500, 4000, 3000, 100, key="summ_len")
                
                if st.button("📝 Generar Resumen Global", key="btn_run_summ"):
                    if df_ai_base.empty:
                        st.warning("No hay datos.")
                    else:
                        with st.spinner("Generando resumen..."):
                            id_col_summ = 'Recommendation ID' if 'Recommendation ID' in df_ai_base.columns else 'Recommendation description'
                            df_summ_unique = df_ai_base.drop_duplicates(subset=[id_col_summ])
                            
                            text_list = []
                            for idx, row in df_summ_unique.iterrows():
                                 desc = str(row.get('Recommendation description', ''))
                                 mgmt = str(row.get('Management response', '')) if 'Management response' in row else "N/A"
                                 comments = str(row.get('Comments', '')) if 'Comments' in row else "N/A"
                                 action = str(row.get('Action plan', '')) if 'Action plan' in row else "N/A"
                                 text_list.append(f"--- Rec ---\nDesc: {desc}\nResp: {mgmt}\nCom: {comments}\nPlan: {action}\n")
                            
                            full_text = "\n".join(text_list)
                            summary_text = generate_executive_summary(full_text, max_output_tokens=max_tokens_val)
                            st.session_state['summary_result'] = summary_text
                            st.success("¡Resumen generado!")

            # --- Display Results Areas (Full Width) ---
            
            # 1. Deep Analysis Results
            if 'deep_analysis_df' in st.session_state:
                st.markdown("---")
                st.subheader("📊 Resultados: Análisis Profundo")
                df_final_deep = st.session_state['deep_analysis_df']
                
                c_m1, c_m2 = st.columns(2)
                c_m1.metric("Coherencia Prom.", f"{df_final_deep['coherence_score'].mean():.2f}")
                c_m2.metric("Calidad Plan Prom.", f"{df_final_deep['plan_quality_score'].mean():.2f}")
                
                st.dataframe(df_final_deep[['Recommendation description', 'coherence_score', 'plan_quality_score', 'rec_innovation_score', 'rejection_difficulty_classification', 'tags']], use_container_width=True)
                
                # Export Deep
                out_deep = BytesIO()
                with pd.ExcelWriter(out_deep, engine='xlsxwriter') as writer:
                    df_final_deep.to_excel(writer, index=False, sheet_name='Analisis_Profundo')
                
                st.download_button("📥 Descargar Reporte Análisis (.xlsx)", out_deep.getvalue(), "analisis_profundo.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

                # Visualizations Deep
                st.markdown("##### 📈 Distribución de Métricas Clave")
                
                col_v1, col_v2 = st.columns(2)
                
                if 'coherence_score' in df_final_deep.columns:
                    with col_v1:
                        fig_d1 = px.histogram(df_final_deep, x="coherence_score", nbins=10, title="Distribución de Coherencia", color_discrete_sequence=['#636EFA'])
                        st.plotly_chart(fig_d1, use_container_width=True)
                
                if 'plan_quality_score' in df_final_deep.columns:
                     with col_v2:
                        fig_d2 = px.histogram(df_final_deep, x="plan_quality_score", nbins=10, title="Distribución de Calidad del Plan", color_discrete_sequence=['#EF553B'])
                        st.plotly_chart(fig_d2, use_container_width=True)
                
                col_v3, col_v4 = st.columns(2)
                
                if 'attention_level_score' in df_final_deep.columns:
                     with col_v3:
                        fig_d3 = px.histogram(df_final_deep, x="attention_level_score", nbins=10, title="Nivel de Atención", color_discrete_sequence=['#00CC96'])
                        st.plotly_chart(fig_d3, use_container_width=True)

                if 'rec_innovation_score' in df_final_deep.columns:
                     with col_v4:
                         # Categorical
                         fig_pie = px.pie(df_final_deep, names='rec_innovation_score', title="Nivel de Innovación", hole=0.3)
                         st.plotly_chart(fig_pie, use_container_width=True)
                
                st.markdown("##### 🚦 Factibilidad e Impacto")
                col_v5, col_v6 = st.columns(2)
                
                if 'rec_operational_feasibility' in df_final_deep.columns:
                    with col_v5:
                        fig_feas = px.histogram(df_final_deep, x='rec_operational_feasibility', title="Factibilidad Operativa", color='rec_operational_feasibility')
                        st.plotly_chart(fig_feas, use_container_width=True)
                
                if 'rec_expected_impact' in df_final_deep.columns:
                    with col_v6:
                         fig_imp = px.histogram(df_final_deep, x='rec_expected_impact', title="Impacto Esperado", color='rec_expected_impact')
                         st.plotly_chart(fig_imp, use_container_width=True)

            # 2. Summary Results
            if 'summary_result' in st.session_state:
                st.markdown("---")
                st.subheader("📄 Resultados: Resumen Ejecutivo")
                st.markdown(st.session_state['summary_result'])
                
                # Export Summary
                out_summ = BytesIO()
                with pd.ExcelWriter(out_summ, engine='xlsxwriter') as writer:
                    # Save filtering data
                    df_ai_base.to_excel(writer, index=False, sheet_name='Datos Base')
                    pd.DataFrame({'Resumen': [st.session_state['summary_result']]}).to_excel(writer, index=False, sheet_name='Resumen')
                
                st.download_button("📥 Descargar Resumen + Datos (.xlsx)", out_summ.getvalue(), "resumen_ejecutivo.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
