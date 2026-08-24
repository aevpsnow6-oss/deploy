"""Generate the Spanish DOCX pilot-support proposal for the ILO GPT assistants."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x00, 0x3E, 0x7E)
CYAN = RGBColor(0x00, 0x72, 0xBC)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT = "Propuesta_Acompanamiento_Piloto_OIT.docx"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, size=10.5, after=6, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    return p


def bullet(doc, text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        r = p.add_run(prefix)
        r.font.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        set_cell_bg(c, "003E7E")
        for para in c.paragraphs:
            for r in para.runs:
                r.font.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cells[i].text = val
            set_cell_bg(cells[i], bg)
            for para in cells[i].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.4)
    s.bottom_margin = Cm(2.4)
    s.left_margin = Cm(2.6)
    s.right_margin = Cm(2.6)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ── Encabezado ───────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("PROPUESTA DE ACOMPAÑAMIENTO TÉCNICO")
r.font.bold = True
r.font.size = Pt(18)
r.font.color.rgb = BLUE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)

p2 = doc.add_paragraph()
r2 = p2.add_run("Fase piloto de los asistentes GPT de valoración de documentos de proyecto")
r2.font.size = Pt(12.5)
r2.font.color.rgb = CYAN
r2.font.bold = True
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(10)

meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate([
    ("Presentado a", "Organización Internacional del Trabajo (OIT) — Oficina Regional para América Latina y el Caribe"),
    ("Presentado por", "Ahmed Guillermo Eid Valdiviezo — Consultor externo"),
    ("Objeto", "Acompañamiento técnico y monitoreo durante la fase piloto (presupuesto de USD 500 en consumo de API)"),
    ("Fecha", "Julio de 2026"),
]):
    row = meta.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], "003E7E")
    set_cell_bg(row.cells[1], "F0F4F8")
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(9.5)
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(12.5)

doc.add_paragraph()

# ── 1. Antecedentes ──────────────────────────────────────────────────────
heading(doc, "1. Antecedentes")
body(doc,
     "La OIT cuenta con tres asistentes GPT operativos para la valoración asistida de documentos de "
     "proyecto: valoración preliminar de calidad de PRODOCs, diagnóstico de atributos específicos "
     "(género, metodologías participativas y transición justa) y diagnóstico de sostenibilidad. Los tres "
     "comparten un backend desplegado en la nube y devuelven resultados en formato Excel, con la "
     "valoración por criterio, la evidencia citada y las áreas que requieren revisión experta.")
body(doc,
     "Una vez concluido el desarrollo y la puesta en marcha, el siguiente paso natural es una fase piloto "
     "controlada que permita validar los asistentes en uso real, medir su desempeño y su costo, y generar "
     "la evidencia necesaria para decidir sobre un eventual despliegue a mayor escala.")

# ── 2. Objetivo ──────────────────────────────────────────────────────────
heading(doc, "2. Objetivo de la fase piloto")
body(doc,
     "Validar en condiciones reales los tres asistentes GPT con un presupuesto acotado de USD 500 "
     "destinado al consumo de la API de OpenAI, asegurando que dicho consumo se traduzca en aprendizajes "
     "institucionales concretos y en una base sólida de evidencia para la toma de decisiones.")

# ── 3. Enfoque de la propuesta (cost framing) ────────────────────────────
heading(doc, "3. Enfoque de la propuesta")
body(doc,
     "El presupuesto de USD 500 asignado a esta fase se destina en su totalidad al consumo de la API de "
     "OpenAI —el insumo que hace posible la evaluación automática—. El acompañamiento técnico del "
     "consultor (soporte, monitoreo del consumo y actualización de rúbricas) se incorpora como parte de "
     "esta propuesta con el fin de asegurar que la inversión en cómputo rinda resultados medibles y "
     "accionables.")
body(doc,
     "Como responsable del desarrollo entregado, el consultor se encuentra en la mejor posición para dar "
     "continuidad al piloto y garantizar su estabilidad. Por ello, la propuesta prioriza deliberadamente "
     "que el presupuesto disponible se concentre en el insumo crítico —el cómputo— y no en tareas de "
     "soporte y seguimiento, que el consultor asume como parte de este acompañamiento.")

body(doc, "Asignación del presupuesto de la fase piloto:", bold=True, after=4)
table(doc,
      ["Concepto", "Monto", "Función"],
      [["Consumo de API de OpenAI", "USD 500 (100%)", "Insumo que habilita la evaluación automática de documentos"],
       ["Acompañamiento técnico del consultor\n(soporte · monitoreo · rúbricas)", "Incluido", "Asegura la estabilidad del servicio y el rendimiento de la inversión"]],
      widths=[6.0, 3.0, 7.0])

doc.add_paragraph()

# ── 4. Alcance del acompañamiento ────────────────────────────────────────
heading(doc, "4. Alcance del acompañamiento del consultor")
body(doc, "Durante la fase piloto, el consultor proveerá:", after=4)
for pfx, txt in [
    ("Soporte técnico básico: ", "atención a incidencias de uso de los asistentes y del servicio, resolución de fallos y orientación puntual a las personas usuarias."),
    ("Monitoreo y análisis del consumo de API: ", "seguimiento del gasto en el tablero de OpenAI, con proyección de agotamiento del presupuesto y alertas oportunas."),
    ("Actualización de rúbricas: ", "ajustes a las rúbricas cargadas en el servidor cuando el equipo técnico de la OIT lo requiera."),
    ("Seguimiento del estado del servicio: ", "verificación de disponibilidad del backend y revisión de evaluaciones fallidas o interrumpidas."),
    ("Recolección de retroalimentación: ", "consolidación ligera de observaciones de las personas usuarias para orientar mejoras."),
]:
    bullet(doc, txt, prefix=pfx)

body(doc,
     "El alcance del soporte es de nivel básico: no comprende el desarrollo de nuevas funcionalidades, "
     "nuevas integraciones ni migraciones de infraestructura, que serían objeto de un acuerdo aparte.",
     italic=True, size=9.5, after=6)

# ── 5. Monitoreo del consumo de API ──────────────────────────────────────
heading(doc, "5. Monitoreo y análisis del consumo de API")
body(doc,
     "El seguimiento del consumo de OpenAI es un componente central de esta propuesta: permite cuidar el "
     "presupuesto de USD 500 y, a la vez, generar métricas de valor para la decisión de escala. OpenAI "
     "expone estas métricas tanto en el tablero de uso (platform.openai.com/usage) como mediante su API "
     "de uso y costos, lo que permite un seguimiento sistemático.")

body(doc,
     "Acceso a los datos: estas métricas se obtienen de forma directa desde el tablero de uso de OpenAI, "
     "incluida su exportación a CSV, sin necesidad de credenciales de administrador ni de desarrollos "
     "adicionales. Un seguimiento programático más automatizado es posible como mejora futura, pero no es "
     "condición para este piloto.",
     size=9.5, after=6)

body(doc, "Métricas directas disponibles en OpenAI:", bold=True, after=4)
table(doc,
      ["Métrica", "Qué informa"],
      [["Gasto (USD)", "Costo acumulado, diario y saldo restante frente al presupuesto de USD 500."],
       ["Tokens de entrada y de salida", "Volumen de texto procesado en cada llamada; es la base del costo."],
       ["Tokens en caché y de razonamiento", "Reportados por separado; los modelos de razonamiento detallan estos tokens, que inciden en el costo."],
       ["Número de solicitudes", "Total de llamadas al modelo en el período."],
       ["Desglose por modelo, proyecto y clave de API", "Segmenta el gasto según esas dimensiones."],
       ["Serie temporal (por día)", "Permite observar el ritmo de gasto a lo largo del piloto."]],
      widths=[5.5, 10.5])
doc.add_paragraph()

body(doc, "Análisis que el consultor derivará de esas métricas:", bold=True, after=4)
for txt in [
    "Costo promedio por evaluación y por documento procesado.",
    "Ritmo de consumo (USD por día) y proyección de la fecha estimada de agotamiento del presupuesto.",
    "Estimación del costo de un despliegue a mayor escala, por extrapolación del costo unitario observado.",
    "Número de evaluaciones y de documentos procesados durante el piloto.",
    "Recomendaciones de optimización (p. ej., proporción entre tokens de entrada, salida y razonamiento).",
    "Alertas al aproximarse a umbrales del presupuesto, para evitar interrupciones no planificadas.",
]:
    bullet(doc, txt)

body(doc,
     "Nota sobre la segmentación por asistente: el tablero desagrega el gasto por modelo, proyecto y clave "
     "de API. Para distinguir el consumo por asistente (calidad, atributos y sostenibilidad) conviene usar "
     "claves o proyectos separados, o etiquetar las solicitudes; en su defecto, el backend ya registra el "
     "uso de tokens por evaluación, lo que habilita un análisis complementario.",
     italic=True, size=9.5, after=6)

# ── 6. Entregables ───────────────────────────────────────────────────────
heading(doc, "6. Entregables")
table(doc,
      ["Entregable", "Descripción", "Momento"],
      [["Soporte durante el piloto", "Atención de incidencias y orientación a usuarios; registro de solicitudes atendidas.", "Continuo"],
       ["Reporte de consumo y uso", "Consumo de API, proyección de presupuesto, número de evaluaciones y distribución por asistente.", "Quincenal / bajo demanda"],
       ["Actualizaciones de rúbrica", "Ajustes aplicados en el servidor, con nota breve de los cambios realizados.", "Según necesidad"],
       ["Informe de cierre del piloto", "Síntesis de resultados, consumo real, aprendizajes y recomendaciones para el escalamiento.", "Al finalizar"]],
      widths=[4.0, 8.5, 3.5])

doc.add_paragraph()

# ── 7. Duración ──────────────────────────────────────────────────────────
heading(doc, "7. Duración")
body(doc,
     "La fase piloto se extenderá hasta agotar el presupuesto de USD 500 en consumo de API o por un "
     "período máximo de tres (3) meses desde su inicio, lo que ocurra primero. Las fechas exactas se "
     "acordarán con la OIT al momento de la aprobación.")

# ── 8. Indicadores de éxito ──────────────────────────────────────────────
heading(doc, "8. Indicadores de éxito de la fase piloto")
for txt in [
    "El servicio se mantiene disponible y estable durante todo el período.",
    "El presupuesto de API se ejecuta de forma controlada, sin interrupciones por agotamiento imprevisto.",
    "Se produce, a partir de los datos de consumo de OpenAI, una estimación del costo de un despliegue mayor.",
    "Las personas usuarias completan evaluaciones y descargan resultados sin bloqueos técnicos.",
    "El informe de cierre entrega recomendaciones claras para la decisión de escalamiento.",
]:
    bullet(doc, txt)

# ── 9. Supuestos ─────────────────────────────────────────────────────────
heading(doc, "9. Supuestos y condiciones")
for txt in [
    "La OIT provee y administra la cuenta y la clave de API de OpenAI, y asigna el presupuesto de USD 500.",
    "Para el monitoreo, la OIT otorga al consultor acceso de lectura al tablero de uso o le comparte la exportación CSV periódica; el seguimiento no requiere una clave de administrador de la organización.",
    "Los asistentes permanecen desplegados en la infraestructura actual durante el piloto.",
    "El acompañamiento aquí descrito corresponde a la fase piloto; un despliegue institucional posterior se acordaría por separado.",
]:
    bullet(doc, txt)

# ── Cierre ───────────────────────────────────────────────────────────────
doc.add_paragraph()
body(doc,
     "Esta propuesta busca que la fase piloto se ejecute con la mayor solidez técnica posible, "
     "concentrando el presupuesto en el cómputo y aportando el acompañamiento necesario para convertir "
     "el uso real en aprendizaje institucional. Quedo a disposición para ajustar el alcance según las "
     "prioridades de la OIT.", after=14)

p = doc.add_paragraph()
r = p.add_run("Ahmed Guillermo Eid Valdiviezo\nConsultor externo")
r.font.size = Pt(10.5)
r.font.bold = True
r.font.color.rgb = BLUE

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
