from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
EN_DIR = BASE / "en"
UPDATED = "15 de julio de 2026"
UPDATED_EN = "July 15, 2026"
APP = "oli_v6_deploy.py + gpt_action_api.py"

# Asistentes GPT publicados en ChatGPT (backend FastAPI compartido).
GPT_BACKEND = "https://ilo-prodoc-appraisal-v3.onrender.com"
GPT_V3 = "https://chatgpt.com/g/g-6a2643b11e708191adc1c03e64260a25-ilo-prodoc-quality-appraisal"
GPT_TAB2 = "https://chatgpt.com/g/g-6a43cc82d4d08191bf6e60357453e336-oit-diagnostico-de-atributos-especificos"
GPT_SUS = "https://chatgpt.com/g/g-6a43d24307b88191bb362632f133c0f3-oit-diagnostico-de-sostenibilidad-del-proyecto"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(3)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    keep_row_together(table.rows[0])
    repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255))
        shade_cell(hdr[i], "002F6C")
    for row in rows:
        table_row = table.add_row()
        keep_row_together(table_row)
        cells = table_row.cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        doc.add_paragraph(item, style=style)


def add_numbered(doc, items):
    for idx, item in enumerate(items, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.add_run(f"{idx}. {item}")


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text="", style=None):
    para = doc.add_paragraph(style=style)
    if text:
        para.add_run(text)
    return para


def codeblock(doc, lines):
    if isinstance(lines, str):
        lines = lines.splitlines()
    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        para.paragraph_format.space_after = Pt(0)


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, "002F6C"),
        ("Heading 1", 16, "002F6C"),
        ("Heading 2", 13, "0072CE"),
        ("Heading 3", 11.5, "333333"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def cover(doc, title, subtitle, language="es"):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 47, 108)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Project Performance Toolkit"
        if language == "en"
        else "Caja de Herramientas para el Mejor Desempeño de los Proyectos"
    )
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 47, 108)

    source_label = "Technical source" if language == "en" else "Fuente técnica"
    updated_label = "Updated" if language == "en" else "Actualizado"
    updated_value = UPDATED_EN if language == "en" else UPDATED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{subtitle}\n{source_label}: {APP}\n{updated_label}: {updated_value}")
    doc.add_paragraph()


def save_doc(filename, title, subtitle, build, output_dir=BASE, language="es"):
    doc = Document()
    set_styles(doc)
    cover(doc, title, subtitle, language=language)
    build(doc)
    doc.save(output_dir / filename)


def doc1(doc):
    h(doc, "1. Alcance y fuente de verdad", 1)
    p(doc, "Esta documentación técnica describe exclusivamente la aplicación principal oli_v6_deploy.py. Los archivos oli_v6_deploy_core.py y oli_v6_deploy_recommendations.py son variantes separadas y no se usan como base de esta actualización documental.")
    p(doc, "La aplicación principal es una app Streamlit monolítica, con funciones auxiliares, lógica de extracción documental, evaluación por IA, clasificación de recomendaciones, visualización y exportación en un único archivo Python.")
    p(doc, "A partir de junio de 2026 el sistema tiene una segunda superficie de uso: tres asistentes GPT publicados en ChatGPT, servidos por un backend FastAPI independiente (gpt_action_api.py) desplegado en Render. Esa capa se documenta en la sección 9 y no comparte proceso ni estado con la app Streamlit.")

    h(doc, "2. Arquitectura de alto nivel", 1)
    add_table(doc, ["Capa", "Implementación en el código", "Responsabilidad"], [
        ["Interfaz", "Streamlit con st.tabs, st.session_state, st.cache_data y CSS institucional ILO", "Carga de documentos, selección de criterios, visualizaciones y descargas."],
        ["Procesamiento documental", "python-docx, docx2python, BeautifulSoup y extracción jerárquica de encabezados", "Convierte DOCX/TXT en texto trazable por sección, párrafo, tabla y metadatos."],
        ["IA generativa", "OpenAI SDK mediante OpenAI(api_key=os.getenv('OPENAI_API_KEY'))", "Evaluaciones, síntesis, chat documental, crítica de respuestas y resúmenes ejecutivos."],
        ["Búsqueda vectorial", "text-embedding-3-large, numpy, torch y FAISS IndexFlatIP", "Carga un índice FAISS de recomendaciones al iniciar la app y usa FAISS para RAG en documentos grandes de Tab 4; Tabs 5 y 6 clasifican con similitud coseno en numpy, no con FAISS."],
        ["Datos", "Archivos Excel, CSV pipe-separated y tensores .pt en el directorio de la app o rutas definidas por st.secrets", "No hay base de datos SQL; los archivos son el almacén operativo."],
        ["Exportación", "xlsxwriter, BytesIO y zipfile", "Genera XLSX y ZIP con resultados, evidencias, estructura extraída y reportes."],
    ], widths=[1.3, 2.4, 3.1])

    h(doc, "3. Componentes funcionales activos", 1)
    add_table(doc, ["Pestaña", "Nombre visible", "Función principal"], [
        ["1", "Valoración Preliminar de Calidad de Proyectos", "Evalúa documentos de diseño contra Appraisal Checklist_2025 es-419.xlsx, aplica análisis A-E y regla A/B para temas transversales, genera crítica y síntesis por pregunta, subsección y sección."],
        ["2", "Diagnóstico de Atributos Específicos", "Evalúa secciones seleccionadas con rúbricas de Rubricas_6ago2025.xlsx para participación durante evaluación, género y transición justa moderna."],
        ["3", "Diagnóstico de Sostenibilidad del Proyecto", "Evalúa sostenibilidad con Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx y visualiza puntajes por dimensión y criterio."],
        ["4", "Pregúntale a tus Documentos", "Chat con uno o varios DOCX/TXT, usando contexto completo para documentos pequeños y RAG con FAISS para documentos grandes."],
        ["5", "Clasificación de Recomendaciones", "Clasifica recomendaciones desde una interfaz en español contra Frame_Recommendations_English.xlsx; el campo requerido del archivo de recomendaciones es Recommendation description."],
        ["6", "Recommendation Classification", "Versión en inglés de la clasificación de recomendaciones con la misma lógica de Tab 5."],
    ], widths=[0.7, 2.2, 4.1])

    h(doc, "4. Dependencias", 1)
    p(doc, "requirements.txt fija o declara las siguientes dependencias principales:")
    codeblock(doc, [
        "streamlit==1.27.0",
        "pandas==2.0.3",
        "numpy==1.24.3",
        "faiss-cpu==1.7.4",
        "openai>=1.0.0",
        "matplotlib==3.7.2",
        "seaborn==0.12.2",
        "torch==2.0.1",
        "plotly==5.15.0",
        "xlsxwriter==3.1.2",
        "python-dotenv==1.0.0",
        "openpyxl",
        "rich==13.9.0",
        "python-docx==0.8.11",
        "docx2python",
        "tiktoken>=0.5.0",
    ])

    h(doc, "5. Modelos de IA y límites de contexto", 1)
    add_bullets(doc, [
        "Modelo principal de evaluación, crítica, síntesis y chat: gpt-5-mini.",
        "Embeddings: text-embedding-3-large.",
        "Análisis profundo de recomendaciones: la UI muestra un selector entre gpt-4o-mini y gpt-4o, pero el código actual no pasa esa selección a run_row_analysis; analyze_recommendation_plan_pair usa gpt-4o-mini por defecto.",
        "Límite documental operativo: truncate_to_token_limit corta a aproximadamente 110,000 tokens para dejar espacio al prompt y respuesta.",
        "En el chat documental, textos totales de hasta 100,000 caracteres usan contexto completo, luego se recortan a 110,000 tokens; textos mayores generan fragmentos de 2,000 caracteres con solape de 300 y RAG con FAISS.",
    ])

    h(doc, "5.1 Regla de preguntas en dos partes en Tab 1", 2)
    p(doc, "Tab 1 detecta preguntas compuestas mediante parse_two_part_question, con patrones para una declaración seguida de pregunta, separadores explícitos como 'específicamente', dos signos de pregunta, y casos unidos sin espacio. Cuando se detectan dos partes, la Parte 2 es el foco evaluativo y la Parte 1 es solo marco contextual.")
    add_bullets(doc, [
        "El razonamiento del primer análisis debe iniciar con 'Se identificaron 2 partes en esta pregunta.'.",
        "La evidencia se etiqueta como [DEDICATED] cuando aporta elementos específicos al sujeto evaluado y como [FRAMING] cuando es solo contexto, lista de grupos o lenguaje general.",
        "El crítico de Tab 1 vuelve a evaluar con el documento completo y el código aplica una compuerta mecánica sobre el veredicto.",
        "Si el modelo no mantiene el foco en Parte 2, la fila puede quedar con Status = Partial aunque el análisis haya terminado.",
    ])

    h(doc, "5.2 Regla de temas transversales en Tab 1", 2)
    p(doc, "Tab 1 mantiene el marco A-E general para preguntas con sujeto específico, pero lo omite por completo cuando la pregunta corresponde a un tema transversal configurado.")
    add_table(doc, ["Tema transversal reconocido", "Criterios de detección en código"], [
        ["Género", "Aliases como género, igualdad de género, enfoque de género, mujeres y niñas."],
        ["No discriminación", "Aliases como no discriminación, antidiscriminación, discriminación, igualdad de oportunidades y trato igualitario."],
        ["Discapacidad", "Aliases como discapacidad, personas con discapacidad, PCD, accesibilidad e inclusión de personas con discapacidad."],
        ["Diálogo social y tripartismo", "Aliases como diálogo social, tripartismo, tripartita, sindicatos, organizaciones de trabajadores y organizaciones de empleadores."],
        ["Sostenibilidad medioambiental", "Aliases como sostenibilidad ambiental, medio ambiente, ambiental, cambio climático, acción climática y economía verde."],
    ], widths=[2.1, 4.7])
    add_table(doc, ["Criterio transversal", "Definición operativa"], [
        ["A", "Presencia operacional del tema en objetivo, producto o actividad."],
        ["B", "Presupuesto, recursos o línea presupuestaria para alguna actividad correspondiente."],
    ], widths=[1.3, 5.5])
    p(doc, "La calificación automática para estos temas es: Yes cuando A y B están presentes; Partial cuando solo A o solo B está presente; No cuando no están presentes ni A ni B. Indicadores y metas no cuentan para la regla transversal. El crítico puede sobrescribir mecánicamente un veredicto inconsistente con A/B.")

    h(doc, "5.3 Comportamiento preciso de Tab 4", 2)
    add_bullets(doc, [
        "DOCX se extrae con python-docx leyendo párrafos; TXT se lee como texto plano UTF-8 con errors='ignore'. Este flujo no usa el extractor jerárquico enriquecido, por lo que tablas, metadatos de encabezado, páginas y citas por sección no son confiables en Tab 4.",
        "Para documentos grandes, los embeddings de fragmentos se guardan en st.session_state; en cada pregunta se reconstruye un FAISS IndexFlatIP, se normalizan los vectores y se recuperan hasta 15 fragmentos vectoriales.",
        "El contexto RAG agrega además hasta 5 fragmentos por coincidencia exacta de palabras clave, deduplicados contra los fragmentos vectoriales.",
        "La llamada al modelo incluye solo los últimos 5 mensajes del chat.",
        "La detección de cambio de archivos usa la lista de nombres, no hash de contenido; si se sube otro archivo con el mismo nombre puede quedar estado anterior hasta reiniciar sesión o cambiar nombres.",
        "La ruta de documentos pequeños exige responder solo con el documento; la ruta RAG permite inferencia con pistas contextuales, por lo que las conclusiones deben verificarse contra el archivo fuente.",
    ])

    h(doc, "6. Inicialización y caché", 1)
    add_bullets(doc, [
        "load_data() carga df_complete_all_full.xlsx, df_split_actions.xlsx y prepara columnas normalizadas.",
        "load_extended_data() agrega analyzed_recommendations_plans_v5.csv cuando está disponible y guarda analyzed_df en st.session_state.",
        "load_embeddings() carga emb_Recomm_rec_cl_4.pt y Recommendation_RAG_Metadata.pt y construye un índice FAISS en memoria. Aunque la interfaz visible de búsqueda de recomendaciones está comentada, esta carga sigue siendo una dependencia activa porque se ejecuta durante el arranque.",
        "Las funciones de carga usan st.cache_data; limpiar caché de Streamlit fuerza una recarga de archivos.",
        "El bloque principal inicializa datos y embeddings dos veces en el archivo; el comportamiento final es equivalente, pero una futura limpieza puede consolidarlo.",
        "Tabs 5 y 6 usan embeddings_cache.pkl para cachear embeddings de clasificación y analysis_cache.pkl para análisis profundo de recomendaciones.",
        "Tab 4 guarda documentos, fragmentos y embeddings solo en st.session_state durante la sesión.",
    ])

    h(doc, "7. Despliegue", 1)
    p(doc, "Ejecución local estándar:")
    codeblock(doc, "streamlit run oli_v6_deploy.py")
    p(doc, "En producción, si STREAMLIT_ENV=production, el código toma rutas de archivos desde st.secrets. La clave de OpenAI se lee desde la variable de entorno OPENAI_API_KEY.")
    add_table(doc, ["Secreto o variable", "Uso en el código"], [
        ["OPENAI_API_KEY", "Autenticación del cliente OpenAI."],
        ["df_path", "Ruta de df_complete_all_full.xlsx en producción."],
        ["df_raw_path", "Ruta de df_split_actions.xlsx en producción."],
        ["embeddings_path", "Ruta de emb_Recomm_rec_cl_4.pt en producción."],
        ["structured_embeddings_path", "Ruta de Recommendation_RAG_Metadata.pt en producción."],
        ["analyzed_recommendations_path", "Ruta opcional de analyzed_recommendations_plans_v5.csv en producción."],
        ["lessons_embeddings_path", "Ruta de emb_LL_ll_cl_4.pt para una función de lecciones que no está conectada a una pestaña activa."],
        ["structured_lessons_path", "Ruta de lessons_metadata.pt para una función de lecciones que no está conectada a una pestaña activa."],
    ], widths=[2.1, 4.7])

    h(doc, "8. Riesgos técnicos operativos", 1)
    add_bullets(doc, [
        "La app no implementa autenticación, autorización, auditoría, aislamiento por tenant, escaneo de archivos cargados ni redacción automática de secretos; depende del control del entorno de hosting.",
        "Los documentos cargados y sus extractos se envían a OpenAI para análisis; deben tratarse como datos potencialmente sensibles.",
        "El almacén de datos es file-based; reemplazar un archivo de entrada cambia el comportamiento de la app sin migraciones de esquema.",
        "Las rutas locales son relativas al directorio desde el cual se ejecuta Streamlit.",
        "Los archivos de salida XLSX/ZIP pueden contener citas textuales del documento original y deben manejarse con la misma confidencialidad.",
        "Los cachés pickle locales no están cifrados ni verificados por integridad.",
    ])

    h(doc, "9. Capa de asistentes GPT en ChatGPT", 1)
    p(doc, "Tres GPTs personalizados publicados en ChatGPT exponen los motores de evaluación sin necesidad de que el usuario abra Streamlit. Los tres consumen un mismo backend FastAPI desplegado en Render; el usuario sube el DOCX dentro del chat y recibe un XLSX descargable.")
    add_table(doc, ["Asistente", "Motor y rúbrica", "Escala", "Endpoints"], [
        ["ILO PRODOC Quality Appraisal (Tab 1 v3)", "tab1_v3_core.py + Rubrica_Tab1_Detallada_Full_v3.xlsx", "Yes / Partial / No / Not Found / N/A sobre 76 criterios en 5 secciones", "POST /v3/jobs, GET /v3/jobs/{id}, GET /v3/jobs/{id}/result"],
        ["OIT - Diagnóstico de Atributos Específicos (Tab 2)", "tab2_core.py + Rubricas_6ago2025.xlsx", "1 a 5 por criterio", "POST /attributes/jobs, GET /attributes/jobs/{id}, GET /attributes/jobs/{id}/result"],
        ["OIT - Diagnóstico de Sostenibilidad del Proyecto (Tab 3)", "sustainability_core.py + Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "0 a 3 por indicador, en 28 indicadores y 3 dimensiones", "POST /sustainability/jobs, GET /sustainability/jobs/{id}, GET /sustainability/jobs/{id}/result"],
    ], widths=[1.7, 2.0, 1.6, 1.5])

    h(doc, "9.1 Enlaces de acceso publicados", 2)
    add_table(doc, ["Asistente", "Enlace"], [
        ["Valoración Preliminar de Calidad (Tab 1 v3)", GPT_V3],
        ["Diagnóstico de Atributos Específicos (Tab 2)", GPT_TAB2],
        ["Diagnóstico de Sostenibilidad (Tab 3)", GPT_SUS],
        ["Backend compartido (FastAPI en Render)", GPT_BACKEND],
    ], widths=[2.4, 4.4])
    p(doc, "Los tres GPTs están configurados como \"Anyone with a link\": no figuran en el directorio público de GPTs, pero cualquier persona que reciba el enlace y tenga una cuenta de ChatGPT puede abrirlos y ejecutar evaluaciones. El enlace funciona como credencial al portador: reenviarlo transfiere el acceso.")
    add_bullets(doc, [
        "No hay lista de personas autorizadas ni verificación de identidad: el backend no registra quién lanzó cada evaluación.",
        "El consumo de la API se factura contra una única OPENAI_API_KEY compartida por los tres asistentes, de modo que un enlace difundido fuera del círculo previsto se traduce en gasto no atribuible.",
        "La única forma de revocar el acceso es cambiar la configuración de compartición del GPT o republicarlo para generar un enlace nuevo, y redistribuirlo.",
        "La distribución de los enlaces es, en la práctica, el único control de acceso: debe gestionarse con el mismo cuidado que una credencial.",
    ])

    h(doc, "9.2 Ciclo de vida de un trabajo de evaluación", 2)
    add_numbered(doc, [
        "El usuario sube un DOCX en el chat de ChatGPT y describe el alcance deseado.",
        "El GPT llama al endpoint POST correspondiente, pasando la referencia del archivo alojado en los servidores de OpenAI y los filtros solicitados.",
        "El backend descarga el archivo, extrae el texto, carga la rúbrica desde disco y lanza las llamadas al modelo en paralelo.",
        "El GPT sondea GET /{familia}/jobs/{id} hasta que el estado pasa a succeeded o failed.",
        "Con el trabajo terminado, el GPT llama a GET /{familia}/jobs/{id}/result y entrega al usuario un resumen narrativo más el XLSX en base64.",
    ])
    p(doc, "El estado de los trabajos vive en el diccionario JOBS de gpt_action_api.py, protegido por threading.Lock. Es almacenamiento en memoria del proceso: un reinicio o redespliegue del servicio pierde los trabajos en curso. Para uso institucional sostenido debe migrarse a Redis o equivalente.")

    h(doc, "9.3 Paralelismo y esquema de estabilidad", 2)
    add_table(doc, ["Motor", "Modelo", "MAX_WORKERS", "Corridas por criterio"], [
        ["tab1_v3_core.py", "gpt-5-mini con esfuerzo de razonamiento adaptativo (medium/minimal según subjetividad del criterio)", "48", "10 (STABILITY_REPEATS), con umbral de estabilidad de 80%"],
        ["tab2_core.py", "gpt-5-mini", "8", "5 (STABILITY_REPEATS), consolidadas por moda con porcentaje de estabilidad"],
        ["sustainability_core.py", "gpt-5-mini", "8", "5 (STABILITY_REPEATS), consolidadas por moda con porcentaje de estabilidad"],
    ], widths=[1.6, 2.6, 1.1, 1.5])
    p(doc, "Los tres motores repiten cada evaluación y consolidan las corridas al valor modal, reportando qué porcentaje de corridas coincidió con la moda. Ese porcentaje viaja al XLSX y permite distinguir un criterio evaluado con consenso de uno inestable.")
    p(doc, "La implementación difiere: tab2_core.py y sustainability_core.py usan el módulo compartido stability.py, mientras que tab1_v3_core.py trae su propio agregador (aggregate_repeated_criterion_results) con 10 repeticiones y un umbral de estabilidad de 80%, por debajo del cual el criterio se marca como inestable. Una valoración completa de los 76 criterios implica por tanto del orden de 760 llamadas al modelo, lo que domina su costo y su duración.")

    h(doc, "9.4 Despliegue del backend", 2)
    add_table(doc, ["Elemento", "Valor"], [
        ["Plataforma", "Render.com, plan free"],
        ["Servicio", "ilo-prodoc-appraisal-v3, definido en render.yaml como Blueprint"],
        ["Runtime", "Docker, imagen construida desde Dockerfile.gpt-action sobre Python 3.10-slim"],
        ["Framework", "FastAPI con uvicorn[standard]; dependencias en requirements.gpt-action.txt"],
        ["Health check", "GET /health, configurado como healthCheckPath en render.yaml"],
        ["CI/CD", "Redespliegue automático al hacer push a la rama main"],
        ["Especificación para ChatGPT", "openapi_gpt_action_v3.yaml, openapi_gpt_action_tab2.yaml y openapi_gpt_action_sustainability.yaml"],
    ], widths=[2.0, 4.8])
    p(doc, "El plan free de Render suspende el servicio tras un período de inactividad; la primera petición después de la suspensión tarda más de lo habitual mientras el contenedor arranca. Es el comportamiento esperado, no una falla.")

    h(doc, "9.5 Riesgos técnicos específicos de esta capa", 2)
    add_bullets(doc, [
        "El estado de trabajos es en memoria: un redespliegue durante una evaluación la pierde sin posibilidad de recuperación.",
        "El plan free de Render impone latencia de arranque en frío y no ofrece garantías de disponibilidad.",
        "Los documentos suben primero a los servidores de OpenAI y luego los descarga el backend; ambos tránsitos deben considerarse en el análisis de confidencialidad.",
        "Las rúbricas viven en el sistema de archivos de la imagen Docker: actualizarlas exige reconstruir y redesplegar, no basta con reemplazar un archivo.",
        "Los tres GPTs comparten un único backend y una única clave de API: una interrupción afecta simultáneamente a los tres.",
    ])


def doc2(doc):
    h(doc, "1. Alcance del código fuente", 1)
    p(doc, "El código documentado es oli_v6_deploy.py. Las variantes oli_v6_deploy_core.py y oli_v6_deploy_recommendations.py existen en el proyecto, pero esta documentación no infiere comportamiento desde ellas.")
    p(doc, "Regla de mantenimiento: cuando se cambie la app principal, revisar si el mismo cambio debe propagarse a las variantes separadas antes de entregar una versión operativa al equipo ILO.")

    h(doc, "2. Estructura relevante del directorio", 1)
    add_table(doc, ["Archivo", "Rol operativo"], [
        ["oli_v6_deploy.py", "Aplicación Streamlit principal y fuente de verdad para esta documentación."],
        ["requirements.txt", "Dependencias Python usadas por la app."],
        ["df_complete_all_full.xlsx", "Dataset principal de recomendaciones enriquecidas."],
        ["df_split_actions.xlsx", "Dataset crudo/desagregado usado para recomendaciones y años."],
        ["analyzed_recommendations_plans_v5.csv", "Análisis adicional de recomendaciones y planes; separador pipe."],
        ["emb_Recomm_rec_cl_4.pt", "Embeddings de recomendaciones cargados en arranque para el índice FAISS."],
        ["Recommendation_RAG_Metadata.pt", "Metadatos alineados al embedding de recomendaciones cargado en arranque."],
        ["Appraisal Checklist_2025 es-419.xlsx", "Preguntas de valoración preliminar, hoja rubric."],
        ["Rubricas_6ago2025.xlsx", "Rúbricas de atributos específicos usadas por Tab 2."],
        ["Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "Rúbrica activa de sostenibilidad para Tab 3."],
        ["Recommendations_World.xlsx", "Recomendaciones mundiales para clasificación en Tabs 5 y 6."],
        ["Frame_Recommendations_English.xlsx", "Marco de dimensiones/subdimensiones para clasificación."],
    ], widths=[2.7, 4.1])

    h(doc, "3. Instalación local", 1)
    add_numbered(doc, [
        "Usar Python compatible con las dependencias fijadas del proyecto.",
        "Crear y activar un entorno virtual.",
        "Instalar dependencias con pip install -r requirements.txt.",
        "Configurar OPENAI_API_KEY como variable de entorno.",
        "Ejecutar streamlit run oli_v6_deploy.py desde el directorio del proyecto.",
    ])
    codeblock(doc, [
        "python -m venv .venv",
        "source .venv/bin/activate",
        "pip install -r requirements.txt",
        "streamlit run oli_v6_deploy.py",
    ])
    p(doc, "La clave OPENAI_API_KEY debe configurarse en el entorno o gestor de secretos elegido por ILO; no se debe registrar ni compartir su valor real en la documentación.")

    h(doc, "4. Mapa de funciones principales", 1)
    add_table(doc, ["Bloque", "Funciones o clases", "Descripción"], [
        ["Excel/exportación", "to_excel", "Convierte DataFrames en XLSX usando xlsxwriter y desactiva strings_to_urls."],
        ["Tokens", "truncate_to_token_limit", "Recorta texto a límites de tokens con tiktoken."],
        ["Embeddings y RAG", "get_embedding_with_retry, find_similar_recommendations, find_recommendations_by_term_matching, load_embeddings", "Genera embeddings y carga un índice FAISS en arranque; la búsqueda visible asociada a recomendaciones está comentada, pero la dependencia sigue activa."],
        ["Chat documental", "Tab 4, truncate_to_token_limit, client.embeddings.create, faiss.IndexFlatIP", "Usa contexto completo para textos pequeños y RAG con fragmentos cacheados en sesión para textos grandes."],
        ["Análisis de recomendaciones", "AnalysisCache, analyze_recommendation_plan_pair, run_row_analysis, generate_executive_summary", "Evalúa coherencia de planes, calidad, factibilidad, impacto, innovación y genera resumen ejecutivo; el modelo efectivo del análisis profundo es gpt-4o-mini aunque la UI muestre selector."],
        ["Carga de datos", "prepare_additional_data, load_data, load_extended_data", "Normaliza columnas, años, categorías y fusiona análisis adicional."],
        ["Extracción documental", "extract_docx_structure_enhanced, validate_extraction, extract_document_content", "Extrae jerarquía, tablas, métricas y texto desde DOCX."],
        ["Valoración preliminar", "load_appraisal_questions, analyze_question_with_llm_tab1, detect_transversal_matter, _critic_impl, _apply_critic_gate_and_render, _apply_transversal_gate_and_render, synthesize_subsection_analysis, synthesize_section_analysis, create_results_download_with_sections", "Pregunta por pregunta, aplica crítica A-E o regla transversal A/B y genera XLSX/ZIP multinivel."],
        ["Sostenibilidad y atributos", "evaluate_criterion_with_llm, synthesize_evaluations", "Evalúa criterios con puntuación 1-5, análisis y evidencia."],
        ["Clasificación", "EmbeddingsCache, classify_recommendations, classify_recommendations_en, verify_match_with_llm", "Clasifica recomendaciones contra marco por similitud coseno en numpy, con caché persistente y verificación opcional por LLM."],
    ], widths=[1.6, 2.5, 2.7])

    h(doc, "5. Convenciones de estado", 1)
    add_bullets(doc, [
        "Cada pestaña usa claves propias de st.session_state, por ejemplo tab1_results_df, document_extracted_tab2, tab3_results, doc_chat_docs, doc_chat_embeddings, classified_world_df y deep_analysis_df.",
        "Los resultados se mantienen durante la sesión y pueden limpiarse con botones de limpiar resultados.",
        "La carga de nuevos archivos se detecta con hash(uploaded_file.getvalue()) y reinicia el estado de extracción cuando corresponde.",
        "Excepción: Tab 4 detecta cambios de documentos por lista de nombres, no por hash de contenido.",
        "Los outputs descargables se generan en memoria con BytesIO; no se guardan automáticamente en disco.",
    ])

    h(doc, "5.1 Lógica especial de preguntas en dos partes", 2)
    p(doc, "parse_two_part_question separa preguntas compuestas y privilegia la Parte 2. analyze_question_with_llm_tab1 instruye al modelo a iniciar el razonamiento con 'Se identificaron 2 partes en esta pregunta.' y a usar la Parte 1 solo como contexto. _critic_impl repite la evaluación con el documento completo y _apply_critic_gate_and_render ajusta el veredicto con base en el conteo de elementos dedicados.")
    p(doc, "El sistema distingue evidencia [DEDICATED] de [FRAMING]. La evidencia de framing, como listas amplias de grupos o lenguaje general de inclusión, no puede justificar Partial o Yes por sí sola.")

    h(doc, "5.2 Lógica especial de temas transversales", 2)
    p(doc, "TRANSVERSAL_MATTERS es un diccionario en código que reconoce género, no discriminación, discapacidad, diálogo social y tripartismo, y sostenibilidad medioambiental mediante aliases normalizados sin acentos. Cuando detect_transversal_matter encuentra uno de esos temas en la pregunta, _critic_impl usa CRITIC_SCHEMA_TRANSVERSAL y _apply_transversal_gate_and_render en vez del esquema A-E general.")
    p(doc, "La salida de razonamiento para esos casos inicia con una línea de auditoría de criterios transversales: A corresponde a objetivo/producto/actividad y B a presupuesto. La regla aplicada es Yes=A+B, Partial=A o B, No=sin A ni B. Indicadores y metas no cuentan para la regla transversal.")

    h(doc, "6. Manejo de errores", 1)
    add_bullets(doc, [
        "Falta de archivos críticos: la app muestra st.error o st.warning y detiene el flujo de esa pestaña cuando el recurso es obligatorio.",
        "Falta de OPENAI_API_KEY: la app muestra warning global y las funciones que requieren OpenAI devuelven error o advertencia.",
        "Errores de análisis por IA: se capturan y se devuelven en columnas Error o mensajes de chat.",
        "Respuestas JSON inválidas: varios flujos intentan limpiar fences de Markdown; si falla el parseo, se devuelve score 0 o score por defecto con error explicativo.",
        "Rate limits: algunas evaluaciones reintentan con backoff; en otros casos se reduce concurrencia por diseño.",
    ])

    h(doc, "7. Pruebas y verificación recomendadas", 1)
    add_bullets(doc, [
        "Ejecutar test_two_part_parsing.py al cambiar parse_two_part_question o la lógica de preguntas en dos partes.",
        "Probar manualmente una pregunta por cada tema transversal para confirmar detección por alias y aplicación de la regla A/B.",
        "Probar manualmente un DOCX con Heading 1/2 y tablas en Tabs 1, 2 y 3.",
        "Probar chat con TXT pequeño y DOCX grande para validar los dos caminos: contexto completo y RAG; confirmar que cambiar contenido con el mismo nombre de archivo requiere reiniciar sesión o renombrar el archivo.",
        "Probar Tab 5 con una muestra limitada y LLM Verification apagado y encendido.",
        "Verificar que cada descarga XLSX/ZIP abre correctamente y conserva columnas de evidencia.",
    ])

    h(doc, "8. Deuda técnica visible", 1)
    add_bullets(doc, [
        "Hay bloques comentados extensos de versiones anteriores; no son funcionales pero dificultan mantenimiento.",
        "La inicialización de datos y embeddings aparece dos veces en el bloque principal.",
        "Tab 4 promete trazabilidad por sección/página, pero la extracción de DOCX para chat solo lee párrafos con python-docx; no garantiza tablas, páginas ni metadatos jerárquicos.",
        "Tabs 5 y 6 tienen un selector de modelo para análisis profundo que no controla el modelo efectivo; run_row_analysis usa gpt-4o-mini.",
        "Algunos mensajes de error de Tab 3 todavía mencionan PRODOC_rubric.xlsx aunque la rúbrica activa es Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx.",
        "Existe un bloque heredado comentado de Tab 2 que también carga gender_rubric; no es el flujo activo. La implementación activa sí carga la hoja rubric_gender_ y la expone como Integración del Enfoque de Género.",
        "La lista de aliases de TRANSVERSAL_MATTERS está definida en código; si ILO cambia terminología oficial o agrega temas, debe actualizarse ese diccionario y probarse Tab 1.",
        "El archivo lessons_metadata.pt es esperado por load_lessons_embeddings, pero no aparece en el inventario local observado; si esa funcionalidad se activa, debe suministrarse.",
        "La app no separa claramente capa de interfaz, dominio y servicios externos; cambios grandes deben hacerse con pruebas manuales por pestaña.",
    ])

    h(doc, "9. Código del backend de los asistentes GPT", 1)
    p(doc, "El backend que sirve a los tres GPTs es independiente de oli_v6_deploy.py: no importa Streamlit ni comparte estado con la app. Los motores de evaluación se extrajeron a módulos core sin dependencia de interfaz, de modo que el mismo código pueda ejecutarse desde Streamlit o desde la API.")

    h(doc, "9.1 Archivos del backend", 2)
    add_table(doc, ["Archivo", "Rol"], [
        ["gpt_action_api.py", "Aplicación FastAPI: autenticación por cabecera, creación y sondeo de trabajos, descarga del DOCX desde OpenAI, endpoints /health y /privacy."],
        ["tab1_v3_core.py", "Motor de valoración preliminar v3: carga la rúbrica de 76 criterios, construye prompts, ejecuta la evaluación y arma el XLSX."],
        ["tab1_v3.py", "Envoltura Streamlit del mismo motor, usada por el Tab 7 experimental de la app."],
        ["tab2_core.py", "Motor de atributos específicos: tres rúbricas seleccionables, escala 1-5, consolidación por estabilidad."],
        ["sustainability_core.py", "Motor de sostenibilidad: 28 indicadores en tres dimensiones, escala 0-3, consolidación por estabilidad."],
        ["stability.py", "Módulo compartido de repetición y consolidación: ejecuta cada ítem N veces, colapsa a la moda y calcula el porcentaje de estabilidad."],
        ["openapi_gpt_action_v3.yaml", "Especificación OpenAPI que se pega en la Action del GPT de valoración."],
        ["openapi_gpt_action_tab2.yaml", "Especificación OpenAPI del GPT de atributos específicos."],
        ["openapi_gpt_action_sustainability.yaml", "Especificación OpenAPI del GPT de sostenibilidad."],
        ["Dockerfile.gpt-action", "Imagen Docker del backend sobre Python 3.10-slim."],
        ["docker-compose.gpt-action.yml", "Ejecución local del backend en contenedor."],
        ["requirements.gpt-action.txt", "Dependencias del backend, separadas de requirements.txt de Streamlit."],
        ["render.yaml", "Blueprint de Render: servicio, runtime, health check y variables de entorno."],
    ], widths=[2.3, 4.5])

    h(doc, "9.2 Ejecución local del backend", 2)
    codeblock(doc, [
        "pip install -r requirements.gpt-action.txt",
        "export OPENAI_API_KEY=...",
        "export ILO_GPT_ACTION_API_KEY=...",
        "uvicorn gpt_action_api:app --reload --port 8000",
        "",
        "# o en contenedor:",
        "docker compose -f docker-compose.gpt-action.yml up",
    ])
    p(doc, "El proceso debe ejecutarse desde el directorio que contiene los archivos de rúbrica, porque los motores los resuelven por ruta relativa.")

    h(doc, "9.3 Mapa de funciones del backend", 2)
    add_table(doc, ["Función o endpoint", "Responsabilidad"], [
        ["require_api_key", "Compara la cabecera X-API-Key contra ILO_GPT_ACTION_API_KEY; protege todos los endpoints de trabajos."],
        ["POST /v3/jobs, /attributes/jobs, /sustainability/jobs", "Registran un trabajo, devuelven su identificador y lanzan la evaluación en segundo plano."],
        ["GET /{familia}/jobs/{id}", "Devuelve el estado del trabajo: queued, running, succeeded o failed, con progreso cuando está disponible."],
        ["GET /{familia}/jobs/{id}/result", "Devuelve el resumen estructurado y el XLSX codificado en base64."],
        ["GET /health", "Sonda de disponibilidad usada por Render. No requiere autenticación."],
        ["GET /privacy", "Política de privacidad en HTML. Requisito de ChatGPT para publicar un GPT con Action. No requiere autenticación."],
        ["stability.evaluate_with_stability", "Ejecuta cada ítem repeats veces en paralelo y delega la consolidación al agregador del motor."],
        ["stability.aggregate_runs", "Colapsa las corridas al valor modal y calcula estabilidad, distribución y deriva."],
        ["results_to_xlsx_bytes", "Presente en los tres motores; construye el XLSX final con xlsxwriter."],
    ], widths=[2.5, 4.3])

    h(doc, "9.4 Deuda técnica del backend", 2)
    add_bullets(doc, [
        "El diccionario JOBS es almacenamiento en memoria del proceso: no sobrevive a reinicios ni admite más de una instancia del servicio.",
        "No hay expiración de trabajos: un proceso de larga vida acumula resultados en memoria indefinidamente.",
        "Los tres motores duplican la función results_to_xlsx_bytes con variaciones menores; podrían converger si el formato de salida se unifica.",
        "Las rúbricas se resuelven por ruta relativa al directorio de ejecución, lo que acopla el proceso a su working directory.",
        "stability.py incluye un _demo() ejecutable como autocomprobación; es el único código del backend con verificación automática.",
    ])


def doc3(doc):
    h(doc, "1. Modelo de datos", 1)
    p(doc, "La aplicación no utiliza una base de datos relacional. El modelo de datos es file-based: Excel, CSV delimitado por pipe, tensores PyTorch y cachés pickle locales.")
    p(doc, "Las relaciones son lógicas y se realizan en memoria con pandas. No hay migraciones, constraints SQL ni control transaccional.")

    h(doc, "2. Inventario de archivos de datos", 1)
    add_table(doc, ["Archivo", "Cargado por", "Uso"], [
        ["df_complete_all_full.xlsx", "load_data", "Base principal. Crea index_df desde ID_Recomendacion, normaliza nombres de columnas y dimension/subdim."],
        ["df_split_actions.xlsx", "load_data", "Base raw con Recommendation date; aporta year e index_df para completar faltantes."],
        ["analyzed_recommendations_plans_v5.csv", "load_extended_data", "CSV con sep='|'. Agrega análisis adicional, tags, scores y métricas para visualizaciones."],
        ["emb_Recomm_rec_cl_4.pt", "load_embeddings", "Matriz de embeddings de recomendaciones; dependencia activa de arranque porque load_embeddings se ejecuta aunque la búsqueda visible esté comentada."],
        ["Recommendation_RAG_Metadata.pt", "load_embeddings", "Metadatos alineados a emb_Recomm_rec_cl_4.pt; dependencia activa de arranque."],
        ["emb_LL_ll_cl_4.pt", "load_lessons_embeddings", "Embeddings para lecciones aprendidas; helper no conectado a una pestaña activa."],
        ["lessons_metadata.pt", "load_lessons_embeddings", "Metadatos esperados para lecciones aprendidas; helper no conectado a una pestaña activa."],
        ["Appraisal Checklist_2025 es-419.xlsx", "load_appraisal_questions", "Hoja rubric; requiere Pregunta_Realizada y usa Tema para renumerar preguntas."],
        ["TRANSVERSAL_MATTERS", "Tab 1", "Diccionario en código, no archivo externo. Reconoce género, no discriminación, discapacidad, diálogo social y tripartismo, y sostenibilidad medioambiental para aplicar regla A/B."],
        ["Rubricas_6ago2025.xlsx", "Tab 2", "La UI activa evalúa rubric_parteval, rubric_gender_ y rubric_TJ_TJ. Otras hojas existen o se cargan en código heredado, pero no son opciones activas."],
        ["Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "Tab 3", "Hoja rubric; fuente activa de criterios de sostenibilidad."],
        ["Recommendations_World.xlsx", "Tabs 5 y 6", "Dataset mundial para clasificación; puede reemplazarse por upload XLSX con Recommendation description."],
        ["Frame_Recommendations_English.xlsx", "Tabs 5 y 6", "Marco de referencia; debe contener texto_merged, dimension y subdim."],
        ["embeddings_cache.pkl", "EmbeddingsCache", "Caché local persistente de embeddings generados en clasificación de Tabs 5 y 6; no está cifrado ni verificado."],
        ["analysis_cache.pkl", "AnalysisCache", "Caché local persistente de análisis profundo de recomendaciones; no está cifrado ni verificado."],
    ], widths=[2.5, 1.7, 2.8])

    h(doc, "3. Relaciones lógicas", 1)
    add_bullets(doc, [
        "df_complete_all_full.xlsx y df_split_actions.xlsx se alinean por index_df.",
        "analyzed_recommendations_plans_v5.csv se fusiona con df por index_df cuando hay columnas nuevas.",
        "emb_Recomm_rec_cl_4.pt debe mantener el mismo orden que Recommendation_RAG_Metadata.pt para el índice de recomendaciones cargado en memoria.",
        "Recommendations_World.xlsx se clasifica contra cada fila de Frame_Recommendations_English.xlsx usando texto_merged como definición semántica; este flujo usa similitud coseno con numpy, no FAISS.",
        "Tabs 5 y 6 deduplican embeddings por texto de recomendación, pero el resultado vuelve al nivel de fila para conservar atributos repetidos.",
        "Los datasets de recomendaciones pueden tener filas duplicadas por múltiples atributos; Tabs 5 y 6 muestran conteos de registros y recomendaciones únicas.",
    ])

    h(doc, "4. Transformaciones principales", 1)
    add_table(doc, ["Transformación", "Detalle"], [
        ["Normalización de columnas", "Los espacios y puntos se reemplazan por guion bajo en load_data."],
        ["Años", "Recommendation_date y Recommendation date se convierten con pandas; años previos a 2018 en analyzed_df se fijan a 2018."],
        ["Categorías", "prepare_additional_data estandariza campos categóricos, listas/tags y clasificaciones cuando existen."],
        ["FAISS", "El índice de recomendaciones se construye al inicio. Tab 4 reconstruye un IndexFlatIP por pregunta para RAG de documentos grandes, con fragmentos ya embebidos y guardados en sesión."],
        ["Temas transversales", "En Tab 1, las preguntas detectadas por TRANSVERSAL_MATTERS omiten A-E y se califican con A=objetivo/producto/actividad y B=presupuesto; Yes requiere A+B, Partial requiere A o B, No requiere ausencia de ambos."],
        ["Clasificación", "Cada recomendación única se embebe una vez, se normaliza y se compara por similitud coseno contra el marco usando numpy. La asignación base usa top 3; la verificación opcional por LLM reranquea top 10; dimensiones/subdimensiones secundarias se conservan si similitud >= 0.60."],
        ["Exportación", "Los resultados se escriben a XLSX con columnas ordenadas y strings_to_urls desactivado."],
    ], widths=[2.0, 4.8])

    h(doc, "5. Campos mínimos por flujo", 1)
    add_table(doc, ["Flujo", "Campos mínimos requeridos"], [
        ["Valoración preliminar", "Appraisal Checklist: Pregunta_Realizada; se recomienda Tema. Las preguntas de temas transversales se reconocen por aliases en código."],
        ["Atributos específicos", "Rubricas_6ago2025: hojas activas rubric_parteval, rubric_gender_ y rubric_TJ_TJ, con Indicador, Dimensión y niveles de desempeño."],
        ["Sostenibilidad", "Rúbrica activa: columnas de dimensión, criterio, indicador y niveles según hoja rubric."],
        ["Clasificación", "Recommendations_World: Recommendation description. Frame: texto_merged, dimension, subdim."],
        ["Análisis profundo", "Recommendation description y Action plan; Comments mejora el análisis si existe."],
        ["Resumen ejecutivo", "Recommendation description; usa Management response, Comments y Action plan si existen."],
    ], widths=[2.0, 4.8])

    h(doc, "6. Respaldo y actualización de datos", 1)
    add_bullets(doc, [
        "Antes de reemplazar un Excel, CSV o tensor .pt, conservar una copia fechada fuera del directorio runtime.",
        "Después de reemplazar un archivo de datos, reiniciar la app o limpiar caché de Streamlit.",
        "Si se regeneran embeddings, regenerar también el archivo de metadatos alineado; no mezclar tensores y metadatos de corridas distintas.",
        "Si cambia el esquema de Recommendations_World.xlsx o Frame_Recommendations_English.xlsx, validar Tabs 5 y 6 con una muestra pequeña antes del procesamiento completo.",
        "No editar manualmente embeddings_cache.pkl o analysis_cache.pkl; si se sospecha corrupción, borrar el archivo con la app detenida para que se regenere.",
        "No asumir que emb_LL_ll_cl_4.pt o lessons_metadata.pt están en uso operativo si la función de lecciones no se reconecta a una pestaña activa.",
    ])

    h(doc, "7. Datos y rúbricas de los asistentes GPT", 1)
    p(doc, "El backend de los GPTs es igualmente file-based: las rúbricas se leen desde el sistema de archivos de la imagen Docker en cada arranque. No hay base de datos ni almacenamiento persistente de documentos ni de resultados.")

    h(doc, "7.1 Rúbricas consumidas por el backend", 2)
    add_table(doc, ["Archivo", "Motor que lo consume", "Contenido"], [
        ["Rubrica_Tab1_Detallada_Full_v3.xlsx", "tab1_v3_core.py", "76 criterios de valoración preliminar en 5 secciones, con subsecciones 1.1 a 5.2 y metadatos de subjetividad por criterio."],
        ["Rubricas_6ago2025.xlsx", "tab2_core.py", "Tres hojas de rúbrica: rubric_parteval (metodologías participativas), rubric_gender_ (género) y rubric_TJ_TJ (transición justa moderna)."],
        ["Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "sustainability_core.py", "28 indicadores en tres dimensiones del ciclo: Diseño (6), Implementación (10) y Pre-Cierre (12)."],
    ], widths=[2.3, 1.7, 2.8])
    p(doc, "sustainability_core.py resuelve su rúbrica con un glob tolerante a la tilde del nombre de archivo, porque la codificación del carácter acentuado varía entre sistemas de archivos. Al reemplazar ese archivo debe conservarse el patrón del nombre.")

    h(doc, "7.2 Modelo de estado de trabajos", 2)
    p(doc, "gpt_action_api.py mantiene un diccionario JOBS en memoria, con acceso serializado por threading.Lock. Cada entrada representa una evaluación en curso o terminada.")
    add_table(doc, ["Campo", "Contenido"], [
        ["job_id", "Identificador generado al crear el trabajo; es la clave del diccionario."],
        ["status", "queued, running, succeeded o failed."],
        ["progress", "Avance informado por el motor durante la evaluación, cuando está disponible."],
        ["result", "Resumen estructurado más el XLSX en base64, disponible sólo cuando status es succeeded."],
        ["error", "Mensaje de fallo, disponible sólo cuando status es failed."],
    ], widths=[1.6, 5.2])
    p(doc, "Este estado es volátil por diseño: un reinicio o redespliegue del servicio lo borra por completo. No constituye un registro histórico y no debe usarse como fuente de auditoría. El registro auditable de una evaluación es el XLSX que descarga el usuario.")

    h(doc, "7.3 Actualización de rúbricas en el backend", 2)
    add_numbered(doc, [
        "Reemplazar el XLSX en el repositorio, conservando el nombre de archivo y la estructura de columnas y hojas.",
        "Verificar la carga localmente ejecutando el backend y lanzando una evaluación de prueba acotada.",
        "Hacer push a la rama main: Render reconstruye la imagen y redespliega automáticamente.",
        "Confirmar con GET /health que el servicio quedó disponible, y con una evaluación corta que la rúbrica nueva se aplica.",
    ])
    p(doc, "No basta con subir el archivo al servidor: la rúbrica forma parte de la imagen Docker, de modo que cualquier cambio exige reconstrucción y redespliegue.")


def doc4(doc):
    h(doc, "1. Configuración obligatoria", 1)
    add_table(doc, ["Elemento", "Valor esperado"], [
        ["Comando local", "streamlit run oli_v6_deploy.py"],
        ["Variable de IA", "OPENAI_API_KEY disponible como variable de entorno."],
        ["Modo producción", "STREAMLIT_ENV=production cuando las rutas se toman de st.secrets."],
        ["Directorio de ejecución", "Directorio que contiene los Excel, CSV y .pt requeridos, salvo que se usen rutas de producción."],
        ["Dependencias", "Instaladas desde requirements.txt."],
    ], widths=[2.0, 4.8])

    h(doc, "2. Secretos y rutas", 1)
    p(doc, "El código lee OPENAI_API_KEY con os.getenv. En Streamlit Cloud u otro hosting, la clave debe exponerse como variable de entorno efectiva para el proceso Python.")
    add_table(doc, ["Nombre", "Cuándo se usa", "Contenido esperado"], [
        ["df_path", "Producción", "Ruta accesible a df_complete_all_full.xlsx."],
        ["df_raw_path", "Producción", "Ruta accesible a df_split_actions.xlsx."],
        ["embeddings_path", "Producción", "Ruta accesible a emb_Recomm_rec_cl_4.pt."],
        ["structured_embeddings_path", "Producción", "Ruta accesible a Recommendation_RAG_Metadata.pt."],
        ["analyzed_recommendations_path", "Producción opcional", "Ruta accesible a analyzed_recommendations_plans_v5.csv."],
        ["lessons_embeddings_path", "Producción para helper no activo", "Ruta accesible a emb_LL_ll_cl_4.pt si se reconecta la función de lecciones."],
        ["structured_lessons_path", "Producción para helper no activo", "Ruta accesible a lessons_metadata.pt si se reconecta la función de lecciones."],
    ], widths=[2.0, 1.8, 3.0])

    h(doc, "3. Control de acceso", 1)
    p(doc, "oli_v6_deploy.py no implementa usuarios, roles, contraseñas, permisos por funcionalidad, auditoría, aislamiento por tenant, escaneo de archivos cargados ni redacción de secretos. El control debe aplicarse fuera de la app, en Streamlit Cloud, servidor institucional, proxy, red privada o mecanismo equivalente.")
    add_bullets(doc, [
        "Solo usuarios autorizados deben poder abrir la aplicación.",
        "El repositorio y los archivos de datos deben tener permisos separados: lectura para operación, escritura solo para mantenimiento.",
        "El acceso a OPENAI_API_KEY debe limitarse al entorno de ejecución y administradores técnicos.",
        "Las descargas generadas por la app son responsabilidad del usuario que las produce.",
    ])

    h(doc, "4. Seguridad de datos", 1)
    add_bullets(doc, [
        "Los documentos subidos se procesan en memoria y mediante archivos temporales; los temporales se eliminan al finalizar la extracción cuando el flujo llega a limpieza.",
        "El contenido de documentos, preguntas y evidencias puede enviarse a OpenAI para generar evaluaciones y respuestas.",
        "st.session_state conserva texto extraído y resultados durante la sesión activa.",
        "Los archivos descargables XLSX/ZIP pueden contener citas textuales, razonamientos y evidencia del documento fuente.",
        "Los cachés pickle locales, incluidas embeddings_cache.pkl y analysis_cache.pkl, no están cifrados ni verificados por integridad.",
        "La aplicación no cifra salidas locales ni implementa retención de datos; estas políticas deben definirse en el entorno institucional.",
    ])

    h(doc, "4.1 Capas de caché activas", 2)
    add_table(doc, ["Capa", "Ubicación", "Riesgo operativo"], [
        ["st.cache_data", "Memoria/proceso Streamlit", "Puede mantener archivos, rúbricas o datos anteriores hasta limpiar caché o reiniciar."],
        ["Tab 4", "st.session_state", "Guarda documentos, fragmentos y embeddings por sesión; cambios con el mismo nombre de archivo pueden no detectarse."],
        ["Tabs 5 y 6", "embeddings_cache.pkl", "Persistente local, sin cifrado ni control de integridad."],
        ["Análisis profundo", "analysis_cache.pkl", "Persistente local, sin cifrado ni control de integridad."],
        ["SimpleHierarchicalStore", "~/document_store/embedding_cache.pkl", "Helper heredado para evaluación jerárquica; no es el mecanismo principal de Tab 4."],
    ], widths=[1.7, 2.3, 2.8])

    h(doc, "5. Parámetros de operación segura", 1)
    add_table(doc, ["Parámetro", "Valor en código", "Impacto"], [
        ["MAX_WORKERS global para valoración preliminar", "48", "Alta concurrencia para preguntas; puede exigir límites de API suficientes."],
        ["MAX_WORKERS Tab 2", "3", "Reduce riesgo de rate limit en evaluación por rúbricas."],
        ["MAX_WORKERS Tab 3", "8", "Procesa criterios de sostenibilidad en paralelo."],
        ["Chat RAG", ">100,000 caracteres", "Evita enviar documentos grandes completos en cada consulta."],
        ["Token limit", "110,000 tokens", "Recorta documentos muy extensos antes de llamar al modelo."],
        ["Clasificación batch", "20 textos por lote lógico", "Controla generación de embeddings de recomendaciones únicas."],
    ], widths=[2.2, 1.6, 3.0])

    h(doc, "6. Lista de verificación antes de producción", 1)
    add_bullets(doc, [
        "Confirmar que OPENAI_API_KEY existe en el proceso runtime y que la cuenta tiene presupuesto suficiente.",
        "Confirmar que todos los archivos de datos obligatorios abren con pandas/torch en el entorno de producción.",
        "Configurar acceso privado a la app en el hosting.",
        "Ejecutar una prueba por pestaña con documentos no sensibles.",
        "Verificar que las descargas se generan, se abren y no exponen más información de la necesaria.",
        "Documentar internamente quién puede reemplazar archivos de datos y quién puede desplegar código.",
    ])

    h(doc, "7. Respuesta a incidentes", 1)
    add_numbered(doc, [
        "Deshabilitar el acceso externo a la app desde el hosting o proxy.",
        "Revocar y rotar OPENAI_API_KEY si se sospecha exposición.",
        "Conservar copias de archivos descargados o logs disponibles solo si la política institucional lo permite.",
        "Restaurar archivos de datos desde respaldo verificado.",
        "Reactivar la app después de una prueba completa de Tabs 1 a 6.",
    ])

    h(doc, "8. Configuración y seguridad de los asistentes GPT", 1)
    p(doc, "El backend FastAPI que sirve a los tres GPTs tiene su propio modelo de configuración y de control de acceso, distinto del de la app Streamlit. A diferencia de oli_v6_deploy.py, este backend sí implementa autenticación.")

    h(doc, "8.1 Variables de entorno del backend", 2)
    add_table(doc, ["Variable", "Uso", "Dónde se configura"], [
        ["OPENAI_API_KEY", "Autenticación del cliente OpenAI para todas las llamadas de evaluación.", "Panel de Render, marcada como sync:false en render.yaml para no versionarla."],
        ["ILO_GPT_ACTION_API_KEY", "Valor esperado de la cabecera X-API-Key; protege todos los endpoints de trabajos.", "Panel de Render, marcada como sync:false; el mismo valor se carga en la Action de cada GPT."],
    ], widths=[2.0, 2.8, 2.0])
    p(doc, "Ambas variables deben existir en el proceso runtime. render.yaml las declara sin valor, de modo que Render las solicita en el primer despliegue y nunca quedan en el repositorio.")

    h(doc, "8.2 Control de acceso en tres capas", 2)
    add_table(doc, ["Capa", "Mecanismo", "Qué protege"], [
        ["Acceso al GPT", "Enlace privado de ChatGPT: sólo abre quien tiene el enlace, y requiere cuenta de ChatGPT.", "Quién puede usar el asistente."],
        ["Acceso a la API", "Cabecera X-API-Key comparada contra ILO_GPT_ACTION_API_KEY en require_api_key.", "Quién puede lanzar evaluaciones contra el backend."],
        ["Endpoints abiertos", "/health y /privacy no requieren autenticación, por requisito de Render y de ChatGPT respectivamente.", "No exponen datos: sólo estado del servicio y texto de política."],
    ], widths=[1.6, 3.0, 2.2])
    add_bullets(doc, [
        "La distribución de los enlaces de los GPTs es el control de acceso efectivo para las personas usuarias: debe gestionarse como se gestionaría una lista de distribución.",
        "ILO_GPT_ACTION_API_KEY es un secreto compartido entre el backend y las tres Actions; rotarla obliga a actualizar los tres GPTs.",
        "No hay usuarios, roles ni auditoría por persona: el backend no distingue quién lanzó cada evaluación.",
    ])

    h(doc, "8.3 Seguridad de datos en el flujo GPT", 2)
    add_bullets(doc, [
        "El documento se sube primero a los servidores de OpenAI, dentro de la conversación de ChatGPT, y luego el backend lo descarga para procesarlo. Ambos tránsitos deben considerarse al clasificar la sensibilidad del documento.",
        "El texto extraído y los fragmentos de evidencia se envían a la API de OpenAI durante la evaluación.",
        "El backend no persiste el documento: lo procesa y lo descarta al terminar el trabajo.",
        "Los resultados viven en memoria del proceso hasta el reinicio; no hay base de datos ni registro histórico.",
        "El XLSX entregado contiene citas textuales del documento original y debe manejarse con la misma confidencialidad que el documento fuente.",
        "La conversación de ChatGPT queda en la cuenta del usuario, con el documento adjunto: la retención depende de la configuración de la cuenta y del plan de ChatGPT usado.",
    ])

    h(doc, "8.4 Parámetros de operación del backend", 2)
    add_table(doc, ["Parámetro", "Valor", "Impacto"], [
        ["MAX_WORKERS tab1_v3_core", "48", "Alta concurrencia sobre 76 criterios, multiplicada por las 10 corridas de estabilidad; exige límites de API holgados en la cuenta OpenAI."],
        ["MAX_WORKERS tab2_core", "8", "Concurrencia moderada, multiplicada por las 5 corridas de estabilidad."],
        ["MAX_WORKERS sustainability_core", "8", "Concurrencia moderada, multiplicada por las 5 corridas de estabilidad."],
        ["STABILITY_REPEATS tab1_v3_core", "10", "Cada criterio se evalúa diez veces; una valoración completa implica del orden de 760 llamadas al modelo. Es el flujo más caro de los tres."],
        ["STABILITY_REPEATS tab2 y sostenibilidad", "5", "Cada criterio se evalúa cinco veces; multiplica por cinco el costo y la duración de esas dos evaluaciones."],
        ["STABILITY_THRESHOLD_PCT tab1_v3_core", "80.0", "Por debajo de ese porcentaje de coincidencia entre corridas, el criterio se marca como inestable y merece revisión humana."],
        ["Modelo de los tres motores", "gpt-5-mini", "Costo por evaluación dominado por este modelo y por el número de repeticiones; tab1_v3 varía el esfuerzo de razonamiento según la subjetividad del criterio."],
        ["Plan de Render", "free", "El servicio se suspende por inactividad; la primera petición tras la suspensión sufre latencia de arranque en frío."],
    ], widths=[2.0, 1.2, 3.6])

    h(doc, "8.5 Lista de verificación del backend antes de producción", 2)
    add_bullets(doc, [
        "Confirmar que OPENAI_API_KEY e ILO_GPT_ACTION_API_KEY están cargadas en Render y que la cuenta OpenAI tiene presupuesto.",
        "Confirmar que GET /health responde correctamente.",
        "Confirmar que los tres archivos de rúbrica están presentes en la imagen desplegada.",
        "Lanzar una evaluación corta por cada uno de los tres GPTs con un documento no sensible.",
        "Verificar que la Action de cada GPT tiene cargada la clave correcta y apunta al servidor correcto.",
        "Confirmar que /privacy responde: ChatGPT lo exige para mantener publicado un GPT con Action.",
        "Documentar internamente quién distribuye los enlaces de los GPTs y quién puede rotar la clave de API.",
    ])

    h(doc, "8.6 Respuesta a incidentes en la capa GPT", 2)
    add_numbered(doc, [
        "Si se sospecha uso indebido, rotar ILO_GPT_ACTION_API_KEY en Render: eso invalida de inmediato las tres Actions hasta que se actualicen.",
        "Si se sospecha exposición de la clave de OpenAI, rotarla en el panel de OpenAI y actualizarla en Render.",
        "Si el problema es de disponibilidad, revisar el estado del servicio en Render y los logs del contenedor.",
        "Si un enlace de GPT se difundió fuera del círculo previsto, republicar el GPT para generar un enlace nuevo y redistribuirlo.",
        "Restaurar el servicio redesplegando desde la rama main y verificar con una evaluación de prueba por asistente.",
    ])


def doc5(doc):
    h(doc, "1. Descripción funcional", 1)
    p(doc, "La aplicación ayuda a revisar documentos de proyecto y evaluación, consultar documentos cargados, clasificar recomendaciones y analizar respuestas institucionales. La interfaz activa tiene seis pestañas.")
    p(doc, "El sistema tiene además una segunda superficie de uso: tres asistentes GPT publicados en ChatGPT que replican los flujos de valoración de calidad, atributos específicos y sostenibilidad sin necesidad de abrir Streamlit. Se documentan en la sección 11.")

    h(doc, "2. Requisitos para usuarios", 1)
    add_bullets(doc, [
        "Usar navegador moderno con conexión estable.",
        "Contar con autorización institucional para cargar documentos en la aplicación.",
        "Subir DOCX con estilos de encabezado de Word cuando se use valoración, atributos o sostenibilidad.",
        "Revisar que los resultados de IA sean evidencia de apoyo y no sustituyan revisión técnica humana.",
    ])

    h(doc, "3. Navegación principal", 1)
    add_table(doc, ["Pestaña", "Cuándo usarla", "Salida principal"], [
        ["Valoración Preliminar de Calidad de Proyectos", "Revisión integral de documento de diseño de proyecto.", "ZIP con XLSX de preguntas, análisis por subsección/sección, plantilla y resumen TXT."],
        ["Diagnóstico de Atributos Específicos", "Evaluación focalizada de participación en evaluación, género o transición justa.", "ZIP con XLSX por rúbrica seleccionada."],
        ["Diagnóstico de Sostenibilidad del Proyecto", "Revisión de sostenibilidad de PRODOC o documento afín.", "ZIP con XLSX y gráficos de puntajes."],
        ["Pregúntale a tus Documentos", "Consulta conversacional sobre uno o varios DOCX/TXT.", "Respuesta en pantalla con memoria durante la sesión."],
        ["Clasificación de Recomendaciones", "Clasificar recomendaciones desde interfaz en español usando la columna Recommendation description.", "Treemaps, evolución temporal, análisis profundo XLSX y resumen ejecutivo XLSX."],
        ["Recommendation Classification", "Mismo flujo de clasificación en inglés.", "English outputs and downloads."],
    ], widths=[2.2, 2.5, 2.1])

    h(doc, "4. Uso de Tab 1: Valoración Preliminar", 1)
    add_numbered(doc, [
        "Descargar o revisar la rúbrica Appraisal Checklist_2025 es-419.xlsx desde la pestaña si es necesario.",
        "Subir un archivo DOCX con encabezados de Word.",
        "Presionar Extraer Documento y revisar la estructura extraída.",
        "Seleccionar secciones y filtros de preguntas si aplica.",
        "Presionar Analizar documento.",
        "Revisar métricas, tabla de resultados, razonamiento, evidencia y evaluación crítica.",
        "Descargar appraisal_checklist_results.zip.",
    ])
    p(doc, "Interpretación: las respuestas son Yes, No, Partial o Not Found. Para preguntas con sujeto específico no transversal, la app aplica el marco A-E: sub-objetivo/output, indicador, actividad, presupuesto y meta cuantificable. El total A-E puede ajustar automáticamente el veredicto.")
    p(doc, "Preguntas en dos partes: la Parte 2 es la pregunta evaluada y la Parte 1 solo enmarca. El razonamiento puede mostrar 'Se identificaron 2 partes en esta pregunta.' y la evidencia puede venir marcada como [DEDICATED] o [FRAMING]. Solo la evidencia dedicada puede sostener un Partial o Yes.")
    p(doc, "Excepción metodológica: cuando la pregunta corresponde a un tema transversal configurado (género, no discriminación, discapacidad, diálogo social y tripartismo, o sostenibilidad medioambiental), la app no usa A-E. Aplica la regla reducida A/B: A significa presencia del tema en objetivo, producto o actividad; B significa presupuesto o recursos asociados. Indicadores y metas no cuentan. La calificación es Yes si A y B están presentes, Partial si solo A o solo B está presente, y No si no aparece ninguno.")
    p(doc, "Distinguir género en Tab 1 y Tab 2: en Tab 1 género es un tema transversal con regla A/B; en Tab 2 'Integración del Enfoque de Género' es una rúbrica Excel de Rubricas_6ago2025.xlsx, hoja rubric_gender_, evaluada con puntaje 1-5 por criterios seleccionados.")
    p(doc, "La descarga principal de Tab 1 es appraisal_checklist_results.zip e incluye appraisal_checklist_results.xlsx, appraisal_checklist_rubric_template.xlsx cuando el archivo fuente está disponible, y appraisal_checklist_summary.txt. La estructura extraída se descarga aparte como estructura_documento_tab1_<archivo>.xlsx.")

    h(doc, "5. Uso de Tab 2: Diagnóstico de Atributos Específicos", 1)
    add_numbered(doc, [
        "Subir DOCX y extraer estructura.",
        "Seleccionar secciones del documento que se evaluarán.",
        "Elegir una o más rúbricas activas: participación durante evaluación/metodología, integración del enfoque de género o transición justa enfoque moderno.",
        "Seleccionar criterios dentro de cada rúbrica.",
        "Presionar Procesar y Evaluar.",
        "Revisar columnas Criterio, Dimensión, Score, Análisis, Evidencia, Error y Rúbrica.",
        "Descargar resultados_rubricas.zip.",
    ])
    p(doc, "La estructura extraída se puede descargar como estructura_documento_tab2_<archivo>.xlsx antes de evaluar.")

    h(doc, "6. Uso de Tab 3: Diagnóstico de Sostenibilidad", 1)
    add_numbered(doc, [
        "Descargar la rúbrica de sostenibilidad si se necesita revisar criterios.",
        "Subir DOCX, extraer estructura y seleccionar secciones.",
        "Seleccionar dimensiones y criterios.",
        "Presionar Procesar y Evaluar.",
        "Revisar resultados por Dimensión, Criterio, Indicador, Score, Análisis y Evidencia.",
        "Usar gráficos de promedio por dimensión y puntaje por criterio para priorizar mejoras.",
        "Descargar resultados_evaluacion_prodoc.zip.",
    ])
    p(doc, "La estructura extraída se puede descargar como estructura_documento_tab3_<archivo>.xlsx antes de evaluar.")

    h(doc, "7. Uso de Tab 4: Pregúntale a tus Documentos", 1)
    add_numbered(doc, [
        "Subir uno o más archivos DOCX o TXT.",
        "Confirmar que aparecen como documentos activos.",
        "Escribir preguntas específicas; pedir citas breves cuando se necesite trazabilidad, pero verificar manualmente sección/página porque Tab 4 no usa el extractor jerárquico enriquecido.",
        "Para documentos pequeños (hasta 100,000 caracteres totales), la app usa contexto completo recortado a 110,000 tokens; para grandes, usa RAG automáticamente.",
        "Mantener la sesión abierta si se necesita conservar la memoria del chat.",
    ])
    p(doc, "En RAG, la app usa fragmentos de 2,000 caracteres con solape de 300, recupera hasta 15 fragmentos por FAISS y hasta 5 por coincidencia de palabras clave, y envía solo los últimos 5 mensajes del chat. Si se reemplaza un archivo por otro con el mismo nombre, reiniciar sesión o renombrarlo para evitar estado anterior.")
    p(doc, "La respuesta debe basarse en los archivos cargados. La ruta RAG permite inferencias desde pistas contextuales; el usuario debe tratarlas como apoyo preliminar y verificarlas en el documento.")

    h(doc, "8. Uso de Tabs 5 y 6: Clasificación de Recomendaciones", 1)
    add_numbered(doc, [
        "Elegir archivos predeterminados o subir un XLSX con la columna Recommendation description.",
        "Aplicar filtros previos por ubicación, tiempo, temática, unidad técnica, fuente, progreso u otros campos disponibles.",
        "Revisar conteo de registros y recomendaciones únicas antes de iniciar.",
        "Decidir si activar LLM Verification: mayor precisión, más lento y más costoso.",
        "Presionar Iniciar Clasificación.",
        "Explorar treemaps por dimensión/subdimensión y evolución temporal por categorías. La asignación base usa top 3 por similitud coseno; LLM Verification reranquea top 10; dimensiones secundarias se conservan si similitud >= 0.60.",
        "Usar Herramientas Avanzadas de IA para análisis profundo o resumen ejecutivo sobre el subconjunto filtrado. Nota: el selector de modelo de análisis profundo no controla el modelo efectivo en el código actual; se usa gpt-4o-mini.",
        "Descargar analisis_profundo.xlsx, resumen_ejecutivo.xlsx, deep_analysis.xlsx o executive_summary.xlsx según el idioma de la pestaña.",
    ])

    h(doc, "9. Buenas prácticas operativas", 1)
    add_bullets(doc, [
        "Empezar con secciones o muestras pequeñas cuando se evalúa un documento nuevo.",
        "Verificar evidencias textuales antes de usar resultados en informes formales.",
        "Guardar descargas con nombre de proyecto, fecha y pestaña usada.",
        "No subir documentos que no estén autorizados para procesamiento por servicios externos de IA.",
        "Limpiar resultados o recargar la app al cambiar de caso de análisis.",
    ])

    h(doc, "10. Problemas frecuentes", 1)
    add_table(doc, ["Síntoma", "Causa probable", "Acción"], [
        ["No detecta secciones", "El DOCX no usa estilos Heading de Word.", "Corregir encabezados en Word y volver a extraer."],
        ["OPENAI API key not found", "OPENAI_API_KEY no está disponible.", "Configurar variable de entorno y reiniciar app."],
        ["La clasificación no inicia", "Falta Recommendations_World.xlsx o Frame_Recommendations_English.xlsx.", "Restaurar archivos o subir XLSX válido."],
        ["Resultados con Error", "Respuesta JSON inválida, rate limit o fallo de API.", "Reintentar con menos criterios/secciones o revisar cuota de API."],
        ["Descarga no contiene lo esperado", "Resultados anteriores persistieron en sesión.", "Usar Limpiar resultados y repetir el flujo."],
        ["Tab 4 responde sobre archivo anterior", "Se subió otro archivo con el mismo nombre y el estado se detecta por nombre.", "Renombrar el archivo o reiniciar la sesión de Streamlit."],
    ], widths=[2.0, 2.3, 2.5])

    h(doc, "11. Asistentes GPT en ChatGPT", 1)
    p(doc, "Además de la app Streamlit, tres de los flujos de evaluación están disponibles como asistentes GPT dentro de ChatGPT. No requieren instalar nada ni abrir Streamlit: se accede por enlace, se sube el DOCX en el chat y se recibe un XLSX descargable. Las rúbricas están cargadas en el servidor, de modo que el usuario nunca las sube.")
    p(doc, "Existe un manual de usuario dedicado a esta modalidad, Manual_Usuario_GPTs_OIT.docx, orientado a personas usuarias finales sin perfil técnico.")

    h(doc, "11.1 Enlaces y equivalencia con las pestañas de Streamlit", 2)
    add_table(doc, ["Asistente GPT", "Equivale a", "Enlace"], [
        ["ILO PRODOC Quality Appraisal", "Tab 1, con la rúbrica v3 de 76 criterios", GPT_V3],
        ["OIT - Diagnóstico de Atributos Específicos", "Tab 2", GPT_TAB2],
        ["OIT - Diagnóstico de Sostenibilidad del Proyecto", "Tab 3", GPT_SUS],
    ], widths=[1.9, 1.6, 3.3])
    p(doc, "Los Tabs 4, 5 y 6 de la app Streamlit no tienen equivalente en ChatGPT: el chat documental y la clasificación de recomendaciones siguen siendo exclusivos de la aplicación.")
    p(doc, "El acceso es por enlace privado y requiere cuenta de ChatGPT. Quien administre el servicio debe tratar la distribución de estos enlaces como el control de acceso efectivo.")

    h(doc, "11.2 Flujo de uso común a los tres asistentes", 2)
    add_numbered(doc, [
        "Abrir el enlace del asistente. Al escribir un saludo o pulsar un botón de inicio, el asistente se presenta y explica qué puede evaluar, sin consumir una evaluación.",
        "Adjuntar un único archivo .docx. Si se suben varios, el asistente pide elegir uno.",
        "Indicar el alcance en el mismo mensaje: sección o subsección para valoración de calidad, rúbrica temática para atributos específicos, dimensión del ciclo para sostenibilidad. Si no se indica, el asistente lo pregunta antes de empezar.",
        "Esperar mientras el asistente lanza el trabajo y consulta su estado. Una evaluación completa puede tardar varios minutos.",
        "Revisar el resumen en pantalla y descargar el XLSX, que es el registro auditable de la evaluación.",
    ])

    h(doc, "11.3 Alcance y escalas por asistente", 2)
    add_table(doc, ["Asistente", "Alcance seleccionable", "Escala"], [
        ["Valoración de Calidad (Tab 1 v3)", "Rúbrica completa, una o varias secciones (1 a 5), o subsecciones concretas (1.1 a 5.2)", "Yes / Partial / No / Not Found / N/A sobre 76 criterios"],
        ["Atributos Específicos (Tab 2)", "Una o varias de tres rúbricas: metodologías participativas, género, transición justa", "1 a 5 por criterio"],
        ["Sostenibilidad (Tab 3)", "Dimensión Diseño (PRODOCs), Implementación (informes de avance) o Pre-Cierre (cierre y evaluación final)", "0 a 3 por indicador"],
    ], widths=[1.8, 3.0, 2.0])
    p(doc, "Las escalas no son intercambiables: 0-3 en sostenibilidad y 1-5 en atributos específicos miden cosas distintas y no deben promediarse ni compararse directamente.")
    p(doc, "Los tres asistentes repiten cada evaluación y consolidan el resultado: 10 corridas por criterio en valoración de calidad, 5 en atributos específicos y 5 en sostenibilidad. El XLSX incluye el porcentaje de estabilidad, que indica cuántas de esas corridas coincidieron: un porcentaje bajo señala un criterio donde el modelo no fue consistente y la revisión humana es más necesaria.")
    p(doc, "El veredicto \"Not Found\" en valoración de calidad significa que el documento no contiene información suficiente para evaluar el criterio. Es distinto de \"No\": este último afirma que el criterio no se cumple, mientras que \"Not Found\" afirma que no se pudo determinar.")

    h(doc, "11.4 Botones de inicio configurados", 2)
    add_table(doc, ["Asistente", "Botones de inicio"], [
        ["Valoración de Calidad", "¿Qué puedes hacer y cómo empiezo? · Evalúa este PRODOC con la rúbrica completa · Evalúa solo la sección 3 (Marco de resultados) · ¿Qué secciones y subsecciones puedo filtrar?"],
        ["Atributos Específicos", "¿Qué rúbricas puedes aplicar? · Evalúa este documento con la rúbrica de género · Aplica la rúbrica de Transición Justa · Evalúa participación y género y compara resultados"],
        ["Sostenibilidad", "¿Qué dimensiones evalúas y cuál me corresponde? · Evalúa este PRODOC con la dimensión de Diseño · Es un informe de avance: aplica Implementación · Aplica la rúbrica completa de sostenibilidad"],
    ], widths=[1.7, 5.1])
    p(doc, "El texto de las instrucciones y de estos botones se mantiene en docs/gpt_onboarding_es.md. Al modificarlos en el editor de ChatGPT debe actualizarse también ese archivo, para que documentación y asistentes no se separen.")

    h(doc, "11.5 Problemas frecuentes en los asistentes GPT", 2)
    add_table(doc, ["Síntoma", "Causa probable", "Acción"], [
        ["La primera evaluación del día tarda mucho en arrancar", "El backend está en plan free de Render y se suspende por inactividad.", "Esperar el arranque en frío; es comportamiento esperado, no una falla."],
        ["El asistente responde que no puede autenticarse", "La clave de la Action no coincide con ILO_GPT_ACTION_API_KEY en el servidor.", "Verificar la clave en la Action del GPT y en el panel de Render."],
        ["La evaluación falla a mitad de camino", "Rate limit de la API, documento muy extenso o fallo transitorio del servicio.", "Reintentar con un alcance más acotado: una sección o una sola rúbrica."],
        ["El asistente pide subir la rúbrica", "Deriva del modelo respecto de sus instrucciones.", "Aclarar que use la rúbrica cargada en el servidor; si se repite, revisar las instrucciones del GPT."],
        ["Se perdió una evaluación en curso", "El servicio se reinició o redesplegó: el estado de trabajos vive en memoria.", "Volver a lanzar la evaluación. No hay recuperación posible del trabajo perdido."],
        ["Los tres asistentes fallan a la vez", "Interrupción del backend compartido.", "Verificar GET /health del servidor; es un problema de servicio, no del documento."],
    ], widths=[2.0, 2.3, 2.5])


def doc6(doc):
    h(doc, "1. Transferencia operativa", 1)
    p(doc, "La operación diaria debe quedar en manos del equipo ILO. El desarrollador original queda como soporte de último recurso para dudas técnicas que no puedan resolverse con esta documentación, revisión del código y pruebas locales.")
    add_table(doc, ["Rol", "Responsabilidad"], [
        ["Equipo ILO propietario", "Administrar acceso, datos, ejecución, resultados y decisiones metodológicas."],
        ["Administrador técnico ILO", "Gestionar despliegue, variables de entorno, archivos de datos y respaldos."],
        ["Usuarios analistas", "Ejecutar flujos, revisar evidencia, descargar resultados y aplicar juicio técnico."],
        ["Desarrollador original", "Soporte excepcional para arquitectura, debugging complejo o cambios que excedan mantenimiento ordinario."],
    ], widths=[2.1, 4.7])

    h(doc, "2. Contacto de último recurso", 1)
    add_bullets(doc, [
        "Canal heredado de la documentación previa: ageidv@gmail.com.",
        "Uso recomendado: solo después de reproducir el problema localmente, revisar logs/mensajes de la app y aislar la pestaña o función afectada.",
        "Incluir al contactar: fecha, pestaña, archivo de entrada no sensible o muestra, pasos para reproducir, mensaje de error y captura de pantalla si la política interna lo permite.",
    ])

    h(doc, "3. Rutina de mantenimiento", 1)
    add_table(doc, ["Frecuencia", "Actividad"], [
        ["Semanal durante operación activa", "Confirmar que la app abre, OPENAI_API_KEY funciona y un documento de prueba procesa en Tab 4."],
        ["Mensual", "Probar una muestra en Tabs 1, 2, 3, 5 y 6; verificar que descargas abren."],
        ["Antes de cada actualización de datos", "Respaldar Excel, CSV, .pt y cachés relevantes."],
        ["Después de cada actualización de código", "Ejecutar pruebas manuales por pestaña y test_two_part_parsing.py si cambió lógica de preguntas."],
        ["Cada 6-12 meses", "Revisar modelos OpenAI, costos, límites, dependencias y compatibilidad de Streamlit."],
    ], widths=[2.2, 4.6])

    h(doc, "4. Recomendaciones para futuras actualizaciones", 1)
    add_bullets(doc, [
        "Separar gradualmente el monolito en módulos: carga de datos, extracción documental, clientes IA, evaluación y UI.",
        "Eliminar bloques comentados obsoletos después de confirmar que no son necesarios.",
        "Consolidar la doble inicialización de load_extended_data y load_embeddings.",
        "Revisar si load_embeddings debe seguir ejecutándose en arranque si la búsqueda visible de recomendaciones permanece comentada.",
        "Corregir Tab 4 si ILO necesita citas por sección/página: debe usar extractor jerárquico, no solo python-docx paragraphs.",
        "Conectar el selector de modelo de análisis profundo a run_row_analysis o retirarlo de la UI.",
        "Agregar pruebas unitarias para parse_two_part_question, _apply_critic_gate_and_render, clasificación y exportación XLSX.",
        "Agregar pruebas unitarias para detect_transversal_matter y _apply_transversal_gate_and_render si la regla A/B se vuelve crítica para reportes institucionales.",
        "Crear un set pequeño de documentos de prueba no sensibles para validación reproducible.",
        "Documentar cambios de esquema de cada Excel junto con fecha y responsable ILO.",
        "Si se mantienen las variantes core y recommendations, sincronizar cambios desde oli_v6_deploy.py y probar cada variante por separado.",
    ])

    h(doc, "5. Recomendaciones de IA y costos", 1)
    add_bullets(doc, [
        "Mantener gpt-5-mini como modelo principal mientras sus resultados sean consistentes con validación humana.",
        "Usar LLM Verification en clasificación solo para subconjuntos donde la precisión adicional justifique costo y tiempo.",
        "Usar filtros previos y límites de filas antes de análisis profundo de recomendaciones.",
        "Conservar cachés embeddings_cache.pkl y analysis_cache.pkl cuando se procesan lotes recurrentes; eliminarlos solo si se sospecha corrupción o cambio sustantivo de datos.",
        "Revisar si el índice FAISS de recomendaciones cargado por load_embeddings sigue siendo necesario en la app principal, porque la interfaz activa de búsqueda de recomendaciones está comentada y las pestañas de clasificación no lo usan aunque la carga se ejecuta en arranque.",
        "Registrar internamente estimaciones de costo por flujo con generate_cost_report.py si se institucionaliza uso frecuente.",
    ])

    h(doc, "6. Criterios de aceptación de cambios", 1)
    add_bullets(doc, [
        "La app inicia sin errores con streamlit run oli_v6_deploy.py.",
        "Tab 1 procesa un DOCX de prueba, genera tabla y descarga ZIP.",
        "Tab 1 reconoce al menos un tema transversal y muestra la línea de auditoría de criterios transversales A/B en el razonamiento.",
        "Tab 2 evalúa al menos un criterio seleccionado y descarga ZIP.",
        "Tab 3 genera resultados y gráficos para una muestra.",
        "Tab 4 responde sobre un TXT pequeño y un DOCX grande.",
        "Tab 4 se prueba reemplazando un archivo por otro con el mismo nombre para confirmar el procedimiento operativo de reinicio o renombrado.",
        "Tabs 5 y 6 clasifican una muestra limitada y muestran treemaps.",
        "No aparecen campos por completar, credenciales reales ni rutas personales innecesarias en documentación de entrega.",
    ])

    h(doc, "7. Paquete mínimo de handoff", 1)
    add_bullets(doc, [
        "Código principal: oli_v6_deploy.py.",
        "Dependencias: requirements.txt.",
        "Documentación actualizada: los seis archivos DOCX en documentation.",
        "Datos y rúbricas activos listados en la documentación de base de datos.",
        "Instrucciones internas de despliegue del hosting elegido por ILO.",
        "Registro ILO de responsables de acceso, datos y despliegue.",
    ])

    h(doc, "8. Transferencia y mantenimiento de los asistentes GPT", 1)
    p(doc, "Los tres asistentes GPT publicados en ChatGPT y su backend FastAPI forman un entregable separado de la app Streamlit, con su propio ciclo de mantenimiento, sus propias credenciales y su propio modo de fallo.")

    h(doc, "8.1 Activos a transferir", 2)
    add_table(doc, ["Activo", "Dónde vive", "Quién debe controlarlo"], [
        ["Los tres GPTs publicados", "Cuenta de ChatGPT donde fueron creados", "Debe migrarse a una cuenta institucional de la OIT; mientras siga en una cuenta personal, la continuidad depende de esa cuenta."],
        ["Servicio backend", "Render.com, servicio ilo-prodoc-appraisal-v3", "Administrador técnico OIT."],
        ["OPENAI_API_KEY", "Variable de entorno en Render", "Administrador técnico OIT; determina a qué cuenta se factura el consumo."],
        ["ILO_GPT_ACTION_API_KEY", "Variable de entorno en Render y en las tres Actions", "Administrador técnico OIT."],
        ["Enlaces privados de los GPTs", "Distribución interna OIT", "Equipo propietario. Los tres están en \"Anyone with a link\": el enlace es una credencial al portador y la lista de distribución es el único control de acceso."],
        ["  · Valoración de Calidad (Tab 1 v3)", GPT_V3, "Enlace privado; requiere cuenta de ChatGPT."],
        ["  · Atributos Específicos (Tab 2)", GPT_TAB2, "Enlace privado; requiere cuenta de ChatGPT."],
        ["  · Sostenibilidad (Tab 3)", GPT_SUS, "Enlace privado; requiere cuenta de ChatGPT."],
        ["  · Backend compartido", GPT_BACKEND, "FastAPI en Render; sonda pública en /health."],
        ["Código y rúbricas", "Repositorio del proyecto", "Administrador técnico OIT."],
    ], widths=[1.8, 2.2, 2.8])
    p(doc, "La migración de los GPTs a una cuenta institucional es la acción de transferencia más importante y la que no puede posponerse indefinidamente: un GPT publicado desde una cuenta personal deja de estar disponible si esa cuenta se cierra o cambia de plan.")

    h(doc, "8.2 Rutina de mantenimiento de la capa GPT", 2)
    add_table(doc, ["Frecuencia", "Actividad"], [
        ["Semanal durante operación activa", "Confirmar que GET /health responde y lanzar una evaluación corta en uno de los tres asistentes."],
        ["Mensual", "Probar los tres asistentes con un documento de prueba y verificar que el XLSX se descarga y abre correctamente."],
        ["Mensual durante la fase piloto", "Revisar el consumo en el tablero de OpenAI y proyectar la fecha de agotamiento del presupuesto."],
        ["Al actualizar una rúbrica", "Reemplazar el XLSX, probar localmente, hacer push a main y verificar el redespliegue con una evaluación acotada."],
        ["Al modificar instrucciones o botones de inicio", "Actualizar también docs/gpt_onboarding_es.md para que la documentación no se separe de los asistentes."],
        ["Cada 6-12 meses", "Revisar modelos disponibles, costos, límites de la API y vigencia del plan de Render."],
    ], widths=[2.2, 4.6])

    h(doc, "8.3 Criterios de aceptación de cambios en la capa GPT", 2)
    add_bullets(doc, [
        "GET /health responde correctamente tras el despliegue.",
        "GET /privacy responde: ChatGPT lo exige para mantener publicado un GPT con Action.",
        "Cada uno de los tres asistentes completa una evaluación acotada de prueba y entrega el XLSX.",
        "El XLSX de atributos específicos y el de sostenibilidad incluyen la columna de estabilidad.",
        "Las tres Actions apuntan al servidor correcto y autentican con la clave vigente.",
        "Al saludar sin adjuntar documento, cada asistente se presenta en lugar de esperar en silencio.",
    ])

    h(doc, "8.4 Recomendaciones de continuidad y costo", 2)
    add_bullets(doc, [
        "Migrar los tres GPTs a una cuenta institucional de la OIT antes de ampliar el uso.",
        "Sustituir el almacenamiento de trabajos en memoria por Redis o equivalente si el uso deja de ser piloto: hoy un redespliegue pierde las evaluaciones en curso.",
        "Evaluar el paso a un plan de pago en Render si la latencia de arranque en frío afecta la experiencia de uso.",
        "Vigilar el costo de las repeticiones de estabilidad: 10 corridas por criterio en valoración de calidad y 5 en los otros dos asistentes. Una valoración completa implica del orden de 760 llamadas al modelo y es, con diferencia, el flujo más caro. Reducir STABILITY_REPEATS es la palanca más directa si el presupuesto aprieta, a costa de fiabilidad; bajarlo en tab1_v3_core es donde más ahorra.",
        "Usar claves de API separadas por asistente si se necesita atribuir el consumo a cada uno: el tablero de OpenAI desagrega por clave, no por GPT.",
        "Priorizar la evaluación por secciones o dimensiones frente a la rúbrica completa cuando el objetivo sea verificar un punto concreto: reduce costo y tiempo sin perder valor.",
        "Registrar el costo unitario observado durante el piloto para poder estimar un despliegue a mayor escala.",
    ])


def doc_en1(doc):
    h(doc, "1. Scope and source of truth", 1)
    p(doc, "This technical documentation describes only the main Streamlit application, oli_v6_deploy.py. The files oli_v6_deploy_core.py and oli_v6_deploy_recommendations.py are split variants and are not used as factual sources for this documentation.")
    p(doc, "The main app is a monolithic Streamlit file that includes document extraction, LLM evaluation, recommendation classification, visualization, caching, and export logic.")
    p(doc, "Since June 2026 the system has a second delivery surface: three custom GPTs published in ChatGPT, served by a separate FastAPI backend (gpt_action_api.py) deployed on Render. That layer is documented in section 9 and shares neither process nor state with the Streamlit app.")

    h(doc, "2. High-level architecture", 1)
    add_table(doc, ["Layer", "Implementation", "Responsibility"], [
        ["Interface", "Streamlit with st.tabs, st.session_state, st.cache_data, and ILO styling", "Document uploads, rubric selection, charts, result tables, and downloads."],
        ["Document processing", "python-docx, docx2python, BeautifulSoup, and enhanced heading extraction", "Turns DOCX/TXT content into text and structured extraction tables for Tabs 1-3."],
        ["Generative AI", "OpenAI SDK with OpenAI(api_key=os.getenv('OPENAI_API_KEY'))", "Evaluations, critique passes, synthesis, document chat, deep analysis, and executive summaries."],
        ["Vector search", "text-embedding-3-large, numpy, torch, and FAISS IndexFlatIP", "Loads a recommendation FAISS index at startup and uses FAISS for large-document RAG in Tab 4. Tabs 5 and 6 use numpy cosine similarity, not FAISS."],
        ["Data", "Excel files, pipe-separated CSV, PyTorch tensors, and local pickle caches", "File-based operational store; there is no SQL database."],
        ["Exports", "xlsxwriter, BytesIO, and zipfile", "Creates XLSX and ZIP outputs with answers, evidence, extracted structure, and reports."],
    ], widths=[1.3, 2.4, 3.1])

    h(doc, "3. Active functional components", 1)
    add_table(doc, ["Tab", "Visible name", "Main function"], [
        ["1", "Valoración Preliminar de Calidad de Proyectos", "Evaluates project design documents against Appraisal Checklist_2025 es-419.xlsx, including A-E critique, two-part question handling, and transversal A/B scoring."],
        ["2", "Diagnóstico de Atributos Específicos", "Evaluates selected document sections using active rubrics for evaluation participation/methodology, gender integration, and modern just transition."],
        ["3", "Diagnóstico de Sostenibilidad del Proyecto", "Evaluates sustainability with Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx and charts scores by dimension and criterion."],
        ["4", "Pregúntale a tus Documentos", "Chat over one or more DOCX/TXT files, using full context for small uploads and FAISS RAG for large uploads."],
        ["5", "Clasificación de Recomendaciones", "Spanish-language UI for classifying records with a required Recommendation description column against Frame_Recommendations_English.xlsx."],
        ["6", "Recommendation Classification", "English-language recommendation classification workflow using the same underlying logic as Tab 5."],
    ], widths=[0.7, 2.2, 4.1])

    h(doc, "4. Dependencies", 1)
    p(doc, "requirements.txt declares the operating dependency set, including Streamlit, pandas, numpy, faiss-cpu, OpenAI, torch, plotly, xlsxwriter, python-docx, docx2python, openpyxl, and tiktoken.")

    h(doc, "5. AI models and context limits", 1)
    add_bullets(doc, [
        "Main evaluation, critique, synthesis, and chat model: gpt-5-mini.",
        "Embedding model: text-embedding-3-large.",
        "Deep recommendation analysis shows a UI selector for gpt-4o-mini and gpt-4o, but the current code does not pass that selection into run_row_analysis; analyze_recommendation_plan_pair effectively uses gpt-4o-mini.",
        "truncate_to_token_limit keeps document context to about 110,000 tokens.",
        "Tab 4 uses full context for total uploaded text up to 100,000 characters; larger uploads are split into 2,000-character chunks with 300-character overlap and processed with FAISS RAG.",
    ])

    h(doc, "5.1 Tab 1 two-part questions", 2)
    p(doc, "parse_two_part_question detects compound questions through regex patterns such as statement-then-question, explicit separators, two question marks, and run-together clauses. When two parts are detected, Part 2 drives the answer and Part 1 is framing only.")
    add_bullets(doc, [
        "The first-pass reasoning is instructed to begin with 'Se identificaron 2 partes en esta pregunta.'.",
        "Evidence may be marked [DEDICATED] or [FRAMING]. Framing evidence alone cannot justify Partial or Yes.",
        "The critic re-checks the answer against the full document and the code layer mechanically applies the A-E gate.",
        "If the model fails to focus on Part 2, the row can receive Status = Partial even when the call completes.",
    ])

    h(doc, "5.2 Tab 1 transversal matters", 2)
    p(doc, "TRANSVERSAL_MATTERS recognizes five configured themes: gender, non-discrimination, disability, social dialogue and tripartism, and environmental sustainability. These questions bypass the A-E framework entirely.")
    add_table(doc, ["Transversal criterion", "Operational definition"], [
        ["A", "Operational presence of the theme in an objective, product/output, or activity."],
        ["B", "Budget, resources, or a budget line for a corresponding activity."],
    ], widths=[1.5, 5.3])
    p(doc, "The verdict is mechanical: Yes = A and B; Partial = A or B; No = neither A nor B. Indicators and targets do not count for transversal scoring. Gender in Tab 1 is this A/B transversal rule; gender in Tab 2 is a separate 1-5 Excel rubric from Rubricas_6ago2025.xlsx, sheet rubric_gender_.")

    h(doc, "5.3 Tab 4 document chat", 2)
    add_bullets(doc, [
        "DOCX chat extraction reads paragraph text with python-docx; it does not use the enhanced hierarchical extractor. Tables, page numbers, section metadata, and precise section/page citations are not reliable in Tab 4.",
        "For large uploads, chunk embeddings are stored in st.session_state. Each question rebuilds a FAISS IndexFlatIP, normalizes vectors, retrieves up to 15 vector chunks, then adds up to 5 keyword chunks.",
        "Only the last 5 chat messages are sent to the model.",
        "File changes are detected by filename list, not content hash. Replacing a file with another file of the same name can leave stale session state until the app session is reset or the file is renamed.",
        "The small-document path is stricter about using only the document. The RAG path allows inference from contextual clues, so any inferred answer must be manually verified.",
    ])

    h(doc, "6. Startup and cache", 1)
    add_bullets(doc, [
        "load_data() reads df_complete_all_full.xlsx and df_split_actions.xlsx.",
        "load_extended_data() optionally merges analyzed_recommendations_plans_v5.csv and stores analyzed_df in st.session_state.",
        "load_embeddings() reads emb_Recomm_rec_cl_4.pt and Recommendation_RAG_Metadata.pt and builds a FAISS index during startup. This remains an active runtime dependency even though the visible recommendation-search UI is commented.",
        "Tabs 5 and 6 use embeddings_cache.pkl for classification embeddings and analysis_cache.pkl for deep recommendation-plan analysis.",
        "Tab 4 stores uploaded document text, chunks, and chunk embeddings in st.session_state only.",
        "The main file initializes data and embeddings twice; the behavior is equivalent but should be consolidated in maintenance.",
    ])

    h(doc, "7. Deployment", 1)
    p(doc, "Run locally with:")
    codeblock(doc, "streamlit run oli_v6_deploy.py")
    p(doc, "OPENAI_API_KEY is read from the process environment. Production file paths are read from st.secrets only when STREAMLIT_ENV=production.")
    add_table(doc, ["Secret or variable", "Use"], [
        ["OPENAI_API_KEY", "OpenAI client authentication."],
        ["df_path", "Production path to df_complete_all_full.xlsx."],
        ["df_raw_path", "Production path to df_split_actions.xlsx."],
        ["embeddings_path", "Production path to emb_Recomm_rec_cl_4.pt."],
        ["structured_embeddings_path", "Production path to Recommendation_RAG_Metadata.pt."],
        ["analyzed_recommendations_path", "Optional production path to analyzed_recommendations_plans_v5.csv."],
        ["lessons_embeddings_path", "Path for an inactive lessons helper if that feature is reconnected."],
        ["structured_lessons_path", "Path for an inactive lessons helper if that feature is reconnected."],
    ], widths=[2.2, 4.6])

    h(doc, "8. Operational risks", 1)
    add_bullets(doc, [
        "The app has no built-in authentication, authorization, audit log, tenant isolation, upload scanning, or secret redaction.",
        "Uploaded documents, extracted text, evidence, questions, and generated outputs may be sent to OpenAI.",
        "The data store is file-based; replacing an input file changes app behavior without migrations or schema controls.",
        "Generated XLSX/ZIP files can contain verbatim source-document evidence and must be handled as sensitive outputs.",
        "Local pickle caches are not encrypted and are not integrity-checked.",
    ])

    h(doc, "9. ChatGPT assistant layer", 1)
    p(doc, "Three custom GPTs published in ChatGPT expose the evaluation engines without requiring the user to open Streamlit. All three consume a single FastAPI backend deployed on Render: the user uploads a DOCX inside the chat and receives a downloadable XLSX.")
    add_table(doc, ["Assistant", "Engine and rubric", "Scale", "Endpoints"], [
        ["ILO PRODOC Quality Appraisal (Tab 1 v3)", "tab1_v3_core.py + Rubrica_Tab1_Detallada_Full_v3.xlsx", "Yes / Partial / No / Not Found / N/A across 76 criteria in 5 sections", "POST /v3/jobs, GET /v3/jobs/{id}, GET /v3/jobs/{id}/result"],
        ["OIT - Diagnóstico de Atributos Específicos (Tab 2)", "tab2_core.py + Rubricas_6ago2025.xlsx", "1 to 5 per criterion", "POST /attributes/jobs, GET /attributes/jobs/{id}, GET /attributes/jobs/{id}/result"],
        ["OIT - Diagnóstico de Sostenibilidad del Proyecto (Tab 3)", "sustainability_core.py + Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "0 to 3 per indicator, 28 indicators across 3 dimensions", "POST /sustainability/jobs, GET /sustainability/jobs/{id}, GET /sustainability/jobs/{id}/result"],
    ], widths=[1.7, 2.0, 1.6, 1.5])

    h(doc, "9.1 Published access links", 2)
    add_table(doc, ["Assistant", "Link"], [
        ["Preliminary Quality Appraisal (Tab 1 v3)", GPT_V3],
        ["Specific Attributes Diagnosis (Tab 2)", GPT_TAB2],
        ["Sustainability Diagnosis (Tab 3)", GPT_SUS],
        ["Shared backend (FastAPI on Render)", GPT_BACKEND],
    ], widths=[2.4, 4.4])
    p(doc, "All three GPTs are configured as \"Anyone with a link\": they are not listed in the public GPT directory, but anyone who receives the link and holds a ChatGPT account can open them and run evaluations. The link acts as a bearer credential — forwarding it transfers access.")
    add_bullets(doc, [
        "There is no authorised-user list and no identity check: the backend does not record who launched an evaluation.",
        "API usage bills against a single OPENAI_API_KEY shared by all three assistants, so a link circulated beyond the intended group becomes unattributable spend.",
        "The only way to revoke access is to change the GPT sharing setting or republish it to mint a new link, then redistribute it.",
        "Link distribution is in practice the only access control and should be managed with the same care as a credential.",
    ])

    h(doc, "9.2 Evaluation job lifecycle", 2)
    add_numbered(doc, [
        "The user uploads a DOCX in the ChatGPT conversation and describes the desired scope.",
        "The GPT calls the matching POST endpoint, passing the reference to the file hosted on OpenAI servers plus any requested filters.",
        "The backend downloads the file, extracts its text, loads the rubric from disk, and dispatches model calls in parallel.",
        "The GPT polls GET /{family}/jobs/{id} until the status becomes succeeded or failed.",
        "Once finished, the GPT calls GET /{family}/jobs/{id}/result and delivers a narrative summary plus the base64-encoded XLSX.",
    ])
    p(doc, "Job state lives in the JOBS dictionary in gpt_action_api.py, guarded by a threading.Lock. This is in-process memory: a restart or redeploy loses any job in flight. Sustained institutional use should migrate this to Redis or an equivalent store.")

    h(doc, "9.3 Parallelism and the stability scheme", 2)
    add_table(doc, ["Engine", "Model", "MAX_WORKERS", "Runs per criterion"], [
        ["tab1_v3_core.py", "gpt-5-mini with adaptive reasoning effort (medium/minimal depending on criterion subjectivity)", "48", "10 (STABILITY_REPEATS), with an 80% stability threshold"],
        ["tab2_core.py", "gpt-5-mini", "8", "5 (STABILITY_REPEATS), collapsed to the modal value with a stability percentage"],
        ["sustainability_core.py", "gpt-5-mini", "8", "5 (STABILITY_REPEATS), collapsed to the modal value with a stability percentage"],
    ], widths=[1.6, 2.6, 1.1, 1.5])
    p(doc, "All three engines repeat each evaluation and collapse the runs to the modal value, reporting what percentage of runs agreed with that mode. The percentage reaches the XLSX and distinguishes a criterion scored by consensus from an unstable one.")
    p(doc, "The implementation differs: tab2_core.py and sustainability_core.py use the shared stability.py module, while tab1_v3_core.py carries its own aggregator (aggregate_repeated_criterion_results) with 10 repeats and an 80% stability threshold, below which a criterion is flagged unstable. A full 76-criterion appraisal therefore issues on the order of 760 model calls, which dominates its cost and duration.")

    h(doc, "9.4 Backend deployment", 2)
    add_table(doc, ["Item", "Value"], [
        ["Platform", "Render.com, free plan"],
        ["Service", "ilo-prodoc-appraisal-v3, declared in render.yaml as a Blueprint"],
        ["Runtime", "Docker image built from Dockerfile.gpt-action on Python 3.10-slim"],
        ["Framework", "FastAPI with uvicorn[standard]; dependencies in requirements.gpt-action.txt"],
        ["Health check", "GET /health, set as healthCheckPath in render.yaml"],
        ["CI/CD", "Automatic redeploy on push to the main branch"],
        ["ChatGPT specifications", "openapi_gpt_action_v3.yaml, openapi_gpt_action_tab2.yaml, and openapi_gpt_action_sustainability.yaml"],
    ], widths=[2.0, 4.8])
    p(doc, "Render's free plan suspends the service after a period of inactivity; the first request after suspension takes noticeably longer while the container boots. This is expected behaviour, not a fault.")

    h(doc, "9.5 Risks specific to this layer", 2)
    add_bullets(doc, [
        "Job state is held in memory: a redeploy during an evaluation loses it with no way to recover.",
        "Render's free plan imposes cold-start latency and offers no availability guarantee.",
        "Documents are uploaded to OpenAI servers first and then downloaded by the backend; both transfers matter for confidentiality analysis.",
        "Rubrics live in the Docker image filesystem: updating one requires a rebuild and redeploy, not just replacing a file.",
        "All three GPTs share one backend and one API key, so an outage affects all three at once.",
    ])


def doc_en2(doc):
    h(doc, "1. Source-code scope", 1)
    p(doc, "The documented source file is oli_v6_deploy.py. The split variants exist in the repository but are not used as factual sources for this documentation.")
    p(doc, "Maintenance rule: when the main app changes, separately decide whether the change must be propagated to the core and recommendations variants before delivery.")

    h(doc, "2. Relevant project files", 1)
    add_table(doc, ["File", "Operational role"], [
        ["oli_v6_deploy.py", "Main Streamlit app and source of truth for these documents."],
        ["requirements.txt", "Python dependencies."],
        ["df_complete_all_full.xlsx", "Main enriched recommendation dataset."],
        ["df_split_actions.xlsx", "Raw/split recommendation dataset used for years and missing records."],
        ["analyzed_recommendations_plans_v5.csv", "Pipe-separated recommendation/action-plan analysis add-on."],
        ["emb_Recomm_rec_cl_4.pt", "Recommendation embeddings loaded at startup for FAISS."],
        ["Recommendation_RAG_Metadata.pt", "Metadata aligned to recommendation embeddings."],
        ["Appraisal Checklist_2025 es-419.xlsx", "Tab 1 appraisal questions, sheet rubric."],
        ["Rubricas_6ago2025.xlsx", "Tab 2 active sheets: rubric_parteval, rubric_gender_, and rubric_TJ_TJ."],
        ["Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "Active Tab 3 sustainability rubric."],
        ["Recommendations_World.xlsx", "Default recommendation records for Tabs 5 and 6."],
        ["Frame_Recommendations_English.xlsx", "Reference frame with texto_merged, dimension, and subdim."],
    ], widths=[2.8, 4.0])

    h(doc, "3. Local installation", 1)
    add_numbered(doc, [
        "Use a Python version compatible with the pinned dependencies.",
        "Create and activate a virtual environment.",
        "Install dependencies with pip install -r requirements.txt.",
        "Set OPENAI_API_KEY in the runtime environment.",
        "Run streamlit run oli_v6_deploy.py from the project directory.",
    ])
    codeblock(doc, [
        "python -m venv .venv",
        "source .venv/bin/activate",
        "pip install -r requirements.txt",
        "streamlit run oli_v6_deploy.py",
    ])

    h(doc, "4. Main function map", 1)
    add_table(doc, ["Block", "Functions or classes", "Description"], [
        ["Excel/export", "to_excel, create_results_download_with_sections", "Writes XLSX/ZIP outputs in memory and disables automatic URL conversion."],
        ["Tokens", "truncate_to_token_limit", "Cuts prompt context to token limits with tiktoken."],
        ["Embeddings and RAG", "get_embedding_with_retry, load_embeddings, faiss.IndexFlatIP", "Loads a startup FAISS index and supports large-document RAG in Tab 4."],
        ["Document extraction", "extract_docx_structure_enhanced, validate_extraction, extract_document_content", "Extracts hierarchy, tables, and text for Tabs 1-3. Tab 4 chat does not use the enhanced extractor."],
        ["Tab 1 appraisal", "load_appraisal_questions, parse_two_part_question, analyze_question_with_llm_tab1, _critic_impl", "Evaluates each question, handles two-part focus, A-E gates, and transversal A/B gates."],
        ["Tab 2/3 rubric evaluation", "evaluate_criterion_with_llm, synthesize_evaluations", "Scores selected criteria 1-5 with evidence."],
        ["Recommendation classification", "EmbeddingsCache, classify_recommendations, classify_recommendations_en, verify_match_with_llm", "Classifies records against the frame with numpy cosine similarity and optional LLM verification."],
        ["Deep recommendation analysis", "AnalysisCache, analyze_recommendation_plan_pair, run_row_analysis", "Uses local pickle cache; effective model is gpt-4o-mini in current code."],
    ], widths=[1.6, 2.6, 2.6])

    h(doc, "5. State conventions", 1)
    add_bullets(doc, [
        "Tabs use separate st.session_state keys such as tab1_results_df, document_extracted_tab2, tab3_results, doc_chat_docs, doc_chat_embeddings, classified_world_df, and deep_analysis_df.",
        "Tabs 1-3 detect uploaded-file changes with file-content hashes and reset extraction state.",
        "Tab 4 detects document changes only by filename list.",
        "Downloadable outputs are generated in memory and are not automatically saved to disk.",
    ])

    h(doc, "6. Error handling and verification", 1)
    add_bullets(doc, [
        "Missing critical files produce Streamlit errors or warnings and stop the affected flow.",
        "Missing OPENAI_API_KEY produces a global warning; OpenAI-dependent functions then fail or return errors.",
        "Invalid JSON responses are cleaned where possible; if parsing fails, flows return default error records.",
        "Run test_two_part_parsing.py when changing two-part parsing or Tab 1 question logic.",
        "Manually test one question per configured transversal theme after changing aliases or scoring.",
        "Test Tab 4 with both a small TXT and a large DOCX; also test replacing a file with the same filename.",
        "Test Tabs 5 and 6 with a limited sample, with LLM Verification off and on.",
    ])

    h(doc, "7. Visible technical debt", 1)
    add_bullets(doc, [
        "Large commented legacy blocks remain in the main file.",
        "Startup data and embeddings are initialized twice.",
        "load_embeddings remains active at startup although the visible recommendation-search UI is commented.",
        "Tab 4 asks users for citations and metadata, but its DOCX chat extraction only reads paragraphs.",
        "Tabs 5 and 6 expose a deep-analysis model selector that does not control the effective model.",
        "Some Tab 3 error messages still mention PRODOC_rubric.xlsx even though the active rubric is Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx.",
        "The app lacks a clean separation between UI, domain logic, extraction, and external-service clients.",
    ])

    h(doc, "8. GPT assistant backend code", 1)
    p(doc, "The backend serving the three GPTs is independent of oli_v6_deploy.py: it imports no Streamlit and shares no state with the app. The evaluation engines were extracted into interface-free core modules so the same code can run from Streamlit or from the API.")

    h(doc, "8.1 Backend files", 2)
    add_table(doc, ["File", "Role"], [
        ["gpt_action_api.py", "FastAPI application: header authentication, job creation and polling, DOCX download from OpenAI, /health and /privacy endpoints."],
        ["tab1_v3_core.py", "Preliminary appraisal v3 engine: loads the 76-criterion rubric, builds prompts, runs the evaluation, assembles the XLSX."],
        ["tab1_v3.py", "Streamlit wrapper around the same engine, used by the experimental Tab 7 of the app."],
        ["tab2_core.py", "Specific-attributes engine: three selectable rubrics, 1-5 scale, stability consolidation."],
        ["sustainability_core.py", "Sustainability engine: 28 indicators across three dimensions, 0-3 scale, stability consolidation."],
        ["stability.py", "Shared repetition and consolidation module: runs each item N times, collapses to the mode, computes the stability percentage."],
        ["openapi_gpt_action_v3.yaml", "OpenAPI specification pasted into the appraisal GPT's Action."],
        ["openapi_gpt_action_tab2.yaml", "OpenAPI specification for the specific-attributes GPT."],
        ["openapi_gpt_action_sustainability.yaml", "OpenAPI specification for the sustainability GPT."],
        ["Dockerfile.gpt-action", "Backend Docker image on Python 3.10-slim."],
        ["docker-compose.gpt-action.yml", "Local containerised run of the backend."],
        ["requirements.gpt-action.txt", "Backend dependencies, kept separate from the Streamlit requirements.txt."],
        ["render.yaml", "Render Blueprint: service, runtime, health check, and environment variables."],
    ], widths=[2.3, 4.5])

    h(doc, "8.2 Running the backend locally", 2)
    codeblock(doc, [
        "pip install -r requirements.gpt-action.txt",
        "export OPENAI_API_KEY=...",
        "export ILO_GPT_ACTION_API_KEY=...",
        "uvicorn gpt_action_api:app --reload --port 8000",
        "",
        "# or containerised:",
        "docker compose -f docker-compose.gpt-action.yml up",
    ])
    p(doc, "The process must run from the directory holding the rubric files, because the engines resolve them by relative path.")

    h(doc, "8.3 Backend function map", 2)
    add_table(doc, ["Function or endpoint", "Responsibility"], [
        ["require_api_key", "Compares the X-API-Key header against ILO_GPT_ACTION_API_KEY; guards every job endpoint."],
        ["POST /v3/jobs, /attributes/jobs, /sustainability/jobs", "Register a job, return its identifier, and start the evaluation in the background."],
        ["GET /{family}/jobs/{id}", "Returns job status: queued, running, succeeded, or failed, with progress when available."],
        ["GET /{family}/jobs/{id}/result", "Returns the structured summary and the base64-encoded XLSX."],
        ["GET /health", "Availability probe used by Render. Unauthenticated."],
        ["GET /privacy", "HTML privacy policy. Required by ChatGPT to publish a GPT with an Action. Unauthenticated."],
        ["stability.evaluate_with_stability", "Runs each item `repeats` times in parallel and delegates consolidation to the engine's aggregator."],
        ["stability.aggregate_runs", "Collapses runs to the modal value and computes stability, distribution, and drift."],
        ["results_to_xlsx_bytes", "Present in all three engines; builds the final XLSX with xlsxwriter."],
    ], widths=[2.5, 4.3])

    h(doc, "8.4 Backend technical debt", 2)
    add_bullets(doc, [
        "The JOBS dictionary is in-process memory: it does not survive restarts and does not support more than one service instance.",
        "Jobs never expire: a long-lived process accumulates results in memory indefinitely.",
        "All three engines duplicate results_to_xlsx_bytes with minor variations; these could converge if the output format is unified.",
        "Rubrics resolve by path relative to the working directory, coupling the process to where it is launched.",
        "stability.py ships a runnable _demo() self-check; it is the only backend code with automated verification.",
    ])


def doc_en3(doc):
    h(doc, "1. Data model", 1)
    p(doc, "The application does not use a relational database. Its data model is file-based: Excel workbooks, pipe-separated CSV, PyTorch tensors, local pickle caches, and in-memory Streamlit session state.")

    h(doc, "2. Data-file inventory", 1)
    add_table(doc, ["File or object", "Loaded by", "Use"], [
        ["df_complete_all_full.xlsx", "load_data", "Main recommendation base; creates index_df from ID_Recomendacion and normalizes dimension/subdim fields."],
        ["df_split_actions.xlsx", "load_data", "Raw base with Recommendation date; contributes year and missing index_df records."],
        ["analyzed_recommendations_plans_v5.csv", "load_extended_data", "Pipe-separated add-on with tags, scores, and analysis metrics."],
        ["emb_Recomm_rec_cl_4.pt", "load_embeddings", "Recommendation embedding matrix; active startup dependency."],
        ["Recommendation_RAG_Metadata.pt", "load_embeddings", "Metadata aligned to recommendation embeddings; active startup dependency."],
        ["emb_LL_ll_cl_4.pt", "load_lessons_embeddings", "Lessons embeddings for an inactive helper, not an active UI tab."],
        ["lessons_metadata.pt", "load_lessons_embeddings", "Lessons metadata for an inactive helper, not an active UI tab."],
        ["Appraisal Checklist_2025 es-419.xlsx", "load_appraisal_questions", "Tab 1 question source; requires Pregunta_Realizada and uses Tema for renumbering."],
        ["TRANSVERSAL_MATTERS", "Tab 1", "In-code dictionary for gender, non-discrimination, disability, social dialogue/tripartism, and environmental sustainability aliases."],
        ["Rubricas_6ago2025.xlsx", "Tab 2", "Active sheets are rubric_parteval, rubric_gender_, and rubric_TJ_TJ."],
        ["Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "Tab 3", "Active sustainability rubric, sheet rubric."],
        ["Recommendations_World.xlsx", "Tabs 5 and 6", "Default classification dataset; can be replaced with an uploaded XLSX containing Recommendation description."],
        ["Frame_Recommendations_English.xlsx", "Tabs 5 and 6", "Reference frame; requires texto_merged, dimension, and subdim."],
        ["embeddings_cache.pkl", "EmbeddingsCache", "Persistent local cache for classification embeddings; not encrypted or integrity-checked."],
        ["analysis_cache.pkl", "AnalysisCache", "Persistent local cache for deep analysis; not encrypted or integrity-checked."],
    ], widths=[2.6, 1.7, 2.7])

    h(doc, "3. Logical relationships", 1)
    add_bullets(doc, [
        "df_complete_all_full.xlsx and df_split_actions.xlsx align by index_df.",
        "analyzed_recommendations_plans_v5.csv is merged into the main dataframe by index_df when new columns exist.",
        "emb_Recomm_rec_cl_4.pt must stay aligned with Recommendation_RAG_Metadata.pt.",
        "Recommendations_World.xlsx records are compared against Frame_Recommendations_English.xlsx rows using texto_merged as semantic reference text.",
        "Recommendation datasets may contain repeated rows because a recommendation can have multiple attributes; Tabs 5 and 6 show both row counts and unique recommendation counts.",
    ])

    h(doc, "4. Main transformations", 1)
    add_table(doc, ["Transformation", "Details"], [
        ["Column normalization", "load_data replaces spaces and dots with underscores."],
        ["Years", "Recommendation_date and Recommendation date are parsed with pandas; years before 2018 in analyzed_df are forced to 2018."],
        ["FAISS", "A recommendation FAISS index is built at startup. Tab 4 rebuilds an IndexFlatIP per large-document query using session-cached chunk embeddings."],
        ["Transversal scoring", "Configured themes bypass A-E and use A=objective/product/activity and B=budget/resources. Yes=A+B, Partial=A or B, No=neither."],
        ["Classification", "Embeddings are normalized and compared with numpy cosine similarity. Base assignment uses top 3; optional LLM verification reranks top 10; secondary dimensions/subdimensions require similarity >=0.60."],
        ["Exports", "Results are written to XLSX and ZIP files with stable column ordering."],
    ], widths=[2.0, 4.8])

    h(doc, "5. Minimum fields by workflow", 1)
    add_table(doc, ["Workflow", "Minimum required fields"], [
        ["Tab 1 appraisal", "Appraisal Checklist: Pregunta_Realizada; Tema recommended. Transversal questions are recognized by in-code aliases."],
        ["Tab 2 attributes", "Rubricas_6ago2025 active sheets with Indicador, Dimensión, and performance-level columns."],
        ["Tab 3 sustainability", "Active rubric fields for dimension, criterion, indicator, and levels from sheet rubric."],
        ["Tabs 5 and 6 classification", "Recommendation file: Recommendation description. Frame: texto_merged, dimension, subdim."],
        ["Deep analysis", "Recommendation description and Action plan; Comments improves analysis when present."],
        ["Executive summary", "Recommendation description; Management response, Comments, and Action plan are used when available."],
    ], widths=[2.0, 4.8])

    h(doc, "6. Backups and data refresh", 1)
    add_bullets(doc, [
        "Keep a dated copy of each Excel, CSV, tensor, and cache before replacing it.",
        "After replacing data files, restart the app or clear Streamlit cache.",
        "If embeddings are regenerated, regenerate aligned metadata in the same run.",
        "Validate Tabs 5 and 6 with a small sample when either Recommendations_World.xlsx or Frame_Recommendations_English.xlsx changes schema.",
        "Do not manually edit embeddings_cache.pkl or analysis_cache.pkl; delete them with the app stopped if corruption is suspected.",
    ])

    h(doc, "7. GPT assistant data and rubrics", 1)
    p(doc, "The GPT backend is equally file-based: rubrics are read from the Docker image filesystem at startup. There is no database and no persistent storage of documents or results.")

    h(doc, "7.1 Rubrics consumed by the backend", 2)
    add_table(doc, ["File", "Consuming engine", "Contents"], [
        ["Rubrica_Tab1_Detallada_Full_v3.xlsx", "tab1_v3_core.py", "76 preliminary appraisal criteria across 5 sections, subsections 1.1 to 5.2, with per-criterion subjectivity metadata."],
        ["Rubricas_6ago2025.xlsx", "tab2_core.py", "Three rubric sheets: rubric_parteval (participatory methods), rubric_gender_ (gender), rubric_TJ_TJ (modern just transition)."],
        ["Evaluación de sostenibilidad del proyecto_rubric_9feb26.xlsx", "sustainability_core.py", "28 indicators across three project-cycle dimensions: Design (6), Implementation (10), Pre-Closure (12)."],
    ], widths=[2.3, 1.7, 2.8])
    p(doc, "sustainability_core.py resolves its rubric through an accent-tolerant glob, because the encoding of the accented character varies across filesystems. Any replacement file must preserve that name pattern.")

    h(doc, "7.2 Job state model", 2)
    p(doc, "gpt_action_api.py keeps an in-memory JOBS dictionary, serialised by a threading.Lock. Each entry represents one evaluation in flight or finished.")
    add_table(doc, ["Field", "Contents"], [
        ["job_id", "Identifier generated on job creation; the dictionary key."],
        ["status", "queued, running, succeeded, or failed."],
        ["progress", "Progress reported by the engine during evaluation, when available."],
        ["result", "Structured summary plus base64 XLSX, present only when status is succeeded."],
        ["error", "Failure message, present only when status is failed."],
    ], widths=[1.6, 5.2])
    p(doc, "This state is volatile by design: a restart or redeploy erases it entirely. It is not a historical record and must not be used for audit. The auditable record of an evaluation is the XLSX the user downloads.")

    h(doc, "7.3 Updating a rubric in the backend", 2)
    add_numbered(doc, [
        "Replace the XLSX in the repository, preserving the filename and the column and sheet structure.",
        "Verify loading locally by running the backend and launching a small scoped test evaluation.",
        "Push to the main branch: Render rebuilds the image and redeploys automatically.",
        "Confirm with GET /health that the service is back, and with a short evaluation that the new rubric applies.",
    ])
    p(doc, "Uploading the file to the server is not enough: the rubric is part of the Docker image, so any change requires a rebuild and redeploy.")


def doc_en4(doc):
    h(doc, "1. Required configuration", 1)
    add_table(doc, ["Item", "Expected value"], [
        ["Local command", "streamlit run oli_v6_deploy.py"],
        ["AI variable", "OPENAI_API_KEY available in the process environment."],
        ["Production mode", "STREAMLIT_ENV=production when file paths should come from st.secrets."],
        ["Working directory", "Directory containing required Excel, CSV, and .pt files unless production secrets provide paths."],
        ["Dependencies", "Installed from requirements.txt."],
    ], widths=[2.0, 4.8])

    h(doc, "2. Secrets and paths", 1)
    p(doc, "OPENAI_API_KEY is read through os.getenv. Data paths are read from st.secrets only when STREAMLIT_ENV=production.")
    add_table(doc, ["Name", "When used", "Expected content"], [
        ["df_path", "Production", "Path to df_complete_all_full.xlsx."],
        ["df_raw_path", "Production", "Path to df_split_actions.xlsx."],
        ["embeddings_path", "Production", "Path to emb_Recomm_rec_cl_4.pt."],
        ["structured_embeddings_path", "Production", "Path to Recommendation_RAG_Metadata.pt."],
        ["analyzed_recommendations_path", "Optional production", "Path to analyzed_recommendations_plans_v5.csv."],
        ["lessons_embeddings_path", "Inactive helper", "Path to emb_LL_ll_cl_4.pt if lessons are reconnected."],
        ["structured_lessons_path", "Inactive helper", "Path to lessons_metadata.pt if lessons are reconnected."],
    ], widths=[2.0, 1.8, 3.0])

    h(doc, "3. Access control", 1)
    p(doc, "The app does not implement users, roles, passwords, per-feature permissions, audit logs, tenant isolation, upload scanning, or secret redaction. Those controls must be provided by the hosting environment, institutional network, proxy, or equivalent access layer.")
    add_bullets(doc, [
        "Only authorized ILO users should be able to open the app.",
        "Repository write access and operational data write access should be restricted to technical maintainers.",
        "OPENAI_API_KEY must be visible only to the runtime and authorized administrators.",
        "Users are responsible for protecting any downloads they generate.",
    ])

    h(doc, "4. Data security", 1)
    add_bullets(doc, [
        "Uploaded documents are processed in memory and, in extraction flows, through temporary files that are removed when cleanup completes.",
        "Document text, questions, evidence, and prompts can be sent to OpenAI.",
        "st.session_state holds extracted text and results during the active session.",
        "Generated XLSX/ZIP outputs can contain verbatim source text, reasoning, and evidence.",
        "Local pickle caches are not encrypted and are not integrity-checked.",
        "The app does not implement output encryption or retention policy; ILO must define those policies in the runtime environment.",
    ])

    h(doc, "4.1 Active cache layers", 2)
    add_table(doc, ["Layer", "Location", "Operational risk"], [
        ["st.cache_data", "Streamlit process memory/cache", "May keep older data/rubric loads until cache clear or restart."],
        ["Tab 4", "st.session_state", "Stores documents, chunks, and embeddings for the session; same-name file replacement can leave stale state."],
        ["Tabs 5 and 6", "embeddings_cache.pkl", "Persistent local cache, unencrypted and unchecked."],
        ["Deep analysis", "analysis_cache.pkl", "Persistent local cache, unencrypted and unchecked."],
        ["SimpleHierarchicalStore", "~/document_store/embedding_cache.pkl", "Legacy helper cache; not the main Tab 4 mechanism."],
    ], widths=[1.7, 2.3, 2.8])

    h(doc, "5. Safe operation parameters", 1)
    add_table(doc, ["Parameter", "Code value", "Impact"], [
        ["MAX_WORKERS global for Tab 1", "48", "High concurrency; requires sufficient API limits."],
        ["MAX_WORKERS Tab 2", "3", "Reduces rate-limit risk for rubric evaluation."],
        ["MAX_WORKERS Tab 3", "8", "Parallel sustainability criteria processing."],
        ["Tab 4 RAG threshold", ">100,000 characters", "Avoids sending very large full documents on each question."],
        ["Token limit", "110,000 tokens", "Truncates long context before model calls."],
        ["Classification batch concept", "20 texts", "Controls embedding generation for unique recommendations."],
    ], widths=[2.2, 1.6, 3.0])

    h(doc, "6. Production checklist", 1)
    add_bullets(doc, [
        "Confirm OPENAI_API_KEY exists in the runtime process and account budget is available.",
        "Confirm all required files open with pandas or torch in the production environment.",
        "Configure private access to the app.",
        "Run one non-sensitive test per tab before release.",
        "Verify downloads open and do not expose more information than expected.",
        "Document who may replace data files and who may deploy code.",
    ])

    h(doc, "7. Incident response", 1)
    add_numbered(doc, [
        "Disable external access to the app at the hosting or proxy layer.",
        "Revoke and rotate OPENAI_API_KEY if exposure is suspected.",
        "Retain downloads or logs only if internal policy allows it.",
        "Restore data files from verified backups.",
        "Reactivate the app after testing Tabs 1 through 6.",
    ])

    h(doc, "8. GPT assistant configuration and security", 1)
    p(doc, "The FastAPI backend serving the three GPTs has its own configuration and access-control model, distinct from the Streamlit app. Unlike oli_v6_deploy.py, this backend does implement authentication.")

    h(doc, "8.1 Backend environment variables", 2)
    add_table(doc, ["Variable", "Use", "Where it is set"], [
        ["OPENAI_API_KEY", "Authenticates the OpenAI client for every evaluation call.", "Render dashboard, marked sync:false in render.yaml so it is never committed."],
        ["ILO_GPT_ACTION_API_KEY", "Expected value of the X-API-Key header; guards every job endpoint.", "Render dashboard, marked sync:false; the same value is loaded into each GPT's Action."],
    ], widths=[2.0, 2.8, 2.0])
    p(doc, "Both variables must exist in the runtime process. render.yaml declares them without values, so Render prompts for them on first deploy and they never reach the repository.")

    h(doc, "8.2 Three layers of access control", 2)
    add_table(doc, ["Layer", "Mechanism", "What it protects"], [
        ["GPT access", "ChatGPT private link: only link holders can open it, and a ChatGPT account is required.", "Who can use the assistant."],
        ["API access", "X-API-Key header compared against ILO_GPT_ACTION_API_KEY in require_api_key.", "Who can launch evaluations against the backend."],
        ["Open endpoints", "/health and /privacy are unauthenticated, required by Render and ChatGPT respectively.", "They expose no data: only service status and policy text."],
    ], widths=[1.6, 3.0, 2.2])
    add_bullets(doc, [
        "Distribution of the GPT links is the effective access control for end users and should be managed like a distribution list.",
        "ILO_GPT_ACTION_API_KEY is a shared secret between the backend and all three Actions; rotating it requires updating all three GPTs.",
        "There are no users, roles, or per-person audit trails: the backend cannot tell who launched an evaluation.",
    ])

    h(doc, "8.3 Data security in the GPT workflow", 2)
    add_bullets(doc, [
        "The document is uploaded to OpenAI servers first, inside the ChatGPT conversation, and then downloaded by the backend for processing. Both transfers matter when classifying document sensitivity.",
        "Extracted text and evidence excerpts are sent to the OpenAI API during evaluation.",
        "The backend does not persist the document: it processes and discards it when the job ends.",
        "Results live in process memory until restart; there is no database and no historical record.",
        "The delivered XLSX contains verbatim quotations from the source document and must be handled with the same confidentiality.",
        "The ChatGPT conversation, with the attached document, remains in the user's account: retention depends on that account's settings and ChatGPT plan.",
    ])

    h(doc, "8.4 Backend operating parameters", 2)
    add_table(doc, ["Parameter", "Value", "Impact"], [
        ["MAX_WORKERS tab1_v3_core", "48", "High concurrency across 76 criteria, multiplied by the 10 stability runs; requires generous API rate limits on the OpenAI account."],
        ["MAX_WORKERS tab2_core", "8", "Moderate concurrency, multiplied by the 5 stability runs."],
        ["MAX_WORKERS sustainability_core", "8", "Moderate concurrency, multiplied by the 5 stability runs."],
        ["STABILITY_REPEATS tab1_v3_core", "10", "Each criterion is evaluated ten times; a full appraisal issues on the order of 760 model calls. It is the most expensive of the three workflows."],
        ["STABILITY_REPEATS, Tab 2 and sustainability", "5", "Each criterion is evaluated five times, multiplying cost and duration of those two evaluations by five."],
        ["STABILITY_THRESHOLD_PCT tab1_v3_core", "80.0", "Below that agreement rate across runs, a criterion is flagged unstable and warrants human review."],
        ["Model across all three engines", "gpt-5-mini", "Dominates per-evaluation cost together with the repeat count; tab1_v3 varies reasoning effort by criterion subjectivity."],
        ["Render plan", "free", "The service suspends on inactivity; the first request after suspension pays a cold-start delay."],
    ], widths=[2.0, 1.2, 3.6])

    h(doc, "8.5 Backend production checklist", 2)
    add_bullets(doc, [
        "Confirm OPENAI_API_KEY and ILO_GPT_ACTION_API_KEY are set in Render and that the OpenAI account has budget.",
        "Confirm GET /health responds.",
        "Confirm all three rubric files are present in the deployed image.",
        "Run a short evaluation through each of the three GPTs with a non-sensitive document.",
        "Verify each GPT's Action holds the correct key and points at the correct server.",
        "Confirm /privacy responds: ChatGPT requires it to keep a GPT with an Action published.",
        "Document internally who distributes the GPT links and who may rotate the API key.",
    ])

    h(doc, "8.6 Incident response for the GPT layer", 2)
    add_numbered(doc, [
        "On suspected misuse, rotate ILO_GPT_ACTION_API_KEY in Render: this immediately invalidates all three Actions until they are updated.",
        "On suspected exposure of the OpenAI key, rotate it in the OpenAI dashboard and update it in Render.",
        "For availability problems, check the Render service status and container logs.",
        "If a GPT link spread beyond the intended circle, republish the GPT to mint a new link and redistribute it.",
        "Restore service by redeploying from the main branch and verify with one test evaluation per assistant.",
    ])


def doc_en5(doc):
    h(doc, "1. Functional description", 1)
    p(doc, "The app helps users review project and evaluation documents, chat with uploaded documents, classify recommendations, and analyze institutional action plans. The active interface has six top-level tabs.")
    p(doc, "The system also has a second delivery surface: three custom GPTs published in ChatGPT that reproduce the quality appraisal, specific attributes, and sustainability workflows without opening Streamlit. They are documented in section 10.")

    h(doc, "2. User requirements", 1)
    add_bullets(doc, [
        "Use a modern browser and stable connection.",
        "Have institutional authorization to upload documents to the app.",
        "Use DOCX files with real Word heading styles for Tabs 1-3.",
        "Treat AI outputs as evidence-supported assistance, not a replacement for technical review.",
    ])

    h(doc, "3. Main navigation", 1)
    add_table(doc, ["Tab", "Use it for", "Main output"], [
        ["Valoración Preliminar de Calidad de Proyectos", "Integrated project design review.", "ZIP with question-level XLSX, subsection/section analyses, rubric template, and TXT summary."],
        ["Diagnóstico de Atributos Específicos", "Focused evaluation of participation methodology, gender integration, or modern just transition.", "ZIP with XLSX files by selected rubric."],
        ["Diagnóstico de Sostenibilidad del Proyecto", "Sustainability review of a PRODOC or related project document.", "ZIP with XLSX files and score charts."],
        ["Pregúntale a tus Documentos", "Conversational review of one or more DOCX/TXT files.", "On-screen chat with session memory."],
        ["Clasificación de Recomendaciones", "Spanish UI for recommendation classification using Recommendation description.", "Treemaps, trends, deep-analysis XLSX, and executive-summary XLSX."],
        ["Recommendation Classification", "English UI for the same classification workflow.", "English outputs and downloads."],
    ], widths=[2.2, 2.5, 2.1])

    h(doc, "4. Tab 1: Preliminary Project Quality Appraisal", 1)
    add_numbered(doc, [
        "Download or review Appraisal Checklist_2025 es-419.xlsx if needed.",
        "Upload a DOCX with Word headings.",
        "Extract the document and review the extracted structure.",
        "Select sections and optional question filters.",
        "Run document analysis.",
        "Review metrics, result rows, reasoning, evidence, and critical evaluation.",
        "Download appraisal_checklist_results.zip.",
    ])
    p(doc, "For non-transversal specific-subject questions, the app uses the A-E framework: sub-objective/output, indicator, activity, budget, and quantifiable target. The code can mechanically adjust verdicts based on the A-E total.")
    p(doc, "For two-part questions, Part 2 is evaluated and Part 1 is framing only. Reasoning may state 'Se identificaron 2 partes en esta pregunta.' Evidence tagged [FRAMING] cannot by itself support Partial or Yes.")
    p(doc, "For configured transversal themes, A-E is not used. A means operational presence in objective/product/activity; B means budget/resources. Yes requires A and B; Partial requires A or B; No means neither. Indicators and targets do not count.")
    p(doc, "Tab 1 gender is a transversal A/B topic. Tab 2 gender is a separate rubric, 'Integración del Enfoque de Género', from Rubricas_6ago2025.xlsx sheet rubric_gender_, scored 1-5.")
    p(doc, "The main ZIP contains appraisal_checklist_results.xlsx, appraisal_checklist_rubric_template.xlsx when the source workbook is available, and appraisal_checklist_summary.txt. Extracted structure can also be downloaded as estructura_documento_tab1_<filename>.xlsx.")

    h(doc, "5. Tab 2: Specific Attributes Diagnosis", 1)
    add_numbered(doc, [
        "Upload a DOCX and extract structure.",
        "Select document sections for evaluation.",
        "Select one or more active rubrics: evaluation participation/methodology, gender integration, or modern just transition.",
        "Select criteria within each rubric.",
        "Run processing and evaluation.",
        "Review Criterio, Dimensión, Score, Análisis, Evidencia, Error, and Rúbrica columns.",
        "Download resultados_rubricas.zip.",
    ])
    p(doc, "The extracted structure can be downloaded separately as estructura_documento_tab2_<filename>.xlsx.")

    h(doc, "6. Tab 3: Project Sustainability Diagnosis", 1)
    add_numbered(doc, [
        "Download the sustainability rubric if review is needed.",
        "Upload DOCX, extract structure, and select sections.",
        "Select dimensions and criteria.",
        "Run processing and evaluation.",
        "Review Dimensión, Criterio, Indicador, Score, Análisis, and Evidencia.",
        "Use charts to prioritize weak dimensions and criteria.",
        "Download resultados_evaluacion_prodoc.zip.",
    ])
    p(doc, "The extracted structure can be downloaded separately as estructura_documento_tab3_<filename>.xlsx.")

    h(doc, "7. Tab 4: Ask Your Documents", 1)
    add_numbered(doc, [
        "Upload one or more DOCX or TXT files.",
        "Confirm active filenames and preview text.",
        "Ask specific questions. Request short quotes where useful, but verify section/page references manually.",
        "For up to 100,000 total characters, the app uses full context truncated to 110,000 tokens.",
        "For larger uploads, the app uses 2,000-character chunks with 300-character overlap, retrieves up to 15 FAISS chunks plus up to 5 keyword chunks, and sends the last 5 chat messages.",
    ])
    p(doc, "If replacing a document with another file of the same name, restart the session or rename the file. The chat detects changes by filename list, not content hash.")
    p(doc, "RAG answers can include inference from contextual clues; verify them in the source document before using them in formal work.")

    h(doc, "8. Tabs 5 and 6: Recommendation Classification", 1)
    add_numbered(doc, [
        "Use the default file or upload an XLSX with Recommendation description.",
        "Apply filters by location, time, theme, technical unit, source, progress, or available fields.",
        "Review row count and unique recommendation count before starting.",
        "Decide whether to enable LLM Verification; it can improve matching but is slower and more expensive.",
        "Start classification.",
        "Review dimension/subdimension treemaps and time trends.",
        "Use advanced AI tools for deep analysis or executive summary on the filtered subset.",
        "Download analisis_profundo.xlsx, resumen_ejecutivo.xlsx, deep_analysis.xlsx, or executive_summary.xlsx depending on tab language.",
    ])
    p(doc, "Base classification uses normalized embeddings and numpy cosine similarity, with top 3 assignment. Optional LLM verification reranks top 10 candidates. Secondary dimensions/subdimensions are kept only when similarity is at least 0.60. The output remains row-level even though embeddings are deduplicated by recommendation text.")
    p(doc, "Current limitation: the deep-analysis model selector does not control the effective model; the analysis function uses gpt-4o-mini.")

    h(doc, "9. Common problems", 1)
    add_table(doc, ["Symptom", "Likely cause", "Action"], [
        ["Sections are not detected", "DOCX lacks Word Heading styles.", "Fix headings in Word and extract again."],
        ["OPENAI API key not found", "OPENAI_API_KEY is not available to the process.", "Configure the environment variable and restart."],
        ["Classification does not start", "Missing Recommendations_World.xlsx or Frame_Recommendations_English.xlsx, or missing required columns.", "Restore files or upload a valid XLSX."],
        ["Rows show Error", "Invalid JSON, rate limit, or API failure.", "Retry with fewer criteria/sections or check API quota."],
        ["Tab 4 answers from an old file", "Different content was uploaded with the same filename.", "Rename the file or restart the Streamlit session."],
    ], widths=[2.0, 2.3, 2.5])

    h(doc, "10. ChatGPT assistants", 1)
    p(doc, "Beyond the Streamlit app, three of the evaluation workflows are available as custom GPTs inside ChatGPT. They require no installation and no Streamlit session: open a link, upload the DOCX in the chat, and receive a downloadable XLSX. Rubrics are loaded on the server, so users never upload them.")
    p(doc, "A dedicated end-user manual exists for this surface, Manual_Usuario_GPTs_OIT.docx, written for non-technical users.")

    h(doc, "10.1 Links and mapping to the Streamlit tabs", 2)
    add_table(doc, ["GPT assistant", "Equivalent to", "Link"], [
        ["ILO PRODOC Quality Appraisal", "Tab 1, with the 76-criterion v3 rubric", GPT_V3],
        ["OIT - Diagnóstico de Atributos Específicos", "Tab 2", GPT_TAB2],
        ["OIT - Diagnóstico de Sostenibilidad del Proyecto", "Tab 3", GPT_SUS],
    ], widths=[1.9, 1.6, 3.3])
    p(doc, "Tabs 4, 5, and 6 have no ChatGPT equivalent: document chat and recommendation classification remain exclusive to the Streamlit app.")
    p(doc, "Access is by private link and requires a ChatGPT account. Whoever administers the service should treat link distribution as the effective access control.")

    h(doc, "10.2 Workflow shared by all three assistants", 2)
    add_numbered(doc, [
        "Open the assistant link. Typing a greeting or pressing a starter button makes the assistant introduce itself and explain what it evaluates, without spending an evaluation.",
        "Attach a single .docx file. If several are uploaded, the assistant asks which one to use.",
        "State the scope in the same message: section or subsection for quality appraisal, thematic rubric for specific attributes, cycle dimension for sustainability. If omitted, the assistant asks before starting.",
        "Wait while the assistant launches the job and polls its status. A full evaluation can take several minutes.",
        "Review the on-screen summary and download the XLSX, which is the auditable record of the evaluation.",
    ])

    h(doc, "10.3 Scope and scales per assistant", 2)
    add_table(doc, ["Assistant", "Selectable scope", "Scale"], [
        ["Quality Appraisal (Tab 1 v3)", "Full rubric, one or more sections (1 to 5), or specific subsections (1.1 to 5.2)", "Yes / Partial / No / Not Found / N/A across 76 criteria"],
        ["Specific Attributes (Tab 2)", "One or more of three rubrics: participatory methods, gender, just transition", "1 to 5 per criterion"],
        ["Sustainability (Tab 3)", "Design dimension (PRODOCs), Implementation (progress reports), or Pre-Closure (closure and final evaluation)", "0 to 3 per indicator"],
    ], widths=[1.8, 3.0, 2.0])
    p(doc, "The scales are not interchangeable: 0-3 in sustainability and 1-5 in specific attributes measure different things and must not be averaged or compared directly.")
    p(doc, "All three assistants repeat each evaluation and consolidate the result: 10 runs per criterion in quality appraisal, 5 in specific attributes, and 5 in sustainability. The XLSX carries the stability percentage, showing how many of those runs agreed: a low percentage flags a criterion where the model was inconsistent and human review matters more.")
    p(doc, "A \"Not Found\" verdict in quality appraisal means the document lacks the information needed to assess the criterion. It differs from \"No\": the latter asserts the criterion is not met, while \"Not Found\" asserts it could not be determined.")

    h(doc, "10.4 Configured conversation starters", 2)
    add_table(doc, ["Assistant", "Starter buttons"], [
        ["Quality Appraisal", "¿Qué puedes hacer y cómo empiezo? · Evalúa este PRODOC con la rúbrica completa · Evalúa solo la sección 3 (Marco de resultados) · ¿Qué secciones y subsecciones puedo filtrar?"],
        ["Specific Attributes", "¿Qué rúbricas puedes aplicar? · Evalúa este documento con la rúbrica de género · Aplica la rúbrica de Transición Justa · Evalúa participación y género y compara resultados"],
        ["Sustainability", "¿Qué dimensiones evalúas y cuál me corresponde? · Evalúa este PRODOC con la dimensión de Diseño · Es un informe de avance: aplica Implementación · Aplica la rúbrica completa de sostenibilidad"],
    ], widths=[1.7, 5.1])
    p(doc, "Instruction text and starter buttons are maintained in docs/gpt_onboarding_es.md. When they are edited in the ChatGPT editor, that file must be updated too so documentation and assistants do not drift apart.")

    h(doc, "10.5 Common problems with the GPT assistants", 2)
    add_table(doc, ["Symptom", "Likely cause", "Action"], [
        ["The first evaluation of the day is slow to start", "The backend runs on Render's free plan and suspends on inactivity.", "Wait for the cold start; this is expected, not a fault."],
        ["The assistant reports an authentication failure", "The Action key does not match ILO_GPT_ACTION_API_KEY on the server.", "Check the key in the GPT's Action and in the Render dashboard."],
        ["The evaluation fails midway", "API rate limit, a very long document, or a transient service failure.", "Retry with a narrower scope: one section or a single rubric."],
        ["The assistant asks the user to upload the rubric", "Model drift from its instructions.", "Clarify that it should use the server-side rubric; if it recurs, review the GPT instructions."],
        ["An in-flight evaluation was lost", "The service restarted or redeployed: job state lives in memory.", "Relaunch the evaluation. The lost job cannot be recovered."],
        ["All three assistants fail at once", "Shared backend outage.", "Check GET /health on the server; this is a service problem, not a document problem."],
    ], widths=[2.0, 2.3, 2.5])


def doc_en6(doc):
    h(doc, "1. Operational handoff", 1)
    p(doc, "Daily operation should sit with the ILO team. The original developer should remain a last-resort support contact for issues that cannot be resolved through this documentation, code review, and local reproduction.")
    add_table(doc, ["Role", "Responsibility"], [
        ["ILO owner team", "Own access, data, execution, outputs, and methodological decisions."],
        ["ILO technical administrator", "Manage deployment, environment variables, data files, and backups."],
        ["Analyst users", "Run workflows, review evidence, download results, and apply technical judgment."],
        ["Original developer", "Exceptional support for architecture, complex debugging, or changes outside ordinary maintenance."],
    ], widths=[2.1, 4.7])

    h(doc, "2. Last-resort contact", 1)
    add_bullets(doc, [
        "Legacy contact from prior documentation: ageidv@gmail.com.",
        "Recommended use: only after reproducing the issue locally, reviewing app messages/logs, and isolating the affected tab or function.",
        "Include date, tab, non-sensitive input sample if possible, reproduction steps, error message, and screenshot if internal policy permits.",
    ])

    h(doc, "3. Maintenance routine", 1)
    add_table(doc, ["Frequency", "Activity"], [
        ["Weekly during active operation", "Confirm the app opens, OPENAI_API_KEY works, and a test document processes in Tab 4."],
        ["Monthly", "Test a sample in Tabs 1, 2, 3, 5, and 6; confirm downloads open."],
        ["Before each data update", "Back up Excel, CSV, .pt, and relevant cache files."],
        ["After each code update", "Run manual tests by tab and test_two_part_parsing.py if question logic changed."],
        ["Every 6-12 months", "Review OpenAI models, cost, limits, dependencies, and Streamlit compatibility."],
    ], widths=[2.2, 4.6])

    h(doc, "4. Future update recommendations", 1)
    add_bullets(doc, [
        "Gradually split the monolith into data loading, document extraction, AI clients, evaluation logic, and UI modules.",
        "Remove obsolete commented blocks after confirming they are not needed.",
        "Consolidate the duplicate load_extended_data and load_embeddings startup calls.",
        "Decide whether load_embeddings should still run at startup while visible recommendation search remains commented.",
        "Fix Tab 4 if section/page citations are required by using the enhanced hierarchical extractor.",
        "Wire the deep-analysis model selector into run_row_analysis or remove the selector.",
        "Add tests for parse_two_part_question, _apply_critic_gate_and_render, detect_transversal_matter, _apply_transversal_gate_and_render, classification, and XLSX exports.",
        "Maintain a small non-sensitive regression document set for reproducible validation.",
        "Record schema changes to each Excel file with date and ILO owner.",
        "If the split variants remain in use, propagate relevant changes from oli_v6_deploy.py and test each variant separately.",
    ])

    h(doc, "5. AI and cost recommendations", 1)
    add_bullets(doc, [
        "Keep gpt-5-mini as the main evaluation model while human validation remains consistent.",
        "Use LLM Verification in classification only when added precision justifies cost and latency.",
        "Filter and limit rows before deep recommendation analysis.",
        "Keep embeddings_cache.pkl and analysis_cache.pkl for recurring batches; delete them only if corrupt or after substantive data changes.",
        "Review the startup FAISS recommendation index because classification tabs do not use it, although it still loads.",
        "Track internal cost estimates with generate_cost_report.py if frequent institutional use continues.",
    ])

    h(doc, "6. Change acceptance criteria", 1)
    add_bullets(doc, [
        "The app starts with streamlit run oli_v6_deploy.py.",
        "Tab 1 processes a test DOCX, produces results, recognizes at least one transversal theme, and shows the A/B audit line.",
        "Tab 2 evaluates at least one selected criterion and downloads a ZIP.",
        "Tab 3 generates results and charts for a sample.",
        "Tab 4 answers on both a small TXT and a large DOCX, and the same-filename replacement procedure is understood.",
        "Tabs 5 and 6 classify a limited sample and show treemaps.",
        "Documentation contains no fill-in fields, real credentials, or unnecessary personal paths.",
    ])

    h(doc, "7. Minimum handoff package", 1)
    add_bullets(doc, [
        "Main code: oli_v6_deploy.py.",
        "Dependencies: requirements.txt.",
        "Updated Spanish documentation: six DOCX files in documentation.",
        "Updated English documentation: six DOCX files in documentation/en.",
        "Active data files and rubrics listed in the data model documentation.",
        "ILO internal deployment instructions for the chosen hosting environment.",
        "ILO register of access, data, and deployment owners.",
    ])

    h(doc, "8. GPT assistant handoff and maintenance", 1)
    p(doc, "The three GPTs published in ChatGPT and their FastAPI backend form a deliverable separate from the Streamlit app, with their own maintenance cycle, credentials, and failure modes.")

    h(doc, "8.1 Assets to transfer", 2)
    add_table(doc, ["Asset", "Where it lives", "Who should control it"], [
        ["The three published GPTs", "The ChatGPT account where they were created", "Should migrate to an institutional ILO account; while they remain on a personal account, continuity depends on that account."],
        ["Backend service", "Render.com, service ilo-prodoc-appraisal-v3", "ILO technical administrator."],
        ["OPENAI_API_KEY", "Environment variable in Render", "ILO technical administrator; determines which account is billed for usage."],
        ["ILO_GPT_ACTION_API_KEY", "Environment variable in Render and in all three Actions", "ILO technical administrator."],
        ["Private GPT links", "Internal ILO distribution", "Owning team. All three are set to \"Anyone with a link\": the link is a bearer credential and the distribution list is the only access control."],
        ["  · Quality Appraisal (Tab 1 v3)", GPT_V3, "Private link; requires a ChatGPT account."],
        ["  · Specific Attributes (Tab 2)", GPT_TAB2, "Private link; requires a ChatGPT account."],
        ["  · Sustainability (Tab 3)", GPT_SUS, "Private link; requires a ChatGPT account."],
        ["  · Shared backend", GPT_BACKEND, "FastAPI on Render; public probe at /health."],
        ["Code and rubrics", "Project repository", "ILO technical administrator."],
    ], widths=[1.8, 2.2, 2.8])
    p(doc, "Migrating the GPTs to an institutional account is the single most important handoff action and cannot be deferred indefinitely: a GPT published from a personal account stops being available if that account is closed or changes plan.")

    h(doc, "8.2 GPT layer maintenance routine", 2)
    add_table(doc, ["Frequency", "Activity"], [
        ["Weekly during active operation", "Confirm GET /health responds and run a short evaluation on one of the three assistants."],
        ["Monthly", "Test all three assistants with a sample document and confirm the XLSX downloads and opens."],
        ["Monthly during the pilot phase", "Review consumption in the OpenAI dashboard and project the budget exhaustion date."],
        ["When updating a rubric", "Replace the XLSX, test locally, push to main, and verify the redeploy with a scoped evaluation."],
        ["When editing instructions or starter buttons", "Update docs/gpt_onboarding_es.md too, so documentation does not drift from the assistants."],
        ["Every 6-12 months", "Review available models, costs, API limits, and the Render plan."],
    ], widths=[2.2, 4.6])

    h(doc, "8.3 Change acceptance criteria for the GPT layer", 2)
    add_bullets(doc, [
        "GET /health responds after deployment.",
        "GET /privacy responds: ChatGPT requires it to keep a GPT with an Action published.",
        "Each of the three assistants completes a scoped test evaluation and delivers its XLSX.",
        "The specific-attributes and sustainability XLSX outputs include the stability column.",
        "All three Actions point at the correct server and authenticate with the current key.",
        "Greeting an assistant with no attachment makes it introduce itself rather than wait in silence.",
    ])

    h(doc, "8.4 Continuity and cost recommendations", 2)
    add_bullets(doc, [
        "Migrate the three GPTs to an institutional ILO account before widening usage.",
        "Replace in-memory job storage with Redis or an equivalent if usage moves beyond a pilot: today a redeploy loses in-flight evaluations.",
        "Consider a paid Render plan if cold-start latency degrades the user experience.",
        "Watch the cost of the stability repeats: 10 runs per criterion in quality appraisal and 5 in the other two assistants. A full appraisal issues on the order of 760 model calls and is by far the most expensive workflow. Lowering STABILITY_REPEATS is the most direct lever if budget tightens, at the cost of reliability; lowering it in tab1_v3_core saves the most.",
        "Use separate API keys per assistant if consumption must be attributed to each: the OpenAI dashboard breaks down by key, not by GPT.",
        "Prefer section- or dimension-scoped evaluations over the full rubric when verifying a specific point: it cuts cost and time without losing value.",
        "Record the observed unit cost during the pilot so a larger deployment can be estimated.",
    ])


DOCS = [
    ("1_Documentacion_Tecnica.docx", "DOCUMENTACIÓN TÉCNICA DEL SISTEMA", "Arquitectura, componentes y despliegue", doc1),
    ("2_Documentacion_Codigo_Fuente.docx", "DOCUMENTACIÓN DEL CÓDIGO FUENTE", "Guía de mantenimiento para la app principal", doc2),
    ("3_Documentacion_Base_Datos.docx", "DOCUMENTACIÓN DE BASE DE DATOS", "Inventario file-based y relaciones lógicas", doc3),
    ("4_Documentacion_Configuracion_Seguridad.docx", "DOCUMENTACIÓN DE CONFIGURACIÓN Y SEGURIDAD", "Variables, secretos, acceso y protección de datos", doc4),
    ("5_Documentacion_Funcional_Operativa.docx", "DOCUMENTACIÓN FUNCIONAL Y OPERATIVA", "Manual operativo para usuarios y administradores", doc5),
    ("6_Otros_Contacto_Recomendaciones.docx", "INFORMACIÓN ADICIONAL DE CIERRE", "Contacto, mantenimiento y recomendaciones futuras", doc6),
]

DOCS_EN = [
    ("1_Technical_Documentation.docx", "SYSTEM TECHNICAL DOCUMENTATION", "Architecture, components, and deployment", doc_en1),
    ("2_Source_Code_Documentation.docx", "SOURCE CODE DOCUMENTATION", "Maintenance guide for the main app", doc_en2),
    ("3_Data_Model_Documentation.docx", "DATA MODEL DOCUMENTATION", "File-based inventory and logical relationships", doc_en3),
    ("4_Configuration_Security_Documentation.docx", "CONFIGURATION AND SECURITY DOCUMENTATION", "Variables, secrets, access, and data protection", doc_en4),
    ("5_Functional_Operational_Documentation.docx", "FUNCTIONAL AND OPERATIONAL DOCUMENTATION", "User and administrator operating manual", doc_en5),
    ("6_Additional_Contact_Recommendations.docx", "ADDITIONAL CLOSEOUT INFORMATION", "Contact, maintenance, and future recommendations", doc_en6),
]


def main():
    EN_DIR.mkdir(exist_ok=True)
    for filename, title, subtitle, builder in DOCS:
        save_doc(filename, title, subtitle, builder)
        print(f"updated {filename}")
    for filename, title, subtitle, builder in DOCS_EN:
        save_doc(filename, title, subtitle, builder, output_dir=EN_DIR, language="en")
        print(f"updated en/{filename}")


if __name__ == "__main__":
    main()
