"""Generate the Spanish DOCX user manual for the three ILO GPT assistants."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x00, 0x3E, 0x7E)
CYAN = RGBColor(0x00, 0x72, 0xBC)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
RED = RGBColor(0xD6, 0x00, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT = "Manual_Usuario_GPTs_OIT.docx"

GPT_V3 = "https://chatgpt.com/g/g-6a2643b11e708191adc1c03e64260a25-ilo-prodoc-quality-appraisal"
GPT_TAB2 = "https://chatgpt.com/g/g-6a43cc82d4d08191bf6e60357453e336-oit-diagnostico-de-atributos-especificos"
GPT_SUS = "https://chatgpt.com/g/g-6a43d24307b88191bb362632f133c0f3-oit-diagnostico-de-sostenibilidad-del-proyecto"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, size=10.5, after=6, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    return p


def bullet(doc, text, prefix=None, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    if prefix:
        r = p.add_run(prefix)
        r.font.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_hyperlink(paragraph, url, text, size=9.5):
    """Insert a clickable external link (python-docx has no native support)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0072BC")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(color)
    rPr.append(u)
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


def table(doc, headers, rows, widths=None, font=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, "003E7E")
        for para in c.paragraphs:
            for r in para.runs:
                r.font.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(font)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cells[i].text = val
            set_cell_bg(cells[i], bg)
            for para in cells[i].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(font)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def link_table(doc, rows):
    """Access table whose middle column holds a clickable link."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Asistente", "Enlace de acceso", "Para qué documento"]):
        c = t.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, "003E7E")
        for para in c.paragraphs:
            for r in para.runs:
                r.font.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(9.5)
    for ri, (name, url, label, use) in enumerate(rows):
        cells = t.add_row().cells
        bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        cells[0].text = name
        cells[2].text = use
        for idx in (0, 2):
            for para in cells[idx].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9.5)
        add_hyperlink(cells[1].paragraphs[0], url, label)
        for c in cells:
            set_cell_bg(c, bg)
    for row in t.rows:
        for i, w in enumerate([4.2, 6.3, 5.5]):
            row.cells[i].width = Cm(w)
    return t


def starters(doc, items):
    body(doc, "Botones de inicio disponibles en el chat:", bold=True, after=4)
    for s in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run("«" + s + "»")
        r.font.size = Pt(10.5)
        r.font.italic = True
        r.font.color.rgb = CYAN
        p.paragraph_format.space_after = Pt(2)


doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.4)
    s.bottom_margin = Cm(2.4)
    s.left_margin = Cm(2.6)
    s.right_margin = Cm(2.6)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ── Portada ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("MANUAL DE USUARIO")
r.font.bold = True
r.font.size = Pt(20)
r.font.color.rgb = BLUE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)

p2 = doc.add_paragraph()
r2 = p2.add_run("Asistentes GPT de valoración de documentos de proyecto")
r2.font.size = Pt(13)
r2.font.color.rgb = CYAN
r2.font.bold = True
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(10)

meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate([
    ("Organización", "Organización Internacional del Trabajo (OIT) — Oficina Regional para América Latina y el Caribe"),
    ("Dirigido a", "Personal de la OIT que diseña, revisa o evalúa documentos de proyecto"),
    ("Alcance", "Tres asistentes GPT publicados en ChatGPT, con backend compartido"),
    ("Versión / Fecha", "1.0 — Julio de 2026"),
]):
    row = meta.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], "003E7E")
    set_cell_bg(row.cells[1], "F0F4F8")
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(9.5)
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(12.5)

doc.add_paragraph()

# ── 1. Qué son ───────────────────────────────────────────────────────────
heading(doc, "1. Qué son estos asistentes")
body(doc,
     "La OIT cuenta con tres asistentes GPT que evalúan documentos de proyecto contra rúbricas "
     "institucionales. Cada asistente lee el documento que usted sube, aplica la rúbrica criterio por "
     "criterio, cita la evidencia encontrada en el texto y devuelve un archivo Excel descargable con la "
     "valoración completa.")
body(doc,
     "Las rúbricas ya están cargadas en el servidor: usted nunca necesita subirlas. Sólo sube el documento "
     "a evaluar e indica el alcance.", bold=True)

body(doc, "Qué hacen:", bold=True, after=4)
for txt in [
    "Evalúan un documento .docx contra una rúbrica institucional, criterio por criterio.",
    "Citan la evidencia textual que sustenta cada valoración.",
    "Señalan los criterios débiles y los que requieren revisión humana.",
    "Entregan un Excel estructurado, listo para trabajar sobre él.",
]:
    bullet(doc, txt)

body(doc, "Qué NO hacen:", bold=True, after=4)
for txt in [
    "No sustituyen el juicio experto: los resultados son asistidos por IA y requieren validación.",
    "No constituyen una determinación oficial de la OIT.",
    "No corrigen ni reescriben el documento evaluado.",
    "No procesan varios documentos a la vez: un archivo por evaluación.",
]:
    bullet(doc, txt)

# ── 2. Acceso ────────────────────────────────────────────────────────────
heading(doc, "2. Acceso a los asistentes")
body(doc,
     "Los asistentes no aparecen en el directorio público de ChatGPT: se accede únicamente por enlace. "
     "Cualquier persona que reciba el enlace y tenga una cuenta de ChatGPT puede abrirlos y ejecutar "
     "evaluaciones. Haga clic en el enlace correspondiente o cópielo en el navegador.",
     after=4)
body(doc,
     "Por esa razón, el enlace debe tratarse como una credencial: al reenviarlo se transfiere el acceso, "
     "y el consumo se factura contra un presupuesto común. Compártalo sólo con las personas que deban usar "
     "los asistentes.", bold=True, after=6)

link_table(doc, [
    ("Valoración Preliminar de Calidad\n(Tab 1 · rúbrica v3)", GPT_V3,
     "ILO PRODOC Quality Appraisal",
     "PRODOCs y documentos de diseño"),
    ("Diagnóstico de Atributos Específicos\n(Tab 2)", GPT_TAB2,
     "OIT - Diagnóstico de Atributos Específicos",
     "Cualquier documento de proyecto"),
    ("Diagnóstico de Sostenibilidad\n(Tab 3)", GPT_SUS,
     "OIT - Diagnóstico de Sostenibilidad del Proyecto",
     "Diseño, avance o cierre, según la etapa"),
])

doc.add_paragraph()
body(doc,
     "Recomendación: guarde los tres enlaces en favoritos. Si no está seguro de cuál usar, la sección 8 "
     "incluye una guía de selección.", italic=True, size=9.5, color=GRAY)

# ── 3. Cómo usarlos ──────────────────────────────────────────────────────
heading(doc, "3. Cómo usarlos — flujo general")
body(doc,
     "Los tres asistentes siguen el mismo flujo. La evaluación se ejecuta en un servidor externo, por lo "
     "que puede tardar varios minutos; el asistente le informará del avance.", after=6)

for i, (t, d) in enumerate([
    ("Abrir el asistente",
     "Haga clic en el enlace de la sección 2. Si es la primera vez, escriba «hola» o pulse un botón de "
     "inicio: el asistente se presentará y le explicará qué puede evaluar."),
    ("Subir el documento",
     "Adjunte un único archivo .docx en el chat. Si sube varios, el asistente le pedirá elegir uno."),
    ("Indicar el alcance",
     "Escriba en el mismo mensaje qué quiere evaluar (una sección concreta, una rúbrica temática o una "
     "dimensión del ciclo). Si no lo indica, el asistente se lo preguntará antes de empezar."),
    ("Esperar la evaluación",
     "El asistente lanza el trabajo y consulta su estado hasta terminar. No cierre la conversación."),
    ("Revisar y descargar",
     "Recibirá un resumen en pantalla y un archivo Excel descargable con la valoración completa y la "
     "evidencia citada."),
]):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run("Paso %d — %s. " % (i + 1, t))
    r.font.bold = True
    r.font.size = Pt(10.5)
    r2 = p.add_run(d)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)

# ── 4. Asistente 1 ───────────────────────────────────────────────────────
heading(doc, "4. Asistente 1 — Valoración Preliminar de Calidad (Tab 1 · v3)")
body(doc,
     "Evalúa un PRODOC contra la rúbrica v3 de valoración preliminar de calidad: 76 criterios "
     "organizados en 5 secciones. Cada criterio recibe uno de cinco veredictos —Yes, Partial, No, "
     "Not Found o N/A— con la evidencia citada del documento.")
body(doc,
     "«Not Found» significa que el documento no contiene información suficiente para evaluar el criterio. "
     "No es lo mismo que «No»: este último afirma que el criterio no se cumple, mientras que «Not Found» "
     "afirma que no se pudo determinar. Distinguirlos importa al decidir si corregir el documento o "
     "completarlo.")
body(doc,
     "Cada criterio se evalúa diez veces y el resultado se consolida al valor más frecuente. Es el "
     "asistente más lento y más costoso de los tres, y también el más robusto frente a la variabilidad "
     "del modelo.")

body(doc, "Estructura de la rúbrica:", bold=True, after=4)
table(doc,
      ["Sección", "Contenido", "Criterios", "Subsecciones"],
      [["1", "Pertinencia", "20", "1.1 – 1.5"],
       ["2", "Validez del diseño", "13", "2.1 – 2.4"],
       ["3", "Marco de resultados / R&M", "27", "3.1 – 3.7"],
       ["4", "Implementación", "14", "4.1 – 4.4"],
       ["5", "Presentación", "2", "5.1 – 5.2"]],
      widths=[2.0, 7.5, 2.2, 4.3])
doc.add_paragraph()

body(doc, "Filtros disponibles:", bold=True, after=4)
for txt, pfx in [
    ("evalúa los 76 criterios. Es la opción más lenta pero la más completa.", "Evaluación completa: "),
    ("por ejemplo «solo la sección 3». Útil para revisar el marco de resultados de forma focalizada.", "Por sección: "),
    ("por ejemplo «solo 1.1 y 2.3». La opción más rápida para verificar puntos concretos.", "Por subsección: "),
]:
    bullet(doc, txt, prefix=pfx)
body(doc,
     "El filtro debe indicarse en el mensaje al subir el documento.", italic=True, size=9.5, color=GRAY)

starters(doc, [
    "¿Qué puedes hacer y cómo empiezo?",
    "Evalúa este PRODOC con la rúbrica completa",
    "Evalúa solo la sección 3 (Marco de resultados)",
    "¿Qué secciones y subsecciones puedo filtrar?",
])

body(doc, "Qué recibe:", bold=True, after=4)
for txt in [
    "Total de criterios evaluados y conteo por veredicto (Yes / Partial / No / Not Found / N/A).",
    "Lista de criterios de alta subjetividad que requieren revisión humana prioritaria.",
    "El porcentaje de estabilidad por criterio: por debajo de 80% de coincidencia entre corridas, el "
    "criterio se marca como inestable y conviene revisarlo con atención.",
    "Archivo Excel con la valoración y la evidencia citada por criterio.",
]:
    bullet(doc, txt)

# ── 5. Asistente 2 ───────────────────────────────────────────────────────
heading(doc, "5. Asistente 2 — Diagnóstico de Atributos Específicos (Tab 2)")
body(doc,
     "Evalúa un documento de proyecto contra una o varias rúbricas temáticas, en escala de 1 a 5 por "
     "criterio. Usted elige qué rúbrica aplicar; si no la indica, el asistente se lo preguntará antes de "
     "empezar.")

body(doc, "Rúbricas disponibles:", bold=True, after=4)
table(doc,
      ["Rúbrica", "Qué mide"],
      [["Metodologías participativas", "Enfoque participativo en el diseño y la implementación del proyecto"],
       ["Género", "Integración del enfoque de género a lo largo del documento"],
       ["Transición Justa", "Alineación con el enfoque moderno de Transición Justa"]],
      widths=[5.5, 10.5])
doc.add_paragraph()

body(doc,
     "Cada criterio se evalúa cinco veces y las puntuaciones se consolidan (esquema de estabilidad). Esto "
     "reduce la variabilidad del modelo y hace la evaluación más fiable, pero también más lenta que una "
     "simple lectura. El Excel incluye un porcentaje de estabilidad que indica cuántas de las cinco "
     "corridas coincidieron: un porcentaje bajo señala un criterio donde conviene revisar con más "
     "atención. Puede aplicar varias rúbricas en una misma consulta y comparar los resultados.")

starters(doc, [
    "¿Qué rúbricas puedes aplicar?",
    "Evalúa este documento con la rúbrica de género",
    "Aplica la rúbrica de Transición Justa",
    "Evalúa participación y género y compara resultados",
])

body(doc, "Qué recibe:", bold=True, after=4)
for txt in [
    "Puntuación de 1 a 5 por criterio, con la evidencia que la sustenta.",
    "Los criterios con puntuación baja (1–2) señalados como oportunidades de mejora.",
    "Archivo Excel con el detalle completo.",
]:
    bullet(doc, txt)

# ── 6. Asistente 3 ───────────────────────────────────────────────────────
heading(doc, "6. Asistente 3 — Diagnóstico de Sostenibilidad (Tab 3)")
body(doc,
     "Evalúa la sostenibilidad del proyecto contra la rúbrica institucional: 28 criterios organizados en "
     "tres dimensiones del ciclo del proyecto. La dimensión que corresponde depende del tipo de documento "
     "que esté evaluando.")

body(doc, "Dimensiones y correspondencia:", bold=True, after=4)
table(doc,
      ["Dimensión", "Criterios", "Aplíquela a"],
      [["Diseño", "6", "PRODOCs y documentos de diseño"],
       ["Implementación", "10", "Informes de avance o de medio término"],
       ["Pre-Cierre", "12", "Documentos de cierre o de evaluación final"]],
      widths=[4.0, 2.5, 9.5])
doc.add_paragraph()

body(doc, "Escala de valoración (0 a 3):", bold=True, after=4)
table(doc,
      ["Puntuación", "Significado"],
      [["0", "Ausente — el documento no aborda el indicador"],
       ["1", "Incipiente — mención superficial, sin desarrollo"],
       ["2", "Parcial — abordado, pero con vacíos relevantes"],
       ["3", "Sólido — abordado de forma completa y verificable"]],
      widths=[3.0, 13.0])
doc.add_paragraph()

body(doc,
     "Atención: esta escala es 0–3, distinta de la escala 1–5 del asistente de Atributos Específicos.",
     bold=True, color=RED)

starters(doc, [
    "¿Qué dimensiones evalúas y cuál me corresponde?",
    "Evalúa este PRODOC con la dimensión de Diseño",
    "Es un informe de avance: aplica Implementación",
    "Aplica la rúbrica completa de sostenibilidad",
])

body(doc, "Qué recibe:", bold=True, after=4)
for txt in [
    "Puntuación de 0 a 3 por indicador; los indicadores en 0–1 se destacan como alertas.",
    "Un Excel con dos hojas: «Lectura amigable» para gestores y «Auditoría técnica» para revisores.",
    "El porcentaje de estabilidad de cada indicador: al igual que en Atributos Específicos, cada "
    "indicador se evalúa cinco veces y se reporta cuántas corridas coincidieron.",
]:
    bullet(doc, txt)

# ── 7. Interpretar resultados ────────────────────────────────────────────
heading(doc, "7. Cómo leer los resultados")
body(doc,
     "Cada asistente usa una escala distinta. Antes de comparar o consolidar resultados, verifique cuál "
     "está leyendo:", after=6)

table(doc,
      ["Asistente", "Escala", "Cómo se lee"],
      [["Valoración de Calidad (Tab 1)", "Yes / Partial / No / Not Found / N/A",
        "«Partial» indica que el criterio está presente pero incompleto; «Not Found» indica que el documento no permite evaluarlo"],
       ["Atributos Específicos (Tab 2)", "1 a 5",
        "5 es el máximo; 1–2 señalan oportunidades de mejora prioritarias"],
       ["Sostenibilidad (Tab 3)", "0 a 3",
        "3 es el máximo; 0–1 son alertas que requieren atención"]],
      widths=[5.0, 3.2, 7.8])
doc.add_paragraph()

body(doc, "Recomendaciones de lectura:", bold=True, after=4)
for txt, pfx in [
    ("no acepte una valoración sin contrastarla con el texto citado. La evidencia es lo que hace auditable el resultado.",
     "Verifique siempre la evidencia: "),
    ("son aquellos donde el juicio depende del contexto institucional; el asistente los señala explícitamente.",
     "Priorice los criterios de alta subjetividad: "),
    ("un criterio valorado «No» puede reflejar una omisión real del documento o una redacción que el modelo no reconoció. Ambas cosas son información útil.",
     "Distinga ausencia de redacción deficiente: "),
    ("los resultados son un insumo de trabajo, no un dictamen. La decisión final es siempre humana.",
     "Use el resultado como punto de partida: "),
]:
    bullet(doc, txt, prefix=pfx)

# ── 8. Qué asistente usar ────────────────────────────────────────────────
heading(doc, "8. Qué asistente usar en cada caso")
table(doc,
      ["Si necesita…", "Use", "Alcance sugerido"],
      [["Revisar la calidad general de un PRODOC antes de aprobarlo",
        "Valoración de Calidad (Tab 1)", "Rúbrica completa"],
       ["Verificar un aspecto concreto del diseño (p. ej. marco de resultados)",
        "Valoración de Calidad (Tab 1)", "Sólo la sección correspondiente"],
       ["Saber si el proyecto integra el enfoque de género",
        "Atributos Específicos (Tab 2)", "Rúbrica de género"],
       ["Evaluar participación o Transición Justa",
        "Atributos Específicos (Tab 2)", "Rúbrica correspondiente"],
       ["Valorar la sostenibilidad de un proyecto en diseño",
        "Sostenibilidad (Tab 3)", "Dimensión Diseño"],
       ["Valorar sostenibilidad a medio término o al cierre",
        "Sostenibilidad (Tab 3)", "Implementación o Pre-Cierre"]],
      widths=[7.0, 4.5, 4.5])
doc.add_paragraph()

# ── 9. Buenas prácticas ──────────────────────────────────────────────────
heading(doc, "9. Buenas prácticas y limitaciones")
body(doc, "Para obtener mejores resultados:", bold=True, after=4)
for txt in [
    "Suba el documento en formato .docx. Otros formatos pueden no procesarse correctamente.",
    "Un solo documento por evaluación: los asistentes no comparan varios archivos entre sí.",
    "Indique el alcance en el mismo mensaje en que sube el archivo: ahorra una ronda de preguntas.",
    "Si el documento es extenso, evalúe por secciones: el resultado llega antes y es más fácil de revisar.",
    "Conserve el Excel: es el registro auditable de la evaluación, con la evidencia citada.",
]:
    bullet(doc, txt)

body(doc, "Limitaciones que conviene tener presentes:", bold=True, after=4)
for txt in [
    "Los resultados son asistidos por IA y pueden contener errores de interpretación.",
    "Dos evaluaciones del mismo documento pueden diferir ligeramente. Los tres asistentes mitigan esto "
    "repitiendo cada criterio y consolidando el resultado: diez veces en Valoración de Calidad, cinco en "
    "Atributos Específicos y cinco en Sostenibilidad. El Excel reporta el porcentaje de estabilidad, es "
    "decir cuántas corridas coincidieron.",
    "El asistente evalúa lo que está escrito en el documento, no el proyecto en la realidad.",
    "Ninguna salida constituye una determinación oficial de la OIT.",
]:
    bullet(doc, txt)

# ── 10. Problemas frecuentes ─────────────────────────────────────────────
heading(doc, "10. Problemas frecuentes")
table(doc,
      ["Situación", "Qué hacer"],
      [["El asistente no responde al abrir el chat",
        "Escriba «hola» o pulse uno de los botones de inicio; se presentará y explicará qué puede evaluar"],
       ["La evaluación tarda mucho",
        "Es normal en documentos extensos o rúbricas completas. Si supera lo razonable, reintente "
        "acotando el alcance a una sola sección o dimensión"],
       ["La evaluación falla",
        "El asistente mostrará el mensaje de error. Reintente con un alcance más acotado; si persiste, "
        "reporte el mensaje al soporte técnico"],
       ["Subí varios archivos y no procesa",
        "Los asistentes admiten un solo .docx por evaluación. Indique cuál usar"],
       ["Me pide subir la rúbrica",
        "No es necesario: las rúbricas están en el servidor. Aclare que use la rúbrica cargada"],
       ["No sé qué dimensión o rúbrica aplicar",
        "Pregúnteselo al asistente describiendo el tipo de documento; le sugerirá la opción adecuada"]],
      widths=[5.5, 10.5])
doc.add_paragraph()

# ── 11. Soporte ──────────────────────────────────────────────────────────
heading(doc, "11. Soporte y continuidad")
body(doc,
     "Durante la fase piloto, el consultor responsable del desarrollo provee soporte técnico básico: "
     "atención a incidencias de uso, resolución de fallos, actualización de rúbricas a solicitud del "
     "equipo técnico de la OIT y seguimiento de la disponibilidad del servicio.")
body(doc,
     "Al reportar una incidencia, incluya: el asistente utilizado, el alcance solicitado, el mensaje de "
     "error tal como apareció y, si es posible, el nombre del documento evaluado. Eso permite reproducir "
     "el problema sin necesidad de compartir el archivo.")

doc.add_paragraph()
body(doc,
     "Los tres asistentes comparten un mismo backend. Si ninguno responde, es probable que se trate de "
     "una interrupción del servicio y no de un problema del documento.",
     italic=True, size=9.5, color=GRAY)

doc.save(OUTPUT)
print("Manual saved: " + OUTPUT)
